import sys
import os
import csv
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QTabWidget, QWidget, QVBoxLayout,
    QHBoxLayout, QLabel, QLineEdit, QComboBox, QPushButton, QFileDialog
)
from docx import Document
from docx.shared import Pt  # 用於設定字型大小
from docx.oxml.ns import qn  # 用於設定字型類型

class CreditorEditor(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("消債蘿蔔特-債權人清冊")
        self.setGeometry(100, 100, 800, 600)
        self.tabs = QTabWidget()
        self.setCentralWidget(self.tabs)
        self.row_inputs = []
        self.total_amount_label = QLabel("0")  # 初始化債權總金額標籤
        self.bank_data, self.company_data = self.read_csv_files()
        self.init_tab()

    def read_csv_files(self):
        bank_data = []
        company_data = []
        
        # 取得目前指令稿所在的目錄
        current_dir = os.path.dirname(os.path.abspath(__file__))
        document_dir = os.path.join(current_dir, "document")
        bank_file = os.path.join(document_dir, "all adress - bank.csv")
        company_file = os.path.join(document_dir, "all adress - company.csv")

        # 讀取銀行CSV檔案
        with open(bank_file, newline='', encoding='utf-8') as csvfile:
            reader = csv.reader(csvfile)
            bank_data = [(row[0], row[1]) for row in reader]

        # 讀取公司CSV檔案
        with open(company_file, newline='', encoding='utf-8') as csvfile:
            reader = csv.reader(csvfile)
            company_data = [(row[0], row[1]) for row in reader]

        return bank_data, company_data

    def init_tab(self):
        tab = QWidget()
        layout = QVBoxLayout()

        button_layout = QHBoxLayout()
        add_button = QPushButton("＋")
        save_button = QPushButton("另存新檔")
        button_layout.addWidget(add_button)
        button_layout.addWidget(save_button)
        layout.addLayout(button_layout)

        # 顯示債權總金額
        total_layout = QHBoxLayout()
        total_label = QLabel("債權總金額：")
        total_layout.addWidget(total_label)
        total_layout.addWidget(self.total_amount_label)
        layout.addLayout(total_layout)

        layout.addStretch()
        self.row_inputs = []
        add_button.clicked.connect(lambda: self.row_inputs.append(self.add_creditor_row(layout)))
        save_button.clicked.connect(self.save_document)

        tab.setLayout(layout)
        self.tabs.addTab(tab, "債權人清冊")

    def add_creditor_row(self, layout):
        row_layout = QHBoxLayout()
        creditor_name = QComboBox()
        creditor_name.addItems([name for name, _ in self.bank_data + self.company_data])
        creditor_name.setEditable(True)
        address_label = QLabel("")
        address_input = QLineEdit()  # 用於手動輸入地址的輸入框
        address_input.setPlaceholderText("輸入地址")
        address_input.setVisible(False)  # 預設隱藏

        def update_address():
            name = creditor_name.currentText()
            address = next((addr for n, addr in self.bank_data + self.company_data if n == name), "")
            if address:
                address_label.setText(address)
                address_input.setText(address)  # 確保輸入框也有地址
                address_input.setVisible(False)
            else:
                address_label.setText("")
                address_input.setVisible(True)

        creditor_name.currentIndexChanged.connect(update_address)
        creditor_name.editTextChanged.connect(update_address)

        amount_edit = QLineEdit()
        amount_edit.setPlaceholderText("債權額")
        amount_edit.textChanged.connect(self.update_total_sum)

        debt_type = QComboBox()
        debt_type.addItems(["信用卡", "信用貸款", "信用卡及信用貸款"])

        row_layout.addWidget(creditor_name)
        row_layout.addWidget(address_label)
        row_layout.addWidget(address_input)  # 加入用於手動輸入地址的輸入框
        row_layout.addWidget(amount_edit)
        row_layout.addWidget(debt_type)

        layout.insertLayout(layout.count() - 2, row_layout)
        return [creditor_name, address_label, address_input, amount_edit, debt_type]

    def update_total_sum(self):
        total_sum = 0
        for row in self.row_inputs:
            amount = row[3].text()  # 更新 index 以取得正確的輸入框
            if amount.strip():
                try:
                    total_sum += int(amount)
                except ValueError:
                    pass
        self.total_amount_label.setText(str(total_sum))
        print(f"總金額更新為：{total_sum}")

    def save_document(self):
        current_dir = os.path.dirname(os.path.abspath(__file__))
        document_dir = os.path.join(current_dir, "document")
        file_path = os.path.join(document_dir, "C.docx")

        doc = Document(file_path)

        # 更新总计金额到第一个表格的A0单元格
        first_table = doc.tables[0]
        cell = first_table.cell(0, 1)
        cell.text = self.total_amount_label.text()  # 填写总计金额

        # 设置总计金额字体为标楷体，字号为16
        run = cell.paragraphs[0].runs[0]
        run.font.name = 'DFKai-SB'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')  # 设置中文字体
        run.font.size = Pt(16)

        # 取得第二個表格
        second_table = doc.tables[1]

        # 删除未使用的代号行
        markers = ['B0', 'B1', 'B2', 'B3', 'C0', 'C1', 'C2', 'C3', 
                   'D0', 'D1', 'D2', 'D3', 'E0', 'E1', 'E2', 'E3', 
                   'F0', 'F1', 'F2', 'F3', 'G0', 'G1', 'G2', 'G3', 
                   'H0', 'H1', 'H2', 'H3']
        rows_to_remove = []
        for row in second_table.rows:
            if any(cell.text.strip() in markers for cell in row.cells):
                rows_to_remove.append(row)

        for row in rows_to_remove:
            second_table._element.remove(row._element)

        # 填入有效數據到第二个表格
        for row_inputs in self.row_inputs:
            name, address_label, address_input, amount_edit, debt_type = row_inputs
            new_row = second_table.add_row().cells

            new_row[0].text = name.currentText()
            new_row[1].text = address_input.text()  # 使用輸入框中的地址
            new_row[2].text = amount_edit.text()
            new_row[3].text = debt_type.currentText()

            # 更新CSV檔案
            self.update_csv(name.currentText(), address_input.text())

        default_filename = "03_債權人清冊.docx"
        save_path, _ = QFileDialog.getSaveFileName(self, "存檔", default_filename, "Word Files (*.docx);;All Files (*)")
        
        if save_path:
            if os.path.exists(save_path):
                os.remove(save_path)  # 確保目標檔案未被佔用
            try:
                doc.save(save_path)
                print(f"檔案已存檔到: {save_path}")
            except PermissionError:
                print(f"無法存檔，權限不足: {save_path}")

    def update_csv(self, name, address):
        current_dir = os.path.dirname(os.path.abspath(__file__))
        document_dir = os.path.join(current_dir, "document")
        bank_file = os.path.join(document_dir, "all adress - bank.csv")
        company_file = os.path.join(document_dir, "all adress - company.csv")
        
        # 確定應該更新哪個CSV檔案
        if "銀行" in name:
            file_to_update = bank_file
            existing_data = self.bank_data
        else:
            file_to_update = company_file
            existing_data = self.company_data

        # 检查是否已经存在
        if any(existing_name == name for existing_name, _ in existing_data):
            print(f"{name} 已經存在於 {file_to_update} 中，未新增。")
            return

        # 附加新的內容
        with open(file_to_update, 'a', newline='', encoding='utf-8') as csvfile:
            writer = csv.writer(csvfile)
            writer.writerow([name, address])
            print(f"已更新CSV檔案: {file_to_update} with {name}, {address}")

        # 更新内存中的数据以避免重复
        if "銀行" in name:
            self.bank_data.append((name, address))
        else:
            self.company_data.append((name, address))

if __name__ == '__main__':
    app = QApplication(sys.argv)
    window = CreditorEditor()
    window.show()
    sys.exit(app.exec_())
