import sys
import os
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QTabWidget, QWidget, QVBoxLayout,
    QHBoxLayout, QLabel, QLineEdit, QComboBox, QPushButton, QFileDialog, QGridLayout
)
from docx import Document

class TableEditor(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("消債蘿蔔特-財產及收入狀況說明書")
        self.setGeometry(100, 100, 800, 600)
        self.tabs = QTabWidget()
        self.setCentralWidget(self.tabs)
        self.row_inputs = []
        self.expense_total_label = QLabel("0")  # 初始化支出總計標籤
        self.income_total_label = QLabel("0")  # 初始化收入總計標籤
        self.init_tabs()

    def init_tabs(self):
        self.add_tab("保險", ['保險公司名稱', '險種', '保單號碼', '金額'], ['一般壽險', '實支實付險'])
        self.add_tab("土地及建物", ['地號或建號', '公告現值總價額', '持有比例', '備註'], ['公同共有難以變賣', '持份過低難以變賣'], is_select_for_notes=True)
        self.add_tab("車輛", ['車輛種類', '車牌號碼', '出廠年份', '備註'], ['普通重型機車', '大型重型機車', '汽車'], is_select_for_notes=False)
        self.add_tab("股票", ['公司名稱', '持股數', '備註'], [], is_select_for_notes=False)
        self.add_income_tab("收入", ['種類', '來源', '總額／元'], ['薪資', '執行業務所得', '其它'])
        self.add_expense_tab("支出", ['種類', '金額（月）', '總額（兩年）'])
        self.add_tab("受扶養人", ['姓名', '關係', '扶養比例', '數額／元／月'], [])

    def add_tab(self, tab_name, headers, select_options, is_select_for_notes=True):
        tab = QWidget()
        layout = QVBoxLayout()

        button_layout = QHBoxLayout()
        add_button = QPushButton("＋")
        save_button = QPushButton("另存新檔")
        button_layout.addWidget(add_button)
        button_layout.addWidget(save_button)
        layout.addLayout(button_layout)

        header_layout = QHBoxLayout()
        for header in headers:
            header_label = QLabel(header)
            header_layout.addWidget(header_label)
        layout.addLayout(header_layout)

        layout.addStretch()
        row_inputs = []
        self.row_inputs.append(row_inputs)
        add_button.clicked.connect(lambda: row_inputs.append(self.add_generic_row(layout, headers, select_options, is_select_for_notes)))
        save_button.clicked.connect(self.save_document)

        tab.setLayout(layout)
        self.tabs.addTab(tab, tab_name)

    def add_income_tab(self, tab_name, headers, select_options):
        tab = QWidget()
        layout = QVBoxLayout()

        button_layout = QHBoxLayout()
        add_button = QPushButton("＋")
        save_button = QPushButton("另存新檔")
        button_layout.addWidget(add_button)
        button_layout.addWidget(save_button)
        layout.addLayout(button_layout)

        header_layout = QHBoxLayout()
        for header in headers:
            header_label = QLabel(header)
            header_layout.addWidget(header_label)
        layout.addLayout(header_layout)

        # 新增顯示總計的標籤
        total_sum_layout = QHBoxLayout()
        total_sum_label = QLabel("總計：")
        total_sum_layout.addWidget(total_sum_label)
        total_sum_layout.addWidget(self.income_total_label)
        layout.addLayout(total_sum_layout)

        layout.addStretch()
        row_inputs = []
        self.row_inputs.append(row_inputs)
        add_button.clicked.connect(lambda: row_inputs.append(self.add_generic_row(layout, headers, select_options)))
        save_button.clicked.connect(self.save_document)

        tab.setLayout(layout)
        self.tabs.addTab(tab, tab_name)

    def add_generic_row(self, layout, headers, select_options, is_select_for_notes=True):
        row_layout = QHBoxLayout()
        inputs = []
        for header in headers:
            if header in ['險種', '種類', '車輛種類']:
                combo = QComboBox()
                combo.addItems(select_options)
                inputs.append(combo)
                row_layout.addWidget(combo)
            elif header == '備註' and is_select_for_notes:
                combo = QComboBox()
                combo.addItems(select_options)
                inputs.append(combo)
                row_layout.addWidget(combo)
            else:
                line_edit = QLineEdit()
                line_edit.textChanged.connect(self.update_total_sum)  # 新增文字變更事件來更新總計
                inputs.append(line_edit)
                row_layout.addWidget(line_edit)
        layout.insertLayout(layout.count() - 1, row_layout)
        return inputs

    def add_expense_tab(self, tab_name, headers):
        tab = QWidget()
        layout = QVBoxLayout()

        button_layout = QHBoxLayout()
        add_button = QPushButton("＋")
        save_button = QPushButton("另存新檔")
        button_layout.addWidget(add_button)
        button_layout.addWidget(save_button)
        layout.addLayout(button_layout)

        grid_layout = QGridLayout()
        layout.addLayout(grid_layout)

        # 新增標題行
        for i, header in enumerate(headers):
            grid_layout.addWidget(QLabel(header), 0, i)

        row_inputs = []
        self.row_inputs.append(row_inputs)

        expense_categories = ['餐費', '水費', '電費', '網路費', '電話費', '勞保費', '健保費', '通勤費', '房租', '扶養費', '日常用品費', '雜支']
        row_num = 1
        for category in expense_categories:
            inputs = self.add_expense_row(grid_layout, row_num, category)
            row_inputs.append(inputs)
            row_num += 1

        add_button.clicked.connect(lambda: row_inputs.append(self.add_expense_row(grid_layout, row_num)))
        save_button.clicked.connect(self.save_document)

        # 新增顯示總計的標籤
        total_sum_layout = QHBoxLayout()
        total_sum_label = QLabel("總計：")
        total_sum_layout.addWidget(total_sum_label)
        total_sum_layout.addWidget(self.expense_total_label)
        layout.addLayout(total_sum_layout)

        tab.setLayout(layout)
        self.tabs.addTab(tab, tab_name)

    def add_expense_row(self, grid_layout, row_num, category="", amount=""):
        category_label = QLabel(category)
        amount_edit = QLineEdit(amount)
        total_label = QLabel("")  # 初始化為空

        # 計算每行的總額
        def update_total():
            try:
                monthly_amount = float(amount_edit.text())
                total_value = int(monthly_amount * 24)  # 取整數，不顯示小數點
                total_label.setText(str(total_value))
                self.update_total_sum()  # 更新總計
            except ValueError:
                total_label.setText("")  # 沒有輸入有效數字時不顯示

        amount_edit.textChanged.connect(update_total)

        grid_layout.addWidget(category_label, row_num, 0)
        grid_layout.addWidget(amount_edit, row_num, 1)
        grid_layout.addWidget(total_label, row_num, 2)

        return [category_label, amount_edit, total_label]

    def update_total_sum(self):
        # 更新收入總計標籤
        income_total_sum = 0
        for row_inputs in self.row_inputs[4]:  # 針對收入表格（第五個表格）
            for input_field in row_inputs:
                if isinstance(input_field, QLineEdit) and input_field.text().strip():
                    try:
                        total_value = int(input_field.text())
                        income_total_sum += total_value
                    except ValueError:
                        print(f"收入總額欄位的值無效：{input_field.text().strip()}，跳過。")
                        pass
        self.income_total_label.setText(str(income_total_sum))  # 更新收入總計標籤的顯示
        print(f"收入總計更新為：{income_total_sum}")

        # 更新支出總計標籤
        expense_total_sum = 0
        for row_inputs in self.row_inputs[5]:  # 針對支出表格（第六個表格）
            for input_field in row_inputs:
                if isinstance(input_field, QLabel) and input_field.text().strip():
                    try:
                        total_value = int(input_field.text())
                        expense_total_sum += total_value
                    except ValueError:
                        print(f"支出總額欄位的值無效：{input_field.text().strip()}，跳過。")
                        pass
        self.expense_total_label.setText(str(expense_total_sum))  # 更新支出總計標籤的顯示
        print(f"支出總計更新為：{expense_total_sum}")

    def save_document(self):
        current_dir = os.path.dirname(os.path.abspath(__file__))
        file_path = os.path.join(current_dir, "document", "B.docx")
        
        doc = Document(file_path)

        headers_list = [
            ['保險公司名稱', '險種', '保單號碼', '金額'],
            ['地號或建號', '公告現值總價額', '持有比例', '備註'],
            ['車輛種類', '車牌號碼', '出廠年份', '備註'],
            ['公司名稱', '持股數', '備註'],
            ['種類', '來源', '總額／元'],
            ['種類', '金額（月）', '總額（兩年）'],
            ['姓名', '關係', '扶養比例', '數額／元／月']
        ]

        # 完全移除所有包含標註的行
        markers = ['A0', 'A1', 'A2', 'A3', 'B0', 'B1', 'B2', 'B3', 'C0', 'C1', 'C2', 'C3', 
                   'D0', 'D1', 'D2', 'F', 'R0', 'S0', 'S1', 'S2', 'S3', 'E0', 'E1', 'E2', 'Q0', 'Q1', 'P0', 'P1', 'O1', 'O0', 'F0', 'F1'
                   , 'G0', 'G1', 'H0', 'H1', 'I0', 'I1', 'J0', 'J1', 'K0', 'K1', 'L0', 'L1', 'M0', 'M1'
                   , 'N0', 'N1']

        for table in doc.tables:
            rows_to_remove = []
            for row in table.rows:
                if any(cell.text.strip() in markers for cell in row.cells):
                    rows_to_remove.append(row)

            for row in rows_to_remove:
                table._element.remove(row._element)

        for headers, row_inputs in zip(headers_list, self.row_inputs):
            table_found = False
            for table in doc.tables:
                if len(table.columns) == len(headers) and all(cell.text.strip() == header for cell, header in zip(table.rows[0].cells, headers)):
                    table_found = True

                    # 清除除標題行和總計行外的所有行
                    while len(table.rows) > 2:
                        table._element.remove(table.rows[1]._element)
                    
                    # 檢查輸入框是否有填寫內容，未填寫則刪除該行
                    valid_rows = [inputs for inputs in row_inputs if any(self.get_input_text(field) for field in inputs)]
                    if not valid_rows:
                        new_row = table.add_row().cells
                        new_row[0].text = "無"
                    else:
                        for inputs in valid_rows:
                            row_cells = table.add_row().cells
                            for j, input_field in enumerate(inputs):
                                row_cells[j].text = self.get_input_text(input_field)
                        # 在表格末尾插入一列用於顯示總計
                        total_row = table.add_row().cells
                        total_row[0].text = "總計"
                        if headers[0] == '種類' and headers == ['種類', '來源', '總額／元']:
                            total_row[2].text = self.income_total_label.text()  # 填入表格五（F 欄位）
                        elif headers[0] == '種類' and headers == ['種類', '金額（月）', '總額（兩年）']:
                            total_row[2].text = self.expense_total_label.text()  # 填入表格六（R0 欄位）

                    break

            if not table_found:
                print(f"Debug: 找不到匹配的表格更新標題: {headers}")  # 調試打印

        default_filename = "02_財產及收入狀況說明書.docx"
        save_path, _ = QFileDialog.getSaveFileName(self, "存檔", default_filename, "Word Files (*.docx);;All Files (*)")
        if save_path:
            doc.save(save_path)
            print(f"檔案已存到: {save_path}")

    def get_input_text(self, input_field):
        # 根據控制項類型回傳文字
        if isinstance(input_field, QLineEdit):
            return input_field.text().strip()
        elif isinstance(input_field, QComboBox):
            return input_field.currentText().strip()
        elif isinstance(input_field, QLabel):
            return input_field.text().strip()
        return ""

if __name__ == '__main__':
    app = QApplication(sys.argv)
    window = TableEditor()
    window.show()
    sys.exit(app.exec_())
