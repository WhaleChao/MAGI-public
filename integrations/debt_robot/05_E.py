import sys
import os
from datetime import datetime
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QScrollArea, QLabel,
    QLineEdit, QTextEdit, QPushButton, QFileDialog, QHBoxLayout, QMessageBox, QComboBox
)
from PyQt5.QtCore import Qt
from docx import Document
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
import re

class DocumentGenerator(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("消債蘿蔔特 - 陳報狀產生器")
        self.setGeometry(200, 100, 1000, 900)
        self.inputs = {}
        self.init_ui()

    def handle_c3_selection(self):
        if self.c3_box.currentIndex() == 1:
            self.c3_input.setVisible(True)
        else:
            self.c3_input.setVisible(False)
            self.c3_input.setText(self.c3_box.currentText())

    def handle_d3_selection(self):
        if self.d3_box.currentIndex() == 1:
            self.d3_text.setVisible(True)
        else:
            self.d3_text.setVisible(False)
#標題
    def init_ui(self):
        container = QWidget()
        layout = QVBoxLayout()
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll_content = QWidget()
        self.scroll_layout = QVBoxLayout(scroll_content)

        def add_input(key, label):
            self.scroll_layout.addWidget(QLabel(label))
            line_edit = QLineEdit()
            self.scroll_layout.addWidget(line_edit)
            self.inputs[key] = line_edit

        def add_textarea(key, label, default_text=""):
            self.scroll_layout.addWidget(QLabel(label))
            text_edit = QTextEdit()
            text_edit.setPlainText(default_text)
            self.scroll_layout.addWidget(text_edit)
            self.inputs[key] = text_edit

        add_input("A1", "一、陳報狀號碼：")
        add_input("A2", "二、案號：")
        add_input("A3", "三、股別：")
        add_input("A4", "四、聲請人名稱：")

        add_textarea("B1", "五、聲請人借款原因：", "聲請人因OO原因而陸續向債權人商借款項，以解其燃眉之急，然因債務及利息循環增生，以致聲請人陷入難以償還之境地，故僅能向 鈞院聲請更生。")
        add_textarea("B2", "六、聲請人調解不成立原因：", "就調解不成立之原因，係因聲請人債權人繁多，且除銀行外，尚有民間債權人，雖銀行同意以每月新臺幣（下同）Ｏ萬塊之方式償還，惟前開民間債權人並未出席調解，如依照聲請人與前開民間債權人之約定，聲請人依然須以每月約Ｏ萬元之金額償還之，前開債務實已讓聲請人不堪負荷，且已積欠數額未繳納，而僅能向 鈞院聲請更生。")
        add_textarea("B3", "七、聲請人目前得履行之更生方案：", "就更生方案部分，聲請人每月現收入為O元；必要性支出為O元，如依消費者債務清理條例第64條之1之規定，聲請人之收入扣除必要支出後之餘額為O元，現聲請人願以72期，每月1期，每期O元之方式償還債務，前開每月期數之金額已超出消費者債務清理條例第64條之1之規定，而已為盡力清償，故懇請 鈞院得依法准許聲請人之聲請。")

        self.scroll_layout.addWidget(QLabel("八、聲請人聲請前兩年之收入："))
        self.c1_box = QComboBox()
        self.c1_options = [
            "聲請人聲請前兩年並無收入。",
            "聲請人聲請前兩年之收入為現金收入，此有收入切結書可稽【聲證O】。",
            "聲請人聲請前兩年之收入詳參薪資單【聲證O】。",
            "聲請人聲請前兩年之收入詳參後述之綜合所得稅各類所得清單。"
        ]
        self.c1_box.addItems(self.c1_options)
        self.scroll_layout.addWidget(self.c1_box)
        self.inputs["C1"] = self.c1_box

        self.scroll_layout.addWidget(QLabel("九、聲請人目前之工作情形："))
        self.c2_box = QComboBox()
        self.c2_options = [
            "聲請人目前並無收入。",
            "聲請人目前之收入為現金收入，此有收入切結書可稽【聲證O】。",
            "聲請人目前之收入詳參薪資單【聲證O】。",
            "聲請人目前之收入詳參薪資現金袋【聲證O】。"
        ]
        self.c2_box.addItems(self.c2_options)
        self.scroll_layout.addWidget(self.c2_box)
        self.inputs["C2"] = self.c2_box

        self.scroll_layout.addWidget(QLabel("十、聲請人受資助情形："))
        self.c3_box = QComboBox()
        self.c3_box.addItems([
            "聲請人目前並未受親友資助生活。",
            "聲請人現由OO關係之親友協助一同生活。"
        ])
        self.c3_input = QLineEdit("聲請人現由OO關係之親友協助一同生活。")
        self.c3_input.setVisible(False)
        self.scroll_layout.addWidget(self.c3_box)
        self.scroll_layout.addWidget(self.c3_input)
        self.inputs["C3"] = self.c3_input
        self.c3_box.currentIndexChanged.connect(self.handle_c3_selection)

        self.scroll_layout.addWidget(QLabel("十一、聲請人現居地資料："))
        self.d1_box = QComboBox()
        self.d1_options = [
            "聲請人現居於OO關係之親友名下之房屋內，此有該房屋之謄本【聲證O】可稽。",
            "聲請人係租屋居住，此有租約及租金支付證明【聲證O】可稽。",
            "聲請人現居於親友之房屋，係無償居住，此有居住同意書【聲證O】可稽。"
        ]
        self.d1_box.addItems(self.d1_options)
        self.scroll_layout.addWidget(self.d1_box)
        self.inputs["D1"] = self.d1_box

        self.scroll_layout.addWidget(QLabel("十二、與聲請人同住之人資料："))
        self.d2_text = QTextEdit("聲請人現與『OO關係之人』一同居住，此有聲請人之全戶戶籍謄本【聲證O】可稽。")
        self.scroll_layout.addWidget(self.d2_text)
        self.inputs["D2"] = self.d2_text

        self.scroll_layout.addWidget(QLabel("十三、聲請人家庭成員資料："))
        self.d3_box = QComboBox()
        self.d3_box.addItems(["刪除此項", "留下空格，供填寫"])
        self.d3_text = QTextEdit("聲請人現家庭成員分別為『OO關係之人』，共O人。")
        self.d3_text.setVisible(False)
        self.scroll_layout.addWidget(self.d3_box)
        self.scroll_layout.addWidget(self.d3_text)
        self.inputs["D3"] = self.d3_text
        self.d3_box.currentIndexChanged.connect(self.handle_d3_selection)

        # D4：全戶戶籍謄本
        self.scroll_layout.addWidget(QLabel("十四、聲請人全戶戶籍謄本："))
        self.d4_box = QComboBox()
        self.d4_box.addItems(["刪除此項", "陳報聲請人之全戶戶籍謄本【聲證O】。"])
        self.scroll_layout.addWidget(self.d4_box)
        self.inputs["D4"] = self.d4_box

        # D5：家族系統表
        self.scroll_layout.addWidget(QLabel("十五、聲請人家族系統表："))
        self.d5_box = QComboBox()
        self.d5_box.addItems(["刪除此項", "陳報聲請人之家族系統表【聲證O】。"])
        self.scroll_layout.addWidget(self.d5_box)
        self.inputs["D5"] = self.d5_box

        # D7：年金與健保
        self.scroll_layout.addWidget(QLabel("十六、聲請人及受扶養人年金及健保資料"))
        self.d7_box = QComboBox()
        self.d7_box.addItems([
            "刪除此項",
            "聲請人尚在申請中，經全數收集完竣將立即陳報 鈞院。",
            "謹陳報聲請人及受其扶養之人之歷年全戶國民年金保險及全民健康保險之投保資料如【聲證O】。"
        ])
        self.scroll_layout.addWidget(self.d7_box)
        self.inputs["D7"] = self.d7_box

        # D8：勞保資料
        self.scroll_layout.addWidget(QLabel("十七、聲請人及受期扶養之人之勞保資料："))
        self.d8_box = QComboBox()
        self.d8_box.addItems([
            "經查詢結果，謹陳報聲請人之勞保資料如【聲證O】",
            "經查詢結果，謹陳報聲請人及受其扶養之人之勞工保險資料詳如【聲證O】。",
            "經聲請人告知，聲請人尚與勞工保險局查詢中，一經確定將立即陳報 鈞院，尚祈 鈞院諒查。"
        ])
        self.scroll_layout.addWidget(self.d8_box)
        self.inputs["D8"] = self.d8_box


        # D10：集保公司資料
        self.scroll_layout.addWidget(QLabel("十八、聲請人集保公司資料"))
        self.d10_box = QComboBox()
        self.d10_box.addItems([
            "謹陳報聲請人之集保公司資料如【聲證O】。",
            "經聲請人告知，聲請人尚在向集保公司申請中，一經取得相關資料將立即陳報 鈞院，尚祈 鈞院諒查。"
        ])
        self.scroll_layout.addWidget(self.d10_box)
        self.inputs["D10"] = self.d10_box

        # D11：壽險資料
        self.scroll_layout.addWidget(QLabel("十九、聲請人壽險資料："))
        self.d11_box = QComboBox()
        self.d11_box.addItems([
            "經查詢結果，聲請人並無投保任何壽險，此有聲請人之人壽保險資料查詢結果可稽【聲證O】。",
            "經查詢結果，聲請人名下之人壽保險及解約金資料詳如【聲證O】。",
            "經聲請人告知，聲請人確有保險，並有聲請人之人壽保險資料查詢結果可稽【聲證O】惟解約金之部分，聲請人尚與保險公司查詢中，一經確定將立即陳報 鈞院，尚祈 鈞院諒查。"
        ])
        self.scroll_layout.addWidget(self.d11_box)
        self.inputs["D11"] = self.d11_box

        # D12：社福津貼
        self.scroll_layout.addWidget(QLabel("二十、聲請人社福津貼情形："))
        self.d12_box = QComboBox()
        self.d12_box.addItems([
            "聲請人及受其扶養之人並無領取任何社會補助或津貼。",
            "聲請人及其受扶養之人領取之社會補助或津貼資料詳如【聲證O】。"
        ])
        self.scroll_layout.addWidget(self.d12_box)
        self.inputs["D12"] = self.d12_box

        # D14：公司營運情形
        self.scroll_layout.addWidget(QLabel("二十一、聲請人公司營運情形："))
        self.d14_box = QComboBox()
        self.d14_box.addItems([
            "聲請人並無擔任商號或公司負責人。",
            "聲請人雖有擔任商號或公司負責人，惟該商號或公司之營業額未超過每月20萬，此有聲請人經營事業之國稅局報表【聲證O】可稽。"
        ])
        self.scroll_layout.addWidget(self.d14_box)
        self.inputs["D14"] = self.d14_box

        # D15：財產書明書等資料
        self.scroll_layout.addWidget(QLabel("二十二、聲請人財產書明書等資料："))
        self.d15_box = QComboBox()
        self.d15_box.addItems(["刪除此項", "謹陳報聲請人之財產及收入狀況說明書、債權人清冊及債務人清冊如【聲證O】。"])
        self.scroll_layout.addWidget(self.d15_box)
        self.inputs["D15"] = self.d15_box

        # D15：財產書明書等資料
        self.scroll_layout.addWidget(QLabel("二十三、審理法院"))
        self.e1_box = QComboBox()
        self.e1_box.addItems(["",  # 默認為空白
            "臺灣臺北地方法院", "臺灣新北地方法院", "臺灣士林地方法院", "臺灣桃園地方法院", 
            "臺灣新竹地方法院", "臺灣苗栗地方法院", "臺灣臺中地方法院", "臺灣南投地方法院", 
            "臺灣彰化地方法院", "臺灣雲林地方法院", "臺灣嘉義地方法院", "臺灣臺南地方法院", 
            "臺灣高雄地方法院", "臺灣橋頭地方法院", "臺灣屏東地方法院", "臺灣臺東地方法院", 
            "臺灣花蓮地方法院", "臺灣宜蘭地方法院", "臺灣基隆地方法院", "臺灣澎湖地方法院", "臺灣高雄少年及家事法院", "福建金門地方法院", "福建連江地方法院"])
        self.scroll_layout.addWidget(self.e1_box)
        self.inputs["E1"] = self.e1_box

                # D13：財產變動
        self.scroll_layout.addWidget(QLabel("二十四、聲請人財產變動情形："))
        self.d13_label = QLabel("書狀將自動寫入「聲請人於聲請前兩年內並無財產變動。」")
        self.scroll_layout.addWidget(self.d13_label)
        self.inputs["D13"] = self.d13_label

                # D9：存摺
        self.scroll_layout.addWidget(QLabel("二十五、聲請人存摺資料"))
        self.d9_label = QLabel("書狀將自動寫入「謹陳報聲情人之所有金融機構之存摺影本或交易明細表如【聲證O】。」")
        self.scroll_layout.addWidget(self.d9_label)
        self.inputs["D9"] = self.d9_label

                # D6：財所清單（自動顯示）
        self.scroll_layout.addWidget(QLabel("二十六、聲請人及受扶養人財所清單"))
        self.d6_label = QLabel("書狀將自動寫入「謹陳報聲情人及受其扶養之人最近兩年之綜合所得稅各類所得清單及財產歸屬資料清單如【聲證O】。」")
        self.scroll_layout.addWidget(self.d6_label)
        self.inputs["D6"] = self.d6_label

#按鈕
        btn_layout = QHBoxLayout()
        save_btn = QPushButton("存檔")
        save_btn.clicked.connect(self.save_doc)
        btn_layout.addWidget(save_btn)
        self.scroll_layout.addLayout(btn_layout)

        scroll.setWidget(scroll_content)
        layout.addWidget(scroll)
        container.setLayout(layout)
        self.setCentralWidget(container)
#自動填寫邏輯區
    
    def apply_inputs_to_doc(self, doc):
        from datetime import datetime
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        import re

        values = {}
        proof_fields_order = [
            "C1", "C2", "D1", "D2", "D3", "D4", "D5", "D6", "D7", "D8",
            "D9", "D10", "D11", "D12", "D13", "D14", "D15"
        ]
        paragraph_hint_map = {
            "D4": "全戶戶籍謄本", "D5": "家族系統表", "D6": "綜合所得稅各類所得清單",
            "D7": "年金及健保資料", "D8": "勞保資料", "D9": "存摺影本或交易明細表",
            "D10": "集保公司資料", "D11": "壽險查詢結果", "D12": "社會補助或津貼",
            "D13": "財產變動情形", "D14": "公司營運情形", "D15": "財產及收入狀況說明書"
        }
        # -------- 收集欄位值 --------
        def get_text(widget):
            if isinstance(widget, QLineEdit):
                return widget.text().strip()
            elif isinstance(widget, QTextEdit):
                return widget.toPlainText().strip()
            elif isinstance(widget, QLabel):
                text = widget.text().strip()
                if "「" in text and "」" in text:
                    return text.split("「", 1)[-1].rsplit("」", 1)[0].strip()
                return text
            elif isinstance(widget, QComboBox):
                return widget.currentText().strip()
            return ""

        proof_targets = []
        for key in self.inputs:
            val = get_text(self.inputs[key])
            if val == "刪除此項":
                values[key] = None
            else:
                values[key] = val
                if "【聲證O】" in val:
                    proof_targets.append(key)

        # -------- 編號聲證 --------
        def num_to_chinese(n):
            numerals = "零一二三四五六七八九十"
            if n <= 10:
                return numerals[n]
            elif n < 20:
                return "十" + (numerals[n % 10] if n % 10 != 0 else "")
            return str(n)

        ordered_proof_targets = [k for k in proof_fields_order if k in proof_targets]
        label_to_proof = {
            key: f"【聲證{num_to_chinese(idx)}】"
            for idx, key in enumerate(ordered_proof_targets, 1)
        }

        for key in values:
            if isinstance(values[key], str) and "【聲證O】" in values[key]:
                values[key] = values[key].replace("【聲證O】", label_to_proof.get(key, ""))

        # -------- 日期欄位 --------
        today = datetime.today()
        roc_year = today.year - 1911
        values["G1"] = f"中華民國{roc_year}年{today.month:02d}月{today.day:02d}日"
        print("✅ G1 民國年結果：", values["G1"])
        # -------- 段落文字替換 --------
        def replace_in_paragraph(para, values):
            original_text = "".join(run.text for run in para.runs)
            new_text = original_text
            changed = False
            for key in sorted(values, key=len, reverse=True):
                val = values[key]
                if val is not None and re.search(rf"\b{re.escape(key)}\b", new_text):
                    new_text = re.sub(rf"\b{re.escape(key)}\b", val, new_text)
                    changed = True
            if not changed:
                return
            base_run = para.runs[0] if para.runs else para.add_run("")
            while para.runs:
                para.runs[0]._element.getparent().remove(para.runs[0]._element)
            new_run = para.add_run(new_text)
            try:
                if base_run.font and base_run.font.name:
                    new_run.font.name = base_run.font.name
                    rPr = new_run._element.get_or_add_rPr()
                    rFonts = rPr.find(qn('w:rFonts'))
                    if rFonts is None:
                        rFonts = OxmlElement('w:rFonts')
                        rPr.append(rFonts)
                    rFonts.set(qn('w:eastAsia'), base_run.font.name)
                new_run.font.size = base_run.font.size
                new_run.bold = base_run.font.bold
                new_run.italic = base_run.font.italic
                new_run.underline = base_run.font.underline
            except Exception as e:
                print(f"⚠️ 樣式錯誤：{e}")
            if "G1" in original_text:
                print("🌀 替換前 G1 段落：", original_text)
                print("🌀 替換後 G1 段落：", new_text)
        # -------- 🧹 刪除段落與標題，並插入空段避免表格黏合 --------
        sorted_keys = sorted(values.keys(), key=len, reverse=True)
        para_indices_to_delete = set()
        insert_blank_after_indices = []

        for i, para in enumerate(doc.paragraphs):
            full_text = "".join(run.text for run in para.runs).strip()
            for key in sorted_keys:
                if values[key] is None:
                    hint = paragraph_hint_map.get(key, "")
                    if key in full_text or hint in full_text:
                        para_indices_to_delete.add(i)

                        # 如果上一段是標題，則也一併刪除，並插入空段
                        if i > 0:
                            prev_para = doc.paragraphs[i - 1]
                            prev_text = "".join(run.text for run in prev_para.runs).strip()
                            if prev_text.startswith(f"{proof_fields_order.index(key)+1}.") or hint in prev_text:
                                para_indices_to_delete.add(i - 1)
                                insert_blank_after_indices.append(i - 1)
                        break
            else:
                replace_in_paragraph(para, values)

        # 實際刪除段落（從後往前）
        for idx in sorted(para_indices_to_delete, reverse=True):
            doc.paragraphs[idx]._element.getparent().remove(doc.paragraphs[idx]._element)

        # 插入空段（避免表格黏合）
        for idx in sorted(insert_blank_after_indices, reverse=True):
            insert_after = doc.paragraphs[idx]._element
            blank_para = OxmlElement("w:p")
            insert_after.addnext(blank_para)


        # -------- 表格處理 --------
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for key in sorted_keys:
                            if values[key] is None and key in para.text:
                                para.clear()
                                break
                        else:
                            replace_in_paragraph(para, values)
        return ordered_proof_targets, label_to_proof
        # -------- 📌 F/H 聲證標籤與附件敘述 --------
    def apply_proof_attachments(self, doc, label_to_proof, ordered_proof_targets):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        # 附件說明詞庫

        proof_description_map = {
            "C1": "聲請人聲請前兩年之證明文件乙份。",
            "C2": "聲請人現工作之證明文件乙份。",
            "D1": "聲請人之現居地資料乙份。",
            "D2": "聲請人之全戶戶籍謄本乙份。",
            "D4": "聲請人之全戶戶籍謄本乙份。",
            "D5": "聲請人之家族系統表乙份。",
            "D6": "聲請人及受其扶養之人最近兩年之綜合所得稅各類所得清單及財產歸屬資料清單乙份。",
            "D7": "聲請人及受其扶養之人歷年全戶國民年金保險及全民健康保險之投保資料乙份。",
            "D8": "聲起人及受其扶養之人之勞保資料乙份。",
            "D9": "聲請人之所有金融機構之存摺影本或交易明細表乙份。",
            "D10": "聲請人之集保公司查詢資料乙份。",
            "D11": "聲請人之壽險資料乙份。",
            "D12": "聲請人之社會補助或津貼請領相關資料乙份。",
            "D14": "聲請人經營之公司行號之國稅局營業稅申報資料表乙份。",
            "D15": "聲請人聲請時之財產及收入狀況說明書、債權人清冊及債務人清冊各乙份。"
        }

        # 建立實際 F/H 替換值
        fh_values = {}
        for i, key in enumerate(ordered_proof_targets):
            fh_values[f"F{i+1}"] = label_to_proof.get(key, "")
            fh_values[f"H{i+1}"] = proof_description_map.get(key, "")

        # 🔸 整理未使用的 F/H（F8~F15、H8~H15 之類）
        unused_fh = []
        for i in range(len(ordered_proof_targets) + 1, 16):
            unused_fh.append(f"F{i}")
            unused_fh.append(f"H{i}")

        # ✅ 將所有段落集中處理（doc.paragraphs + 表格內）
        all_paragraphs = list(doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_paragraphs.extend(cell.paragraphs)

        # ❌ 刪除包含未使用 F/H 的段落
        def safe_remove_paragraph(para):
            from docx.oxml import OxmlElement
            parent = para._element.getparent()
            parent.remove(para._element)

            # 🔍 如果這段是表格儲存格，且已經沒內容，就補一個空段
            if parent.tag.endswith("tc") and not any(c.tag.endswith("p") for c in parent):
                empty_p = OxmlElement("w:p")
                parent.append(empty_p)

        
        for para in all_paragraphs[:]:  # 複製避免錯誤
            if any(unused in para.text for unused in unused_fh):
                try:
                    safe_remove_paragraph(para)
                except Exception as e:
                    print(f"⚠️ 無法刪除段落：{e}")

        # 🔁 替換 F/H 區段內容
        for para in all_paragraphs:
            full_text = "".join(run.text for run in para.runs)
            if not full_text:
                continue

            new_text = full_text
            replaced = False
            for key in sorted(fh_values.keys(), key=lambda k: int(k[1:]), reverse=True):
                if key in new_text:
                    new_text = new_text.replace(key, fh_values[key])
                    replaced = True
            if replaced:
                        base_run = para.runs[0] if para.runs else para.add_run("")
                        while para.runs:
                            para.runs[0]._element.getparent().remove(para.runs[0]._element)

                        new_run = para.add_run(new_text)
                        try:
                            new_run.font.name = "新細明體"
                            rPr = new_run._element.get_or_add_rPr()
                            rFonts = rPr.find(qn('w:rFonts'))
                            if rFonts is None:
                                rFonts = OxmlElement('w:rFonts')
                                rPr.append(rFonts)
                            rFonts.set(qn('w:eastAsia'), "新細明體")

                            new_run.font.size = base_run.font.size
                            new_run.bold = base_run.font.bold
                            new_run.italic = base_run.font.italic
                            new_run.underline = base_run.font.underline
                        except Exception as e:
                            print("⚠️ F/H 替換樣式錯誤：", e)
    def save_doc(self):
        import os
        import re
        from datetime import datetime
        from PyQt5.QtWidgets import QFileDialog, QMessageBox
        from docx import Document

        # 📁 模板位置
        template_path = os.path.join(os.path.dirname(__file__), "document", "D.docx")
        if not os.path.exists(template_path):
            QMessageBox.critical(self, "錯誤", "找不到 D.docx 模板")
            return

        try:
            # 📄 載入模板
            doc = Document(template_path)

            # 🧠 套用欄位與聲證替換邏輯
            ordered_proof_targets, label_to_proof = self.apply_inputs_to_doc(doc)
            self.apply_proof_attachments(doc, label_to_proof, ordered_proof_targets)

            # 📆 民國日期
            today = datetime.today()
            roc_year = today.year - 1911
            date_str = f"{roc_year:03d}{today.month:02d}{today.day:02d}"

            # 🔍 取得 A1 / A4 值
            def get_text(widget):
                if hasattr(widget, "text"):
                    return widget.text().strip()
                elif hasattr(widget, "toPlainText"):
                    return widget.toPlainText().strip()
                elif hasattr(widget, "currentText"):
                    return widget.currentText().strip()
                return ""

            a1_raw = get_text(self.inputs.get("A1"))
            a4_raw = get_text(self.inputs.get("A4"))

            # ✅ 避免非法檔名字元
            def clean_filename(text):
                return re.sub(r'[\\/*?:"<>|]', '', text or '')

            a1 = clean_filename(a1_raw or "未填A1")
            a4 = clean_filename(a4_raw or "未填A4")

            filename = f"{date_str}_消費者債務清理陳報({a1})狀({a4}).docx"

            # 💾 使用者選擇儲存位置
            save_result = QFileDialog.getSaveFileName(
                self,
                "另存 Word 文件",
                filename,
                "Word 文件 (*.docx)"
            )

            if not save_result or not save_result[0]:  # 使用者按取消
                QMessageBox.information(self, "取消儲存", "你尚未選擇儲存路徑，動作已取消。")
                return

            save_path = save_result[0]

            # 📥 儲存 Word 文件
            doc.save(save_path)
            QMessageBox.information(self, "完成", f"文件已儲存至：\n{save_path}")

        except PermissionError:
            QMessageBox.warning(self, "儲存失敗", "檔案正在使用中，請關閉檔案後再試。")

        except Exception as e:
            import traceback
            print("❗ 發生錯誤：", e)
            print(traceback.format_exc())
            QMessageBox.critical(self, "儲存錯誤", f"無法儲存文件：\n{str(e)}")

#GUI介面
if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = DocumentGenerator()
    window.show()
    sys.exit(app.exec_())
