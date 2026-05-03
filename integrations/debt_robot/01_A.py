import sys
import os
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QLabel, QLineEdit, QComboBox, QPushButton, QFileDialog, QTextEdit
)
from docx import Document

class DocumentEditor(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("消債羅伯特-聲請狀")
        self.setGeometry(100, 100, 800, 600)
        self.init_ui()

    def init_ui(self):
        main_layout = QVBoxLayout()

        # 姓名和地址輸入框
        name_address_layout = QHBoxLayout()
        self.name_input = QLineEdit()
        self.name_input.setPlaceholderText("輸入姓名")
        self.address_input = QLineEdit()
        self.address_input.setPlaceholderText("輸入地址")
        name_address_layout.addWidget(QLabel("姓名："))
        name_address_layout.addWidget(self.name_input)
        name_address_layout.addWidget(QLabel("地址："))
        name_address_layout.addWidget(self.address_input)
        main_layout.addLayout(name_address_layout)

        # 開啟資料夾按鈕
        self.open_button = QPushButton("開啟聲請人資料夾，讀取財產及債權人清冊")
        self.open_button.clicked.connect(self.open_folder)
        main_layout.addWidget(self.open_button)

        # 顯示總額
        self.total_income_label = QLabel("資產總價值：0")
        self.total_debt_label = QLabel("債務總金額：0")
        main_layout.addWidget(self.total_income_label)
        main_layout.addWidget(self.total_debt_label)

        # 顯示銀行名稱
        self.bank_name_label = QLabel("最大債權銀行：")
        main_layout.addWidget(self.bank_name_label)

        # 選擇法院資訊
        court_layout = QVBoxLayout()
        court_options = [
            "",  # 預設為空白
            "臺灣臺北地方法院", "臺灣新北地方法院", "臺灣士林地方法院", "臺灣桃園地方法院", 
            "臺灣新竹地方法院", "臺灣苗栗地方法院", "臺灣臺中地方法院", "臺灣南投地方法院", 
            "臺灣彰化地方法院", "臺灣雲林地方法院", "臺灣嘉義地方法院", "臺灣臺南地方法院", 
            "臺灣高雄地方法院", "臺灣橋頭地方法院", "臺灣屏東地方法院", "臺灣臺東地方法院", 
            "臺灣花蓮地方法院", "臺灣宜蘭地方法院", "臺灣基隆地方法院", "臺灣澎湖地方法院", 
            "臺灣高雄少年及家事法院", "福建金門地方法院", "福建連江地方法院"
        ]
        self.court_combo_box_d0 = QComboBox()
        self.court_combo_box_d0.addItems(court_options)
        self.court_input_d1 = QLineEdit()
        self.court_combo_box_e0 = QComboBox()
        self.court_combo_box_e0.addItems(court_options)
        
        court_layout.addWidget(QLabel("選擇執行法院名稱："))
        court_layout.addWidget(self.court_combo_box_d0)
        court_layout.addWidget(QLabel("輸入執行案號："))
        court_layout.addWidget(self.court_input_d1)
        court_layout.addWidget(QLabel("選擇聲請法院名稱："))
        court_layout.addWidget(self.court_combo_box_e0)
        main_layout.addLayout(court_layout)

        # F1 輸入框
        f1_layout = QVBoxLayout()
        self.f1_input = QTextEdit()
        self.f1_input.setPlaceholderText("已有財產收入狀況說明書、債權人清冊、債務人清冊、債務人戶籍謄本、綜合所得資料清單影本、財產資料清單影本、金融機構債權人清冊影本，請輸入其餘附件名稱。")
        f1_layout.addWidget(QLabel("附件名稱："))
        f1_layout.addWidget(self.f1_input)
        main_layout.addLayout(f1_layout)

        # 儲存按鈕
        self.save_button = QPushButton("存檔")
        self.save_button.clicked.connect(self.save_document)
        main_layout.addWidget(self.save_button)

        container = QWidget()
        container.setLayout(main_layout)
        self.setCentralWidget(container)

    def open_folder(self):
        folder_path = QFileDialog.getExistingDirectory(self, "選擇資料夾")
        if folder_path:
            self.load_documents(folder_path)

    def load_documents(self, folder_path):
        income_doc_path = os.path.join(folder_path, "02_財產及收入狀況說明書.docx")
        debt_doc_path = os.path.join(folder_path, "03_債權人清冊.docx")

        # 讀取 02_財產及收入狀況說明書.docx 的表格五總額
        if os.path.exists(income_doc_path):
            income_doc = Document(income_doc_path)
            income_total = "0"
            if len(income_doc.tables) >= 5:
                table_five = income_doc.tables[4]  # 取得第五個表格
                for row in table_five.rows:
                    if "總計" in row.cells[0].text:
                        income_total = row.cells[-1].text.strip()  # 取得最後一個單元格的值
                        break
                self.total_income_label.setText(f"資產總價值：{income_total}")

        # 讀取 03_債權人清冊.docx 的表格1和表格2
        if os.path.exists(debt_doc_path):
            debt_doc = Document(debt_doc_path)
            if len(debt_doc.tables) >= 2:
                table1 = debt_doc.tables[0]
                debt_total = "0"
                for row in table1.rows:
                    if row.cells and len(row.cells) > 1:
                        debt_total = row.cells[1].text.strip()
                        break
                self.total_debt_label.setText(f"債務總金額：{debt_total}")

                # 查找第二個表格中的最大債權
                table2 = debt_doc.tables[1]
                max_value = -float('inf')
                bank_name = ""
                for row in table2.rows:
                    try:
                        value = int(row.cells[2].text.strip().replace(',', ''))  # 將第三欄的文本轉換為整數
                        if value > max_value and row.cells[0].text.strip():  # 確保第一欄有銀行名稱
                            max_value = value
                            bank_name = row.cells[0].text.strip()  # 記錄第一欄的銀行名稱
                    except ValueError:
                        continue  # 如果轉換失敗（不是數字），跳過此行
                self.bank_name_label.setText(f"最大債權銀行：{bank_name}")

    def save_document(self):
        # 打開並編輯 A.docx
        current_dir = os.path.dirname(os.path.abspath(__file__))
        a_doc_path = os.path.join(current_dir, "document", "A.docx")
        doc = Document(a_doc_path)

        # 更新姓名和地址
        table1 = doc.tables[0]
        table1.cell(0, 1).text = self.name_input.text()  # A0
        table1.cell(0, 2).text = self.address_input.text()  # A1

        # 更新資產總價值和債務總金額
        table2 = doc.tables[1]
        table2.cell(0, 2).text = self.total_income_label.text().split("：")[-1]  # B0
        table2.cell(1, 2).text = self.total_debt_label.text().split("：")[-1]  # B1

        # 更新最大債權銀行名稱
        table3 = doc.tables[2]
        table3.cell(0, 0).text = self.bank_name_label.text().split("：")[-1]  # C0

        # 更新法院資訊和其他資訊
        table4 = doc.tables[3]
        table4.cell(0, 0).text = self.court_combo_box_d0.currentText() if self.court_combo_box_d0.currentText() else "無"  # D0
        table4.cell(0, 1).text = self.court_input_d1.text()  # D1

        table5 = doc.tables[4]
        table5.cell(0, 0).text = self.court_combo_box_e0.currentText()  # E0

        # 更新 F1 內容
        table6 = doc.tables[5]
        table6.cell(1, 0).text = self.f1_input.toPlainText()  # F1

        # 更新具狀人姓名
        table7 = doc.tables[6]
        table7.cell(0, 2).text = self.name_input.text() 
        # G0

        # 預設檔名為 01_消費者債務清理聲請狀（表格內A0的文字）.docx
        default_filename = f"01_消費者債務清理聲請狀（{self.name_input.text()}）.docx"
        save_path, _ = QFileDialog.getSaveFileName(self, "存檔", default_filename, "Word 文件 (*.docx);;所有文件 (*)")

        if save_path:
            doc.save(save_path)
            print(f"文件已存到: {save_path}")

if __name__ == '__main__':
    app = QApplication(sys.argv)
    window = DocumentEditor()
    window.show()
    sys.exit(app.exec_())
