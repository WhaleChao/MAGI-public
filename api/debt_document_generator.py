# -*- coding: utf-8 -*-
"""
MAGI 消費者債務清理文件產生器 (Debt Document Generator)

整合自 Robot 資料夾的五個 PyQt5 桌面工具，改為純後端 API 模式：
  01_A.py → generate_application()        聲請狀
  02_B.py → generate_asset_statement()     財產及收入狀況說明書
  03_C.py → generate_creditor_list()       債權人清冊
  04_D.py → merge_debt_pdfs()             合併 PDF
  05_E.py → generate_report()             陳報狀

改善紀錄：
  - 移除 PyQt5 GUI 依賴，改為純資料驅動函式
  - if/else 判讀改為字典映射 (dict-based dispatch)
  - 法院清單、費用類別等硬編碼改為常數配置
  - CSV 地址查詢改為可擴展的資料層
  - 聲證編號邏輯改為自動化管線
  - 統一錯誤處理與日誌
  - 新增驗證函式、批次生成、自動匯入、地址儲存等功能
"""

from __future__ import annotations

import csv
import json
import logging
import os
import re
from copy import deepcopy
from datetime import datetime
from typing import Any, Optional

logger = logging.getLogger("DebtDocumentGenerator")

_MAGI_ROOT = os.path.abspath(os.path.dirname(os.path.dirname(__file__)))
_TEMPLATE_DIR = os.path.join(_MAGI_ROOT, "templates", "debt_templates")

# ═══════════════════════════════════════════════════════════════
# 常數配置（原本散落在各 PY 檔的硬編碼）
# ═══════════════════════════════════════════════════════════════

COURT_OPTIONS = [
    # 地方法院
    "臺灣臺北地方法院", "臺灣新北地方法院", "臺灣士林地方法院", "臺灣桃園地方法院",
    "臺灣新竹地方法院", "臺灣苗栗地方法院", "臺灣臺中地方法院", "臺灣南投地方法院",
    "臺灣彰化地方法院", "臺灣雲林地方法院", "臺灣嘉義地方法院", "臺灣臺南地方法院",
    "臺灣高雄地方法院", "臺灣橋頭地方法院", "臺灣屏東地方法院", "臺灣臺東地方法院",
    "臺灣花蓮地方法院", "臺灣宜蘭地方法院", "臺灣基隆地方法院", "臺灣澎湖地方法院",
    "福建金門地方法院", "福建連江地方法院",
    # 專業法院
    "臺灣高雄少年及家事法院", "智慧財產及商業法院",
    # 高等法院
    "臺灣高等法院", "臺灣高等法院臺中分院", "臺灣高等法院臺南分院",
    "臺灣高等法院高雄分院", "臺灣高等法院花蓮分院",
    # 高等行政法院
    "臺北高等行政法院", "臺中高等行政法院", "高雄高等行政法院",
    # 最高法院
    "最高法院", "最高行政法院",
]

EXPENSE_CATEGORIES = [
    "餐費", "水費", "電費", "網路費", "電話費", "勞保費",
    "健保費", "通勤費", "房租", "扶養費", "日常用品費", "雜支",
]

INSURANCE_TYPES = ["一般壽險", "實支實付險"]
LAND_NOTES = ["公同共有難以變賣", "持份過低難以變賣"]
VEHICLE_TYPES = ["普通重型機車", "大型重型機車", "汽車"]
INCOME_TYPES = ["薪資", "執行業務所得", "其它"]
DEBT_TYPES = ["信用卡", "信用貸款", "信用卡及信用貸款"]

# 法定費用參考額度（勞保費/健保費）
STATUTORY_EXPENSE_REFERENCES = {
    "勞保費": 1755,      # 2024年一級保險費基數月額
    "健保費": 2659,      # 2024年基層保費月額（二類被保險人）
}

# 陳報狀下拉選項映射 (key → list of options)
REPORT_OPTIONS = {
    "C1": [
        "聲請人聲請前兩年並無收入。",
        "聲請人聲請前兩年之收入為現金收入，此有收入切結書可稽【聲證O】。",
        "聲請人聲請前兩年之收入詳參薪資單【聲證O】。",
        "聲請人聲請前兩年之收入詳參後述之綜合所得稅各類所得清單。",
    ],
    "C2": [
        "聲請人目前並無收入。",
        "聲請人目前之收入為現金收入，此有收入切結書可稽【聲證O】。",
        "聲請人目前之收入詳參薪資單【聲證O】。",
        "聲請人目前之收入詳參薪資現金袋【聲證O】。",
    ],
    "C3": [
        "聲請人目前並未受親友資助生活。",
        "聲請人現由OO關係之親友協助一同生活。",
    ],
    "D1": [
        "聲請人現居於OO關係之親友名下之房屋內，此有該房屋之謄本【聲證O】可稽。",
        "聲請人係租屋居住，此有租約及租金支付證明【聲證O】可稽。",
        "聲請人現居於親友之房屋，係無償居住，此有居住同意書【聲證O】可稽。",
    ],
    "D4": ["刪除此項", "陳報聲請人之全戶戶籍謄本【聲證O】。"],
    "D5": ["刪除此項", "陳報聲請人之家族系統表【聲證O】。"],
    "D7": [
        "刪除此項",
        "聲請人尚在申請中，經全數收集完竣將立即陳報 鈞院。",
        "謹陳報聲請人及受其扶養之人之歷年全戶國民年金保險及全民健康保險之投保資料如【聲證O】。",
    ],
    "D8": [
        "經查詢結果，謹陳報聲請人之勞保資料如【聲證O】",
        "經查詢結果，謹陳報聲請人及受其扶養之人之勞工保險資料詳如【聲證O】。",
        "經聲請人告知，聲請人尚與勞工保險局查詢中，一經確定將立即陳報 鈞院，尚祈 鈞院諒查。",
    ],
    "D10": [
        "謹陳報聲請人之集保公司資料如【聲證O】。",
        "經聲請人告知，聲請人尚在向集保公司申請中，一經取得相關資料將立即陳報 鈞院，尚祈 鈞院諒查。",
    ],
    "D11": [
        "經查詢結果，聲請人並無投保任何壽險，此有聲請人之人壽保險資料查詢結果可稽【聲證O】。",
        "經查詢結果，聲請人名下之人壽保險及解約金資料詳如【聲證O】。",
        "經聲請人告知，聲請人確有保險，並有聲請人之人壽保險資料查詢結果可稽【聲證O】惟解約金之部分，聲請人尚與保險公司查詢中，一經確定將立即陳報 鈞院，尚祈 鈞院諒查。",
    ],
    "D12": [
        "聲請人及受其扶養之人並無領取任何社會補助或津貼。",
        "聲請人及其受扶養之人領取之社會補助或津貼資料詳如【聲證O】。",
    ],
    "D14": [
        "聲請人並無擔任商號或公司負責人。",
        "聲請人雖有擔任商號或公司負責人，惟該商號或公司之營業額未超過每月20萬，此有聲請人經營事業之國稅局報表【聲證O】可稽。",
    ],
    "D15": ["刪除此項", "謹陳報聲請人之財產及收入狀況說明書、債權人清冊及債務人清冊如【聲證O】。"],
}

# 陳報狀聲證附件說明對照表
PROOF_DESCRIPTION_MAP = {
    "C1": "聲請人聲請前兩年之證明文件乙份。",
    "C2": "聲請人現工作之證明文件乙份。",
    "D1": "聲請人之現居地資料乙份。",
    "D2": "聲請人之全戶戶籍謄本乙份。",
    "D4": "聲請人之全戶戶籍謄本乙份。",
    "D5": "聲請人之家族系統表乙份。",
    "D6": "聲請人及受其扶養之人最近兩年之綜合所得稅各類所得清單及財產歸屬資料清單乙份。",
    "D7": "聲請人及受其扶養之人歷年全戶國民年金保險及全民健康保險之投保資料乙份。",
    "D8": "聲請人及受其扶養之人之勞保資料乙份。",
    "D9": "聲請人之所有金融機構之存摺影本或交易明細表乙份。",
    "D10": "聲請人之集保公司查詢資料乙份。",
    "D11": "聲請人之壽險資料乙份。",
    "D12": "聲請人之社會補助或津貼請領相關資料乙份。",
    "D14": "聲請人經營之公司行號之國稅局營業稅申報資料表乙份。",
    "D15": "聲請人聲請時之財產及收入狀況說明書、債權人清冊及債務人清冊各乙份。",
}

# 自動填入的固定欄位（陳報狀）
REPORT_AUTO_FIELDS = {
    "D6": "謹陳報聲請人及受其扶養之人最近兩年之綜合所得稅各類所得清單及財產歸屬資料清單如【聲證O】。",
    "D9": "謹陳報聲請人之所有金融機構之存摺影本或交易明細表如【聲證O】。",
    "D13": "聲請人於聲請前兩年內並無財產變動。",
}

# 聲證欄位排序
PROOF_FIELDS_ORDER = [
    "C1", "C2", "D1", "D2", "D3", "D4", "D5", "D6", "D7", "D8",
    "D9", "D10", "D11", "D12", "D13", "D14", "D15",
]

# ═══════════════════════════════════════════════════════════════
# 證據資料夾檔名 → 陳報狀欄位自動偵測映射
# ═══════════════════════════════════════════════════════════════
#
# 每個欄位對應一組關鍵字，掃描案件的「證據資料」子資料夾時，
# 若檔名包含任何一個關鍵字，即判定該欄位對應的證據已備齊。
#
# suggested_option_index: 當偵測到證據時，建議選擇 REPORT_OPTIONS 中的第幾個選項
#   （0-based；若為 None 表示該欄位為自動填入欄位，無需使用者選擇）
#
EVIDENCE_SCAN_MAP: dict[str, dict] = {
    "C1": {
        "keywords": ["收入切結", "切結書", "薪資證明", "扣繳憑單", "薪資單", "所得清單"],
        "suggested_option_index": 1,  # 收入切結書
        "description": "聲請前兩年收入證明",
    },
    "C2": {
        "keywords": ["在職證明", "薪資單", "薪資袋", "薪資現金袋", "僱用證明", "勞動契約"],
        "suggested_option_index": 2,  # 薪資單
        "description": "目前工作收入證明",
    },
    "D1": {
        "keywords": ["謄本", "租約", "租賃", "居住同意", "房屋"],
        "suggested_option_index": 0,  # 親友房屋謄本
        "description": "現居地資料",
    },
    "D2": {
        "keywords": ["戶籍謄本", "全戶戶籍"],
        "suggested_option_index": None,  # 文字欄位，非選項
        "description": "同住之人（全戶戶籍謄本）",
    },
    "D4": {
        "keywords": ["戶籍謄本", "全戶戶籍"],
        "suggested_option_index": 1,  # 陳報聲請人之全戶戶籍謄本
        "description": "全戶戶籍謄本",
    },
    "D5": {
        "keywords": ["家族系統", "系統表", "親屬系統"],
        "suggested_option_index": 1,  # 陳報聲請人之家族系統表
        "description": "家族系統表",
    },
    "D6": {
        "keywords": ["所得清單", "所得稅", "財產歸屬", "綜合所得", "財產清單"],
        "suggested_option_index": None,  # REPORT_AUTO_FIELDS 自動帶入
        "description": "財所清單",
    },
    "D7": {
        "keywords": ["年金", "國民年金", "健保", "全民健康保險", "投保資料"],
        "suggested_option_index": 2,  # 謹陳報聲請人及受其扶養之人之歷年全戶國民年金...
        "description": "年金及健保資料",
    },
    "D8": {
        "keywords": ["勞保", "勞工保險", "勞保資料"],
        "suggested_option_index": 0,  # 經查詢結果，謹陳報聲請人之勞保資料
        "description": "勞保資料",
    },
    "D9": {
        "keywords": ["存摺", "交易明細", "帳戶明細", "銀行對帳"],
        "suggested_option_index": None,  # REPORT_AUTO_FIELDS 自動帶入
        "description": "存摺或交易明細",
    },
    "D10": {
        "keywords": ["集保", "集中保管", "有價證券"],
        "suggested_option_index": 0,  # 謹陳報聲請人之集保公司資料
        "description": "集保公司資料",
    },
    "D11": {
        "keywords": ["壽險", "人壽保險", "保險查詢", "解約金", "保單價值"],
        "suggested_option_index": 1,  # 聲請人名下之人壽保險及解約金資料
        "description": "壽險資料",
    },
    "D12": {
        "keywords": ["社福", "津貼", "社會補助", "補助金", "低收入"],
        "suggested_option_index": 1,  # 領取之社會補助或津貼資料
        "description": "社福津貼",
    },
    "D14": {
        "keywords": ["國稅局", "營業稅", "營業額", "公司登記", "商業登記", "負責人"],
        "suggested_option_index": 1,  # 有擔任負責人，附國稅局報表
        "description": "公司營運情形",
    },
    "D15": {
        "keywords": ["財產及收入", "財產說明", "債權人清冊", "債務人清冊"],
        "suggested_option_index": 1,  # 謹陳報聲請人之財產及收入狀況說明書...
        "description": "財產說明書等資料",
    },
}


def scan_evidence_folder(folder_path: str) -> dict:
    """
    掃描案件的證據資料子資料夾，根據檔名匹配 EVIDENCE_SCAN_MAP。

    Args:
        folder_path: 案件根資料夾路徑（Windows canonical 或本地路徑皆可）

    Returns:
        {
            "ok": True/False,
            "evidence_folder": "實際掃描的路徑",
            "files": ["檔名1.pdf", ...],
            "matches": {
                "D4": {"found": True, "matched_files": ["全戶戶籍謄本.pdf"], "suggested_option_index": 1, "description": "..."},
                "D7": {"found": False, "matched_files": [], ...},
                ...
            },
            "summary": "找到 8/14 項證據資料"
        }
    """
    # 找到「證據資料」子資料夾（不同案件類別編號不同）
    evidence_folder = ""
    if os.path.isdir(folder_path):
        for entry in os.listdir(folder_path):
            full = os.path.join(folder_path, entry)
            if os.path.isdir(full) and "證據資料" in entry:
                evidence_folder = full
                break

    if not evidence_folder or not os.path.isdir(evidence_folder):
        return {
            "ok": False,
            "error": f"找不到證據資料夾（搜尋路徑: {folder_path}）",
            "evidence_folder": "",
            "files": [],
            "matches": {},
        }

    # 列出所有檔案（遞迴掃描，含子資料夾）
    all_files = []
    for root, _dirs, files in os.walk(evidence_folder):
        for f in files:
            if f.startswith("."):
                continue
            all_files.append(f)

    # 執行匹配
    matches = {}
    found_count = 0
    for field_key, config in EVIDENCE_SCAN_MAP.items():
        matched_files = []
        for filename in all_files:
            name_lower = filename.lower()
            for kw in config["keywords"]:
                if kw.lower() in name_lower:
                    matched_files.append(filename)
                    break
        found = len(matched_files) > 0
        if found:
            found_count += 1
        matches[field_key] = {
            "found": found,
            "matched_files": matched_files,
            "suggested_option_index": config["suggested_option_index"],
            "description": config["description"],
        }

    total = len(EVIDENCE_SCAN_MAP)
    return {
        "ok": True,
        "evidence_folder": evidence_folder,
        "files": all_files,
        "file_count": len(all_files),
        "matches": matches,
        "found_count": found_count,
        "total_fields": total,
        "summary": f"找到 {found_count}/{total} 項證據資料",
    }


# ═══════════════════════════════════════════════════════════════
# 結構化錯誤回傳
# ═══════════════════════════════════════════════════════════════

class ValidationError(Exception):
    """驗證錯誤，包含欄位級別的錯誤訊息"""
    def __init__(self, errors: dict[str, list[str]]):
        self.errors = errors
        super().__init__(json.dumps(errors, ensure_ascii=False))


def create_error_response(field_errors: dict[str, list[str]]) -> dict:
    """建立結構化的錯誤回應"""
    return {
        "success": False,
        "errors": field_errors,
        "message": f"驗證失敗，共有 {len(field_errors)} 個欄位有問題"
    }


# ═══════════════════════════════════════════════════════════════
# 驗證函式
# ═══════════════════════════════════════════════════════════════

def validate_application_data(data: dict[str, Any]) -> tuple[bool, dict[str, list[str]]]:
    """驗證聲請狀資料"""
    errors = {}

    # 必填欄位
    if not data.get("name", "").strip():
        errors.setdefault("name", []).append("聲請人姓名為必填欄位")

    if not data.get("address", "").strip():
        errors.setdefault("address", []).append("聲請人地址為必填欄位")

    # 數值驗證
    try:
        asset = float(str(data.get("asset_total", 0)))
        if asset < 0:
            errors.setdefault("asset_total", []).append("資產總價值不能為負數")
    except (ValueError, TypeError):
        errors.setdefault("asset_total", []).append("資產總價值必須為數字")

    try:
        debt = float(str(data.get("debt_total", 0)))
        if debt < 0:
            errors.setdefault("debt_total", []).append("債務總金額不能為負數")
    except (ValueError, TypeError):
        errors.setdefault("debt_total", []).append("債務總金額必須為數字")

    return len(errors) == 0, errors


def validate_asset_statement_data(data: dict[str, Any]) -> tuple[bool, dict[str, list[str]]]:
    """驗證財產及收入狀況說明書資料"""
    errors = {}

    # 驗證保險清單
    for i, item in enumerate(data.get("insurance", [])):
        prefix = f"insurance[{i}]"
        if item.get("amount"):
            try:
                float(item["amount"])
            except (ValueError, TypeError):
                errors.setdefault(prefix, []).append("保險金額必須為數字")

    # 驗證土地清單
    for i, item in enumerate(data.get("land", [])):
        prefix = f"land[{i}]"
        if item.get("value"):
            try:
                float(item["value"])
            except (ValueError, TypeError):
                errors.setdefault(prefix, []).append("公告現值必須為數字")

    # 驗證車輛清單
    for i, item in enumerate(data.get("vehicles", [])):
        prefix = f"vehicles[{i}]"
        if item.get("type") and item["type"] not in VEHICLE_TYPES:
            errors.setdefault(prefix, []).append(f"車輛種類必須為: {', '.join(VEHICLE_TYPES)}")

    # 驗證收入清單
    for i, item in enumerate(data.get("income", [])):
        prefix = f"income[{i}]"
        if item.get("amount"):
            try:
                float(item["amount"])
            except (ValueError, TypeError):
                errors.setdefault(prefix, []).append("收入總額必須為數字")

    # 驗證支出清單
    for i, item in enumerate(data.get("expenses", [])):
        prefix = f"expenses[{i}]"
        if item.get("monthly"):
            try:
                float(item["monthly"])
            except (ValueError, TypeError):
                errors.setdefault(prefix, []).append("月支出必須為數字")

    # 驗證受扶養人清單
    for i, item in enumerate(data.get("dependents", [])):
        prefix = f"dependents[{i}]"
        if item.get("amount"):
            try:
                float(item["amount"])
            except (ValueError, TypeError):
                errors.setdefault(prefix, []).append("扶養金額必須為數字")

    return len(errors) == 0, errors


def validate_creditor_list_data(data: dict[str, Any]) -> tuple[bool, dict[str, list[str]]]:
    """驗證債權人清冊資料"""
    errors = {}

    if not data.get("creditors"):
        errors.setdefault("creditors", []).append("至少需要填寫一個債權人")
        return False, errors

    for i, item in enumerate(data.get("creditors", [])):
        prefix = f"creditors[{i}]"

        if not item.get("name", "").strip():
            errors.setdefault(prefix, []).append("債權人名稱為必填欄位")

        if item.get("amount"):
            try:
                float(str(item["amount"]).replace(",", ""))
            except (ValueError, TypeError):
                errors.setdefault(prefix, []).append("債權額必須為數字")
        else:
            errors.setdefault(prefix, []).append("債權額為必填欄位")

    return len(errors) == 0, errors


def validate_report_data(data: dict[str, Any]) -> tuple[bool, dict[str, list[str]]]:
    """驗證陳報狀資料"""
    errors = {}

    if not data.get("A4", "").strip():
        errors.setdefault("A4", []).append("聲請人名稱為必填欄位")

    return len(errors) == 0, errors


# ═══════════════════════════════════════════════════════════════
# 工具函式
# ═══════════════════════════════════════════════════════════════

def _num_to_chinese(n: int) -> str:
    """數字轉中文（一～二十）"""
    numerals = "零一二三四五六七八九十"
    if n <= 10:
        return numerals[n]
    elif n < 20:
        return "十" + (numerals[n % 10] if n % 10 != 0 else "")
    return str(n)


def _roc_date_str(dt: datetime | None = None) -> str:
    """產生民國日期字串"""
    dt = dt or datetime.now()
    roc_year = dt.year - 1911
    return f"中華民國{roc_year}年{dt.month:02d}月{dt.day:02d}日"


def _set_font_style(run, font_name='標楷體', size_pt=12, bold=False):
    """統一設定 run 的字型樣式"""
    from docx.shared import Pt
    from docx.oxml.ns import qn
    font = run.font
    font.name = font_name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    font.size = Pt(size_pt)
    font.bold = bold


def load_address_data() -> dict:
    """
    載入銀行 / 公司地址 CSV，回傳 {名稱: 地址} 字典。
    改善：原本分散在 03_C.py 的 CSV 讀取邏輯，統一為單一函式。
    """
    address_map = {}
    for filename in ["all adress - bank.csv", "all adress - company.csv"]:
        filepath = os.path.join(_TEMPLATE_DIR, filename)
        if not os.path.exists(filepath):
            logger.warning("地址 CSV 不存在: %s", filepath)
            continue
        try:
            with open(filepath, newline='', encoding='utf-8') as f:
                reader = csv.reader(f)
                header = next(reader, None)  # skip header
                for row in reader:
                    if len(row) >= 2 and row[0].strip():
                        address_map[row[0].strip()] = row[1].strip()
        except Exception as e:
            logger.warning("讀取地址 CSV 失敗 %s: %s", filename, e)
    return address_map


def get_address_options() -> dict:
    """回傳前端下拉選項所需的地址資料"""
    banks = []
    companies = []
    for filename, target in [("all adress - bank.csv", banks), ("all adress - company.csv", companies)]:
        filepath = os.path.join(_TEMPLATE_DIR, filename)
        if not os.path.exists(filepath):
            continue
        try:
            with open(filepath, newline='', encoding='utf-8') as f:
                reader = csv.reader(f)
                next(reader, None)
                for row in reader:
                    if len(row) >= 2:
                        target.append({"name": row[0].strip(), "address": row[1].strip()})
        except Exception as _e:
            logging.getLogger("magi.debt_doc").warning("Failed to load CSV %s: %s", filepath, _e)
    return {"banks": banks, "companies": companies}


def get_expense_reference(expense_type: str) -> Optional[int]:
    """
    返回指定支出類型的法定參考金額。
    例如：勞保費、健保費等有法定基數。
    """
    return STATUTORY_EXPENSE_REFERENCES.get(expense_type)


def save_address_to_csv(creditor_name: str, address: str, csv_type: str = "bank") -> bool:
    """
    將新增的債權人地址保存回 CSV 檔案。

    Args:
        creditor_name: 債權人/銀行名稱
        address: 地址
        csv_type: "bank" 或 "company"

    Returns:
        是否成功保存
    """
    filename = "all adress - bank.csv" if csv_type == "bank" else "all adress - company.csv"
    filepath = os.path.join(_TEMPLATE_DIR, filename)

    if not os.path.exists(filepath):
        logger.warning("地址 CSV 檔案不存在: %s", filepath)
        return False

    try:
        # 讀取現有資料
        existing_rows = []
        with open(filepath, newline='', encoding='utf-8') as f:
            reader = csv.reader(f)
            header = next(reader, None)
            for row in reader:
                existing_rows.append(row)

        # 檢查是否已存在
        for row in existing_rows:
            if len(row) >= 1 and row[0].strip() == creditor_name:
                # 更新地址
                row[1] = address if len(row) > 1 else address
                break
        else:
            # 新增記錄
            existing_rows.append([creditor_name, address])

        # 寫回 CSV
        with open(filepath, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            if header:
                writer.writerow(header)
            writer.writerows(existing_rows)

        logger.info("地址已保存: %s -> [REDACTED]", creditor_name[:1] + "**")
        return True
    except Exception as e:
        logger.error("保存地址到 CSV 失敗: %s", e)
        return False


# ═══════════════════════════════════════════════════════════════
# 1. 聲請狀產生器 (原 01_A.py)
# ═══════════════════════════════════════════════════════════════

def generate_application(data: dict[str, Any]) -> "Document":
    """
    產生消費者債務清理聲請狀。

    data 欄位:
      - name: 聲請人姓名
      - address: 聲請人地址
      - asset_total: 資產總價值
      - debt_total: 債務總金額
      - max_creditor_bank: 最大債權銀行名稱
      - execution_court: 執行法院名稱
      - execution_case_no: 執行案號
      - application_court: 聲請法院名稱
      - attachments: 附件名稱（文字）
    """
    from docx import Document

    template_path = os.path.join(_TEMPLATE_DIR, "A.docx")
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"找不到聲請狀範本: {template_path}")

    doc = Document(template_path)

    name = data.get("name", "")
    address = data.get("address", "")
    asset_total = str(data.get("asset_total", "0"))
    debt_total = str(data.get("debt_total", "0"))
    max_bank = data.get("max_creditor_bank", "")
    exec_court = data.get("execution_court", "")
    exec_case_no = data.get("execution_case_no", "")
    app_court = data.get("application_court", "")
    attachments = data.get("attachments", "")

    # 表格映射（改善：原本是 table1/table2/... 硬編碼索引）
    table_mapping = [
        (0, [(0, 1, name), (0, 2, address)]),                    # 姓名、地址
        (1, [(0, 2, asset_total), (1, 2, debt_total)]),          # 資產、債務
        (2, [(0, 0, max_bank)]),                                  # 最大債權銀行
        (3, [(0, 0, exec_court or "無"), (0, 1, exec_case_no)]), # 執行法院
        (4, [(0, 0, app_court)]),                                 # 聲請法院
        (5, [(1, 0, attachments)]),                               # 附件
        (6, [(0, 2, name)]),                                      # 具狀人
    ]

    for table_idx, cells in table_mapping:
        if table_idx < len(doc.tables):
            table = doc.tables[table_idx]
            for row, col, value in cells:
                if row < len(table.rows) and col < len(table.rows[row].cells):
                    table.cell(row, col).text = value

    return doc


# ═══════════════════════════════════════════════════════════════
# 2. 財產及收入狀況說明書 (原 02_B.py)
# ═══════════════════════════════════════════════════════════════

def generate_asset_statement(data: dict[str, Any]) -> "Document":
    """
    產生財產及收入狀況說明書。

    data 欄位:
      - insurance: [{company, type, policy_no, amount}, ...]
      - land: [{location, value, ratio, note}, ...]
      - vehicles: [{type, plate, year, note}, ...]
      - stocks: [{company, shares, note}, ...]
      - income: [{type, source, amount}, ...]
      - expenses: [{category, monthly}, ...]       月支出，自動×24
      - dependents: [{name, relation, ratio, amount}, ...]
    """
    from docx import Document

    template_path = os.path.join(_TEMPLATE_DIR, "B.docx")
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"找不到財產說明書範本: {template_path}")

    doc = Document(template_path)

    # 定義表格結構（改善：原本用 headers_list 陣列對比表頭文字）
    SECTION_CONFIG = {
        "insurance":  {"headers": ['保險公司名稱', '險種', '保單號碼', '金額'], "fields": ["company", "type", "policy_no", "amount"]},
        "land":       {"headers": ['地號或建號', '公告現值總價額', '持有比例', '備註'], "fields": ["location", "value", "ratio", "note"]},
        "vehicles":   {"headers": ['車輛種類', '車牌號碼', '出廠年份', '備註'], "fields": ["type", "plate", "year", "note"]},
        "stocks":     {"headers": ['公司名稱', '持股數', '備註'], "fields": ["company", "shares", "note"]},
        "income":     {"headers": ['種類', '來源', '總額／元'], "fields": ["type", "source", "amount"]},
        "expenses":   {"headers": ['種類', '金額（月）', '總額（兩年）'], "fields": ["category", "monthly", "total"]},
        "dependents": {"headers": ['姓名', '關係', '扶養比例', '數額／元／月'], "fields": ["name", "relation", "ratio", "amount"]},
    }

    # 先移除所有標記行（原本用 markers 列表）
    markers = set()
    for prefix in "ABCDEFGHIJKLMNOPQRS":
        for suffix in range(4):
            markers.add(f"{prefix}{suffix}")

    for table in doc.tables:
        rows_to_remove = []
        for row in table.rows:
            if any(cell.text.strip() in markers for cell in row.cells):
                rows_to_remove.append(row)
        for row in rows_to_remove:
            table._element.remove(row._element)

    # 處理每個表格區塊
    for section_key, config in SECTION_CONFIG.items():
        headers = config["headers"]
        fields = config["fields"]
        items = data.get(section_key, [])

        # 找到對應的表格
        target_table = None
        for table in doc.tables:
            if len(table.columns) == len(headers):
                first_row_texts = [cell.text.strip() for cell in table.rows[0].cells]
                if first_row_texts == headers:
                    target_table = table
                    break

        if not target_table:
            logger.warning("找不到表格: %s", headers)
            continue

        # 清除資料行（保留表頭和最後一行）
        while len(target_table.rows) > 2:
            target_table._element.remove(target_table.rows[1]._element)

        if not items:
            new_row = target_table.add_row().cells
            new_row[0].text = "無"
        else:
            # 計算支出的兩年總額
            if section_key == "expenses":
                for item in items:
                    try:
                        monthly = float(item.get("monthly", 0) or 0)
                        item["total"] = str(int(monthly * 24))
                    except (ValueError, TypeError):
                        item["total"] = ""

            for item in items:
                row_cells = target_table.add_row().cells
                for j, field in enumerate(fields):
                    row_cells[j].text = str(item.get(field, "") or "")

            # 加入總計行
            total_row = target_table.add_row().cells
            total_row[0].text = "總計"

            if section_key == "income":
                total = sum(int(item.get("amount", 0) or 0) for item in items if str(item.get("amount", "")).strip())
                total_row[2].text = str(total)
            elif section_key == "expenses":
                total = sum(int(item.get("total", 0) or 0) for item in items if str(item.get("total", "")).strip())
                total_row[2].text = str(total)

    return doc


# ═══════════════════════════════════════════════════════════════
# 3. 債權人清冊 (原 03_C.py)
# ═══════════════════════════════════════════════════════════════

def generate_creditor_list(data: dict[str, Any]) -> "Document":
    """
    產生債權人清冊。

    data 欄位:
      - creditors: [{name, address, amount, debt_type}, ...]
      - auto_lookup_address: bool (是否自動查詢地址)
    """
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn

    template_path = os.path.join(_TEMPLATE_DIR, "C.docx")
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"找不到債權人清冊範本: {template_path}")

    doc = Document(template_path)
    creditors = data.get("creditors", [])
    auto_lookup = data.get("auto_lookup_address", True)

    # 地址自動補全（改善：原本在 GUI callback 裡做）
    if auto_lookup:
        address_map = load_address_data()
        for c in creditors:
            if not c.get("address") and c.get("name"):
                c["address"] = address_map.get(c["name"], "")

    # 計算總金額
    total_amount = 0
    for c in creditors:
        try:
            total_amount += int(str(c.get("amount", "0")).replace(",", ""))
        except (ValueError, TypeError):
            pass

    # 填入第一個表格的總金額
    if doc.tables:
        cell = doc.tables[0].cell(0, 1)
        cell.text = str(total_amount)
        if cell.paragraphs[0].runs:
            run = cell.paragraphs[0].runs[0]
            run.font.name = 'DFKai-SB'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
            run.font.size = Pt(16)

    # 處理第二個表格（債權人明細）
    if len(doc.tables) >= 2:
        second_table = doc.tables[1]

        # 移除標記行
        markers = set()
        for prefix in "BCDEFGH":
            for suffix in range(4):
                markers.add(f"{prefix}{suffix}")

        rows_to_remove = [
            row for row in second_table.rows
            if any(cell.text.strip() in markers for cell in row.cells)
        ]
        for row in rows_to_remove:
            second_table._element.remove(row._element)

        # 填入債權人資料
        for c in creditors:
            new_row = second_table.add_row().cells
            new_row[0].text = c.get("name", "")
            new_row[1].text = c.get("address", "")
            new_row[2].text = str(c.get("amount", ""))
            new_row[3].text = c.get("debt_type", "")

    return doc


# ═══════════════════════════════════════════════════════════════
# 4. PDF 合併工具 (原 04_D.py)
# ═══════════════════════════════════════════════════════════════

def merge_debt_pdfs(file_paths: list[str], output_path: str | None = None,
                   add_bookmarks: bool = False) -> str:
    """
    合併多個 PDF/DOCX 檔案，奇數頁自動補空白頁。

    改善：
      - 原本用 tkinter GUI 選檔案，改為接收路徑列表
      - DOCX 轉 PDF 改為可選依賴
      - 回傳輸出路徑
      - 新增 add_bookmarks 參數支援 PDF 書籤/TOC

    Args:
        file_paths: PDF 或 DOCX 檔案路徑列表
        output_path: 輸出路徑（預設自動產生）
        add_bookmarks: 是否為合併後的 PDF 加入書籤（需要 PyPDF2 >= 2.0）

    Returns:
        輸出的 PDF 檔案路徑
    """
    from PyPDF2 import PdfReader, PdfWriter

    if not output_path:
        stamp = datetime.now().strftime("%Y%m%d")
        export_dir = os.path.join(_MAGI_ROOT, "exports")
        os.makedirs(export_dir, exist_ok=True)
        output_path = os.path.join(export_dir, f"{stamp}_消費者債務清理調解聲請狀及附件.pdf")

    pdf_writer = PdfWriter()
    page_count = 0

    for fpath in file_paths:
        fpath = str(fpath).strip()
        if not os.path.exists(fpath):
            logger.warning("檔案不存在，跳過: %s", fpath)
            continue

        # DOCX → PDF 轉換
        if fpath.lower().endswith(".docx"):
            try:
                from docx2pdf import convert
                pdf_path = fpath.replace(".docx", ".pdf")
                convert(fpath, pdf_path)
                fpath = pdf_path
            except ImportError:
                logger.warning("docx2pdf 未安裝，無法轉換 DOCX: %s", fpath)
                continue
            except Exception as e:
                logger.warning("DOCX 轉 PDF 失敗 %s: %s", fpath, e)
                continue

        try:
            reader = PdfReader(fpath)

            # 可選：為每份文件的起始頁添加書籤
            if add_bookmarks:
                doc_name = os.path.splitext(os.path.basename(fpath))[0]
                try:
                    pdf_writer.add_bookmark(doc_name, page_count)
                except Exception as e:
                    logger.debug("新增書籤失敗: %s", e)

            for page in reader.pages:
                pdf_writer.add_page(page)
                page_count += 1

            # 奇數頁補空白頁（改善：提取為清楚的條件判斷）
            if len(reader.pages) % 2 != 0:
                pdf_writer.add_blank_page()
                page_count += 1
        except Exception as e:
            logger.warning("讀取 PDF 失敗 %s: %s", fpath, e)
            continue

    with open(output_path, 'wb') as f:
        pdf_writer.write(f)

    logger.info("PDF 合併完成: %s (%d 頁)", output_path, page_count)
    return output_path


# ═══════════════════════════════════════════════════════════════
# 5. 陳報狀產生器 (原 05_E.py)
# ═══════════════════════════════════════════════════════════════

def _apply_inputs_to_doc_improved(doc, values: dict, paragraph_hint_map: dict):
    """
    改進的段落與表格內容替換邏輯。
    相容 05_E.py 的複雜邏輯，處理段落、表格、多行文本等。
    """
    from docx.oxml.ns import qn

    def _replace_in_paragraph(para, vals):
        """在段落內進行文字替換，保留樣式"""
        original_text = "".join(run.text for run in para.runs)
        new_text = original_text
        changed = False
        for key in sorted(vals, key=len, reverse=True):
            val = vals[key]
            if val is not None and re.search(rf"\b{re.escape(key)}\b", new_text):
                new_text = re.sub(rf"\b{re.escape(key)}\b", val, new_text)
                changed = True
        if not changed:
            return
        base_run = para.runs[0] if para.runs else para.add_run("")
        while para.runs:
            para.runs[0]._element.getparent().remove(para.runs[0]._element)
        new_run = para.add_run(new_text)
        try:
            if base_run.font and base_run.font.name:
                new_run.font.name = base_run.font.name
                rPr = new_run._element.get_or_add_rPr()
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    from docx.oxml import OxmlElement
                    rFonts = OxmlElement('w:rFonts')
                    rPr.append(rFonts)
                rFonts.set(qn('w:eastAsia'), base_run.font.name)
            new_run.font.size = base_run.font.size
            new_run.bold = base_run.font.bold
            new_run.italic = base_run.font.italic
            new_run.underline = base_run.font.underline
        except Exception as e:
            logger.debug("樣式複製失敗: %s", e)

    # 刪除不需要的段落
    sorted_keys = sorted(values.keys(), key=len, reverse=True)
    para_indices_to_delete = set()

    for i, para in enumerate(doc.paragraphs):
        full_text = "".join(run.text for run in para.runs).strip()
        for key in sorted_keys:
            if values[key] is None:
                hint = paragraph_hint_map.get(key, "")
                if key in full_text or (hint and hint in full_text):
                    para_indices_to_delete.add(i)
                    # 上一段如果是標題也刪除
                    if i > 0:
                        prev_text = "".join(run.text for run in doc.paragraphs[i - 1].runs).strip()
                        if hint and hint in prev_text:
                            para_indices_to_delete.add(i - 1)
                    break
        else:
            _replace_in_paragraph(para, values)

    # 從後往前刪除段落
    for idx in sorted(para_indices_to_delete, reverse=True):
        doc.paragraphs[idx]._element.getparent().remove(doc.paragraphs[idx]._element)

    # 表格內替換
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for key in sorted_keys:
                        if values[key] is None and key in para.text:
                            para.clear()
                            break
                    else:
                        _replace_in_paragraph(para, values)


def auto_import_from_docs(asset_statement_path: str, creditor_list_path: str) -> dict:
    """
    從已生成的財產說明書和債權人清冊自動匯入資料到聲請狀。
    這模仿原 01_A.py 的 `load_documents()` 方法。

    Args:
        asset_statement_path: 財產及收入狀況說明書 DOCX 檔案路徑
        creditor_list_path: 債權人清冊 DOCX 檔案路徑

    Returns:
        包含 asset_total, debt_total 等的資料字典
    """
    from docx import Document

    extracted_data = {}

    try:
        # 從財產說明書抽取資產總額（對應原 01_A.py 的 load_documents）
        # 原邏輯：讀取第五個表格（index 4），找到「總計」行的最後一個 cell
        if asset_statement_path and os.path.exists(asset_statement_path):
            asset_doc = Document(asset_statement_path)
            if len(asset_doc.tables) >= 5:
                table_five = asset_doc.tables[4]
                for row in table_five.rows:
                    if "總計" in row.cells[0].text:
                        extracted_data["asset_total"] = row.cells[-1].text.strip()
                        break
            logger.info("從財產說明書讀取資料: %s", asset_statement_path)
    except Exception as e:
        logger.error("讀取財產說明書失敗: %s", e)

    try:
        # 從債權人清冊抽取總債務金額和最大債權銀行（對應原 01_A.py 的 load_documents）
        # 原邏輯：
        #   - 表格1(index 0) cell(0,1) = 債務總金額
        #   - 表格2(index 1) 找第三欄(index 2)金額最大的行，取第一欄(index 0)為銀行名
        if creditor_list_path and os.path.exists(creditor_list_path):
            creditor_doc = Document(creditor_list_path)
            if creditor_doc.tables:
                # 讀取總金額
                first_table = creditor_doc.tables[0]
                try:
                    raw = first_table.cell(0, 1).text.strip().replace(",", "")
                    extracted_data["debt_total"] = int(raw) if raw.isdigit() else raw
                except Exception:
                    logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 1061, exc_info=True)

                # 讀取最大債權銀行
                if len(creditor_doc.tables) >= 2:
                    second_table = creditor_doc.tables[1]
                    max_value = -1
                    bank_name = ""
                    for row in second_table.rows:
                        try:
                            value = int(row.cells[2].text.strip().replace(",", ""))
                            if value > max_value and row.cells[0].text.strip():
                                max_value = value
                                bank_name = row.cells[0].text.strip()
                        except (ValueError, IndexError):
                            continue
                    if bank_name:
                        extracted_data["max_bank"] = bank_name

            logger.info("從債權人清冊讀取資料: %s", creditor_list_path)
    except Exception as e:
        logger.error("讀取債權人清冊失敗: %s", e)

    return extracted_data


def generate_report(data: dict[str, Any]) -> "Document":
    """
    產生消費者債務清理陳報狀。

    這是最複雜的文件，包含：
      - 多個下拉選項對應不同段落文字
      - 自動聲證編號（【聲證一】、【聲證二】...）
      - 刪除不需要的段落
      - F/H 附件區替換

    data 欄位:
      - A1: 陳報狀號碼
      - A2: 案號
      - A3: 股別
      - A4: 聲請人名稱
      - B1: 借款原因（長文）
      - B2: 調解不成立原因（長文）
      - B3: 更生方案（長文）
      - C1: 聲請前兩年收入 (index or text)
      - C2: 目前工作情形 (index or text)
      - C3: 受資助情形 (text)
      - D1: 現居地 (index or text)
      - D2: 同住之人（長文）
      - D3: 家庭成員（text or null=刪除）
      - D4~D15: 各項陳報（index or text, null=刪除）
      - E1: 審理法院
    """
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    template_path = os.path.join(_TEMPLATE_DIR, "D.docx")
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"找不到陳報狀範本: {template_path}")

    doc = Document(template_path)

    # ── 收集所有欄位值 ──
    values = {}
    proof_targets = []

    # 簡單文字欄位
    for key in ["A1", "A2", "A3", "A4"]:
        values[key] = str(data.get(key, "") or "").strip()

    # 長文欄位（改善：原本用 QTextEdit 的 default text，現在用 data 直接帶入）
    for key in ["B1", "B2", "B3", "D2"]:
        values[key] = str(data.get(key, "") or "").strip()

    # 下拉選項欄位（改善：原本 if/else 判斷 index，現在接受 index 或直接文字）
    for key in ["C1", "C2", "C3", "D1", "D4", "D5", "D7", "D8", "D10", "D11", "D12", "D14", "D15"]:
        val = data.get(key)
        if val is None:
            values[key] = None  # 刪除此項
        elif isinstance(val, int) and key in REPORT_OPTIONS:
            options = REPORT_OPTIONS[key]
            if 0 <= val < len(options):
                text = options[val]
                values[key] = None if text == "刪除此項" else text
            else:
                values[key] = None
        else:
            text = str(val).strip()
            values[key] = None if text == "刪除此項" else text

    # D3 特殊處理（原 05_E.py: 0=刪除此項, 1=留空讓使用者編輯的文字模板）
    d3_val = data.get("D3")
    d3_options = ["刪除此項", "聲請人現家庭成員分別為『OO關係之人』，共O人。"]
    if d3_val is None or d3_val == "刪除此項" or (isinstance(d3_val, int) and d3_val == 0):
        values["D3"] = None
    elif isinstance(d3_val, int):
        values["D3"] = d3_options[d3_val] if 0 <= d3_val < len(d3_options) else None
    else:
        text = str(d3_val).strip()
        values["D3"] = None if text == "刪除此項" else text

    # 自動填入欄位
    for key, text in REPORT_AUTO_FIELDS.items():
        if key not in values or values[key] is None:
            if key not in data or data.get(key) is None:
                values[key] = text
            # 如果使用者明確設了 None/刪除，保留 None

    # 法院
    values["E1"] = str(data.get("E1", "") or "").strip()

    # ── 聲證自動編號（改善：原本在 apply_inputs_to_doc 裡用複雜的迴圈）──
    for key, val in values.items():
        if isinstance(val, str) and "【聲證O】" in val:
            proof_targets.append(key)

    ordered_proof_targets = [k for k in PROOF_FIELDS_ORDER if k in proof_targets]
    label_to_proof = {
        key: f"【聲證{_num_to_chinese(idx)}】"
        for idx, key in enumerate(ordered_proof_targets, 1)
    }

    # 替換聲證標記
    for key in values:
        if isinstance(values[key], str) and "【聲證O】" in values[key]:
            values[key] = values[key].replace("【聲證O】", label_to_proof.get(key, ""))

    # 日期
    values["G1"] = _roc_date_str()

    # ── 段落替換（使用改進的邏輯） ──
    paragraph_hint_map = {
        "D4": "全戶戶籍謄本", "D5": "家族系統表", "D6": "綜合所得稅各類所得清單",
        "D7": "年金及健保資料", "D8": "勞保資料", "D9": "存摺影本或交易明細表",
        "D10": "集保公司資料", "D11": "壽險查詢結果", "D12": "社會補助或津貼",
        "D13": "財產變動情形", "D14": "公司營運情形", "D15": "財產及收入狀況說明書",
    }

    # 呼叫改進的替換邏輯
    _apply_inputs_to_doc_improved(doc, values, paragraph_hint_map)

    # ── F/H 聲證附件替換 ──
    fh_values = {}
    for i, key in enumerate(ordered_proof_targets):
        fh_values[f"F{i+1}"] = label_to_proof.get(key, "")
        fh_values[f"H{i+1}"] = PROOF_DESCRIPTION_MAP.get(key, "")

    unused_fh = []
    for i in range(len(ordered_proof_targets) + 1, 16):
        unused_fh.append(f"F{i}")
        unused_fh.append(f"H{i}")

    # 收集所有段落
    all_paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paragraphs.extend(cell.paragraphs)

    # 刪除未使用的 F/H 段落
    for para in all_paragraphs[:]:
        if any(unused in para.text for unused in unused_fh):
            try:
                parent = para._element.getparent()
                parent.remove(para._element)
                if parent.tag.endswith("tc") and not any(c.tag.endswith("p") for c in parent):
                    empty_p = OxmlElement("w:p")
                    parent.append(empty_p)
            except Exception as e:
                logger.debug("無法刪除 F/H 段落: %s", e)

    # 替換 F/H 值
    for para in all_paragraphs:
        full_text = "".join(run.text for run in para.runs)
        if not full_text:
            continue
        new_text = full_text
        replaced = False
        for key in sorted(fh_values.keys(), key=lambda k: int(k[1:]), reverse=True):
            if key in new_text:
                new_text = new_text.replace(key, fh_values[key])
                replaced = True
        if replaced:
            base_run = para.runs[0] if para.runs else para.add_run("")
            while para.runs:
                para.runs[0]._element.getparent().remove(para.runs[0]._element)
            new_run = para.add_run(new_text)
            try:
                new_run.font.name = "新細明體"
                rPr = new_run._element.get_or_add_rPr()
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = OxmlElement('w:rFonts')
                    rPr.append(rFonts)
                rFonts.set(qn('w:eastAsia'), "新細明體")
                new_run.font.size = base_run.font.size
                new_run.bold = base_run.font.bold
                new_run.italic = base_run.font.italic
                new_run.underline = base_run.font.underline
            except Exception as e:
                logger.debug("F/H 樣式錯誤: %s", e)

    return doc


# ═══════════════════════════════════════════════════════════════
# 批次生成函式
# ═══════════════════════════════════════════════════════════════

def generate_all_documents(base_data: dict[str, Any], output_dir: str | None = None) -> dict[str, str]:
    """
    一鍵生成所有五份消債文件。

    Args:
        base_data: 包含所有表單資料的字典，key 為表單類型
                   例如: {
                       "application": {...},
                       "asset_statement": {...},
                       "creditor_list": {...},
                       "report": {...}
                   }
        output_dir: 輸出目錄（預設為 exports/）

    Returns:
        {document_type: filepath, ...} 生成的檔案路徑映射
    """
    from docx import Document

    if output_dir is None:
        output_dir = os.path.join(_MAGI_ROOT, "exports")
    os.makedirs(output_dir, exist_ok=True)

    results = {}

    try:
        # 聲請狀
        app_data = base_data.get("application", {})
        valid, errors = validate_application_data(app_data)
        if not valid:
            logger.error("聲請狀驗證失敗: %s", errors)
        else:
            doc = generate_application(app_data)
            filepath = os.path.join(output_dir, "01_聲請狀.docx")
            doc.save(filepath)
            results["application"] = filepath
            logger.info("聲請狀已生成: %s", filepath)
    except Exception as e:
        logger.error("聲請狀生成失敗: %s", e)

    try:
        # 財產及收入狀況說明書
        asset_data = base_data.get("asset_statement", {})
        valid, errors = validate_asset_statement_data(asset_data)
        if not valid:
            logger.error("財產說明書驗證失敗: %s", errors)
        else:
            doc = generate_asset_statement(asset_data)
            filepath = os.path.join(output_dir, "02_財產及收入狀況說明書.docx")
            doc.save(filepath)
            results["asset_statement"] = filepath
            logger.info("財產說明書已生成: %s", filepath)
    except Exception as e:
        logger.error("財產說明書生成失敗: %s", e)

    try:
        # 債權人清冊
        creditor_data = base_data.get("creditor_list", {})
        valid, errors = validate_creditor_list_data(creditor_data)
        if not valid:
            logger.error("債權人清冊驗證失敗: %s", errors)
        else:
            doc = generate_creditor_list(creditor_data)
            filepath = os.path.join(output_dir, "03_債權人清冊.docx")
            doc.save(filepath)
            results["creditor_list"] = filepath
            logger.info("債權人清冊已生成: %s", filepath)
    except Exception as e:
        logger.error("債權人清冊生成失敗: %s", e)

    try:
        # 陳報狀
        report_data = base_data.get("report", {})
        valid, errors = validate_report_data(report_data)
        if not valid:
            logger.error("陳報狀驗證失敗: %s", errors)
        else:
            doc = generate_report(report_data)
            filepath = os.path.join(output_dir, "05_陳報狀.docx")
            doc.save(filepath)
            results["report"] = filepath
            logger.info("陳報狀已生成: %s", filepath)
    except Exception as e:
        logger.error("陳報狀生成失敗: %s", e)

    try:
        # PDF 合併
        pdf_files = base_data.get("pdf_files", [])
        if pdf_files:
            pdf_output = merge_debt_pdfs(pdf_files, output_dir=output_dir)
            results["merged_pdf"] = pdf_output
            logger.info("PDF 已合併: %s", pdf_output)
    except Exception as e:
        logger.error("PDF 合併失敗: %s", e)

    return results


# ═══════════════════════════════════════════════════════════════
# API 回傳用：取得各表單的欄位定義（供前端渲染表單）
# ═══════════════════════════════════════════════════════════════

def get_form_schema(form_type: str) -> dict:
    """
    回傳指定表單類型的欄位定義，供 OSC 前端動態渲染表單。

    改善：原本每個 PyQt5 視窗各自 hardcode UI，
    現在統一為 JSON schema，前端可動態生成。
    """
    schemas = {
        "application": {
            "title": "消費者債務清理聲請狀",
            "description": "填寫聲請人基本資料，系統將自動產生聲請狀文件",
            "fields": [
                {"key": "name", "label": "聲請人姓名", "type": "text", "required": True},
                {"key": "address", "label": "聲請人地址", "type": "text", "required": True},
                {"key": "asset_total", "label": "資產總價值", "type": "number", "default": "0"},
                {"key": "debt_total", "label": "債務總金額", "type": "number", "default": "0"},
                {"key": "max_creditor_bank", "label": "最大債權銀行", "type": "text"},
                {"key": "execution_court", "label": "執行法院名稱", "type": "select", "options": [""] + COURT_OPTIONS},
                {"key": "execution_case_no", "label": "執行案號", "type": "text"},
                {"key": "application_court", "label": "聲請法院名稱", "type": "select", "options": [""] + COURT_OPTIONS},
                {"key": "attachments", "label": "附件名稱", "type": "textarea",
                 "placeholder": "已有財產收入狀況說明書、債權人清冊、債務人清冊、債務人戶籍謄本、綜合所得資料清單影本、財產資料清單影本、金融機構債權人清冊影本，請輸入其餘附件名稱。"},
            ],
        },
        "asset_statement": {
            "title": "財產及收入狀況說明書",
            "description": "填寫財產、收入、支出及受扶養人資料",
            "sections": [
                {"key": "insurance", "label": "保險", "type": "table",
                 "columns": [
                     {"key": "company", "label": "保險公司名稱", "type": "text"},
                     {"key": "type", "label": "險種", "type": "select", "options": INSURANCE_TYPES},
                     {"key": "policy_no", "label": "保單號碼", "type": "text"},
                     {"key": "amount", "label": "金額", "type": "number"},
                 ]},
                {"key": "land", "label": "土地及建物", "type": "table",
                 "columns": [
                     {"key": "location", "label": "地號或建號", "type": "text"},
                     {"key": "value", "label": "公告現值總價額", "type": "number"},
                     {"key": "ratio", "label": "持有比例", "type": "text"},
                     {"key": "note", "label": "備註", "type": "select", "options": LAND_NOTES},
                 ]},
                {"key": "vehicles", "label": "車輛", "type": "table",
                 "columns": [
                     {"key": "type", "label": "車輛種類", "type": "select", "options": VEHICLE_TYPES},
                     {"key": "plate", "label": "車牌號碼", "type": "text"},
                     {"key": "year", "label": "出廠年份", "type": "text"},
                     {"key": "note", "label": "備註", "type": "text"},
                 ]},
                {"key": "stocks", "label": "股票", "type": "table",
                 "columns": [
                     {"key": "company", "label": "公司名稱", "type": "text"},
                     {"key": "shares", "label": "持股數", "type": "text"},
                     {"key": "note", "label": "備註", "type": "text"},
                 ]},
                {"key": "income", "label": "收入", "type": "table", "show_total": True, "total_column": "amount",
                 "columns": [
                     {"key": "type", "label": "種類", "type": "select", "options": INCOME_TYPES},
                     {"key": "source", "label": "來源", "type": "text"},
                     {"key": "amount", "label": "總額／元", "type": "number"},
                 ]},
                {"key": "expenses", "label": "支出", "type": "table", "show_total": True, "total_column": "monthly",
                 "total_multiplier": 24, "preset_categories": EXPENSE_CATEGORIES,
                 "columns": [
                     {"key": "category", "label": "種類", "type": "text", "readonly_if_preset": True},
                     {"key": "monthly", "label": "金額（月）", "type": "number"},
                     {"key": "total", "label": "總額（兩年）", "type": "computed", "formula": "monthly * 24"},
                 ]},
                {"key": "dependents", "label": "受扶養人", "type": "table",
                 "columns": [
                     {"key": "name", "label": "姓名", "type": "text"},
                     {"key": "relation", "label": "關係", "type": "text"},
                     {"key": "ratio", "label": "扶養比例", "type": "text"},
                     {"key": "amount", "label": "數額／元／月", "type": "number"},
                 ]},
            ],
        },
        "creditor_list": {
            "title": "債權人清冊",
            "description": "填寫債權人資料，系統將自動計算總金額並查詢地址",
            "fields": [
                {"key": "creditors", "label": "債權人", "type": "table",
                 "columns": [
                     {"key": "name", "label": "債權人名稱", "type": "autocomplete", "source": "address_data"},
                     {"key": "address", "label": "地址", "type": "text", "auto_fill": True},
                     {"key": "amount", "label": "債權額", "type": "number"},
                     {"key": "debt_type", "label": "債務類型", "type": "select", "options": DEBT_TYPES},
                 ]},
            ],
        },
        "pdf_merge": {
            "title": "合併PDF檔案",
            "description": "上傳多個PDF或DOCX檔案，合併為單一PDF（奇數頁自動補空白頁）",
            "fields": [
                {"key": "files", "label": "選擇檔案", "type": "file_list", "accept": ".pdf,.docx"},
                {"key": "add_bookmarks", "label": "添加書籤", "type": "checkbox", "default": False},
            ],
        },
        "report": {
            "title": "陳報狀產生器",
            "description": "填寫陳報內容，系統自動編號聲證並產生文件",
            "fields": [
                {"key": "A1", "label": "陳報狀號碼", "type": "text"},
                {"key": "A2", "label": "案號", "type": "text"},
                {"key": "A3", "label": "股別", "type": "text"},
                {"key": "A4", "label": "聲請人名稱", "type": "text", "required": True},
                {"key": "B1", "label": "聲請人借款原因", "type": "textarea",
                 "default": "聲請人因OO原因而陸續向債權人商借款項，以解其燃眉之急，然因債務及利息循環增生，以致聲請人陷入難以償還之境地，故僅能向 鈞院聲請更生。"},
                {"key": "B2", "label": "調解不成立原因", "type": "textarea",
                 "default": "就調解不成立之原因，係因聲請人債權人繁多，且除銀行外，尚有民間債權人，雖銀行同意以每月新臺幣（下同）Ｏ萬塊之方式償還，惟前開民間債權人並未出席調解，如依照聲請人與前開民間債權人之約定，聲請人依然須以每月約Ｏ萬元之金額償還之，前開債務實已讓聲請人不堪負荷，且已積欠數額未繳納，而僅能向 鈞院聲請更生。"},
                {"key": "B3", "label": "更生方案", "type": "textarea",
                 "default": "就更生方案部分，聲請人每月現收入為O元；必要支出為O元，如依消費者債務清理條例第64條之1之規定，聲請人之收入扣除必要支出後之餘額為O元，現聲請人願以72期，每月1期，每期O元之方式償還債務，前開每月期數之金額已超出消費者債務清理條例第64條之1之規定，而已為盡力清償，故懇請 鈞院得依法准許聲請人之聲請。"},
                {"key": "C1", "label": "聲請前兩年收入", "type": "select", "options": REPORT_OPTIONS["C1"]},
                {"key": "C2", "label": "目前工作情形", "type": "select", "options": REPORT_OPTIONS["C2"]},
                {"key": "C3", "label": "受資助情形", "type": "select_or_text", "options": REPORT_OPTIONS["C3"]},
                {"key": "D1", "label": "現居地資料", "type": "select", "options": REPORT_OPTIONS["D1"]},
                {"key": "D2", "label": "同住之人資料", "type": "textarea",
                 "default": "聲請人現與『OO關係之人』一同居住，此有聲請人之全戶戶籍謄本【聲證O】可稽。"},
                {"key": "D3", "label": "家庭成員資料", "type": "select_or_text",
                 "options": ["刪除此項", "聲請人現家庭成員分別為『OO關係之人』，共O人。"]},
                {"key": "D4", "label": "全戶戶籍謄本", "type": "select", "options": REPORT_OPTIONS["D4"]},
                {"key": "D5", "label": "家族系統表", "type": "select", "options": REPORT_OPTIONS["D5"]},
                {"key": "D7", "label": "年金及健保資料", "type": "select", "options": REPORT_OPTIONS["D7"]},
                {"key": "D8", "label": "勞保資料", "type": "select", "options": REPORT_OPTIONS["D8"]},
                {"key": "D10", "label": "集保公司資料", "type": "select", "options": REPORT_OPTIONS["D10"]},
                {"key": "D11", "label": "壽險資料", "type": "select", "options": REPORT_OPTIONS["D11"]},
                {"key": "D12", "label": "社福津貼", "type": "select", "options": REPORT_OPTIONS["D12"]},
                {"key": "D14", "label": "公司營運情形", "type": "select", "options": REPORT_OPTIONS["D14"]},
                {"key": "D15", "label": "財產說明書等資料", "type": "select", "options": REPORT_OPTIONS["D15"]},
                {"key": "E1", "label": "審理法院", "type": "select", "options": [""] + COURT_OPTIONS},
            ],
        },
    }
    return schemas.get(form_type, {})


def get_all_form_types() -> list[dict]:
    """回傳所有可用的消債文件類型"""
    return [
        {"key": "application", "label": "聲請狀", "icon": "📋", "description": "消費者債務清理聲請狀"},
        {"key": "asset_statement", "label": "財產及收入狀況說明書", "icon": "💰", "description": "資產、收入、支出明細"},
        {"key": "creditor_list", "label": "債權人清冊", "icon": "🏦", "description": "債權人資料與金額"},
        {"key": "pdf_merge", "label": "合併PDF", "icon": "📎", "description": "合併多個PDF/DOCX為單一檔案"},
        {"key": "report", "label": "陳報狀", "icon": "📝", "description": "消費者債務清理陳報狀"},
    ]
