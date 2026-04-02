import json
import difflib
import hashlib
import hmac
import mimetypes
import os
import re
import sys
import time
import threading
import uuid
import platform
import subprocess
import tempfile
import html as ihtml
from datetime import date, datetime, time as dt_time, timedelta
from decimal import Decimal
import urllib.request
import urllib.error
from urllib.parse import urlparse
from collections import defaultdict
from pathlib import Path
import shutil
from concurrent.futures import ThreadPoolExecutor
from dotenv import load_dotenv

# Ensure MAGI root is in sys.path (needed before importing skills.*)
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))
from api.runtime_paths import ensure_path_on_sys_path, get_config_path, get_orch_dir
from api.product_runtime import PRODUCT_RUNTIME_PATH, product_profile_report, update_product_runtime
from api.case_path_mapper import (
    local_synology_path_candidates,
    preferred_case_roots,
    preferred_synology_share_roots,
    translate_local_path_to_canonical,
)
_MAGI_ROOT = os.path.abspath(os.path.dirname(os.path.dirname(__file__)))

# Auto-reap zombie children (cloudflared, skill subprocesses, etc.)
import signal as _signal
def _sigchld_handler(_signum, _frame):
    while True:
        try:
            pid, _ = os.waitpid(-1, os.WNOHANG)
            if pid == 0:
                break
        except ChildProcessError:
            break
_signal.signal(_signal.SIGCHLD, _sigchld_handler)

# Load Env — always use explicit path to guarantee .env is found regardless of cwd
load_dotenv(os.path.join(_MAGI_ROOT, ".env"))

# Validate required config before anything else
from skills.ops.config import validate_config
validate_config()

import logging
from logging.handlers import RotatingFileHandler
from flask import Flask, request, abort, render_template, redirect, url_for, flash, jsonify, Response, send_file, send_from_directory
from api.line_compat import (
    AudioMessage,
    FileMessage,
    ImageMessage,
    ImageSendMessage,
    InvalidSignatureError,
    LINE_SDK_AVAILABLE,
    LineBotApiError,
    MessageEvent,
    TextMessage,
    TextSendMessage,
    build_line_clients,
    line_feature_enabled,
)

# Initialize Logger — structured JSON for file, human-readable for console
_agent_dir_for_logs = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".agent"))
os.makedirs(_agent_dir_for_logs, exist_ok=True)
_server_log_path = os.path.join(_agent_dir_for_logs, "server.log")

from skills.ops.structured_log import JSONFormatter, HybridFormatter, RequestContextFilter

_root = logging.getLogger()
_root.setLevel(logging.INFO)
_root.addFilter(RequestContextFilter())

_file = RotatingFileHandler(_server_log_path, maxBytes=2 * 1024 * 1024, backupCount=5, encoding="utf-8")
_file.setFormatter(JSONFormatter())
_root.addHandler(_file)

_console = logging.StreamHandler()
_console.setFormatter(HybridFormatter())
_root.addHandler(_console)

logger = logging.getLogger("Server")
from api.thread_pools import channel_pool as _CHANNEL_BG_EXECUTOR, io_pool as _ATTACHMENT_BG_EXECUTOR
_CHANNEL_DELIVERY_AUDIT_FILE = os.path.join(_agent_dir_for_logs, "channel_delivery_audit.jsonl")
_channel_audit_lock = threading.Lock()

# (sys.path already set at top of file)
# Stability-first default: avoid distributed inference unless explicitly enabled.
os.environ.setdefault("MAGI_AVOID_DISTRIBUTED", "1")

from api.orchestrator import Orchestrator
try:
    from api.tw_output_guard import normalize_output_text as _normalize_output_text
except Exception:
    _normalize_output_text = None

# Auth Modules
import mysql.connector
from api.mysql_connector_guard import patch_mysql_connector_for_stability
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from werkzeug.security import generate_password_hash, check_password_hash

# DB connector stability guard:
# default to pure-python mysql-connector path to avoid C-extension segfaults under threaded load.
os.environ.setdefault("MAGI_MYSQL_USE_PURE", "1")
if patch_mysql_connector_for_stability():
    logger.info("✅ MySQL connector guard enabled (MAGI_MYSQL_USE_PURE=%s)", os.environ.get("MAGI_MYSQL_USE_PURE", "1"))

# Configuration
_SERVER_START_TIME = time.time()
app = Flask(__name__, template_folder='../templates', static_folder='../static')
app.config['TEMPLATES_AUTO_RELOAD'] = True

# P0-13 / P1-04: Session cookie hardening & security headers
app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'
# 若部署在 HTTPS 後方，啟用 Secure cookie:
if os.environ.get("MAGI_FORCE_HTTPS", "").strip().lower() in {"1", "true", "yes"}:
    app.config['SESSION_COOKIE_SECURE'] = True


@app.after_request
def _add_security_headers(response):
    """補齊安全標頭基線。"""
    response.headers.setdefault("X-Content-Type-Options", "nosniff")
    response.headers.setdefault("X-Frame-Options", "SAMEORIGIN")
    response.headers.setdefault("X-XSS-Protection", "1; mode=block")
    response.headers.setdefault("Referrer-Policy", "strict-origin-when-cross-origin")
    response.headers.setdefault(
        "Content-Security-Policy",
        "default-src 'self'; script-src 'self' 'unsafe-inline'; style-src 'self' 'unsafe-inline'; img-src 'self' data:;",
    )
    response.headers.setdefault("Strict-Transport-Security", "max-age=31536000; includeSubDomains")
    # 不對 static 資源設 no-store，僅對 API/HTML response
    if not request.path.startswith("/static/"):
        response.headers.setdefault("Cache-Control", "no-store")
    return response

# ── CSRF Protection ──
try:
    from api.csrf_guard import middleware_apply_csrf
    middleware_apply_csrf(app)
    logger.info("CSRF protection enabled")
except Exception as _csrf_err:
    logger.warning("CSRF protection not loaded: %s", _csrf_err)

# ── Blueprints ───────────────────────────────────────────────────────
from api.blueprints.osc_settings import osc_settings_bp
from api.blueprints.osc_accounting import osc_accounting_bp
from api.blueprints.osc_debt import osc_debt_bp
app.register_blueprint(osc_settings_bp)
app.register_blueprint(osc_accounting_bp)
app.register_blueprint(osc_debt_bp)
try:
    app.secret_key = os.environ["FLASK_SECRET_KEY"]
except KeyError:
    raise RuntimeError("Missing required env var: FLASK_SECRET_KEY. Set it in .env")

# Initialize Login Manager
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'

# ── Lightweight rate limiter (no external dependency) ──────────────────
_rate_limit_store: dict = {}  # key -> (count, window_start)
_RATE_LIMIT_WINDOW = 60  # seconds
_RATE_LIMIT_MAX = {
    "webhook": 120,      # LINE/Telegram webhook: 120 req/min per IP
    "api": 60,           # General API: 60 req/min per IP
}


def _check_rate_limit(category: str = "webhook") -> bool:
    """Return True if request should be rejected (rate exceeded)."""
    ip = request.remote_addr or "unknown"
    key = f"{category}:{ip}"
    now = time.time()
    limit = _RATE_LIMIT_MAX.get(category, 60)
    entry = _rate_limit_store.get(key)
    if entry:
        count, window_start = entry
        if now - window_start < _RATE_LIMIT_WINDOW:
            if count >= limit:
                return True
            _rate_limit_store[key] = (count + 1, window_start)
        else:
            _rate_limit_store[key] = (1, now)
    else:
        _rate_limit_store[key] = (1, now)
    # Prune stale entries periodically (trigger earlier, prevent unbounded growth)
    if len(_rate_limit_store) > 500:
        cutoff = now - _RATE_LIMIT_WINDOW * 2
        stale = [k for k, (_, ws) in _rate_limit_store.items() if ws < cutoff]
        for k in stale:
            _rate_limit_store.pop(k, None)
    return False


# Public hardening: never expose OpenClaw control surface via public hostname.
OPENCLAW_BLOCKED_PREFIXES = (
    "/openclaw",
    "/openclaw-gateway",
)


def _is_local_host(host: str) -> bool:
    h = (host or "").strip().split(",")[0].strip()
    if ":" in h:
        h = h.split(":", 1)[0]
    h = h.lower()
    return h in {"localhost", "127.0.0.1", "::1"}


def _is_cloudflare_tunnel_request() -> bool:
    host = (
        request.headers.get("X-Forwarded-Host")
        or request.host
        or ""
    ).lower()
    if host.endswith(".trycloudflare.com"):
        return True
    return bool(
        request.headers.get("Cf-Connecting-Ip")
        or request.headers.get("Cf-Ray")
    )


@app.before_request
def _block_public_openclaw_routes():
    path = (request.path or "").strip().lower()
    if not path:
        return None
    blocked = any(path == p or path.startswith(p + "/") for p in OPENCLAW_BLOCKED_PREFIXES)
    if not blocked:
        return None

    host = (
        request.headers.get("X-Forwarded-Host")
        or request.host
        or ""
    )
    if _is_local_host(host):
        return None

    logger.warning("Blocked public request to OpenClaw route: host=%s path=%s", host, path)
    abort(404)


@app.before_request
def _limit_cloudflare_tunnel_surface():
    if not _is_cloudflare_tunnel_request():
        return None

    path = (request.path or "").strip().lower()
    allowed_prefixes = ("/line/webhook", "/telegram/webhook", "/callback", "/health")
    allowed = any(path == p or path.startswith(p + "/") for p in allowed_prefixes)
    if allowed:
        return None

    logger.warning(
        "Blocked Cloudflare tunnel request outside LINE surface: host=%s path=%s",
        request.headers.get("X-Forwarded-Host") or request.host or "",
        path,
    )
    abort(403)

# Database Config
DB_CONFIG = {
    'host': os.environ.get("DB_HOST", "100.121.61.74"),
    'user': os.environ.get("DB_USER", "casper_service"),
    'password': os.environ.get("DB_PASSWORD", ""),
    'database': 'magi_brain'
}
if not DB_CONFIG['password']:
    logger.error("Missing required env var: DB_PASSWORD. Set it in .env")


def _load_runtime_config():
    cfg = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "config.json"))
    try:
        if os.path.exists(cfg):
            with open(cfg, "r", encoding="utf-8") as f:
                data = json.load(f)
                if isinstance(data, dict):
                    return data
    except Exception:
        logger.warning("Failed to load runtime config from %s", cfg)
    return {}


RUNTIME_CONFIG = _load_runtime_config()
SKILLS_ROOT = Path(os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "skills")))
SKILL_NAME_RE = re.compile(r"^[A-Za-z0-9][A-Za-z0-9._-]{0,63}$")
NERV_PRODUCT_NAMES = ("file_review", "transcript", "laf")

# Shared runtime dir for webhook sidecar state (captcha request/response, last sender, etc.)
AGENT_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".agent"))
os.makedirs(AGENT_DIR, exist_ok=True)


def _skill_doc_path(skill_name: str) -> Path:
    name = str(skill_name or "").strip()
    if not SKILL_NAME_RE.fullmatch(name):
        raise ValueError("invalid_skill_name")
    path = (SKILLS_ROOT / name / "SKILL.md").resolve()
    if SKILLS_ROOT.resolve() not in path.parents:
        raise ValueError("invalid_skill_path")
    return path


def _skill_action_path(skill_name: str) -> Path:
    name = str(skill_name or "").strip()
    if not SKILL_NAME_RE.fullmatch(name):
        raise ValueError("invalid_skill_name")
    path = (SKILLS_ROOT / name / "action.py").resolve()
    if SKILLS_ROOT.resolve() not in path.parents:
        raise ValueError("invalid_skill_path")
    return path


def _skill_summary(content: str) -> str:
    for line in str(content or "").splitlines():
        stripped = line.strip().lstrip("#").strip()
        if stripped:
            return stripped[:120]
    return ""


def _nerv_product_runtime_payload() -> dict:
    reports = {name: product_profile_report(name) for name in NERV_PRODUCT_NAMES}
    return {
        "ok": True,
        "runtime_path": str(PRODUCT_RUNTIME_PATH),
        "can_edit": bool(getattr(current_user, "is_admin", False)),
        "products": reports,
    }


def _list_skill_docs() -> list[dict]:
    items: list[dict] = []
    root = SKILLS_ROOT.resolve()
    if not root.exists():
        return items
    for child in sorted(root.iterdir(), key=lambda p: p.name.lower()):
        if not child.is_dir() or child.name.startswith("."):
            continue
        skill_doc = child / "SKILL.md"
        action_file = child / "action.py"
        if not skill_doc.exists() and not action_file.exists():
            continue
        content = ""
        if skill_doc.exists():
            try:
                content = skill_doc.read_text(encoding="utf-8")
            except Exception:
                content = ""
        stat_target = skill_doc if skill_doc.exists() else action_file
        updated_at = ""
        try:
            updated_at = datetime.fromtimestamp(stat_target.stat().st_mtime).isoformat()
        except Exception:
            updated_at = ""
        items.append(
            {
                "name": child.name,
                "path": str(child),
                "skill_doc_path": str(skill_doc),
                "has_skill_doc": skill_doc.exists(),
                "has_action": action_file.exists(),
                "summary": _skill_summary(content),
                "updated_at": updated_at,
            }
        )
    return items


def _nerv_skill_interview_user_id() -> str:
    return f"nerv:{getattr(current_user, 'id', '') or 'unknown'}"


def _extract_interview_skill_name(message: str) -> str:
    match = re.search(r"資料夾：`([^`]+)`", str(message or ""))
    if match:
        return str(match.group(1) or "").strip()
    match = re.search(r"資料夾：([A-Za-z0-9._-]+)", str(message or ""))
    if match:
        return str(match.group(1) or "").strip()
    return ""


def _json_auth_error(status_code: int, error: str):
    return jsonify({"ok": False, "error": error}), status_code


def _require_json_auth(admin: bool = False):
    if not getattr(current_user, "is_authenticated", False):
        return _json_auth_error(401, "auth_required")
    if admin and not current_user.is_admin():
        return _json_auth_error(403, "admin_required")
    return None

# --- User Model ---
class User(UserMixin):
    def __init__(self, id, username, role):
        self.id = id
        self.username = username
        self.role = role

    def is_admin(self):
        return self.role == 'admin'

@login_manager.user_loader
def load_user(user_id):
    try:
        from api.db_helper import get_cursor
        with get_cursor(config=DB_CONFIG, dictionary=True) as (_conn, cursor):
            cursor.execute("SELECT * FROM users WHERE id = %s", (user_id,))
            user_data = cursor.fetchone()
            if user_data:
                return User(user_data['id'], user_data['username'], user_data['role'])
        return None
    except Exception as e:
        logger.error("DB Error: %s", e)
        return None

LINE_CHANNEL_ACCESS_TOKEN = (
    os.environ.get("MAGI_LINE_CHANNEL_ACCESS_TOKEN")
    or os.environ.get("LINE_CHANNEL_ACCESS_TOKEN")
    or ""
).strip()
LINE_CHANNEL_SECRET = (
    os.environ.get("MAGI_LINE_CHANNEL_SECRET")
    or os.environ.get("LINE_CHANNEL_SECRET")
    or ""
).strip()
line_bot_api, handler, LINE_BOT_ENABLED, _line_bot_reason = build_line_clients(
    LINE_CHANNEL_ACCESS_TOKEN,
    LINE_CHANNEL_SECRET,
)
if not LINE_BOT_ENABLED:
    if not line_feature_enabled():
        logger.info("LINE webhook disabled by MAGI_ENABLE_LINE.")
    elif not LINE_SDK_AVAILABLE:
        logger.warning("LINE webhook disabled: %s", _line_bot_reason)
    else:
        logger.warning("LINE webhook disabled: %s", _line_bot_reason)

# Initialize Brain
orchestrator = Orchestrator()

# Web Notifications Buffer
WEB_NOTIFICATIONS = defaultdict(list)

# Register Iron Dome Sync Routes
try:
    from skills.ops.iron_dome_sync import register_iron_dome_routes
    register_iron_dome_routes(app)
    logger.info("🛡️ Iron Dome Sync routes registered")
except Exception as e:
    logger.warning(f"⚠️ Iron Dome Sync routes not registered: {e}")

# Initialize Auto-Skill (Self-Healing)
try:
    from skills.management.auto_skill import AutoSkill
    auto_skill = AutoSkill()
    logger.info("🧠 Auto-Skill Engine Online")
except Exception as e:
    logger.error(f"❌ Auto-Skill Init Failed: {e}")

# --- Auth Routes ---
# KEEP: Internal-only redirect. Caddy blocks /intel* externally (see Caddyfile_openclaw).
# Used by internal tools and worldmonitor skill that writes to static/worldmonitor_reports.
# Dashboard openIntelPanel() also references /intel for local access.
@app.route('/static/worldmonitor_reports')
@app.route('/static/worldmonitor_reports/')
def legacy_worldmonitor_redirect():
    return redirect("/intel")

@app.route('/dashboard')
@login_required
def dashboard():
    return render_template('dashboard.html', user=current_user)

@app.route('/dashboard/pixel')
@login_required
def dashboard_pixel():
    return render_template('dashboard_pixel.html', user=current_user)


@app.route('/dashboard/nerv')
@login_required
def dashboard_nerv():
    return render_template('dashboard_nerv.html', user=current_user)


@app.route('/dashboard/nerv/api/health')
@login_required
def nerv_api_health():
    """Real-time health check for all MAGI subsystems."""
    import requests as _rq
    results = {}

    def _check(name, fn):
        try:
            results[name] = fn()
        except Exception as e:
            results[name] = {"status": "error", "detail": str(e)[:120]}

    def _omlx():
        try:
            r = _rq.get("http://127.0.0.1:8080/v1/models", timeout=3)
            if r.status_code == 200:
                models = [m.get("id", "?") for m in (r.json().get("data") or [])]
                return {"status": "online", "models": models, "count": len(models)}
        except Exception:
            logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 528, exc_info=True)
        return {"status": "error", "detail": "unreachable"}

    def _ollama():
        # Ollama 已退役 (2026-03-11)，保留檢查供監控確認已關閉
        try:
            r = _rq.get("http://127.0.0.1:11434/api/tags", timeout=2)
            if r.status_code == 200:
                models = [m.get("name", "?") for m in (r.json().get("models") or [])]
                return {"status": "online", "models": models, "count": len(models)}
        except Exception:
            logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 539, exc_info=True)
        return {"status": "retired", "detail": "已退役，推理走 oMLX"}

    def _melchior():
        # MELCHIOR 已整併為本地邏輯模組，不再是獨立遠端節點
        return {"status": "local", "detail": "oMLX 本地推理"}

    def _balthasar():
        # BALTHASAR 已整併為本地邏輯模組
        return {"status": "local", "detail": "oMLX 本地摘要"}

    def _watcher():
        # Watcher 功能由 worldmonitor 取代
        return {"status": "retired", "detail": "由 Worldmonitor 取代"}

    def _mysql():
        try:
            import mysql.connector as _mc
            c = _mc.connect(
                host=os.environ.get("DB_HOST", "100.121.61.74"),
                port=int(os.environ.get("DB_PORT", "3306")),
                user=os.environ.get("DB_USER", "casper_service"),
                password=os.environ.get("DB_PASSWORD") or os.environ.get("MAGI_REMOTE_DB_PASSWORD", ""),
                connection_timeout=4,
                use_pure=True,
            )
            c.close()
            return {"status": "online"}
        except Exception as e:
            return {"status": "error", "detail": str(e)[:80]}

    def _cloudflared():
        try:
            if _is_cloudflared_alive():
                return {"status": "online"}
            return {"status": "offline"}
        except Exception:
            return {"status": "error", "detail": "check failed"}

    def _line_webhook():
        try:
            wh = os.environ.get("MAGI_LINE_WEBHOOK_ENDPOINT", "")
            if not wh:
                return {"status": "offline", "detail": "no endpoint configured"}
            r = _rq.get(wh.replace("/line/webhook", "/health"), timeout=5)
            return {"status": "online" if r.status_code == 200 else "error"}
        except Exception:
            return {"status": "error", "detail": "unreachable"}

    def _worldmonitor():
        try:
            import subprocess as _sp
            r = _sp.run(["pgrep", "-f", "worldmonitor"], capture_output=True, timeout=3)
            return {"status": "online" if r.returncode == 0 else "offline"}
        except Exception:
            return {"status": "error"}

    def _office_app():
        try:
            r = _rq.get("http://127.0.0.1:4200/office", timeout=4)
            return {"status": "online" if r.status_code == 200 else "error", "detail": f"HTTP {r.status_code}"}
        except Exception as e:
            return {"status": "error", "detail": str(e)[:80]}

    def _caddy_proxy():
        # Caddy proxy removed (Phase 0: OpenClaw migration). Cloudflared now connects directly to port 5002.
        return {"status": "skipped", "detail": "removed (direct cloudflared→5002)"}

    def _skills():
        docs = _list_skill_docs()
        found = [
            item["name"]
            for item in docs
            if not item["name"].startswith(("_", "."))
            and item["name"] not in {"bridge", "ops", "memory", "evolution", "brain_manager"}
        ]
        return {"status": "online", "skills": found, "count": len(found)}

    # Run checks in parallel
    from concurrent.futures import ThreadPoolExecutor as _TP
    checks = {
        "omlx": _omlx,
        "ollama": _ollama,
        "melchior": _melchior,
        "balthasar": _balthasar,
        "watcher": _watcher,
        "mysql": _mysql,
        "cloudflared": _cloudflared,
        "line_webhook": _line_webhook,
        "worldmonitor": _worldmonitor,
        "office_app": _office_app,
        "caddy_proxy": _caddy_proxy,
        "skills": _skills,
    }
    with _TP(max_workers=10) as pool:
        futs = {name: pool.submit(fn) for name, fn in checks.items()}
        for name, fut in futs.items():
            try:
                results[name] = fut.result(timeout=8)
            except Exception as e:
                results[name] = {"status": "error", "detail": str(e)[:80]}

    results["magi_server"] = {"status": "online", "pid": os.getpid()}
    results["timestamp"] = datetime.now().isoformat()
    return jsonify(results)

@app.route('/dashboard/pixel/api/status')
@login_required
def pixel_api_status():
    """將 MAGI 狀態轉譯為 Star-Office-UI 格式"""
    import json as _json
    status_file = os.path.join(os.path.dirname(app.root_path), 'static', 'magi_status.json')
    try:
        with open(status_file) as f:
            magi = _json.load(f)

        casper = magi.get('nodes', {}).get('casper', {})
        tasks = magi.get('tasks', {})

        if not casper.get('online'):
            state, detail = 'error', '系統離線'
        elif tasks:
            first_task = list(tasks.values())[0]
            state = 'executing'
            detail = first_task.get('name', '執行任務中')
        else:
            model = (casper.get('model') or '').strip()
            if model.lower() in ('idle', '', 'service down'):
                state, detail = 'idle', '待命中'
            elif 'llama' in model.lower() or 'ollama' in model.lower():
                state, detail = 'writing', f'模型運作中: {model}'
            else:
                state, detail = 'writing', model

        return jsonify({
            'state': state,
            'detail': detail,
            'progress': 0,
            'updated_at': magi.get('timestamp', '')
        })
    except Exception:
        return jsonify({'state': 'error', 'detail': '無法讀取狀態', 'progress': 0})

@app.route('/dashboard/pixel/api/agents')
@login_required
def pixel_api_agents():
    """將 MAGI 節點轉譯為 Star-Office-UI agents"""
    import json as _json
    status_file = os.path.join(os.path.dirname(app.root_path), 'static', 'magi_status.json')
    try:
        with open(status_file) as f:
            magi = _json.load(f)

        agents = []
        node_configs = {
            'melchior': {'name': 'Melchior', 'emoji': '\U0001f52c'},
            'balthasar': {'name': 'Balthasar', 'emoji': '\U0001f469'},
            'keeper': {'name': 'Keeper', 'emoji': '\U0001f5c4\ufe0f'},
            'watcher': {'name': 'Watcher', 'emoji': '\U0001f441\ufe0f'},
        }

        for node_id, cfg in node_configs.items():
            node = magi.get('nodes', {}).get(node_id)
            if not node:
                continue

            online = node.get('online', False)
            model = (node.get('model') or '').strip()

            if not online:
                state, auth = 'idle', 'offline'
            elif model in ('Service Down',):
                state, auth = 'idle', 'approved'  # 在線但服務待機
            elif model in ('Audit Only', ''):
                state, auth = 'idle', 'pending'
            elif model == 'Idle':
                state, auth = 'idle', 'approved'
            else:
                state, auth = 'writing', 'approved'

            area_map = {'idle': 'breakroom', 'writing': 'writing', 'error': 'error'}

            agents.append({
                'agentId': node_id,
                'name': cfg['name'],
                'isMain': False,
                'state': state,
                'detail': model or '---',
                'area': area_map.get(state, 'breakroom'),
                'authStatus': auth,
                'updated_at': node.get('last_check', '')
            })

        return jsonify(agents)
    except Exception:
        return jsonify([])

@app.route('/osc')
@login_required
def osc_interface():
    return render_template('osc.html', user=current_user)

@app.route('/osc/debt')
@login_required
def osc_debt_interface():
    return render_template('osc_debt.html', user=current_user)

@app.route('/exports/<path:filename>')
@login_required
def serve_exports(filename):
    """Serve generated documents from the exports directory."""
    export_dir = os.path.join(_MAGI_ROOT, "exports")
    return send_from_directory(export_dir, filename)


# KEEP: Active proxy used by dashboard.html (TOOLS_API = '/toolsapi') for audit_log
# and health endpoints. Caddy routes /toolsapi/health and /toolsapi/api/audit_log*
# through server.py; all other /toolsapi/* paths are blocked by Caddy (403).
_TOOLSAPI_COMPAT_ALLOW_PREFIX = (
    "api/audit_log",
    "health",
)


@app.route('/toolsapi/<path:subpath>', methods=['GET', 'POST', 'PUT', 'PATCH', 'DELETE', 'OPTIONS'])
def toolsapi_compat_proxy(subpath):
    """
    Backward-compatible proxy for legacy dashboard paths:
      /toolsapi/api/audit_log...
    Proxies to MAGI Tools API (default http://127.0.0.1:5003).
    """
    target = str(subpath or "").strip().lstrip("/")
    if not target:
        return jsonify({"success": False, "error": "missing target path"}), 404

    if not any(target == p or target.startswith(p + "/") for p in _TOOLSAPI_COMPAT_ALLOW_PREFIX):
        return jsonify({"success": False, "error": "toolsapi path not allowed"}), 404

    base = (os.environ.get("MAGI_TOOLS_API") or "http://127.0.0.1:5003").rstrip("/")
    qs = request.query_string.decode("utf-8", errors="ignore")
    url = f"{base}/{target}"
    if qs:
        url = f"{url}?{qs}"

    fwd_headers = {}
    ctype = (request.headers.get("Content-Type") or "").strip()
    if ctype:
        fwd_headers["Content-Type"] = ctype

    data = request.get_data() if request.method in {"POST", "PUT", "PATCH", "DELETE"} else None
    req_obj = urllib.request.Request(url, data=data if data else None, headers=fwd_headers, method=request.method)
    try:
        with urllib.request.urlopen(req_obj, timeout=20) as resp:
            body = resp.read()
            status = int(getattr(resp, "status", 200))
            resp_ct = resp.headers.get("Content-Type", "application/json; charset=utf-8")
        return Response(body, status=status, content_type=resp_ct)
    except urllib.error.HTTPError as e:
        try:
            err_body = e.read()
        except Exception:
            err_body = b""
        err_ct = getattr(e, "headers", {}).get("Content-Type", "application/json; charset=utf-8") if getattr(e, "headers", None) else "application/json; charset=utf-8"
        return Response(err_body or json.dumps({"success": False, "error": f"toolsapi_http_{getattr(e, 'code', 500)}"}).encode("utf-8"), status=int(getattr(e, "code", 500)), content_type=err_ct)
    except Exception as e:
        return jsonify({"success": False, "error": f"toolsapi_proxy_failed: {type(e).__name__}: {e}"}), 502


# ---------------------------------------------------------------------------
# Catch-all 404 → Tools API (5003) fallback proxy
# ---------------------------------------------------------------------------
# Any route not matched on server.py (5002) is automatically forwarded to
# tools_api.py (5003). This eliminates the 5002/5003 confusion — callers
# can always use port 5002 as the single entry point.
# ---------------------------------------------------------------------------
_TOOLS_API_FALLBACK_PATHS = {
    # Only proxy paths that tools_api actually serves (prefix match).
    # Static dashboard / HTML paths should NOT be proxied.
    "health", "summarize", "search", "research", "fetch", "vision",
    "melchior", "skills", "collab", "council", "remember", "recall",
    "clients", "meetings", "legal", "alert", "definitions", "laf",
    "iron-dome", "code", "connections", "sages", "osc/external",
}


@app.errorhandler(404)
def _fallback_to_tools_api(error):
    """Forward unmatched routes to Tools API (port 5003)."""
    path = request.path.lstrip("/")

    # Only proxy if the first path segment matches a known tools_api prefix.
    first_seg = path.split("/")[0] if path else ""
    first_two = "/".join(path.split("/")[:2]) if "/" in path else ""

    if first_seg not in _TOOLS_API_FALLBACK_PATHS and first_two not in _TOOLS_API_FALLBACK_PATHS:
        return jsonify({"error": "not_found", "path": request.path}), 404

    base = (os.environ.get("MAGI_TOOLS_API") or "http://127.0.0.1:5003").rstrip("/")
    qs = request.query_string.decode("utf-8", errors="ignore")
    url = f"{base}/{path}"
    if qs:
        url = f"{url}?{qs}"

    fwd_headers = {}
    ctype = (request.headers.get("Content-Type") or "").strip()
    if ctype:
        fwd_headers["Content-Type"] = ctype
    auth = (request.headers.get("Authorization") or "").strip()
    if auth:
        fwd_headers["Authorization"] = auth
    api_key = (request.headers.get("X-API-Key") or "").strip()
    if api_key:
        fwd_headers["X-API-Key"] = api_key

    data = request.get_data() if request.method in {"POST", "PUT", "PATCH", "DELETE"} else None
    req_obj = urllib.request.Request(url, data=data if data else None, headers=fwd_headers, method=request.method)
    try:
        with urllib.request.urlopen(req_obj, timeout=300) as resp:
            body = resp.read()
            status = int(getattr(resp, "status", 200))
            resp_ct = resp.headers.get("Content-Type", "application/json; charset=utf-8")
        return Response(body, status=status, content_type=resp_ct)
    except urllib.error.HTTPError as e:
        try:
            err_body = e.read()
        except Exception:
            err_body = b""
        err_ct = "application/json; charset=utf-8"
        if getattr(e, "headers", None):
            err_ct = e.headers.get("Content-Type", err_ct)
        return Response(
            err_body or json.dumps({"success": False, "error": f"tools_api_http_{getattr(e, 'code', 500)}"}).encode("utf-8"),
            status=int(getattr(e, "code", 500)),
            content_type=err_ct,
        )
    except Exception as e:
        logger.warning("tools_api fallback proxy failed: %s", e)
        return jsonify({"success": False, "error": f"tools_api_unreachable: {type(e).__name__}"}), 502


PROCESS_MONITOR_STATE_PATH = os.path.abspath(
    os.path.join(os.path.dirname(__file__), "..", "static", "process_guardian_state.json")
)
PROCESS_MONITOR_WORKER_MARKERS = [
    "skills/judgment-collector/action.py",
    "skills/file-review-orchestrator/action.py",
    "skills/transcript-downloader/action.py",
    "skills/laf-portal-automation/action.py",
    "skills/laf-orchestrator/action.py",
    "skills/laf-withdrawal-report/action.py",
    "skills/laf-refine-case/action.py",
    "skills/osc-orchestrator/action.py",
    "skills/osc-scan-folder/action.py",
    "skills/pdf-namer/action.py",
    "skills/crawler-targets/action.py",
    "skills/statutes-vdb/action.py",
    "skills/magi-autopilot/action.py",
]
try:
    from daemon import REAPER_NEVER_KILL as _DAEMON_NEVER_KILL
    PROCESS_MONITOR_CORE_MARKERS = list(_DAEMON_NEVER_KILL)
except ImportError:
    PROCESS_MONITOR_CORE_MARKERS = [
        f"{_MAGI_ROOT}/daemon.py",
        "api/server.py",
        "api/discord_bot.py",
        "skills/ops/openclaw_cron_runner.py",
        "openclaw-gateway",
        "rpc-server",
    ]
PROCESS_MONITOR_CORE_LABELS = {
    f"{_MAGI_ROOT}/daemon.py": "Daemon",
    "api/server.py": "API/LINE Webhook",
    "api/discord_bot.py": "Discord Bot",
    "skills/ops/openclaw_cron_runner.py": "OpenClaw Cron Bridge",
    "openclaw-gateway": "OpenClaw Gateway",
    "rpc-server": "RPC Worker",
}


def _parse_etime_to_sec(s: str) -> int:
    t = (s or "").strip()
    if not t:
        return 0
    m = re.match(r"^(?:(\d+)-)?(?:(\d+):)?(\d+):(\d+)$", t)
    if not m:
        return 0
    dd = int(m.group(1) or 0)
    hh = int(m.group(2) or 0)
    mm = int(m.group(3) or 0)
    ss = int(m.group(4) or 0)
    return (dd * 86400) + (hh * 3600) + (mm * 60) + ss


def _collect_process_monitor():
    rows = []
    try:
        out = subprocess.run(
            ["ps", "-axo", "pid=,ppid=,etime=,command="],
            stdout=subprocess.PIPE,
            stderr=subprocess.DEVNULL,
            text=True,
            timeout=8,
        ).stdout or ""
        for raw in out.splitlines():
            line = (raw or "").strip()
            if not line:
                continue
            parts = line.split(None, 3)
            if len(parts) < 4:
                continue
            try:
                pid = int(parts[0]); ppid = int(parts[1])
            except Exception:
                continue
            rows.append(
                {
                    "pid": pid,
                    "ppid": ppid,
                    "age_sec": _parse_etime_to_sec(parts[2]),
                    "age": parts[2],
                    "cmd": parts[3],
                }
            )
    except Exception as e:
        return {"ok": False, "error": str(e), "summary": {}, "core": [], "workers": [], "orphans": [], "duplicates": []}

    core = []
    workers = []
    orphans = []
    grouped = defaultdict(list)
    for r in rows:
        cmd = str(r.get("cmd") or "")
        label = None
        for m in PROCESS_MONITOR_CORE_MARKERS:
            if m in cmd:
                label = PROCESS_MONITOR_CORE_LABELS.get(m, m)
                break
        if label:
            x = dict(r)
            x["label"] = label
            core.append(x)
        is_worker = any(m in cmd for m in PROCESS_MONITOR_WORKER_MARKERS)
        if is_worker:
            workers.append(r)
            key = cmd
            grouped[key].append(r)
            if int(r.get("ppid") or 0) == 1:
                orphans.append(r)

    duplicates = []
    for key, items in grouped.items():
        if len(items) <= 1:
            continue
        duplicates.append(
            {
                "count": len(items),
                "pids": [int(it["pid"]) for it in items],
                "cmd": key[:320],
            }
        )

    guardian_state = {}
    try:
        if os.path.exists(PROCESS_MONITOR_STATE_PATH):
            with open(PROCESS_MONITOR_STATE_PATH, "r", encoding="utf-8") as f:
                guardian_state = json.load(f) or {}
    except Exception:
        guardian_state = {}

    return {
        "ok": True,
        "ts": datetime.now().isoformat(timespec="seconds"),
        "summary": {
            "core_count": len(core),
            "worker_count": len(workers),
            "orphan_count": len(orphans),
            "duplicate_groups": len(duplicates),
        },
        "core": sorted(core, key=lambda x: (x.get("label", ""), x.get("pid", 0))),
        "workers": sorted(workers, key=lambda x: (x.get("age_sec", 0), x.get("pid", 0)), reverse=True),
        "orphans": sorted(orphans, key=lambda x: (x.get("age_sec", 0), x.get("pid", 0)), reverse=True),
        "duplicates": sorted(duplicates, key=lambda x: x.get("count", 0), reverse=True),
        "guardian_state": guardian_state,
    }


@app.route('/ops/process-monitor')
@login_required
def process_monitor_page():
    return render_template('process_monitor.html', user=current_user)


# ============== Vector Memory API (Dashboard) ==============

@app.route('/api/memory/stats', methods=['GET'])
@login_required
def api_memory_stats():
    """Return vector memory + obsidian statistics."""
    import json as _json
    stats = {"doc_count": 0, "last_ingest": None, "obsidian": {}, "faiss_size": 0}
    try:
        idx_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), ".agent", "doc_vector_index.json")
        if os.path.exists(idx_path):
            with open(idx_path, "r") as f:
                idx = _json.load(f)
            entries = idx if isinstance(idx, list) else list(idx.values()) if isinstance(idx, dict) else []
            stats["doc_count"] = len(entries)
            dates = [e.get("updated_at") or e.get("created_at", "") for e in entries if isinstance(e, dict)]
            dates = sorted([d for d in dates if d], reverse=True)
            if dates:
                stats["last_ingest"] = dates[0]
    except Exception as e:
        stats["doc_index_error"] = str(e)
    try:
        obs_cfg = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), ".agent", "obsidian_vault_config.json")
        obs_idx = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), ".agent", "obsidian_index.json")
        if os.path.exists(obs_cfg):
            with open(obs_cfg, "r") as f:
                cfg = _json.load(f)
            stats["obsidian"]["vault_path"] = cfg.get("vault_path", "")
            stats["obsidian"]["vault_name"] = cfg.get("vault_name", "")
        if os.path.exists(obs_idx):
            with open(obs_idx, "r") as f:
                oidx = _json.load(f)
            stats["obsidian"]["notes_indexed"] = len(oidx.get("notes", {}))
            stats["obsidian"]["last_update"] = oidx.get("updated_at", "")
    except Exception as e:
        stats["obsidian_error"] = str(e)
    try:
        faiss_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "skills", "memory", "index_cache", "mem_index.faiss")
        if os.path.exists(faiss_path):
            stats["faiss_size"] = os.path.getsize(faiss_path)
    except Exception:
        logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 1004, exc_info=True)
    return jsonify(stats)


@app.route('/api/memory/recall', methods=['POST'])
@login_required
def api_memory_recall():
    """Search vector memory from dashboard."""
    data = request.get_json() or {}
    query = str(data.get("query", "")).strip()
    top_k = min(20, max(1, int(data.get("top_k", 5))))
    source_filter = str(data.get("source", "")).strip() or None
    if not query:
        return jsonify({"error": "請輸入搜尋關鍵字"}), 400
    try:
        from skills.memory.mem_bridge import recall
        results = recall(query, top_k=top_k, source_contains=source_filter)
        return jsonify({"memories": results or [], "query": query})
    except Exception as e:
        logger.error(f"Memory recall error: {e}")
        return jsonify({"error": str(e)}), 500


@app.route('/api/memory/remember', methods=['POST'])
@login_required
def api_memory_remember():
    """Store content into vector memory from dashboard."""
    data = request.get_json() or {}
    content = str(data.get("content", "")).strip()
    source = str(data.get("source", "dashboard-manual")).strip() or "dashboard-manual"
    if not content:
        return jsonify({"error": "請輸入要記憶的內容"}), 400
    if len(content) > 50000:
        return jsonify({"error": "內容過長（上限 50,000 字元）"}), 400
    try:
        from skills.memory.mem_bridge import remember
        remember(content, source)
        return jsonify({"success": True, "message": f"已儲存 {len(content)} 字元至向量記憶庫"})
    except Exception as e:
        logger.error(f"Memory remember error: {e}")
        return jsonify({"error": str(e)}), 500


@app.route('/api/memory/obsidian-sync', methods=['POST'])
@login_required
def api_memory_obsidian_sync():
    """Trigger Obsidian vault re-index."""
    import threading
    def _run_ingest():
        try:
            from skills.obsidian.action import task_ingest
            task_ingest({})
        except Exception as e:
            logger.error(f"Obsidian ingest error: {e}")
    t = threading.Thread(target=_run_ingest, daemon=True)
    t.start()
    return jsonify({"success": True, "message": "Obsidian 重新索引已啟動（背景執行中）"})


@app.route('/api/ops/process-monitor', methods=['GET'])
@login_required
def process_monitor_api():
    data = _collect_process_monitor()
    
    # Also attach the active control state so UI knows what the intended state is
    ctrl_path = os.path.join(os.path.dirname(__file__), "..", "static", "guardian_control.json")
    ctrl_enabled = True
    if os.path.exists(ctrl_path):
        try:
            with open(ctrl_path, "r", encoding="utf-8") as f:
                ctrl_enabled = json.load(f).get("enabled", True)
        except Exception:
            logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 1076, exc_info=True)
    data["guardian_control_enabled"] = ctrl_enabled
    
    code = 200 if data.get("ok") else 500
    return jsonify(data), code

@app.route('/api/ops/process-guardian/toggle', methods=['POST'])
@login_required
def process_guardian_toggle_api():
    ctrl_path = os.path.join(os.path.dirname(__file__), "..", "static", "guardian_control.json")
    try:
        ctrl = {"enabled": True}
        if os.path.exists(ctrl_path):
            with open(ctrl_path, "r", encoding="utf-8") as f:
                ctrl = json.load(f)
        
        ctrl["enabled"] = not ctrl.get("enabled", True)
        
        with open(ctrl_path, "w", encoding="utf-8") as f:
            json.dump(ctrl, f, ensure_ascii=False, indent=2)
            
        return jsonify({"ok": True, "enabled": ctrl["enabled"]})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500

@app.route('/api/osc/chat', methods=['POST'])
@login_required
def osc_chat_api():
    data = request.get_json() or {}
    msg = (data.get("message") or "").strip()
    if not msg:
        return jsonify({"error": "Empty message"}), 400
    
    # Process synchronously (or return initial ack if async)
    reply = orchestrator.process_message(
        user_id=str(current_user.id),
        message=msg,
        platform="WEB",
        role=current_user.role
    )
    try:
        if _normalize_output_text:
            reply = _normalize_output_text(str(reply or ""), platform="WEB")
    except Exception:
        logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 1120, exc_info=True)
    return jsonify({"reply": reply})

@app.route('/api/osc/poll', methods=['GET'])
@login_required
def osc_poll_api():
    uid = str(current_user.id)
    msgs = []
    if uid in WEB_NOTIFICATIONS:
        msgs = list(WEB_NOTIFICATIONS[uid])
        WEB_NOTIFICATIONS[uid].clear()
    return jsonify({"messages": msgs})

# DEPRECATED: No frontend callers found. Superseded by /api/osc/judgments (osc_judgments_compat_api)
# which merges DB insights + judgments.json. Remove after confirming no external consumers (2026-Q2).
@app.route('/api/osc/judgments_legacy', methods=['GET'])
@login_required
def osc_judgments_api():
    try:
        # Path to judgments.json in skills/judgment-collector
        json_path = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "skills", "judgment-collector", "judgments.json"))
        if os.path.exists(json_path):
            with open(json_path, "r", encoding="utf-8") as f:
                data = json.load(f)
            return jsonify(data)
        return jsonify([])
    except Exception as e:
        logger.error(f"Error serving judgments: {e}")
        return jsonify([])


# -----------------------------------------------------------------------------
# OSC Web API (law_firm_data CRUD + insights merge)
# -----------------------------------------------------------------------------
OSC_WEB_DB_CONFIG = {
    "host": os.environ.get("OSC_DB_HOST", "") or os.environ.get("DB_HOST", "127.0.0.1"),
    "port": int(os.environ.get("OSC_DB_PORT", "") or os.environ.get("DB_PORT", "3306")),
    "user": os.environ.get("OSC_DB_USER", "") or os.environ.get("DB_USER", ""),
    "password": os.environ.get("OSC_DB_PASSWORD", "") or os.environ.get("DB_PASSWORD", ""),
    "database": os.environ.get("OSC_DB_NAME", "law_firm_data"),
}


def _load_code_db_profile(profile_name: str = "Studio_VPN_Remote"):
    cfg_paths = [
        str(get_config_path("config.json")),
        os.path.join(str(get_orch_dir()), "config.json"),
    ]
    for p in cfg_paths:
        try:
            if not os.path.exists(p):
                continue
            with open(p, "r", encoding="utf-8") as f:
                obj = json.load(f) or {}
            for it in (obj.get("mariadb_profiles") or []):
                if not isinstance(it, dict):
                    continue
                if str(it.get("profile_name") or "").strip() != profile_name:
                    continue
                c = (it.get("config") or {}) if isinstance(it.get("config"), dict) else {}
                return {
                    "host": str(c.get("host") or "100.121.61.74"),
                    "port": int(c.get("port") or 3306),
                    "user": str(c.get("user") or os.environ.get("OSC_DB_USER", "python_user")),
                    "password": str(c.get("password") or os.environ.get("OSC_DB_PASSWORD", "")),
                    "database": str(c.get("database") or "law_firm_data"),
                }
        except Exception:
            continue
    return None


def _resolve_osc_web_db_config():
    # Explicit OSC_WEB_DB_* always wins.
    explicit = any((os.environ.get(k, "") or "").strip() for k in ["OSC_WEB_DB_HOST", "OSC_WEB_DB_PORT", "OSC_WEB_DB_USER", "OSC_WEB_DB_PASSWORD", "OSC_WEB_DB_NAME"])
    if explicit:
        return {
            "host": os.environ.get("OSC_WEB_DB_HOST") or "100.121.61.74",
            "port": int((os.environ.get("OSC_WEB_DB_PORT") or "3306").strip()),
            "user": os.environ.get("OSC_WEB_DB_USER") or "python_user",
            "password": os.environ.get("OSC_WEB_DB_PASSWORD") or "",
            "database": os.environ.get("OSC_WEB_DB_NAME") or "law_firm_data",
        }

    # Default behavior: prefer CRUD-capable profile from code config.
    prefer_crud_user = (os.environ.get("OSC_WEB_PREFER_CRUD_USER", "1") or "1").strip().lower() in {"1", "true", "yes", "on"}
    if prefer_crud_user:
        prof = _load_code_db_profile("Studio_VPN_Remote")
        if prof:
            return prof

    # Fallback to existing MAGI env chain.
    return {
        "host": os.environ.get("OSC_DB_HOST") or os.environ.get("MAGI_REMOTE_DB_HOST") or "100.121.61.74",
        "port": int((os.environ.get("OSC_DB_PORT") or os.environ.get("MAGI_REMOTE_DB_PORT") or "3306").strip()),
        "user": os.environ.get("OSC_DB_USER") or os.environ.get("MAGI_REMOTE_DB_USER") or "python_user",
        "password": os.environ.get("OSC_DB_PASSWORD") or os.environ.get("MAGI_REMOTE_DB_PASSWORD") or "",
        "database": os.environ.get("OSC_DB_NAME") or os.environ.get("MAGI_REMOTE_DB_NAME") or "law_firm_data",
    }


OSC_WEB_DB_CONFIG = _resolve_osc_web_db_config()


def _osc_web_db_candidates():
    primary = dict(OSC_WEB_DB_CONFIG)
    cands = [primary]
    local_host = (os.environ.get("MAGI_LOCAL_DB_HOST") or "127.0.0.1").strip()
    local_port = int((os.environ.get("MAGI_LOCAL_DB_PORT") or "3307").strip())
    local_user = (os.environ.get("MAGI_LOCAL_DB_USER") or primary["user"]).strip()
    local_pass = os.environ.get("MAGI_LOCAL_DB_PASSWORD") or primary["password"]
    local_name = (os.environ.get("MAGI_LOCAL_DB_NAME") or primary["database"]).strip()
    if (local_host, local_port, local_name, local_user) != (primary["host"], primary["port"], primary["database"], primary["user"]):
        cands.append(
            {
                "host": local_host,
                "port": local_port,
                "user": local_user,
                "password": local_pass,
                "database": local_name,
            }
        )
    return cands


def _osc_web_connect():
    last_err = None
    for cfg in _osc_web_db_candidates():
        try:
            conn = mysql.connector.connect(
                host=cfg["host"],
                port=int(cfg["port"]),
                user=cfg["user"],
                password=cfg["password"],
                database=cfg["database"],
                autocommit=False,
                charset="utf8mb4",
                collation="utf8mb4_unicode_ci",
                connection_timeout=5,
            )
            return conn, cfg
        except Exception as e:
            last_err = e
            continue
    raise last_err or RuntimeError("osc_web_db_connect_failed")


def _osc_json_value(v):
    if isinstance(v, datetime):
        return v.isoformat(sep=" ")
    if isinstance(v, date):
        return v.isoformat()
    if isinstance(v, dt_time):
        return v.isoformat()
    if isinstance(v, timedelta):
        total = int(v.total_seconds())
        sign = "-" if total < 0 else ""
        total = abs(total)
        h = total // 3600
        m = (total % 3600) // 60
        s = total % 60
        return f"{sign}{h:02d}:{m:02d}:{s:02d}"
    if isinstance(v, Decimal):
        return float(v)
    return v


def _osc_row_json(row):
    if not isinstance(row, dict):
        return row
    return {k: _osc_json_value(v) for k, v in row.items()}


def _osc_parse_dt(v):
    if isinstance(v, datetime):
        return v
    if isinstance(v, date):
        return datetime(v.year, v.month, v.day)
    s = str(v or "").strip()
    if not s:
        return None
    s = s.replace("T", " ")
    fmts = [
        "%Y-%m-%d %H:%M:%S",
        "%Y-%m-%d %H:%M",
        "%Y-%m-%d",
        "%Y/%m/%d %H:%M:%S",
        "%Y/%m/%d %H:%M",
        "%Y/%m/%d",
    ]
    for f in fmts:
        try:
            return datetime.strptime(s, f)
        except Exception:
            continue
    return None


def _osc_norm_path(path_str: str) -> str:
    s = str(path_str or "").strip()
    if not s:
        return s
    if s.startswith("/Users/") or s.startswith("/Volumes/") or s.startswith("smb://"):
        return s
    s2 = s.replace("/", "\\")
    up = s2.upper()
    if up.startswith("K:\\SYNOLOGYDRIVE"):
        return s2.replace("K:\\SynologyDrive", "Z:\\lumi63181107").replace("K:\\SYNOLOGYDRIVE", "Z:\\lumi63181107")
    if up.startswith("K:\\LUMI"):
        return "Z:" + s2[2:]
    if up.startswith("K:"):
        return "Z:" + s2[2:]
    return s2


def _osc_local_path_candidates(path_str: str) -> list[str]:
    """
    Convert legacy Windows/NAS paths into local synced Synology paths on macOS.
    """
    return local_synology_path_candidates(_osc_norm_path(path_str))


def _osc_allowed_local_roots() -> list[str]:
    roots = preferred_synology_share_roots(include_closed=True) + ["/Volumes"]
    out = []
    for root in roots:
        rp = os.path.realpath(root)
        if rp not in out:
            out.append(rp)
    return out


def _osc_is_safe_local_path(path_str: str, *, allow_missing: bool = False) -> bool:
    p = str(path_str or "").strip()
    if not p:
        return False
    try:
        real = os.path.realpath(p)
    except Exception:
        return False
    if not allow_missing and not os.path.exists(real):
        return False
    for root in _osc_allowed_local_roots():
        if real == root or real.startswith(root + os.sep):
            return True
    return False


def _osc_resolve_existing_local_path(path_str: str, *, prefer_dir: bool | None = None) -> str:
    candidates = _osc_local_path_candidates(path_str)
    norm = _osc_norm_path(path_str).replace("\\", "/")
    if norm and norm not in candidates:
        candidates.append(norm)
    for cand in candidates:
        try:
            real = os.path.realpath(cand)
            if not _osc_is_safe_local_path(real):
                continue
            if not os.path.exists(real):
                continue
            if prefer_dir is True and not os.path.isdir(real):
                continue
            if prefer_dir is False and not os.path.isfile(real):
                continue
            return real
        except Exception:
            continue
    return ""


def _osc_relpath_under(base_path: str, target_path: str) -> str:
    try:
        rel = os.path.relpath(os.path.realpath(target_path), os.path.realpath(base_path))
    except Exception:
        return ""
    if rel in {".", ""}:
        return ""
    return rel.replace("\\", "/")


def _osc_human_size(size: int) -> str:
    units = ["B", "KB", "MB", "GB", "TB"]
    n = float(max(0, int(size or 0)))
    unit = units[0]
    for unit in units:
        if n < 1024 or unit == units[-1]:
            break
        n /= 1024.0
    return f"{n:.1f}{unit}" if unit != "B" else f"{int(n)}B"


def _osc_folder_entries(base_path: str, relative_path: str = "", limit: int = 240) -> dict:
    base_real = os.path.realpath(base_path)
    target_real = os.path.realpath(os.path.join(base_real, relative_path or ""))
    if not _osc_is_safe_local_path(base_real):
        return {"ok": False, "error": "base_not_allowed"}
    if target_real != base_real and not target_real.startswith(base_real + os.sep):
        return {"ok": False, "error": "path_escape"}
    if not os.path.isdir(target_real):
        return {"ok": False, "error": "folder_not_found"}
    entries = []
    try:
        names = sorted(os.listdir(target_real), key=lambda n: (not os.path.isdir(os.path.join(target_real, n)), n.lower()))
    except Exception as e:
        return {"ok": False, "error": str(e)}
    for name in names[:max(1, int(limit or 240))]:
        full = os.path.join(target_real, name)
        try:
            is_dir = os.path.isdir(full)
            stat = os.stat(full)
        except Exception:
            continue
        rel = _osc_relpath_under(base_real, full)
        entries.append(
            {
                "name": name,
                "relative_path": rel,
                "type": "dir" if is_dir else "file",
                "ext": "" if is_dir else os.path.splitext(name)[1].lower(),
                "size": None if is_dir else int(stat.st_size),
                "size_label": "" if is_dir else _osc_human_size(int(stat.st_size)),
                "modified_at": datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d %H:%M:%S"),
            }
        )
    parent_relative = _osc_relpath_under(base_real, os.path.dirname(target_real)) if target_real != base_real else ""
    return {
        "ok": True,
        "base_path": base_real,
        "current_path": target_real,
        "current_relative_path": _osc_relpath_under(base_real, target_real),
        "parent_relative_path": parent_relative,
        "entries": entries,
    }


_OSC_TEXT_EXTENSIONS = {
    ".txt", ".md", ".json", ".csv", ".tsv", ".yaml", ".yml",
    ".xml", ".html", ".htm", ".log", ".py", ".js", ".ts", ".css",
}


def _osc_is_editable_text_path(path_str: str) -> bool:
    ext = os.path.splitext(str(path_str or "").strip())[1].lower()
    return ext in _OSC_TEXT_EXTENSIONS


def _osc_read_text_file(path_str: str, max_bytes: int = 2 * 1024 * 1024) -> tuple[str, str]:
    size = os.path.getsize(path_str)
    if size > max_bytes:
        raise ValueError(f"file_too_large:{size}")
    raw = Path(path_str).read_bytes()
    for enc in ("utf-8-sig", "utf-8", "cp950"):
        try:
            return raw.decode(enc), enc
        except Exception:
            continue
    return raw.decode("utf-8", errors="replace"), "utf-8-replace"


def _osc_smb_candidates(path_str: str) -> list[str]:
    """
    Return ordered SMB candidate URLs for NAS.
    """
    host = (os.environ.get("MAGI_NAS_HOST") or "192.168.1.3").strip()
    p = _osc_norm_path(path_str).replace("\\", "/")
    if p.startswith("/Users/") or p.startswith("/Volumes/"):
        p = translate_local_path_to_canonical(p).replace("\\", "/")
    out: list[str] = []
    rel = ""
    if p.startswith("Z:/lumi63181107"):
        rel = p[len("Z:/lumi63181107"):].lstrip("/")
        for base in [f"smb://{host}/SynologyDrive", f"smb://{host}/home", f"smb://{host}/homes/lumi63181107"]:
            out.append(f"{base}/{rel}" if rel else base)
    elif p.startswith("Y:/"):
        rel = p[len("Y:/"):].lstrip("/")
        if rel.startswith("lumi/"):
            rel = rel[len("lumi/"):]
        for base in [f"smb://{host}/lumi/lumi", f"smb://{host}/lumi", f"smb://{host}/home"]:
            out.append(f"{base}/{rel}" if rel else base)
    elif p.lower().startswith("smb://"):
        out.append(p)
    elif re.match(r"^[A-Za-z]:/", p):
        rel = p[3:].lstrip("/")
        out.append(f"smb://{host}/{rel}" if rel else f"smb://{host}")
    else:
        out.append(f"smb://{host}")

    # always keep NAS root as last resort
    out.append(f"smb://{host}")
    seen = set()
    uniq = []
    for x in out:
        if x not in seen:
            seen.add(x)
            uniq.append(x)
    return uniq


def _osc_path_to_smb(path_str: str) -> str:
    """
    Convert canonical Windows path to smb:// URL for NAS browsing.
    """
    cands = _osc_smb_candidates(path_str)
    return cands[0] if cands else str(path_str or "")


def _osc_try_open_path(path_str: str) -> dict:
    """
    Best-effort open folder in host OS. Returns execution result only.
    """
    p = str(path_str or "").strip()
    if not p:
        return {"ok": False, "error": "empty_path"}
    try:
        if platform.system() == "Darwin":
            subprocess.run(["open", p], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, timeout=10)
            return {"ok": True, "method": "open"}
        if platform.system().lower().startswith("win"):
            os.startfile(p)  # type: ignore[attr-defined]
            return {"ok": True, "method": "startfile"}
        subprocess.run(["xdg-open", p], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, timeout=10)
        return {"ok": True, "method": "xdg-open"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def _osc_case_folder_from_doc_path(case_number: str, file_path: str) -> str:
    """
    Best-effort derive case folder from a document path.
    """
    p = str(file_path or "").strip().replace("\\", "/")
    if not p:
        return ""
    cn = str(case_number or "").strip()
    parts = [x for x in p.split("/") if x]
    if cn and cn in parts:
        idx = parts.index(cn)
        return "/" + "/".join(parts[: idx + 1])
    if cn:
        for i, seg in enumerate(parts):
            if cn in seg:
                return "/" + "/".join(parts[: i + 1])
    return str(Path(p).parent)


def _osc_guess_case_folder(case_number: str) -> str:
    """
    Fallback when cases.folder_path is empty:
    infer from document_index/case_documents latest file path.
    """
    cn = (case_number or "").strip()
    if not cn:
        return ""
    try:
        row, _ = _osc_exec(
            """
            SELECT file_path
            FROM document_index
            WHERE case_number=%s AND file_path IS NOT NULL AND file_path != ''
            ORDER BY modified_date DESC, id DESC
            LIMIT 1
            """,
            (cn,),
            fetch="one",
        )
        if row and row.get("file_path"):
            return _osc_case_folder_from_doc_path(cn, row.get("file_path"))
    except Exception:
        logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 1588, exc_info=True)
    try:
        row, _ = _osc_exec(
            """
            SELECT file_path
            FROM case_documents
            WHERE file_path IS NOT NULL AND file_path != '' AND (description LIKE %s OR file_name LIKE %s)
            ORDER BY upload_date DESC, id DESC
            LIMIT 1
            """,
            (f"%{cn}%", f"%{cn}%"),
            fetch="one",
        )
        if row and row.get("file_path"):
            return _osc_case_folder_from_doc_path(cn, row.get("file_path"))
    except Exception:
        logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 1604, exc_info=True)
    return ""


def _osc_strip_html_to_text(html: str) -> str:
    s = html or ""
    s = re.sub(r"(?is)<script[^>]*>.*?</script>", " ", s)
    s = re.sub(r"(?is)<style[^>]*>.*?</style>", " ", s)
    s = re.sub(r"(?is)<noscript[^>]*>.*?</noscript>", " ", s)
    s = re.sub(r"(?is)<br\\s*/?>", "\n", s)
    s = re.sub(r"(?is)</p>", "\n", s)
    s = re.sub(r"(?is)<[^>]+>", " ", s)
    s = ihtml.unescape(s)
    s = re.sub(r"[ \t\r\f\v]+", " ", s)
    s = re.sub(r"\n\\s*\n+", "\n\n", s)
    return s.strip()


def _osc_fetch_url_text(url: str, timeout: int = 20) -> dict:
    import ssl as _ssl
    u = str(url or "").strip()
    if not u.lower().startswith(("http://", "https://")):
        return {"ok": False, "error": "invalid_url"}
    headers = {
        "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 14_0) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122 Safari/537.36",
        "Accept-Language": "zh-TW,zh;q=0.9,en;q=0.8",
    }
    req = urllib.request.Request(u, headers=headers)
    to = max(5, int(timeout))
    # Try normal SSL first, then fallback to unverified context
    for ctx in (None, _ssl.create_default_context()):
        try:
            if ctx is not None:
                ctx.check_hostname = False
                ctx.verify_mode = _ssl.CERT_NONE
            with urllib.request.urlopen(req, timeout=to, context=ctx) as resp:
                raw = resp.read()
                ct = (resp.headers.get("Content-Type") or "").lower()
            if "charset=" in ct:
                enc = ct.split("charset=", 1)[1].split(";", 1)[0].strip()
            else:
                enc = "utf-8"
            try:
                html_str = raw.decode(enc, errors="ignore")
            except Exception:
                html_str = raw.decode("utf-8", errors="ignore")
            text = _osc_strip_html_to_text(html_str)
            if len(text) < 80:
                return {"ok": False, "error": "content_too_short", "text": text}
            return {"ok": True, "text": text, "html_len": len(html_str)}
        except _ssl.SSLError:
            if ctx is not None:
                return {"ok": False, "error": "ssl_failed_even_unverified"}
            continue  # retry with unverified context
        except urllib.error.HTTPError as e:
            return {"ok": False, "error": f"http_{e.code}"}
        except Exception as e:
            if "ssl" in str(e).lower() or "certificate" in str(e).lower():
                if ctx is not None:
                    return {"ok": False, "error": str(e)}
                continue  # retry with unverified context
            return {"ok": False, "error": str(e)}
    return {"ok": False, "error": "fetch_exhausted"}


def _osc_lookup_fulltext_fallback(title: str = "", case_number: str = "", url: str = "") -> dict:
    """
    Fallback source for sites that block direct scraping (e.g. login-gated pages).
    Try retrieving full text from local DB mirrors first.
    """
    t = (title or "").strip()
    cn = (case_number or "").strip()
    u = (url or "").strip()
    try:
        params = []
        clauses = []
        if cn:
            clauses.append("case_number LIKE %s")
            params.append(f"%{cn}%")
        if t:
            clauses.append("(document_name LIKE %s OR case_reason LIKE %s)")
            params.extend([f"%{t[:80]}%", f"%{t[:80]}%"])
        if u:
            clauses.append("source_file LIKE %s")
            params.append(f"%{u}%")
        if clauses:
            row, _ = _osc_exec(
                f"""
                SELECT id, document_name, case_number, case_reason, source_file, insight_text, raw_text
                FROM legal_insights
                WHERE {' OR '.join(clauses)}
                ORDER BY CHAR_LENGTH(COALESCE(insight_text, raw_text, '')) DESC, extracted_date DESC, id DESC
                LIMIT 1
                """,
                tuple(params),
                fetch="one",
            )
            if row:
                txt = (row.get("insight_text") or row.get("raw_text") or "").strip()
                if len(txt) >= 300:
                    return {
                        "ok": True,
                        "text": txt,
                        "source": "fallback_legal_insights",
                        "matched": {
                            "id": row.get("id"),
                            "title": row.get("document_name") or "",
                            "case_number": row.get("case_number") or "",
                        },
                    }
    except Exception:
        logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 1715, exc_info=True)
    try:
        params = []
        clauses = []
        if cn:
            clauses.append("case_number LIKE %s")
            params.append(f"%{cn}%")
        if t:
            clauses.append("(case_number LIKE %s OR summary LIKE %s)")
            params.extend([f"%{t[:40]}%", f"%{t[:120]}%"])
        if u:
            clauses.append("source_url LIKE %s")
            params.append(f"%{u}%")
        if clauses:
            row, _ = _osc_exec(
                f"""
                SELECT id, court_name, case_number, summary, full_text, source_url
                FROM court_judgments
                WHERE {' OR '.join(clauses)}
                ORDER BY CHAR_LENGTH(COALESCE(full_text, summary, '')) DESC, crawled_at DESC, id DESC
                LIMIT 1
                """,
                tuple(params),
                fetch="one",
            )
            if row:
                txt = (row.get("full_text") or row.get("summary") or "").strip()
                if len(txt) >= 300:
                    return {
                        "ok": True,
                        "text": txt,
                        "source": "fallback_court_judgments",
                        "matched": {
                            "id": row.get("id"),
                            "title": f"{row.get('court_name') or ''} {row.get('case_number') or ''}".strip(),
                            "case_number": row.get("case_number") or "",
                        },
                    }
    except Exception:
        logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 1754, exc_info=True)
    return {"ok": False, "error": "fallback_not_found"}


def _osc_run_skill(skill: str, task: str, timeout_sec: int = 180, route_key: str = "") -> dict:
    tools_api = (os.environ.get("MAGI_TOOLS_API") or "http://127.0.0.1:5003").rstrip("/")
    payload = {
        "skill": skill,
        "task": task,
        "timeout_sec": int(timeout_sec),
        "auto_repair": False,
        "rollback_on_fail": True,
        "auto_install_deps": False,
        "route_key": route_key or f"osc:{skill}",
    }
    data = json.dumps(payload, ensure_ascii=False).encode("utf-8")
    req = urllib.request.Request(
        tools_api + "/skills/run",
        data=data,
        headers={"Content-Type": "application/json"},
    )
    try:
        with urllib.request.urlopen(req, timeout=max(20, int(timeout_sec) + 30)) as resp:
            raw = resp.read().decode("utf-8", errors="replace")
            return json.loads(raw or "{}")
    except urllib.error.HTTPError as e:
        try:
            body = e.read().decode("utf-8", errors="replace")
        except Exception:
            body = ""
        return {"success": False, "error": f"http_{getattr(e, 'code', 0)}", "body": body[:800]}
    except Exception as e:
        return {"success": False, "error": str(e)[:240]}


def _osc_skill_json_task(command: str, payload: dict) -> str:
    return f"{str(command or '').strip()}{json.dumps(payload or {}, ensure_ascii=False, separators=(',', ':'))}"


def _osc_parse_skill_output(run_result: dict) -> dict:
    if (not isinstance(run_result, dict)) or (not run_result.get("success")):
        return {"success": False, "error": (run_result.get("error") if isinstance(run_result, dict) else "run_failed")}
    raw = str(run_result.get("output") or "").strip()
    if not raw:
        return {"success": False, "error": "empty_output"}
    try:
        obj = json.loads(raw)
    except Exception as e:
        return {"success": False, "error": f"output_json_parse_failed: {e}"}
    return obj if isinstance(obj, dict) else {"success": False, "error": "output_not_dict"}


def _osc_title_norm(s: str) -> str:
    x = str(s or "")
    x = x.replace("（", "(").replace("）", ")")
    x = re.sub(r"\s+", "", x)
    x = re.sub(r"[，。、「」『』【】《》:：;；,./\\|_+\-\(\)\[\]{}·]", "", x)
    return x.lower()


_OSC_JUDICIAL_COURT_SEARCH_LABELS = {
    "最高法院": "最高法院",
    "最高行政法院": "最高行政法院(含改制前行政法院)",
    "臺北高等行政法院": "臺北高等行政法院",
    "臺中高等行政法院": "臺中高等行政法院",
    "高雄高等行政法院": "高雄高等行政法院",
    "智慧財產及商業法院": "智慧財產及商業法院",
    "懲戒法院": "懲戒法院",
    "臺灣高等法院": "臺灣高等法院",
    "臺灣高等法院臺中分院": "臺灣高等法院臺中分院",
    "臺灣高等法院臺南分院": "臺灣高等法院臺南分院",
    "臺灣高等法院高雄分院": "臺灣高等法院高雄分院",
    "臺灣高等法院花蓮分院": "臺灣高等法院花蓮分院",
    "臺灣臺北地方法院": "臺灣臺北地方法院",
    "臺灣士林地方法院": "臺灣士林地方法院",
    "臺灣新北地方法院": "臺灣新北地方法院",
    "臺灣桃園地方法院": "臺灣桃園地方法院",
    "臺灣新竹地方法院": "臺灣新竹地方法院",
    "臺灣苗栗地方法院": "臺灣苗栗地方法院",
    "臺灣臺中地方法院": "臺灣臺中地方法院",
    "臺灣南投地方法院": "臺灣南投地方法院",
    "臺灣彰化地方法院": "臺灣彰化地方法院",
    "臺灣雲林地方法院": "臺灣雲林地方法院",
    "臺灣嘉義地方法院": "臺灣嘉義地方法院",
    "臺灣臺南地方法院": "臺灣臺南地方法院",
    "臺灣高雄地方法院": "臺灣高雄地方法院",
    "臺灣橋頭地方法院": "臺灣橋頭地方法院",
    "臺灣屏東地方法院": "臺灣屏東地方法院",
    "臺灣臺東地方法院": "臺灣臺東地方法院",
    "臺灣花蓮地方法院": "臺灣花蓮地方法院",
    "臺灣宜蘭地方法院": "臺灣宜蘭地方法院",
    "臺灣基隆地方法院": "臺灣基隆地方法院",
    "臺灣澎湖地方法院": "臺灣澎湖地方法院",
    "福建金門地方法院": "福建金門地方法院",
    "福建連江地方法院": "福建連江地方法院",
}

_OSC_JUDICIAL_COURT_ALIASES = {
    "台灣高等法院": "臺灣高等法院",
    "高等法院": "臺灣高等法院",
    "台灣高等法院台中分院": "臺灣高等法院臺中分院",
    "台灣高等法院台南分院": "臺灣高等法院臺南分院",
    "台灣高等法院高雄分院": "臺灣高等法院高雄分院",
    "台灣高等法院花蓮分院": "臺灣高等法院花蓮分院",
    "臺灣高等法院台中分院": "臺灣高等法院臺中分院",
    "臺灣高等法院台南分院": "臺灣高等法院臺南分院",
    "高等法院台中分院": "臺灣高等法院臺中分院",
    "高等法院台南分院": "臺灣高等法院臺南分院",
    "高等法院高雄分院": "臺灣高等法院高雄分院",
    "高等法院花蓮分院": "臺灣高等法院花蓮分院",
    "台北高等行政法院": "臺北高等行政法院",
    "台中高等行政法院": "臺中高等行政法院",
    "台灣台北地方法院": "臺灣臺北地方法院",
    "台灣台中地方法院": "臺灣臺中地方法院",
    "台灣台南地方法院": "臺灣臺南地方法院",
    "台灣台東地方法院": "臺灣臺東地方法院",
}


def _osc_unique_keep_order(values: list[str]) -> list[str]:
    seen = set()
    out = []
    for v in values:
        s = str(v or "").strip()
        if (not s) or (s in seen):
            continue
        seen.add(s)
        out.append(s)
    return out


def _osc_normalize_court_name(s: str) -> str:
    text = str(s or "").strip()
    if not text:
        return ""
    text = text.replace("台", "臺")
    return _OSC_JUDICIAL_COURT_ALIASES.get(text, text)


def _osc_extract_court_names(s: str) -> list[str]:
    text = str(s or "").strip()
    if not text:
        return []
    text2 = text.replace("台", "臺")
    hits: list[str] = []

    # Match more specific court names first to avoid a generic "臺灣高等法院"
    # hit swallowing a branch court name.
    for court_name in sorted(_OSC_JUDICIAL_COURT_SEARCH_LABELS.keys(), key=len, reverse=True):
        if court_name and (court_name in text2):
            hits.append(court_name)

    if not hits:
        m = re.search(r"(臺灣高等法院[一-龥]{1,4}分院)", text2)
        if m:
            hits.append(m.group(1))
        m = re.search(r"(臺灣[一-龥]{1,4}地方法院)", text2)
        if m:
            hits.append(m.group(1))
        m = re.search(r"(福建[一-龥]{1,4}地方法院)", text2)
        if m:
            hits.append(m.group(1))

    return _osc_unique_keep_order([_osc_normalize_court_name(x) for x in hits])


def _osc_normalize_case_word(s: str) -> str:
    word = str(s or "").strip()
    if not word:
        return ""
    word = word.replace("臺", "台")
    word = re.sub(r"\s+", "", word)
    word = re.sub(r"[字第號()（）【】\[\]{}「」『』《》,，.。:：;；/\\|_+\-]", "", word)
    return word


def _osc_parse_structured_case_spec(*, title: str = "", case_number: str = "") -> dict:
    texts = [
        ("case_number", str(case_number or "").strip()),
        ("title", str(title or "").strip()),
    ]
    out = {
        "case_year": "",
        "case_word": "",
        "case_no": "",
        "case_number_query": "",
        "case_marker": "",
        "courts": _osc_extract_court_names(title),
        "source": "",
    }
    patterns = [
        r"(?P<year>\d{2,3})\s*年度\s*(?P<word>[A-Za-z\u4e00-\u9fff]{1,16}?)\s*字?\s*第?\s*(?P<no>0*\d{1,8})\s*號?",
        r"(?P<year>\d{2,3})\s*(?P<word>[A-Za-z\u4e00-\u9fff]{1,16}?)\s*(?:字)?\s*第?\s*(?P<no>0*\d{1,8})\s*號?",
    ]
    for source, raw in texts:
        if not raw:
            continue
        text = raw.replace("（", "(").replace("）", ")")
        text = re.sub(r"\s+", "", text)
        for pat in patterns:
            m = re.search(pat, text)
            if not m:
                continue
            year = str(m.group("year") or "").strip()
            word = _osc_normalize_case_word(m.group("word") or "")
            no_raw = str(m.group("no") or "").strip()
            if (not year) or (not word) or (not no_raw):
                continue
            try:
                no = str(int(no_raw))
            except Exception:
                no = no_raw.lstrip("0") or "0"
            out.update(
                {
                    "case_year": year,
                    "case_word": word,
                    "case_no": no,
                    "case_number_query": f"{year}年度{word}字第{no}號",
                    "case_marker": f"{year}{word}{no}",
                    "source": source,
                }
            )
            return out
    return out


def _osc_extract_case_markers(s: str) -> set[str]:
    text = str(s or "")
    if not text:
        return set()
    pats = [
        r"\d{1,3}\s*[台臺]\s*[\u4e00-\u9fff]{1,10}\s*字?\s*第?\s*\d{1,8}\s*號?",
        r"\d{2,3}\s*年度\s*[\u4e00-\u9fff]{1,10}\s*字\s*第?\s*\d{1,8}\s*號?",
        r"\d{2,3}\s*年度\s*[\u4e00-\u9fff]{1,10}\s*第?\s*\d{1,8}\s*號?",
    ]
    out: set[str] = set()
    for p in pats:
        for m in re.findall(p, text):
            t = re.sub(r"\s+", "", m or "").replace("第", "").replace("號", "").replace("臺", "台")
            if t:
                out.add(t)
    return out


def _osc_load_judicial_search_results(result_obj: dict) -> list[dict]:
    items = []
    results_path = str(result_obj.get("results_path") or "").strip()
    if results_path and os.path.exists(results_path):
        try:
            with open(results_path, "r", encoding="utf-8") as f:
                cached = json.load(f) or {}
            items = cached.get("results") or []
        except Exception:
            items = []
    if not items:
        items = result_obj.get("results") or []
    return [it for it in items if isinstance(it, dict)]


def _osc_pick_exact_judicial_search_result(items: list[dict], *, title: str = "", case_number: str = "") -> dict:
    if not isinstance(items, list) or not items:
        return {}
    target = _osc_parse_structured_case_spec(title=title, case_number=case_number)
    if not (target.get("case_year") and target.get("case_word") and target.get("case_no")):
        return {}
    target_courts = {_osc_normalize_court_name(x) for x in (target.get("courts") or []) if str(x or "").strip()}
    target_title_norm = _osc_title_norm(title)
    target_word = _osc_normalize_case_word(target.get("case_word") or "")
    target_no = str(target.get("case_no") or "").lstrip("0") or "0"
    best = None
    best_score = -1.0
    for it in items:
        title2 = str(it.get("title") or "").strip()
        url2 = str(it.get("url") or "").strip()
        if not title2 or not url2:
            continue
        spec2 = _osc_parse_structured_case_spec(title=title2)
        if not (spec2.get("case_year") and spec2.get("case_word") and spec2.get("case_no")):
            continue
        if str(spec2.get("case_year") or "") != str(target.get("case_year") or ""):
            continue
        if _osc_normalize_case_word(spec2.get("case_word") or "") != target_word:
            continue
        no2 = str(spec2.get("case_no") or "").lstrip("0") or "0"
        if no2 != target_no:
            continue
        item_courts = {_osc_normalize_court_name(x) for x in _osc_extract_court_names(title2)}
        if target_courts and item_courts and target_courts.isdisjoint(item_courts):
            continue
        score = 10.0
        if target_courts and item_courts and (not target_courts.isdisjoint(item_courts)):
            score += 3.0
        if target_title_norm:
            title_norm2 = _osc_title_norm(title2)
            if title_norm2:
                ratio = difflib.SequenceMatcher(None, target_title_norm, title_norm2).ratio()
                score += ratio * 2.0
                if title_norm2 == target_title_norm:
                    score += 1.5
        if score > best_score:
            best_score = score
            best = it
    if not best:
        return {}
    out = dict(best)
    out["match_score"] = round(best_score, 4)
    return out


def _osc_fetch_fulltext_from_exact_case_search(*, title: str = "", case_number: str = "", timeout_sec: int = 180) -> dict:
    target = _osc_parse_structured_case_spec(title=title, case_number=case_number)
    if not (target.get("case_year") and target.get("case_word") and target.get("case_no")):
        return {"ok": False, "error": "structured_case_unavailable"}

    search_payload = {
        "keywords": "",
        "max_results": 20,
        "headless": True,
        "timeout_sec": max(60, min(240, int(timeout_sec))),
        "case_year": target.get("case_year") or "",
        "case_word": target.get("case_word") or "",
        "case_no": target.get("case_no") or "",
    }
    search_courts = [
        _OSC_JUDICIAL_COURT_SEARCH_LABELS.get(court_name, court_name)
        for court_name in (target.get("courts") or [])
        if str(court_name or "").strip()
    ]
    search_courts = _osc_unique_keep_order(search_courts)
    if search_courts:
        search_payload["courts"] = search_courts

    rr = _osc_run_skill(
        "judicial-web-search",
        _osc_skill_json_task("search", search_payload),
        timeout_sec=max(120, int(timeout_sec) + 60),
        route_key="osc:insights:fetch_full:exact_case_search",
    )
    rp = _osc_parse_skill_output(rr)
    if not rp.get("success"):
        return {"ok": False, "error": rp.get("error") or "judicial_exact_search_failed"}

    items = _osc_load_judicial_search_results(rp)
    best = _osc_pick_exact_judicial_search_result(items, title=title, case_number=case_number)
    if not best:
        return {"ok": False, "error": "judicial_exact_case_not_found"}

    fetch_payload = {
        "url": str(best.get("url") or "").strip(),
        "headless": True,
        "timeout_sec": max(45, min(180, int(timeout_sec))),
        "max_chars": 180000,
    }
    rr2 = _osc_run_skill(
        "judicial-web-search",
        _osc_skill_json_task("fetch_text", fetch_payload),
        timeout_sec=max(120, int(timeout_sec) + 60),
        route_key="osc:insights:fetch_full:exact_case_fetch",
    )
    rp2 = _osc_parse_skill_output(rr2)
    text = ""
    text_path = str(rp2.get("text_path") or "").strip() if isinstance(rp2, dict) else ""
    if text_path and os.path.exists(text_path):
        try:
            with open(text_path, "r", encoding="utf-8", errors="replace") as f:
                text = (f.read() or "").strip()
        except Exception:
            text = ""
    if (not text) and isinstance(rp2, dict):
        text = str(rp2.get("text") or "").strip()
    if (not text) and fetch_payload["url"]:
        direct = _osc_fetch_url_text(fetch_payload["url"], timeout=max(20, min(60, int(timeout_sec))))
        if direct.get("ok"):
            text = str(direct.get("text") or "").strip()
    if len(text) < 120:
        return {"ok": False, "error": "judicial_exact_case_text_not_found"}

    return {
        "ok": True,
        "source": "fallback_judicial_exact_case",
        "text": text,
        "matched": {
            "title": str(best.get("title") or ""),
            "url": str(best.get("url") or ""),
            "case_number_query": str(target.get("case_number_query") or ""),
        },
    }


def _osc_pick_best_manifest_item(items: list[dict], *, title: str = "", case_number: str = "") -> dict:
    if not isinstance(items, list) or not items:
        return {}
    target_title = str(title or "").strip()
    target_case = re.sub(r"\s+", "", str(case_number or ""))
    target_norm = _osc_title_norm(target_title)
    target_markers = _osc_extract_case_markers(target_title + " " + target_case)
    best = None
    best_score = -1.0
    for it in items:
        if not isinstance(it, dict):
            continue
        p = str(it.get("archived_text_path") or it.get("text_path") or "").strip()
        if not p or (not os.path.exists(p)):
            continue
        title2 = str(it.get("title") or "").strip()
        score = 0.0
        item_markers = _osc_extract_case_markers(title2)
        if target_case and (target_case in re.sub(r"\s+", "", title2)):
            score += 6.0
        if target_markers and item_markers:
            inter = target_markers.intersection(item_markers)
            if inter:
                score += 5.0
        n2 = _osc_title_norm(title2)
        if target_norm and n2:
            score += difflib.SequenceMatcher(None, target_norm, n2).ratio() * 2.5
        score += min(float(os.path.getsize(p) if os.path.exists(p) else 0) / 200000.0, 0.8)
        if score > best_score:
            best_score = score
            best = it
    if (not best) or best_score < 1.0:
        return {}
    out = dict(best)
    out["match_score"] = round(best_score, 4)
    return out


def _osc_summarize_legal_insight(full_text: str) -> str:
    text = str(full_text or "").strip()
    if not text:
        return ""
    prompt = (
        "你是臺灣法律實務見解萃取器。\n"
        "以下是一份裁判全文，請從中萃取「法院對法律問題的解釋與見解」，\n"
        "重點不是判決結果（誰勝誰敗、刑度多少），而是法院在判決理由中\n"
        "對法律爭點的論述、法條的解釋適用、以及可供其他案件援引的法律見解。\n\n"
        "輸出語言：繁體中文（臺灣用語）。\n"
        "輸出格式固定為：\n"
        "1) 法律爭點：本案涉及哪些法律問題\n"
        "2) 法院見解：法院對各爭點的法律解釋與論理（逐點摘錄，保留原文關鍵用語）\n"
        "3) 可援引要旨（條列）：可直接引用於書狀中的法院見解要旨，每條標註出處段落\n\n"
        "注意：\n"
        "- 不要摘要案件事實經過或判決主文\n"
        "- 聚焦於法院的法律論理、法條解釋、證據法則適用等「見解」部分\n"
        "- 若判決中引用其他判例或決議，請一併摘錄\n"
        "- 不要回覆語言偏好確認，不要加入前言，請勿杜撰。\n\n"
        f"【全文開始】\n{text[:180000]}\n【全文結束】"
    )
    bad_markers = ("近期行程", "婦女節", "確認鄭羢允案", "法扶開辦末日")

    def _clean_output(raw: str) -> str:
        cleaned = str(raw or "").strip()
        if not cleaned:
            return ""
        if _normalize_output_text:
            try:
                cleaned = _normalize_output_text(cleaned, platform="WEB")
            except Exception:
                logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 2212, exc_info=True)
        return cleaned.strip()

    def _usable(raw: str) -> bool:
        cleaned = _clean_output(raw)
        if not cleaned:
            return False
        if any(marker in cleaned for marker in bad_markers):
            return False
        if "爭點" in cleaned and ("法院見解" in cleaned or "可援引" in cleaned or "可直接引用" in cleaned):
            return True
        return len(cleaned) >= 80

    try:
        from skills.bridge.inference_gateway import InferenceGateway
        _gw = InferenceGateway()

        rr = _gw.chat(
            prompt,
            task_type="legal_analysis",
            timeout=int(os.environ.get("OSC_INSIGHT_SUMMARY_TIMEOUT_SEC", "120") or "120"),
        )
        out = _clean_output(rr.get("response") or "")
        if rr.get("success") and _usable(out):
            return out
    except Exception:
        logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 2238, exc_info=True)

    # Fallback: try shorter prompt
    fallback_prompt = (
        "請從以下裁判全文中萃取法院的法律見解（非判決結果）。\n"
        "格式：1) 法律爭點 2) 法院見解 3) 可援引要旨（條列）\n"
        "輸出繁體中文，不要摘要事實經過。\n\n"
        f"【全文開始】\n{text[:60000]}\n【全文結束】"
    )
    try:
        from skills.bridge.inference_gateway import InferenceGateway
        _gw2 = InferenceGateway()

        rr = _gw2.chat(
            fallback_prompt,
            task_type="legal_analysis",
            timeout=int(os.environ.get("OSC_INSIGHT_SUMMARY_FALLBACK_TIMEOUT_SEC", "120") or "120"),
        )
        out = _clean_output(rr.get("response") or "")
        if rr.get("success") and _usable(out):
            return out
    except Exception:
        logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 2260, exc_info=True)

    return "摘要失敗：本地摘要模型未產出可用內容。"


def _osc_fetch_fulltext_from_judicial(*, title: str = "", case_number: str = "", case_reason: str = "", timeout_sec: int = 180) -> dict:
    """
    來源被登入保護/反爬阻擋時，
    先走司法院案號精準查詢，再退回司法院全文搜尋歸檔，
    最後才退回 judgment-collector。
    """
    reason = (title or "").strip() or (case_number or "").strip() or (case_reason or "").strip()
    if not reason:
        return {"ok": False, "error": "missing_query"}
    archive_payload = {
        "query": reason,
        "max_results": 5,
        "max_chars": 180000,
        "headless": True,
        "timeout_sec": max(90, min(240, int(timeout_sec))),
    }
    collect_payload = {
        "case_reason": reason,
        "case_number": (case_number or "").strip(),
        "max_results": 5,
        "max_chars": 180000,
        "headless": True,
        "timeout_sec": max(90, min(420, int(timeout_sec))),
        "save_to_db": True,
        "notify": False,
    }

    def _try_skill(skill: str, task: str, *, route_key: str, ok_source: str, summary_source: str, inline_source: str) -> dict:
        rr = _osc_run_skill(
            skill,
            task,
            timeout_sec=max(120, int(timeout_sec) + 60),
            route_key=route_key,
        )
        rp = _osc_parse_skill_output(rr)
        if not rp.get("success"):
            return {"ok": False, "error": rp.get("error") or f"{skill}_failed"}
        items = []
        manifest_path = str(rp.get("manifest_path") or "").strip()
        if manifest_path and os.path.exists(manifest_path):
            try:
                with open(manifest_path, "r", encoding="utf-8") as f:
                    mf = json.load(f) or {}
                items = mf.get("items") or []
            except Exception:
                items = []
        if not items:
            items = rp.get("items_preview") or rp.get("items") or []
        best = _osc_pick_best_manifest_item(items, title=title, case_number=case_number)
        path = str(best.get("archived_text_path") or best.get("text_path") or "").strip()
        if path and os.path.exists(path):
            try:
                with open(path, "r", encoding="utf-8", errors="replace") as f:
                    text = (f.read() or "").strip()
                if len(text) >= 120:
                    return {
                        "ok": True,
                        "source": ok_source,
                        "text": text,
                        "matched": {
                            "title": str(best.get("title") or ""),
                            "url": str(best.get("url") or ""),
                            "path": path,
                        },
                    }
            except Exception:
                logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 2331, exc_info=True)
        summary_path = str(rp.get("summary_path") or "").strip()
        if summary_path and os.path.exists(summary_path):
            try:
                with open(summary_path, "r", encoding="utf-8", errors="replace") as f:
                    text = (f.read() or "").strip()
                if len(text) >= 120:
                    return {"ok": True, "source": summary_source, "text": text}
            except Exception:
                logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 2340, exc_info=True)
        for it in (items or []):
            txt = str(it.get("full_text") or it.get("text") or it.get("summary") or "").strip()
            if len(txt) >= 300:
                return {"ok": True, "source": inline_source, "text": txt}
        return {"ok": False, "error": f"{skill}_fulltext_not_found"}

    last_error = "judicial_fulltext_not_found"

    exact = _osc_fetch_fulltext_from_exact_case_search(
        title=title,
        case_number=case_number,
        timeout_sec=max(90, min(240, int(timeout_sec))),
    )
    if exact.get("ok"):
        return exact
    last_error = str(exact.get("error") or last_error)

    archive = _try_skill(
        "judicial-flow-search-archive",
        _osc_skill_json_task("search_archive", archive_payload),
        route_key="osc:insights:fetch_full:search_archive",
        ok_source="fallback_judicial_archive",
        summary_source="fallback_judicial_archive_summary",
        inline_source="fallback_judicial_archive_inline",
    )
    if archive.get("ok"):
        return archive
    last_error = str(archive.get("error") or last_error)

    collector = _try_skill(
        "judgment-collector",
        _osc_skill_json_task("collect", collect_payload),
        route_key="osc:insights:fetch_full:judgment_collector",
        ok_source="fallback_judgment_collector",
        summary_source="fallback_judgment_collector_summary",
        inline_source="fallback_judgment_collector_inline",
    )
    if collector.get("ok"):
        return collector
    last_error = str(collector.get("error") or last_error)
    return {"ok": False, "error": last_error}


def _osc_exec(sql, params=(), fetch="none"):
    conn, cfg = _osc_web_connect()
    cur = conn.cursor(dictionary=True)
    try:
        cur.execute(sql, params)
        if fetch == "one":
            row = cur.fetchone()
            conn.commit()
            return (_osc_row_json(row) if row else None), cfg
        if fetch == "all":
            rows = cur.fetchall() or []
            conn.commit()
            return [_osc_row_json(r) for r in rows], cfg
        conn.commit()
        return {"rowcount": cur.rowcount, "lastrowid": getattr(cur, "lastrowid", None)}, cfg
    finally:
        try:
            cur.close()
        except Exception:
            logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 2403, exc_info=True)
        try:
            conn.close()
        except Exception:
            logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 2407, exc_info=True)


def _osc_norm_case_category(v: str) -> str:
    s = str(v or "").strip()
    if not s:
        return ""
    alias = {
        "法扶案件": "法律扶助案件",
        "法扶": "法律扶助案件",
        "法律扶助": "法律扶助案件",
    }
    return alias.get(s, s)


def _osc_resolve_case_id(ref: str) -> str:
    r = str(ref or "").strip()
    if not r:
        return ""
    row, _ = _osc_exec("SELECT id FROM cases WHERE id=%s LIMIT 1", (r,), fetch="one")
    if row and row.get("id"):
        return str(row.get("id"))
    row, _ = _osc_exec("SELECT id FROM cases WHERE case_number=%s LIMIT 1", (r,), fetch="one")
    if row and row.get("id"):
        return str(row.get("id"))
    return r


def _osc_safe_int(v, default=0) -> int:
    try:
        return int(v)
    except Exception:
        return int(default)


def _osc_truthy(v) -> bool:
    return str(v or "").strip().lower() in {"1", "true", "yes", "on", "y"}


def _osc_text(v):
    s = str(v if v is not None else "").strip()
    return s or None


def _osc_current_actor() -> str:
    try:
        if current_user.is_authenticated and getattr(current_user, "username", None):
            return str(current_user.username)
    except Exception:
        logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 2456, exc_info=True)
    return "web"


def _osc_log_activity(action: str, entity_type: str = "", entity_id: str = "", details=None) -> None:
    try:
        payload = details
        if isinstance(payload, (dict, list)):
            payload = json.dumps(payload, ensure_ascii=False)
        payload = _osc_text(payload)
        _osc_exec(
            "INSERT INTO activity_logs (action, entity_type, entity_id, details, user) VALUES (%s,%s,%s,%s,%s)",
            (
                _osc_text(action) or "osc_action",
                _osc_text(entity_type),
                _osc_text(entity_id),
                payload,
                _osc_current_actor(),
            ),
            fetch="none",
        )
    except Exception as e:
        logger.warning("OSC activity log write failed: %s", e)


def _osc_accounting_window(today: date | None = None) -> tuple[date, date]:
    today = today or date.today()
    if today.day <= 25:
        end_date = date(today.year, today.month, 25)
        if today.month == 1:
            start_date = date(today.year - 1, 12, 26)
        else:
            start_date = date(today.year, today.month - 1, 26)
    else:
        start_date = date(today.year, today.month, 26)
        if today.month == 12:
            end_date = date(today.year + 1, 1, 25)
        else:
            end_date = date(today.year, today.month + 1, 25)
    return start_date, end_date


def _osc_get_setting_value(setting_key: str, default: str = "") -> str:
    key = str(setting_key or "").strip()
    if not key:
        return str(default or "")
    try:
        row, _ = _osc_exec("SELECT value FROM settings WHERE `key`=%s", (key,), fetch="one")
        value = (row or {}).get("value")
        if value is None:
            return str(default or "")
        text = str(value).strip()
        return text if text else str(default or "")
    except Exception:
        return str(default or "")


def _osc_unique_strings(values) -> list[str]:
    seen = set()
    out = []
    for raw in values or []:
        s = str(raw or "").strip()
        if not s or s in seen:
            continue
        seen.add(s)
        out.append(s)
    return out


def _osc_read_plain_text(path: str) -> str:
    for enc in ("utf-8", "utf-8-sig", "cp950", "big5"):
        try:
            with open(path, "r", encoding=enc) as f:
                return f.read()
        except Exception:
            continue
    try:
        with open(path, "rb") as f:
            return f.read().decode("utf-8", errors="ignore")
    except Exception:
        return ""


def _osc_read_docx_text(path: str) -> str:
    try:
        from docx import Document  # type: ignore

        doc = Document(path)
        return "\n".join(p.text for p in doc.paragraphs if str(p.text or "").strip())
    except Exception:
        return ""


def _osc_read_pdf_text(path: str) -> str:
    try:
        from pypdf import PdfReader  # type: ignore

        reader = PdfReader(path)
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    except Exception:
        logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 2556, exc_info=True)
    try:
        from PyPDF2 import PdfReader  # type: ignore

        reader = PdfReader(path)
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    except Exception:
        logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 2563, exc_info=True)
    tool = shutil.which("pdftotext")
    if tool:
        try:
            proc = subprocess.run(
                [tool, "-layout", path, "-"],
                capture_output=True,
                text=True,
                encoding="utf-8",
                errors="ignore",
                check=False,
            )
            if proc.returncode == 0:
                return proc.stdout or ""
        except Exception:
            logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 2578, exc_info=True)
    return ""


def _osc_read_textutil_text(path: str) -> str:
    tool = shutil.which("textutil")
    if not tool:
        return ""
    try:
        proc = subprocess.run(
            [tool, "-convert", "txt", "-stdout", path],
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="ignore",
            check=False,
        )
        if proc.returncode == 0:
            return proc.stdout or ""
    except Exception:
        return ""
    return ""


def _osc_resolve_existing_local_path_with_candidates(raw_path: str) -> tuple[str, list[str]]:
    candidates = []
    norm = _osc_norm_path(raw_path).replace("\\", "/")
    if norm:
        candidates.append(norm)
    candidates.extend(_osc_local_path_candidates(raw_path))
    uniq = _osc_unique_strings(candidates)
    for cand in uniq:
        if cand and os.path.exists(cand):
            return cand, uniq
    return "", uniq


def _osc_read_reference_document(raw_path: str, max_chars: int = 9000) -> dict:
    resolved, candidates = _osc_resolve_existing_local_path_with_candidates(raw_path)
    meta = {
        "file_path": str(raw_path or "").strip(),
        "resolved_path": resolved,
        "candidates": candidates,
        "ok": False,
        "text": "",
        "error": "",
    }
    if not resolved:
        meta["error"] = "file_not_found"
        return meta

    ext = Path(resolved).suffix.lower()
    text = ""
    if ext == ".docx":
        text = _osc_read_docx_text(resolved) or _osc_read_textutil_text(resolved)
    elif ext == ".pdf":
        text = _osc_read_pdf_text(resolved)
    elif ext in {".txt", ".md", ".text", ".log", ".json", ".yaml", ".yml", ".csv", ".tsv", ".rst"}:
        text = _osc_read_plain_text(resolved)
    elif ext in {".doc", ".rtf", ".odt", ".html", ".htm"}:
        text = _osc_read_textutil_text(resolved)
    else:
        try:
            if os.path.getsize(resolved) <= 2_000_000:
                text = _osc_read_plain_text(resolved)
        except Exception:
            logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 2644, exc_info=True)
        if not text:
            text = _osc_read_textutil_text(resolved)

    text = re.sub(r"\n{3,}", "\n\n", str(text or "")).strip()
    if len(text) > max_chars:
        text = text[:max_chars].rstrip() + "\n...[截斷]"
    meta["ok"] = bool(text)
    meta["text"] = text
    if not text:
        meta["error"] = "text_extract_failed"
    return meta


def _osc_clean_draft_output(text: str) -> str:
    cleaned = ihtml.unescape(str(text or ""))
    cleaned = cleaned.replace("<br>", "\n").replace("<br/>", "\n").replace("<br />", "\n")
    cleaned = re.sub(r"^#+\s*", "", cleaned, flags=re.MULTILINE)
    cleaned = re.sub(r"\*\*(.+?)\*\*", r"\1", cleaned)
    cleaned = re.sub(r"\*(.+?)\*", r"\1", cleaned)
    cleaned = re.sub(r"__(.+?)__", r"\1", cleaned)
    cleaned = re.sub(r"_(.+?)_", r"\1", cleaned)
    cleaned = re.sub(r"^[-*_]{3,}\s*$", "", cleaned, flags=re.MULTILINE)
    cleaned = re.sub(r"```(?:[\s\S]*?)```", "", cleaned)
    cleaned = re.sub(r"`(.+?)`", r"\1", cleaned)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip()


def _osc_render_draft_template(template: str, values: dict) -> str:
    rendered = str(template or "")
    for key, value in (values or {}).items():
        rendered = rendered.replace("{" + str(key) + "}", str(value or ""))
    return rendered


def _osc_draft_enabled_flag() -> bool:
    raw = _osc_get_setting_value("enable_draft_generation", "")
    if raw:
        return _osc_truthy(raw)
    cfg_val = RUNTIME_CONFIG.get("enable_draft_generation")
    if isinstance(cfg_val, bool):
        return cfg_val
    if cfg_val is None:
        return True
    return _osc_truthy(cfg_val)


def _osc_draft_defendant(case_row: dict, payload: dict) -> str:
    explicit = str(payload.get("defendant") or payload.get("opponent_name") or "").strip()
    if explicit:
        return explicit
    row_value = str((case_row or {}).get("opponent_name") or "").strip()
    if row_value:
        return row_value
    case_number = str((payload.get("case_number") or (case_row or {}).get("case_number") or "")).strip()
    if not case_number:
        return ""
    try:
        rows, _ = _osc_exec(
            "SELECT name FROM opponents WHERE case_number=%s AND (is_active=1 OR is_active IS NULL) ORDER BY updated_date DESC, id DESC LIMIT 5",
            (case_number,),
            fetch="all",
        )
        names = _osc_unique_strings((r.get("name") for r in (rows or [])))
        return "、".join(names)
    except Exception:
        return ""


def _osc_resolve_draft_insights(payload: dict) -> list[dict]:
    selected = payload.get("selected_insights") or payload.get("insights") or []
    if not selected:
        selected = [{"id": x} for x in (payload.get("selected_insight_ids") or [])]
    lookup = {str(it.get("id")): it for it in _osc_collect_insights()}
    out = []
    for raw in (selected or [])[:10]:
        if isinstance(raw, str):
            raw = {"id": raw}
        if not isinstance(raw, dict):
            continue
        sid = str(raw.get("id") or "").strip()
        base = lookup.get(sid, {})
        title = str(raw.get("title") or raw.get("reference") or base.get("title") or "").strip()
        summary = str(raw.get("summary") or raw.get("insight_text") or base.get("summary") or base.get("insight_text") or "").strip()
        full_text = str(raw.get("full_text") or raw.get("text") or base.get("full_text") or "").strip()
        case_number = str(raw.get("case_number") or base.get("case_number") or "").strip()
        case_reason = str(raw.get("case_reason") or raw.get("reason") or base.get("case_reason") or "").strip()
        court = str(raw.get("court") or base.get("court") or "").strip()
        if not full_text and summary:
            full_text = summary
        if not title:
            title = "實務見解"
        if not (summary or full_text):
            continue
        out.append(
            {
                "id": sid or f"manual-{len(out) + 1}",
                "title": title,
                "summary": summary or full_text[:350],
                "full_text": full_text,
                "case_number": case_number,
                "case_reason": case_reason,
                "court": court,
            }
        )
    return out


def _osc_collect_draft_reference_style(payload: dict) -> tuple[str, list[dict], list[str]]:
    selected = payload.get("selected_documents") or payload.get("reference_documents") or []
    blocks = []
    refs = []
    warnings = []
    for raw in (selected or [])[:3]:
        if not isinstance(raw, dict):
            continue
        file_path = str(raw.get("file_path") or raw.get("path") or "").strip()
        file_name = str(raw.get("file_name") or os.path.basename(file_path or "") or "參考文件").strip()
        provided = str(raw.get("text") or raw.get("content") or "").strip()
        read_meta = None
        text = provided
        if not text and file_path:
            read_meta = _osc_read_reference_document(file_path, max_chars=9000)
            text = str(read_meta.get("text") or "").strip()
        if text:
            excerpt = text[:3000].strip()
            blocks.append(f"--- 參考範本：{file_name} ---\n{excerpt}")
        else:
            reason = str((read_meta or {}).get("error") or "no_content")
            warnings.append(f"{file_name}: {reason}")
        refs.append(
            {
                "id": str(raw.get("id") or "").strip(),
                "file_name": file_name,
                "file_path": file_path,
                "resolved_path": str((read_meta or {}).get("resolved_path") or "").strip(),
                "loaded": bool(text),
            }
        )
    return ("\n\n".join(blocks).strip() or "(無參考範本)"), refs, warnings


def _osc_build_draft_context(payload: dict) -> dict:
    body = payload or {}
    case_row = {}
    case_id = str(body.get("case_id") or body.get("selected_case_id") or "").strip()
    lookup_number = str(body.get("case_lookup_number") or body.get("selected_case_number") or body.get("case_number") or "").strip()

    if case_id:
        case_row, _ = _osc_exec("SELECT * FROM cases WHERE id=%s LIMIT 1", (case_id,), fetch="one")
    elif lookup_number:
        case_row, _ = _osc_exec(
            """
            SELECT * FROM cases
            WHERE case_number=%s OR court_case_no=%s OR court_case_number=%s
            ORDER BY updated_at DESC, created_date DESC
            LIMIT 1
            """,
            (lookup_number, lookup_number, lookup_number),
            fetch="one",
        )
    case_row = case_row or {}

    doc_type = str(body.get("doc_type") or body.get("document_type") or "").strip()
    case_number = str(
        body.get("case_number")
        or case_row.get("court_case_number")
        or case_row.get("court_case_no")
        or case_row.get("case_number")
        or ""
    ).strip()
    division = str(body.get("division") or case_row.get("court_division") or "").strip()
    court_name = str(body.get("court_name") or case_row.get("court_name") or "").strip()
    reason = str(body.get("reason") or case_row.get("case_reason") or "").strip()
    plaintiff = str(body.get("plaintiff") or case_row.get("client_name") or "").strip()
    defendant = _osc_draft_defendant(case_row, body)
    case_facts = str(body.get("case_facts") or body.get("facts") or case_row.get("description") or case_row.get("notes") or "").strip()

    selected_insights = _osc_resolve_draft_insights(body)
    legal_insights = ""
    for i, insight in enumerate(selected_insights[:5], 1):
        ref = str(insight.get("title") or "實務見解").strip()
        court = str(insight.get("court") or "").strip()
        summary = str(insight.get("summary") or insight.get("insight_text") or insight.get("full_text") or "").strip()[:2000]
        label = f"{ref}"
        if court:
            label = f"{court}｜{label}"
        legal_insights += f"\n{i}. 【{label}】\n{summary}\n"

    reference_style, references, warnings = _osc_collect_draft_reference_style(body)
    custom_template = _osc_get_setting_value("draft_prompt_template", "").strip()
    template = custom_template if custom_template else _OSC_DRAFT_PROMPT_TEMPLATE
    values = {
        "doc_type": doc_type or "(未指定)",
        "case_number": case_number or "(待填)",
        "division": division or "(待填)",
        "court_name": court_name or "(待填)",
        "reason": reason or "(未指定)",
        "plaintiff": plaintiff or "(待填)",
        "defendant": defendant or "(待填)",
        "case_facts": case_facts or "(未提供)",
        "legal_insights": legal_insights or "(無)",
        "reference_style": reference_style or "(無參考範本)",
    }
    prompt = _osc_render_draft_template(template, values)
    suggested_filename = str(body.get("suggested_filename") or "").strip()
    if not suggested_filename:
        parts = [doc_type or "書狀草稿", case_number or case_row.get("case_number") or "未命名"]
        suggested_filename = "_".join(str(p).strip() for p in parts if str(p).strip())
    return {
        "case": case_row,
        "doc_type": doc_type,
        "case_number": case_number,
        "division": division,
        "court_name": court_name,
        "reason": reason,
        "plaintiff": plaintiff,
        "defendant": defendant,
        "case_facts": case_facts,
        "selected_insights": selected_insights,
        "selected_documents": references,
        "warnings": warnings,
        "prompt": prompt,
        "template_source": "custom" if custom_template else "default",
        "suggested_filename": suggested_filename,
        "export_title": doc_type or "書狀草稿",
    }


def _osc_generate_draft_with_casper(prompt: str) -> str:
    skill_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "skills", "casper-client"))
    if skill_dir not in sys.path:
        sys.path.append(skill_dir)
    try:
        from casper_tools_client import casper_chat  # type: ignore
    except Exception as e:
        raise RuntimeError(f"CASPER 客戶端載入失敗: {e}")

    result = casper_chat(prompt, timeout_sec=300)
    if not isinstance(result, dict) or not result.get("success"):
        err = (result.get("error") if isinstance(result, dict) else "") or "unknown_error"
        raise RuntimeError(f"CASPER 生成失敗: {err}")
    text = str(result.get("response") or "").strip()
    if not text:
        raise RuntimeError("CASPER 生成失敗: empty response")
    return text


def _osc_generate_draft_with_ollama(prompt: str, model: str, ollama_url: str) -> str:
    """透過 oMLX (OpenAI-compatible API) 生成草稿。保留函式名以相容既有呼叫端。"""
    base = (ollama_url or "http://127.0.0.1:8080").rstrip("/")
    url = base + "/v1/chat/completions"
    body = {
        "model": model or "TAIDE-12b-Chat-mlx-4bit",
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.3,
        "max_tokens": 2048,
        "stream": False,
    }
    payload = json.dumps(body).encode("utf-8")
    req = urllib.request.Request(url, data=payload, headers={"Content-Type": "application/json"}, method="POST")
    try:
        with urllib.request.urlopen(req, timeout=300) as resp:
            raw = resp.read().decode("utf-8", errors="ignore")
    except urllib.error.URLError as e:
        raise RuntimeError(f"oMLX 連線失敗: {e}")
    try:
        data = json.loads(raw or "{}")
    except Exception as e:
        raise RuntimeError(f"oMLX 回應解析失敗: {e}")
    choices = data.get("choices") or []
    text = (choices[0].get("message", {}).get("content", "") if choices else "").strip()
    if not text:
        raise RuntimeError("oMLX 未回傳內容")
    return text


def _osc_generate_draft_with_gemini(prompt: str) -> tuple[str, str]:
    allow_cloud = str(os.environ.get("MAGI_ALLOW_CLOUD_MODELS", "0") or "").strip().lower() in {"1", "true", "yes", "on"}
    if not allow_cloud:
        return _osc_generate_draft_with_casper(prompt), "casper"
    api_key = (
        os.environ.get("GEMINI_API_KEY")
        or _osc_get_setting_value("gemini_api_key", "")
        or ""
    ).strip()
    model_name = (
        os.environ.get("GEMINI_MODEL")
        or _osc_get_setting_value("gemini_model", "")
        or str(RUNTIME_CONFIG.get("gemini_model") or "").strip()
        or "gemini-2.0-flash"
    ).strip()
    if not api_key:
        raise RuntimeError("未設定 Gemini API Key")
    try:
        import google.generativeai as genai  # type: ignore
    except Exception as e:
        raise RuntimeError(f"google-generativeai 套件不可用: {e}")
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(model_name)
        response = model.generate_content(prompt)
        text = str(getattr(response, "text", "") or "").strip()
        if not text:
            raise RuntimeError("Gemini 未回傳內容")
        return text, model_name
    except Exception as e:
        raise RuntimeError(f"Gemini API 錯誤: {e}")


def _osc_get_case_identity_by_payload(payload: dict) -> dict:
    p = payload or {}
    row = None
    row_id = str(p.get("case_id") or p.get("id") or "").strip()
    case_number = str(p.get("case_number") or "").strip()
    laf_case_no = str(p.get("laf_case_no") or p.get("laf_case_number") or "").strip()
    client_name = str(p.get("client_name") or "").strip()

    if row_id:
        row, _ = _osc_exec(
            """
            SELECT id, case_number, client_name, case_category, case_stage, case_reason, status, folder_path,
                   laf_case_no, application_no, court_case_no
            FROM cases
            WHERE id=%s
            LIMIT 1
            """,
            (row_id,),
            fetch="one",
        )
    if (not row) and case_number:
        row, _ = _osc_exec(
            """
            SELECT id, case_number, client_name, case_category, case_stage, case_reason, status, folder_path,
                   laf_case_no, application_no, court_case_no
            FROM cases
            WHERE case_number=%s
            LIMIT 1
            """,
            (case_number,),
            fetch="one",
        )
    if (not row) and laf_case_no:
        row, _ = _osc_exec(
            """
            SELECT id, case_number, client_name, case_category, case_stage, case_reason, status, folder_path,
                   laf_case_no, application_no, court_case_no
            FROM cases
            WHERE laf_case_no=%s
            LIMIT 1
            """,
            (laf_case_no,),
            fetch="one",
        )
    if (not row) and client_name:
        row, _ = _osc_exec(
            """
            SELECT id, case_number, client_name, case_category, case_stage, case_reason, status, folder_path,
                   laf_case_no, application_no, court_case_no
            FROM cases
            WHERE client_name=%s
            ORDER BY updated_at DESC, created_date DESC
            LIMIT 1
            """,
            (client_name,),
            fetch="one",
        )
    return row or {}


def _osc_build_form_preview(form_type: str, case_row: dict, fields: dict) -> dict:
    ftype = str(form_type or "").strip().lower()
    if ftype in {"poa", "power_of_attorney", "委任狀", "委任状"}:
        ftype = "power_of_attorney"
    elif ftype in {"receipt", "收據", "收据"}:
        ftype = "receipt"
    elif ftype in {"contract", "契約", "契約書", "契约", "契约书"}:
        ftype = "contract"
    else:
        raise ValueError("unsupported_form_type")

    c = case_row or {}
    f = fields or {}
    today = datetime.now().strftime("%Y-%m-%d")
    
    if ftype == "contract":
        title = "契約書草稿"
        doc = (
            f"{title}\n\n"
            f"日期：{f.get('date') or today}\n"
            f"當事人：{f.get('client_name') or c.get('client_name') or ''}\n"
            f"案件編號：{f.get('case_number') or c.get('case_number') or ''}\n"
            f"法院案號：{f.get('court_case_no') or c.get('court_case_no') or ''}\n"
            f"法扶案號：{f.get('laf_case_no') or c.get('laf_case_no') or ''}\n"
            f"受任律師：{f.get('lawyer_name') or '＿＿＿＿'}\n"
            f"費用項目：{f.get('item') or ''}\n"
            f"金額：{f.get('amount') or ''}\n"
            f"備註：{f.get('notes') or ''}\n"
            f"\n\n（以下為契約條文草稿，請自行替換正文）"
        )
        filename = f"契約書草稿_{c.get('case_number') or '未指定位案件'}"
        return {"form_type": ftype, "title": title, "preview_text": doc, "suggested_filename": filename}

    if ftype == "power_of_attorney":
        title = "委任狀草稿"
        doc = (
            f"{title}\n\n"
            f"日期：{f.get('date') or today}\n"
            f"當事人：{f.get('client_name') or c.get('client_name') or ''}\n"
            f"案件編號：{f.get('case_number') or c.get('case_number') or ''}\n"
            f"法院案號：{f.get('court_case_no') or c.get('court_case_no') or ''}\n"
            f"法扶案號：{f.get('laf_case_no') or c.get('laf_case_no') or ''}\n"
            f"案由：{f.get('case_reason') or c.get('case_reason') or ''}\n"
            f"受任律師：{f.get('lawyer_name') or '＿＿＿＿'}\n"
            f"備註：{f.get('notes') or ''}\n"
        )
        filename = f"委任狀草稿_{c.get('case_number') or '未指定位案件'}"
        return {"form_type": ftype, "title": title, "preview_text": doc, "suggested_filename": filename}

    amount = f.get("amount") or ""
    if amount:
        try:
            amount = f"{float(amount):,.0f}"
        except Exception:
            amount = str(amount)
    title = "收據草稿"
    doc = (
        f"{title}\n\n"
        f"日期：{f.get('date') or today}\n"
        f"收據編號：{f.get('receipt_no') or ''}\n"
        f"當事人：{f.get('client_name') or c.get('client_name') or ''}\n"
        f"案件編號：{f.get('case_number') or c.get('case_number') or ''}\n"
        f"法扶案號：{f.get('laf_case_no') or c.get('laf_case_no') or ''}\n"
        f"費用項目：{f.get('item') or '法律服務費'}\n"
        f"金額：{amount}\n"
        f"付款方式：{f.get('payment_method') or ''}\n"
        f"備註：{f.get('notes') or ''}\n"
    )
    filename = f"收據草稿_{c.get('case_number') or '未指定位案件'}"
    return {"form_type": ftype, "title": title, "preview_text": doc, "suggested_filename": filename}


def _osc_import_laf_orchestrator():
    ensure_path_on_sys_path(get_orch_dir())
    from laf_orchestrator import LAFOrchestrator  # type: ignore
    return LAFOrchestrator


def _osc_map_laf_action(action: str) -> str:
    a = str(action or "").strip().lower()
    aliases = {
        "開辦": "go_live",
        "go_live": "go_live",
        "golive": "go_live",
        "疑義": "inquiry",
        "inquiry": "inquiry",
        "訴訟中費用支付": "fee",
        "fee": "fee",
        "二階段": "condition",
        "condition": "condition",
        "撤回": "withdrawal",
        "withdrawal": "withdrawal",
        "結案": "closing",
        "closing": "closing",
    }
    return aliases.get(a, a)


def _osc_prepare_laf_identity(payload: dict) -> dict:
    case_row = _osc_get_case_identity_by_payload(payload)
    return {
        "case_row": case_row,
        "laf_case_number": str(payload.get("laf_case_no") or payload.get("laf_case_number") or case_row.get("laf_case_no") or "").strip(),
        "case_number": str(payload.get("case_number") or case_row.get("case_number") or "").strip(),
        "client_name": str(payload.get("client_name") or case_row.get("client_name") or "").strip(),
    }


def _osc_enrich_portal_preview(artifact: dict) -> dict:
    art = dict(artifact or {})
    png = str(art.get("png") or "").strip()
    html = str(art.get("html") or "").strip()
    png_export = art.get("png_export") if isinstance(art.get("png_export"), dict) else {}
    html_export = art.get("html_export") if isinstance(art.get("html_export"), dict) else {}

    if (not png_export) and png:
        u = _public_url_for_local_file(png)
        if u:
            png_export = {"url": u, "path": png}
            art["png_export"] = png_export
    if (not html_export) and html:
        u = _public_url_for_local_file(html)
        if u:
            html_export = {"url": u, "path": html}
            art["html_export"] = html_export
    return art


def _osc_get_closed_archive_base() -> str:
    env_base = (os.environ.get("MAGI_CLOSED_CASE_ARCHIVE_PATH") or "").strip()
    if env_base:
        return env_base
    try:
        ensure_path_on_sys_path(get_orch_dir())
        from osc_core.paths import get_closed_case_archive_path  # type: ignore
        p = (get_closed_case_archive_path() or "").strip()
        if p:
            return p
    except Exception:
        logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 3154, exc_info=True)
    roots = preferred_case_roots(include_closed=True)
    if len(roots) > 1:
        return roots[1]
    if roots:
        return roots[0]
    return str(Path.home() / "Library" / "CloudStorage" / "SynologyDrive-homes" / "99_結案案件")


def _osc_build_archive_preview(limit: int = 300) -> dict:
    rows, _ = _osc_exec(
        """
        SELECT id, case_number, client_name, status, folder_path, updated_at
        FROM cases
        WHERE (status LIKE %s OR status LIKE %s OR LOWER(status)='closed')
        ORDER BY updated_at DESC, created_date DESC
        LIMIT %s
        """,
        ("%結案%", "%Closed%", int(limit)),
        fetch="all",
    )
    archive_base = _osc_get_closed_archive_base()
    archive_local_candidates = _osc_local_path_candidates(_osc_norm_path(archive_base))
    archive_local = ""
    for c in archive_local_candidates:
        if os.path.exists(c):
            archive_local = c
            break
    if (not archive_local) and archive_local_candidates:
        archive_local = archive_local_candidates[0]
    items = []
    for r in (rows or []):
        source_raw = (r.get("folder_path") or "").strip() or _osc_guess_case_folder(r.get("case_number") or "")
        source_norm = _osc_norm_path(source_raw)
        local_candidates = _osc_local_path_candidates(source_norm)
        source_local = ""
        for c in local_candidates:
            if c and os.path.exists(c):
                source_local = c
                break
        folder_name = os.path.basename(source_local.rstrip("/")) if source_local else os.path.basename(source_norm.rstrip("/"))
        target_local = os.path.join(archive_local, folder_name) if archive_local and folder_name else ""
        target_exists = bool(target_local and os.path.exists(target_local))
        source_exists = bool(source_local and os.path.exists(source_local))
        item = {
            "id": r.get("id"),
            "case_number": r.get("case_number") or "",
            "client_name": r.get("client_name") or "",
            "status": r.get("status") or "",
            "source_path": source_norm,
            "source_local": source_local,
            "source_exists": source_exists,
            "target_local": target_local,
            "target_exists": target_exists,
            "ready": bool(source_exists and target_local and (not target_exists)),
            "updated_at": r.get("updated_at"),
        }
        if not source_exists:
            item["reason"] = "來源資料夾不存在或未同步到本機"
        elif target_exists:
            item["reason"] = "封存目標已存在"
        else:
            item["reason"] = "可搬移"
        items.append(item)
    return {
        "ok": True,
        "archive_base": archive_base,
        "archive_local": archive_local,
        "items": items,
        "summary": {
            "total": len(items),
            "ready": len([x for x in items if x.get("ready")]),
            "missing_source": len([x for x in items if (not x.get("source_exists"))]),
            "target_exists": len([x for x in items if x.get("target_exists")]),
        },
    }


def _osc_template_data_json_or_wrap(v: str | None) -> str | None:
    """
    document_templates.template_data has CHECK(json_valid(...)).
    Accept plain text from UI and wrap into JSON to avoid 4025 failures.
    """
    s = (v or "").strip()
    if not s:
        return None
    try:
        parsed = json.loads(s)
        return json.dumps(parsed, ensure_ascii=False)
    except Exception:
        return json.dumps({"content": s}, ensure_ascii=False)


def _osc_json_or_wrap(v, fallback_key: str = "content") -> str | None:
    s = ("" if v is None else str(v)).strip()
    if not s:
        return None
    try:
        parsed = json.loads(s)
        return json.dumps(parsed, ensure_ascii=False)
    except Exception:
        return json.dumps({fallback_key: s}, ensure_ascii=False)


def _osc_collect_insights():
    items = []
    conn, _cfg = _osc_web_connect()
    cur = conn.cursor(dictionary=True)
    try:
        try:
            cur.execute(
                """
                SELECT id, case_number, document_name, court_reference, court_type,
                       insight_type, insight_text, case_reason, source_file, extracted_date, raw_text
                FROM legal_insights
                ORDER BY extracted_date DESC, id DESC
                LIMIT 500
                """
            )
            for r in (cur.fetchall() or []):
                title = (r.get("document_name") or r.get("insight_type") or "實務見解").strip()
                # insight_text = 結構化法律見解萃取結果；raw_text = 判決原文
                insight_text = (r.get("insight_text") or "").strip()
                raw_text = (r.get("raw_text") or "").strip()
                full_text = raw_text or insight_text
                summary = (insight_text or full_text[:500])[:350]
                ts = r.get("extracted_date")
                source_file = str(r.get("source_file") or "").strip()
                source_url = source_file if source_file.lower().startswith(("http://", "https://")) else ""
                items.append(
                    {
                        "id": f"li-{r.get('id')}",
                        "source_type": "legal_insights",
                        "source": "見解庫",
                        "title": title,
                        "summary": summary,
                        "insight_text": insight_text,
                        "full_text": full_text,
                        "url": source_url,
                        "case_number": r.get("case_number") or "",
                        "case_reason": r.get("case_reason") or "",
                        "court": r.get("court_reference") or r.get("court_type") or "",
                        "timestamp": _osc_json_value(ts) if ts else "",
                        "sort_ts": _osc_parse_dt(ts).timestamp() if _osc_parse_dt(ts) else 0,
                    }
                )
        except Exception:
            logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 3301, exc_info=True)

        try:
            cur.execute(
                """
                SELECT id, jid, court_name, case_number, case_type, judgment_date,
                       summary, full_text, source_url, crawled_at
                FROM court_judgments
                ORDER BY crawled_at DESC, id DESC
                LIMIT 500
                """
            )
            for r in (cur.fetchall() or []):
                title = f"{(r.get('court_name') or '').strip()} {(r.get('case_number') or '').strip()}".strip() or "裁判見解"
                full_text = (r.get("full_text") or r.get("summary") or "").strip()
                summary = (r.get("summary") or full_text[:350] or "").strip()
                ts = r.get("crawled_at") or r.get("judgment_date")
                items.append(
                    {
                        "id": f"cj-{r.get('id')}",
                        "source_type": "court_judgments",
                        "source": "裁判書",
                        "title": title,
                        "summary": summary,
                        "full_text": full_text,
                        "url": r.get("source_url") or "",
                        "case_number": r.get("case_number") or "",
                        "case_reason": r.get("case_type") or "",
                        "court": r.get("court_name") or "",
                        "timestamp": _osc_json_value(ts) if ts else "",
                        "sort_ts": _osc_parse_dt(ts).timestamp() if _osc_parse_dt(ts) else 0,
                    }
                )
        except Exception:
            logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 3335, exc_info=True)
    finally:
        try:
            cur.close()
        except Exception:
            logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 3340, exc_info=True)
        try:
            conn.close()
        except Exception:
            logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 3344, exc_info=True)

    # Merge legacy judgments json so old workflow remains visible.
    try:
        json_path = os.path.abspath(
            os.path.join(os.path.dirname(__file__), "..", "skills", "judgment-collector", "judgments.json")
        )
        if os.path.exists(json_path):
            with open(json_path, "r", encoding="utf-8") as f:
                data = json.load(f) or []
            if isinstance(data, list):
                for i, r in enumerate(data):
                    if not isinstance(r, dict):
                        continue
                    full_text = (r.get("full_text") or r.get("summary") or "").strip()
                    ts = r.get("timestamp")
                    items.append(
                        {
                            "id": f"json-{i}",
                            "source_type": "judgments_json",
                            "source": r.get("source") or "爬蟲快照",
                            "title": r.get("title") or "裁判資料",
                            "summary": (r.get("summary") or "")[:350],
                            "full_text": full_text,
                            "url": r.get("url") or "",
                            "case_number": r.get("case_number") or "",
                            "case_reason": r.get("case_reason") or "",
                            "court": r.get("court_name") or "",
                            "timestamp": ts or "",
                            "sort_ts": _osc_parse_dt(ts).timestamp() if _osc_parse_dt(ts) else 0,
                        }
                    )
    except Exception as e:
        logger.warning(f"osc insights json merge failed: {e}")

    items.sort(key=lambda x: x.get("sort_ts") or 0, reverse=True)
    for it in items:
        it.pop("sort_ts", None)
    return items


_OSC_DRAFT_DOC_TYPES = [
    "民事起訴狀",
    "民事答辯狀",
    "民事準備書狀",
    "民事上訴狀",
    "民事聲請狀",
    "刑事告訴狀",
    "刑事答辯狀",
    "刑事上訴狀",
    "刑事聲請狀",
    "刑事陳報狀",
    "行政起訴狀",
    "行政答辯狀",
    "抗告狀",
    "聲明異議狀",
    "強制執行聲請狀",
    "假扣押聲請狀",
    "假處分聲請狀",
    "支付命令聲請狀",
    "本票裁定聲請狀",
]


_OSC_DRAFT_PROMPT_TEMPLATE = """你是一位專業的台灣律師助理，請根據以下資料協助草擬法律文書。

## 書狀類型
{doc_type}

## 案件基本資訊
- 案號：{case_number}
- 股別：{division}
- 法院/地檢署：{court_name}
- 案由：{reason}
- 原告/聲請人：{plaintiff}
- 被告/相對人：{defendant}

## 案件事實
{case_facts}

## 參考實務見解
{legal_insights}

## 書寫風格參考（以下為過往類似書狀的格式範例）
{reference_style}

## 要求
1. 請按照上述參考風格撰寫完整的{doc_type}
2. 格式需符合台灣法院規範
3. 適當引用提供的實務見解（如有提供）
4. 確保案號、股別、法院名稱正確填入狀頭
5. 論述需有邏輯、條理分明
6. 請加入常見的法律用語和格式

請直接輸出完整書狀內容：
"""


_OSC_DOC_KIND_KEYWORDS = {
    "all": [],
    "poa": ["委任", "委託", "委任狀", "委任书", "委託書"],
    "receipt": ["收據", "收执", "收執", "繳費", "訴訟中費用", "粉紅"],
    "laf": ["法扶", "法律扶助", "接案通知", "開辦資料", "開辦通知"],
    "judgment": ["判決", "裁定", "調解不成立", "和解", "決定書"],
    "court_notice": ["通知", "庭期", "開庭", "法院通知"],
}


def _osc_doc_kind_match(kind: str, blob: str) -> bool:
    k = str(kind or "all").strip().lower()
    if k in {"", "all"}:
        return True
    kws = _OSC_DOC_KIND_KEYWORDS.get(k, [])
    if not kws:
        return True
    b = str(blob or "")
    return any(x in b for x in kws)


def _osc_doc_kind_label(blob: str) -> str:
    b = str(blob or "")
    if _osc_doc_kind_match("poa", b):
        return "委任狀/委託書"
    if _osc_doc_kind_match("receipt", b):
        return "收據/繳費"
    if _osc_doc_kind_match("laf", b):
        return "法扶資料"
    if _osc_doc_kind_match("court_notice", b):
        return "法院通知"
    if _osc_doc_kind_match("judgment", b):
        return "判決/裁定"
    return "一般文件"


@app.route("/api/osc/meta", methods=["GET"])
@login_required
def osc_meta_api():
    try:
        conn, cfg = _osc_web_connect()
        cur = conn.cursor(dictionary=True)
        try:
            cur.execute("SELECT CURRENT_USER() AS current_user_name")
            who = cur.fetchone() or {}
            counts = {}
            for tbl in [
                "cases",
                "clients",
                "meetings",
                "case_todos",
                "legal_insights",
                "court_judgments",
                "case_transactions",
                "document_index",
                "document_templates",
                "document_keywords",
                "document_replacements",
                "expense_defaults",
                "recurring_expenses",
                "quotations",
                "quotation_templates",
                "calendar_events",
                "legal_aid_checklists",
                "laf_lifecycle_log",
                "laf_email_records",
            ]:
                try:
                    cur.execute(f"SELECT COUNT(*) AS c FROM `{tbl}`")
                    counts[tbl] = int((cur.fetchone() or {}).get("c") or 0)
                except Exception:
                    counts[tbl] = None
            return jsonify(
                {
                    "ok": True,
                    "db": {
                        "host": cfg["host"],
                        "port": int(cfg["port"]),
                        "database": cfg["database"],
                        "user": cfg["user"],
                        "current_user": who.get("current_user_name") or "",
                    },
                    "counts": counts,
                }
            )
        finally:
            try:
                cur.close()
            except Exception:
                logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 3531, exc_info=True)
            try:
                conn.close()
            except Exception:
                logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 3535, exc_info=True)
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


def _osc_auto_create_folder_for_case(row_id: str, payload: dict, case_category: str) -> dict:
    """建立案件資料夾並更新 DB，回傳結果 dict。供 POST /api/osc/cases 使用。"""
    from casper_ecosystem.law_firm_orchestrators.osc.folder_utils import (
        build_full_case_path,
        create_folder_structure,
    )

    case_roots = preferred_case_roots()
    if not case_roots or not os.path.isdir(case_roots[0]):
        return {"ok": False, "error": "no_case_root"}

    case_number = (payload.get("case_number") or payload.get("case_no") or payload.get("caseNumber") or "").strip()
    client_name = (payload.get("client_name") or payload.get("name") or payload.get("client") or "").strip()
    case_type = (payload.get("case_type") or payload.get("type") or "").strip()
    case_stage = (payload.get("case_stage") or "").strip()
    case_reason = (payload.get("case_reason") or "").strip()

    if not case_number or not client_name:
        return {"ok": False, "error": "missing_case_number_or_client_name"}

    full_path = build_full_case_path(
        case_roots[0], case_number, client_name,
        case_type=case_type, case_category=case_category or "一般案件",
        case_stage=case_stage, case_reason=case_reason,
    )
    result = create_folder_structure(full_path, case_category or "一般案件")
    if not result.get("ok"):
        return result

    canonical = translate_local_path_to_canonical(full_path)
    try:
        _osc_exec("UPDATE cases SET folder_path=%s, updated_at=NOW() WHERE id=%s", (canonical, row_id), fetch="none")
    except Exception as e:
        return {"ok": True, "path": full_path, "canonical": canonical, "db_update_error": str(e)}
    return {"ok": True, "path": full_path, "canonical": canonical, "subfolders": result.get("subfolders", [])}


@app.route("/api/osc/cases", methods=["GET", "POST"])
@login_required
def osc_cases_api():
    if request.method == "GET":
        q = (request.args.get("q") or "").strip()
        category = (request.args.get("category") or "").strip()
        limit = max(1, min(500, int(request.args.get("limit") or "200")))
        where = []
        params = []
        if q:
            like = f"%{q}%"
            where.append(
                """
                (
                    case_number LIKE %s
                    OR client_name LIKE %s
                    OR court_case_no LIKE %s
                    OR laf_case_no LIKE %s
                    OR application_no LIKE %s
                )
                """
            )
            params.extend([like, like, like, like, like])
        if category and category not in {"全部", "all", "ALL"}:
            if category == "消費者債務清理":
                where.append("(case_category = %s OR case_type = %s)")
                params.extend([category, category])
            else:
                where.append("case_category = %s")
                params.append(category)
        sql = """
            SELECT id, case_number, client_name, case_category, case_type, case_stage, case_reason,
                   laf_case_no, application_no, court_case_no, status, notes, updated_at, created_date
            FROM cases
        """
        if where:
            sql += " WHERE " + " AND ".join(where)
        sql += " ORDER BY updated_at DESC, created_date DESC LIMIT %s"
        params.append(limit)
        rows, _ = _osc_exec(sql, tuple(params), fetch="all")
        return jsonify({"ok": True, "items": rows})

    payload = request.get_json() or {}
    row_id = (payload.get("id") or f"web-{uuid.uuid4().hex[:12]}").strip()
    case_number = (
        payload.get("case_number")
        or payload.get("case_no")
        or payload.get("caseNumber")
        or ""
    ).strip()
    client_name = (
        payload.get("client_name")
        or payload.get("name")
        or payload.get("client")
        or ""
    ).strip()
    if not client_name:
        return jsonify({"ok": False, "error": "client_name required"}), 400
    case_category = _osc_norm_case_category(payload.get("case_category") or payload.get("category") or "")
    cols = [
        "id", "case_number", "client_name", "client_phone", "client_email", "client_id_number",
        "case_category", "case_type", "case_stage", "case_reason",
        "laf_case_no", "application_no", "court_case_no", "status", "notes", "folder_path"
    ]
    vals = [
        row_id,
        case_number or None,
        client_name,
        (payload.get("client_phone") or "").strip() or None,
        (payload.get("client_email") or "").strip() or None,
        (payload.get("client_id_number") or "").strip() or None,
        case_category or None,
        (payload.get("case_type") or payload.get("type") or "").strip() or None,
        (payload.get("case_stage") or "").strip() or None,
        (payload.get("case_reason") or "").strip() or None,
        (payload.get("laf_case_no") or payload.get("legal_aid_number") or "").strip() or None,
        (payload.get("application_no") or "").strip() or None,
        (payload.get("court_case_no") or payload.get("court_case_number") or "").strip() or None,
        (payload.get("status") or "Active").strip() or "Active",
        (payload.get("notes") or "").strip() or None,
        translate_local_path_to_canonical((payload.get("folder_path") or "").strip()) or None,
    ]
    auto_create_folder = str(payload.get("auto_create_folder") or "").strip().lower() in {"1", "true", "yes", "on"}
    sql_insert = f"INSERT INTO cases ({','.join(cols)}) VALUES ({','.join(['%s'] * len(cols))})"
    try:
        result, _ = _osc_exec(sql_insert, tuple(vals), fetch="none")
        resp = {"ok": True, "result": result, "id": row_id, "mode": "insert"}
        if auto_create_folder:
            folder_resp = _osc_auto_create_folder_for_case(row_id, payload, case_category)
            resp["folder"] = folder_resp
        return jsonify(resp)
    except Exception as e:
        # 容錯：同案號/同 id 重複時改為更新，不讓前端看起來像「無法新增」。
        msg = str(e)
        is_dup = ("1062" in msg) or ("Duplicate entry" in msg)
        if not is_dup:
            return jsonify({"ok": False, "error": msg}), 500

        target = None
        if case_number:
            target, _ = _osc_exec("SELECT id FROM cases WHERE case_number=%s LIMIT 1", (case_number,), fetch="one")
        if not target and row_id:
            target, _ = _osc_exec("SELECT id FROM cases WHERE id=%s LIMIT 1", (row_id,), fetch="one")
        if not target:
            return jsonify({"ok": False, "error": msg}), 500

        update_payload = {
            "client_name": client_name,
            "case_category": case_category or None,
            "case_type": (payload.get("case_type") or payload.get("type") or "").strip() or None,
            "case_stage": (payload.get("case_stage") or "").strip() or None,
            "case_reason": (payload.get("case_reason") or "").strip() or None,
            "laf_case_no": (payload.get("laf_case_no") or payload.get("legal_aid_number") or "").strip() or None,
            "application_no": (payload.get("application_no") or "").strip() or None,
            "court_case_no": (payload.get("court_case_no") or payload.get("court_case_number") or "").strip() or None,
            "status": (payload.get("status") or "Active").strip() or "Active",
            "notes": (payload.get("notes") or "").strip() or None,
            "folder_path": translate_local_path_to_canonical((payload.get("folder_path") or "").strip()) or None,
        }
        if case_number:
            update_payload["case_number"] = case_number
        sets = []
        vals2 = []
        for k, v in update_payload.items():
            sets.append(f"{k}=%s")
            vals2.append(v)
        sets.append("updated_at=NOW()")
        vals2.append(target.get("id"))
        result, _ = _osc_exec(f"UPDATE cases SET {','.join(sets)} WHERE id=%s", tuple(vals2), fetch="none")
        return jsonify({"ok": True, "result": result, "id": target.get("id"), "mode": "upsert"})


@app.route("/api/osc/cases/<row_id>", methods=["GET", "PUT", "DELETE"])
@login_required
def osc_case_detail_api(row_id):
    row_id = (row_id or "").strip()
    if not row_id:
        return jsonify({"ok": False, "error": "invalid id"}), 400
    if request.method == "GET":
        row, _ = _osc_exec("SELECT * FROM cases WHERE id=%s", (row_id,), fetch="one")
        if not row:
            return jsonify({"ok": False, "error": "not found"}), 404
        return jsonify({"ok": True, "item": row})
    if request.method == "DELETE":
        result, _ = _osc_exec("DELETE FROM cases WHERE id=%s", (row_id,), fetch="none")
        return jsonify({"ok": True, "result": result})
    payload = request.get_json() or {}
    allowed = [
        "case_number", "client_name", "client_name_en", "client_phone", "client_email", "client_id_number",
        "case_category", "case_type", "case_stage", "case_reason",
        "laf_case_no", "application_no", "court_case_no", "status", "notes", "folder_path",
        "legal_aid_status", "court_case_number", "court_name",
    ]
    sets = []
    vals = []
    for k in allowed:
        if k in payload:
            sets.append(f"{k}=%s")
            v = (payload.get(k) or "").strip() or None
            if k == "case_category":
                v = _osc_norm_case_category(v or "")
            if k == "court_case_no" and not v:
                v = (payload.get("court_case_number") or "").strip() or None
            if k == "laf_case_no" and not v:
                v = (payload.get("legal_aid_number") or "").strip() or None
            if k == "folder_path" and v:
                v = translate_local_path_to_canonical(v) or v
            vals.append(v)
    if not sets:
        return jsonify({"ok": False, "error": "no fields"}), 400
    sets.append("updated_at=NOW()")
    vals.append(row_id)
    result, _ = _osc_exec(f"UPDATE cases SET {','.join(sets)} WHERE id=%s", tuple(vals), fetch="none")
    return jsonify({"ok": True, "result": result})


@app.route("/api/osc/cases/<row_id>/open-folder", methods=["POST"])
@login_required
def osc_case_open_folder_api(row_id):
    row_id = (row_id or "").strip()
    row, _ = _osc_exec("SELECT id, case_number, client_name, folder_path FROM cases WHERE id=%s", (row_id,), fetch="one")
    if not row:
        return jsonify({"ok": False, "error": "case_not_found"}), 404
    folder_path = (row.get("folder_path") or "").strip()
    if not folder_path:
        folder_path = _osc_guess_case_folder(row.get("case_number") or "")
    if not folder_path:
        return jsonify({"ok": False, "error": "folder_path_empty"}), 400
    norm = _osc_norm_path(folder_path)
    smb_candidates = _osc_smb_candidates(norm)
    smb = smb_candidates[0] if smb_candidates else ""
    local_candidates = _osc_local_path_candidates(norm)

    chosen_open_path = ""
    open_result = {"ok": False, "error": "open_failed"}

    # prefer local synced path first (more stable than direct SMB)
    for lp in local_candidates:
        try:
            if lp and os.path.exists(lp):
                r = _osc_try_open_path(lp)
                chosen_open_path = lp
                open_result = r
                if r.get("ok"):
                    break
        except Exception:
            continue

    # fallback to SMB candidates
    if not open_result.get("ok"):
        for sp in smb_candidates:
            r = _osc_try_open_path(sp)
            chosen_open_path = sp
            open_result = r
            if r.get("ok"):
                break

    return jsonify(
        {
            "ok": True,
            "case": {"id": row.get("id"), "case_number": row.get("case_number"), "client_name": row.get("client_name")},
            "folder_path": norm,
            "smb_url": smb,
            "smb_candidates": smb_candidates,
            "local_candidates": local_candidates,
            "chosen_open_path": chosen_open_path,
            "open_result": open_result,
            "browser_supported": True,
            "browser_url": f"/api/osc/cases/{row_id}/folder-browser",
        }
    )


@app.route("/api/osc/cases/<row_id>/create-folder", methods=["POST"])
@login_required
def osc_case_create_folder_api(row_id):
    """建立案件資料夾結構並更新 DB folder_path。"""
    from casper_ecosystem.law_firm_orchestrators.osc.folder_utils import (
        build_full_case_path,
        create_folder_structure,
    )

    row_id = (row_id or "").strip()
    row, _ = _osc_exec(
        "SELECT id, case_number, client_name, case_category, case_type, case_stage, case_reason, folder_path FROM cases WHERE id=%s",
        (row_id,),
        fetch="one",
    )
    if not row:
        return jsonify({"ok": False, "error": "case_not_found"}), 404

    # 決定基礎路徑（使用 preferred_case_roots）
    case_roots = preferred_case_roots()
    if not case_roots:
        return jsonify({"ok": False, "error": "no_case_root_configured"}), 500
    base_path = case_roots[0]
    if not os.path.isdir(base_path):
        return jsonify({"ok": False, "error": f"base_path_not_found: {base_path}"}), 500

    case_number = row.get("case_number") or ""
    client_name = row.get("client_name") or ""
    case_category = row.get("case_category") or "一般案件"
    case_type = row.get("case_type") or ""
    case_stage = row.get("case_stage") or ""
    case_reason = row.get("case_reason") or ""

    if not case_number or not client_name:
        return jsonify({"ok": False, "error": "case_number and client_name are required"}), 400

    full_path = build_full_case_path(
        base_path, case_number, client_name,
        case_type=case_type, case_category=case_category,
        case_stage=case_stage, case_reason=case_reason,
    )

    result = create_folder_structure(full_path, case_category)
    if not result.get("ok"):
        return jsonify(result), 500

    # 更新 DB folder_path（存 canonical Windows 路徑）
    canonical = translate_local_path_to_canonical(full_path)
    _osc_exec("UPDATE cases SET folder_path=%s, updated_at=NOW() WHERE id=%s", (canonical, row_id), fetch="none")

    return jsonify({
        "ok": True,
        "folder_path": full_path,
        "canonical_path": canonical,
        "subfolders": result.get("subfolders", []),
    })


@app.route("/api/osc/cases/<row_id>/folder-browser", methods=["GET"])
@login_required
def osc_case_folder_browser_api(row_id):
    row_id = (row_id or "").strip()
    row, _ = _osc_exec("SELECT id, case_number, client_name, folder_path FROM cases WHERE id=%s", (row_id,), fetch="one")
    if not row:
        return jsonify({"ok": False, "error": "case_not_found"}), 404
    folder_path = (row.get("folder_path") or "").strip()
    if not folder_path:
        folder_path = _osc_guess_case_folder(row.get("case_number") or "")
    if not folder_path:
        return jsonify({"ok": False, "error": "folder_path_empty"}), 400
    norm = _osc_norm_path(folder_path)
    smb_candidates = _osc_smb_candidates(norm)
    local_candidates = _osc_local_path_candidates(norm)
    local_folder = _osc_resolve_existing_local_path(norm, prefer_dir=True)
    rel = (request.args.get("path") or "").strip().strip("/")
    payload = {
        "ok": True,
        "case": {"id": row.get("id"), "case_number": row.get("case_number"), "client_name": row.get("client_name")},
        "folder_path": norm,
        "local_candidates": local_candidates,
        "smb_candidates": smb_candidates,
        "local_folder": local_folder,
        "folder_exists": bool(local_folder),
    }
    if not local_folder:
        payload["entries"] = []
        payload["current_relative_path"] = ""
        payload["parent_relative_path"] = ""
        payload["error"] = "folder_not_synced"
        return jsonify(payload)
    listing = _osc_folder_entries(local_folder, rel)
    if not listing.get("ok"):
        return jsonify({**payload, **listing}), 400
    payload.update(listing)
    return jsonify(payload)


@app.route("/api/osc/cases/<row_id>/quick-action", methods=["POST"])
@login_required
def osc_case_quick_action_api(row_id):
    row_id = (row_id or "").strip()
    if not row_id:
        return jsonify({"ok": False, "error": "invalid_id"}), 400
    case, _ = _osc_exec(
        """
        SELECT id, case_number, client_name, case_category, case_reason, case_stage, court_case_no, laf_case_no
        FROM cases
        WHERE id=%s
        """,
        (row_id,),
        fetch="one",
    )
    if not case:
        return jsonify({"ok": False, "error": "case_not_found"}), 404
    payload = request.get_json() or {}
    action = (payload.get("action") or "").strip()
    action_map = {
        "generate_power_of_attorney": "請針對此案件產生委任狀草稿，並列出欄位缺漏供人工確認。",
        "generate_receipt": "請針對此案件產生收據草稿，並列出必填欄位。",
        "closing_overview": "請彙整此案件結案回報需要的進度、文件與風險缺漏，輸出待辦清單。",
        "laf_progress_summary": "請整理此案件目前法扶進度、補件狀態與卡點，輸出下一步建議。",
        "laf_closing_status": "請整理此案件結案狀況（已完成/待補/風險），並列出缺漏文件。",
    }
    if action not in action_map:
        return jsonify({"ok": False, "error": "unsupported_action"}), 400
    prompt = (
        f"{action_map[action]}\n\n"
        f"案件編號: {case.get('case_number') or ''}\n"
        f"當事人: {case.get('client_name') or ''}\n"
        f"案件種類: {case.get('case_category') or ''}\n"
        f"案由: {case.get('case_reason') or ''}\n"
        f"審級/階段: {case.get('case_stage') or ''}\n"
        f"法院案號: {case.get('court_case_no') or ''}\n"
        f"法扶案號: {case.get('laf_case_no') or ''}\n"
    )
    try:
        reply = orchestrator.process_message(
            user_id=str(current_user.id),
            message=prompt,
            platform="WEB",
            role=current_user.role,
        )
        if _normalize_output_text:
            reply = _normalize_output_text(str(reply or ""), platform="WEB")
        return jsonify({"ok": True, "action": action, "case": case, "reply": str(reply or "")})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/api/osc/clients/<row_id>/workbench", methods=["GET"])
@login_required
def osc_client_workbench_api(row_id):
    row_id = (row_id or "").strip()
    client, _ = _osc_exec("SELECT * FROM clients WHERE id=%s", (row_id,), fetch="one")
    if not client:
        return jsonify({"ok": False, "error": "client_not_found"}), 404
    name = (client.get("name") or "").strip()
    like = f"%{name}%"
    cases, _ = _osc_exec(
        """
        SELECT id, case_number, client_name, case_category, case_type, case_stage, case_reason, status, folder_path,
               laf_case_no, application_no, court_case_no, legal_aid_status, updated_at
        FROM cases
        WHERE client_name LIKE %s
        ORDER BY updated_at DESC, created_date DESC
        LIMIT 200
        """,
        (like,),
        fetch="all",
    )
    case_numbers = [str(c.get("case_number") or "").strip() for c in cases if (c.get("case_number") or "").strip()]
    todos = []
    meetings = []
    legal_aid_checklist = []
    case_checklist = []
    lifecycle = []
    opponents = []
    pdf_generation_log = []
    if case_numbers:
        ph = ",".join(["%s"] * len(case_numbers))
        todos, _ = _osc_exec(
            f"""
            SELECT id, case_number, client_name, todo_type, todo_date, todo_time, description, status, source_file, created_date
            FROM case_todos
            WHERE case_number IN ({ph})
            ORDER BY todo_date DESC, id DESC
            LIMIT 500
            """,
            tuple(case_numbers),
            fetch="all",
        )
        meetings, _ = _osc_exec(
            f"""
            SELECT id, case_number, client_name, type, datetime, duration, location, notes, status, reminder, reminder_time
            FROM meetings
            WHERE case_number IN ({ph})
            ORDER BY datetime DESC, id DESC
            LIMIT 500
            """,
            tuple(case_numbers),
            fetch="all",
        )
        legal_aid_checklist, _ = _osc_exec(
            f"""
            SELECT id, case_number, item_key, item_label, status, notes, last_updated
            FROM legal_aid_checklists
            WHERE case_number IN ({ph})
            ORDER BY last_updated DESC, id DESC
            LIMIT 500
            """,
            tuple(case_numbers),
            fetch="all",
        )
        case_checklist, _ = _osc_exec(
            f"""
            SELECT id, case_number, item_label, status, notes, is_active
            FROM case_checklists
            WHERE case_number IN ({ph})
            ORDER BY id DESC
            LIMIT 500
            """,
            tuple(case_numbers),
            fetch="all",
        )
        lifecycle, _ = _osc_exec(
            f"""
            SELECT id, case_number, event_type, status, created_at, completed_at, event_data
            FROM laf_lifecycle_log
            WHERE case_number IN ({ph})
            ORDER BY created_at DESC, id DESC
            LIMIT 500
            """,
            tuple(case_numbers),
            fetch="all",
        )
        opponents, _ = _osc_exec(
            f"""
            SELECT id, case_number, name, address, created_date, updated_date, is_active
            FROM opponents
            WHERE case_number IN ({ph})
            ORDER BY updated_date DESC, id DESC
            LIMIT 500
            """,
            tuple(case_numbers),
            fetch="all",
        )
        pdf_generation_log, _ = _osc_exec(
            f"""
            SELECT id, case_number, file_name, log_timestamp, status, error_message
            FROM pdf_generation_log
            WHERE case_number IN ({ph})
            ORDER BY log_timestamp DESC, id DESC
            LIMIT 500
            """,
            tuple(case_numbers),
            fetch="all",
        )
    return jsonify(
        {
            "ok": True,
            "client": client,
            "cases": cases,
            "todos": todos,
            "meetings": meetings,
            "legal_aid_checklist": legal_aid_checklist,
            "case_checklist": case_checklist,
            "laf_progress": lifecycle,
            "opponents": opponents,
            "pdf_generation_log": pdf_generation_log,
        }
    )


@app.route("/api/osc/cases/<row_id>/workbench", methods=["GET"])
@login_required
def osc_case_workbench_api(row_id):
    case, _ = _osc_exec("SELECT * FROM cases WHERE id=%s", ((row_id or "").strip(),), fetch="one")
    if not case:
        return jsonify({"ok": False, "error": "case_not_found"}), 404
    case_number = (case.get("case_number") or "").strip()
    todos, _ = _osc_exec(
        """
        SELECT id, case_number, client_name, todo_type, todo_date, todo_time, description, status, source_file, created_date, completed_date
        FROM case_todos WHERE case_number=%s ORDER BY todo_date DESC, id DESC LIMIT 800
        """,
        (case_number,),
        fetch="all",
    )
    meetings, _ = _osc_exec(
        """
        SELECT id, case_number, client_name, type, datetime, duration, location, notes, status
        FROM meetings WHERE case_number=%s ORDER BY datetime DESC, id DESC LIMIT 800
        """,
        (case_number,),
        fetch="all",
    )
    legal_aid, _ = _osc_exec(
        """
        SELECT id, case_number, item_key, item_label, status, notes, last_updated
        FROM legal_aid_checklists WHERE case_number=%s ORDER BY last_updated DESC, id DESC LIMIT 1000
        """,
        (case_number,),
        fetch="all",
    )
    lifecycle, _ = _osc_exec(
        """
        SELECT id, case_number, event_type, status, created_at, completed_at, event_data
        FROM laf_lifecycle_log WHERE case_number=%s ORDER BY created_at DESC, id DESC LIMIT 1000
        """,
        (case_number,),
        fetch="all",
    )
    docs, _ = _osc_exec(
        """
        SELECT id, case_number, file_name, file_path, subfolder_name, party, reason, modified_date
        FROM document_index WHERE case_number=%s ORDER BY modified_date DESC, id DESC LIMIT 1000
        """,
        (case_number,),
        fetch="all",
    )
    opponents, _ = _osc_exec(
        """
        SELECT id, case_number, name, address, created_date, updated_date, is_active
        FROM opponents WHERE case_number=%s ORDER BY updated_date DESC, id DESC LIMIT 300
        """,
        (case_number,),
        fetch="all",
    )
    pdf_generation_log, _ = _osc_exec(
        """
        SELECT id, case_number, file_name, log_timestamp, status, error_message
        FROM pdf_generation_log WHERE case_number=%s ORDER BY log_timestamp DESC, id DESC LIMIT 300
        """,
        (case_number,),
        fetch="all",
    )
    stats = {
        "todo_total": len(todos),
        "todo_pending": len([t for t in todos if str(t.get("status") or "").lower() not in {"completed", "done", "已完成"}]),
        "todo_completed": len([t for t in todos if str(t.get("status") or "").lower() in {"completed", "done", "已完成"}]),
        "meeting_total": len(meetings),
        "laf_items": len(legal_aid),
        "docs_indexed": len(docs),
        "opponents_total": len(opponents),
        "pdf_logs_total": len(pdf_generation_log),
    }
    return jsonify(
        {
            "ok": True,
            "case": case,
            "stats": stats,
            "todos": todos,
            "meetings": meetings,
            "legal_aid_checklist": legal_aid,
            "laf_progress": lifecycle,
            "documents": docs,
            "opponents": opponents,
            "pdf_generation_log": pdf_generation_log,
        }
    )


@app.route("/api/osc/dashboard", methods=["GET"])
@login_required
def osc_dashboard_api():
    start_date, end_date = _osc_accounting_window()
    active_cases_row, _ = _osc_exec(
        """
        SELECT COUNT(*) AS c FROM cases
        WHERE status NOT IN ('已結案', '已結案，待報結') OR status IS NULL OR status=''
        """,
        fetch="one",
    )
    legal_aid_cases_row, _ = _osc_exec(
        """
        SELECT COUNT(*) AS c FROM cases
        WHERE (case_category='法律扶助案件' OR case_reason LIKE '%法扶%' OR case_reason LIKE '%法律扶助%')
          AND (status NOT IN ('已結案', '已結案，待報結') OR status IS NULL OR status='')
        """,
        fetch="one",
    )
    monthly_revenue_row, _ = _osc_exec(
        "SELECT COALESCE(SUM(amount),0) AS total FROM case_transactions WHERE date >= %s AND date <= %s AND type='收入'",
        (start_date, end_date),
        fetch="one",
    )
    monthly_expense_row, _ = _osc_exec(
        "SELECT COALESCE(SUM(amount),0) AS total FROM case_transactions WHERE date >= %s AND date <= %s AND type='支出'",
        (start_date, end_date),
        fetch="one",
    )
    closed_regular_row, _ = _osc_exec(
        """
        SELECT COUNT(*) AS c FROM cases
        WHERE status IN ('已結案', '已結案，待報結')
          AND NOT (case_category='法律扶助案件' OR case_reason LIKE '%法扶%' OR case_reason LIKE '%法律扶助%')
        """,
        fetch="one",
    )
    closed_laf_row, _ = _osc_exec(
        """
        SELECT COUNT(*) AS c FROM cases
        WHERE status IN ('已結案', '已結案，待報結')
          AND (case_category='法律扶助案件' OR case_reason LIKE '%法扶%' OR case_reason LIKE '%法律扶助%')
        """,
        fetch="one",
    )
    recent_cases, _ = _osc_exec(
        """
        SELECT id, case_number, client_name, case_category, case_type, case_stage, case_reason, status, updated_at, created_date
        FROM cases
        ORDER BY updated_at DESC, created_date DESC
        LIMIT 12
        """,
        fetch="all",
    )
    pending_todos, _ = _osc_exec(
        """
        SELECT id, case_number, client_name, todo_type, todo_date, todo_time, description, status
        FROM case_todos
        WHERE status IS NULL OR status='' OR LOWER(status) NOT IN ('completed', 'done')
        ORDER BY COALESCE(todo_date, CURDATE()) ASC, id DESC
        LIMIT 20
        """,
        fetch="all",
    )
    upcoming_calendar, _ = _osc_exec(
        """
        SELECT id, case_number, title, start_date, end_date, description, location, color, is_all_day
        FROM calendar_events
        WHERE start_date >= %s
        ORDER BY start_date ASC, id ASC
        LIMIT 20
        """,
        (date.today(),),
        fetch="all",
    )
    recent_activity, _ = _osc_exec(
        """
        SELECT id, action, entity_type, entity_id, details, user, timestamp
        FROM activity_logs
        ORDER BY timestamp DESC, id DESC
        LIMIT 20
        """,
        fetch="all",
    )
    recent_pdf_logs, _ = _osc_exec(
        """
        SELECT id, case_number, file_name, log_timestamp, status, error_message
        FROM pdf_generation_log
        ORDER BY log_timestamp DESC, id DESC
        LIMIT 20
        """,
        fetch="all",
    )
    return jsonify(
        {
            "ok": True,
            "window": {"start_date": str(start_date), "end_date": str(end_date)},
            "stats": {
                "active_cases": int((active_cases_row or {}).get("c") or 0),
                "legal_aid_cases": int((legal_aid_cases_row or {}).get("c") or 0),
                "monthly_revenue": float((monthly_revenue_row or {}).get("total") or 0),
                "monthly_expense": float((monthly_expense_row or {}).get("total") or 0),
                "closed_regular": int((closed_regular_row or {}).get("c") or 0),
                "closed_legal_aid": int((closed_laf_row or {}).get("c") or 0),
            },
            "recent_cases": recent_cases or [],
            "pending_todos": pending_todos or [],
            "upcoming_calendar": upcoming_calendar or [],
            "recent_activity": recent_activity or [],
            "recent_pdf_logs": recent_pdf_logs or [],
        }
    )


@app.route("/api/osc/case-reason-templates", methods=["GET", "POST"])
@login_required
def osc_case_reason_templates_api():
    if request.method == "GET":
        q = (request.args.get("q") or "").strip()
        case_type = (request.args.get("case_type") or "").strip()
        common_only = _osc_truthy(request.args.get("common_only"))
        limit = max(1, min(2000, int(request.args.get("limit") or "300")))
        where = ["1=1"]
        params = []
        if case_type:
            where.append("case_type=%s")
            params.append(case_type)
        if common_only:
            where.append("is_common=1")
        if q:
            like = f"%{q}%"
            where.append("(case_type LIKE %s OR reason LIKE %s)")
            params.extend([like, like])
        params.append(limit)
        rows, _ = _osc_exec(
            f"""
            SELECT id, case_type, reason, is_common, created_date
            FROM case_reason_templates
            WHERE {' AND '.join(where)}
            ORDER BY is_common DESC, case_type ASC, id DESC
            LIMIT %s
            """,
            tuple(params),
            fetch="all",
        )
        return jsonify({"ok": True, "items": rows or []})
    payload = request.get_json() or {}
    case_type = _osc_text(payload.get("case_type"))
    reason = _osc_text(payload.get("reason"))
    if not case_type or not reason:
        return jsonify({"ok": False, "error": "case_type/reason required"}), 400
    is_common = 1 if _osc_truthy(payload.get("is_common")) else 0
    result, _ = _osc_exec(
        """
        INSERT INTO case_reason_templates (case_type, reason, is_common)
        VALUES (%s,%s,%s)
        ON DUPLICATE KEY UPDATE is_common=VALUES(is_common)
        """,
        (case_type, reason, is_common),
        fetch="none",
    )
    _osc_log_activity("case_reason_template:save", "case_reason_templates", f"{case_type}:{reason}", payload)
    return jsonify({"ok": True, "result": result})


@app.route("/api/osc/case-reason-templates/<int:row_id>", methods=["GET", "PUT", "DELETE"])
@login_required
def osc_case_reason_template_detail_api(row_id):
    if request.method == "GET":
        row, _ = _osc_exec("SELECT * FROM case_reason_templates WHERE id=%s", (row_id,), fetch="one")
        if not row:
            return jsonify({"ok": False, "error": "not found"}), 404
        return jsonify({"ok": True, "item": row})
    if request.method == "DELETE":
        result, _ = _osc_exec("DELETE FROM case_reason_templates WHERE id=%s", (row_id,), fetch="none")
        _osc_log_activity("case_reason_template:delete", "case_reason_templates", str(row_id))
        return jsonify({"ok": True, "result": result})
    payload = request.get_json() or {}
    sets, vals = [], []
    for key in ["case_type", "reason", "is_common"]:
        if key not in payload:
            continue
        sets.append(f"{key}=%s")
        if key == "is_common":
            vals.append(1 if _osc_truthy(payload.get(key)) else 0)
        else:
            vals.append(_osc_text(payload.get(key)))
    if not sets:
        return jsonify({"ok": False, "error": "no fields"}), 400
    vals.append(row_id)
    result, _ = _osc_exec(f"UPDATE case_reason_templates SET {','.join(sets)} WHERE id=%s", tuple(vals), fetch="none")
    _osc_log_activity("case_reason_template:update", "case_reason_templates", str(row_id), payload)
    return jsonify({"ok": True, "result": result})


## settings routes → moved to api.blueprints.osc_settings


@app.route("/api/osc/activity-logs", methods=["GET", "POST"])
@login_required
def osc_activity_logs_api():
    if request.method == "GET":
        q = (request.args.get("q") or "").strip()
        entity_type = (request.args.get("entity_type") or "").strip()
        user_name = (request.args.get("user") or "").strip()
        limit = max(1, min(2000, int(request.args.get("limit") or "300")))
        sql = "SELECT id, action, entity_type, entity_id, details, user, timestamp FROM activity_logs WHERE 1=1 "
        params = []
        if entity_type:
            sql += "AND entity_type=%s "
            params.append(entity_type)
        if user_name:
            sql += "AND user=%s "
            params.append(user_name)
        if q:
            like = f"%{q}%"
            sql += "AND (action LIKE %s OR entity_type LIKE %s OR entity_id LIKE %s OR details LIKE %s OR user LIKE %s) "
            params.extend([like, like, like, like, like])
        sql += "ORDER BY timestamp DESC, id DESC LIMIT %s"
        params.append(limit)
        rows, _ = _osc_exec(sql, tuple(params), fetch="all")
        return jsonify({"ok": True, "items": rows or []})
    payload = request.get_json() or {}
    action = _osc_text(payload.get("action"))
    if not action:
        return jsonify({"ok": False, "error": "action required"}), 400
    result, _ = _osc_exec(
        "INSERT INTO activity_logs (action, entity_type, entity_id, details, user) VALUES (%s,%s,%s,%s,%s)",
        (
            action,
            _osc_text(payload.get("entity_type")),
            _osc_text(payload.get("entity_id")),
            _osc_text(payload.get("details")),
            _osc_text(payload.get("user")) or _osc_current_actor(),
        ),
        fetch="none",
    )
    return jsonify({"ok": True, "result": result})


@app.route("/api/osc/activity-logs/<int:row_id>", methods=["GET", "DELETE"])
@login_required
def osc_activity_log_detail_api(row_id):
    if request.method == "GET":
        row, _ = _osc_exec("SELECT * FROM activity_logs WHERE id=%s", (row_id,), fetch="one")
        if not row:
            return jsonify({"ok": False, "error": "not found"}), 404
        return jsonify({"ok": True, "item": row})
    result, _ = _osc_exec("DELETE FROM activity_logs WHERE id=%s", (row_id,), fetch="none")
    return jsonify({"ok": True, "result": result})


@app.route("/api/osc/user-settings", methods=["GET", "POST"])
@login_required
def osc_user_settings_api():
    if request.method == "GET":
        q = (request.args.get("q") or "").strip()
        hostname = (request.args.get("hostname") or "").strip()
        limit = max(1, min(2000, int(request.args.get("limit") or "300")))
        sql = "SELECT id, hostname, setting_key, setting_value, last_updated FROM user_settings WHERE 1=1 "
        params = []
        if hostname:
            sql += "AND hostname=%s "
            params.append(hostname)
        if q:
            like = f"%{q}%"
            sql += "AND (hostname LIKE %s OR setting_key LIKE %s OR setting_value LIKE %s) "
            params.extend([like, like, like])
        sql += "ORDER BY hostname ASC, setting_key ASC, id DESC LIMIT %s"
        params.append(limit)
        rows, _ = _osc_exec(sql, tuple(params), fetch="all")
        return jsonify({"ok": True, "items": rows or []})
    payload = request.get_json() or {}
    hostname = _osc_text(payload.get("hostname"))
    setting_key = _osc_text(payload.get("setting_key"))
    if not hostname or not setting_key:
        return jsonify({"ok": False, "error": "hostname/setting_key required"}), 400
    result, _ = _osc_exec(
        """
        INSERT INTO user_settings (hostname, setting_key, setting_value)
        VALUES (%s,%s,%s)
        ON DUPLICATE KEY UPDATE setting_value=VALUES(setting_value)
        """,
        (hostname, setting_key, _osc_text(payload.get("setting_value"))),
        fetch="none",
    )
    _osc_log_activity("user_setting:save", "user_settings", f"{hostname}:{setting_key}", payload)
    return jsonify({"ok": True, "result": result})


@app.route("/api/osc/user-settings/<int:row_id>", methods=["GET", "PUT", "DELETE"])
@login_required
def osc_user_setting_detail_api(row_id):
    if request.method == "GET":
        row, _ = _osc_exec("SELECT * FROM user_settings WHERE id=%s", (row_id,), fetch="one")
        if not row:
            return jsonify({"ok": False, "error": "not found"}), 404
        return jsonify({"ok": True, "item": row})
    if request.method == "DELETE":
        result, _ = _osc_exec("DELETE FROM user_settings WHERE id=%s", (row_id,), fetch="none")
        _osc_log_activity("user_setting:delete", "user_settings", str(row_id))
        return jsonify({"ok": True, "result": result})
    payload = request.get_json() or {}
    sets, vals = [], []
    for key in ["hostname", "setting_key", "setting_value"]:
        if key not in payload:
            continue
        sets.append(f"{key}=%s")
        vals.append(_osc_text(payload.get(key)))
    if not sets:
        return jsonify({"ok": False, "error": "no fields"}), 400
    vals.append(row_id)
    result, _ = _osc_exec(f"UPDATE user_settings SET {','.join(sets)} WHERE id=%s", tuple(vals), fetch="none")
    _osc_log_activity("user_setting:update", "user_settings", str(row_id), payload)
    return jsonify({"ok": True, "result": result})


@app.route("/api/osc/memory-keywords", methods=["GET", "POST"])
@login_required
def osc_memory_keywords_api():
    if request.method == "GET":
        q = (request.args.get("q") or "").strip()
        case_number = (request.args.get("case_number") or "").strip()
        limit = max(1, min(2000, int(request.args.get("limit") or "300")))
        sql = "SELECT case_number, hotkey, name, value FROM memory_keywords WHERE 1=1 "
        params = []
        if case_number:
            sql += "AND case_number=%s "
            params.append(case_number)
        if q:
            like = f"%{q}%"
            sql += "AND (case_number LIKE %s OR hotkey LIKE %s OR name LIKE %s OR value LIKE %s) "
            params.extend([like, like, like, like])
        sql += "ORDER BY case_number ASC, hotkey ASC LIMIT %s"
        params.append(limit)
        rows, _ = _osc_exec(sql, tuple(params), fetch="all")
        return jsonify({"ok": True, "items": rows or []})
    payload = request.get_json() or {}
    case_number = _osc_text(payload.get("case_number"))
    hotkey = _osc_text(payload.get("hotkey"))
    if not case_number or not hotkey:
        return jsonify({"ok": False, "error": "case_number/hotkey required"}), 400
    result, _ = _osc_exec(
        """
        INSERT INTO memory_keywords (case_number, hotkey, name, value)
        VALUES (%s,%s,%s,%s)
        ON DUPLICATE KEY UPDATE name=VALUES(name), value=VALUES(value)
        """,
        (case_number, hotkey, _osc_text(payload.get("name")), _osc_text(payload.get("value"))),
        fetch="none",
    )
    _osc_log_activity("memory_keyword:save", "memory_keywords", f"{case_number}:{hotkey}", payload)
    return jsonify({"ok": True, "result": result})


@app.route("/api/osc/memory-keywords/<path:case_number>/<path:hotkey>", methods=["GET", "PUT", "DELETE"])
@login_required
def osc_memory_keyword_detail_api(case_number, hotkey):
    if request.method == "GET":
        row, _ = _osc_exec(
            "SELECT case_number, hotkey, name, value FROM memory_keywords WHERE case_number=%s AND hotkey=%s",
            (case_number, hotkey),
            fetch="one",
        )
        if not row:
            return jsonify({"ok": False, "error": "not found"}), 404
        return jsonify({"ok": True, "item": row})
    if request.method == "DELETE":
        result, _ = _osc_exec("DELETE FROM memory_keywords WHERE case_number=%s AND hotkey=%s", (case_number, hotkey), fetch="none")
        _osc_log_activity("memory_keyword:delete", "memory_keywords", f"{case_number}:{hotkey}")
        return jsonify({"ok": True, "result": result})
    payload = request.get_json() or {}
    sets, vals = [], []
    for key in ["name", "value"]:
        if key not in payload:
            continue
        sets.append(f"{key}=%s")
        vals.append(_osc_text(payload.get(key)))
    if not sets:
        return jsonify({"ok": False, "error": "no fields"}), 400
    vals.extend([case_number, hotkey])
    result, _ = _osc_exec(
        f"UPDATE memory_keywords SET {','.join(sets)} WHERE case_number=%s AND hotkey=%s",
        tuple(vals),
        fetch="none",
    )
    _osc_log_activity("memory_keyword:update", "memory_keywords", f"{case_number}:{hotkey}", payload)
    return jsonify({"ok": True, "result": result})


## courts + legal-aid-branches routes → moved to api.blueprints.osc_settings


@app.route("/api/osc/opponents", methods=["GET", "POST"])
@login_required
def osc_opponents_api():
    if request.method == "GET":
        q = (request.args.get("q") or "").strip()
        case_number = (request.args.get("case_number") or "").strip()
        active_only = _osc_truthy(request.args.get("active_only"))
        limit = max(1, min(2000, int(request.args.get("limit") or "300")))
        sql = "SELECT id, case_number, name, address, created_date, updated_date, is_active FROM opponents WHERE 1=1 "
        params = []
        if case_number:
            sql += "AND case_number=%s "
            params.append(case_number)
        if active_only:
            sql += "AND is_active=1 "
        if q:
            like = f"%{q}%"
            sql += "AND (case_number LIKE %s OR name LIKE %s OR address LIKE %s) "
            params.extend([like, like, like])
        sql += "ORDER BY updated_date DESC, id DESC LIMIT %s"
        params.append(limit)
        rows, _ = _osc_exec(sql, tuple(params), fetch="all")
        return jsonify({"ok": True, "items": rows or []})
    payload = request.get_json() or {}
    case_number = _osc_text(payload.get("case_number"))
    name = _osc_text(payload.get("name"))
    if not case_number or not name:
        return jsonify({"ok": False, "error": "case_number/name required"}), 400
    result, _ = _osc_exec(
        """
        INSERT INTO opponents (case_number, name, address, is_active)
        VALUES (%s,%s,%s,%s)
        """,
        (case_number, name, _osc_text(payload.get("address")), 1 if _osc_truthy(payload.get("is_active", 1)) else 0),
        fetch="none",
    )
    _osc_log_activity("opponent:create", "opponents", case_number, payload)
    return jsonify({"ok": True, "result": result})


@app.route("/api/osc/opponents/<int:row_id>", methods=["GET", "PUT", "DELETE"])
@login_required
def osc_opponent_detail_api(row_id):
    if request.method == "GET":
        row, _ = _osc_exec("SELECT * FROM opponents WHERE id=%s", (row_id,), fetch="one")
        if not row:
            return jsonify({"ok": False, "error": "not found"}), 404
        return jsonify({"ok": True, "item": row})
    if request.method == "DELETE":
        result, _ = _osc_exec("DELETE FROM opponents WHERE id=%s", (row_id,), fetch="none")
        _osc_log_activity("opponent:delete", "opponents", str(row_id))
        return jsonify({"ok": True, "result": result})
    payload = request.get_json() or {}
    sets, vals = [], []
    for key in ["case_number", "name", "address", "is_active"]:
        if key not in payload:
            continue
        sets.append(f"{key}=%s")
        if key == "is_active":
            vals.append(1 if _osc_truthy(payload.get(key)) else 0)
        else:
            vals.append(_osc_text(payload.get(key)))
    if not sets:
        return jsonify({"ok": False, "error": "no fields"}), 400
    vals.append(row_id)
    result, _ = _osc_exec(f"UPDATE opponents SET {','.join(sets)} WHERE id=%s", tuple(vals), fetch="none")
    _osc_log_activity("opponent:update", "opponents", str(row_id), payload)
    return jsonify({"ok": True, "result": result})


@app.route("/api/osc/pdf-generation-log", methods=["GET"])
@login_required
def osc_pdf_generation_log_api():
    q = (request.args.get("q") or "").strip()
    case_number = (request.args.get("case_number") or "").strip()
    status = (request.args.get("status") or "").strip()
    limit = max(1, min(2000, int(request.args.get("limit") or "300")))
    sql = "SELECT id, case_number, file_name, log_timestamp, status, error_message FROM pdf_generation_log WHERE 1=1 "
    params = []
    if case_number:
        sql += "AND case_number=%s "
        params.append(case_number)
    if status:
        sql += "AND status=%s "
        params.append(status)
    if q:
        like = f"%{q}%"
        sql += "AND (case_number LIKE %s OR file_name LIKE %s OR status LIKE %s OR error_message LIKE %s) "
        params.extend([like, like, like, like])
    sql += "ORDER BY log_timestamp DESC, id DESC LIMIT %s"
    params.append(limit)
    rows, _ = _osc_exec(sql, tuple(params), fetch="all")
    return jsonify({"ok": True, "items": rows or []})


@app.route("/api/osc/pdf-generation-log/<int:row_id>", methods=["GET", "DELETE"])
@login_required
def osc_pdf_generation_log_detail_api(row_id):
    if request.method == "GET":
        row, _ = _osc_exec("SELECT * FROM pdf_generation_log WHERE id=%s", (row_id,), fetch="one")
        if not row:
            return jsonify({"ok": False, "error": "not found"}), 404
        return jsonify({"ok": True, "item": row})
    result, _ = _osc_exec("DELETE FROM pdf_generation_log WHERE id=%s", (row_id,), fetch="none")
    _osc_log_activity("pdf_log:delete", "pdf_generation_log", str(row_id))
    return jsonify({"ok": True, "result": result})


@app.route("/api/osc/drafts/meta", methods=["GET"])
@login_required
def osc_drafts_meta_api():
    provider = _osc_get_setting_value("ai_draft_provider", "casper") or "casper"
    model = _osc_get_setting_value("ai_draft_ollama_model", "taide-12b") or "taide-12b"
    ollama_url = _osc_get_setting_value("ollama_url", "http://127.0.0.1:8080") or "http://127.0.0.1:8080"
    custom_template = _osc_get_setting_value("draft_prompt_template", "")
    allow_cloud_models = str(os.environ.get("MAGI_ALLOW_CLOUD_MODELS", "0") or "").strip().lower() in {"1", "true", "yes", "on"}
    effective_provider = provider
    if provider == "gemini" and not allow_cloud_models:
        effective_provider = "casper"
    return jsonify(
        {
            "ok": True,
            "meta": {
                "enabled": _osc_draft_enabled_flag(),
                "provider": provider,
                "effective_provider": effective_provider,
                "ollama_model": model,
                "ollama_url": ollama_url,
                "allow_cloud_models": allow_cloud_models,
                "template_source": "custom" if custom_template.strip() else "default",
                "has_custom_template": bool(custom_template.strip()),
                "template_length": len(custom_template.strip() or _OSC_DRAFT_PROMPT_TEMPLATE),
            },
            "doc_types": _OSC_DRAFT_DOC_TYPES,
        }
    )


@app.route("/api/osc/drafts/generate", methods=["POST"])
@login_required
def osc_drafts_generate_api():
    payload = request.get_json() or {}
    ctx = _osc_build_draft_context(payload)
    doc_type = str(ctx.get("doc_type") or "").strip()
    case_facts = str(ctx.get("case_facts") or "").strip()
    if not doc_type:
        return jsonify({"ok": False, "error": "doc_type required"}), 400
    if not case_facts:
        return jsonify({"ok": False, "error": "case_facts required"}), 400

    prompt = str(ctx.get("prompt") or "").strip()
    provider = str(payload.get("provider") or _osc_get_setting_value("ai_draft_provider", "casper") or "casper").strip().lower()
    ollama_model = str(payload.get("ollama_model") or _osc_get_setting_value("ai_draft_ollama_model", "taide-12b") or "taide-12b").strip()
    ollama_url = str(payload.get("ollama_url") or _osc_get_setting_value("ollama_url", "http://127.0.0.1:8080") or "http://127.0.0.1:8080").strip()
    dry_run = _osc_truthy(payload.get("dry_run") or payload.get("preview_only"))

    if dry_run:
        return jsonify(
            {
                "ok": True,
                "dry_run": True,
                "provider": provider,
                "ollama_model": ollama_model,
                "prompt_preview": prompt,
                "warnings": ctx.get("warnings") or [],
                "case": ctx.get("case") or {},
                "selected_documents": ctx.get("selected_documents") or [],
                "selected_insights": ctx.get("selected_insights") or [],
                "suggested_filename": ctx.get("suggested_filename") or "",
            }
        )

    try:
        actual_provider = provider
        actual_model = ""
        if provider == "ollama":
            draft_text = _osc_generate_draft_with_ollama(prompt, ollama_model, ollama_url)
            actual_model = ollama_model
        elif provider == "gemini":
            draft_text, actual_model = _osc_generate_draft_with_gemini(prompt)
            if actual_model == "casper":
                actual_provider = "casper"
                actual_model = ""
        else:
            draft_text = _osc_generate_draft_with_casper(prompt)
        cleaned = _osc_clean_draft_output(draft_text)
        _osc_log_activity(
            "draft:generate",
            "drafts",
            str((ctx.get("case") or {}).get("id") or ctx.get("case_number") or ""),
            {
                "provider": actual_provider,
                "model": actual_model,
                "doc_type": ctx.get("doc_type"),
                "case_number": ctx.get("case_number"),
                "documents": len(ctx.get("selected_documents") or []),
                "insights": len(ctx.get("selected_insights") or []),
            },
        )
        return jsonify(
            {
                "ok": True,
                "provider": actual_provider,
                "model": actual_model,
                "draft_text": cleaned,
                "prompt_preview": prompt,
                "warnings": ctx.get("warnings") or [],
                "case": ctx.get("case") or {},
                "selected_documents": ctx.get("selected_documents") or [],
                "selected_insights": ctx.get("selected_insights") or [],
                "suggested_filename": ctx.get("suggested_filename") or "",
                "export_title": ctx.get("export_title") or "書狀草稿",
            }
        )
    except Exception as e:
        _osc_log_activity(
            "draft:generate_error",
            "drafts",
            str((ctx.get("case") or {}).get("id") or ctx.get("case_number") or ""),
            {"provider": provider, "error": str(e)},
        )
        return jsonify({"ok": False, "error": str(e), "prompt_preview": prompt, "warnings": ctx.get("warnings") or []}), 500


@app.route("/api/osc/drafts/export", methods=["POST"])
@login_required
def osc_drafts_export_api():
    payload = request.get_json() or {}
    text = str(payload.get("draft_text") or payload.get("text") or "").strip()
    if not text:
        return jsonify({"ok": False, "error": "draft_text required"}), 400
    title = str(payload.get("title") or payload.get("doc_type") or "書狀草稿").strip() or "書狀草稿"
    case_number = str(payload.get("case_number") or "").strip()
    suggested = str(payload.get("suggested_filename") or "").strip()
    if not suggested:
        pieces = [title, case_number or "未命名"]
        suggested = "_".join(p for p in pieces if p)
    exported = _export_osc_form_files(title, text, suggested)
    status = "success" if exported.get("success") else "failed"
    if exported.get("success") and exported.get("errors"):
        status = "partial_success"
    preferred = exported.get("export") or {}
    error_text = ""
    if exported.get("errors"):
        error_text = "; ".join(str(x.get("error") or "") for x in exported.get("errors") or [] if str(x.get("error") or "").strip())
    try:
        _osc_exec(
            "INSERT INTO pdf_generation_log (case_number, file_name, status, error_message) VALUES (%s,%s,%s,%s)",
            (
                case_number or "draft",
                str(preferred.get("filename") or suggested or title),
                status,
                error_text or None,
            ),
            fetch="none",
        )
    except Exception as e:
        logger.warning("draft export log write failed: %s", e)
    _osc_log_activity(
        "draft:export",
        "drafts",
        case_number or "draft",
        {"title": title, "status": status, "filename": str(preferred.get("filename") or suggested or title)},
    )
    http_status = 200 if exported.get("success") else 500
    return jsonify({"ok": bool(exported.get("success")), **exported, "status": status}), http_status


@app.route("/api/osc/documents", methods=["GET"])
@login_required
def osc_documents_api():
    q = (request.args.get("q") or "").strip().lower()
    case_number = (request.args.get("case_number") or "").strip()
    kind = (request.args.get("kind") or "all").strip()
    limit = max(1, min(1000, int(request.args.get("limit") or "300")))

    items = []
    di_limit = max(200, limit * 3)
    cd_limit = max(200, limit * 3)

    di_where = []
    di_params = []
    if case_number:
        di_where.append("case_number = %s")
        di_params.append(case_number)
    if q:
        like = f"%{q}%"
        di_where.append("(file_name LIKE %s OR file_path LIKE %s OR reason LIKE %s OR party LIKE %s OR subfolder_name LIKE %s)")
        di_params.extend([like, like, like, like, like])
    di_sql = """
        SELECT id, case_number, file_name, file_path, subfolder_name, reason, party, modified_date
        FROM document_index
    """
    if di_where:
        di_sql += " WHERE " + " AND ".join(di_where)
    di_sql += " ORDER BY modified_date DESC, id DESC LIMIT %s"
    di_params.append(di_limit)
    di_rows, _ = _osc_exec(di_sql, tuple(di_params), fetch="all")
    for r in di_rows:
        blob = " ".join(
            [
                str(r.get("file_name") or ""),
                str(r.get("subfolder_name") or ""),
                str(r.get("reason") or ""),
                str(r.get("party") or ""),
            ]
        )
        if not _osc_doc_kind_match(kind, blob):
            continue
        ts = r.get("modified_date")
        items.append(
            {
                "id": f"di-{r.get('id')}",
                "source": "document_index",
                "case_number": r.get("case_number") or "",
                "file_name": r.get("file_name") or "",
                "file_path": r.get("file_path") or "",
                "subfolder_name": r.get("subfolder_name") or "",
                "reason": r.get("reason") or "",
                "party": r.get("party") or "",
                "kind_label": _osc_doc_kind_label(blob),
                "timestamp": _osc_json_value(ts) if ts else "",
                "sort_ts": _osc_parse_dt(ts).timestamp() if _osc_parse_dt(ts) else 0,
            }
        )

    cd_where = []
    cd_params = []
    if case_number:
        cd_where.append("(cd.case_id = %s OR cd.case_id IN (SELECT id FROM cases WHERE case_number=%s))")
        cd_params.extend([case_number, case_number])
    if q:
        like = f"%{q}%"
        cd_where.append("(cd.file_name LIKE %s OR cd.file_path LIKE %s OR cd.document_type LIKE %s OR cd.description LIKE %s)")
        cd_params.extend([like, like, like, like])
    cd_sql = """
        SELECT cd.id, cd.case_id, c.case_number AS case_number_ref, cd.document_type, cd.file_name, cd.file_path, cd.description, cd.upload_date
        FROM case_documents cd
        LEFT JOIN cases c ON c.id = cd.case_id
    """
    if cd_where:
        cd_sql += " WHERE " + " AND ".join(cd_where)
    cd_sql += " ORDER BY upload_date DESC, id DESC LIMIT %s"
    cd_params.append(cd_limit)
    cd_rows, _ = _osc_exec(cd_sql, tuple(cd_params), fetch="all")
    for r in cd_rows:
        blob = " ".join(
            [
                str(r.get("document_type") or ""),
                str(r.get("file_name") or ""),
                str(r.get("description") or ""),
            ]
        )
        if not _osc_doc_kind_match(kind, blob):
            continue
        ts = r.get("upload_date")
        items.append(
            {
                "id": f"cd-{r.get('id')}",
                "source": "case_documents",
                "case_number": r.get("case_number_ref") or r.get("case_id") or "",
                "file_name": r.get("file_name") or "",
                "file_path": r.get("file_path") or "",
                "subfolder_name": r.get("document_type") or "",
                "reason": r.get("description") or "",
                "party": "",
                "kind_label": _osc_doc_kind_label(blob),
                "timestamp": _osc_json_value(ts) if ts else "",
                "sort_ts": _osc_parse_dt(ts).timestamp() if _osc_parse_dt(ts) else 0,
            }
        )

    items.sort(key=lambda x: x.get("sort_ts") or 0, reverse=True)
    out = items[:limit]
    for it in out:
        it.pop("sort_ts", None)
    return jsonify({"ok": True, "items": out})


@app.route("/api/osc/documents/open", methods=["POST"])
@login_required
def osc_documents_open_api():
    payload = request.get_json() or {}
    raw = str(payload.get("path") or "").strip()
    if not raw:
        return jsonify({"ok": False, "error": "path required"}), 400
    norm = _osc_norm_path(raw)
    local_candidates = _osc_local_path_candidates(norm)
    smb_candidates = _osc_smb_candidates(norm)
    chosen_open_path = ""
    open_result = {"ok": False, "error": "open_failed"}

    for lp in local_candidates:
        try:
            if lp and os.path.exists(lp):
                r = _osc_try_open_path(lp)
                chosen_open_path = lp
                open_result = r
                if r.get("ok"):
                    break
        except Exception:
            continue
    if not open_result.get("ok"):
        for sp in smb_candidates:
            r = _osc_try_open_path(sp)
            chosen_open_path = sp
            open_result = r
            if r.get("ok"):
                break
    return jsonify(
        {
            "ok": True,
            "path": norm,
            "local_candidates": local_candidates,
            "smb_candidates": smb_candidates,
            "chosen_open_path": chosen_open_path,
            "open_result": open_result,
        }
    )


@app.route("/api/osc/files/content", methods=["GET"])
@login_required
def osc_file_content_api():
    raw = str(request.args.get("path") or "").strip()
    if not raw:
        return jsonify({"ok": False, "error": "path required"}), 400
    local_file = _osc_resolve_existing_local_path(raw, prefer_dir=False)
    if not local_file:
        return jsonify({"ok": False, "error": "file_not_found"}), 404
    if not _osc_is_safe_local_path(local_file):
        return jsonify({"ok": False, "error": "path_not_allowed"}), 403
    inline = str(request.args.get("inline") or "").strip() in {"1", "true", "yes"}
    mime, _ = mimetypes.guess_type(local_file)
    # Always read into memory first to avoid Errno 11 (Resource deadlock avoided)
    # from Synology Drive sync locks. The deadlock occurs in werkzeug's WSGI
    # streaming layer (wsgi.py __next__) which is outside try/except scope when
    # using send_file with a file path directly.
    import io
    try:
        with open(local_file, "rb") as f:
            buf = io.BytesIO(f.read())
    except OSError as e:
        log.error("osc_file_content_api read error (errno=%s): %s – file=%s", e.errno, e, local_file)
        return jsonify({"ok": False, "error": f"send_file_error: {e}"}), 500
    try:
        resp = send_file(
            buf,
            mimetype=mime or "application/octet-stream",
            as_attachment=not inline,
            download_name=os.path.basename(local_file),
        )
        # Add ETag based on file mtime+size for caching
        try:
            st = os.stat(local_file)
            resp.headers["ETag"] = f'"{int(st.st_mtime)}-{st.st_size}"'
            resp.headers["Cache-Control"] = "private, max-age=300"
        except OSError:
            pass
        return resp
    except Exception as e:
        log.error("osc_file_content_api send_file error: %s – file=%s", e, local_file)
        return jsonify({"ok": False, "error": f"send_file_error: {e}"}), 500


@app.route("/api/osc/files/text", methods=["GET", "PUT"])
@login_required
def osc_file_text_api():
    if request.method == "GET":
        raw = str(request.args.get("path") or "").strip()
        if not raw:
            return jsonify({"ok": False, "error": "path required"}), 400
        if not _osc_is_editable_text_path(raw):
            return jsonify({"ok": False, "error": "not_editable_text"}), 400
        local_file = _osc_resolve_existing_local_path(raw, prefer_dir=False)
        if not local_file:
            return jsonify({"ok": False, "error": "file_not_found"}), 404
        if not _osc_is_safe_local_path(local_file):
            return jsonify({"ok": False, "error": "path_not_allowed"}), 403
        try:
            content, encoding = _osc_read_text_file(local_file)
        except ValueError as e:
            return jsonify({"ok": False, "error": str(e)}), 400
        return jsonify(
            {
                "ok": True,
                "path": raw,
                "local_path": local_file,
                "content": content,
                "encoding": encoding,
                "size": os.path.getsize(local_file),
            }
        )

    payload = request.get_json() or {}
    raw = str(payload.get("path") or "").strip()
    content = payload.get("content")
    if not raw:
        return jsonify({"ok": False, "error": "path required"}), 400
    if content is None:
        return jsonify({"ok": False, "error": "content required"}), 400
    if not _osc_is_editable_text_path(raw):
        return jsonify({"ok": False, "error": "not_editable_text"}), 400
    local_file = _osc_resolve_existing_local_path(raw, prefer_dir=False)
    if not local_file:
        return jsonify({"ok": False, "error": "file_not_found"}), 404
    if not _osc_is_safe_local_path(local_file):
        return jsonify({"ok": False, "error": "path_not_allowed"}), 403
    text = str(content)
    Path(local_file).write_text(text, encoding="utf-8")
    return jsonify({"ok": True, "path": raw, "local_path": local_file, "size": len(text.encode('utf-8'))})


@app.route("/api/osc/files/upload", methods=["POST"])
@login_required
def osc_file_upload_api():
    folder_path = str(request.form.get("folder_path") or request.args.get("folder_path") or "").strip()
    relative_path = str(request.form.get("relative_path") or request.args.get("relative_path") or "").strip().strip("/")
    overwrite = str(request.form.get("overwrite") or request.args.get("overwrite") or "").strip().lower() in {"1", "true", "yes", "on"}
    if not folder_path:
        return jsonify({"ok": False, "error": "folder_path required"}), 400
    base_folder = _osc_resolve_existing_local_path(folder_path, prefer_dir=True)
    if not base_folder:
        return jsonify({"ok": False, "error": "folder_not_found"}), 404
    target_dir = os.path.realpath(os.path.join(base_folder, relative_path or ""))
    base_real = os.path.realpath(base_folder)
    if target_dir != base_real and not target_dir.startswith(base_real + os.sep):
        return jsonify({"ok": False, "error": "path_escape"}), 400
    if not os.path.isdir(target_dir):
        return jsonify({"ok": False, "error": "target_dir_not_found"}), 404
    uploads = request.files.getlist("file") or request.files.getlist("files")
    if not uploads:
        return jsonify({"ok": False, "error": "file required"}), 400

    saved = []
    for uploaded in uploads:
        name = os.path.basename(str(uploaded.filename or "").strip())
        if not name:
            continue
        dest = os.path.join(target_dir, name)
        if os.path.exists(dest) and not overwrite:
            return jsonify({"ok": False, "error": "file_exists", "file_name": name, "target_path": dest}), 409
        uploaded.save(dest)
        saved.append(
            {
                "file_name": name,
                "target_path": dest,
                "size": os.path.getsize(dest),
            }
        )
    if not saved:
        return jsonify({"ok": False, "error": "no_valid_files"}), 400
    return jsonify({"ok": True, "saved": saved, "target_dir": target_dir, "overwrite": overwrite})


@app.route("/api/osc/document-templates", methods=["GET", "POST"])
@login_required
def osc_document_templates_api():
    if request.method == "GET":
        q = (request.args.get("q") or "").strip()
        case_number = (request.args.get("case_number") or "").strip()
        doc_type = (request.args.get("doc_type") or "").strip()
        limit = max(1, min(1000, int(request.args.get("limit") or "300")))
        sql = (
            "SELECT id, doc_type, party_name, case_number, division, template_data, created_date, last_used, use_count "
            "FROM document_templates WHERE 1=1 "
        )
        params = []
        if case_number:
            sql += "AND case_number=%s "
            params.append(case_number)
        if doc_type:
            sql += "AND doc_type=%s "
            params.append(doc_type)
        if q:
            like = f"%{q}%"
            sql += "AND (doc_type LIKE %s OR party_name LIKE %s OR case_number LIKE %s OR division LIKE %s OR template_data LIKE %s) "
            params.extend([like, like, like, like, like])
        sql += "ORDER BY COALESCE(last_used, created_date) DESC, id DESC LIMIT %s"
        params.append(limit)
        rows, _ = _osc_exec(sql, tuple(params), fetch="all")
        return jsonify({"ok": True, "items": rows})

    payload = request.get_json() or {}
    row_id = str(payload.get("id") or "").strip()
    body = {
        "doc_type": (payload.get("doc_type") or "").strip() or None,
        "party_name": (payload.get("party_name") or "").strip() or None,
        "case_number": (payload.get("case_number") or "").strip() or None,
        "division": (payload.get("division") or "").strip() or None,
        "template_data": _osc_template_data_json_or_wrap(payload.get("template_data")),
        "use_count": _osc_safe_int(payload.get("use_count"), 0),
    }
    if row_id:
        sets = [f"{k}=%s" for k in body.keys()]
        vals = list(body.values()) + [row_id]
        result, _ = _osc_exec(f"UPDATE document_templates SET {','.join(sets)} WHERE id=%s", tuple(vals), fetch="none")
        return jsonify({"ok": True, "mode": "update", "id": row_id, "result": result})

    cols = list(body.keys())
    vals = [body[c] for c in cols]
    result, _ = _osc_exec(
        f"INSERT INTO document_templates ({','.join(cols)}) VALUES ({','.join(['%s'] * len(cols))})",
        tuple(vals),
        fetch="none",
    )
    return jsonify({"ok": True, "mode": "insert", "id": result.get("lastrowid"), "result": result})


@app.route("/api/osc/document-templates/<int:row_id>", methods=["GET", "PUT", "DELETE"])
@login_required
def osc_document_template_detail_api(row_id):
    if request.method == "GET":
        row, _ = _osc_exec("SELECT * FROM document_templates WHERE id=%s", (row_id,), fetch="one")
        if not row:
            return jsonify({"ok": False, "error": "not found"}), 404
        return jsonify({"ok": True, "item": row})
    if request.method == "DELETE":
        result, _ = _osc_exec("DELETE FROM document_templates WHERE id=%s", (row_id,), fetch="none")
        return jsonify({"ok": True, "result": result})
    payload = request.get_json() or {}
    allowed = ["doc_type", "party_name", "case_number", "division", "template_data", "last_used", "use_count"]
    sets, vals = [], []
    for k in allowed:
        if k in payload:
            sets.append(f"{k}=%s")
            if k == "use_count":
                vals.append(_osc_safe_int(payload.get(k), 0))
            elif k == "template_data":
                vals.append(_osc_template_data_json_or_wrap(payload.get(k)))
            else:
                vals.append((payload.get(k) or "").strip() or None)
    if not sets:
        return jsonify({"ok": False, "error": "no fields"}), 400
    vals.append(row_id)
    result, _ = _osc_exec(f"UPDATE document_templates SET {','.join(sets)} WHERE id=%s", tuple(vals), fetch="none")
    return jsonify({"ok": True, "result": result})


@app.route("/api/osc/document-keywords", methods=["GET", "POST"])
@login_required
def osc_document_keywords_api():
    if request.method == "GET":
        q = (request.args.get("q") or "").strip()
        case_number = (request.args.get("case_number") or "").strip()
        category = (request.args.get("category") or "").strip()
        limit = max(1, min(1000, int(request.args.get("limit") or "300")))
        sql = (
            "SELECT id, case_number, keyword_name, keyword_content, category, hotkey, is_case_specific, usage_count, created_date, modified_date "
            "FROM document_keywords WHERE 1=1 "
        )
        params = []
        if case_number:
            sql += "AND case_number=%s "
            params.append(case_number)
        if category:
            sql += "AND category=%s "
            params.append(category)
        if q:
            like = f"%{q}%"
            sql += "AND (case_number LIKE %s OR keyword_name LIKE %s OR keyword_content LIKE %s OR category LIKE %s OR hotkey LIKE %s) "
            params.extend([like, like, like, like, like])
        sql += "ORDER BY modified_date DESC, created_date DESC, id DESC LIMIT %s"
        params.append(limit)
        rows, _ = _osc_exec(sql, tuple(params), fetch="all")
        return jsonify({"ok": True, "items": rows})

    payload = request.get_json() or {}
    row_id = str(payload.get("id") or "").strip()
    body = {
        "case_number": (payload.get("case_number") or "").strip() or None,
        "keyword_name": (payload.get("keyword_name") or "").strip() or None,
        "keyword_content": (payload.get("keyword_content") or "").strip() or None,
        "category": (payload.get("category") or "").strip() or None,
        "hotkey": (payload.get("hotkey") or "").strip() or None,
        "is_case_specific": 1 if str(payload.get("is_case_specific") or "").strip().lower() in {"1", "true", "yes", "on"} else 0,
        "usage_count": _osc_safe_int(payload.get("usage_count"), 0),
    }
    if not body["keyword_name"]:
        return jsonify({"ok": False, "error": "keyword_name required"}), 400
    if row_id:
        sets = [f"{k}=%s" for k in body.keys()]
        vals = list(body.values()) + [row_id]
        result, _ = _osc_exec(
            f"UPDATE document_keywords SET {','.join(sets)}, modified_date=NOW() WHERE id=%s",
            tuple(vals),
            fetch="none",
        )
        return jsonify({"ok": True, "mode": "update", "id": row_id, "result": result})

    cols = list(body.keys())
    vals = [body[c] for c in cols]
    result, _ = _osc_exec(
        f"INSERT INTO document_keywords ({','.join(cols)}) VALUES ({','.join(['%s'] * len(cols))})",
        tuple(vals),
        fetch="none",
    )
    return jsonify({"ok": True, "mode": "insert", "result": result})


@app.route("/api/osc/document-keywords/<int:row_id>", methods=["GET", "PUT", "DELETE"])
@login_required
def osc_document_keyword_detail_api(row_id):
    if request.method == "GET":
        row, _ = _osc_exec("SELECT * FROM document_keywords WHERE id=%s", (row_id,), fetch="one")
        if not row:
            return jsonify({"ok": False, "error": "not found"}), 404
        return jsonify({"ok": True, "item": row})
    if request.method == "DELETE":
        result, _ = _osc_exec("DELETE FROM document_keywords WHERE id=%s", (row_id,), fetch="none")
        return jsonify({"ok": True, "result": result})
    payload = request.get_json() or {}
    allowed = ["case_number", "keyword_name", "keyword_content", "category", "hotkey", "is_case_specific", "usage_count"]
    sets, vals = [], []
    for k in allowed:
        if k not in payload:
            continue
        sets.append(f"{k}=%s")
        if k in {"usage_count", "is_case_specific"}:
            vals.append(_osc_safe_int(payload.get(k), 0))
        else:
            vals.append((payload.get(k) or "").strip() or None)
    if not sets:
        return jsonify({"ok": False, "error": "no fields"}), 400
    vals.append(row_id)
    result, _ = _osc_exec(
        f"UPDATE document_keywords SET {','.join(sets)}, modified_date=NOW() WHERE id=%s",
        tuple(vals),
        fetch="none",
    )
    return jsonify({"ok": True, "result": result})


@app.route("/api/osc/document-replacements", methods=["GET", "POST"])
@login_required
def osc_document_replacements_api():
    if request.method == "GET":
        q = (request.args.get("q") or "").strip()
        case_number = (request.args.get("case_number") or "").strip()
        limit = max(1, min(1000, int(request.args.get("limit") or "300")))
        sql = (
            "SELECT id, template_file, new_case_number, old_client_name, new_client_name, old_data, new_data, replaced_date "
            "FROM document_replacements WHERE 1=1 "
        )
        params = []
        if case_number:
            sql += "AND new_case_number=%s "
            params.append(case_number)
        if q:
            like = f"%{q}%"
            sql += "AND (template_file LIKE %s OR new_case_number LIKE %s OR old_client_name LIKE %s OR new_client_name LIKE %s OR old_data LIKE %s OR new_data LIKE %s) "
            params.extend([like, like, like, like, like, like])
        sql += "ORDER BY replaced_date DESC, id DESC LIMIT %s"
        params.append(limit)
        rows, _ = _osc_exec(sql, tuple(params), fetch="all")
        return jsonify({"ok": True, "items": rows})

    payload = request.get_json() or {}
    body = {
        "template_file": (payload.get("template_file") or "").strip() or None,
        "new_case_number": (payload.get("new_case_number") or payload.get("case_number") or "").strip() or None,
        "old_client_name": (payload.get("old_client_name") or "").strip() or None,
        "new_client_name": (payload.get("new_client_name") or "").strip() or None,
        "old_data": (payload.get("old_data") or "").strip() or None,
        "new_data": (payload.get("new_data") or "").strip() or None,
    }
    cols = list(body.keys())
    vals = [body[c] for c in cols]
    result, _ = _osc_exec(
        f"INSERT INTO document_replacements ({','.join(cols)}) VALUES ({','.join(['%s'] * len(cols))})",
        tuple(vals),
        fetch="none",
    )
    return jsonify({"ok": True, "result": result})


@app.route("/api/osc/document-replacements/<int:row_id>", methods=["GET", "DELETE"])
@login_required
def osc_document_replacement_detail_api(row_id):
    if request.method == "GET":
        row, _ = _osc_exec("SELECT * FROM document_replacements WHERE id=%s", (row_id,), fetch="one")
        if not row:
            return jsonify({"ok": False, "error": "not found"}), 404
        return jsonify({"ok": True, "item": row})
    result, _ = _osc_exec("DELETE FROM document_replacements WHERE id=%s", (row_id,), fetch="none")
    return jsonify({"ok": True, "result": result})


@app.route("/api/osc/laf", methods=["GET"])
@login_required
def osc_laf_api():
    q = (request.args.get("q") or "").strip()
    case_number = (request.args.get("case_number") or "").strip()
    limit = max(1, min(1000, int(request.args.get("limit") or "300")))

    like = f"%{q}%"
    where_case = "case_number=%s" if case_number else "1=1"
    case_params = [case_number] if case_number else []

    checklist_sql = (
        "SELECT id, case_number, item_key, item_label, status, notes, last_updated "
        "FROM legal_aid_checklists "
        f"WHERE {where_case} "
    )
    checklist_params = list(case_params)
    if q:
        checklist_sql += "AND (case_number LIKE %s OR item_key LIKE %s OR item_label LIKE %s OR status LIKE %s OR notes LIKE %s) "
        checklist_params.extend([like, like, like, like, like])
    checklist_sql += "ORDER BY last_updated DESC, id DESC LIMIT %s"
    checklist_params.append(limit)
    checklist, _ = _osc_exec(checklist_sql, tuple(checklist_params), fetch="all")

    lifecycle_sql = (
        "SELECT id, case_number, event_type, status, created_at, completed_at, event_data "
        "FROM laf_lifecycle_log "
        f"WHERE {where_case} "
    )
    lifecycle_params = list(case_params)
    if q:
        lifecycle_sql += "AND (case_number LIKE %s OR event_type LIKE %s OR status LIKE %s OR event_data LIKE %s) "
        lifecycle_params.extend([like, like, like, like])
    lifecycle_sql += "ORDER BY created_at DESC, id DESC LIMIT %s"
    lifecycle_params.append(limit)
    lifecycle, _ = _osc_exec(lifecycle_sql, tuple(lifecycle_params), fetch="all")

    email_sql = "SELECT id, gmail_message_id, subject, sender, received_at, processed_at, status, case_number, created_case_id, error_message FROM laf_email_records WHERE 1=1 "
    email_params = []
    if case_number:
        email_sql += "AND case_number=%s "
        email_params.append(case_number)
    if q:
        email_sql += "AND (subject LIKE %s OR sender LIKE %s OR case_number LIKE %s OR status LIKE %s OR error_message LIKE %s) "
        email_params.extend([like, like, like, like, like])
    email_sql += "ORDER BY received_at DESC, id DESC LIMIT %s"
    email_params.append(limit)
    emails, _ = _osc_exec(email_sql, tuple(email_params), fetch="all")

    return jsonify(
        {
            "ok": True,
            "items": {
                "checklist": checklist or [],
                "lifecycle": lifecycle or [],
                "emails": emails or [],
            },
            "counts": {
                "checklist": len(checklist or []),
                "lifecycle": len(lifecycle or []),
                "emails": len(emails or []),
            },
        }
    )


## accounting routes → moved to api.blueprints.osc_accounting


@app.route("/api/osc/quotations", methods=["GET", "POST"])
@login_required
def osc_quotations_api():
    if request.method == "GET":
        q = (request.args.get("q") or "").strip()
        status = (request.args.get("status") or "").strip()
        limit = max(1, min(1000, int(request.args.get("limit") or "300")))
        sql = (
            "SELECT id, client_name, project_name, phone, email, date, expiry, subtotal, discount, tax, total, status, updated_date, created_date "
            "FROM quotations WHERE 1=1 "
        )
        params = []
        if status:
            sql += "AND status=%s "
            params.append(status)
        if q:
            like = f"%{q}%"
            sql += "AND (id LIKE %s OR client_name LIKE %s OR project_name LIKE %s OR phone LIKE %s OR email LIKE %s OR notes LIKE %s) "
            params.extend([like, like, like, like, like, like])
        sql += "ORDER BY updated_date DESC, created_date DESC LIMIT %s"
        params.append(limit)
        rows, _ = _osc_exec(sql, tuple(params), fetch="all")
        return jsonify({"ok": True, "items": rows})

    payload = request.get_json() or {}
    row_id = (payload.get("id") or "").strip() or f"q-{datetime.now().strftime('%Y%m%d')}-{uuid.uuid4().hex[:8]}"
    client_name = (payload.get("client_name") or "").strip()
    project_name = (payload.get("project_name") or "").strip()
    if not client_name or not project_name:
        return jsonify({"ok": False, "error": "client_name/project_name required"}), 400
    def _fnum(x, d=0.0):
        try:
            return float(x if x is not None and str(x).strip() != "" else d)
        except Exception:
            return float(d)
    cols = [
        "id", "client_name", "project_name", "contact", "phone", "email", "address", "tax_id",
        "date", "expiry", "items", "subtotal", "discount", "tax", "total", "status", "notes", "extended_data"
    ]
    vals = [
        row_id,
        client_name,
        project_name,
        (payload.get("contact") or "").strip() or None,
        (payload.get("phone") or "").strip() or None,
        (payload.get("email") or "").strip() or None,
        (payload.get("address") or "").strip() or None,
        (payload.get("tax_id") or "").strip() or None,
        (payload.get("date") or "").strip() or None,
        (payload.get("expiry") or "").strip() or None,
        _osc_json_or_wrap(payload.get("items"), fallback_key="items"),
        _fnum(payload.get("subtotal"), 0),
        _fnum(payload.get("discount"), 0),
        _fnum(payload.get("tax"), 0),
        _fnum(payload.get("total"), 0),
        (payload.get("status") or "draft").strip() or "draft",
        (payload.get("notes") or "").strip() or None,
        _osc_json_or_wrap(payload.get("extended_data"), fallback_key="extended_data"),
    ]
    try:
        result, _ = _osc_exec(
            f"INSERT INTO quotations ({','.join(cols)}) VALUES ({','.join(['%s'] * len(cols))})",
            tuple(vals),
            fetch="none",
        )
        return jsonify({"ok": True, "mode": "insert", "id": row_id, "result": result})
    except Exception as e:
        msg = str(e)
        is_dup = ("1062" in msg) or ("Duplicate entry" in msg)
        if not is_dup:
            return jsonify({"ok": False, "error": msg}), 500
        # duplicate id -> update
        sets = [
            "client_name=%s", "project_name=%s", "contact=%s", "phone=%s", "email=%s", "address=%s", "tax_id=%s",
            "date=%s", "expiry=%s", "items=%s", "subtotal=%s", "discount=%s", "tax=%s", "total=%s", "status=%s",
            "notes=%s", "extended_data=%s"
        ]
        vals2 = [
            client_name,
            project_name,
            (payload.get("contact") or "").strip() or None,
            (payload.get("phone") or "").strip() or None,
            (payload.get("email") or "").strip() or None,
            (payload.get("address") or "").strip() or None,
            (payload.get("tax_id") or "").strip() or None,
            (payload.get("date") or "").strip() or None,
            (payload.get("expiry") or "").strip() or None,
            _osc_json_or_wrap(payload.get("items"), fallback_key="items"),
            _fnum(payload.get("subtotal"), 0),
            _fnum(payload.get("discount"), 0),
            _fnum(payload.get("tax"), 0),
            _fnum(payload.get("total"), 0),
            (payload.get("status") or "draft").strip() or "draft",
            (payload.get("notes") or "").strip() or None,
            _osc_json_or_wrap(payload.get("extended_data"), fallback_key="extended_data"),
            row_id,
        ]
        result, _ = _osc_exec(f"UPDATE quotations SET {','.join(sets)} WHERE id=%s", tuple(vals2), fetch="none")
        return jsonify({"ok": True, "mode": "upsert", "id": row_id, "result": result})


@app.route("/api/osc/quotations/<row_id>", methods=["GET", "PUT", "DELETE"])
@login_required
def osc_quotation_detail_api(row_id):
    row_id = (row_id or "").strip()
    if request.method == "GET":
        row, _ = _osc_exec("SELECT * FROM quotations WHERE id=%s", (row_id,), fetch="one")
        if not row:
            return jsonify({"ok": False, "error": "not found"}), 404
        return jsonify({"ok": True, "item": row})
    if request.method == "DELETE":
        result, _ = _osc_exec("DELETE FROM quotations WHERE id=%s", (row_id,), fetch="none")
        return jsonify({"ok": True, "result": result})
    payload = request.get_json() or {}
    allowed = [
        "client_name", "project_name", "contact", "phone", "email", "address", "tax_id", "date", "expiry",
        "items", "subtotal", "discount", "tax", "total", "status", "notes", "extended_data"
    ]
    sets, vals = [], []
    for k in allowed:
        if k not in payload:
            continue
        sets.append(f"{k}=%s")
        if k in {"subtotal", "discount", "tax", "total"}:
            try:
                vals.append(float(payload.get(k) or 0))
            except Exception:
                return jsonify({"ok": False, "error": f"{k} invalid"}), 400
        elif k in {"items", "extended_data"}:
            vals.append(_osc_json_or_wrap(payload.get(k), fallback_key=k))
        else:
            vals.append((payload.get(k) or "").strip() or None)
    if not sets:
        return jsonify({"ok": False, "error": "no fields"}), 400
    vals.append(row_id)
    result, _ = _osc_exec(f"UPDATE quotations SET {','.join(sets)} WHERE id=%s", tuple(vals), fetch="none")
    return jsonify({"ok": True, "result": result})


@app.route("/api/osc/quotation-templates", methods=["GET", "POST"])
@login_required
def osc_quotation_templates_api():
    if request.method == "GET":
        q = (request.args.get("q") or "").strip()
        limit = max(1, min(1000, int(request.args.get("limit") or "300")))
        sql = "SELECT id, name, description, items, notes, is_default, updated_date, created_date FROM quotation_templates WHERE 1=1 "
        params = []
        if q:
            like = f"%{q}%"
            sql += "AND (name LIKE %s OR description LIKE %s OR notes LIKE %s) "
            params.extend([like, like, like])
        sql += "ORDER BY is_default DESC, updated_date DESC, created_date DESC LIMIT %s"
        params.append(limit)
        rows, _ = _osc_exec(sql, tuple(params), fetch="all")
        return jsonify({"ok": True, "items": rows})
    payload = request.get_json() or {}
    name = (payload.get("name") or "").strip()
    if not name:
        return jsonify({"ok": False, "error": "name required"}), 400
    is_default = 1 if str(payload.get("is_default") or "").strip().lower() in {"1", "true", "yes", "on"} else 0
    result, _ = _osc_exec(
        "INSERT INTO quotation_templates (name, description, items, notes, is_default) VALUES (%s,%s,%s,%s,%s)",
        (
            name,
            (payload.get("description") or "").strip() or None,
            _osc_json_or_wrap(payload.get("items"), fallback_key="items"),
            (payload.get("notes") or "").strip() or None,
            is_default,
        ),
        fetch="none",
    )
    return jsonify({"ok": True, "id": result.get("lastrowid"), "result": result})


@app.route("/api/osc/quotation-templates/<int:row_id>", methods=["GET", "PUT", "DELETE"])
@login_required
def osc_quotation_template_detail_api(row_id):
    if request.method == "GET":
        row, _ = _osc_exec("SELECT * FROM quotation_templates WHERE id=%s", (row_id,), fetch="one")
        if not row:
            return jsonify({"ok": False, "error": "not found"}), 404
        return jsonify({"ok": True, "item": row})
    if request.method == "DELETE":
        result, _ = _osc_exec("DELETE FROM quotation_templates WHERE id=%s", (row_id,), fetch="none")
        return jsonify({"ok": True, "result": result})
    payload = request.get_json() or {}
    allowed = ["name", "description", "items", "notes", "is_default"]
    sets, vals = [], []
    for k in allowed:
        if k not in payload:
            continue
        sets.append(f"{k}=%s")
        if k == "is_default":
            vals.append(1 if str(payload.get(k) or "").strip().lower() in {"1", "true", "yes", "on"} else 0)
        elif k == "items":
            vals.append(_osc_json_or_wrap(payload.get("items"), fallback_key="items"))
        else:
            vals.append((payload.get(k) or "").strip() or None)
    if not sets:
        return jsonify({"ok": False, "error": "no fields"}), 400
    vals.append(row_id)
    result, _ = _osc_exec(f"UPDATE quotation_templates SET {','.join(sets)} WHERE id=%s", tuple(vals), fetch="none")
    return jsonify({"ok": True, "result": result})


@app.route("/api/osc/calendar/events", methods=["GET", "POST"])
@login_required
def osc_calendar_events_api():
    if request.method == "GET":
        q = (request.args.get("q") or "").strip()
        case_number = (request.args.get("case_number") or "").strip()
        start_date = (request.args.get("start_date") or "").strip()
        end_date = (request.args.get("end_date") or "").strip()
        limit = max(1, min(1000, int(request.args.get("limit") or "300")))
        sql = (
            "SELECT id, event_id, title, summary, description, start_date, end_date, color, location, is_all_day, reminder_minutes, case_number, created_date, updated_date "
            "FROM calendar_events WHERE 1=1 "
        )
        params = []
        if case_number:
            sql += "AND case_number=%s "
            params.append(case_number)
        if start_date:
            sql += "AND start_date >= %s "
            params.append(start_date)
        if end_date:
            sql += "AND end_date <= %s "
            params.append(end_date)
        if q:
            like = f"%{q}%"
            sql += "AND (title LIKE %s OR summary LIKE %s OR description LIKE %s OR location LIKE %s OR case_number LIKE %s) "
            params.extend([like, like, like, like, like])
        sql += "ORDER BY start_date DESC, id DESC LIMIT %s"
        params.append(limit)
        rows, _ = _osc_exec(sql, tuple(params), fetch="all")
        return jsonify({"ok": True, "items": rows})

    payload = request.get_json() or {}
    title = (payload.get("title") or "").strip()
    start = (payload.get("start_date") or "").strip()
    end = (payload.get("end_date") or "").strip()
    if not title or not start or not end:
        return jsonify({"ok": False, "error": "title/start_date/end_date required"}), 400
    event_id = (payload.get("event_id") or "").strip() or f"osc-{uuid.uuid4().hex[:20]}"
    is_all_day = 1 if str(payload.get("is_all_day") or "").strip().lower() in {"1", "true", "yes", "on"} else 0
    reminder = _osc_safe_int(payload.get("reminder_minutes"), 0)
    result, _ = _osc_exec(
        "INSERT INTO calendar_events (event_id, title, summary, description, start_date, end_date, color, location, is_all_day, reminder_minutes, raw_data, case_number) "
        "VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)",
        (
            event_id,
            title,
            (payload.get("summary") or "").strip() or None,
            (payload.get("description") or "").strip() or None,
            start,
            end,
            (payload.get("color") or "#3498db").strip() or "#3498db",
            (payload.get("location") or "").strip() or None,
            is_all_day,
            reminder,
            _osc_json_or_wrap(payload.get("raw_data"), fallback_key="raw_data"),
            (payload.get("case_number") or "").strip() or None,
        ),
        fetch="none",
    )
    return jsonify({"ok": True, "id": result.get("lastrowid"), "event_id": event_id, "result": result})


@app.route("/api/osc/calendar/events/<int:row_id>", methods=["GET", "PUT", "DELETE"])
@login_required
def osc_calendar_event_detail_api(row_id):
    if request.method == "GET":
        row, _ = _osc_exec("SELECT * FROM calendar_events WHERE id=%s", (row_id,), fetch="one")
        if not row:
            return jsonify({"ok": False, "error": "not found"}), 404
        return jsonify({"ok": True, "item": row})
    if request.method == "DELETE":
        result, _ = _osc_exec("DELETE FROM calendar_events WHERE id=%s", (row_id,), fetch="none")
        return jsonify({"ok": True, "result": result})

    payload = request.get_json() or {}
    allowed = [
        "event_id", "title", "summary", "description", "start_date", "end_date",
        "color", "location", "is_all_day", "reminder_minutes", "raw_data", "case_number"
    ]
    sets, vals = [], []
    for k in allowed:
        if k not in payload:
            continue
        sets.append(f"{k}=%s")
        if k in {"is_all_day", "reminder_minutes"}:
            vals.append(_osc_safe_int(payload.get(k), 0))
        elif k == "raw_data":
            vals.append(_osc_json_or_wrap(payload.get("raw_data"), fallback_key="raw_data"))
        else:
            vals.append((payload.get(k) or "").strip() or None)
    if not sets:
        return jsonify({"ok": False, "error": "no fields"}), 400
    vals.append(row_id)
    result, _ = _osc_exec(f"UPDATE calendar_events SET {','.join(sets)} WHERE id=%s", tuple(vals), fetch="none")
    return jsonify({"ok": True, "result": result})


@app.route("/api/osc/clients", methods=["GET", "POST"])
@login_required
def osc_clients_api():
    if request.method == "GET":
        q = (request.args.get("q") or "").strip()
        limit = max(1, min(500, int(request.args.get("limit") or "200")))
        if q:
            like = f"%{q}%"
            rows, _ = _osc_exec(
                """
                SELECT id, name, contact_person, phone, email, address, tax_id, notes, status, updated_date, created_date
                FROM clients
                WHERE name LIKE %s OR phone LIKE %s OR email LIKE %s
                ORDER BY updated_date DESC, created_date DESC
                LIMIT %s
                """,
                (like, like, like, limit),
                fetch="all",
            )
        else:
            rows, _ = _osc_exec(
                """
                SELECT id, name, contact_person, phone, email, address, tax_id, notes, status, updated_date, created_date
                FROM clients
                ORDER BY updated_date DESC, created_date DESC
                LIMIT %s
                """,
                (limit,),
                fetch="all",
            )
        return jsonify({"ok": True, "items": rows})
    payload = request.get_json() or {}
    row_id = (payload.get("id") or f"webc-{uuid.uuid4().hex[:12]}").strip()
    name = (payload.get("name") or "").strip()
    if not name:
        return jsonify({"ok": False, "error": "name required"}), 400
    cols = ["id", "name", "contact_person", "phone", "email", "address", "tax_id", "notes", "status"]
    vals = [
        row_id,
        name,
        (payload.get("contact_person") or "").strip() or None,
        (payload.get("phone") or "").strip() or None,
        (payload.get("email") or "").strip() or None,
        (payload.get("address") or "").strip() or None,
        (payload.get("tax_id") or "").strip() or None,
        (payload.get("notes") or "").strip() or None,
        (payload.get("status") or "Active").strip() or "Active",
    ]
    result, _ = _osc_exec(
        f"INSERT INTO clients ({','.join(cols)}) VALUES ({','.join(['%s'] * len(cols))})",
        tuple(vals),
        fetch="none",
    )
    return jsonify({"ok": True, "result": result, "id": row_id})


@app.route("/api/osc/clients/<row_id>", methods=["GET", "PUT", "DELETE"])
@login_required
def osc_client_detail_api(row_id):
    row_id = (row_id or "").strip()
    if request.method == "GET":
        row, _ = _osc_exec("SELECT * FROM clients WHERE id=%s", (row_id,), fetch="one")
        if not row:
            return jsonify({"ok": False, "error": "not found"}), 404
        return jsonify({"ok": True, "item": row})
    if request.method == "DELETE":
        result, _ = _osc_exec("DELETE FROM clients WHERE id=%s", (row_id,), fetch="none")
        return jsonify({"ok": True, "result": result})
    payload = request.get_json() or {}
    allowed = ["name", "contact_person", "phone", "email", "address", "tax_id", "notes", "status"]
    sets = []
    vals = []
    for k in allowed:
        if k in payload:
            sets.append(f"{k}=%s")
            vals.append((payload.get(k) or "").strip() or None)
    if not sets:
        return jsonify({"ok": False, "error": "no fields"}), 400
    sets.append("updated_date=NOW()")
    vals.append(row_id)
    result, _ = _osc_exec(f"UPDATE clients SET {','.join(sets)} WHERE id=%s", tuple(vals), fetch="none")
    return jsonify({"ok": True, "result": result})


@app.route("/api/osc/meetings", methods=["GET", "POST"])
@login_required
def osc_meetings_api():
    if request.method == "GET":
        q = (request.args.get("q") or "").strip()
        limit = max(1, min(500, int(request.args.get("limit") or "200")))
        if q:
            like = f"%{q}%"
            rows, _ = _osc_exec(
                """
                SELECT id, case_number, client_name, type, datetime, duration, location, notes, reminder, reminder_time, status, todo_id
                FROM meetings
                WHERE case_number LIKE %s OR client_name LIKE %s OR type LIKE %s OR notes LIKE %s
                ORDER BY datetime DESC, id DESC
                LIMIT %s
                """,
                (like, like, like, like, limit),
                fetch="all",
            )
        else:
            rows, _ = _osc_exec(
                """
                SELECT id, case_number, client_name, type, datetime, duration, location, notes, reminder, reminder_time, status, todo_id
                FROM meetings
                ORDER BY datetime DESC, id DESC
                LIMIT %s
                """,
                (limit,),
                fetch="all",
            )
        return jsonify({"ok": True, "items": rows})
    payload = request.get_json() or {}
    client_name = (payload.get("client_name") or "").strip()
    meeting_type = (payload.get("type") or "").strip()
    when = (payload.get("datetime") or "").strip()
    if not client_name or not meeting_type or not when:
        return jsonify({"ok": False, "error": "client_name/type/datetime required"}), 400
    when = when.replace("T", " ")
    cols = ["case_number", "client_name", "type", "datetime", "duration", "location", "notes", "reminder", "reminder_time", "status"]
    vals = [
        (payload.get("case_number") or "").strip() or None,
        client_name,
        meeting_type,
        when,
        int(payload.get("duration") or 60),
        (payload.get("location") or "").strip() or None,
        (payload.get("notes") or "").strip() or None,
        int(payload.get("reminder") if payload.get("reminder") is not None else 1),
        int(payload.get("reminder_time") or 30),
        (payload.get("status") or "scheduled").strip() or "scheduled",
    ]
    result, _ = _osc_exec(
        f"INSERT INTO meetings ({','.join(cols)}) VALUES ({','.join(['%s'] * len(cols))})",
        tuple(vals),
        fetch="none",
    )
    return jsonify({"ok": True, "result": result})


@app.route("/api/osc/meetings/<int:row_id>", methods=["GET", "PUT", "DELETE"])
@login_required
def osc_meeting_detail_api(row_id):
    if request.method == "GET":
        row, _ = _osc_exec("SELECT * FROM meetings WHERE id=%s", (row_id,), fetch="one")
        if not row:
            return jsonify({"ok": False, "error": "not found"}), 404
        return jsonify({"ok": True, "item": row})
    if request.method == "DELETE":
        result, _ = _osc_exec("DELETE FROM meetings WHERE id=%s", (row_id,), fetch="none")
        return jsonify({"ok": True, "result": result})
    payload = request.get_json() or {}
    allowed = ["case_number", "client_name", "type", "datetime", "duration", "location", "notes", "reminder", "reminder_time", "status", "todo_id"]
    sets = []
    vals = []
    for k in allowed:
        if k in payload:
            sets.append(f"{k}=%s")
            val = payload.get(k)
            if k == "datetime" and val:
                val = str(val).replace("T", " ")
            vals.append(val)
    if not sets:
        return jsonify({"ok": False, "error": "no fields"}), 400
    vals.append(row_id)
    result, _ = _osc_exec(f"UPDATE meetings SET {','.join(sets)} WHERE id=%s", tuple(vals), fetch="none")
    return jsonify({"ok": True, "result": result})


@app.route("/api/osc/todos", methods=["GET", "POST"])
@login_required
def osc_todos_api():
    if request.method == "GET":
        q = (request.args.get("q") or "").strip()
        limit = max(1, min(500, int(request.args.get("limit") or "200")))
        if q:
            like = f"%{q}%"
            rows, _ = _osc_exec(
                """
                SELECT id, case_number, client_name, todo_type, todo_date, todo_time, description, status, source_file, created_date, completed_date
                FROM case_todos
                WHERE case_number LIKE %s OR client_name LIKE %s OR todo_type LIKE %s OR description LIKE %s
                ORDER BY todo_date DESC, id DESC
                LIMIT %s
                """,
                (like, like, like, like, limit),
                fetch="all",
            )
        else:
            rows, _ = _osc_exec(
                """
                SELECT id, case_number, client_name, todo_type, todo_date, todo_time, description, status, source_file, created_date, completed_date
                FROM case_todos
                ORDER BY todo_date DESC, id DESC
                LIMIT %s
                """,
                (limit,),
                fetch="all",
            )
        return jsonify({"ok": True, "items": rows})
    payload = request.get_json() or {}
    case_number = (payload.get("case_number") or "").strip()
    todo_type = (payload.get("todo_type") or "").strip()
    if not case_number or not todo_type:
        return jsonify({"ok": False, "error": "case_number/todo_type required"}), 400
    cols = ["case_number", "client_name", "todo_type", "todo_date", "todo_time", "description", "status", "source_file"]
    vals = [
        case_number,
        (payload.get("client_name") or "").strip() or None,
        todo_type,
        (payload.get("todo_date") or "").strip() or None,
        (payload.get("todo_time") or "").strip() or None,
        (payload.get("description") or "").strip() or None,
        (payload.get("status") or "pending").strip() or "pending",
        (payload.get("source_file") or "").strip() or None,
    ]
    result, _ = _osc_exec(
        f"INSERT INTO case_todos ({','.join(cols)}) VALUES ({','.join(['%s'] * len(cols))})",
        tuple(vals),
        fetch="none",
    )
    return jsonify({"ok": True, "result": result})


@app.route("/api/osc/todos/<int:row_id>", methods=["GET", "PUT", "DELETE"])
@login_required
def osc_todo_detail_api(row_id):
    if request.method == "GET":
        row, _ = _osc_exec("SELECT * FROM case_todos WHERE id=%s", (row_id,), fetch="one")
        if not row:
            return jsonify({"ok": False, "error": "not found"}), 404
        return jsonify({"ok": True, "item": row})
    if request.method == "DELETE":
        result, _ = _osc_exec("DELETE FROM case_todos WHERE id=%s", (row_id,), fetch="none")
        return jsonify({"ok": True, "result": result})
    payload = request.get_json() or {}
    allowed = ["case_number", "client_name", "todo_type", "todo_date", "todo_time", "description", "status", "source_file", "google_calendar_id", "google_calendar_event_id"]
    sets = []
    vals = []
    for k in allowed:
        if k in payload:
            sets.append(f"{k}=%s")
            vals.append((payload.get(k) or "").strip() or None)
    if "status" in payload and str(payload.get("status")).strip().lower() == "completed":
        sets.append("completed_date=NOW()")
    if not sets:
        return jsonify({"ok": False, "error": "no fields"}), 400
    vals.append(row_id)
    result, _ = _osc_exec(f"UPDATE case_todos SET {','.join(sets)} WHERE id=%s", tuple(vals), fetch="none")
    return jsonify({"ok": True, "result": result})


@app.route("/api/osc/insights", methods=["GET", "POST"])
@login_required
def osc_insights_api():
    if request.method == "GET":
        q = (request.args.get("q") or "").strip().lower()
        case_number = (request.args.get("case_number") or "").strip().lower()
        case_reason = (request.args.get("case_reason") or "").strip().lower()
        limit = max(1, min(500, int(request.args.get("limit") or "300")))
        items = _osc_collect_insights()
        if q:
            def _hit(it):
                blob = " ".join(
                    [
                        str(it.get("title") or ""),
                        str(it.get("summary") or ""),
                        str(it.get("full_text") or ""),
                        str(it.get("case_number") or ""),
                        str(it.get("case_reason") or ""),
                        str(it.get("court") or ""),
                    ]
                ).lower()
                return q in blob
            items = [it for it in items if _hit(it)]
        if case_number:
            items = [it for it in items if case_number in str(it.get("case_number") or "").lower()]
        if case_reason:
            items = [
                it
                for it in items
                if case_reason in " ".join(
                    [
                        str(it.get("case_reason") or ""),
                        str(it.get("title") or ""),
                        str(it.get("summary") or ""),
                    ]
                ).lower()
            ]
        items = items[:limit]
        return jsonify({"ok": True, "items": items})
    payload = request.get_json() or {}
    insight_text = (payload.get("insight_text") or payload.get("full_text") or "").strip()
    if not insight_text:
        return jsonify({"ok": False, "error": "insight_text required"}), 400
    cols = ["case_number", "document_name", "court_reference", "court_type", "insight_type", "insight_text", "case_reason", "source_file", "raw_text"]
    vals = [
        (payload.get("case_number") or "").strip() or None,
        (payload.get("document_name") or payload.get("title") or "手動新增見解").strip(),
        (payload.get("court_reference") or payload.get("court") or "").strip() or None,
        (payload.get("court_type") or "").strip() or None,
        (payload.get("insight_type") or payload.get("source_type") or "manual").strip(),
        insight_text,
        (payload.get("case_reason") or "").strip() or None,
        (payload.get("source_file") or "").strip() or None,
        (payload.get("raw_text") or "").strip() or None,
    ]
    result, _ = _osc_exec(
        f"INSERT INTO legal_insights ({','.join(cols)}) VALUES ({','.join(['%s'] * len(cols))})",
        tuple(vals),
        fetch="none",
    )
    return jsonify({"ok": True, "result": result})


@app.route("/api/osc/insights/<insight_id>", methods=["GET"])
@login_required
def osc_insight_detail_api(insight_id):
    sid = (insight_id or "").strip()
    if sid.startswith("li-"):
        row_id = sid.split("-", 1)[1]
        row, _ = _osc_exec("SELECT * FROM legal_insights WHERE id=%s", (row_id,), fetch="one")
        if not row:
            return jsonify({"ok": False, "error": "not found"}), 404
        return jsonify({"ok": True, "item": row})
    if sid.startswith("cj-"):
        row_id = sid.split("-", 1)[1]
        row, _ = _osc_exec("SELECT * FROM court_judgments WHERE id=%s", (row_id,), fetch="one")
        if not row:
            return jsonify({"ok": False, "error": "not found"}), 404
        return jsonify({"ok": True, "item": row})
    # json-* fallback from merged set
    for it in _osc_collect_insights():
        if str(it.get("id")) == sid:
            return jsonify({"ok": True, "item": it})
    return jsonify({"ok": False, "error": "not found"}), 404


@app.route("/api/osc/insights/fetch-full", methods=["POST"])
@login_required
def osc_insights_fetch_full_api():
    payload = request.get_json() or {}
    url = (payload.get("url") or "").strip()
    raw_title = (payload.get("title") or "").strip()
    title = raw_title or "裁判見解全文"
    case_number = (payload.get("case_number") or "").strip() or None
    case_reason = (payload.get("case_reason") or "").strip() or None
    if not url and not case_number and not raw_title:
        return jsonify({"ok": False, "error": "url, title or case_number required"}), 400
    full_text = ""
    fallback_source = ""
    fetch_error = ""
    # Step 1: Try direct URL fetch (if URL provided)
    if url:
        fetched = _osc_fetch_url_text(url, timeout=30)
        if fetched.get("ok"):
            full_text = (fetched.get("text") or "").strip()
        else:
            fetch_error = fetched.get("error") or "fetch_failed"
    # Step 2: Try local DB lookup
    if not full_text:
        fallback = _osc_lookup_fulltext_fallback(title=title, case_number=case_number or "", url=url or "")
        if fallback.get("ok"):
            full_text = (fallback.get("text") or "").strip()
            fallback_source = str(fallback.get("source") or "")
    # Step 3: Try judgment-collector (collect + archive)
    if not full_text:
        jy = _osc_fetch_fulltext_from_judicial(
            title=title,
            case_number=case_number or "",
            case_reason=case_reason or "",
            timeout_sec=600,
        )
        if jy.get("ok"):
            full_text = (jy.get("text") or "").strip()
            fallback_source = str(jy.get("source") or "")
    if not full_text:
        error_detail = fetch_error or "all_sources_exhausted"
        return jsonify({"ok": False, "error": error_detail, "detail": "URL 抓取失敗、本地 DB 無紀錄、判決收集器也未找到結果。請確認 URL 正確或直接貼上全文。"}), 400
    actor_id = str(getattr(current_user, "id", "") or "osc_web")
    _ = actor_id  # reserved for future audit usage
    try:
        summary = _osc_summarize_legal_insight(full_text)
    except Exception as e:
        summary = f"摘要失敗：{e}"
    cols = ["case_number", "document_name", "court_reference", "insight_type", "insight_text", "case_reason", "source_file", "raw_text"]
    vals = [case_number, title, None, "web_fetch_fulltext", str(summary or "").strip(), case_reason, url, full_text]
    r, _ = _osc_exec(
        f"INSERT INTO legal_insights ({','.join(cols)}) VALUES ({','.join(['%s'] * len(cols))})",
        tuple(vals),
        fetch="none",
    )
    return jsonify(
        {
            "ok": True,
            "inserted": r,
            "item": {
                "source": "網頁全文擷取" if not fallback_source else f"網頁全文擷取（{fallback_source}）",
                "title": title,
                "case_number": case_number or "",
                "case_reason": case_reason or "",
                "url": url,
                "summary": str(summary or ""),
                "full_text": full_text,
            },
        }
    )


# KEEP: Canonical judgments endpoint. No active frontend callers found in templates,
# but may be used by OSC tools or external integrations. Supersedes judgments_legacy.
@app.route("/api/osc/judgments", methods=["GET"])
@login_required
def osc_judgments_compat_api():
    """
    Canonical judgments endpoint: returns merged insights from DB + judgments.json.
    """
    try:
        return jsonify(_osc_collect_insights())
    except Exception as e:
        logger.error(f"Error serving merged judgments: {e}")
        return jsonify([])


@app.route("/api/osc/forms/preview", methods=["POST"])
@login_required
def osc_forms_preview_api():
    payload = request.get_json() or {}
    form_type = (payload.get("form_type") or "").strip()
    if not form_type:
        return jsonify({"ok": False, "error": "form_type required"}), 400
    case_row = _osc_get_case_identity_by_payload(payload)
    fields = payload.get("fields") or {}
    if form_type == "legal_attest":
        content = fields.get("notes") or "(內文空白)"
        doc = (
            f"存證信函預覽\n\n"
            f"寄件人：{fields.get('sender_name')}\n"
            f"寄件地址：{fields.get('sender_addr')}\n"
            f"收件人：{fields.get('receiver_name')}\n"
            f"收件地址：{fields.get('receiver_addr')}\n"
            f"內文預覽：\n{content}\n\n（按下「匯出 WORD + PDF」即會產生符合郵局格式之對齊版式 PDF 歸檔）"
        )
        return jsonify({
            "ok": True,
            "case": case_row,
            "form_type": "legal_attest",
            "title": "存證信函草稿",
            "preview_text": doc,
            "suggested_filename": "legal_attest"
        })

    try:
        out = _osc_build_form_preview(form_type, case_row, fields if isinstance(fields, dict) else {})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 400
    return jsonify({"ok": True, "case": case_row, **out})


@app.route("/api/osc/forms/export", methods=["POST"])
@login_required
def osc_forms_export_api():
    try:
        _record_last_public_base_url()
    except Exception:
        logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 6187, exc_info=True)
    payload = request.get_json() or {}
    form_type = (payload.get("form_type") or "").strip()
    if not form_type:
        return jsonify({"ok": False, "error": "form_type required"}), 400
    case_row = _osc_get_case_identity_by_payload(payload)
    fields = payload.get("fields") or {}
    if form_type == "legal_attest":
        import uuid
        import os
        from skills.legal_attest.generator import core
        export_dir = f"{_MAGI_ROOT}/exports"
        os.makedirs(export_dir, exist_ok=True)
        filename_base = f"legal_attest_{uuid.uuid4().hex[:8]}"
        pdf_path = os.path.join(export_dir, f"{filename_base}.pdf")
        
        sender_name_list = [[fields.get("sender_name") or ""]]
        sender_addr_list = [fields.get("sender_addr") or ""]
        receiver_name_list = [[fields.get("receiver_name") or ""]]
        receiver_addr_list = [fields.get("receiver_addr") or ""]
        content = fields.get("notes") or "(內文空白)"
        
        try:
            core.generate_text_and_letter(
                sender_name_list, sender_addr_list,
                receiver_name_list, receiver_addr_list,
                [], [],
                content
            )
            core.merge_text_and_letter(pdf_path)
            core.clean_temp_files()
        except Exception as e:
            return jsonify({"ok": False, "error": f"產生存證信函失敗: {e}"}), 500
            
        public_url = f"{_get_public_base_url()}/exports/{filename_base}.pdf"
        doc = (
            f"存證信函已產出！\n\n"
            f"寄件人：{fields.get('sender_name')}\n"
            f"收件人：{fields.get('receiver_name')}\n"
            f"內文預覽：\n{content}"
        )
        return jsonify(
            {
                "ok": True,
                "case": case_row,
                "form_type": "legal_attest",
                "title": "存證信函預覽",
                "preview_text": doc,
                "export": {"success": True},
                "export_pdf": {"success": True, "url": public_url},
                "export_docx": {"success": False},
                "export_errors": [],
            }
        )

    try:
        out = _osc_build_form_preview(form_type, case_row, fields if isinstance(fields, dict) else {})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 400
        
    actual_form_type = out.get("form_type")
    if actual_form_type in ["power_of_attorney", "receipt", "contract"]:
        import uuid
        import os
        import time
        from api.osc_document_generator import generate_receipt, generate_poa, generate_engagement_agreement
        
        export_dir = f"{_MAGI_ROOT}/exports"
        os.makedirs(export_dir, exist_ok=True)
        stamp = time.strftime("%Y%m%d_%H%M%S")
        token = uuid.uuid4().hex[:8]
        filename_base = f"{actual_form_type}_{stamp}_{token}"
        docx_path = os.path.join(export_dir, f"{filename_base}.docx")
        
        data = dict(case_row)
        for k, v in (fields if isinstance(fields, dict) else {}).items():
            if v: data[k] = v
            
        data['案號'] = data.get('court_case_no', '')
        data['股別'] = data.get('court_branch', '')
        data['委任人/當事人'] = data.get('client_name', '')
        data['案由/事件'] = data.get('case_reason', '')
        data['受任律師'] = data.get('lawyer_name', '')
        data['通訊地址'] = data.get('address', '')
        data['聯絡電話'] = data.get('phone', '')
        data['身分證字號'] = data.get('tax_id', '')
        data['委任範圍'] = data.get('item', '')
        data['金額'] = data.get('amount', '')
        data['委任費用(數字)'] = data.get('amount', '')
        data['法院/檢察署'] = data.get('court_name', '')
        data['取代日期'] = data.get('date', '')

        config = {}
        try:
            # Fallback mappings for some standard names
            config['company_name'] = '偵理法律事務所'
            config['default_lawyer'] = '喬政翔律師'
        except Exception:
            logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 6285, exc_info=True)

        try:
            if actual_form_type == "receipt":
                doc = generate_receipt(data, data.get('item') or '法律服務費', config)
            elif actual_form_type == "power_of_attorney":
                case_type = '民事'
                role = '代理人'
                cat = str(data.get('case_category', ''))
                if '刑' in cat:
                    case_type = '刑事'
                    role = '辯護人' if '被告' in str(data.get('client_role', '')) else '告訴代理人'
                elif '行' in cat:
                    case_type = '行政'
                doc = generate_poa(data, case_type, role, config)
            elif actual_form_type == "contract":
                doc = generate_engagement_agreement(data, config)
                
            doc.save(docx_path)
            docx_meta = _export_file_meta(docx_path)
            exported = {
                "success": docx_meta.get("success"),
                "export": docx_meta,
                "export_docx": docx_meta,
                "export_pdf": {"success": False, "error": "pdf_conversion_skip"},
                "errors": [] if docx_meta.get("success") else [{"type": "docx", "error": docx_meta.get("error")}]
            }
        except Exception as e:
            exported = {"success": False, "errors": [{"type": "generator", "error": str(e)}], "export_docx": {}, "export_pdf": {}}
    else:
        exported = _export_osc_form_files(
            out.get("title") or out.get("form_type") or "OSC 文件",
            out.get("preview_text") or "",
            out.get("suggested_filename") or "osc_form",
        )
    return jsonify(
        {
            "ok": bool(exported.get("success")),
            "case": case_row,
            **out,
            "export": exported.get("export") or {"success": False},
            "export_docx": exported.get("export_docx") or {"success": False},
            "export_pdf": exported.get("export_pdf") or {"success": False},
            "export_errors": exported.get("errors") or [],
        }
    )


@app.route("/api/osc/laf-wizard/run", methods=["POST"])
@login_required
def osc_laf_wizard_run_api():
    payload = request.get_json() or {}
    mode = (payload.get("mode") or "preview").strip().lower()
    if mode not in {"preview", "draft", "submit"}:
        return jsonify({"ok": False, "error": "mode must be preview|draft|submit"}), 400
    action = _osc_map_laf_action(payload.get("action") or "")
    if action not in {"go_live", "inquiry", "fee", "condition", "withdrawal", "closing"}:
        return jsonify({"ok": False, "error": "unsupported action"}), 400
    if mode == "submit" and (not getattr(current_user, "is_admin", lambda: False)()):
        return jsonify({"ok": False, "error": "admin_required_for_submit"}), 403

    ident = _osc_prepare_laf_identity(payload)
    fields = payload.get("fields") or {}
    if not isinstance(fields, dict):
        fields = {}
    reason = str(payload.get("reason") or "").strip()
    try:
        LAFOrchestrator = _osc_import_laf_orchestrator()
        orchestrator_inst = LAFOrchestrator(dry_run=(mode == "preview"))
        if mode == "submit":
            result = orchestrator_inst.execute_portal_action_submit(
                action=action,
                laf_case_number=ident["laf_case_number"],
                case_number=ident["case_number"],
                client_name=ident["client_name"],
                reason=reason,
                fields=fields,
            )
        else:
            result = orchestrator_inst.execute_portal_action_draft(
                action=action,
                laf_case_number=ident["laf_case_number"],
                case_number=ident["case_number"],
                client_name=ident["client_name"],
                reason=reason,
                fields=fields,
            )
        artifact = _osc_enrich_portal_preview(orchestrator_inst._last_portal_artifact if hasattr(orchestrator_inst, "_last_portal_artifact") else {})
        return jsonify(
            {
                "ok": bool(isinstance(result, dict) and result.get("ok")),
                "mode": mode,
                "action": action,
                "identity": ident,
                "result": result if isinstance(result, dict) else {"ok": bool(result)},
                "artifact": artifact,
            }
        )
    except Exception as e:
        return jsonify({"ok": False, "error": str(e), "mode": mode, "action": action, "identity": ident}), 500


@app.route("/api/osc/laf-backfill", methods=["POST"])
@login_required
def osc_laf_backfill_api():
    """手動觸發法扶案號補填（資料夾 + 接案清冊）。"""
    try:
        import sys
        sys.path.insert(0, os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "scripts"))
        from laf_nightly_audit import run_backfill_only
        result = run_backfill_only(notify=False)
        return jsonify(result)
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/api/osc/archive-wizard/preview", methods=["GET"])
@login_required
def osc_archive_wizard_preview_api():
    limit = max(1, min(1000, int(request.args.get("limit") or "300")))
    try:
        out = _osc_build_archive_preview(limit=limit)
        return jsonify(out)
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/api/osc/archive-wizard/execute", methods=["POST"])
@login_required
def osc_archive_wizard_execute_api():
    payload = request.get_json() or {}
    if not bool(payload.get("confirm")):
        return jsonify({"ok": False, "error": "confirm_required"}), 400
    force = bool(payload.get("force"))
    case_ids = payload.get("case_ids") or []
    if isinstance(case_ids, str):
        case_ids = [x.strip() for x in case_ids.split(",") if x.strip()]
    case_ids = [str(x).strip() for x in case_ids if str(x).strip()]

    preview = _osc_build_archive_preview(limit=1000)
    items = preview.get("items") or []
    pick = [it for it in items if (not case_ids) or (str(it.get("id")) in set(case_ids))]
    moved = []
    skipped = []
    errors = []

    for it in pick:
        cid = str(it.get("id") or "").strip()
        src = str(it.get("source_local") or "").strip()
        dst = str(it.get("target_local") or "").strip()
        if not src or not os.path.exists(src):
            skipped.append({"id": cid, "case_number": it.get("case_number"), "reason": "source_missing"})
            continue
        if not dst:
            skipped.append({"id": cid, "case_number": it.get("case_number"), "reason": "target_missing"})
            continue
        if os.path.exists(dst) and not force:
            skipped.append({"id": cid, "case_number": it.get("case_number"), "reason": "target_exists"})
            continue
        try:
            os.makedirs(os.path.dirname(dst), exist_ok=True)
            if os.path.abspath(src) != os.path.abspath(dst):
                shutil.move(src, dst)
            _osc_exec("UPDATE cases SET folder_path=%s, updated_at=NOW() WHERE id=%s", (dst, cid), fetch="none")
            moved.append({"id": cid, "case_number": it.get("case_number"), "from": src, "to": dst})
        except Exception as e:
            errors.append({"id": cid, "case_number": it.get("case_number"), "error": str(e)})

    return jsonify(
        {
            "ok": not errors,
            "summary": {"selected": len(pick), "moved": len(moved), "skipped": len(skipped), "errors": len(errors)},
            "moved": moved,
            "skipped": skipped,
            "errors": errors,
        }
    )

# ── 勞動基準法計算器 ────────────────────────────────────────────────────────────
@app.route("/api/osc/labor-law/calc", methods=["POST"])
@login_required
def osc_labor_law_calc():
    """
    勞動基準法計算器 API。
    支援：
      - 加班費（單日）
      - 特休假天數
      - 資遣費
      - 本地 xlsx/pdf 批次計算（以 file_paths 傳入本機路徑清單）
      - 檔案上傳計算（multipart/form-data with files[]）

    Request JSON:
      {
        "task": "自然語言指令，例如：月薪45000，休息日加班3小時",
        "monthly_wage": 45000,
        "monthly_wage_by_year": {"2021": 45000, "2022": 47000},
        "file_paths": ["/path/to/file.xlsx", "/path/to/holiday.pdf"],
        "mode": "overtime|annual_leave|severance|calc_file|calc_dir"
      }

    或 multipart/form-data:
      - task (str)
      - monthly_wage (float)
      - files[] (file uploads)
    """
    import sys
    import tempfile
    skill_path = os.path.join(
        os.path.dirname(__file__), "..", "skills", "labor-law-calculator", "action.py"
    )
    skill_dir = os.path.dirname(skill_path)
    if skill_dir not in sys.path:
        sys.path.insert(0, skill_dir)

    try:
        import importlib.util
        spec = importlib.util.spec_from_file_location("labor_law_action", os.path.abspath(skill_path))
        mod = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(mod)
    except Exception as e:
        return jsonify({"ok": False, "error": f"無法載入 skill：{e}"}), 500

    # Parse request – support both JSON and multipart
    uploaded_paths: list = []
    if request.content_type and "multipart" in request.content_type:
        task = request.form.get("task", "")
        try:
            monthly_wage = float(request.form.get("monthly_wage") or 0) or None
        except Exception:
            monthly_wage = None
        wage_by_year_raw = request.form.get("monthly_wage_by_year")
        # Save uploaded files to temp dir
        temp_dir = tempfile.mkdtemp(prefix="labor_law_")
        for f in request.files.getlist("files[]") + request.files.getlist("file"):
            dest = os.path.join(temp_dir, f.filename)
            f.save(dest)
            uploaded_paths.append(dest)
    else:
        data = request.get_json() or {}
        task = data.get("task", "")
        try:
            monthly_wage = float(data.get("monthly_wage") or 0) or None
        except Exception:
            monthly_wage = None
        wage_by_year_raw = data.get("monthly_wage_by_year")
        uploaded_paths = [str(p) for p in (data.get("file_paths") or [])]

    # Parse wage_by_year
    wage_by_year = None
    if wage_by_year_raw:
        try:
            raw = wage_by_year_raw if isinstance(wage_by_year_raw, dict) else json.loads(wage_by_year_raw)
            wage_by_year = {int(k): float(v) for k, v in raw.items()}
        except Exception:
            logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 6525, exc_info=True)

    kwargs = {}
    if monthly_wage:
        kwargs["monthly_wage"] = monthly_wage
    if wage_by_year:
        kwargs["monthly_wage_by_year"] = wage_by_year
    if uploaded_paths:
        kwargs["file_paths"] = uploaded_paths
        kwargs.setdefault("mode", "calc_file")

    try:
        result_text = mod.run(task, **kwargs)
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500

    return jsonify({"ok": True, "result": result_text})


@app.route("/api/osc/labor-law/parse-files", methods=["POST"])
@login_required
def osc_labor_law_parse_files():
    """
    解析指定路徑的出勤 Excel/PDF，回傳每日加班明細（不計算金額）。
    Request JSON: {"file_paths": [...], "monthly_wage": 45000}
    """
    import sys
    skill_path = os.path.join(
        os.path.dirname(__file__), "..", "skills", "labor-law-calculator", "action.py"
    )
    try:
        import importlib.util
        spec = importlib.util.spec_from_file_location("labor_law_action", os.path.abspath(skill_path))
        mod = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(mod)
    except Exception as e:
        return jsonify({"ok": False, "error": f"無法載入 skill：{e}"}), 500

    data = request.get_json() or {}
    file_paths = [str(p) for p in (data.get("file_paths") or [])]
    monthly_wage = float(data.get("monthly_wage") or 0) or None

    if not file_paths:
        return jsonify({"ok": False, "error": "請提供 file_paths"}), 400

    all_records = []
    errors = []
    for fp in file_paths:
        ext = fp.lower().split(".")[-1]
        try:
            if ext in ("xlsx", "xls"):
                recs = mod._parse_attendance_excel(fp)
            elif ext == "pdf":
                recs = mod._parse_holiday_pdf(fp)
            else:
                errors.append(f"不支援：{fp}")
                continue
            all_records.extend([{
                "date": r.date_str,
                "weekday": r.weekday,
                "day_type": r.day_type,
                "pre_ot_min": r.pre_ot_min,
                "post_ot_min": r.post_ot_min,
                "total_ot_min": r.total_ot_min,
                "source": r.source,
                "note": r.note,
                "ot_pay": mod._calc_ot_pay_for_record(r, monthly_wage) if monthly_wage else None,
            } for r in recs])
        except Exception as e:
            errors.append(f"{fp}: {e}")

    return jsonify({
        "ok": True,
        "total_records": len(all_records),
        "total_ot_hours": round(sum(r["total_ot_min"] for r in all_records) / 60, 2),
        "total_ot_pay": round(sum(r["ot_pay"] for r in all_records if r["ot_pay"]), 2) if monthly_wage else None,
        "records": all_records,
        "errors": errors,
    })


@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        return callback()
    return redirect(url_for('dashboard'))

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        try:
            from api.db_helper import get_cursor
            with get_cursor(config=DB_CONFIG, dictionary=True) as (_conn, cursor):
                cursor.execute("SELECT * FROM users WHERE username = %s", (username,))
                user_data = cursor.fetchone()

                if user_data and check_password_hash(user_data['password_hash'], password):
                    user = User(user_data['id'], user_data['username'], user_data['role'])
                    login_user(user)
                    return redirect(url_for('dashboard'))
                else:
                    flash('Invalid username or password')
        except Exception as e:
            flash(f"Login Error: {str(e)}")

    return render_template('login.html')

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        hashed_pw = generate_password_hash(password)
        
        try:
            from api.db_helper import get_cursor
            with get_cursor(config=DB_CONFIG) as (conn, cursor):
                cursor.execute("SELECT COUNT(*) FROM users")
                count = cursor.fetchone()[0]
                role = 'admin' if count == 0 else 'user'

                cursor.execute("INSERT INTO users (username, password_hash, role) VALUES (%s, %s, %s)",
                               (username, hashed_pw, role))
                conn.commit()

            flash(f'Registration successful! You are now an {role}. Please login.')
            return redirect(url_for('login'))
        except mysql.connector.Error as err:
            flash(f"Error: {err}")
            
    return render_template('register.html')

@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('login'))

# --- System Test & Self-Repair API ---

@app.route('/api/system-test', methods=['POST'])
@login_required
def api_system_test():
    """Run comprehensive system health check."""
    try:
        from skills.ops.system_test import run_all_tests
        report = run_all_tests()
        return jsonify(report)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/self-repair', methods=['POST'])
@login_required
def api_self_repair():
    """Trigger self-repair for failed test items."""
    try:
        data = request.get_json() or {}
        targets = data.get("targets")  # optional list of IDs
        import importlib.util
        base_dir = os.path.join(os.path.dirname(__file__), "..", "skills")
        candidates = [
            os.path.join(base_dir, "magi-self-repair", "action.py"),
            os.path.join(base_dir, "magi-doctor", "action.py"),
        ]
        repair_mod = None
        for action_path in candidates:
            if not os.path.exists(action_path):
                continue
            spec = importlib.util.spec_from_file_location("magi_self_repair", action_path)
            if spec is None or spec.loader is None:
                continue
            repair_mod = importlib.util.module_from_spec(spec)
            spec.loader.exec_module(repair_mod)
            break
        if repair_mod is None:
            raise FileNotFoundError(
                "No self-repair module found. Tried: " + ", ".join(candidates)
            )
        if not hasattr(repair_mod, "repair_targets"):
            raise AttributeError("self-repair module missing repair_targets()")
        report = repair_mod.repair_targets(targets)
        return jsonify(report)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/api/nerv/skill-interview', methods=['GET'])
def api_nerv_skill_interview_status():
    auth_error = _require_json_auth()
    if auth_error:
        return auth_error
    try:
        state = orchestrator.get_skill_interview_state(_nerv_skill_interview_user_id(), "NERV")
        return jsonify({
            "ok": True,
            "can_edit": bool(getattr(current_user, "is_admin", False)),
            "interview": state,
        })
    except Exception as e:
        logger.error("NERV skill interview status failed: %s", e, exc_info=True)
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route('/api/nerv/skill-interview/start', methods=['POST'])
def api_nerv_skill_interview_start():
    auth_error = _require_json_auth(admin=True)
    if auth_error:
        return auth_error
    payload = request.get_json(silent=True) or {}
    initial_request = str(payload.get("request") or "").strip()
    if not initial_request:
        return jsonify({"ok": False, "error": "empty_request"}), 400
    try:
        message = orchestrator.start_skill_interview(
            _nerv_skill_interview_user_id(),
            "NERV",
            getattr(current_user, "role", "user"),
            initial_request,
            trigger_reason="manual",
        )
        state = orchestrator.get_skill_interview_state(_nerv_skill_interview_user_id(), "NERV")
        return jsonify({
            "ok": True,
            "message": message,
            "interview": state,
        })
    except Exception as e:
        logger.error("NERV skill interview start failed: %s", e, exc_info=True)
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route('/api/nerv/skill-interview/reply', methods=['POST'])
def api_nerv_skill_interview_reply():
    auth_error = _require_json_auth(admin=True)
    if auth_error:
        return auth_error
    payload = request.get_json(silent=True) or {}
    reply_text = str(payload.get("message") or "").strip()
    if not reply_text:
        return jsonify({"ok": False, "error": "empty_message"}), 400
    try:
        handled, message = orchestrator.reply_skill_interview(
            _nerv_skill_interview_user_id(),
            "NERV",
            getattr(current_user, "role", "user"),
            reply_text,
        )
        if not handled:
            return jsonify({"ok": False, "error": "no_active_interview"}), 400
        state = orchestrator.get_skill_interview_state(_nerv_skill_interview_user_id(), "NERV")
        finalized = (not state.get("active")) and ("新 SKILL 已建立並啟用" in str(message or ""))
        cancelled = (not state.get("active")) and ("已取消這次 SKILL 訪談" in str(message or ""))
        return jsonify({
            "ok": True,
            "message": message,
            "interview": state,
            "finalized": finalized,
            "cancelled": cancelled,
            "skill_name": _extract_interview_skill_name(message),
        })
    except Exception as e:
        logger.error("NERV skill interview reply failed: %s", e, exc_info=True)
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route('/api/skills/interview-history', methods=['GET'])
def api_skill_interview_history():
    auth_error = _require_json_auth()
    if auth_error:
        return auth_error
    limit = request.args.get('limit', default=10, type=int) or 10
    limit = max(1, min(limit, 50))
    try:
        from skills.management.skill_interview import list_interview_history

        return jsonify({"ok": True, "history": list_interview_history(limit=limit)})
    except Exception as e:
        logger.error("Skill interview history failed: %s", e, exc_info=True)
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route('/api/skills/<skill_name>/versions', methods=['GET'])
def api_skill_versions(skill_name):
    auth_error = _require_json_auth()
    if auth_error:
        return auth_error
    try:
        _skill_doc_path(skill_name)
    except ValueError as e:
        return jsonify({"ok": False, "error": str(e)}), 400
    try:
        from skills.evolution.skill_genesis import list_skill_versions

        result = list_skill_versions(str(skill_name).strip())
        if not result.get("success"):
            return jsonify({"ok": False, "error": result.get("error") or "versions_unavailable"}), 404
        return jsonify({"ok": True, "versions": result.get("versions") or []})
    except Exception as e:
        logger.error("Skill versions failed for %s: %s", skill_name, e, exc_info=True)
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route('/api/skills/<skill_name>/rollback', methods=['POST'])
def api_skill_rollback(skill_name):
    auth_error = _require_json_auth(admin=True)
    if auth_error:
        return auth_error
    try:
        _skill_doc_path(skill_name)
    except ValueError as e:
        return jsonify({"ok": False, "error": str(e)}), 400
    payload = request.get_json(silent=True) or {}
    version_id = str(payload.get("version_id") or "").strip()
    try:
        from skills.evolution.skill_genesis import rollback_skill_version
        from skills.bridge.embedding_router import get_router
        import skills.bridge.semantic_router as semantic_router

        result = rollback_skill_version(str(skill_name).strip(), version_id=version_id)
        if not result.get("success"):
            return jsonify({"ok": False, "error": result.get("error") or "rollback_failed"}), 400
        try:
            router = get_router()
            if router.is_ready:
                router.rebuild_cache()
            else:
                router.initialize()
        except Exception:
            logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 6855, exc_info=True)
        try:
            semantic_router._SKILLS_CACHE = None
            semantic_router._SKILLS_CACHE_TS = 0.0
        except Exception:
            logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 6860, exc_info=True)
        return jsonify({"ok": True, "result": result})
    except Exception as e:
        logger.error("Skill rollback failed for %s: %s", skill_name, e, exc_info=True)
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route('/api/nerv/skills', methods=['GET'])
def api_nerv_skills():
    auth_error = _require_json_auth()
    if auth_error:
        return auth_error
    try:
        return jsonify({"ok": True, "skills": _list_skill_docs()})
    except Exception as e:
        logger.error("NERV skill list failed: %s", e, exc_info=True)
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route('/api/nerv/product-runtime', methods=['GET', 'POST'])
def api_nerv_product_runtime():
    auth_error = _require_json_auth(admin=request.method == 'POST')
    if auth_error:
        return auth_error

    if request.method == 'GET':
        try:
            return jsonify(_nerv_product_runtime_payload())
        except Exception as e:
            logger.error("NERV product runtime load failed: %s", e, exc_info=True)
            return jsonify({"ok": False, "error": str(e)}), 500

    payload = request.get_json(silent=True) or {}
    product = str(payload.get("product") or "").strip().lower()
    if product not in NERV_PRODUCT_NAMES:
        return jsonify({"ok": False, "error": "unsupported_product"}), 400

    allowed_keys = {"codex_mode"}
    if product == "laf":
        allowed_keys |= {"portal_env", "prod_base_url", "test_base_url", "compare_base_url"}

    updates = {}
    for key in allowed_keys:
        value = payload.get(key)
        if value is None:
            continue
        text = str(value).strip()
        if not text:
            continue
        updates[key] = text

    if not updates:
        return jsonify({"ok": False, "error": "empty_updates"}), 400

    try:
        updated = update_product_runtime(product, **updates)
        response = _nerv_product_runtime_payload()
        response["updated_product"] = product
        response["updated_profile"] = updated
        return jsonify(response)
    except Exception as e:
        logger.error("NERV product runtime save failed: %s", e, exc_info=True)
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route('/api/nerv/skills/<skill_name>', methods=['GET', 'POST'])
def api_nerv_skill_detail(skill_name):
    auth_error = _require_json_auth(admin=request.method != 'GET')
    if auth_error:
        return auth_error

    try:
        skill_doc = _skill_doc_path(skill_name)
        action_file = _skill_action_path(skill_name)
    except ValueError as e:
        return jsonify({"ok": False, "error": str(e)}), 400

    if request.method == 'GET':
        exists = skill_doc.exists()
        content = ""
        if exists:
            try:
                content = skill_doc.read_text(encoding="utf-8")
            except Exception as e:
                return jsonify({"ok": False, "error": f"read_failed: {e}"}), 500
        updated_at = ""
        stat_target = skill_doc if exists else action_file
        if stat_target.exists():
            try:
                updated_at = datetime.fromtimestamp(stat_target.stat().st_mtime).isoformat()
            except Exception:
                updated_at = ""
        return jsonify(
            {
                "ok": True,
                "skill": {
                    "name": str(skill_name).strip(),
                    "content": content,
                    "has_skill_doc": exists,
                    "has_action": action_file.exists(),
                    "updated_at": updated_at,
                    "summary": _skill_summary(content),
                },
            }
        )

    payload = request.get_json(silent=True) or {}
    content = str(payload.get("content") or "")
    if not content.strip():
        return jsonify({"ok": False, "error": "empty_skill_content"}), 400

    try:
        skill_doc.parent.mkdir(parents=True, exist_ok=True)
        normalized = content.replace("\r\n", "\n")
        if not normalized.endswith("\n"):
            normalized += "\n"
        skill_doc.write_text(normalized, encoding="utf-8")
    except Exception as e:
        logger.error("NERV skill save failed for %s: %s", skill_name, e, exc_info=True)
        return jsonify({"ok": False, "error": f"save_failed: {e}"}), 500

    return jsonify(
        {
            "ok": True,
            "saved": True,
            "skill": {
                "name": str(skill_name).strip(),
                "content": normalized,
                "has_skill_doc": True,
                "has_action": action_file.exists(),
                "updated_at": datetime.now().isoformat(),
                "summary": _skill_summary(normalized),
            },
        }
    )


@app.route('/api/codex-distributed/status', methods=['GET'])
def api_codex_distributed_status():
    auth_error = _require_json_auth()
    if auth_error:
        return auth_error
    try:
        from skills.bridge.llm_direct import public_status_report

        return jsonify({"status": public_status_report(), "can_toggle": current_user.is_admin()})
    except Exception as e:
        logger.error("Codex distributed status failed: %s", e, exc_info=True)
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route('/api/codex-distributed/toggle', methods=['POST'])
def api_codex_distributed_toggle():
    auth_error = _require_json_auth(admin=True)
    if auth_error:
        return auth_error
    try:
        from skills.bridge.llm_direct import apply_manual_command, public_status_report

        payload = request.get_json(silent=True) or {}
        command = str(payload.get("command") or "").strip().lower()
        features = payload.get("features")
        apply_manual_command(command, features=features)
        return jsonify({"status": public_status_report(), "can_toggle": True})
    except ValueError as e:
        return jsonify({"ok": False, "error": str(e)}), 400
    except Exception as e:
        logger.error("Codex distributed toggle failed: %s", e, exc_info=True)
        return jsonify({"ok": False, "error": str(e)}), 500


# --- Status API ---
@app.route('/api/status')
def api_status():
    """Returns MAGI node status as JSON for dashboard polling."""
    status_file = os.path.join(os.path.dirname(__file__), '..', 'static', 'magi_status.json')
    try:
        with open(status_file, 'r') as f:
            data = json.load(f)
        return data
    except Exception as e:
        return {"error": str(e)}, 500

@app.route('/api/live-log')
@login_required
def api_live_log():
    """Return recent server log lines for dashboard live-log panel."""
    limit = min(int(request.args.get('limit', 40)), 100)
    log_path = os.path.join(os.path.dirname(__file__), '..', '.agent', 'server.log')
    lines = []
    try:
        with open(log_path, 'rb') as f:
            f.seek(0, 2)
            size = f.tell()
            read_size = min(size, 32768)
            f.seek(size - read_size)
            raw = f.read().decode('utf-8', errors='replace')
            all_lines = raw.strip().splitlines()
            lines = all_lines[-limit:]
    except Exception as e:
        lines = [f"[LOG READ ERROR] {e}"]
    return jsonify({'lines': lines})

@app.route("/callback", methods=['GET', 'POST'])
@app.route("/line/webhook", methods=['GET', 'POST'])
def callback():
    # LINE webhook verification may use GET probes.
    if request.method == "GET":
        return "OK", 200

    if _check_rate_limit("webhook"):
        logger.warning(f"⚠️ Rate limit exceeded for LINE webhook from {request.remote_addr}")
        return "Too Many Requests", 429

    if not LINE_BOT_ENABLED:
        return "LINE webhook disabled: missing credentials", 503

    # Record base URL for exporting long responses as downloadable TXT links.
    try:
        _record_last_public_base_url()
    except Exception:
        logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 7081, exc_info=True)

    # get X-Line-Signature header value
    signature = request.headers.get('X-Line-Signature')
    if not signature:
        logger.error("Missing X-Line-Signature header.")
        abort(400)

    # get request body as text
    body = request.get_data(as_text=True)
    try:
        ua = (request.headers.get("User-Agent") or "").strip()
        xff = (request.headers.get("X-Forwarded-For") or "").strip()
        path = (request.path or "").strip()
        logger.info(f"LINE callback received ({len(body)} bytes) path={path!r} ua={ua!r} xff={xff!r}")
    except Exception:
        logger.info(f"LINE callback received ({len(body)} bytes).")

    # handle webhook body
    try:
        handler.handle(body, signature)
    except InvalidSignatureError:
        logger.error("Invalid signature. Please check LINE_CHANNEL_SECRET / MAGI_LINE_CHANNEL_SECRET.")
        abort(400)
    try:
        _record_last_line_callback(request.path or "/callback")
    except Exception:
        logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 7108, exc_info=True)

    return 'OK'

TELEGRAM_CONTEXT_TTL_SECONDS = int(os.environ.get("MAGI_TELEGRAM_CONTEXT_TTL_SECONDS", "300") or "300")
_TELEGRAM_SEEN_UPDATES: dict[int, int] = {}


def _load_telegram_webhook_secret() -> str:
    """讀取 TG webhook secret — 純環境變數，不再依賴 openclaw.json。"""
    return (
        os.environ.get("TELEGRAM_WEBHOOK_SECRET")
        or os.environ.get("OPENCLAW_TELEGRAM_WEBHOOK_SECRET")
        or ""
    ).strip()


def _telegram_verify_webhook_secret() -> bool:
    expected = _load_telegram_webhook_secret()
    if not expected:
        return True
    received = (request.headers.get("X-Telegram-Bot-Api-Secret-Token") or "").strip()
    return bool(received) and hmac.compare_digest(received, expected)


def _telegram_mark_seen_update(update_id: int | None) -> bool:
    if update_id is None:
        return False
    now = int(time.time())
    # prune old ids
    stale_before = now - 3600
    for k, ts in list(_TELEGRAM_SEEN_UPDATES.items()):
        if ts < stale_before:
            _TELEGRAM_SEEN_UPDATES.pop(k, None)
    if update_id in _TELEGRAM_SEEN_UPDATES:
        return True
    _TELEGRAM_SEEN_UPDATES[update_id] = now
    return False


def _telegram_api_post(token: str, method: str, payload: dict | None = None, files: dict | None = None):
    try:
        from skills.bridge.http_pool import get_session
        sess = get_session()
        url = f"https://api.telegram.org/bot{token}/{method}"
        if files:
            return sess.post(url, data=payload or {}, files=files, timeout=20)
        return sess.post(url, json=payload or {}, timeout=20)
    except Exception:
        return None


def _audit_preview(text: str, limit: int = 180) -> str:
    s = " ".join(str(text or "").strip().split())
    if len(s) <= limit:
        return s
    return s[:limit] + "..."


def _audit_sha1(text: str) -> str:
    return hashlib.sha1(str(text or "").encode("utf-8", "ignore")).hexdigest()


def _append_channel_delivery_audit(event: dict) -> None:
    try:
        payload = {"ts": time.time()}
        payload.update(event or {})
        with _channel_audit_lock:
            with open(_CHANNEL_DELIVERY_AUDIT_FILE, "a", encoding="utf-8") as f:
                f.write(json.dumps(payload, ensure_ascii=False) + "\n")
            # Auto-prune: keep last 30K lines when file exceeds 5MB
            try:
                if os.path.getsize(_CHANNEL_DELIVERY_AUDIT_FILE) > 5 * 1024 * 1024:
                    with open(_CHANNEL_DELIVERY_AUDIT_FILE, "r", encoding="utf-8") as f:
                        lines = f.readlines()
                    with open(_CHANNEL_DELIVERY_AUDIT_FILE, "w", encoding="utf-8") as f:
                        f.writelines(lines[-30000:])
            except Exception:
                logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 7192, exc_info=True)
    except Exception:
        logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 7194, exc_info=True)


def _telegram_send_text_to(chat_id: str, text: str, reply_to_message_id: int | None = None) -> bool:
    token = _load_openclaw_telegram_token()
    if not token or not str(chat_id or "").strip():
        return False
    ok_all = True
    for chunk in _chunk_text_for_line(text, limit=3900):
        payload = {
            "chat_id": str(chat_id),
            "text": chunk,
        }
        if reply_to_message_id:
            payload["reply_to_message_id"] = int(reply_to_message_id)
        resp = _telegram_api_post(token, "sendMessage", payload=payload)
        ok = bool(resp and resp.status_code == 200)
        message_id = None
        try:
            data = resp.json() if resp is not None else {}
            if isinstance(data, dict):
                message_id = ((data.get("result") or {}) or {}).get("message_id")
        except Exception:
            message_id = None
        _append_channel_delivery_audit(
            {
                "platform": "telegram",
                "kind": "text",
                "chat_id": str(chat_id),
                "reply_to_message_id": int(reply_to_message_id) if reply_to_message_id else None,
                "message_id": int(message_id) if message_id else None,
                "ok": ok,
                "status_code": int(resp.status_code) if resp is not None else None,
                "text_sha1": _audit_sha1(chunk),
                "preview": _audit_preview(chunk),
            }
        )
        if not ok:
            ok_all = False
    return ok_all


def _telegram_send_document(chat_id: str, file_path: str, caption: str = "", reply_to_message_id: int | None = None) -> bool:
    token = _load_openclaw_telegram_token()
    p = (file_path or "").strip()
    if not token or not str(chat_id or "").strip() or (not p) or (not os.path.exists(p)):
        return False
    payload = {"chat_id": str(chat_id), "caption": (caption or "")[:900]}
    if reply_to_message_id:
        payload["reply_to_message_id"] = int(reply_to_message_id)
    try:
        with open(p, "rb") as f:
            files = {"document": (os.path.basename(p), f)}
            resp = _telegram_api_post(token, "sendDocument", payload=payload, files=files)
        ok = bool(resp and resp.status_code == 200)
        message_id = None
        try:
            data = resp.json() if resp is not None else {}
            if isinstance(data, dict):
                message_id = ((data.get("result") or {}) or {}).get("message_id")
        except Exception:
            message_id = None
        _append_channel_delivery_audit(
            {
                "platform": "telegram",
                "kind": "document",
                "chat_id": str(chat_id),
                "reply_to_message_id": int(reply_to_message_id) if reply_to_message_id else None,
                "message_id": int(message_id) if message_id else None,
                "ok": ok,
                "status_code": int(resp.status_code) if resp is not None else None,
                "file_path": p,
                "file_name": os.path.basename(p),
                "file_size": os.path.getsize(p),
                "caption_sha1": _audit_sha1(caption or ""),
                "preview": _audit_preview(caption or ""),
            }
        )
        return ok
    except Exception as e:
        logger.warning(f"⚠️ Telegram sendDocument failed: {e}")
        _append_channel_delivery_audit(
            {
                "platform": "telegram",
                "kind": "document",
                "chat_id": str(chat_id),
                "reply_to_message_id": int(reply_to_message_id) if reply_to_message_id else None,
                "ok": False,
                "file_path": p,
                "file_name": os.path.basename(p),
                "preview": _audit_preview(caption or ""),
                "error": str(e)[:500],
            }
        )
        return False


def _telegram_send_photo(chat_id: str, image_path: str, caption: str = "", reply_to_message_id: int | None = None) -> bool:
    token = _load_openclaw_telegram_token()
    p = (image_path or "").strip()
    if not token or not str(chat_id or "").strip() or (not p) or (not os.path.exists(p)):
        return False
    payload = {"chat_id": str(chat_id), "caption": (caption or "")[:900]}
    if reply_to_message_id:
        payload["reply_to_message_id"] = int(reply_to_message_id)
    try:
        with open(p, "rb") as f:
            files = {"photo": (os.path.basename(p), f)}
            resp = _telegram_api_post(token, "sendPhoto", payload=payload, files=files)
        ok = bool(resp and resp.status_code == 200)
        message_id = None
        try:
            data = resp.json() if resp is not None else {}
            if isinstance(data, dict):
                message_id = ((data.get("result") or {}) or {}).get("message_id")
        except Exception:
            message_id = None
        _append_channel_delivery_audit(
            {
                "platform": "telegram",
                "kind": "photo",
                "chat_id": str(chat_id),
                "reply_to_message_id": int(reply_to_message_id) if reply_to_message_id else None,
                "message_id": int(message_id) if message_id else None,
                "ok": ok,
                "status_code": int(resp.status_code) if resp is not None else None,
                "file_path": p,
                "file_name": os.path.basename(p),
                "file_size": os.path.getsize(p),
                "caption_sha1": _audit_sha1(caption or ""),
                "preview": _audit_preview(caption or ""),
            }
        )
        return ok
    except Exception as e:
        logger.warning(f"⚠️ Telegram sendPhoto failed: {e}")
        _append_channel_delivery_audit(
            {
                "platform": "telegram",
                "kind": "photo",
                "chat_id": str(chat_id),
                "reply_to_message_id": int(reply_to_message_id) if reply_to_message_id else None,
                "ok": False,
                "file_path": p,
                "file_name": os.path.basename(p),
                "preview": _audit_preview(caption or ""),
                "error": str(e)[:500],
            }
        )
        return False


def _telegram_download_file(file_id: str, suggested_name: str = "") -> str:
    token = _load_openclaw_telegram_token()
    fid = (file_id or "").strip()
    if not token or not fid:
        return ""
    out_path = ""
    try:
        from skills.bridge.http_pool import get_session
        sess = get_session()
        r = sess.get(f"https://api.telegram.org/bot{token}/getFile", params={"file_id": fid}, timeout=20)
        if r.status_code != 200:
            return ""
        obj = r.json() if r.content else {}
        if not isinstance(obj, dict) or not obj.get("ok"):
            return ""
        file_path = str((obj.get("result") or {}).get("file_path") or "").strip()
        if not file_path:
            return ""
        ext = os.path.splitext(suggested_name or file_path)[1] or ""
        out_path = f"/tmp/tg_{fid}{ext}"
        file_url = f"https://api.telegram.org/file/bot{token}/{file_path}"
        rr = sess.get(file_url, timeout=60)
        if rr.status_code != 200:
            return ""
        with open(out_path, "wb") as f:
            f.write(rr.content)
        return out_path
    except Exception as e:
        logger.warning(f"⚠️ Telegram file download failed: {e}")
        # Clean up partial temp file on failure
        if out_path:
            try:
                os.unlink(out_path)
            except OSError:
                pass
        return ""


def _normalize_telegram_output_text(text: str) -> str:
    s = (text or "").strip()
    if not s:
        return s
    try:
        if _normalize_output_text:
            return _normalize_output_text(s, platform="TELEGRAM")
    except Exception as e:
        logger.warning(f"⚠️ Taiwan wording guard skipped (Telegram): {e}")
    return s


def _telegram_send_orchestrator_response(chat_id: str, response_text: str, reply_to_message_id: int | None = None) -> None:
    text = _normalize_telegram_output_text(response_text)
    if not text:
        _telegram_send_text_to(chat_id, "⚠️ 任務完成，但沒有可用輸出。", reply_to_message_id=reply_to_message_id)
        return

    if "|||FILE_PATH|||" in text:
        try:
            text_part, file_path = text.split("|||FILE_PATH|||", 1)
            text_part = (text_part or "").strip()
            file_path = (file_path or "").strip()
            if _telegram_send_document(chat_id, file_path, caption=text_part, reply_to_message_id=reply_to_message_id):
                return
            _telegram_send_text_to(chat_id, f"{text_part}\n⚠️ 檔案傳送失敗：{file_path}", reply_to_message_id=reply_to_message_id)
            return
        except Exception as e:
            _telegram_send_text_to(chat_id, f"⚠️ 檔案回傳解析失敗：{e}", reply_to_message_id=reply_to_message_id)
            return

    if "|||IMAGE_PATH|||" in text:
        try:
            text_part, image_path = text.split("|||IMAGE_PATH|||", 1)
            text_part = (text_part or "").strip()
            image_path = (image_path or "").strip()
            if _telegram_send_photo(chat_id, image_path, caption=text_part, reply_to_message_id=reply_to_message_id):
                return
            _telegram_send_text_to(chat_id, f"{text_part}\n⚠️ 圖片傳送失敗：{image_path}", reply_to_message_id=reply_to_message_id)
            return
        except Exception as e:
            _telegram_send_text_to(chat_id, f"⚠️ 圖片回傳解析失敗：{e}", reply_to_message_id=reply_to_message_id)
            return

    _telegram_send_text_to(chat_id, text, reply_to_message_id=reply_to_message_id)


def _telegram_process_async(
    chat_id: str,
    user_id: str,
    role: str,
    user_text: str,
    attachment: dict | None = None,
    reply_to_message_id: int | None = None,
    channel_context: dict | None = None,
) -> None:
    # OBS-1: correlation ID  /  OBS-2: latency tracking
    correlation_id = f"magi-{uuid.uuid4().hex[:12]}"
    _start_ts = time.monotonic()
    tmp_path = ""
    try:
        if attachment:
            tmp_path = str(attachment.get("path") or "").strip()
            response_text = orchestrator.process_message(
                user_id=user_id,
                message=user_text,
                platform="Telegram",
                role=role,
                attachment=attachment,
                correlation_id=correlation_id,
                channel_context=channel_context,
            )
        else:
            response_text = orchestrator.process_message(
                user_id=user_id,
                message=user_text,
                platform="Telegram",
                role=role,
                correlation_id=correlation_id,
                channel_context=channel_context,
            )
        _telegram_send_orchestrator_response(chat_id, str(response_text or ""), reply_to_message_id=reply_to_message_id)
    except Exception as e:
        logger.error(f"❌ Telegram processing error: {e}")
        _telegram_send_text_to(chat_id, "⚠️ 系統暫時忙碌中，請稍後再試一次。")
    finally:
        if tmp_path and os.path.exists(tmp_path):
            _safe_remove_tmp(tmp_path)
        # OBS-2: record processing latency
        elapsed_ms = int((time.monotonic() - _start_ts) * 1000)
        _append_channel_delivery_audit({
            "platform": "Telegram",
            "kind": "latency",
            "user_id": str(user_id or ""),
            "chat_id": str(chat_id or ""),
            "correlation_id": correlation_id,
            "latency_ms": elapsed_ms,
        })


def _telegram_handle_update(update: dict, from_poll: bool = False) -> dict:
    # Inject request context for structured logging
    from skills.ops.structured_log import set_request_context
    _tg_user = ((update or {}).get("message") or {}).get("from") or {}
    set_request_context(
        request_id=uuid.uuid4().hex[:12],
        user_id=str(_tg_user.get("id", "")),
        platform="Telegram",
    )
    if not isinstance(update, dict):
        update = {}
    update_id = update.get("update_id")
    try:
        update_id = int(update_id) if update_id is not None else None
    except Exception:
        update_id = None
    if (not from_poll) and _telegram_mark_seen_update(update_id):
        return {"ok": True, "deduped": True}

    msg = (
        update.get("message")
        or update.get("edited_message")
        or update.get("channel_post")
        or update.get("edited_channel_post")
        or {}
    )
    if not isinstance(msg, dict) or not msg:
        return {"ok": True, "ignored": "no_message"}

    chat = msg.get("chat") or {}
    sender_chat = msg.get("sender_chat") or {}
    user = msg.get("from") or {}
    chat_id = str(chat.get("id") or "").strip()
    chat_type = str(chat.get("type") or "").strip().lower()
    chat_title = str(chat.get("title") or "").strip()
    user_id_raw = str(user.get("id") or "").strip()
    sender_chat_id = str(sender_chat.get("id") or "").strip()
    sender_id = user_id_raw or sender_chat_id or chat_id
    if not chat_id:
        return {"ok": True, "ignored": "missing_ids"}

    allowed_admin_ids = set(_load_admin_telegram_ids() or [])
    allowed_admin_ids |= set(_load_notify_telegram_ids() or [])
    candidate_ids = {x for x in [sender_id, chat_id, user_id_raw, sender_chat_id] if str(x or "").strip()}

    def _is_allowed() -> bool:
        return (not allowed_admin_ids) or any(cid in allowed_admin_ids for cid in candidate_ids)

    if not _is_allowed():
        # Bootstrap path: if this is a MAGI-named group/channel, bind once then re-check.
        try:
            auto_pair = str(os.environ.get("MAGI_TG_AUTO_PAIR_MAGI_GROUP", "1")).strip().lower() in {"1", "true", "yes", "on"}
            title_hit = "magi" in chat_title.lower()
            if auto_pair and chat_type in {"group", "supergroup", "channel"} and title_hit:
                _telegram_apply_group_notify_binding(chat_id=chat_id, sender_id=(user_id_raw or sender_chat_id or chat_id))
                allowed_admin_ids = set(_load_admin_telegram_ids() or [])
                allowed_admin_ids |= set(_load_notify_telegram_ids() or [])
        except Exception as pair_err:
            logger.warning(f"⚠️ Telegram auto-pair failed: {pair_err}")

    if not _is_allowed():
        logger.warning(
            "⛔ Telegram blocked by allowlist chat_id=%s chat_type=%s sender_id=%s user_id=%s sender_chat_id=%s title=%s candidates=%s",
            chat_id,
            chat_type,
            sender_id,
            user_id_raw,
            sender_chat_id,
            chat_title[:80],
            list(candidate_ids),
        )
        _telegram_send_text_to(chat_id, "⛔ 此 Telegram 帳號不在允許清單中。")
        return {"ok": True, "blocked": "allowlist"}

    role = "admin" if any(cid in allowed_admin_ids for cid in candidate_ids) else "user"
    user_id = f"telegram_{sender_id}"
    user_text = str(msg.get("text") or msg.get("caption") or "").strip()
    try:
        message_thread_id = int(msg.get("message_thread_id")) if msg.get("message_thread_id") is not None else None
    except Exception:
        message_thread_id = None
    attachment = None

    # Auto-bind: if admin talks in a group/supergroup, make this chat a notify target.
    try:
        auto_bind = str(os.environ.get("MAGI_TG_AUTO_BIND_GROUP_NOTIFY", "1")).strip().lower() in {"1", "true", "yes", "on"}
        if auto_bind and role == "admin" and chat_type in {"group", "supergroup"}:
            ok, msg_bind = _telegram_apply_group_notify_binding(chat_id=chat_id, sender_id=sender_id)
            if ok:
                logger.info(f"✅ Telegram auto-bind notify target chat_id={chat_id}")
            else:
                logger.warning(f"⚠️ Telegram auto-bind failed chat_id={chat_id}: {msg_bind}")
    except Exception as bind_err:
        logger.warning(f"⚠️ Telegram auto-bind exception: {bind_err}")

    settings_reply = _handle_telegram_settings_command(
        user_text,
        chat_id=chat_id,
        sender_id=sender_id,
        message_thread_id=message_thread_id,
        role=role,
    )
    if settings_reply is not None:
        _telegram_send_text_to(chat_id, settings_reply, reply_to_message_id=msg.get("message_id"))
        return {"ok": True, "settings_cmd": True}

    # ── Skip MAGI notification messages ──────────────────────────────
    # When LAFNotifier (or other MAGI subsystems) sends notifications via
    # TG bot API, the webhook may receive them in group chats.  Also skip
    # when a user *replies* to a notification — the reply should not be
    # treated as a conversational prompt for the AI.
    _NOTIFICATION_PREFIXES = ("📋", "💰", "📥", "⚠️ 閱卷", "✅ 閱卷", "🔔")

    def _is_notification_text(t: str) -> bool:
        if not t:
            return False
        return any(t.startswith(p) for p in _NOTIFICATION_PREFIXES)

    # Case 1: message itself is from a bot (self-message in group)
    if user.get("is_bot") and _is_notification_text(user_text):
        logger.info("🔕 Telegram: skipping bot-originated notification message")
        return {"ok": True, "ignored": "bot_notification"}

    # Case 2: user replied to a notification message
    reply_to = msg.get("reply_to_message") or {}
    reply_from = reply_to.get("from") or {}
    reply_text = str(reply_to.get("text") or reply_to.get("caption") or "").strip()
    if reply_from.get("is_bot") and _is_notification_text(reply_text):
        logger.info("🔕 Telegram: skipping reply-to-notification (replied to: %s)", reply_text[:60])
        return {"ok": True, "ignored": "reply_to_notification"}
    # ─────────────────────────────────────────────────────────────────

    try:
        if isinstance(msg.get("voice"), dict):
            voice = msg.get("voice") or {}
            file_id = str(voice.get("file_id") or "").strip()
            path = _telegram_download_file(file_id, suggested_name="voice.ogg")
            if path:
                attachment = {"type": "audio", "path": path, "filename": "voice.ogg", "timestamp": time.time()}
                if not user_text:
                    user_text = "請轉換成逐字稿，附上時間戳記，並輸出TXT檔。"
        elif isinstance(msg.get("audio"), dict):
            audio = msg.get("audio") or {}
            file_id = str(audio.get("file_id") or "").strip()
            fname = str(audio.get("file_name") or "audio.m4a").strip()
            path = _telegram_download_file(file_id, suggested_name=fname)
            if path:
                attachment = {"type": "audio", "path": path, "filename": fname, "timestamp": time.time()}
                if not user_text:
                    user_text = "請轉換成逐字稿，附上時間戳記，並輸出TXT檔。"
        elif isinstance(msg.get("photo"), list) and msg.get("photo"):
            photos = msg.get("photo") or []
            best = photos[-1] if isinstance(photos[-1], dict) else {}
            file_id = str(best.get("file_id") or "").strip()
            path = _telegram_download_file(file_id, suggested_name="photo.jpg")
            if path:
                attachment = {"type": "image", "path": path, "filename": "photo.jpg", "timestamp": time.time()}
                if not user_text:
                    user_text = "請分析這張圖片並用繁體中文回覆重點。"
        elif isinstance(msg.get("document"), dict):
            doc = msg.get("document") or {}
            file_id = str(doc.get("file_id") or "").strip()
            fname = str(doc.get("file_name") or "document.bin").strip()
            mime = str(doc.get("mime_type") or "").lower()
            path = _telegram_download_file(file_id, suggested_name=fname)
            if path:
                msg_type = "audio" if mime.startswith("audio/") else "file"
                attachment = {"type": msg_type, "path": path, "filename": fname, "timestamp": time.time()}
                if not user_text:
                    user_text = "請轉換成逐字稿，附上時間戳記，並輸出TXT檔。" if msg_type == "audio" else "請分析這個檔案並回覆重點。"
    except Exception as att_err:
        logger.warning(f"⚠️ Telegram attachment parse failed: {att_err}")

    if not user_text and not attachment:
        # Ignore service/empty updates to avoid triggering generic fallback tasks.
        return {"ok": True, "ignored": "empty_message"}
    if not user_text:
        user_text = "請協助處理這則訊息。"

    recent_followup = False
    try:
        recent_followup = orchestrator.has_recent_attachment_followup(user_id, "Telegram", user_text)
    except Exception as recent_err:
        logger.warning(f"⚠️ Telegram recent attachment probe failed: {recent_err}")

    long_task = _likely_long_task(user_text, attachment) or recent_followup
    if long_task:
        ack_msg = "⏳ 已收到，正在處理中。完成後我會回覆結果。"
        if attachment and attachment.get("type") in ("file", "audio"):
            try:
                att_path = attachment.get("path", "")
                att_size = os.path.getsize(att_path) if att_path and os.path.exists(att_path) else 0
                if att_size > 0:
                    from api.orchestrator import Orchestrator
                    ack_msg = Orchestrator.estimate_file_processing_time(
                        file_size_bytes=att_size,
                        filename=attachment.get("filename", ""),
                        prompt=user_text or "",
                        file_path=att_path,
                    )
            except Exception:
                logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 7685, exc_info=True)
        _telegram_send_text_to(chat_id, ack_msg, reply_to_message_id=msg.get("message_id"))

    if (attachment and attachment.get("type") in ("file", "audio")) or recent_followup:
        try:
            job = _enqueue_attachment_job(
                platform_name="Telegram",
                user_id=user_id,
                role=role,
                user_text=user_text,
                attachment=attachment,
                chat_id=chat_id,
                reply_to_message_id=msg.get("message_id"),
            )
            if attachment:
                att_path = str(attachment.get("path") or "").strip()
                if att_path and os.path.exists(att_path):
                    _safe_remove_tmp(att_path)
            return {"ok": True, "job_id": job.get("job_id")}
        except Exception as enqueue_err:
            logger.error(f"❌ Telegram attachment job enqueue failed: {enqueue_err}")

    # 2026-03-29: Build channel_context from Telegram message_thread_id
    _tg_channel_ctx = None
    try:
        _tg_topic_key = ""
        if message_thread_id is not None:
            _tg_state = _load_telegram_channel_state()
            _tg_topic_map = _tg_state.get("topicMap") or {}
            # Reverse lookup: thread_id → topic_key
            for _tk, _tid in _tg_topic_map.items():
                if int(_tid) == int(message_thread_id):
                    _tg_topic_key = str(_tk)
                    break
        _tg_channel_ctx = {
            "topic_key": _tg_topic_key,
            "thread_id": message_thread_id,
            "channel_id": str(chat_id),
            "platform": "Telegram",
        }
    except Exception as _ctx_err:
        logger.debug(f"Telegram channel_context build skipped: {_ctx_err}")

    _CHANNEL_BG_EXECUTOR.submit(
        _telegram_process_async,
        chat_id,
        user_id,
        role,
        user_text,
        attachment,
        msg.get("message_id"),
        _tg_channel_ctx,
    )
    return {"ok": True}


@app.route("/telegram/webhook", methods=["GET", "POST"])
def telegram_webhook():
    if request.method == "GET":
        return "OK", 200

    if _check_rate_limit("webhook"):
        return jsonify({"ok": False, "error": "rate limited"}), 429

    if not _telegram_verify_webhook_secret():
        return jsonify({"ok": False, "error": "invalid webhook secret"}), 401

    try:
        _record_last_public_base_url()
    except Exception:
        logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 7755, exc_info=True)

    update = request.get_json(silent=True) or {}
    try:
        return jsonify(_telegram_handle_update(update, from_poll=False))
    except Exception as _tg_err:
        logger.error("Telegram webhook handler exception: %s", _tg_err, exc_info=True)
        return jsonify({"ok": False, "error": "internal_error"}), 500


TELEGRAM_POLL_OFFSET_FILE = os.path.join(_agent_dir_for_logs, "telegram_poll_offset.json")
TELEGRAM_POLLING_ENABLED = os.environ.get("MAGI_TELEGRAM_POLLING_ENABLED", "1").strip().lower() in {"1", "true", "yes", "on"}
TELEGRAM_POLLING_FORCE = os.environ.get("MAGI_TELEGRAM_POLLING_FORCE", "0").strip().lower() in {"1", "true", "yes", "on"}
_TELEGRAM_POLL_STARTED = False


def _load_telegram_poll_offset() -> int:
    try:
        if os.path.exists(TELEGRAM_POLL_OFFSET_FILE):
            obj = json.loads(Path(TELEGRAM_POLL_OFFSET_FILE).read_text(encoding="utf-8")) or {}
            return int(obj.get("offset") or -1)
    except Exception:
        return -1
    return -1


def _save_telegram_poll_offset(offset: int) -> None:
    try:
        Path(TELEGRAM_POLL_OFFSET_FILE).write_text(
            json.dumps({"offset": int(offset), "updated_at": int(time.time())}, ensure_ascii=False),
            encoding="utf-8",
        )
    except Exception:
        logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 7788, exc_info=True)


def _telegram_poll_loop() -> None:
    """
    Fallback mode for Telegram command intake when webhook URL is not publicly reachable.
    If webhook mode is active, Telegram usually returns 409 for getUpdates; we back off quietly.
    """
    from skills.bridge.http_pool import get_session
    offset = _load_telegram_poll_offset()
    backoff = 5
    while True:
        token = _load_openclaw_telegram_token()
        if not token:
            time.sleep(30)
            continue
        try:
            sess = get_session()
            params = {"timeout": 25}
            if offset >= 0:
                params["offset"] = offset + 1
            resp = sess.get(f"https://api.telegram.org/bot{token}/getUpdates", params=params, timeout=35)
            if resp.status_code == 409 and (not TELEGRAM_POLLING_FORCE):
                time.sleep(30)
                continue
            if resp.status_code != 200:
                time.sleep(min(backoff, 60))
                backoff = min(backoff * 2, 60)
                continue

            obj = resp.json() if resp.content else {}
            updates = obj.get("result") if isinstance(obj, dict) else []
            if not isinstance(updates, list) or not updates:
                backoff = 5
                continue

            for up in updates:
                try:
                    uid = int((up or {}).get("update_id"))
                    if uid > offset:
                        offset = uid
                except Exception:
                    logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 7830, exc_info=True)
                try:
                    _telegram_handle_update(up, from_poll=True)
                except Exception as e:
                    logger.warning(f"⚠️ Telegram poll update process failed: {e}")
            _save_telegram_poll_offset(offset)
            backoff = 5
        except Exception as e:
            logger.warning(f"⚠️ Telegram polling error: {e}")
            time.sleep(min(backoff, 60))
            backoff = min(backoff * 2, 60)


def _start_telegram_polling_fallback() -> None:
    global _TELEGRAM_POLL_STARTED
    if (not TELEGRAM_POLLING_ENABLED) or _TELEGRAM_POLL_STARTED:
        return
    _TELEGRAM_POLL_STARTED = True
    t = threading.Thread(target=_telegram_poll_loop, name="telegram-poll", daemon=True)
    t.start()
    logger.info("📮 Telegram polling fallback started.")


try:
    from api.admin_allowlist import get_line_admin_user_ids
except Exception:
    def get_line_admin_user_ids():  # type: ignore
        return {
            uid.strip()
            for uid in os.environ.get("MAGI_ADMIN_LINE_IDS", "").split(",")
            if uid.strip()
        }

ADMIN_LINE_USER_IDS = set(get_line_admin_user_ids() or set())
if not ADMIN_LINE_USER_IDS:
    logger.warning("⚠️ No LINE admin allowlist configured (MAGI_ADMIN_LINE_IDS/admin_allowlist.json). LINE users will default to non-admin.")

# -----------------------------------------------------------------------------
# LAF CAPTCHA human-in-the-loop
# -----------------------------------------------------------------------------
# Security note:
# CAPTCHA on production sites is an access-control / anti-bot mechanism.
# We do NOT auto-solve it. Instead, CASPER can push a captcha image link and
# ask the admin to reply with 4 digits. This handler intercepts that reply and
# writes it to a local file for the automation runner to consume.

LAF_CAPTCHA_REQUEST_FILE = os.environ.get(
    "MAGI_LAF_CAPTCHA_REQUEST_FILE",
    os.path.join(AGENT_DIR, "laf_captcha_request.json"),
)
LAF_CAPTCHA_RESPONSE_FILE = os.environ.get(
    "MAGI_LAF_CAPTCHA_RESPONSE_FILE",
    os.path.join(AGENT_DIR, "laf_captcha_response.json"),
)
LAF_CAPTCHA_TTL_SECONDS = int(os.environ.get("MAGI_LAF_CAPTCHA_TTL_SECONDS", "300") or "300")


def _load_json(path: str) -> dict:
    try:
        if path and os.path.exists(path):
            with open(path, "r", encoding="utf-8") as f:
                return json.load(f) or {}
    except Exception:
        return {}
    return {}


def _write_json_atomic(path: str, data: dict) -> None:
    try:
        os.makedirs(os.path.dirname(path) or ".", exist_ok=True)
        tmp = path + ".tmp"
        with open(tmp, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False)
        os.replace(tmp, path)
    except Exception:
        logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 7905, exc_info=True)


def _maybe_handle_laf_captcha_reply(event, user_id: str, user_text: str) -> bool:
    text = (user_text or "").strip()
    # Accept "1234" or "驗證碼 1234" etc.
    m = re.search(r"(^|\\D)(\\d{4})(\\D|$)", text)
    if not m:
        return False

    req = _load_json(LAF_CAPTCHA_REQUEST_FILE)
    if not req:
        return False

    now = int(time.time())
    requested_at = int(req.get("requested_at") or 0)
    expires_at = int(req.get("expires_at") or 0)
    if expires_at and now > expires_at:
        return False
    if requested_at and (now - requested_at) > max(30, LAF_CAPTCHA_TTL_SECONDS):
        return False

    req_id = (req.get("request_id") or "").strip()
    if not req_id:
        return False

    code = m.group(2)
    resp = {
        "request_id": req_id,
        "captcha": code,
        "received_at": now,
        "from_user_id": user_id,
    }
    _write_json_atomic(LAF_CAPTCHA_RESPONSE_FILE, resp)

    # Best-effort ack; avoid reply_token expiry issues by falling back to push.
    _line_send_text(event, user_id, "✅ 已收到驗證碼，CASPER 正在登入法扶，完成後我會再回報。", prefer_push=False)
    return True


# -----------------------------------------------------------------------------
# Generic human CAPTCHA broker (for multiple modules)
# -----------------------------------------------------------------------------
GEN_CAPTCHA_REQUEST_FILE = os.environ.get(
    "MAGI_CAPTCHA_REQUEST_FILE",
    os.path.join(AGENT_DIR, "captcha_request.json"),
)
GEN_CAPTCHA_RESPONSE_FILE = os.environ.get(
    "MAGI_CAPTCHA_RESPONSE_FILE",
    os.path.join(AGENT_DIR, "captcha_response.json"),
)


def _maybe_handle_generic_captcha_reply(event, user_id: str, user_text: str) -> bool:
    """
    Handle replies for human-in-the-loop captcha requests.
    Only triggers if a pending request exists and is not expired.
    """
    req = _load_json(GEN_CAPTCHA_REQUEST_FILE)
    if not req:
        return False

    now = int(time.time())
    expires_at = int(req.get("expires_at") or 0)
    if expires_at and now > expires_at:
        return False

    req_id = (req.get("request_id") or "").strip()
    if not req_id:
        return False

    expected_len = int(req.get("expected_len") or 0)
    text = (user_text or "").strip()

    # Extract digits; accept either exact length or a reasonable range if not specified.
    digits = re.sub(r"[^0-9]", "", text)
    if expected_len and expected_len > 0:
        if len(digits) < expected_len:
            return False
        digits = digits[:expected_len]
    else:
        if not (4 <= len(digits) <= 12):
            return False

    resp = {
        "request_id": req_id,
        "captcha": digits,
        "received_at": now,
        "from_user_id": user_id,
    }
    _write_json_atomic(GEN_CAPTCHA_RESPONSE_FILE, resp)
    _line_send_text(event, user_id, "✅ 已收到驗證碼，CASPER 正在繼續處理。", prefer_push=False)
    return True

# Persist the last seen LINE sender so CASPER can push notifications even when admin IDs aren't preconfigured.
AGENT_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".agent"))
os.makedirs(AGENT_DIR, exist_ok=True)
LINE_LAST_SENDER_FILE = os.environ.get(
    "MAGI_LINE_LAST_SENDER_FILE",
    os.path.join(AGENT_DIR, "line_last_sender.json"),
)
LINE_LAST_CALLBACK_FILE = os.environ.get(
    "MAGI_LINE_LAST_CALLBACK_FILE",
    os.path.join(AGENT_DIR, "line_last_callback.json"),
)
# Safer default: OFF. Admin must be explicitly allowlisted.
LINE_AUTO_ADMIN_LAST_SENDER = os.environ.get("MAGI_LINE_AUTO_ADMIN_LAST_SENDER", "0").strip().lower() in {"1", "true", "yes", "on"}

LINE_LAST_BASE_URL_FILE = os.environ.get(
    "MAGI_LINE_LAST_BASE_URL_FILE",
    os.path.join(AGENT_DIR, "line_last_base_url.json"),
)

EXPORTS_DIR = os.environ.get(
    "MAGI_EXPORTS_DIR",
    os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "static", "exports")),
)
EXPORT_LONG_TEXT = os.environ.get("MAGI_EXPORT_LONG_TEXT", "1").strip().lower() in {"1", "true", "yes", "on"}
EXPORT_TEXT_THRESHOLD = int(os.environ.get("MAGI_EXPORT_TEXT_THRESHOLD", "9000"))


def _is_loopback_base_url(base: str) -> bool:
    s = (base or "").strip()
    if not s:
        return True
    if "://" not in s:
        s = "https://" + s
    try:
        host = (urlparse(s).hostname or "").strip().lower()
    except Exception:
        return True
    if not host:
        return True
    if host == "localhost" or host == "::1":
        return True
    if host.startswith("127."):
        return True
    return False


def _normalize_public_base_url(base: str) -> str:
    s = (base or "").strip().strip("'\"")
    if not s:
        return ""
    if "://" not in s:
        s = "https://" + s
    return s.rstrip("/") + "/"


def _base_from_webhook_url(url: str) -> str:
    s = (url or "").strip().strip("'\"")
    if not s:
        return ""
    if "://" not in s:
        s = "https://" + s
    try:
        p = urlparse(s)
        if not p.scheme or not p.netloc:
            return ""
        return f"{p.scheme}://{p.netloc}/"
    except Exception:
        return ""


def _record_last_public_base_url():
    """
    Record the public base URL from the current request so background tasks can build downloadable links.
    Respects reverse proxies via X-Forwarded-Proto / X-Forwarded-Host.
    """
    try:
        proto = (request.headers.get("X-Forwarded-Proto") or "").split(",")[0].strip()
        host = (request.headers.get("X-Forwarded-Host") or "").split(",")[0].strip()
        if not proto:
            proto = (request.scheme or "http").strip()
        if not host:
            host = (request.host or "").strip()
        base = _normalize_public_base_url(f"{proto}://{host}")
        if not base or _is_loopback_base_url(base):
            return
        with open(LINE_LAST_BASE_URL_FILE, "w", encoding="utf-8") as f:
            json.dump({"base_url": base, "updated_at": int(time.time())}, f, ensure_ascii=False)
    except Exception:
        logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 8087, exc_info=True)


def _build_tailscale_base_url() -> str:
    """Build base URL from Tailscale IP if configured."""
    from skills.ops.export_text import _load_dotenv_value
    ts_ip = (
        os.environ.get("MAGI_TAILSCALE_IP")
        or _load_dotenv_value("MAGI_TAILSCALE_IP")
        or ""
    ).strip()
    if not ts_ip:
        return ""
    ts_port = (
        os.environ.get("MAGI_TAILSCALE_PORT")
        or _load_dotenv_value("MAGI_TAILSCALE_PORT")
        or "5002"
    ).strip()
    return f"http://{ts_ip}:{ts_port}/"


def _load_public_base_url() -> str:
    """
    Priority: explicit override → Tailscale VPN → LINE webhook → cached base URL.
    """
    # 1. Explicit override
    env_base = _normalize_public_base_url(os.environ.get("MAGI_PUBLIC_BASE_URL") or "")
    if env_base and (not _is_loopback_base_url(env_base)):
        return env_base
    # 2. Tailscale (stable, always-on VPN)
    ts_base = _build_tailscale_base_url()
    if ts_base:
        return ts_base
    # 3. LINE webhook domain (Cloudflare tunnel, may rotate)
    webhook_base = _base_from_webhook_url(os.environ.get("MAGI_LINE_WEBHOOK_ENDPOINT") or "")
    if webhook_base and (not _is_loopback_base_url(webhook_base)):
        return webhook_base
    # 4. Cached base URL from last webhook
    try:
        if os.path.exists(LINE_LAST_BASE_URL_FILE):
            with open(LINE_LAST_BASE_URL_FILE, "r", encoding="utf-8") as f:
                data = json.load(f) or {}
            base = _normalize_public_base_url(data.get("base_url") or "")
            if base and (not _is_loopback_base_url(base)):
                return base
    except Exception:
        logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 8133, exc_info=True)
    return ""


def _export_text_to_static(text: str, prefix: str = "casper") -> dict:
    """
    Write a UTF-8 TXT file under /static/exports and return a public URL if available.
    """
    s = (text or "").strip()
    if not s:
        return {"success": False, "error": "empty text"}
    # Strip Markdown formatting — TXT is plain text
    try:
        from api.tw_output_guard import strip_markdown_for_chat
        s = strip_markdown_for_chat(s)
    except Exception:
        logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 8149, exc_info=True)
    try:
        os.makedirs(EXPORTS_DIR, exist_ok=True)
        stamp = time.strftime("%Y%m%d_%H%M%S")
        token = uuid.uuid4().hex[:10]
        filename = f"{prefix}_{stamp}_{token}.txt"
        path = os.path.join(EXPORTS_DIR, filename)
        with open(path, "w", encoding="utf-8") as f:
            f.write(s + "\n")
        base = _load_public_base_url()
        url = (base.rstrip("/") + f"/static/exports/{filename}") if base else ""
        return {"success": True, "path": path, "filename": filename, "url": url}
    except Exception as e:
        return {"success": False, "error": str(e)}


def _safe_export_stem(name: str, fallback: str = "document") -> str:
    raw = str(name or "").strip()
    if not raw:
        raw = fallback
    # Keep CJK characters, strip path separators and invalid filesystem chars.
    raw = re.sub(r'[\\/:*?"<>|]+', "_", raw)
    raw = re.sub(r"\s+", "_", raw).strip(" ._")
    return raw or fallback


def _export_file_meta(path: str) -> dict:
    p = os.path.abspath(path)
    filename = os.path.basename(p)
    base = _load_public_base_url().rstrip("/")
    url = f"{base}/static/exports/{filename}" if base else ""
    return {"success": True, "path": p, "filename": filename, "url": url}


def _find_chrome_binary() -> str:
    candidates = [
        (os.environ.get("MAGI_CHROME_BIN") or "").strip(),
        "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome",
        "/Applications/Chromium.app/Contents/MacOS/Chromium",
        shutil.which("google-chrome"),
        shutil.which("chromium"),
        shutil.which("chromium-browser"),
    ]
    for c in candidates:
        if c and os.path.exists(c):
            return c
    return ""


def _export_form_docx(preview_text: str, stem: str) -> dict:
    txt = str(preview_text or "").strip()
    if not txt:
        return {"success": False, "error": "empty_text"}
    try:
        from docx import Document  # type: ignore
    except Exception as e:
        return {"success": False, "error": f"python_docx_unavailable: {e}"}
    try:
        os.makedirs(EXPORTS_DIR, exist_ok=True)
        filename = f"{stem}.docx"
        path = os.path.join(EXPORTS_DIR, filename)
        doc = Document()
        for line in txt.splitlines():
            doc.add_paragraph(line)
        doc.save(path)
        return _export_file_meta(path)
    except Exception as e:
        return {"success": False, "error": str(e)}


def _render_form_text_to_html(title: str, text: str) -> str:
    safe_title = ihtml.escape(str(title or "OSC 文件"))
    safe_text = ihtml.escape(str(text or ""))
    return (
        "<!doctype html><html><head><meta charset='utf-8'>"
        f"<title>{safe_title}</title>"
        "<style>"
        "body{font-family:'Noto Sans TC','PingFang TC','Microsoft JhengHei',sans-serif;"
        "margin:36px;color:#111;line-height:1.6;}"
        "h1{margin:0 0 16px;font-size:24px;}"
        "pre{white-space:pre-wrap;word-wrap:break-word;font-family:inherit;font-size:15px;margin:0;}"
        "</style></head><body>"
        f"<h1>{safe_title}</h1><pre>{safe_text}</pre></body></html>"
    )


def _export_form_pdf(title: str, preview_text: str, stem: str) -> dict:
    txt = str(preview_text or "").strip()
    if not txt:
        return {"success": False, "error": "empty_text"}
    try:
        os.makedirs(EXPORTS_DIR, exist_ok=True)
        pdf_name = f"{stem}.pdf"
        pdf_path = os.path.join(EXPORTS_DIR, pdf_name)
        
        # Render HTML
        html_content = _render_form_text_to_html(title, txt)
        
        # Generate PDF using weasyprint
        import weasyprint
        weasyprint.HTML(string=html_content).write_pdf(pdf_path)
        
        if (not os.path.exists(pdf_path)) or os.path.getsize(pdf_path) < 64:
            return {"success": False, "error": "pdf_not_generated"}
            
        return _export_file_meta(pdf_path)
    except Exception as e:
        import traceback
        err_msg = traceback.format_exc()
        return {"success": False, "error": f"weasyprint_failed: {e}\n{err_msg}"}


def _export_osc_form_files(title: str, preview_text: str, suggested_filename: str = "") -> dict:
    txt = str(preview_text or "").strip()
    if not txt:
        return {"success": False, "errors": [{"type": "common", "error": "empty_text"}]}
    stamp = time.strftime("%Y%m%d_%H%M%S")
    token = uuid.uuid4().hex[:8]
    stem = _safe_export_stem(suggested_filename, fallback="osc_form")
    full_stem = f"{stem}_{stamp}_{token}"
    docx_meta = _export_form_docx(txt, full_stem)
    pdf_meta = _export_form_pdf(title, txt, full_stem)
    errors = []
    if not docx_meta.get("success"):
        errors.append({"type": "docx", "error": str(docx_meta.get("error") or "docx_failed")})
    if not pdf_meta.get("success"):
        errors.append({"type": "pdf", "error": str(pdf_meta.get("error") or "pdf_failed")})
    ok = bool(docx_meta.get("success") or pdf_meta.get("success"))
    preferred = pdf_meta if pdf_meta.get("success") else (docx_meta if docx_meta.get("success") else {"success": False})
    return {
        "success": ok,
        "export": preferred,
        "export_docx": docx_meta,
        "export_pdf": pdf_meta,
        "errors": errors,
    }


def _public_url_for_local_file(local_path: str) -> str:
    """
    Return a public URL for a local file.
    If the file is already inside /static/, return its URL directly (no copy).
    Otherwise, copy to EXPORTS_DIR and return its URL.
    """
    try:
        p = (local_path or "").strip().strip("'\"")
        if not p or (not os.path.exists(p)):
            return ""
        base = _load_public_base_url().rstrip("/")
        if not base:
            return ""
        abs_p = os.path.abspath(p)
        static_abs = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "static"))
        # If already under /static/, serve directly without copying
        if abs_p.startswith(static_abs + os.sep):
            rel = abs_p[len(static_abs) + 1:]
            return f"{base}/static/{rel}"
        # Otherwise, copy to exports
        os.makedirs(EXPORTS_DIR, exist_ok=True)
        filename = os.path.basename(abs_p)
        stem, ext = os.path.splitext(filename)
        stamp = time.strftime("%Y%m%d_%H%M%S")
        token = uuid.uuid4().hex[:8]
        filename = f"{stem}_{stamp}_{token}{ext}"
        import shutil
        shutil.copy2(abs_p, os.path.join(EXPORTS_DIR, filename))
        return f"{base}/static/exports/{filename}"
    except Exception:
        return ""


def _record_last_line_sender(event):
    try:
        src = getattr(event, "source", None)
        user_id = getattr(src, "user_id", None)
        group_id = getattr(src, "group_id", None)
        room_id = getattr(src, "room_id", None)
        payload = {
            "user_id": user_id,
            "group_id": group_id,
            "room_id": room_id,
            "updated_at": int(time.time()),
        }
        with open(LINE_LAST_SENDER_FILE, "w", encoding="utf-8") as f:
            json.dump(payload, f, ensure_ascii=False)
    except Exception:
        logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 8335, exc_info=True)


def _record_last_line_callback(path: str = ""):
    try:
        _write_json_atomic(
            LINE_LAST_CALLBACK_FILE,
            {
                "updated_at": int(time.time()),
                "path": (path or "").strip() or "/callback",
            },
        )
    except Exception:
        logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 8348, exc_info=True)


def _load_last_line_sender_user_id() -> str:
    try:
        if os.path.exists(LINE_LAST_SENDER_FILE):
            with open(LINE_LAST_SENDER_FILE, "r", encoding="utf-8") as f:
                data = json.load(f) or {}
            return (data.get("user_id") or "").strip()
    except Exception:
        return ""
    return ""

# Context Buffer (Simple in-memory for now)
# { user_id: { "type": "image|audio|file", "path": "/tmp/...", "timestamp": ... } }
user_context = {}
CONTEXT_TTL_SECONDS = int(os.environ.get("LINE_CONTEXT_TTL_SECONDS", "900"))
EXPECTED_MAGI_API_KEY = os.environ.get("MAGI_API_KEY", "").strip()
ATTACHMENT_JOB_DIR = os.path.join(AGENT_DIR, "attachment_jobs")
ATTACHMENT_JOB_FILE_DIR = os.path.join(ATTACHMENT_JOB_DIR, "files")
if not EXPECTED_MAGI_API_KEY:
    logger.warning("⚠️ MAGI_API_KEY not configured. /api/transcribe will require authenticated dashboard session.")

def _safe_remove_tmp(path: str) -> None:
    """
    Safety: never delete Synology Drive artifacts.
    For temporary files (typically /tmp), allow cleanup via safe_fs.
    """
    p = (path or "").strip()
    if not p:
        return
    try:
        ensure_path_on_sys_path(get_orch_dir())
        import safe_fs  # type: ignore
        safe_fs.safe_remove(p, reason="tmp_cleanup", allow_delete=True)
        return
    except Exception:
        logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 8385, exc_info=True)
    try:
        if os.path.exists(p):
            os.remove(p)
    except Exception:
        logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 8390, exc_info=True)


def _cleanup_user_context():
    now = time.time()
    expired = []
    for uid, ctx in user_context.items():
        ts = float(ctx.get("timestamp", 0) or 0)
        if ts and (now - ts > CONTEXT_TTL_SECONDS):
            expired.append((uid, ctx))
    for uid, ctx in expired:
        path = ctx.get("path")
        if path and os.path.exists(path):
            _safe_remove_tmp(path)
        user_context.pop(uid, None)


def cleanup_old_exports(days: int = 30) -> int:
    """刪除 static/exports/ 中超過 N 天的檔案，回傳刪除數量。"""
    exports_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), "static", "exports")
    if not os.path.isdir(exports_dir):
        return 0
    cutoff = time.time() - (days * 86400)
    removed = 0
    for fname in os.listdir(exports_dir):
        fpath = os.path.join(exports_dir, fname)
        try:
            if os.path.isfile(fpath) and os.path.getmtime(fpath) < cutoff:
                os.remove(fpath)
                removed += 1
        except Exception:
            continue
    if removed:
        logging.getLogger(__name__).info("🧹 Cleaned up %d old exports (>%d days)", removed, days)
    return removed


try:
    from skills.memory import job_queue as _jq
except ImportError:
    _jq = None  # type: ignore[assignment]


def _read_attachment_job(job_id: str) -> dict:
    if _jq:
        return _jq.read(job_id)
    # Legacy JSON fallback
    status_path = Path(ATTACHMENT_JOB_DIR) / f"attachment_{job_id}.json"
    if not status_path.exists():
        return {}
    try:
        return json.loads(status_path.read_text(encoding="utf-8") or "{}")
    except Exception:
        return {}


def _write_attachment_job(job_id: str, patch: dict) -> dict:
    if _jq:
        job = _jq.read(job_id)
        status = str(patch.get("status") or "").strip()
        if status == "done":
            _jq.complete(job_id, result=str(patch.get("response_preview") or ""))
        elif status == "failed":
            _jq.fail(job_id, error=str(patch.get("error") or ""))
        elif status == "abandoned":
            _jq.abandon(job_id, reason=str(patch.get("abandon_reason") or patch.get("error") or ""))
        elif status == "running":
            _jq.claim(job_id)
        payload_patch = {}
        for key in (
            "progress",
            "progress_total",
            "progress_phase",
            "progress_message",
            "progress_current",
            "updated_at_iso",
            "finished_at",
            "success",
            "response_preview",
            "error",
        ):
            if key in patch:
                payload_patch[key] = patch.get(key)
        if payload_patch:
            job = _jq.update_payload(job_id, payload_patch)
        else:
            job = _jq.read(job_id)
        job.update(patch)
        return job
    # Legacy JSON fallback
    os.makedirs(ATTACHMENT_JOB_DIR, exist_ok=True)
    status_path = Path(ATTACHMENT_JOB_DIR) / f"attachment_{job_id}.json"
    data = _read_attachment_job(job_id)
    data.update(patch or {})
    data["job_id"] = job_id
    data["updated_at"] = datetime.now().isoformat()
    tmp_path = status_path.with_suffix(".tmp")
    tmp_path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    tmp_path.replace(status_path)
    return data


def _list_attachment_job_ids() -> list[str]:
    if _jq:
        return [j["id"] for j in _jq.list_all(limit=500)]
    base = Path(ATTACHMENT_JOB_DIR)
    if not base.exists():
        return []
    files = sorted(base.glob("attachment_*.json"), key=lambda p: p.stat().st_mtime)
    return [p.stem.replace("attachment_", "", 1) for p in files]


def _persist_attachment_copy(src_path: str, filename: str = "", prefix: str = "attachment") -> str:
    src = str(src_path or "").strip()
    if not src or not os.path.exists(src):
        return ""
    try:
        src_real = os.path.realpath(src)
        dst_root = os.path.realpath(ATTACHMENT_JOB_FILE_DIR)
        if src_real.startswith(dst_root + os.sep) or src_real == dst_root:
            return src
    except Exception:
        logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 8492, exc_info=True)
    os.makedirs(ATTACHMENT_JOB_FILE_DIR, exist_ok=True)
    ext = os.path.splitext(str(filename or src).strip())[1] or os.path.splitext(src)[1] or ".bin"
    safe_prefix = re.sub(r"[^A-Za-z0-9_-]+", "_", str(prefix or "attachment"))[:48] or "attachment"
    dst = os.path.join(ATTACHMENT_JOB_FILE_DIR, f"{safe_prefix}_{uuid.uuid4().hex[:10]}{ext}")
    shutil.copy2(src, dst)
    return dst


def _persist_attachment_payload(attachment: dict | None, *, prefix: str) -> dict | None:
    if not isinstance(attachment, dict):
        return None
    src_path = str(attachment.get("path") or "").strip()
    if not src_path or not os.path.exists(src_path):
        return None
    filename = str(attachment.get("filename") or os.path.basename(src_path) or "attachment").strip()
    dst_path = _persist_attachment_copy(src_path, filename=filename, prefix=prefix)
    if not dst_path:
        return None
    out = dict(attachment)
    out["path"] = dst_path
    out["filename"] = filename
    out["timestamp"] = float(attachment.get("timestamp") or time.time())
    return out


def _line_push_orchestrator_response(user_id: str, response_text: str) -> None:
    text = str(response_text or "").strip()
    if not text:
        _line_push_text(user_id, "⚠️ 任務完成，但沒有可用輸出。")
        return
    if "|||FILE_PATH|||" in text:
        try:
            text_part, file_path = text.split("|||FILE_PATH|||", 1)
            file_url = _public_url_for_local_file((file_path or "").strip())
            if file_url:
                body = (text_part or "").strip()
                msg = (body + "\n\n" if body else "") + f"📎 檔案下載：{file_url}"
                _line_push_text(user_id, msg)
                return
            _line_push_text(user_id, f"{(text_part or '').strip()}\n⚠️ 檔案已產生，但目前無法建立公開下載連結。")
            return
        except Exception as file_err:
            logger.error(f"❌ LINE push file response failed: {file_err}")
            _line_push_text(user_id, "❌ 檔案處理失敗，請稍後再試。")
            return
    _line_push_text(user_id, text)


def _deliver_attachment_job_response(job: dict, response_text: str) -> None:
    platform_name = str(job.get("platform") or "").strip().upper()
    if platform_name == "TELEGRAM":
        _telegram_send_orchestrator_response(
            str(job.get("chat_id") or "").strip(),
            str(response_text or ""),
            reply_to_message_id=int(job.get("reply_to_message_id") or 0) or None,
        )
        return
    if platform_name == "LINE":
        _line_push_orchestrator_response(str(job.get("user_id") or "").strip(), str(response_text or ""))
        return
    logger.warning("⚠️ Unknown attachment job platform: %s", platform_name)


def _run_attachment_job(job_id: str) -> None:
    if _jq:
        if not _jq.claim(job_id):
            return
        job = _jq.read(job_id)
    else:
        job = _read_attachment_job(job_id)
    if not job:
        return

    try:
        # SQLite stores attachment in payload dict; legacy JSON stores it flat
        _payload = job.get("payload") if isinstance(job.get("payload"), dict) else {}
        attachment = (
            _payload.get("attachment")
            if isinstance(_payload.get("attachment"), dict) else
            job.get("attachment") if isinstance(job.get("attachment"), dict) else
            None
        )

        # Progress callback — push intermediate status to user during long tasks.
        import time as _progress_time
        _last_progress = [0.0]
        def _job_progress_cb(phase, current, total, message):
            now = _progress_time.monotonic()
            if now - _last_progress[0] < 15:
                return
            _last_progress[0] = now
            try:
                total_i = max(1, int(total or 1))
                current_i = max(0, min(total_i, int(current or 0)))
                progress = int(max(0, min(100, round((current_i / total_i) * 100))))
            except Exception:
                current_i, total_i, progress = 0, 1, 0
            _write_attachment_job(
                job_id,
                {
                    "status": "running",
                    "progress": progress,
                    "progress_current": current_i,
                    "progress_total": total_i,
                    "progress_phase": str(phase or ""),
                    "progress_message": str(message or ""),
                    "updated_at_iso": datetime.now().isoformat(),
                },
            )
            try:
                _deliver_attachment_job_response(job, str(message or ""))
            except Exception:
                logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 8605, exc_info=True)

        response_text = orchestrator.process_message(
            user_id=str(job.get("user_id") or "").strip(),
            message=str(job.get("user_text") or ""),
            platform=str(job.get("platform") or "LINE"),
            role=str(job.get("role") or "user"),
            attachment=attachment,
            progress_callback=_job_progress_cb,
        )
        if response_text:
            try:
                orchestrator.record_assistant_reply(str(job.get("user_id") or "").strip(), response_text)
            except Exception:
                logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 8619, exc_info=True)
        _deliver_attachment_job_response(job, str(response_text or ""))
        _write_attachment_job(
            job_id,
            {
                "status": "done",
                "success": True,
                "progress": 100,
                "finished_at": datetime.now().isoformat(),
                "response_preview": str(response_text or "")[:1200],
            },
        )
    except Exception as e:
        err = str(e)
        logger.error("❌ Attachment job failed job_id=%s error=%s", job_id, err)
        _write_attachment_job(
            job_id,
            {
                "status": "failed",
                "success": False,
                "progress": 100,
                "finished_at": datetime.now().isoformat(),
                "error": err,
            },
        )
        try:
            _deliver_attachment_job_response(job, f"❌ 系統處理失敗：{err}")
        except Exception:
            logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 8647, exc_info=True)
    finally:
        pass  # SQLite job_queue handles state; no lock file to release


def _enqueue_attachment_job(
    *,
    platform_name: str,
    user_id: str,
    role: str,
    user_text: str,
    attachment: dict | None = None,
    chat_id: str = "",
    reply_to_message_id: int | None = None,
) -> dict:
    durable_attachment = _persist_attachment_payload(attachment, prefix=f"{platform_name.lower()}_att") if attachment else None
    if _jq:
        job_id = _jq.enqueue(
            job_type="attachment",
            platform=platform_name,
            user_id=user_id,
            role=role,
            user_text=user_text,
            chat_id=chat_id,
            reply_to_message_id=reply_to_message_id,
            payload={"attachment": durable_attachment} if durable_attachment else {},
        )
    else:
        job_id = datetime.now().strftime("%Y%m%d_%H%M%S") + "_" + uuid.uuid4().hex[:6]
        payload = {
            "status": "queued",
            "platform": str(platform_name or "LINE"),
            "user_id": str(user_id or "").strip(),
            "role": str(role or "user"),
            "user_text": str(user_text or ""),
            "attachment": durable_attachment,
            "chat_id": str(chat_id or "").strip(),
            "reply_to_message_id": int(reply_to_message_id or 0) or None,
            "created_at": datetime.now().isoformat(),
            "worker_pid": 0,
            "attempts": 0,
        }
        _write_attachment_job(job_id, payload)
    if durable_attachment:
        try:
            orchestrator.remember_recent_attachment(
                user_id=str(user_id or "").strip(),
                platform=str(platform_name or ""),
                attachment=durable_attachment,
                source_message=str(user_text or ""),
            )
        except Exception as recent_err:
            logger.warning(f"⚠️ Failed to remember recent attachment for job {job_id}: {recent_err}")
    _ATTACHMENT_BG_EXECUTOR.submit(_run_attachment_job, job_id)
    return {"success": True, "job_id": job_id}


def _resume_pending_attachment_jobs() -> None:
    _MAX_RESUME_ATTEMPTS = int(os.environ.get("MAGI_ATTACHMENT_MAX_RESUME", "3") or "3")
    if _jq:
        resumed, abandoned = _jq.recover_stale_running(max_attempts=_MAX_RESUME_ATTEMPTS)
        # Re-submit recovered jobs to executor
        for job in _jq.list_by_status("queued"):
            _ATTACHMENT_BG_EXECUTOR.submit(_run_attachment_job, job["id"])
        # Periodic cleanup of old completed jobs
        _jq.cleanup_old(days=30)
        return
    # Legacy JSON fallback
    resumed = 0
    abandoned = 0
    for job_id in _list_attachment_job_ids():
        job = _read_attachment_job(job_id)
        status = str(job.get("status") or "").strip().lower()
        if status not in {"queued", "running"}:
            continue
        attempts = int(job.get("attempts") or 0)
        if attempts >= _MAX_RESUME_ATTEMPTS:
            _write_attachment_job(job_id, {**job, "status": "abandoned", "abandon_reason": f"exceeded {_MAX_RESUME_ATTEMPTS} attempts"})
            abandoned += 1
            continue
        _ATTACHMENT_BG_EXECUTOR.submit(_run_attachment_job, job_id)
        resumed += 1
    if resumed or abandoned:
        logger.info("♻️ Resumed %s pending attachment jobs, abandoned %s.", resumed, abandoned)

# Async Processing Function
def _chunk_text_for_line(text: str, limit: int = 4200) -> list[str]:
    s = (text or "").strip()
    if not s:
        return []
    limit = max(300, int(limit))
    out = []
    i = 0
    while i < len(s):
        out.append(s[i : i + limit])
        i += limit
    return out


def _likely_long_task(user_text: str, attachment: dict | None) -> bool:
    if attachment:
        return True
    t = (user_text or "").lower()
    if re.search(r"https?://", t):
        return True
    if any(
        k in t
        for k in [
            # Reading/translation/summarization
            "翻譯", "translate", "摘要", "總結", "整理", "讀取", "分析", "文件", "檔案", "網頁", "網址",
            "全文", "整篇", "整份", "完整翻譯", "全文翻譯", "逐字稿", "時間戳", "時間碼",
            # Research/fetch
            "搜尋", "search", "抓取", "fetch", "research",
            # Heavier reasoning
            "深度思考", "deep think",
            # Media generation/processing
            "畫", "draw", "產生圖片", "generate image", "製作音樂", "生成音樂",
        ]
    ):
        return True
    if len(t) > 1200:
        return True
    return False


_LINE_PUSH_COUNTER_FILE = os.path.join(AGENT_DIR, "line_push_counter.json")
_LINE_PUSH_DAILY_LIMIT = int(os.environ.get("MAGI_LINE_PUSH_DAILY_LIMIT", "5") or "5")


def _line_push_budget_ok() -> bool:
    """Check if daily LINE push budget still has room (free plan = 200/month ≈ 6/day)."""
    today = time.strftime("%Y-%m-%d")
    try:
        if os.path.exists(_LINE_PUSH_COUNTER_FILE):
            with open(_LINE_PUSH_COUNTER_FILE, "r", encoding="utf-8") as f:
                data = json.load(f) or {}
        else:
            data = {}
        return int(data.get(today, 0)) < _LINE_PUSH_DAILY_LIMIT
    except Exception:
        return True


def _line_push_budget_increment() -> None:
    """Record one push message for today's budget."""
    today = time.strftime("%Y-%m-%d")
    try:
        data = {}
        if os.path.exists(_LINE_PUSH_COUNTER_FILE):
            with open(_LINE_PUSH_COUNTER_FILE, "r", encoding="utf-8") as f:
                data = json.load(f) or {}
        # Clean old dates
        data = {k: v for k, v in data.items() if k >= time.strftime("%Y-%m", time.localtime())}
        data[today] = int(data.get(today, 0)) + 1
        with open(_LINE_PUSH_COUNTER_FILE, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False)
    except Exception:
        logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 8804, exc_info=True)


def _line_push_text(user_id: str, text: str, *, is_chat_reply: bool = False) -> bool:
    """
    Push a (possibly long) text message to LINE user.
    Push is used for long tasks to avoid reply_token expiry.

    Args:
        is_chat_reply: True if this is a fallback for an expired reply_token
                       (user-initiated conversation). Chat replies are never
                       budget-limited — the user should always get a response.
    """
    if not is_chat_reply and not _line_push_budget_ok():
        logger.warning("LINE push skipped — daily notification budget exhausted (%d/%d). "
                       "Message will only appear in TG.", _LINE_PUSH_DAILY_LIMIT, _LINE_PUSH_DAILY_LIMIT)
        return False
    ok_all = True
    safe_text = _normalize_line_output_text(text)
    for part in _chunk_text_for_line(safe_text, limit=4200):
        try:
            line_bot_api.push_message(user_id, TextSendMessage(text=part))
            _line_push_budget_increment()
        except Exception as e:
            logger.error(f"❌ LINE push failed: {e}")
            _handle_line_send_failure(e, user_id=user_id, phase="push")
            ok_all = False
            break
    return ok_all


LINE_QUOTA_ALERT_FILE = os.path.join(AGENT_DIR, "line_quota_alert.json")
LINE_DELAY_QUEUE_FILE = os.path.join(AGENT_DIR, "line_delayed_queue.json")
DISCORD_LAST_CHANNEL_FILE = os.path.join(AGENT_DIR, "discord_last_channel.json")
_LINE_LAST_OUTGOING: dict[str, dict] = {}
_LINE_DELAYED_ALERT_TS: dict[str, int] = {}


def _remember_last_line_outgoing(user_id: str, text: str) -> None:
    try:
        uid = str(user_id or "").strip()
        if not uid:
            return
        body = str(text or "").strip()
        if not body:
            return
        _LINE_LAST_OUTGOING[uid] = {"ts": int(time.time()), "text": body[:12000]}
    except Exception:
        logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 8852, exc_info=True)


def _last_line_outgoing_preview(user_id: str) -> str:
    try:
        uid = str(user_id or "").strip()
        item = _LINE_LAST_OUTGOING.get(uid) or {}
        text = str(item.get("text") or "").strip()
        if not text:
            return ""
        text = re.sub(r"\s+", " ", text)
        return text[:280]
    except Exception:
        return ""


def _load_openclaw_cfg() -> dict:
    def _raw_load() -> dict:
        try:
            p = Path.home() / ".openclaw" / "openclaw.json"
            if p.exists():
                data = json.loads(p.read_text(encoding="utf-8"))
                return data if isinstance(data, dict) else {}
        except Exception:
            return {}
        return {}

    cfg = _raw_load()
    try:
        channels = cfg.setdefault("channels", {})
        tg = channels.setdefault("telegram", {})
        legacy_notify = tg.get("notifyTo")
        legacy_topic_map = tg.get("topicMap")
        if isinstance(legacy_notify, list) or isinstance(legacy_topic_map, dict):
            state = _load_telegram_channel_state()
            changed_cfg = False
            changed_state = False
            if isinstance(legacy_notify, list):
                merged = [str(x).strip() for x in legacy_notify if str(x).strip()]
                for item in merged:
                    if item not in state["notifyTo"]:
                        state["notifyTo"].append(item)
                        changed_state = True
                tg.pop("notifyTo", None)
                changed_cfg = True
            if isinstance(legacy_topic_map, dict):
                for key, value in legacy_topic_map.items():
                    try:
                        tid = int(value or 0)
                    except Exception:
                        tid = 0
                    if key and tid > 0 and int(state["topicMap"].get(str(key)) or 0) != tid:
                        state["topicMap"][str(key)] = tid
                        changed_state = True
                tg.pop("topicMap", None)
                changed_cfg = True
            if changed_state:
                _save_telegram_channel_state(state)
            if changed_cfg:
                _save_openclaw_cfg(cfg)
    except Exception as e:
        logger.warning(f"⚠️ telegram channel-state migration skipped: {e}")
    return cfg


def _load_telegram_channel_state() -> dict:
    path = Path(f"{_MAGI_ROOT}/.agent/telegram_channel_state.json")
    try:
        if path.exists():
            data = json.loads(path.read_text(encoding="utf-8"))
            if isinstance(data, dict):
                notify_to = data.get("notifyTo")
                topic_map = data.get("topicMap")
                return {
                    "notifyTo": [str(x).strip() for x in notify_to if str(x).strip()] if isinstance(notify_to, list) else [],
                    "topicMap": {
                        str(k): int(v)
                        for k, v in (topic_map or {}).items()
                        if str(k).strip() and str(v).strip() and int(v or 0) > 0
                    } if isinstance(topic_map, dict) else {},
                }
    except Exception:
        logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 8934, exc_info=True)
    return {"notifyTo": [], "topicMap": {}}


def _save_telegram_channel_state(state: dict) -> bool:
    try:
        path = Path(f"{_MAGI_ROOT}/.agent/telegram_channel_state.json")
        path.parent.mkdir(parents=True, exist_ok=True)
        payload = {
            "notifyTo": [str(x).strip() for x in (state.get("notifyTo") or []) if str(x).strip()],
            "topicMap": {
                str(k): int(v)
                for k, v in (state.get("topicMap") or {}).items()
                if str(k).strip() and int(v or 0) > 0
            },
        }
        tmp = path.with_name(f"{path.name}.{os.getpid()}.{int(time.time()*1000)}.tmp")
        tmp.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
        tmp.replace(path)
        return True
    except Exception as e:
        logger.warning(f"⚠️ save telegram channel state failed: {e}")
        return False


def _save_openclaw_cfg(cfg: dict) -> bool:
    try:
        p = Path.home() / ".openclaw" / "openclaw.json"
        p.parent.mkdir(parents=True, exist_ok=True)
        tmp = p.with_name(f"{p.name}.{os.getpid()}.{int(time.time()*1000)}.tmp")
        tmp.write_text(
            json.dumps(cfg or {}, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        tmp.replace(p)
        return True
    except Exception as e:
        logger.warning(f"⚠️ save openclaw config failed: {e}")
        return False


def _telegram_topic_key(raw: str) -> str:
    key = str(raw or "").strip().lower()
    aliases = {
        "general": "general",
        "default": "general",
        "預設": "general",
        "主串": "general",
        "一般": "general",
        "filereview": "filereview",
        "file_review": "filereview",
        "file-review": "filereview",
        "閱卷": "filereview",
        "卷宗": "filereview",
        "docket": "filereview",
        "transcript": "transcript",
        "transcripts": "transcript",
        "筆錄": "transcript",
        "laf": "laf",
        "legal_aid": "laf",
        "legal-aid": "laf",
        "法扶": "laf",
        "judgment": "judgment",
        "judgments": "judgment",
        "判決": "judgment",
        "司法院": "judgment",
        "verbatim": "verbatim",
        "逐字稿": "verbatim",
        "音訊": "verbatim",
        "translation": "translation",
        "translate": "translation",
        "翻譯": "translation",
        "summary": "summary",
        "summarize": "summary",
        "摘要": "summary",
        "market": "market",
        "stocks": "market",
        "stock": "market",
        "股票": "market",
        "股市": "market",
        "check": "check",
        "checks": "check",
        "health": "check",
        "檢查": "check",
        "巡檢": "check",
        "nightly": "nightly",
        "夜間": "nightly",
        "改善": "nightly",
        "夜間會議": "nightly",
        "alert": "alert",
        "alerts": "alert",
        "warning": "alert",
        "警報": "alert",
        "警告": "alert",
        "告警": "alert",
        "iron_dome": "alert",
        "irondome": "alert",
        "鐵穹": "alert",
    }
    return aliases.get(key, "")


def _telegram_apply_group_notify_binding(chat_id: str, sender_id: str) -> tuple[bool, str]:
    state = _load_telegram_channel_state()
    changed_state = False

    allow_from = state.get("allowFrom") if isinstance(state.get("allowFrom"), list) else []
    for cid in [str(sender_id or "").strip(), str(chat_id or "").strip()]:
        if cid and cid not in allow_from:
            allow_from.append(cid)
            changed_state = True
    state["allowFrom"] = allow_from

    notify_to = state.get("notifyTo") if isinstance(state.get("notifyTo"), list) else []
    c = str(chat_id or "").strip()
    if c and c not in notify_to:
        notify_to.append(c)
        changed_state = True
    state["notifyTo"] = notify_to

    if not changed_state:
        return True, f"✅ 本群已在通知目標中\nchat_id: {chat_id}\nnotifyTo: {notify_to}"

    if not _save_telegram_channel_state(state):
        return False, "❌ 寫入設定失敗（telegram_channel_state.json）"
    return True, f"✅ 已綁定本群為通知目標\nchat_id: {chat_id}\nnotifyTo: {notify_to}"


def _telegram_bind_topic(chat_id: str, sender_id: str, topic_raw: str, thread_id: int) -> tuple[bool, str]:
    topic_key = _telegram_topic_key(topic_raw)
    if not topic_key:
        return False, "❌ 無法辨識主題，請用：general / filereview / transcript / laf / judgment / verbatim / translation / summary / market / check / nightly / alert"

    state = _load_telegram_channel_state()
    changed_state = False

    allow_from = state.get("allowFrom") if isinstance(state.get("allowFrom"), list) else []
    for cid in [str(sender_id or "").strip(), str(chat_id or "").strip()]:
        if cid and cid not in allow_from:
            allow_from.append(cid)
            changed_state = True
    state["allowFrom"] = allow_from

    notify_to = state.get("notifyTo") if isinstance(state.get("notifyTo"), list) else []
    c = str(chat_id or "").strip()
    if c and c not in notify_to:
        notify_to.append(c)
        changed_state = True
    state["notifyTo"] = notify_to

    topic_map = state.get("topicMap") if isinstance(state.get("topicMap"), dict) else {}
    old_tid = int(topic_map.get(str(topic_key)) or 0)
    new_tid = int(thread_id)
    if old_tid != new_tid:
        topic_map[str(topic_key)] = new_tid
        changed_state = True
    state["topicMap"] = topic_map

    if not changed_state:
        return True, f"✅ 主題 `{topic_key}` 已是 thread_id {int(thread_id)}"

    if not _save_telegram_channel_state(state):
        return False, "❌ 寫入主題設定失敗（telegram_channel_state.json）"
    return True, f"✅ 已綁定主題 `{topic_key}` -> thread_id {int(thread_id)}"


def _telegram_setup_topics(chat_id: str, sender_id: str) -> tuple[bool, str]:
    token = _load_openclaw_telegram_token()
    if not token:
        return False, "❌ 找不到 Telegram bot token。"
    c = str(chat_id or "").strip()
    if not c:
        return False, "❌ 缺少 chat_id。"

    ok_bind, bind_msg = _telegram_apply_group_notify_binding(chat_id=c, sender_id=str(sender_id or "").strip())
    if not ok_bind:
        return False, bind_msg

    # Preflight: forum topics are only available in supergroup with topics enabled.
    try:
        chat_resp = _telegram_api_post(token, "getChat", payload={"chat_id": c})
        if chat_resp and int(getattr(chat_resp, "status_code", 0) or 0) == 200:
            data = chat_resp.json() if hasattr(chat_resp, "json") else {}
            chat_info = (data or {}).get("result") or {}
            chat_type = str(chat_info.get("type") or "").strip().lower()
            is_forum = bool(chat_info.get("is_forum"))
            if chat_type == "channel":
                return False, (
                    "⚠️ 這個聊天是「頻道 (channel)」，不是可開 Topic 的群組。\n"
                    "Telegram Topics 只能用在「超級群組 (supergroup)」。\n\n"
                    "請建立/使用超級群組後再執行：`建立MAGI主題`。\n"
                    "若你想先用目前頻道，通知仍可發送，但會全部在同一串。"
                )
            if chat_type == "supergroup" and not is_forum:
                return False, (
                    "⚠️ 目前是超級群組，但尚未啟用 Topics。\n"
                    "請到 Telegram 群組設定開啟「Topics/主題」，再執行：`建立MAGI主題`。"
                )
    except Exception:
        logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 9151, exc_info=True)

    plan = [
        ("general", "一般"),
        ("filereview", "閱卷"),
        ("transcript", "筆錄"),
        ("laf", "法扶"),
        ("judgment", "判決"),
        ("judicial_api", "司法院API"),
        ("verbatim", "逐字稿"),
        ("translation", "翻譯"),
        ("summary", "摘要"),
        ("market", "股票"),
        ("check", "檢查"),
        ("nightly", "夜間"),
        ("alert", "警告"),
    ]
    state = _load_telegram_channel_state()
    existing_topic_map = state.get("topicMap") if isinstance(state.get("topicMap"), dict) else {}

    def _try_edit_topic_name(thread_id: int, topic_name: str) -> tuple[bool, str]:
        try:
            resp = _telegram_api_post(
                token,
                "editForumTopic",
                payload={
                    "chat_id": c,
                    "message_thread_id": int(thread_id),
                    "name": topic_name,
                },
            )
            if resp is None:
                return False, "no_response"
            try:
                data = resp.json()
            except Exception:
                data = {}
            if int(getattr(resp, "status_code", 0) or 0) == 200 and data.get("ok"):
                return True, ""
            desc = str((data or {}).get("description") or f"http_{getattr(resp, 'status_code', 'ERR')}")
            if "topic_not_modified" in desc.lower():
                return True, ""
            return False, desc
        except Exception as e:
            return False, type(e).__name__

    created: list[tuple[str, int]] = []
    reused: list[tuple[str, int]] = []
    skipped: list[str] = []
    failed: list[str] = []

    for topic_key, topic_name in plan:
        try:
            known_tid = int(existing_topic_map.get(topic_key) or 0)
        except Exception:
            known_tid = 0
        if known_tid > 0:
            ok_edit, why = _try_edit_topic_name(known_tid, topic_name)
            if ok_edit:
                ok_topic, _ = _telegram_bind_topic(c, sender_id, topic_key, known_tid)
                if ok_topic:
                    reused.append((topic_name, known_tid))
                    continue
                failed.append(f"{topic_name}: bind_failed")
                continue
            low_why = str(why or "").lower()
            # Stale/invalid thread mapping, fallback to creating a fresh topic.
            if ("message thread not found" not in low_why) and ("thread not found" not in low_why):
                failed.append(f"{topic_name}: reuse_failed({why or 'unknown'})")
                continue

        try:
            resp = _telegram_api_post(
                token,
                "createForumTopic",
                payload={"chat_id": c, "name": topic_name},
            )
            if resp is None:
                failed.append(f"{topic_name}: no_response")
                continue
            try:
                data = resp.json()
            except Exception:
                data = {}
            if int(getattr(resp, "status_code", 0) or 0) == 200 and data.get("ok"):
                result = data.get("result") or {}
                mtid = int(result.get("message_thread_id") or 0)
                if mtid > 0:
                    ok_topic, _ = _telegram_bind_topic(c, sender_id, topic_key, mtid)
                    if ok_topic:
                        created.append((topic_name, mtid))
                    else:
                        failed.append(f"{topic_name}: bind_failed")
                else:
                    failed.append(f"{topic_name}: no_thread_id")
                continue

            desc = str((data or {}).get("description") or "")
            low = desc.lower()
            if ("chat not found" in low) or ("forbidden" in low):
                return False, f"❌ 建立主題失敗：{desc or 'chat access denied'}"
            if ("not a forum" in low) or ("chat_not_forum" in low):
                return False, (
                    "⚠️ Telegram 回覆此聊天不是 forum。\n"
                    "可能是「頻道」或「尚未啟用 Topics 的超級群組」。\n"
                    "請改用超級群組並啟用 Topics 後重試。"
                )
            if ("topic already exists" in low) or ("already exists" in low):
                skipped.append(topic_name)
                continue
            if ("topic with this name already exists" in low) or ("topic name is already occupied" in low):
                skipped.append(topic_name)
                continue
            failed.append(f"{topic_name}: {desc or ('http_' + str(getattr(resp, 'status_code', 'ERR')))}")
        except Exception as e:
            failed.append(f"{topic_name}: {type(e).__name__}")

    lines = ["✅ 已嘗試建立 MAGI 主題並綁定通知分流。"]
    if reused:
        lines.append("已沿用：")
        for name, tid in reused:
            lines.append(f"- {name}（thread_id={tid}）")
    if created:
        lines.append("已建立：")
        for name, tid in created:
            lines.append(f"- {name}（thread_id={tid}）")
    if skipped:
        lines.append("已存在（略過）：")
        for name in skipped:
            lines.append(f"- {name}")
    if failed:
        lines.append("失敗：")
        for item in failed:
            lines.append(f"- {item}")

    # No longer need a separate "default" topic — "general" serves that role.

    lines.append("")
    lines.append(_telegram_notify_settings_text())
    return True, "\n".join(lines)


def _telegram_notify_settings_text() -> str:
    state = _load_telegram_channel_state()
    allow_from = state.get("allowFrom") if isinstance(state.get("allowFrom"), list) else []
    notify_to = state.get("notifyTo") if isinstance(state.get("notifyTo"), list) else []
    topic_map = state.get("topicMap") if isinstance(state.get("topicMap"), dict) else {}
    lines = [
        "📮 Telegram 通知設定",
        f"allowFrom: {allow_from}",
        f"notifyTo: {notify_to}",
        "topicMap:",
    ]
    if topic_map:
        for k in sorted(topic_map.keys()):
            lines.append(f"- {k}: {topic_map.get(k)}")
    else:
        lines.append("- (empty)")
    return "\n".join(lines)


def _handle_telegram_settings_command(
    user_text: str,
    *,
    chat_id: str,
    sender_id: str,
    message_thread_id: int | None,
    role: str,
) -> str | None:
    text = str(user_text or "").strip()
    low = text.lower()
    if not text:
        return None
    if role != "admin":
        return None

    bind_group_cmds = {"綁定本群通知", "通知綁定本群", "設定通知到本群", "/bind_group_notify"}
    if text in bind_group_cmds or low in bind_group_cmds:
        ok, msg = _telegram_apply_group_notify_binding(chat_id=chat_id, sender_id=sender_id)
        return msg if ok else msg

    if text in {"通知設定", "顯示通知設定", "/notify_status"} or low in {"/notify_status"}:
        return _telegram_notify_settings_text()

    topic_val = ""
    if text.startswith("綁定主題 "):
        topic_val = text.replace("綁定主題 ", "", 1).strip()
    elif low.startswith("/bind_topic "):
        topic_val = text.split(" ", 1)[1].strip() if " " in text else ""
    if topic_val:
        if not message_thread_id:
            return "⚠️ 請在要綁定的 Topic 裡執行此指令（需要 message_thread_id）。"
        ok, msg = _telegram_bind_topic(
            chat_id=chat_id,
            sender_id=sender_id,
            topic_raw=topic_val,
            thread_id=int(message_thread_id),
        )
        return msg if ok else msg

    setup_cmds = {"建立magi主題", "自動建立magi主題", "建立主題", "/setup_topics"}
    if low in setup_cmds or text in {"建立MAGI主題", "自動建立MAGI主題"}:
        ok, msg = _telegram_setup_topics(chat_id=chat_id, sender_id=sender_id)
        return msg if ok else msg

    return None


def _load_openclaw_telegram_token() -> str:
    """讀取 TG bot token — 純環境變數，不再依賴 openclaw.json。"""
    return (os.environ.get("OPENCLAW_TELEGRAM_BOT_TOKEN") or "").strip()


def _load_admin_telegram_ids() -> list[str]:
    """讀取 TG admin IDs — 純環境變數，不再依賴 openclaw.json。"""
    return [
        x.strip()
        for x in (os.environ.get("MAGI_ADMIN_TELEGRAM_IDS") or "").split(",")
        if x.strip()
    ]


def _load_notify_telegram_ids() -> list[str]:
    ids = [
        x.strip()
        for x in (os.environ.get("MAGI_NOTIFY_TELEGRAM_IDS") or "").split(",")
        if x.strip()
    ]
    try:
        state = _load_telegram_channel_state()
        notify_to = state.get("notifyTo") or []
        if isinstance(notify_to, list):
            ids.extend([str(x).strip() for x in notify_to if str(x).strip()])
    except Exception:
        logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 9403, exc_info=True)
    out: list[str] = []
    seen: set[str] = set()
    for x in ids:
        if x and x not in seen:
            seen.add(x)
            out.append(x)
    return out


def _send_telegram_text(text: str) -> bool:
    msg = str(text or "").strip()
    if not msg:
        return False
    try:
        from skills.ops.red_phone import send_telegram_push_with_status  # lazy import

        st = send_telegram_push_with_status(
            msg,
            severity="warning",
            source="api_server",
            topic_key="alert",
            queue_on_fail=True,
        ) or {}
        if bool(st.get("telegram")) or bool(st.get("queued")):
            return True
    except Exception:
        logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 9430, exc_info=True)

    token = _load_openclaw_telegram_token()
    notify_ids = _load_notify_telegram_ids()
    if not token or not notify_ids:
        return False
    payload = {"text": msg}
    body = json.dumps(payload, ensure_ascii=False).encode("utf-8")
    ok_any = False
    for chat_id in notify_ids:
        try:
            req = urllib.request.Request(
                f"https://api.telegram.org/bot{token}/sendMessage?chat_id={chat_id}",
                data=body,
                headers={"Content-Type": "application/json"},
                method="POST",
            )
            with urllib.request.urlopen(req, timeout=8):
                pass
            ok_any = True
        except Exception as e:
            logger.warning(f"⚠️ Telegram notify failed: {e}")
    return ok_any


def _load_discord_notify_target() -> tuple[str, str]:
    token = (os.environ.get("DISCORD_BOT_TOKEN") or "").strip()
    channel_id = (os.environ.get("DISCORD_CHANNEL_ID") or "").strip()
    if not channel_id and os.path.exists(DISCORD_LAST_CHANNEL_FILE):
        try:
            data = json.loads(Path(DISCORD_LAST_CHANNEL_FILE).read_text(encoding="utf-8")) or {}
            channel_id = str(data.get("channel_id") or "").strip()
        except Exception:
            channel_id = ""
    return token, channel_id


def _send_discord_text(text: str) -> bool:
    # Discord 僅用於互動指令與聊天，系統通知不發送到 Discord。
    return False
    # noinspection PyUnreachableCode
    try:
        from skills.ops.red_phone import send_discord_alert  # lazy import
        if send_discord_alert(str(text or "").strip(), severity="warning"):
            return True
    except Exception:
        logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 9476, exc_info=True)

    token, channel_id = _load_discord_notify_target()
    if not token or not channel_id:
        return False
    payload = {"content": str(text or "").strip()[:1900]}
    body = json.dumps(payload, ensure_ascii=False).encode("utf-8")
    try:
        req = urllib.request.Request(
            f"https://discord.com/api/v10/channels/{channel_id}/messages",
            data=body,
            headers={
                "Content-Type": "application/json",
                "Authorization": f"Bot {token}",
            },
            method="POST",
        )
        with urllib.request.urlopen(req, timeout=8):
            pass
        return True
    except Exception as e:
        logger.warning(f"⚠️ Discord notify failed: {e}")
        return False


def _notify_admin_telegram_once(text: str, dedupe_sec: int = 1800) -> None:
    token = _load_openclaw_telegram_token()
    notify_ids = _load_notify_telegram_ids()
    if not token or not notify_ids:
        return

    now = int(time.time())
    try:
        if os.path.exists(LINE_QUOTA_ALERT_FILE):
            prev = json.loads(Path(LINE_QUOTA_ALERT_FILE).read_text(encoding="utf-8"))
            last = int(prev.get("ts") or 0)
            if now - last < dedupe_sec:
                return
    except Exception:
        logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 9515, exc_info=True)

    _send_telegram_text(text)

    try:
        Path(LINE_QUOTA_ALERT_FILE).write_text(
            json.dumps({"ts": now, "reason": "line_quota"}, ensure_ascii=False),
            encoding="utf-8",
        )
    except Exception:
        logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 9525, exc_info=True)


def _enqueue_line_delayed(user_id: str, phase: str, reason: str, preview: str) -> None:
    try:
        items = []
        if os.path.exists(LINE_DELAY_QUEUE_FILE):
            raw = json.loads(Path(LINE_DELAY_QUEUE_FILE).read_text(encoding="utf-8"))
            if isinstance(raw, list):
                items = raw
        items.append(
            {
                "ts": int(time.time()),
                "user_id": str(user_id or "").strip(),
                "phase": str(phase or "").strip(),
                "reason": str(reason or "").strip(),
                "preview": str(preview or "").strip(),
                "status": "queued",
            }
        )
        if len(items) > 300:
            items = items[-300:]
        Path(LINE_DELAY_QUEUE_FILE).write_text(
            json.dumps(items, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
    except Exception as e:
        logger.warning(f"⚠️ Failed to enqueue LINE delayed item: {e}")


def _line_quota_active(window_sec: int | None = None) -> bool:
    """
    Best-effort detector for recent LINE quota exhaustion.
    If quota alert was raised recently, skip repeated LINE send attempts and
    use cross-channel notice + delayed queue instead.
    """
    ws = int(window_sec or os.environ.get("LINE_QUOTA_ACTIVE_WINDOW_SEC", "600") or "600")
    ws = max(60, ws)
    try:
        if not os.path.exists(LINE_QUOTA_ALERT_FILE):
            return False
        obj = json.loads(Path(LINE_QUOTA_ALERT_FILE).read_text(encoding="utf-8")) or {}
        ts = int(obj.get("ts") or 0)
        if ts <= 0:
            return False
        return (int(time.time()) - ts) < ws
    except Exception:
        return False


def _should_emit_line_delay_notice(user_id: str, phase: str, dedupe_sec: int = 180) -> bool:
    key = f"{str(user_id or '').strip()}::{str(phase or '').strip()}"
    now = int(time.time())
    last = int(_LINE_DELAYED_ALERT_TS.get(key) or 0)
    if now - last < max(30, int(dedupe_sec)):
        return False
    _LINE_DELAYED_ALERT_TS[key] = now
    return True


def _fanout_line_delayed_notice(user_id: str, phase: str, preview: str) -> None:
    if not _should_emit_line_delay_notice(user_id=user_id, phase=phase, dedupe_sec=180):
        return
    body = (
        "⚠️ LINE 回覆已延遲（額度用盡，已加入佇列）\n"
        f"來源使用者：{user_id}\n"
        f"流程：{phase}\n"
        f"內容摘要：{(preview or '（空）')[:260]}"
    )
    _send_telegram_text(body)


def _handle_line_send_failure(err: Exception, user_id: str, phase: str, failed_text: str = "") -> None:
    """
    Detect quota/429 failures and alert admin on Telegram so LINE silence is explainable.
    """
    status = None
    try:
        status = int(getattr(err, "status_code", 0) or 0)
    except Exception:
        status = None
    msg = str(err or "")
    low = msg.lower()
    if status == 429 or ("monthly limit" in low):
        logger.error(f"⛔ LINE quota reached (phase={phase}, user={user_id}).")
        preview = (failed_text or "").strip() or _last_line_outgoing_preview(user_id)
        _enqueue_line_delayed(
            user_id=user_id,
            phase=phase,
            reason=f"line_quota_{status or 429}",
            preview=preview[:500],
        )
        _fanout_line_delayed_notice(user_id=user_id, phase=phase, preview=preview)
        _notify_admin_telegram_once(
            "⛔ LINE 額度已達上限（API 429），LINE 可能暫時無法回覆。請先改用 TG/DC 或更換可用 token。",
            dedupe_sec=1800,
        )


def _normalize_line_output_text(text: str, skip_llm: bool = False) -> str:
    s = (text or "").strip()
    if not s:
        return s
    try:
        if _normalize_output_text:
            if skip_llm:
                # Deterministic replacements only — skip slow TAIDE LLM review.
                from api.tw_output_guard import _opencc_s2twp, _replace_mainland_terms, _strip_internal_leaks, _limit_message_for_platform, strip_markdown_for_chat
                s = _opencc_s2twp(s)
                s, _ = _replace_mainland_terms(s)
                s, _ = _strip_internal_leaks(s)
                s = strip_markdown_for_chat(s)
                return _limit_message_for_platform(s, platform="LINE")
            return _normalize_output_text(s, platform="LINE")
    except Exception as e:
        logger.warning(f"⚠️ Taiwan wording guard skipped: {e}")
    return s


def _register_orchestrator_notifications():
    """
    Wire Orchestrator async notifications to LINE push so long-running tasks can
    provide progress updates (without depending on reply_token).
    """
    try:
        def _cb(uid: str, msg: str, platform: str = "LINE"):
            if (platform or "").upper() == "WEB":
                WEB_NOTIFICATIONS[str(uid)].append(msg)
                return
            if (platform or "").upper() != "LINE":
                return
            _line_push_text(uid, msg)

        orchestrator.register_callback(_cb)
        logger.info("🔔 Orchestrator notifications enabled for LINE push.")
    except Exception as e:
        logger.warning(f"⚠️ Failed to register orchestrator notifications: {e}")


def _line_send_text(event, user_id: str, text: str, prefer_push: bool = False, skip_llm: bool = False) -> bool:
    """
    Best-effort delivery:
    - If prefer_push or text is long: push (chunked).
    - Else: reply; if reply_token expired/invalid or any failure, fall back to push.
    This prevents "opened then no response" caused by reply_token expiry.
    """
    s = _normalize_line_output_text(text, skip_llm=skip_llm)
    if not s:
        return True
    _remember_last_line_outgoing(user_id, s)

    # If the content is too long for chat, export to TXT and send a link (best effort).
    if EXPORT_LONG_TEXT and len(s) >= max(2000, int(EXPORT_TEXT_THRESHOLD)):
        exported = _export_text_to_static(s, prefix="casper_reply")
        if exported.get("success"):
            url = (exported.get("url") or "").strip()
            if url:
                msg = (
                    "內容比較長，我先幫你整理成 TXT 檔方便下載：\n"
                    f"{url}\n\n"
                    "如果你點不開這個連結，請把你目前對外的公開網址（例如網域/TS Funnel）貼給我，我會改用那個網址產生連結。"
                )
            else:
                msg = (
                    "內容比較長，我已輸出成 TXT 檔（目前尚未取得可公開下載的網址）。\n"
                    f"檔案位置：{exported.get('path')}\n"
                    "如果你有對外的公開網址（例如網域/TS Funnel），設定 `MAGI_PUBLIC_BASE_URL` 後我就能改用連結傳給你。"
                )
            return _line_push_text(user_id, msg, is_chat_reply=True)

    # Avoid reply_message size limit and reduce chance of token expiry.
    if prefer_push or len(s) > 4200:
        return _line_push_text(user_id, s, is_chat_reply=True)

    try:
        line_bot_api.reply_message(event.reply_token, TextSendMessage(text=s))
        return True
    except Exception as e:
        # Typical failure: reply_token expired (400) or already used. Push is still allowed.
        try:
            if isinstance(e, LineBotApiError):
                logger.warning(f"⚠️ LINE reply failed (status={getattr(e, 'status_code', '?')}), fallback push: {e}")
            else:
                logger.warning(f"⚠️ LINE reply failed, fallback push: {e}")
        except Exception:
            logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 9710, exc_info=True)
        _handle_line_send_failure(e, user_id=user_id, phase="reply", failed_text=s)
        return _line_push_text(user_id, s, is_chat_reply=True)


def _line_send_messages(event, user_id: str, messages, prefer_push: bool = False) -> bool:
    """
    Best-effort delivery for multi-message payloads (e.g., text + image).
    Falls back to push on reply failure.
    """
    if not messages:
        return True
    try:
        guarded = []
        text_parts = []
        for m in messages:
            if isinstance(m, TextSendMessage):
                txt = _normalize_line_output_text(getattr(m, "text", ""))
                guarded.append(TextSendMessage(text=txt))
                if txt:
                    text_parts.append(txt)
            else:
                guarded.append(m)
        messages = guarded
        if text_parts:
            _remember_last_line_outgoing(user_id, "\n".join(text_parts))
    except Exception as guard_err:
        logger.warning(f"⚠️ LINE message guard skipped: {guard_err}")
    try:
        if prefer_push:
            line_bot_api.push_message(user_id, messages)
            return True
        line_bot_api.reply_message(event.reply_token, messages)
        return True
    except Exception as e:
        logger.warning(f"⚠️ LINE send messages failed, fallback push: {e}")
        _handle_line_send_failure(e, user_id=user_id, phase="reply_messages")
        try:
            line_bot_api.push_message(user_id, messages)
            return True
        except Exception as push_err:
            logger.error(f"❌ LINE push messages failed: {push_err}")
            _handle_line_send_failure(push_err, user_id=user_id, phase="push_messages")
            return False


def process_message_async(event, user_id, user_text, attachment, role="user", long_task: bool | None = None, already_acked: bool = False):
    # OBS-1: correlation ID  /  OBS-2: latency tracking
    correlation_id = f"magi-{uuid.uuid4().hex[:12]}"
    _start_ts = time.monotonic()
    try:
        if long_task is None:
            long_task = _likely_long_task(user_text, attachment)
        if long_task and not already_acked:
            # Reply quickly (within reply_token lifetime), then push final result.
            ack_msg = "⏳ 已收到，正在處理中。完成後我會用推播回覆結果。"
            if attachment and attachment.get("type") in ("file", "audio"):
                try:
                    att_path = attachment.get("path", "")
                    att_size = os.path.getsize(att_path) if att_path and os.path.exists(att_path) else 0
                    if att_size > 0:
                        from api.orchestrator import Orchestrator
                        ack_msg = Orchestrator.estimate_file_processing_time(
                            file_size_bytes=att_size,
                            filename=attachment.get("filename", ""),
                            prompt=user_text or "",
                        )
                except Exception:
                    logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 9778, exc_info=True)
            _line_send_text(event, user_id, ack_msg, prefer_push=False)

        if attachment:
            logger.info(f"📎 Processing attachment for {user_id}: {attachment['type']} cid={correlation_id}")
            response_text = orchestrator.process_message(
                user_id,
                user_text,
                platform="LINE",
                attachment=attachment,
                role=role,
                correlation_id=correlation_id,
            )
        else:
            logger.info(f"📩 Processing text for {user_id}: {user_text} cid={correlation_id}")
            response_text = orchestrator.process_message(user_id, user_text, platform="LINE", role=role, correlation_id=correlation_id)

        if response_text:
            try:
                orchestrator.record_assistant_reply(user_id, response_text)
            except Exception as track_err:
                logger.warning(f"⚠️ Failed to track assistant reply for {user_id}: {track_err}")

            # If LINE is still in quota-limited window, avoid repeated failed sends.
            if _line_quota_active():
                preview = _normalize_line_output_text(response_text)
                _enqueue_line_delayed(
                    user_id=user_id,
                    phase="result",
                    reason="line_quota_active",
                    preview=(preview or "")[:500],
                )
                _fanout_line_delayed_notice(user_id=user_id, phase="result", preview=preview)
                return

            if "|||IMAGE_PATH|||" in response_text:
                try:
                    text_part, image_path = response_text.split("|||IMAGE_PATH|||", 1)
                    image_path = (image_path or "").strip()

                    if not image_path or not os.path.exists(image_path):
                        msg = f"{text_part}\n⚠️ Image file not found at path."
                        _line_send_text(event, user_id, msg, prefer_push=long_task)
                        return

                    # Serve image via local static files + Cloudflare Tunnel (no Imgur needed)
                    image_url = _public_url_for_local_file(image_path)
                    if not image_url:
                        msg = f"{text_part}\n⚠️ 無法建立圖片公開連結（tunnel 可能未啟動）。"
                        _line_send_text(event, user_id, msg, prefer_push=long_task)
                        return
                    logger.info(f"🖼️ Serving image via tunnel: {image_url}")
                    messages = [TextSendMessage(text=text_part)]
                    if image_url:
                        messages.append(
                            ImageSendMessage(
                                original_content_url=image_url,
                                preview_image_url=image_url,
                            )
                        )
                    ok = _line_send_messages(event, user_id, messages, prefer_push=long_task)
                    if not ok:
                        _line_push_text(user_id, f"{text_part}\n⚠️ 圖片傳送失敗（已嘗試 reply/push）。")
                except Exception as img_err:
                    logger.error(f"❌ Failed to send image to LINE: {img_err}")
                    msg = f"{response_text}\n(Image send failed: {img_err})"
                    _line_send_text(event, user_id, msg, prefer_push=long_task)
            elif "|||FILE_PATH|||" in response_text:
                try:
                    text_part, file_path = response_text.split("|||FILE_PATH|||", 1)
                    file_path = (file_path or "").strip()
                    file_url = _public_url_for_local_file(file_path)
                    if file_url:
                        body = (text_part or "").strip()
                        msg = (body + "\n\n" if body else "") + f"📎 檔案下載：{file_url}"
                        _line_send_text(event, user_id, msg, prefer_push=long_task)
                    else:
                        msg = f"{(text_part or '').strip()}\n⚠️ 檔案已產生，但目前無法建立公開下載連結。"
                        _line_send_text(event, user_id, msg, prefer_push=long_task)
                except Exception as file_err:
                    logger.error(f"❌ Failed to send file link to LINE: {file_err}")
                    _line_send_text(event, user_id, "❌ 檔案處理失敗，請稍後再試。", prefer_push=long_task)
            else:
                # Static system responses (command tables, status) don't need LLM review.
                _skip = response_text.lstrip().startswith(("🛠️", "📊", "✅ 系統", "⚡", "🔧", "📋"))
                _line_send_text(event, user_id, response_text, prefer_push=long_task, skip_llm=_skip)
    except Exception as e:
        logger.error(f"❌ Async Processing Error: {e}")
        try:
            # Best effort: reply if possible, otherwise push.
            _line_send_text(event, user_id, "❌ 系統暫時忙碌，請稍後再試。", prefer_push=False)
        except Exception:
            logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 9870, exc_info=True)
    finally:
        # OBS-2: record processing latency
        elapsed_ms = int((time.monotonic() - _start_ts) * 1000)
        _append_channel_delivery_audit({
            "platform": "LINE",
            "kind": "latency",
            "user_id": str(user_id or ""),
            "correlation_id": correlation_id,
            "latency_ms": elapsed_ms,
        })

@handler.add(MessageEvent, message=TextMessage)
def handle_message(event):
    user_id = event.source.user_id
    user_text = event.message.text
    user_text_norm = (user_text or "").strip().lower()
    _cleanup_user_context()
    _record_last_line_sender(event)

    # Inject request context for structured logging (cleared at next request)
    from skills.ops.structured_log import set_request_context
    set_request_context(request_id=uuid.uuid4().hex[:12], user_id=user_id, platform="LINE")

    # Fast-path health probe: avoid LLM queueing for simple connectivity checks.
    if user_text_norm in {"連線測試", "連線測試。", "ping", "test", "連線", "測試連線"}:
        try:
            from datetime import datetime
            ts = datetime.now().strftime("%H:%M:%S")
            _line_send_text(event, user_id, f"✅ 連線正常（{ts}）", prefer_push=False)
        except Exception as probe_err:
            logger.warning(f"⚠️ Fast-path LINE probe reply failed: {probe_err}")
            _line_push_text(user_id, "✅ 連線正常")
        return 'OK'

    # Intercept LAF captcha replies (admin human-in-the-loop) before Orchestrator.
    try:
        if _maybe_handle_laf_captcha_reply(event, user_id, user_text):
            return 'OK'
    except Exception as _cap_err:
        logger.warning(f"⚠️ LAF captcha intercept failed: {_cap_err}")

    # Generic captcha broker (used by other modules, e.g. file review / transcripts)
    try:
        if _maybe_handle_generic_captcha_reply(event, user_id, user_text):
            return 'OK'
    except Exception as _cap2_err:
        logger.warning(f"⚠️ Generic captcha intercept failed: {_cap2_err}")

    # If quota was recently hit, queue user request immediately and notify via TG/DC.
    if _line_quota_active():
        preview = re.sub(r"\s+", " ", (user_text or "")).strip()[:500]
        _enqueue_line_delayed(
            user_id=user_id,
            phase="incoming",
            reason="line_quota_active",
            preview=preview,
        )
        _fanout_line_delayed_notice(user_id=user_id, phase="incoming", preview=preview)
    
    # Role Check
    role = "user"
    if user_id in ADMIN_LINE_USER_IDS:
        role = "admin"
    elif LINE_AUTO_ADMIN_LAST_SENDER and (not ADMIN_LINE_USER_IDS) and user_id and user_id == _load_last_line_sender_user_id():
        role = "admin"
        
    # Check Context
    attachment = user_context.get(user_id)
    if attachment:
        ts = float(attachment.get("timestamp", 0) or 0)
        if ts and (time.time() - ts > CONTEXT_TTL_SECONDS):
            stale_path = attachment.get("path")
            if stale_path and os.path.exists(stale_path):
                _safe_remove_tmp(stale_path)
            attachment = None
        user_context.pop(user_id, None)

    recent_followup = False
    try:
        recent_followup = orchestrator.has_recent_attachment_followup(user_id, "LINE", user_text)
    except Exception as recent_err:
        logger.warning(f"⚠️ LINE recent attachment probe failed: {recent_err}")

    # If this looks like a long task, ACK synchronously before returning from the webhook.
    # This avoids "didn't respond" cases due to background-thread scheduling or reply_token expiry.
    long_task = _likely_long_task(user_text, attachment) or recent_followup
    if long_task:
        try:
            ack_msg = "⏳ 已收到，正在處理中。完成後我會用推播回覆結果。"
            if attachment and attachment.get("type") in ("file", "audio", "image"):
                try:
                    att_path = attachment.get("path", "")
                    att_size = os.path.getsize(att_path) if att_path and os.path.exists(att_path) else 0
                    # Use original filename (from LINE event.message.file_name), fall back to basename of temp path
                    att_fname = attachment.get("filename") or os.path.basename(att_path) or "附件"
                    if att_size > 0:
                        from api.orchestrator import Orchestrator
                        ack_msg = Orchestrator.estimate_file_processing_time(
                            file_size_bytes=att_size,
                            filename=att_fname,
                            prompt=user_text or "",
                            file_path=att_path,
                        )
                except Exception:
                    logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 9975, exc_info=True)
            _line_send_text(event, user_id, ack_msg, prefer_push=False)
        except Exception as ack_err:
            logger.warning(f"⚠️ Failed to send immediate ACK: {ack_err}")

        # Ensure Orchestrator can push progress updates during long tasks.
        _register_orchestrator_notifications()

    if (attachment and attachment.get("type") in ("file", "audio")) or recent_followup:
        try:
            _enqueue_attachment_job(
                platform_name="LINE",
                user_id=user_id,
                role=role,
                user_text=user_text,
                attachment=attachment,
            )
            if attachment:
                att_path = str(attachment.get("path") or "").strip()
                if att_path and os.path.exists(att_path):
                    _safe_remove_tmp(att_path)
            return 'OK'
        except Exception as enqueue_err:
            logger.error(f"❌ LINE attachment job enqueue failed: {enqueue_err}")

    # Run in bounded background pool
    _CHANNEL_BG_EXECUTOR.submit(process_message_async, event, user_id, user_text, attachment, role, long_task, True)
    
    # Return immediately to avoid LINE timeout
    return 'OK'

@handler.add(MessageEvent, message=(ImageMessage, AudioMessage, FileMessage))
def handle_content(event):
    user_id = event.source.user_id
    message_id = event.message.id
    msg_type = event.message.type
    _record_last_line_sender(event)

    logger.info(f"📂 Received {msg_type} from {user_id}")

    # specific handling for file name if available
    file_name = getattr(event.message, 'file_name', f"{message_id}.{msg_type}")
    
    # Download Content
    message_content = line_bot_api.get_message_content(message_id)
    
    # Check message type to determine extension/path
    ext = "bin"
    if msg_type == "image": ext = "jpg"
    elif msg_type == "audio": ext = "m4a"
    elif msg_type == "file": 
        # try to get extension from filename
        if "." in file_name:
            ext = file_name.split(".")[-1]
    
    temp_path = f"/tmp/{message_id}.{ext}"
    
    with open(temp_path, 'wb') as fd:
        for chunk in message_content.iter_content():
            fd.write(chunk)
            
    logger.info(f"💾 Saved to {temp_path}")
    
    attachment_payload = {
        "type": msg_type,
        "path": temp_path,
        "filename": file_name,
        "timestamp": time.time(),
    }
    user_context[user_id] = attachment_payload
    if msg_type in {"image", "file"}:
        try:
            durable_recent = _persist_attachment_payload(attachment_payload, prefix=f"line_recent_{message_id}")
            if durable_recent:
                orchestrator.remember_recent_attachment(
                    user_id=user_id,
                    platform="LINE",
                    attachment=durable_recent,
                    source_message="",
                )
        except Exception as recent_err:
            logger.warning(f"⚠️ Failed to persist LINE recent attachment {message_id}: {recent_err}")
    
    # Reply asking for instruction
    reply_map = {
        "image": "📸 圖片已接收。請告訴我您想做什麼？(例如：描述這張圖、翻譯文字)",
        "file": f"DFC 檔案 ({file_name}) 已接收。請下達指令。"
    }
    
    if msg_type == "audio":
        # Route voice through orchestrator so we can consistently output timestamped TXT.
        _line_send_text(event, user_id, "⏳ 已收到語音，正在進行逐字稿處理（含時間戳/TXT）。完成後我會用推播回覆。", prefer_push=False)

        role = "user"
        if user_id in ADMIN_LINE_USER_IDS:
            role = "admin"
        elif LINE_AUTO_ADMIN_LAST_SENDER and (not ADMIN_LINE_USER_IDS) and user_id and user_id == _load_last_line_sender_user_id():
            role = "admin"

        voice_attachment = dict(attachment_payload)
        voice_prompt = "請轉換成逐字稿，附上時間戳記，並輸出TXT檔。"
        try:
            _enqueue_attachment_job(
                platform_name="LINE",
                user_id=user_id,
                role=role,
                user_text=voice_prompt,
                attachment=voice_attachment,
            )
        except Exception as enqueue_err:
            logger.error(f"❌ Voice job enqueue failed: {enqueue_err}")

            def _run_voice_fallback():
                try:
                    process_message_async(
                        event,
                        user_id,
                        voice_prompt,
                        voice_attachment,
                        role=role,
                        long_task=True,
                        already_acked=True,
                    )
                except Exception as e:
                    logger.error(f"❌ Voice processing fallback error: {e}")
                    _line_push_text(user_id, f"❌ 語音處理錯誤: {str(e)}")
                finally:
                    try:
                        if os.path.exists(temp_path):
                            _safe_remove_tmp(temp_path)
                    except Exception:
                        logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 10106, exc_info=True)
                    user_context.pop(user_id, None)

            _CHANNEL_BG_EXECUTOR.submit(_run_voice_fallback)
        else:
            try:
                if os.path.exists(temp_path):
                    _safe_remove_tmp(temp_path)
            except Exception:
                logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 10115, exc_info=True)
            user_context.pop(user_id, None)
        return
    
    else:
        reply_text = reply_map.get(msg_type, "檔案已接收。請指示下一步。")
    
    _line_send_text(event, user_id, reply_text, prefer_push=False)

@app.route("/health", methods=['GET'])
def health():
    """System health check with component diagnostics."""
    import time as _time
    from skills.bridge.http_pool import get_session as _get_sess
    _rq = _get_sess()
    checks = {"status": "operational", "timestamp": _time.time()}

    # oMLX (primary inference engine)
    try:
        _r = _rq.get("http://127.0.0.1:8080/v1/models", timeout=3)
        _models = [m.get("id", "") for m in (_r.json() or {}).get("data", [])]
        checks["omlx"] = {"ok": _r.status_code == 200, "models": _models}
    except Exception:
        checks["omlx"] = {"ok": False}

    # MariaDB
    _conn = None
    try:
        _conn = mysql.connector.connect(**DB_CONFIG, connection_timeout=3, use_pure=True)
        checks["db"] = {"ok": _conn.is_connected()}
    except Exception as _e:
        checks["db"] = {"ok": False, "detail": str(_e)[:80]}
    finally:
        if _conn:
            try:
                _conn.close()
            except Exception:
                logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 10152, exc_info=True)

    # System resources
    try:
        import psutil
        _vm = psutil.virtual_memory()
        _du = psutil.disk_usage("/")
        checks["system"] = {
            "cpu_percent": psutil.cpu_percent(interval=0.1),
            "memory_percent": _vm.percent,
            "memory_available_gb": round(_vm.available / (1024**3), 1),
            "disk_percent": _du.percent,
            "disk_free_gb": round(_du.free / (1024**3), 1),
        }
    except Exception:
        logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 10167, exc_info=True)

    # FAISS vector index — defer loading during startup to avoid OOM
    _uptime = _time.time() - _SERVER_START_TIME
    if _uptime < 60:
        checks["faiss"] = {"ok": True, "deferred": True, "reason": "startup_grace_period"}
    else:
        try:
            from skills.memory.faiss_index import FAISSMemoryIndex
            _idx = FAISSMemoryIndex.get_instance()
            checks["faiss"] = {"ok": True, "vectors": getattr(_idx, 'total', getattr(_idx, 'ntotal', 0))}
        except Exception:
            checks["faiss"] = {"ok": False}

    # Attachment jobs
    try:
        if _jq:
            checks["attachment_jobs"] = _jq.stats()
        else:
            job_ids = _list_attachment_job_ids()
            pending = sum(1 for jid in job_ids if _read_attachment_job(jid).get("status") in ("queued", "running"))
            checks["attachment_jobs"] = {"total": len(job_ids), "active": pending}
    except Exception:
        logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 10190, exc_info=True)

    # NAS mounts
    try:
        from api.nas_mount_guard import _is_mounted, _SHARES
        checks["nas"] = {vol.split("/")[-1]: _is_mounted(vol) for _, vol in _SHARES}
    except Exception:
        logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 10197, exc_info=True)

    # Uptime
    try:
        checks["uptime_seconds"] = round(_time.time() - _SERVER_START_TIME, 0)
    except Exception:
        logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 10203, exc_info=True)

    checks["status"] = "operational" if checks.get("omlx", {}).get("ok") else "degraded"
    return jsonify(checks), 200

@app.route('/api/transcribe', methods=['POST'])
def transcribe_audio():
    """Endpoint for audio transcription (Balthasar)."""
    # API Key Authentication (Bypass login for nodes)
    api_key = (request.headers.get("X-MAGI-API-KEY") or "").strip()
    api_key_ok = bool(EXPECTED_MAGI_API_KEY) and hmac.compare_digest(api_key, EXPECTED_MAGI_API_KEY)
    if not api_key_ok and not current_user.is_authenticated:
        return jsonify({"error": "Unauthorized"}), 401
            
    if 'file' not in request.files:
        return jsonify({"error": "No file part"}), 400
        
    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No selected file"}), 400
        
    if file:
        try:
            # Save temp file
            safe_filename = "".join([c for c in file.filename if c.isalnum() or c in "._-"]) or "audio.wav"
            filename = f"audio_{int(time.time())}_{safe_filename}" 
            filepath = os.path.join("/tmp", filename)
            file.save(filepath)
            
            logger.info(f"🎤 Received audio for transcription: {filepath}")
            
            from skills.bridge.balthasar_bridge import transcribe
            language = str(request.form.get("language") or "").strip() or None
            taigi_hint_raw = str(request.form.get("taigi_hint") or "").strip().lower()
            taigi_hint = taigi_hint_raw in {"1", "true", "yes", "on"}
            result = transcribe(filepath, language=language, taigi_hint=taigi_hint)
            
            # Clean up
            if os.path.exists(filepath):
                _safe_remove_tmp(filepath)
                
            return jsonify(result)
            
        except Exception as e:
            logger.error(f"❌ Transcription endpoint error: {e}")
            return jsonify({"error": str(e)}), 500

# Start Telegram polling fallback after all helpers are defined.
_resume_pending_attachment_jobs()
_start_telegram_polling_fallback()

# Pre-load FAISS index in background so first chat doesn't pay 30s+ penalty
def _preload_faiss():
    try:
        from skills.memory.mem_bridge import _get_faiss_index
        idx = _get_faiss_index()
        if idx:
            logger.info("✅ FAISS index pre-loaded: %d vectors", getattr(idx, 'total', 0))
    except Exception as e:
        logger.warning("⚠️ FAISS pre-load failed (non-fatal): %s", e)

threading.Thread(target=_preload_faiss, daemon=True, name="faiss-preload").start()

# Startup cleanup: remove old export files (>30 days)
try:
    _n_cleaned = cleanup_old_exports(days=30)
    if _n_cleaned:
        logger.info("🧹 Startup: cleaned %d old exports", _n_cleaned)
except Exception:
    pass

# Warm up TAIDE model on oMLX so first chat doesn't pay model-load penalty
def _warmup_omlx():
    try:
        import time as _t
        _t.sleep(2)  # let oMLX finish its own startup
        from skills.bridge.http_pool import get_session
        _model = os.environ.get("CASPER_LOCAL_MODEL", "TAIDE-12b-Chat-mlx-4bit")
        r = get_session().post("http://127.0.0.1:8080/v1/chat/completions", json={
            "model": _model,
            "messages": [{"role": "user", "content": "ping"}],
            "max_tokens": 1, "temperature": 0,
        }, timeout=60)
        if r.status_code == 200:
            logger.info("✅ oMLX TAIDE model warmed up")
        else:
            logger.warning("⚠️ oMLX warmup got %d", r.status_code)
    except Exception as e:
        logger.warning("⚠️ oMLX warmup failed (non-fatal): %s", e)

threading.Thread(target=_warmup_omlx, daemon=True, name="omlx-warmup").start()

# --- Cloudflare Quick Tunnel for LINE webhook ---
def _is_cloudflared_alive() -> bool:
    """Check if cloudflared tunnel process is actually running (not pgrep self-match)."""
    import subprocess
    try:
        result = subprocess.run(
            ["pgrep", "-f", "/opt/homebrew/bin/cloudflared tunnel"],
            capture_output=True, timeout=3,
        )
        return result.returncode == 0
    except Exception:
        return False


def _ensure_cloudflared():
    """Start cloudflared if not running and always register webhook with LINE."""
    import subprocess, re as _re, time as _time
    try:
        log_path = os.path.join(os.path.dirname(__file__), "..", "logs", "cloudflared.log")
        os.makedirs(os.path.dirname(log_path), exist_ok=True)
        already_running = False

        # Count running cloudflared instances; kill all if >1 (prevent accumulation)
        try:
            result = subprocess.run(
                ["pgrep", "-f", "/opt/homebrew/bin/cloudflared tunnel"],
                capture_output=True, text=True, timeout=3,
            )
            cf_pids = [p.strip() for p in (result.stdout or "").strip().splitlines() if p.strip()]
        except Exception:
            cf_pids = []

        if len(cf_pids) > 1:
            logger.warning("☁️ Found %d cloudflared instances, killing all to restart cleanly", len(cf_pids))
            try:
                subprocess.run(["pkill", "-f", "/opt/homebrew/bin/cloudflared tunnel"],
                               capture_output=True, timeout=3)
                _time.sleep(1)
            except Exception:
                pass
            cf_pids = []

        if len(cf_pids) == 1:
            # Check if the log still has the URL (not truncated)
            try:
                with open(log_path) as f:
                    content = f.read()
                if _re.search(r'https://[a-z0-9-]+\.trycloudflare\.com', content):
                    logger.info("☁️ cloudflared already running (pid=%s)", cf_pids[0])
                    already_running = True
                else:
                    logger.warning("☁️ cloudflared running but log empty, restarting")
                    subprocess.run(["kill", cf_pids[0]], capture_output=True, timeout=3)
                    _time.sleep(1)
            except Exception:
                logger.info("☁️ cloudflared already running (pid=%s)", cf_pids[0])
                already_running = True

        if not already_running:
            # 殺掉所有殘留的 cloudflared 實例（防止多重累積）
            try:
                subprocess.run(["pkill", "-f", "/opt/homebrew/bin/cloudflared tunnel"],
                               capture_output=True, timeout=3)
            except Exception:
                logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 10317, exc_info=True)
            logger.info("☁️ Starting cloudflared tunnel...")
            _cf_log_fh = open(log_path, "w")  # kept open for cloudflared's lifetime
            _cf_proc = subprocess.Popen(
                ["/opt/homebrew/bin/cloudflared", "tunnel", "--url", "http://127.0.0.1:5002", "--no-autoupdate"],
                stdout=subprocess.DEVNULL, stderr=_cf_log_fh,
            )
            # Cleanup: close log file handle when cloudflared exits
            def _cleanup_cf_log(proc=_cf_proc, fh=_cf_log_fh):
                try:
                    proc.wait()
                finally:
                    try:
                        fh.close()
                    except Exception:
                        logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 10332, exc_info=True)
            import threading as _thr
            _thr.Thread(target=_cleanup_cf_log, daemon=True, name="cf-log-cleanup").start()

        def _register():
            cf_url = ""
            if already_running:
                # Read tunnel URL from existing log
                try:
                    with open(log_path) as f:
                        m = _re.search(r'https://[a-z0-9-]+\.trycloudflare\.com', f.read())
                        if m:
                            cf_url = m.group(0)
                except Exception:
                    logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 10346, exc_info=True)
            if not cf_url:
                # Wait for URL to appear (new tunnel or log not yet written)
                for _ in range(30):
                    _time.sleep(1)
                    try:
                        with open(log_path) as f:
                            m = _re.search(r'https://[a-z0-9-]+\.trycloudflare\.com', f.read())
                            if m:
                                cf_url = m.group(0)
                                break
                    except Exception:
                        logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 10358, exc_info=True)
            if not cf_url:
                logger.error("☁️ Could not get cloudflare tunnel URL after 30s")
                return
            webhook_url = f"{cf_url}/line/webhook"
            logger.info(f"☁️ Tunnel: {cf_url}")
            # Load LINE token
            token = os.environ.get("MAGI_LINE_CHANNEL_ACCESS_TOKEN", "")
            if not token:
                env_path = os.path.join(os.path.dirname(__file__), "..", ".env")
                try:
                    with open(env_path) as f:
                        for ln in f:
                            if ln.strip().startswith("MAGI_LINE_CHANNEL_ACCESS_TOKEN="):
                                token = ln.strip().split("=", 1)[1].strip().strip("\"'")
                                break
                except Exception:
                    logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 10375, exc_info=True)
            if not token:
                logger.warning("☁️ No LINE token, skipping webhook registration")
                return
            import urllib.request, urllib.parse, json
            if not already_running:
                _time.sleep(3)  # Wait for new tunnel to be routable
            # Check if LINE already points to this URL
            try:
                get_req = urllib.request.Request(
                    "https://api.line.me/v2/bot/channel/webhook/endpoint",
                    method="GET",
                    headers={"Authorization": f"Bearer {token}"},
                )
                with urllib.request.urlopen(get_req, timeout=10) as resp:
                    current = json.loads(resp.read())
                if current.get("endpoint") == webhook_url:
                    logger.info(f"☁️ LINE webhook already correct: {webhook_url}")
                    return
                logger.info(f"☁️ LINE webhook mismatch: {current.get('endpoint')} → {webhook_url}")
            except Exception:
                logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 10396, exc_info=True)  # Proceed to register anyway
            data = json.dumps({"endpoint": webhook_url}).encode()
            registered = False
            for attempt in range(3):
                try:
                    req = urllib.request.Request(
                        "https://api.line.me/v2/bot/channel/webhook/endpoint",
                        data=data, method="PUT",
                        headers={"Authorization": f"Bearer {token}", "Content-Type": "application/json"},
                    )
                    with urllib.request.urlopen(req, timeout=10) as resp:
                        logger.info(f"☁️ LINE webhook registered: {webhook_url} → {resp.status}")
                        registered = True
                        break
                except Exception as e:
                    logger.warning(f"☁️ LINE webhook registration attempt {attempt+1}/3 failed: {e}")
                    _time.sleep(5)
            if not registered:
                logger.error("☁️ LINE webhook registration failed after 3 attempts")
            # ── Telegram webhook auto-registration ──
            try:
                tg_token = _load_openclaw_telegram_token()
                tg_secret = _load_telegram_webhook_secret()
                if tg_token:
                    tg_webhook_url = f"{cf_url}/telegram/webhook"
                    tg_data = urllib.parse.urlencode({"url": tg_webhook_url, **({"secret_token": tg_secret} if tg_secret else {})}).encode()
                    tg_req = urllib.request.Request(f"https://api.telegram.org/bot{tg_token}/setWebhook", data=tg_data)
                    with urllib.request.urlopen(tg_req, timeout=10) as tg_resp:
                        logger.info(f"☁️ Telegram webhook registered: {tg_webhook_url} → {tg_resp.status}")
            except Exception as tg_e:
                logger.warning(f"☁️ Telegram webhook registration failed: {tg_e}")
            # Save URLs
            agent_dir = os.path.join(os.path.dirname(__file__), "..", ".agent")
            os.makedirs(agent_dir, exist_ok=True)
            try:
                with open(os.path.join(agent_dir, "line_webhook_url.txt"), "w") as f:
                    f.write(webhook_url + "\n")
                with open(os.path.join(agent_dir, "cloudflare_tunnel_url.txt"), "w") as f:
                    f.write(cf_url + "\n")
            except Exception:
                logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 10424, exc_info=True)
        import threading
        threading.Thread(target=_register, daemon=True, name="cloudflared-register").start()
    except Exception as e:
        logger.warning(f"☁️ cloudflared startup failed: {e}")

try:
    _ensure_cloudflared()
except Exception:
    logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 10433, exc_info=True)

# ─── Cloudflared 定期健康檢查（每 90 秒偵測，死掉自動重啟 + 重新註冊 LINE webhook）───
def _cloudflared_watchdog():
    import time as _time
    _INTERVAL = 90
    _time.sleep(60)  # 啟動後等 60 秒再開始巡檢
    while True:
        try:
            if not _is_cloudflared_alive():
                logger.warning("☁️ cloudflared died — restarting...")
                _ensure_cloudflared()
        except Exception as e:
            logger.warning(f"☁️ cloudflared watchdog error: {e}")
        _time.sleep(_INTERVAL)

import threading
threading.Thread(target=_cloudflared_watchdog, daemon=True, name="cloudflared-watchdog").start()

# ─── NAS SMB 自動掛載守衛 ────────────────────────────────────
try:
    from api.nas_mount_guard import start_nas_mount_guard
    start_nas_mount_guard(interval=120)
except Exception as e:
    logger.warning(f"NAS mount guard 啟動失敗: {e}")

# ─── LAF Gmail 背景監控 ───────────────────────────────────────
def _start_laf_gmail_monitor():
    """背景執行緒：每 300 秒掃描 Gmail 法扶信件。"""
    try:
        _laf_paths = [
            os.path.join(os.path.dirname(__file__), '..', 'casper_ecosystem', 'law_firm_orchestrators'),
            os.path.join(os.path.dirname(__file__), '..', 'skills', 'legal'),
        ]
        for p in _laf_paths:
            if p not in sys.path:
                sys.path.insert(0, p)
        from laf_orchestrator import LAFOrchestrator
        orchestrator = LAFOrchestrator(dry_run=False)
        orchestrator.run_monitor()  # blocking loop (interval=300s)
    except Exception as e:
        logger.warning(f"📧 LAF Gmail Monitor 啟動失敗: {e}")

try:
    _laf_gmail_thread = threading.Thread(
        target=_start_laf_gmail_monitor,
        daemon=True,
        name="laf-gmail-monitor",
    )
    _laf_gmail_thread.start()
    logger.info("📧 LAF Gmail Monitor 背景執行緒已啟動")
except Exception as e:
    logger.warning(f"📧 LAF Gmail Monitor 啟動失敗: {e}")

if __name__ == "__main__":
    import signal as _sig
    from api.thread_pools import shutdown_all as _shutdown_pools

    def _graceful_shutdown(signum=None, frame=None):
        logger.info("MAGI server shutting down (signal=%s)...", signum)
        _shutdown_pools(wait=True)
        sys.exit(0)

    _sig.signal(_sig.SIGTERM, _graceful_shutdown)
    _sig.signal(_sig.SIGINT, _graceful_shutdown)

    # Run on port 5002 to avoid conflict with AirPlay (5000)
    app.run(host='127.0.0.1', port=5002)
