# <MAGI_ROOT>/api/osc_document_generator.py

from datetime import datetime
import os

_MAGI_ROOT = os.path.abspath(os.path.dirname(os.path.dirname(__file__)))

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    pass

def set_font_style(run, font_name='標楷體', size_pt=12, bold=False):
    from docx.shared import Pt
    from docx.oxml.ns import qn
    font = run.font
    font.name = font_name
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    font.size = Pt(size_pt)
    font.bold = bold

def generate_receipt(data, fee_type, config):
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document()
    section = doc.sections[0]
    for prop in ['top_margin', 'bottom_margin', 'left_margin', 'right_margin']:
        setattr(section, prop, Cm(1.5))
        
    company_name = config.get('company_name') or '偵理法律事務所'
    company_address = config.get('company_address_hl') or ''
    
    p = doc.add_paragraph()
    set_font_style(p.add_run(f'{company_name}  律師酬金收據'), size_pt=20, bold=True, font_name='標楷體')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run(f"案由： {data.get('案由/事件', '')}")
    p.add_run(f"\t\t事務所地址：{company_address}")
    doc.add_paragraph()
    
    doc.add_paragraph(f"本所就 {data.get('委任人/當事人', ' ')} 之{fee_type}總計新台幣 {data.get('金額', ' ')} 整。")
    doc.add_paragraph("已收受前項服務款項，特立此款項收據。")
    doc.add_paragraph("\n請查照\n\n\n")
    
    p = doc.add_paragraph()
    p.add_run(f"{company_name}\t\t").bold = True
    p.add_run(data.get('律師姓名') or config.get('default_lawyer') or '').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    now = datetime.now()
    roc_year = now.year - 1911
    date_str = data.get('取代日期') or f"中华民國　{roc_year}　年　{now.month}　月　{now.day}　日"
    if not data.get('取代日期'):
        date_str = f"中華民國　{roc_year}　年　{now.month}　月　{now.day}　日"
        
    p = doc.add_paragraph(date_str)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    for para in doc.paragraphs:
        para.paragraph_format.line_spacing = 1.5
        for run in para.runs:
            if not run.font.size:
                set_font_style(run, size_pt=14, font_name='標楷體')
    return doc

def generate_poa(data, case_type, role, config):
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document()
    section = doc.sections[0]
    for prop in ['top_margin', 'bottom_margin', 'left_margin', 'right_margin']:
        setattr(section, prop, Cm(1.5))

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font_style(p_title.add_run(f"{case_type}委任狀"), bold=True, size_pt=22)
    
    p_case_info = doc.add_paragraph()
    p_case_info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font_style(p_case_info.add_run(f"案號：{data.get('案號', '')}\t\t股別：{data.get('股別', '')}"), size_pt=12)
    
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    table.width = Cm(18)
    table.allow_autofit = False
    table.columns[0].width = Cm(2.5)
    table.columns[1].width = Cm(3.5)
    table.columns[2].width = Cm(12)
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "稱謂"
    hdr_cells[1].text = "姓名或名稱"
    hdr_cells[2].text = "年籍及地址等資料"
    
    label_map = {'告訴代理人': '委任人 (告訴人)', '辯護人': '委任人 (被告)', '代理人': '委任人'}
    p_client_cells = table.rows[1].cells
    p_client_cells[0].text = label_map.get(role, '委任人')
    p_client_cells[1].text = data.get('委任人/當事人', '')
    
    address = data.get('通訊地址', '')
    phone = data.get('聯絡電話', '')
    tax_id = data.get('身分證字號', '')
    
    details_text = []
    if address:
        details_text.append(f"住：{address}")
    if phone:
        details_text.append(f"電話：{phone}")
    if tax_id:
        details_text.append(f"統一編號/身分證：{tax_id}")
        
    p_client_cells[2].text = "\n".join(details_text)
    
    company_name = config.get('company_name') or '偵理法律事務所'
    company_address = config.get('company_address_hl') or ''
    company_phone = config.get('company_phone') or ''
    company_fax = config.get('company_fax') or ''
    default_lawyer = config.get('default_lawyer') or '喬政翔律師'
    
    p_agent_cells = table.rows[2].cells
    p_agent_cells[0].text = "受任人"
    p_agent_cells[1].text = data.get('受任律師', default_lawyer)
    p_agent_cells[2].text = (f"{company_name}\n"
                            f"地址：{company_address}\n"
                            f"電話：{company_phone}\n"
                            f"傳真：{company_fax}")

    if case_type == '民事':
        main_text = "為委任代理人事\n委任人茲委任受任人為訴訟代理人，就本事件□有□無為一切訴訟行為之代理權，並有民事訴訟法第70條第1項但書及同條第2項所列各行為之特別代理權。爰依同法第69條第1項前段規定，提出委任書如上。"
    elif case_type == '刑事' and role == '告訴代理人':
        main_text = f"為委任告訴代理人事\n為 {data.get('案由/事件', '')} 刑事案件，茲依照刑事訴訟法第236-1條第1項委任受任人為告訴代理人，就本案件有為一切訴訟行為之權，並依同條第2項之規定提出委任書如上。"
    elif case_type == '刑事' and role == '辯護人':
        main_text = f"為委任辯護人事\n為 {data.get('案由/事件', '')} 刑事案件，茲委任受任人為辯護人，就本案件有為一切訴訟行為之權，提出委任書如上。"
    elif case_type == '行政':
        main_text = "為委任代理人事\n委任人茲委任受任人為訴訟代理人，就本案件有為一切訴訟行為之權，且□有□無行政訴訟法第51條第1項但書及第2項特別代理權，並依行政訴訟法第50條前段提出委任狀。"
    else:
        main_text = "為委任事\n委任人茲委任受任人，就本案件有為一切訴訟行為之權，提出委任書如上。"
        
    now = datetime.now()
    roc_year = now.year - 1911

    p_footer_cell = table.rows[3].cells[0]
    p_footer_cell.merge(table.rows[3].cells[1]).merge(table.rows[3].cells[2])
    
    p_footer = p_footer_cell.paragraphs[0]
    p_footer.text = f"{main_text}\n\n謹呈   {data.get('法院/檢察署', '')}    公鑒"
    p_footer.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    p_signature = p_footer_cell.add_paragraph()
    p_signature.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_signature.add_run(f"\n\n\n委任人： {data.get('委任人/當事人', '')}\n")
    p_signature.add_run(f"受任人： {data.get('受任律師', default_lawyer)}\n\n")
    
    date_str = data.get('取代日期')
    if not date_str:
        date_str = f"中華民國　{roc_year}　年　{now.month}　月　{now.day}　日"
    p_signature.add_run(date_str)

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = 1
            for para in cell.paragraphs:
                para.paragraph_format.line_spacing = 1.5
                for run in para.runs:
                    set_font_style(run, size_pt=14)
    
    return doc

def generate_engagement_agreement(data, config):
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document()
    section = doc.sections[0]
    for prop in ['top_margin', 'bottom_margin', 'left_margin', 'right_margin']:
        setattr(section, prop, Cm(2))
        
    company_name = config.get('company_name') or '偵理法律事務所'
    default_lawyer = config.get('default_lawyer') or '喬政翔律師'
    company_phone = config.get('company_phone') or ''
    company_email = config.get('company_email') or ''
    company_address = config.get('company_address_hl') or ''
    bank_name = config.get('bank_name') or ''
    bank_account_name = config.get('bank_account_name') or ''
    bank_account_number = config.get('bank_account_number') or ''
        
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font_style(p.add_run(company_name + '\n'), size_pt=20, bold=True)
    set_font_style(p.add_run('委任契約書'), size_pt=20, bold=True)
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    set_font_style(p.add_run(f"委任人： {data.get('委任人/當事人', ' ')} 。"))
    p = doc.add_paragraph()
    set_font_style(p.add_run(f"受任人：{default_lawyer}。"))
    p = doc.add_paragraph()
    set_font_style(p.add_run(f"茲為 {data.get('案由/事件', ' ')} 事件，委任人委任受任人辦理之條件如下："))
    doc.add_paragraph()
    
    def add_article(title, content_list):
        p = doc.add_paragraph()
        set_font_style(p.add_run(title), bold=True, size_pt=12, font_name='標楷體')
        for line in content_list:
            p = doc.add_paragraph()
            set_font_style(p.add_run(line), size_pt=12, font_name='標楷體')
            
    add_article("第一條【受任人工作內容】", [
        f"委任範圍： {data.get('委任範圍', '')} 。", 
        "受任權限：", 
        "□有□無 捨棄、認諾、撤回、和解、提起反訴、上訴或再審之訴及選任代理人之權。（如為刑事案件則無須勾選）。", 
        "□有□無 為強制執行之行為或領取所爭物之權。（如為刑事案件則無須勾選）。"
    ])
    add_article("第二條【委任契約服務費用】", [
        f"委任人同意就本契約書第一條第一項所列之法律服務給付委任費用總計 新台幣（下同） {data.get('委任費用(數字)', ' ')} 元整（已扣除　仟元諮詢費用） 。", 
        "事務費用：不另計收。", 
        "支付方式與時程：於簽署本契約後 七 日內，以匯款方式為之。", 
        "匯款資料：", 
        f"銀行：{bank_name}（花蓮分行）", 
        f"銀行代碼：007", 
        f"戶名：{bank_account_name}", 
        f"帳號：{bank_account_number}"
    ])
    add_article("第三條 【委任人義務】", [
        "委任人保證對受任人所述事實均係真確，絕無虛假。",
        "委任人同意受任人選任及複委任律師共同或協同完成受託事務，並同意選任及複委任律師之異動。如需委任人另出具委任書時，委任人同意配合辦理。",
        "就案件之處理，委任人應尊重受任人之專業判斷，所採行動應事先通知受任人。",
        "委任人理解律師依法不得保證訴訟或非訟程序之結果，訴訟或非訟程序必然有其風險，且法官、檢察官、公務員可能有法律見解之不同及裁量空間，故同樣或類似之事件不必然會做成相類似之判斷；委任人亦理解程序、訴訟期間長短，開會或開庭頻率，視個別辦案狀況不同，是否提出資料、資料內容、開庭方式應由受任人本其專業判斷決定。"
    ])
    add_article("第四條 【受任人義務】", [
        "對委任人交付之證物及相關文件資料，應妥為保管及保密。保管期間為委任關係結束後兩年，但另有約定者依其約定。惟委任人交付之證物及相關資料於案件終結後，如已逾法定保管期間仍未取回，受任人不負保管義務。",
        "受任人應及時將委任事務進行之狀況報告委任人，但若涉及偵查不公開之內容或第三人之隱私，受任人得不向委任人說明。如委任人要求提供受任人所設置之檔案，受任人應提供檔案影本，不得無故拖延或拒絕；其所需費用，由委任人負擔。但依法律規定不得提供予委任人之文件、資料，不在此限。",
        "受任人就受任事件，應負職務上之保密義務，於委任關係終止後亦同，但委任人同意者不在此限。"
    ])
    add_article("第五條 【契約之終止】", [
        "委任人終止本契約者，委任人應依案件辦理之程度給付第二條約定之委任契約服務費用及已發生之事務費用。受任事項達成和解或委任人自行撤回訴訟者，委任人應給付第二條約定之委任契約服務費用全部及已發生之事務費用。",
        "委任人違反本契約第二、三條或委任事項涉及違反其他禁止受任之法令強行規定者，受任人得隨時終止本契約，並得請求委任人依案件辦理之程度給付第二條約定之委任契約服務費用，如受任人受有損害，並得請求委任人賠償。但受任人終止本契約係因違反其他禁止受任之法令強行規定者，除終止事由於受任後發生者外，不得收取委任契約服務費用。"
    ])
    add_article("第六條 【聯繫方式】", [
        "雙方依本契約約定應為之一切通知、給付，均以契約雙方如下所列地址與電子信箱為據：", 
        f"委任人 ： {data.get('委任人/當事人', '')}", 
        f"身分證字號： {data.get('身分證字號', '')}", 
        f"聯絡電話： {data.get('聯絡電話', '')}", 
        f"電子信箱： {data.get('電子信箱', '')}", 
        f"通訊地址： {data.get('通訊地址', '')}",
        "", 
        f"受任人 ：{default_lawyer}", 
        f"聯絡電話：{company_phone}", 
        f"電子信箱：{company_email}", 
        f"通訊地址：{company_address}", 
        "任何一方之聯繫方式如有變更，亦應以書面及電子郵件方式通知對方變更地址或電子郵件。否則他方按上述聯繫方式所為之一切通知、給付仍然有效，受通知方或受領方均不得異議。"
    ])
    add_article("第七條【契約權利變動】", ["本契約或依本契約所載明之權利或義務均不能讓渡或移轉。"])
    add_article("第八條【代刻印章】", ["委任人同意受任人於辦理案件範圍內代刻印章，交由受任人保管，並於委任關係結束後銷毀。", "（同意簽章：               ）"])
    add_article("第九條【法律和管轄】", ["本契約之適用與解釋應依照中華民國法律辦理，並受中華民國法律主管機關之管轄，本契約未約定者，適用民法債編及其他法律相關規定。若雙方因本契約發生爭執、違約或其他相關問題時，任一方得請求受任人所加入之律師公會調處；如進入訴訟，雙方合意以臺灣花蓮地方法院為第一審訴訟管轄法院。"])
    add_article("第十條【本契約之修訂】", ["本契約除非由雙方合法授權之代表於修訂文件上簽字，不得變更修訂，修訂程序和制定本契約相同。"])
    add_article("第十一條【本契約之確認】", ["茲證明雙方簽署本契約，係經由雙方合法授權之代表於首揭日期簽署，本契約乙式貳份，由雙方各執正本乙份為憑。"])
    
    p = doc.add_paragraph()
    set_font_style(p.add_run("\n雙方約定如上，並立書面，以資信守。"))
    p = doc.add_paragraph()
    set_font_style(p.add_run("確認及同意：\n\n"))
    
    p = doc.add_paragraph()
    p.add_run(f"委任人：\t\t\t\t\t\t受任人：{default_lawyer}")
    doc.add_paragraph(f"\n\n簽署：")
    
    date_str = data.get('取代日期')
    if not date_str:
        now = datetime.now()
        roc_year = now.year - 1911
        date_str = f"中華民國 {roc_year} 年 {now.month} 月 {now.day} 日"
    p = doc.add_paragraph(f"\n{date_str}")
    
    for para in doc.paragraphs:
        for run in para.runs:
            if not run.font.name or run.font.name not in ['標楷體']:
                set_font_style(run, size_pt=run.font.size.pt if run.font.size else 12, font_name='標楷體')
    return doc
