"""Hearing-conflict decisions and leave-request document generation.

The module is deliberately independent from Flask and the database.  Route and
ingestion layers supply normalized calendar rows and case metadata; this module
decides whether an older appointment is another hearing (generate a draft) or a
different business appointment (notify only), then renders a reviewable DOCX.

No document is ever filed with a court automatically.
"""

from __future__ import annotations

import re
import hashlib
import json
from dataclasses import dataclass
from datetime import date, datetime, time, timedelta
from pathlib import Path
from typing import Any, Iterable, Mapping
from zoneinfo import ZoneInfo


_HEARING_MARKERS = (
    "開庭",
    "庭期",
    "言詞辯論",
    "準備程序",
    "審理程序",
    "調解",
    "訊問",
    "協商程序",
    "調查程序",
    "審判程序",
)
_NON_HEARING_MARKERS: tuple[tuple[str, tuple[str, ...]], ...] = (
    ("legal_aid", ("法扶", "法律扶助", "扶助分會", "接案")),
    ("client_visit", ("律見", "接見", "會客", "訪視")),
    ("file_review", ("閱卷", "調卷", "卷證")),
    ("meeting", ("會議", "開會", "訪談", "諮詢")),
)
_NON_HEARING_CONTEXTS = (
    "開庭前會議",
    "庭前會議",
    "庭期協調",
    "開庭通知寄送",
    "開庭通知寄發",
    "開庭通知轉寄",
    "開庭準備會議",
    "調解前會議",
)
_NOT_A_HEARING = ("待確認", "未確認", "宣判", "宣示判決", "判決公告")
_TENTATIVE_MARKERS = (
    "待確認",
    "未確認",
    "tentative",
    "needsaction",
    "needs action",
    "已取消",
    "cancelled",
    "canceled",
)
_CANCELLED = {
    "cancelled",
    "canceled",
    "已取消",
    "取消",
    "deleted",
    "completed",
    "done",
    "已完成",
    "完成",
}
_SAFE_FILENAME_RE = re.compile(r"[\\/:*?\"<>|\x00-\x1f]+")


@dataclass(frozen=True, slots=True)
class NormalizedSchedule:
    event_id: str
    case_number: str
    title: str
    start: datetime
    end: datetime
    created_at: datetime | None
    status: str
    source: str
    location: str
    kind: str
    all_day: bool
    lawyer: str
    excluded: bool
    raw: Mapping[str, Any]


@dataclass(frozen=True, slots=True)
class ConflictDecision:
    candidate: NormalizedSchedule
    existing: NormalizedSchedule
    action: str
    reason: str

    def as_dict(self) -> dict[str, Any]:
        return {
            "action": self.action,
            "reason": self.reason,
            "candidate": schedule_to_dict(self.candidate),
            "existing": schedule_to_dict(self.existing),
        }


def _text(value: Any) -> str:
    return str(value or "").strip()


def _truthy(value: Any) -> bool:
    return _text(value).lower() in {"1", "true", "yes", "on"}


def _parse_datetime(value: Any, *, default_time: time = time.min) -> datetime | None:
    if isinstance(value, datetime):
        return value.astimezone(ZoneInfo("Asia/Taipei")).replace(tzinfo=None) if value.tzinfo else value
    if isinstance(value, date):
        return datetime.combine(value, default_time)
    text = _text(value)
    if not text:
        return None
    normalized = text.replace("Z", "+00:00")
    try:
        parsed = datetime.fromisoformat(normalized)
        return parsed.astimezone(ZoneInfo("Asia/Taipei")).replace(tzinfo=None) if parsed.tzinfo else parsed
    except ValueError:
        pass
    for fmt in ("%Y-%m-%d %H:%M", "%Y-%m-%d %H:%M:%S", "%Y/%m/%d %H:%M", "%Y-%m-%d"):
        try:
            parsed = datetime.strptime(text, fmt)
            if fmt == "%Y-%m-%d":
                return datetime.combine(parsed.date(), default_time)
            return parsed
        except ValueError:
            continue
    return None


def _schedule_text(value: Mapping[str, Any] | str) -> str:
    if isinstance(value, Mapping):
        return " ".join(
            _text(value.get(key))
            for key in (
                "title",
                "summary",
                "description",
                "todo_type",
                "location",
                "status",
                "todo_status",
            )
        )
    return _text(value)


def is_excluded_schedule(value: Mapping[str, Any] | str) -> bool:
    text = _schedule_text(value)
    lowered = text.lower()
    if any(marker in lowered for marker in _TENTATIVE_MARKERS):
        return True
    if isinstance(value, Mapping):
        status = _text(value.get("status") or value.get("todo_status")).lower()
        if status in _CANCELLED:
            return True
        all_day = value.get("is_all_day") is True or _truthy(value.get("is_all_day"))
        if all_day and any(marker in text for marker in _HEARING_MARKERS):
            return True
    return False


def classify_schedule(value: Mapping[str, Any] | str) -> str:
    text = _schedule_text(value)
    if is_excluded_schedule(value):
        return "other"
    if any(marker in text for marker in _NOT_A_HEARING):
        return "other"
    if any(marker in text for marker in _NON_HEARING_CONTEXTS):
        for kind, markers in _NON_HEARING_MARKERS:
            if any(marker in text for marker in markers):
                return kind
        return "meeting"
    if any(marker in text for marker in _HEARING_MARKERS):
        return "hearing"
    for kind, markers in _NON_HEARING_MARKERS:
        if any(marker in text for marker in markers):
            return kind
    return "other"


def normalize_schedule(value: Mapping[str, Any], *, default_duration_minutes: int = 60) -> NormalizedSchedule:
    title = _text(value.get("title") or value.get("todo_type") or value.get("summary") or value.get("description"))
    start = _parse_datetime(value.get("start_date") or value.get("start") or _todo_datetime(value))
    if start is None:
        raise ValueError("schedule_start_required")
    all_day = bool(value.get("is_all_day")) if isinstance(value.get("is_all_day"), bool) else _truthy(value.get("is_all_day"))
    end = _parse_datetime(value.get("end_date") or value.get("end"))
    if end is None or end <= start:
        end = start + (timedelta(days=1) if all_day else timedelta(minutes=max(1, default_duration_minutes)))
    created_at = _parse_datetime(value.get("created_date") or value.get("created_at") or value.get("indexed_date"))
    event_id = _text(value.get("event_id") or value.get("todo_id") or value.get("id"))
    source = _text(value.get("source_kind") or value.get("source_table") or value.get("source_file") or value.get("source"))
    excluded = is_excluded_schedule(value)
    return NormalizedSchedule(
        event_id=event_id,
        case_number=_text(value.get("case_number")),
        title=title or "未命名行程",
        start=start,
        end=end,
        created_at=created_at,
        status=_text(value.get("status") or value.get("todo_status")),
        source=source,
        location=_text(value.get("location")),
        kind=classify_schedule(value),
        all_day=all_day,
        lawyer=_text(value.get("lawyer") or value.get("attorney") or value.get("responsible_lawyer")),
        excluded=excluded,
        raw=dict(value),
    )


def _todo_datetime(value: Mapping[str, Any]) -> str:
    day = _text(value.get("todo_date"))
    clock = _text(value.get("todo_time"))
    return f"{day} {clock}".strip()


def schedule_to_dict(value: NormalizedSchedule) -> dict[str, Any]:
    return {
        "event_id": value.event_id,
        "case_number": value.case_number,
        "title": value.title,
        "start": value.start.isoformat(sep=" "),
        "end": value.end.isoformat(sep=" "),
        "created_at": value.created_at.isoformat(sep=" ") if value.created_at else "",
        "status": value.status,
        "source": value.source,
        "location": value.location,
        "kind": value.kind,
        "is_all_day": value.all_day,
        "lawyer": value.lawyer,
        "excluded": value.excluded,
        "hearing_type": _text(value.raw.get("hearing_type")),
    }


def _same_event(left: NormalizedSchedule, right: NormalizedSchedule) -> bool:
    if left.event_id and right.event_id and left.event_id == right.event_id and left.source == right.source:
        return True
    left_file = _text(left.raw.get("source_file"))
    right_file = _text(right.raw.get("source_file"))
    if (
        left_file
        and right_file
        and left_file == right_file
        and left.case_number
        and left.case_number == right.case_number
        and left.start == right.start
    ):
        # PDF ingestion can materialize the same notice as both a case_todo and
        # a calendar candidate with different display labels.  It must not
        # conflict with itself merely because the two labels differ.
        return True
    return bool(
        left.case_number
        and left.case_number == right.case_number
        and left.start == right.start
        and re.sub(r"\s+", "", left.title) == re.sub(r"\s+", "", right.title)
    )


def _is_earlier(existing: NormalizedSchedule, candidate: NormalizedSchedule) -> bool:
    if existing.created_at and candidate.created_at:
        return existing.created_at < candidate.created_at
    # A manually previewed candidate has not been persisted yet; all persisted
    # schedules supplied by the caller are therefore older.
    if candidate.created_at is None:
        return True
    # Missing legacy timestamps cannot prove priority and must not cause an
    # automatic legal document.  The UI may still show them for review.
    return False


def _lawyer_tokens(value: str) -> set[str]:
    return {
        token.replace("律師", "").strip()
        for token in re.split(r"[\s、,，/;；]+", _text(value))
        if token.replace("律師", "").strip()
    }


def _same_lawyer(left: str, right: str) -> bool | None:
    left_tokens = _lawyer_tokens(left)
    right_tokens = _lawyer_tokens(right)
    if not left_tokens or not right_tokens:
        return None
    return bool(left_tokens & right_tokens)


def find_conflict_decisions(
    candidate_value: Mapping[str, Any],
    existing_values: Iterable[Mapping[str, Any]],
) -> list[ConflictDecision]:
    candidate = normalize_schedule(candidate_value)
    if candidate.kind != "hearing":
        return []
    decisions: list[ConflictDecision] = []
    for raw in existing_values:
        try:
            existing = normalize_schedule(raw)
        except ValueError:
            continue
        decision = _conflict_decision(candidate, existing)
        if decision:
            decisions.append(decision)
    decisions.sort(key=lambda item: (item.existing.start, item.existing.created_at or datetime.min, item.existing.title))
    return decisions


def _conflict_decision(
    candidate: NormalizedSchedule,
    existing: NormalizedSchedule,
) -> ConflictDecision | None:
    if candidate.kind != "hearing":
        return None
    if candidate.excluded or existing.excluded:
        return None
    if existing.status.lower() in _CANCELLED or _same_event(candidate, existing):
        return None
    if not _is_earlier(existing, candidate):
        return None
    if not (candidate.start < existing.end and existing.start < candidate.end):
        return None
    if existing.kind == "hearing" and not existing.all_day:
        same_lawyer = _same_lawyer(candidate.lawyer, existing.lawyer)
        if same_lawyer is False:
            return None
        if same_lawyer is None:
            return ConflictDecision(
                candidate,
                existing,
                "notify_only",
                "開庭時間重疊，但承辦律師資料不完整，未自動產生請假狀",
            )
        return ConflictDecision(candidate, existing, "generate_leave_request", "同一承辦律師較早排定的開庭與新庭期重疊")
    return ConflictDecision(candidate, existing, "notify_only", "較早排定的非開庭行程與新庭期重疊")


def find_all_conflict_decisions(
    values: Iterable[Mapping[str, Any]],
    *,
    max_decisions: int | None = None,
    candidate_case_number: str = "",
) -> tuple[list[NormalizedSchedule], list[ConflictDecision]]:
    """Normalize once and scan all confirmed hearings in one lightweight pass.

    The OSC workbench can inspect up to 1,000 schedules.  Re-parsing that full
    list for every hearing caused avoidable CPU and allocation spikes, so this
    bulk path performs datetime/classification normalization exactly once.
    """

    schedules: list[NormalizedSchedule] = []
    for raw in values:
        try:
            schedules.append(normalize_schedule(raw))
        except ValueError:
            continue
    by_start = sorted(schedules, key=lambda item: (item.start, item.end, item.event_id))
    decisions: list[ConflictDecision] = []
    for candidate in schedules:
        if candidate.kind != "hearing":
            continue
        if candidate_case_number and candidate.case_number != _text(candidate_case_number):
            continue
        for existing in by_start:
            if existing.start >= candidate.end:
                break
            if existing.end <= candidate.start:
                continue
            decision = _conflict_decision(candidate, existing)
            if decision:
                decisions.append(decision)
                if max_decisions is not None and len(decisions) >= max(1, int(max_decisions)):
                    decisions.sort(
                        key=lambda item: (
                            item.candidate.start,
                            item.existing.start,
                            item.existing.created_at or datetime.min,
                            item.existing.title,
                        )
                    )
                    return schedules, decisions
    decisions.sort(
        key=lambda item: (
            item.candidate.start,
            item.existing.start,
            item.existing.created_at or datetime.min,
            item.existing.title,
        )
    )
    return schedules, decisions


def scan_all_conflict_decisions(
    values: Iterable[Mapping[str, Any]],
    *,
    max_decisions: int = 250,
    candidate_case_number: str = "",
) -> tuple[list[NormalizedSchedule], list[ConflictDecision], bool]:
    """Bound a workbench response before pair explosion can exhaust memory."""

    cap = max(1, int(max_decisions))
    schedules, decisions = find_all_conflict_decisions(
        values,
        max_decisions=cap + 1,
        candidate_case_number=candidate_case_number,
    )
    truncated = len(decisions) > cap
    return schedules, decisions[:cap], truncated


def pick_leave_request_conflict(decisions: Iterable[ConflictDecision]) -> ConflictDecision | None:
    hearings = [item for item in decisions if item.action == "generate_leave_request"]
    return min(hearings, key=lambda item: (item.existing.created_at or datetime.max, item.existing.start), default=None)


def is_legal_aid_case(case: Mapping[str, Any]) -> bool:
    text = " ".join(
        _text(case.get(key))
        for key in (
            "case_category",
            "case_type",
            "case_reason",
            "laf_case_no",
            "legal_aid_number",
            "application_no",
            "legal_aid_status",
        )
    )
    return "法扶" in text or "法律扶助" in text or bool(re.search(r"\d{6,8}-[A-Z]-\d{3}", text))


def public_document_title(case: Mapping[str, Any]) -> str:
    text = " ".join(_text(case.get(key)) for key in ("case_category", "case_type", "case_reason"))
    if any(marker in text for marker in ("消費者債務清理", "消債", "更生", "清算")):
        return "消費者債務清理事件聲請改期狀"
    if "行政" in text:
        return "行政聲請改期狀"
    if "刑事" in text:
        return "刑事聲請改期狀"
    return "民事聲請變更期日狀"


def roc_date(value: date | datetime) -> str:
    day = value.date() if isinstance(value, datetime) else value
    return f"中華民國{day.year - 1911}年{day.month}月{day.day}日"


def roc_datetime(value: datetime) -> str:
    noon = "上午" if value.hour < 12 else "下午"
    hour = value.hour if 1 <= value.hour <= 12 else (12 if value.hour in {0, 12} else value.hour - 12)
    minute = f"{value.minute}分" if value.minute else ""
    return f"民國{value.year - 1911}年{value.month}月{value.day}日{noon}{hour}時{minute}"


def leave_request_payload(
    case: Mapping[str, Any],
    decision: ConflictDecision,
    *,
    lawyer_name: str = "",
    party_role: str = "當事人",
    generated_on: date | None = None,
) -> dict[str, Any]:
    existing_case = decision.existing.raw
    aid = is_legal_aid_case(case)
    target_hearing_label = _text(decision.candidate.raw.get("hearing_type")) or _hearing_label(decision.candidate.title)
    prior_hearing_label = _hearing_label(decision.existing.title)
    return {
        "document_title": public_document_title(case),
        "court_case_no": _text(case.get("court_case_no") or case.get("court_case_number")),
        "division": _text(case.get("court_division") or case.get("division")),
        "court_name": _text(case.get("court_name") or decision.candidate.location),
        "party_name": _text(case.get("client_name") or case.get("party_name")),
        "party_role": _text(party_role) or "當事人",
        "lawyer_name": _text(lawyer_name or case.get("lawyer")),
        "is_legal_aid": aid,
        "lawyer_capacity": "扶助律師" if aid else "委任律師",
        "conflict_statement": "兩案庭期重疊。",
        "case_reason": _text(case.get("case_reason")) or "本案",
        "target_start": decision.candidate.start,
        "target_hearing_label": target_hearing_label,
        "prior_start": decision.existing.start,
        "prior_hearing_label": prior_hearing_label,
        "prior_court_name": _text(existing_case.get("court_name") or existing_case.get("location") or decision.existing.location),
        "prior_court_case_no": _text(existing_case.get("court_case_no") or existing_case.get("court_case_number")),
        "prior_case_number": decision.existing.case_number,
        "generated_on": generated_on or date.today(),
        "office_address": _text(case.get("office_address")),
        "office_phone": _text(case.get("office_phone")),
        "generation_mode": "automatic_conflict",
    }


def manual_leave_request_payload(
    case: Mapping[str, Any],
    *,
    target_start: datetime,
    prior_start: datetime,
    prior_court_name: str,
    prior_hearing_label: str = "開庭",
    prior_court_case_no: str = "",
    lawyer_name: str = "",
    party_role: str = "當事人",
    target_hearing_label: str = "開庭",
    conflict_statement: str = "",
    generated_on: date | None = None,
) -> dict[str, Any]:
    """Build a human-requested draft without requiring an automatic conflict.

    This is intentionally separate from ``find_conflict_decisions``: a person
    may generate a draft for an event excluded by automatic classification, but
    the output remains a reviewable draft and is never filed automatically.
    """

    aid = is_legal_aid_case(case)
    return {
        "document_title": public_document_title(case),
        "court_case_no": _text(case.get("court_case_no") or case.get("court_case_number")),
        "division": _text(case.get("court_division") or case.get("division")),
        "court_name": _text(case.get("court_name")),
        "party_name": _text(case.get("client_name") or case.get("party_name")),
        "party_role": _text(party_role) or "當事人",
        "lawyer_name": _text(lawyer_name or case.get("lawyer")),
        "is_legal_aid": aid,
        "lawyer_capacity": "扶助律師" if aid else "委任律師",
        "conflict_statement": _text(conflict_statement) or "前述庭期時間相近或重疊，代理人無法兼顧。",
        "case_reason": _text(case.get("case_reason")) or "本案",
        "target_start": target_start,
        "target_hearing_label": _text(target_hearing_label) or "開庭",
        "prior_start": prior_start,
        "prior_hearing_label": _text(prior_hearing_label) or "開庭",
        "prior_court_name": _text(prior_court_name),
        "prior_court_case_no": _text(prior_court_case_no),
        "prior_case_number": "",
        "generated_on": generated_on or date.today(),
        "office_address": _text(case.get("office_address")),
        "office_phone": _text(case.get("office_phone")),
        "generation_mode": "manual",
    }


def _hearing_label(text: str) -> str:
    for marker in (
        "言詞辯論",
        "準備程序",
        "審理程序",
        "調解",
        "訊問",
        "協商程序",
        "調查程序",
        "審判程序",
        "開庭",
    ):
        if marker in _text(text):
            return marker
    return "開庭"


def _safe_filename(value: str) -> str:
    return _SAFE_FILENAME_RE.sub("_", _text(value)).strip(" ._") or "未命名"


def leave_request_filename(payload: Mapping[str, Any]) -> str:
    day = payload.get("generated_on")
    if isinstance(day, datetime):
        day = day.date()
    if not isinstance(day, date):
        day = date.today()
    title = _safe_filename(_text(payload.get("document_title")) or "聲請改期狀")
    party = _safe_filename(_text(payload.get("party_name")) or "當事人")
    target = payload.get("target_start")
    prior = payload.get("prior_start")
    mode = "人工" if _text(payload.get("generation_mode")) == "manual" else "自動"
    fingerprint_fields = {
        key: (
            value.isoformat(sep=" ")
            if isinstance(value, datetime)
            else value.isoformat()
            if isinstance(value, date)
            else _text(value)
        )
        for key, value in payload.items()
        if key
        in {
            "document_title",
            "court_case_no",
            "division",
            "court_name",
            "party_name",
            "party_role",
            "lawyer_name",
            "lawyer_capacity",
            "case_reason",
            "target_start",
            "target_hearing_label",
            "prior_start",
            "prior_hearing_label",
            "prior_court_name",
            "prior_court_case_no",
            "conflict_statement",
            "generation_mode",
        }
    }
    fingerprint = hashlib.sha256(
        json.dumps(fingerprint_fields, ensure_ascii=False, sort_keys=True).encode("utf-8")
    ).hexdigest()[:8]
    target_suffix = f"_庭期{target:%Y%m%d%H%M}" if isinstance(target, datetime) else ""
    prior_suffix = f"_衝庭{prior:%Y%m%d%H%M}" if isinstance(prior, datetime) else ""
    return f"{day:%Y%m%d} {title}({party})_{mode}{target_suffix}{prior_suffix}_{fingerprint}v1.docx"


def resolve_pleading_output_dir(
    case_folder: str | Path,
    *,
    title: str,
    day: date | None = None,
    case_category: str = "一般案件",
) -> Path:
    root = Path(case_folder).expanduser()
    if not root.is_dir():
        raise FileNotFoundError(f"case_folder_not_found: {root}")
    from api.osc.case_folder_schema import case_subfolders

    category = _text(case_category)
    if category in {"法扶", "法扶案件", "法律扶助"}:
        category = "法律扶助案件"
    expected = next(
        (name for name in case_subfolders(category) if "我方歷次書狀" in name),
        "02_我方歷次書狀",
    )
    candidates = sorted(
        (child for child in root.iterdir() if child.is_dir() and "我方歷次書狀" in child.name),
        key=lambda child: (0 if child.name == expected else 1, child.name),
    )
    pleading_root = candidates[0] if candidates else root / expected
    pleading_root.mkdir(parents=True, exist_ok=True)
    folder = pleading_root / f"{(day or date.today()):%Y%m%d} {_safe_filename(title)}"
    folder.mkdir(parents=True, exist_ok=True)
    return folder


def build_leave_request_docx(payload: Mapping[str, Any], output_path: str | Path) -> Path:
    """Render the reviewed NAS public format into an editable A4 DOCX."""

    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt

    path = Path(output_path).expanduser()
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    table_widths = (4.2, 0.7, 11.9)

    def set_run(run, *, size: float = 13.5, bold: bool = False) -> None:
        # Arial Unicode MS is bundled on the target Mac and renders Traditional
        # Chinese reliably in both Microsoft Word and LibreOffice.  BiauKaiTC
        # produced valid text extraction but blank glyphs in the headless PDF
        # renderer used by the release gate.
        run.font.name = "Arial Unicode MS"
        fonts = run._element.get_or_add_rPr().rFonts
        fonts.set(qn("w:ascii"), "Arial Unicode MS")
        fonts.set(qn("w:hAnsi"), "Arial Unicode MS")
        fonts.set(qn("w:eastAsia"), "Arial Unicode MS")
        fonts.set(qn("w:cs"), "Arial Unicode MS")
        run.font.size = Pt(size)
        run.bold = bold

    def add_text(paragraph, text: str, *, size: float = 14, bold: bool = False):
        run = paragraph.add_run(_text(text))
        set_run(run, size=size, bold=bold)
        return run

    def set_cell_width(cell, cm: float) -> None:
        width = Cm(cm)
        cell.width = width
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_w = tc_pr.find(qn("w:tcW"))
        if tc_w is None:
            tc_w = OxmlElement("w:tcW")
            tc_pr.append(tc_w)
        tc_w.set(qn("w:w"), str(int(width.twips)))
        tc_w.set(qn("w:type"), "dxa")

    def keep_row_together(row) -> None:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(10)
    add_text(title_p, _text(payload.get("document_title")) or "聲請變更期日狀", size=20, bold=True)

    info_rows = [
        ("案號", _text(payload.get("court_case_no")) or "【請填法院案號】"),
        ("股別", _text(payload.get("division")) or "【請填股別】"),
        (_text(payload.get("party_role")) or "當事人", _text(payload.get("party_name")) or "【請填當事人】"),
        ("代理人", _text(payload.get("lawyer_name")) or "【請填律師姓名】"),
    ]
    table = doc.add_table(rows=len(info_rows), cols=3)
    table.style = "Table Grid"
    table.autofit = False
    for column, width in zip(table.columns, table_widths):
        column.width = Cm(width)
    for row, (label, value) in zip(table.rows, info_rows):
        keep_row_together(row)
        for cell, width in zip(row.cells, table_widths):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row.cells[0].text = ""
        row.cells[1].text = ""
        row.cells[2].text = ""
        add_text(row.cells[0].paragraphs[0], label)
        add_text(row.cells[1].paragraphs[0], "：")
        add_text(row.cells[2].paragraphs[0], value)
        if label == "代理人" and _text(payload.get("lawyer_capacity")):
            p = row.cells[0].add_paragraph()
            add_text(p, f"（{_text(payload.get('lawyer_capacity'))}）", size=12)
        if label == "代理人":
            address = _text(payload.get("office_address"))
            phone = _text(payload.get("office_phone"))
            if address:
                add_text(row.cells[2].add_paragraph(), f"設：{address}", size=12)
            if phone:
                add_text(row.cells[2].add_paragraph(), f"電話：{phone}", size=12)

    lead = doc.add_paragraph()
    lead.paragraph_format.space_before = Pt(8)
    lead.paragraph_format.line_spacing = 1.5
    add_text(
        lead,
        f"為{_text(payload.get('case_reason')) or '本案'}事件，謹懇請改定{_text(payload.get('target_hearing_label')) or '庭期'}事：",
        bold=True,
    )

    target_start = payload.get("target_start")
    prior_start = payload.get("prior_start")
    if not isinstance(target_start, datetime) or not isinstance(prior_start, datetime):
        raise ValueError("target_start_and_prior_start_required")
    prior_court = _text(payload.get("prior_court_name")) or "另案法院"
    prior_no = _text(payload.get("prior_court_case_no"))
    prior_case = f"（{prior_no}）" if prior_no else ""
    target_when = _text(payload.get("target_display")) or roc_datetime(target_start)
    prior_when = _text(payload.get("prior_display")) or roc_datetime(prior_start)
    body = (
        f"頃奉　鈞院通知，本案將於{target_when}進行"
        f"{_text(payload.get('target_hearing_label')) or '開庭'}，惟代理人先前已受理並排定於"
        f"{prior_when}至{prior_court}{prior_case}進行"
        f"{_text(payload.get('prior_hearing_label')) or '開庭'}，"
        f"{_text(payload.get('conflict_statement')) or '兩案庭期重疊。'}此有該案開庭通知書可參【附件一】。"
        "為保障當事人權益，實有改定期日之必要，爰懇請　鈞院准予請假並另定期日。如蒙恩准，實無任感禱。"
    )
    body_p = doc.add_paragraph()
    body_p.paragraph_format.first_line_indent = Cm(0.85)
    body_p.paragraph_format.line_spacing = 1.5
    add_text(body_p, body)

    closing = doc.add_paragraph()
    closing.paragraph_format.space_before = Pt(8)
    add_text(closing, "謹　狀")
    court_p = doc.add_paragraph()
    add_text(court_p, f"{_text(payload.get('court_name')) or '【請填法院】'}　公鑒", bold=True)

    attach_p = doc.add_paragraph()
    attach_p.paragraph_format.space_before = Pt(8)
    add_text(attach_p, "【附件名稱及件數】（影本）", bold=True)
    attachment = doc.add_table(rows=1, cols=3)
    attachment.style = "Table Grid"
    attachment.autofit = False
    for column, width in zip(attachment.columns, table_widths):
        column.width = Cm(width)
    keep_row_together(attachment.rows[0])
    for cell, width in zip(attachment.rows[0].cells, table_widths):
        set_cell_width(cell, width)
        cell.text = ""
    add_text(attachment.cell(0, 0).paragraphs[0], "附件一")
    add_text(attachment.cell(0, 1).paragraphs[0], "：")
    add_text(attachment.cell(0, 2).paragraphs[0], f"{prior_court}開庭通知書乙份。")

    generated_on = payload.get("generated_on")
    if isinstance(generated_on, datetime):
        generated_on = generated_on.date()
    if not isinstance(generated_on, date):
        generated_on = date.today()
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_p.paragraph_format.space_before = Pt(12)
    add_text(date_p, roc_date(generated_on))

    sign = doc.add_table(rows=2, cols=3)
    sign.style = "Table Grid"
    sign.autofit = False
    for column, width in zip(sign.columns, table_widths):
        column.width = Cm(width)
    for row in sign.rows:
        keep_row_together(row)
        for cell, width in zip(row.cells, table_widths):
            set_cell_width(cell, width)
            cell.text = ""
    add_text(sign.cell(0, 0).paragraphs[0], _text(payload.get("party_role")) or "當事人")
    add_text(sign.cell(0, 1).paragraphs[0], "：")
    add_text(sign.cell(0, 2).paragraphs[0], _text(payload.get("party_name")) or "【請填當事人】")
    capacity = _text(payload.get("lawyer_capacity"))
    lawyer_label = f"訴訟代理人（{capacity}）" if capacity else "訴訟代理人"
    add_text(sign.cell(1, 0).paragraphs[0], lawyer_label)
    add_text(sign.cell(1, 1).paragraphs[0], "：")
    add_text(sign.cell(1, 2).paragraphs[0], _text(payload.get("lawyer_name")) or "【請填律師姓名】")

    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.line_spacing = paragraph.paragraph_format.line_spacing or 1.5
        for run in paragraph.runs:
            if run.font.size is None:
                set_run(run)
    doc.core_properties.title = _text(payload.get("document_title")) or "聲請變更期日狀"
    doc.core_properties.subject = "MAGI 衝庭請假狀公版；送出前須由律師人工確認"
    doc.core_properties.author = "MAGI"
    settings = doc.settings._element
    zoom = settings.find(qn("w:zoom"))
    if zoom is None:
        zoom = OxmlElement("w:zoom")
        settings.insert(0, zoom)
    zoom.set(qn("w:percent"), "100")
    doc.save(path)
    return path


def conflict_notification_text(decision: ConflictDecision, *, generated_path: str = "") -> str:
    candidate = decision.candidate
    existing = decision.existing
    lines = [
        "⚠️ MAGI 衝庭提醒",
        f"新庭期：{candidate.start:%Y-%m-%d %H:%M}｜{candidate.title}｜{candidate.case_number or '未標案件'}",
        f"既有行程：{existing.start:%Y-%m-%d %H:%M}｜{existing.title}｜{existing.case_number or '非案件行程'}",
    ]
    if decision.action == "generate_leave_request":
        lines.append("處理：較早行程為開庭，已產生聲請改期／請假狀草稿，送出前請人工確認。")
        if generated_path:
            lines.append(f"草稿：{Path(generated_path).name}")
    else:
        if existing.kind == "hearing":
            lines.append("處理：開庭時間重疊，但承辦律師資料不足；未自動產生請假狀，請人工確認。")
        else:
            lines.append("處理：較早行程不是開庭，不自動產生請假狀；請人工決定如何調整並聯繫相關單位。")
    return "\n".join(lines)
