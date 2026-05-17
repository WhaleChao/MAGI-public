"""
MAGI startup helpers: file I/O, captcha brokers, URL utilities,
export functions, and background initialization routines.

Extracted from server.py to reduce monolith size.
"""
from __future__ import annotations

import html as ihtml
import json
import logging
import os
import re
import shutil
import subprocess
import sys
import threading
import time
import uuid
from urllib.parse import urlparse
import urllib.request

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Once-guard: run_startup_hooks() must execute at most once per process
# ---------------------------------------------------------------------------
_STARTUP_HOOKS_DONE = False
_STARTUP_HOOKS_LOCK = threading.Lock()
_SHARE_TUNNEL_LOCK = threading.Lock()

# ---------------------------------------------------------------------------
# Directory / path constants
# ---------------------------------------------------------------------------
AGENT_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".agent"))
os.makedirs(AGENT_DIR, exist_ok=True)

# ---------------------------------------------------------------------------
# LAF captcha constants
# ---------------------------------------------------------------------------
LAF_CAPTCHA_REQUEST_FILE = os.environ.get(
    "MAGI_LAF_CAPTCHA_REQUEST_FILE",
    os.path.join(AGENT_DIR, "laf_captcha_request.json"),
)
LAF_CAPTCHA_RESPONSE_FILE = os.environ.get(
    "MAGI_LAF_CAPTCHA_RESPONSE_FILE",
    os.path.join(AGENT_DIR, "laf_captcha_response.json"),
)
LAF_CAPTCHA_TTL_SECONDS = int(os.environ.get("MAGI_LAF_CAPTCHA_TTL_SECONDS", "300") or "300")

# Generic captcha constants
GEN_CAPTCHA_REQUEST_FILE = os.environ.get(
    "MAGI_CAPTCHA_REQUEST_FILE",
    os.path.join(AGENT_DIR, "captcha_request.json"),
)
GEN_CAPTCHA_RESPONSE_FILE = os.environ.get(
    "MAGI_CAPTCHA_RESPONSE_FILE",
    os.path.join(AGENT_DIR, "captcha_response.json"),
)

# LINE last-sender / callback / base-url persistence
LINE_LAST_SENDER_FILE = os.environ.get(
    "MAGI_LINE_LAST_SENDER_FILE",
    os.path.join(AGENT_DIR, "line_last_sender.json"),
)
LINE_LAST_CALLBACK_FILE = os.environ.get(
    "MAGI_LINE_LAST_CALLBACK_FILE",
    os.path.join(AGENT_DIR, "line_last_callback.json"),
)
LINE_AUTO_ADMIN_LAST_SENDER = os.environ.get(
    "MAGI_LINE_AUTO_ADMIN_LAST_SENDER", "0"
).strip().lower() in {"1", "true", "yes", "on"}

LINE_LAST_BASE_URL_FILE = os.environ.get(
    "MAGI_LINE_LAST_BASE_URL_FILE",
    os.path.join(AGENT_DIR, "line_last_base_url.json"),
)


def _load_dotenv_value(key: str, default: str = "") -> str:
    env_value = os.environ.get(key)
    if env_value is not None:
        return env_value
    env_path = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".env"))
    try:
        with open(env_path, encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line or line.startswith("#") or "=" not in line:
                    continue
                k, v = line.split("=", 1)
                if k.strip() == key:
                    return v.strip().strip('"').strip("'")
    except Exception:
        pass
    return default

# Export directory
EXPORTS_DIR = os.environ.get(
    "MAGI_EXPORTS_DIR",
    os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "static", "exports")),
)
EXPORT_LONG_TEXT = os.environ.get("MAGI_EXPORT_LONG_TEXT", "1").strip().lower() in {"1", "true", "yes", "on"}
EXPORT_TEXT_THRESHOLD = int(os.environ.get("MAGI_EXPORT_TEXT_THRESHOLD", "9000"))


# ============================================================================
# 1. File operation helpers
# ============================================================================

def _load_json(path: str) -> dict:
    try:
        if path and os.path.exists(path):
            with open(path, "r", encoding="utf-8") as f:
                return json.load(f) or {}
    except Exception:
        return {}
    return {}


def _write_json_atomic(path: str, data: dict) -> None:
    try:
        os.makedirs(os.path.dirname(path) or ".", exist_ok=True)
        tmp = path + ".tmp"
        with open(tmp, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False)
        os.replace(tmp, path)
    except Exception:
        logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, "_write_json_atomic", exc_info=True)


# ============================================================================
# 2. Captcha handlers
# ============================================================================

def _maybe_handle_laf_captcha_reply(event, user_id: str, user_text: str, *, _line_send_text=None) -> bool:
    text = (user_text or "").strip()
    # Accept "1234" or "驗證碼 1234" etc.
    m = re.search(r"(^|\D)(\d{4})(\D|$)", text)
    if not m:
        return False

    req = _load_json(LAF_CAPTCHA_REQUEST_FILE)
    if not req:
        return False

    now = int(time.time())
    requested_at = int(req.get("requested_at") or 0)
    expires_at = int(req.get("expires_at") or 0)
    if expires_at and now > expires_at:
        return False
    if requested_at and (now - requested_at) > max(30, LAF_CAPTCHA_TTL_SECONDS):
        return False

    req_id = (req.get("request_id") or "").strip()
    if not req_id:
        return False

    code = m.group(2)
    resp = {
        "request_id": req_id,
        "captcha": code,
        "received_at": now,
        "from_user_id": user_id,
    }
    _write_json_atomic(LAF_CAPTCHA_RESPONSE_FILE, resp)

    # Best-effort ack
    if _line_send_text is not None:
        _line_send_text(event, user_id, "\u2705 \u5df2\u6536\u5230\u9a57\u8b49\u78bc\uff0cCASPER \u6b63\u5728\u767b\u5165\u6cd5\u6276\uff0c\u5b8c\u6210\u5f8c\u6211\u6703\u518d\u56de\u5831\u3002", prefer_push=False)
    return True


def _maybe_handle_generic_captcha_reply(event, user_id: str, user_text: str, *, _line_send_text=None) -> bool:
    """
    Handle replies for human-in-the-loop captcha requests.
    Only triggers if a pending request exists and is not expired.
    """
    req = _load_json(GEN_CAPTCHA_REQUEST_FILE)
    if not req:
        return False

    now = int(time.time())
    expires_at = int(req.get("expires_at") or 0)
    if expires_at and now > expires_at:
        return False

    req_id = (req.get("request_id") or "").strip()
    if not req_id:
        return False

    expected_len = int(req.get("expected_len") or 0)
    text = (user_text or "").strip()

    # Extract digits; accept either exact length or a reasonable range if not specified.
    digits = re.sub(r"[^0-9]", "", text)
    if expected_len and expected_len > 0:
        if len(digits) < expected_len:
            return False
        digits = digits[:expected_len]
    else:
        if not (4 <= len(digits) <= 12):
            return False

    resp = {
        "request_id": req_id,
        "captcha": digits,
        "received_at": now,
        "from_user_id": user_id,
    }
    _write_json_atomic(GEN_CAPTCHA_RESPONSE_FILE, resp)
    if _line_send_text is not None:
        _line_send_text(event, user_id, "\u2705 \u5df2\u6536\u5230\u9a57\u8b49\u78bc\uff0cCASPER \u6b63\u5728\u7e7c\u7e8c\u8655\u7406\u3002", prefer_push=False)
    return True


# ============================================================================
# 3. URL utility functions
# ============================================================================

def _is_loopback_base_url(base: str) -> bool:
    s = (base or "").strip()
    if not s:
        return True
    if "://" not in s:
        s = "https://" + s
    try:
        host = (urlparse(s).hostname or "").strip().lower()
    except Exception:
        return True
    if not host:
        return True
    if host == "localhost" or host == "::1":
        return True
    if host.startswith("127."):
        return True
    return False


def _normalize_public_base_url(base: str) -> str:
    s = (base or "").strip().strip("'\"")
    if not s:
        return ""
    if "://" not in s:
        s = "https://" + s
    return s.rstrip("/") + "/"


def _base_from_webhook_url(url: str) -> str:
    s = (url or "").strip().strip("'\"")
    if not s:
        return ""
    if "://" not in s:
        s = "https://" + s
    try:
        p = urlparse(s)
        if not p.scheme or not p.netloc:
            return ""
        return f"{p.scheme}://{p.netloc}/"
    except Exception:
        return ""


def _record_last_public_base_url():
    """
    Record the public base URL from the current request so background tasks can build downloadable links.
    Respects reverse proxies via X-Forwarded-Proto / X-Forwarded-Host.
    """
    try:
        from flask import request
        proto = (request.headers.get("X-Forwarded-Proto") or "").split(",")[0].strip()
        host = (request.headers.get("X-Forwarded-Host") or "").split(",")[0].strip()
        if not proto:
            proto = (request.scheme or "http").strip()
        if not host:
            host = (request.host or "").strip()
        base = _normalize_public_base_url(f"{proto}://{host}")
        if not base or _is_loopback_base_url(base):
            return
        with open(LINE_LAST_BASE_URL_FILE, "w", encoding="utf-8") as f:
            json.dump({"base_url": base, "updated_at": int(time.time())}, f, ensure_ascii=False)
    except Exception:
        logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, "_record_last_public_base_url", exc_info=True)


def _build_tailscale_base_url() -> str:
    """Build base URL from Tailscale IP if configured."""
    from skills.ops.export_text import _load_dotenv_value
    ts_ip = (
        os.environ.get("MAGI_TAILSCALE_IP")
        or _load_dotenv_value("MAGI_TAILSCALE_IP")
        or ""
    ).strip()
    if not ts_ip:
        return ""
    ts_port = (
        os.environ.get("MAGI_TAILSCALE_PORT")
        or _load_dotenv_value("MAGI_TAILSCALE_PORT")
        or "5002"
    ).strip()
    return f"http://{ts_ip}:{ts_port}/"


def _load_public_base_url() -> str:
    """
    Priority: explicit override -> Tailscale VPN -> LINE webhook -> cached base URL.
    """
    # 1. Explicit override
    env_base = _normalize_public_base_url(os.environ.get("MAGI_PUBLIC_BASE_URL") or "")
    if env_base and (not _is_loopback_base_url(env_base)):
        return env_base
    # 2. Tailscale (stable, always-on VPN)
    ts_base = _build_tailscale_base_url()
    if ts_base:
        return ts_base
    # 3. LINE webhook domain (Cloudflare tunnel, may rotate)
    webhook_base = _base_from_webhook_url(os.environ.get("MAGI_LINE_WEBHOOK_ENDPOINT") or "")
    if webhook_base and (not _is_loopback_base_url(webhook_base)):
        return webhook_base
    # 4. Cached base URL from last webhook
    try:
        if os.path.exists(LINE_LAST_BASE_URL_FILE):
            with open(LINE_LAST_BASE_URL_FILE, "r", encoding="utf-8") as f:
                data = json.load(f) or {}
            base = _normalize_public_base_url(data.get("base_url") or "")
            if base and (not _is_loopback_base_url(base)):
                return base
    except Exception:
        logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, "_load_public_base_url", exc_info=True)
    return ""


# ============================================================================
# 4. Export functions
# ============================================================================

def _export_text_to_static(text: str, prefix: str = "casper") -> dict:
    """
    Write a UTF-8 TXT file under /static/exports and return a public URL if available.
    """
    s = (text or "").strip()
    if not s:
        return {"success": False, "error": "empty text"}
    # Strip Markdown formatting -- TXT is plain text
    try:
        from api.tw_output_guard import strip_markdown_for_chat
        s = strip_markdown_for_chat(s)
    except Exception:
        logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, "_export_text_to_static", exc_info=True)
    try:
        os.makedirs(EXPORTS_DIR, exist_ok=True)
        stamp = time.strftime("%Y%m%d_%H%M%S")
        token = uuid.uuid4().hex[:10]
        filename = f"{prefix}_{stamp}_{token}.txt"
        path = os.path.join(EXPORTS_DIR, filename)
        with open(path, "w", encoding="utf-8") as f:
            f.write(s + "\n")
        base = _load_public_base_url()
        url = (base.rstrip("/") + f"/static/exports/{filename}") if base else ""
        return {"success": True, "path": path, "filename": filename, "url": url}
    except Exception as e:
        return {"success": False, "error": str(e)}


def _safe_export_stem(name: str, fallback: str = "document") -> str:
    raw = str(name or "").strip()
    if not raw:
        raw = fallback
    # Keep CJK characters, strip path separators and invalid filesystem chars.
    raw = re.sub(r'[\\/:*?"<>|]+', "_", raw)
    raw = re.sub(r"\s+", "_", raw).strip(" ._")
    return raw or fallback


def _export_file_meta(path: str) -> dict:
    p = os.path.abspath(path)
    filename = os.path.basename(p)
    base = _load_public_base_url().rstrip("/")
    url = f"{base}/static/exports/{filename}" if base else ""
    return {"success": True, "path": p, "filename": filename, "url": url}


def _find_chrome_binary() -> str:
    import shutil as _shutil
    candidates = [
        (os.environ.get("MAGI_CHROME_BIN") or "").strip(),
        "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome",
        "/Applications/Chromium.app/Contents/MacOS/Chromium",
        _shutil.which("google-chrome"),
        _shutil.which("chromium"),
        _shutil.which("chromium-browser"),
    ]
    for c in candidates:
        if c and os.path.exists(c):
            return c
    return ""


def _clean_document_export_text(text: str) -> str:
    txt = ihtml.unescape(str(text or ""))
    txt = txt.replace("<br>", "\n").replace("<br/>", "\n").replace("<br />", "\n")
    txt = re.sub(r"^\s*```[a-zA-Z0-9_-]*\s*$", "", txt, flags=re.MULTILINE)
    txt = re.sub(r"^#{1,6}\s*", "", txt, flags=re.MULTILINE)
    txt = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1", txt)
    txt = re.sub(r"\*\*(.+?)\*\*", r"\1", txt)
    txt = re.sub(r"__(.+?)__", r"\1", txt)
    txt = re.sub(r"(?<!\*)\*(?!\s)(.+?)(?<!\s)\*(?!\*)", r"\1", txt)
    txt = re.sub(r"(?<!_)_(?!\s)(.+?)(?<!\s)_(?!_)", r"\1", txt)
    txt = re.sub(r"`([^`]+)`", r"\1", txt)
    txt = re.sub(r"^\s*>\s?", "", txt, flags=re.MULTILINE)

    lines: list[str] = []
    for raw in txt.splitlines():
        line = raw.strip()
        if re.match(r"^\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$", line):
            continue
        if "|" in line and line.count("|") >= 2:
            cells = [c.strip() for c in line.strip("|").split("|")]
            line = "　".join(c for c in cells if c)
        line = re.sub(r"^\s*[-*+]\s+", "", line)
        lines.append(line)

    txt = "\n".join(lines)
    txt = re.sub(r"[ \t]+\n", "\n", txt)
    txt = re.sub(r"\n{3,}", "\n\n", txt)
    return txt.strip()


def _set_docx_font(run, *, size_pt: int = 14, bold: bool = False, font_name: str = "標楷體") -> None:
    from docx.shared import Pt  # type: ignore
    from docx.oxml.ns import qn  # type: ignore

    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    try:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    except Exception:
        logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, "_set_docx_font", exc_info=True)


def _looks_like_pleading_title(line: str) -> bool:
    text = str(line or "").strip()
    return bool(text and len(text) <= 24 and text.endswith("狀"))


def _is_meta_or_salutation_line(line: str) -> bool:
    text = str(line or "").strip()
    if not text:
        return False
    return bool(
        re.match(r"^(案號|股別|案由|法院|當事人|原告|被告|聲請人|相對人|上訴人|被上訴人|具狀人|撰狀人|受任人|此致|謹呈|中華民國)", text)
    )


def _is_signature_line(line: str) -> bool:
    text = str(line or "").strip()
    return bool(re.match(r"^(具狀人|撰狀人|受任人|代理人|中華民國)", text))


_PLEADING_META_LABELS = {
    "案號",
    "股別",
    "案由",
    "法院",
    "原告",
    "被告",
    "聲請人",
    "相對人",
    "債權人",
    "債務人",
    "上訴人",
    "被上訴人",
    "抗告人",
    "受任人",
    "當事人",
    "法定代理人",
    "訴訟代理人",
    "代理人",
    "代表人",
    "住",
    "設",
    "住所",
    "居所",
    "電話",
    "傳真",
    "手機",
}


def _split_pleading_meta_line(line: str) -> tuple[str, str] | None:
    text = str(line or "").strip()
    if not text:
        return None
    m = re.match(r"^(.{1,18}?)\s*[：:]\s*(.*)$", text)
    if not m:
        return None
    label = re.sub(r"[\s　]+", "", m.group(1) or "")
    if label not in _PLEADING_META_LABELS:
        return None
    return label, (m.group(2) or "").strip()


def _collect_pleading_meta_rows(lines: list[str], start: int = 0) -> tuple[list[tuple[str, str]], int]:
    rows: list[tuple[str, str]] = []
    i = start
    blank_seen = False
    while i < len(lines):
        line = str(lines[i] or "").strip()
        if not line:
            blank_seen = True
            i += 1
            continue
        split = _split_pleading_meta_line(line)
        if not split:
            break
        if blank_seen and rows and split[0] not in {
            "原告",
            "被告",
            "聲請人",
            "相對人",
            "債權人",
            "債務人",
            "上訴人",
            "被上訴人",
            "抗告人",
        }:
            break
        rows.append(split)
        blank_seen = False
        i += 1
    return rows, i


def _set_paragraph_distribute_alignment(paragraph) -> None:
    try:
        from docx.oxml import OxmlElement  # type: ignore
        from docx.oxml.ns import qn  # type: ignore

        ppr = paragraph._p.get_or_add_pPr()
        jc = ppr.find(qn("w:jc"))
        if jc is None:
            jc = OxmlElement("w:jc")
            ppr.append(jc)
        jc.set(qn("w:val"), "distribute")
    except Exception:
        logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, "_set_paragraph_distribute_alignment", exc_info=True)


def _set_table_borders_none(table) -> None:
    try:
        from docx.oxml import OxmlElement  # type: ignore
        from docx.oxml.ns import qn  # type: ignore

        tbl_pr = table._tbl.tblPr
        borders = tbl_pr.first_child_found_in("w:tblBorders")
        if borders is None:
            borders = OxmlElement("w:tblBorders")
            tbl_pr.append(borders)
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            tag = "w:" + edge
            element = borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                borders.append(element)
            element.set(qn("w:val"), "nil")
    except Exception:
        logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, "_set_table_borders_none", exc_info=True)


def _add_pleading_meta_table(doc, rows: list[tuple[str, str]]) -> None:
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT  # type: ignore
    from docx.shared import Cm, Pt  # type: ignore

    if not rows:
        return
    table = doc.add_table(rows=0, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    _set_table_borders_none(table)
    widths = (Cm(3.25), Cm(0.35), Cm(13.1))
    for label, value in rows:
        cells = table.add_row().cells
        for idx, width in enumerate(widths):
            cells[idx].width = width
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        values = (label, "：", value)
        for idx, cell in enumerate(cells):
            p = cell.paragraphs[0]
            pf = p.paragraph_format
            pf.line_spacing = Pt(26)
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            if idx == 0 and len(label) <= 6:
                _set_paragraph_distribute_alignment(p)
            run = p.add_run(values[idx])
            _set_docx_font(run, size_pt=16)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def _add_pleading_paragraph(doc, text: str, *, align: str = "body", bold: bool = False, size_pt: int = 14) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    from docx.shared import Pt  # type: ignore

    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = Pt(26)
    pf.space_before = Pt(0)
    pf.space_after = Pt(2)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if align == "body":
            pf.first_line_indent = Pt(28)
    run = p.add_run(text)
    _set_docx_font(run, size_pt=size_pt, bold=bold)


def _export_form_docx(preview_text: str, stem: str, title: str = "") -> dict:
    txt = _clean_document_export_text(preview_text)
    if not txt:
        return {"success": False, "error": "empty_text"}
    try:
        from docx import Document  # type: ignore
        from docx.shared import Cm, Pt  # type: ignore
        from docx.oxml.ns import qn  # type: ignore
    except Exception as e:
        return {"success": False, "error": f"python_docx_unavailable: {e}"}
    try:
        os.makedirs(EXPORTS_DIR, exist_ok=True)
        filename = f"{stem}.docx"
        path = os.path.join(EXPORTS_DIR, filename)
        doc = Document()
        section = doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

        normal = doc.styles["Normal"]
        normal.font.name = "標楷體"
        normal.font.size = Pt(14)
        try:
            normal._element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")
        except Exception:
            logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, "_export_form_docx", exc_info=True)

        lines = [ln.strip() for ln in txt.splitlines()]
        nonempty = [ln for ln in lines if ln]
        doc_title = str(title or "").strip() or "OSC 文件"
        skip_first_title = False
        if nonempty and _looks_like_pleading_title(nonempty[0]):
            doc_title = nonempty[0]
            skip_first_title = True
        _add_pleading_paragraph(doc, doc_title, align="center", bold=True, size_pt=26)

        start_idx = 0
        if skip_first_title:
            for idx, line in enumerate(lines):
                if line == doc_title:
                    start_idx = idx + 1
                    break
        meta_rows, meta_end = _collect_pleading_meta_rows(lines, start=start_idx)
        if meta_rows:
            _add_pleading_meta_table(doc, meta_rows)

        blank_pending = False
        first_seen = False
        for idx, line in enumerate(lines):
            if idx < meta_end:
                if skip_first_title and not first_seen and line == doc_title:
                    first_seen = True
                continue
            if skip_first_title and not first_seen and line == doc_title:
                first_seen = True
                continue
            if not line:
                blank_pending = True
                continue
            first_seen = True
            if blank_pending:
                spacer = doc.add_paragraph()
                spacer.paragraph_format.space_after = Pt(2)
                blank_pending = False
            if _is_signature_line(line):
                _add_pleading_paragraph(doc, line, align="right", size_pt=14)
            elif _is_meta_or_salutation_line(line):
                _add_pleading_paragraph(doc, line, align="left", size_pt=14)
            elif _looks_like_pleading_title(line):
                _add_pleading_paragraph(doc, line, align="center", bold=True, size_pt=18)
            else:
                _add_pleading_paragraph(doc, line, align="body", size_pt=14)
        doc.save(path)
        return _export_file_meta(path)
    except Exception as e:
        return {"success": False, "error": str(e)}


def _render_form_text_to_html(title: str, text: str) -> str:
    safe_title = ihtml.escape(str(title or "OSC \u6587\u4ef6"))
    lines = _clean_document_export_text(text).splitlines()
    body_parts = []
    first_title_skipped = False
    start_idx = 0
    for idx, raw in enumerate(lines):
        if _looks_like_pleading_title(raw) and raw.strip() == str(title or "").strip():
            start_idx = idx + 1
            first_title_skipped = True
            break
    meta_rows, meta_end = _collect_pleading_meta_rows(lines, start=start_idx)
    if meta_rows:
        body_parts.append("<table class='meta-table'><tbody>")
        for label, value in meta_rows:
            body_parts.append(
                "<tr>"
                f"<td class='meta-label'>{ihtml.escape(label)}</td>"
                "<td class='meta-colon'>：</td>"
                f"<td>{ihtml.escape(value)}</td>"
                "</tr>"
            )
        body_parts.append("</tbody></table>")
    for idx, raw in enumerate(lines):
        if idx < meta_end:
            continue
        line = raw.strip()
        if not line:
            body_parts.append("<div class='spacer'></div>")
            continue
        if (not first_title_skipped) and _looks_like_pleading_title(line) and line == str(title or "").strip():
            first_title_skipped = True
            continue
        first_title_skipped = True
        cls = "body"
        if _is_signature_line(line):
            cls = "signature"
        elif _is_meta_or_salutation_line(line):
            cls = "meta"
        elif _looks_like_pleading_title(line):
            cls = "subheading"
        body_parts.append(f"<p class='{cls}'>{ihtml.escape(line)}</p>")
    return (
        "<!doctype html><html><head><meta charset='utf-8'>"
        f"<title>{safe_title}</title>"
        "<style>"
        "@page{size:A4;margin:22mm 25mm 20mm 25mm;}"
        "body{font-family:'BiauKai','DFKai-SB','標楷體','Noto Serif CJK TC',serif;"
        "color:#111;line-height:1.85;font-size:16pt;}"
        "h1{margin:0 0 18px;text-align:center;font-size:22pt;font-weight:700;}"
        "p{margin:0 0 6px;}"
        "table.meta-table{width:100%;border-collapse:collapse;margin:0 0 10px;}"
        ".meta-table td{border:0;padding:0 3px 0 0;vertical-align:top;line-height:1.7;}"
        ".meta-label{width:3.2cm;text-align:justify;text-align-last:justify;}"
        ".meta-colon{width:.35cm;}"
        ".body{text-indent:2em;}"
        ".meta{text-indent:0;}"
        ".signature{text-align:right;text-indent:0;}"
        ".subheading{text-align:center;font-size:18pt;font-weight:700;text-indent:0;margin-top:8px;}"
        ".spacer{height:10px;}"
        "</style></head><body>"
        f"<h1>{safe_title}</h1>{''.join(body_parts)}</body></html>"
    )


def _wrap_pdf_line(text: str, max_width: float, font_name: str, size_pt: int) -> list[str]:
    from reportlab.pdfbase.pdfmetrics import stringWidth  # type: ignore

    source = str(text or "")
    lines: list[str] = []
    current = ""
    for char in source:
        candidate = current + char
        if current and stringWidth(candidate, font_name, size_pt) > max_width:
            lines.append(current)
            current = char
        else:
            current = candidate
    if current:
        lines.append(current)
    return lines or [""]


def _export_form_pdf_reportlab(title: str, preview_text: str, pdf_path: str) -> dict:
    from reportlab.lib.pagesizes import A4  # type: ignore
    from reportlab.pdfbase import pdfmetrics  # type: ignore
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont  # type: ignore
    from reportlab.pdfgen import canvas  # type: ignore

    font_name = "MSung-Light"
    try:
        pdfmetrics.registerFont(UnicodeCIDFont(font_name))
    except Exception:
        font_name = "STSong-Light"
        pdfmetrics.registerFont(UnicodeCIDFont(font_name))

    width, height = A4
    left = 72
    right = 72
    top = 62
    bottom = 56
    content_width = width - left - right
    y = height - top
    c = canvas.Canvas(pdf_path, pagesize=A4)

    def ensure_space(amount: float) -> None:
        nonlocal y
        if y - amount < bottom:
            c.showPage()
            c.setFont(font_name, 14)
            y = height - top

    c.setFont(font_name, 20)
    c.drawCentredString(width / 2, y, str(title or "OSC 文件"))
    y -= 34
    c.setFont(font_name, 14)

    for raw in _clean_document_export_text(preview_text).splitlines():
        line = raw.strip()
        if not line:
            y -= 10
            continue
        if _looks_like_pleading_title(line) and line == str(title or "").strip():
            continue
        if _is_signature_line(line):
            ensure_space(20)
            c.drawRightString(width - right, y, line)
            y -= 23
            continue
        prefix = "" if _is_meta_or_salutation_line(line) else "　　"
        wrapped = _wrap_pdf_line(prefix + line, content_width, font_name, 14)
        for part in wrapped:
            ensure_space(20)
            c.drawString(left, y, part)
            y -= 23

    c.save()
    if (not os.path.exists(pdf_path)) or os.path.getsize(pdf_path) < 64:
        return {"success": False, "error": "pdf_not_generated"}
    meta = _export_file_meta(pdf_path)
    meta["renderer"] = "reportlab"
    return meta


def _find_soffice_binary() -> str:
    candidates = [
        (os.environ.get("MAGI_SOFFICE_BIN") or "").strip(),
        shutil.which("soffice"),
        shutil.which("libreoffice"),
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        "/Applications/LibreOffice.app/Contents/MacOS/libreoffice",
    ]
    for candidate in candidates:
        if candidate and os.path.exists(candidate):
            return candidate
    return ""


def _export_docx_pdf(docx_path: str, stem: str) -> dict:
    source = os.path.abspath(str(docx_path or ""))
    if not source or not os.path.exists(source):
        return {"success": False, "error": "docx_missing"}
    soffice = _find_soffice_binary()
    if not soffice:
        return {"success": False, "error": "soffice_unavailable"}
    try:
        os.makedirs(EXPORTS_DIR, exist_ok=True)
        target = os.path.join(EXPORTS_DIR, f"{stem}.pdf")
        default_pdf = os.path.join(EXPORTS_DIR, os.path.splitext(os.path.basename(source))[0] + ".pdf")
        for old in {target, default_pdf}:
            try:
                if old and os.path.exists(old):
                    os.remove(old)
            except OSError:
                pass
        env = dict(os.environ)
        env.setdefault("HOME", os.path.expanduser("~"))
        result = subprocess.run(
            [
                soffice,
                "--headless",
                "--nologo",
                "--nofirststartwizard",
                "--convert-to",
                "pdf",
                "--outdir",
                EXPORTS_DIR,
                source,
            ],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            timeout=int(os.environ.get("MAGI_DOCX_PDF_TIMEOUT", "90") or "90"),
            env=env,
        )
        if result.returncode != 0:
            err = (result.stderr or result.stdout or "").strip()
            return {"success": False, "error": f"soffice_failed: {err}"}
        produced = default_pdf if os.path.exists(default_pdf) else target
        if produced != target and os.path.exists(produced):
            os.replace(produced, target)
        if (not os.path.exists(target)) or os.path.getsize(target) < 64:
            err = (result.stderr or result.stdout or "").strip()
            return {"success": False, "error": f"soffice_pdf_not_generated: {err}"}
        meta = _export_file_meta(target)
        meta["renderer"] = "libreoffice"
        meta["source_docx"] = source
        return meta
    except Exception as e:
        return {"success": False, "error": f"docx_pdf_convert_failed: {e}"}


def _export_form_pdf(title: str, preview_text: str, stem: str) -> dict:
    txt = str(preview_text or "").strip()
    if not txt:
        return {"success": False, "error": "empty_text"}
    pdf_name = f"{stem}.pdf"
    pdf_path = os.path.join(EXPORTS_DIR, pdf_name)
    try:
        os.makedirs(EXPORTS_DIR, exist_ok=True)

        # Render HTML
        html_content = _render_form_text_to_html(title, txt)

        # Generate PDF using weasyprint
        import weasyprint
        weasyprint.HTML(string=html_content).write_pdf(pdf_path)

        if (not os.path.exists(pdf_path)) or os.path.getsize(pdf_path) < 64:
            return {"success": False, "error": "pdf_not_generated"}

        return _export_file_meta(pdf_path)
    except Exception as e:
        import traceback
        err_msg = traceback.format_exc()
        try:
            fallback = _export_form_pdf_reportlab(title, txt, pdf_path)
            if fallback.get("success"):
                fallback["warning"] = f"weasyprint_failed_fallback_reportlab: {e}"
                return fallback
        except Exception as fallback_e:
            return {"success": False, "error": f"weasyprint_failed: {e}\n{err_msg}\nreportlab_failed: {fallback_e}"}
        return {"success": False, "error": f"weasyprint_failed: {e}\n{err_msg}"}


def _export_osc_form_files(title: str, preview_text: str, suggested_filename: str = "") -> dict:
    txt = str(preview_text or "").strip()
    if not txt:
        return {"success": False, "errors": [{"type": "common", "error": "empty_text"}]}
    stamp = time.strftime("%Y%m%d_%H%M%S")
    token = uuid.uuid4().hex[:8]
    stem = _safe_export_stem(suggested_filename, fallback="osc_form")
    full_stem = f"{stem}_{stamp}_{token}"
    txt = _clean_document_export_text(txt)
    docx_meta = _export_form_docx(txt, full_stem, title=title)
    if docx_meta.get("success"):
        pdf_meta = _export_docx_pdf(str(docx_meta.get("path") or ""), full_stem)
    else:
        pdf_meta = {"success": False, "error": "docx_unavailable_for_pdf"}
    if not pdf_meta.get("success"):
        fallback_pdf_meta = _export_form_pdf(title, txt, full_stem)
        if fallback_pdf_meta.get("success"):
            fallback_pdf_meta["warning"] = str(pdf_meta.get("error") or "docx_pdf_convert_failed")
            pdf_meta = fallback_pdf_meta
    errors = []
    if not docx_meta.get("success"):
        errors.append({"type": "docx", "error": str(docx_meta.get("error") or "docx_failed")})
    if not pdf_meta.get("success"):
        errors.append({"type": "pdf", "error": str(pdf_meta.get("error") or "pdf_failed")})
    ok = bool(docx_meta.get("success") or pdf_meta.get("success"))
    preferred = pdf_meta if pdf_meta.get("success") else (docx_meta if docx_meta.get("success") else {"success": False})
    return {
        "success": ok,
        "export": preferred,
        "export_docx": docx_meta,
        "export_pdf": pdf_meta,
        "errors": errors,
    }


def _public_url_for_local_file(local_path: str) -> str:
    """
    Return a public URL for a local file.
    If the file is already inside /static/, return its URL directly (no copy).
    Otherwise, copy to EXPORTS_DIR and return its URL.
    """
    try:
        p = (local_path or "").strip().strip("'\"")
        if not p or (not os.path.exists(p)):
            return ""
        base = _load_public_base_url().rstrip("/")
        if not base:
            return ""
        abs_p = os.path.abspath(p)
        static_abs = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "static"))
        # If already under /static/, serve directly without copying
        if abs_p.startswith(static_abs + os.sep):
            rel = abs_p[len(static_abs) + 1:]
            return f"{base}/static/{rel}"
        # Otherwise, copy to exports
        os.makedirs(EXPORTS_DIR, exist_ok=True)
        filename = os.path.basename(abs_p)
        stem, ext = os.path.splitext(filename)
        stamp = time.strftime("%Y%m%d_%H%M%S")
        token = uuid.uuid4().hex[:8]
        filename = f"{stem}_{stamp}_{token}{ext}"
        import shutil
        shutil.copy2(abs_p, os.path.join(EXPORTS_DIR, filename))
        return f"{base}/static/exports/{filename}"
    except Exception:
        return ""


# ============================================================================
# 5. Cloudflared tunnel management
# ============================================================================

def _cloudflared_pids_for_port(port: str) -> list[str]:
    """Return cloudflared PIDs for the MAGI webhook tunnel port only."""
    import subprocess
    try:
        pattern = f"cloudflared tunnel --url http://127.0.0.1:{str(port).strip()}"
        result = subprocess.run(
            ["pgrep", "-f", pattern],
            capture_output=True, text=True, timeout=3,
        )
        return [p.strip() for p in (result.stdout or "").splitlines() if p.strip()]
    except Exception:
        return []


def _magi_webhook_port() -> str:
    return (
        os.environ.get("MAGI_SERVER_PORT")
        or _load_dotenv_value("MAGI_SERVER_PORT")
        or "5002"
    ).strip()


def _is_cloudflared_alive() -> bool:
    """Check if the MAGI webhook cloudflared tunnel is actually running."""
    try:
        return bool(_cloudflared_pids_for_port(_magi_webhook_port()))
    except Exception:
        return False


def _ensure_cloudflared():
    """Start cloudflared if not running and always register webhook with LINE."""
    import subprocess
    import re as _re
    import time as _time
    try:
        log_path = os.path.join(os.path.dirname(__file__), "..", "logs", "cloudflared.log")
        os.makedirs(os.path.dirname(log_path), exist_ok=True)
        already_running = False
        _cf_local_port = _magi_webhook_port()

        # Count only the MAGI webhook tunnel. Other tunnels, such as Paperclip
        # sharing, may legitimately run on different ports.
        cf_pids = _cloudflared_pids_for_port(_cf_local_port)

        if len(cf_pids) > 1:
            logger.warning("cloudflared: Found %d MAGI webhook instances, restarting port %s cleanly", len(cf_pids), _cf_local_port)
            try:
                for pid in cf_pids:
                    subprocess.run(["kill", pid], capture_output=True, timeout=3)
                _time.sleep(1)
            except Exception:
                pass
            cf_pids = []

        if len(cf_pids) == 1:
            # Check if the log still has the URL (not truncated)
            try:
                with open(log_path) as f:
                    content = f.read()
                if _re.search(r'https://[a-z0-9-]+\.trycloudflare\.com', content):
                    logger.info("cloudflared already running (pid=%s)", cf_pids[0])
                    already_running = True
                else:
                    logger.warning("cloudflared running but log empty, restarting")
                    subprocess.run(["kill", cf_pids[0]], capture_output=True, timeout=3)
                    _time.sleep(1)
            except Exception:
                logger.info("cloudflared already running (pid=%s)", cf_pids[0])
                already_running = True

        if not already_running:
            try:
                subprocess.run(["pkill", "-f", f"cloudflared tunnel --url http://127.0.0.1:{_cf_local_port}"],
                               capture_output=True, timeout=3)
            except Exception:
                logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, "_ensure_cloudflared/pkill", exc_info=True)
            logger.info("Starting cloudflared tunnel...")
            _cf_log_fh = open(log_path, "w")  # kept open for cloudflared's lifetime
            logger.info("cloudflared → local port %s", _cf_local_port)
            _cf_proc = subprocess.Popen(
                ["/opt/homebrew/bin/cloudflared", "tunnel", "--url", f"http://127.0.0.1:{_cf_local_port}", "--no-autoupdate"],
                stdout=subprocess.DEVNULL, stderr=_cf_log_fh,
            )
            # Safety net: register atexit handler to close file handle
            import atexit
            def _atexit_close_cf_log(fh=_cf_log_fh):
                try:
                    if fh and not fh.closed:
                        fh.close()
                except Exception:
                    pass
            atexit.register(_atexit_close_cf_log)

            # Cleanup: close log file handle when cloudflared exits
            def _cleanup_cf_log(proc=_cf_proc, fh=_cf_log_fh):
                try:
                    proc.wait(timeout=3600)  # don't block forever
                except subprocess.TimeoutExpired:
                    logger.warning("cloudflared cleanup wait timed out after 1h")
                except Exception as e:
                    logger.debug("cloudflared wait error: %s", e)
                finally:
                    try:
                        if fh and not fh.closed:
                            fh.close()
                    except Exception:
                        logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, "_cleanup_cf_log", exc_info=True)
            threading.Thread(target=_cleanup_cf_log, daemon=True, name="cf-log-cleanup").start()

        def _register():
            cf_url = ""
            if already_running:
                try:
                    with open(log_path) as f:
                        m = _re.search(r'https://[a-z0-9-]+\.trycloudflare\.com', f.read())
                        if m:
                            cf_url = m.group(0)
                except Exception:
                    logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, "_register/read_log", exc_info=True)
            if not cf_url:
                for _ in range(30):
                    _time.sleep(1)
                    try:
                        with open(log_path) as f:
                            m = _re.search(r'https://[a-z0-9-]+\.trycloudflare\.com', f.read())
                            if m:
                                cf_url = m.group(0)
                                break
                    except Exception:
                        logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, "_register/wait_log", exc_info=True)
            if not cf_url:
                logger.error("Could not get cloudflare tunnel URL after 30s")
                return
            webhook_url = f"{cf_url}/line/webhook"
            logger.info("Tunnel: %s", cf_url)
            # Load LINE token
            token = os.environ.get("MAGI_LINE_CHANNEL_ACCESS_TOKEN", "")
            if not token:
                env_path = os.path.join(os.path.dirname(__file__), "..", ".env")
                try:
                    with open(env_path) as f:
                        for ln in f:
                            if ln.strip().startswith("MAGI_LINE_CHANNEL_ACCESS_TOKEN="):
                                token = ln.strip().split("=", 1)[1].strip().strip("\"'")
                                break
                except Exception:
                    logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, "_register/load_token", exc_info=True)
            if not token:
                logger.warning("No LINE token, skipping webhook registration")
                return
            import urllib.request
            import urllib.parse
            if not already_running:
                _time.sleep(3)  # Wait for new tunnel to be routable
            # Check if LINE already points to this URL
            try:
                get_req = urllib.request.Request(
                    "https://api.line.me/v2/bot/channel/webhook/endpoint",
                    method="GET",
                    headers={"Authorization": f"Bearer {token}"},
                )
                with urllib.request.urlopen(get_req, timeout=10) as resp:
                    current = json.loads(resp.read())
                if current.get("endpoint") == webhook_url:
                    logger.info("LINE webhook already correct: %s", webhook_url)
                    return
                logger.info("LINE webhook mismatch: %s -> %s", current.get("endpoint"), webhook_url)
            except Exception:
                logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, "_register/check_webhook", exc_info=True)
            data = json.dumps({"endpoint": webhook_url}).encode()
            registered = False
            for attempt in range(3):
                try:
                    req = urllib.request.Request(
                        "https://api.line.me/v2/bot/channel/webhook/endpoint",
                        data=data, method="PUT",
                        headers={"Authorization": f"Bearer {token}", "Content-Type": "application/json"},
                    )
                    with urllib.request.urlopen(req, timeout=10) as resp:
                        logger.info("LINE webhook registered: %s -> %s", webhook_url, resp.status)
                        registered = True
                        break
                except Exception as e:
                    logger.warning("LINE webhook registration attempt %d/3 failed: %s", attempt + 1, e)
                    _time.sleep(5)
            if not registered:
                logger.error("LINE webhook registration failed after 3 attempts")
            # Telegram webhook auto-registration
            try:
                from api.webhooks.telegram import _load_telegram_bot_token, _load_telegram_webhook_secret
                tg_token = _load_telegram_bot_token()
                tg_secret = _load_telegram_webhook_secret()
                if tg_token:
                    tg_webhook_url = f"{cf_url}/telegram/webhook"
                    tg_data = urllib.parse.urlencode({"url": tg_webhook_url, **({"secret_token": tg_secret} if tg_secret else {})}).encode()
                    tg_req = urllib.request.Request(f"https://api.telegram.org/bot{tg_token}/setWebhook", data=tg_data)
                    with urllib.request.urlopen(tg_req, timeout=10) as tg_resp:
                        logger.info("Telegram webhook registered: %s -> %s", tg_webhook_url, tg_resp.status)
            except Exception as tg_e:
                logger.warning("Telegram webhook registration failed: %s", tg_e)
            # Save URLs
            agent_dir = os.path.join(os.path.dirname(__file__), "..", ".agent")
            os.makedirs(agent_dir, exist_ok=True)
            try:
                with open(os.path.join(agent_dir, "line_webhook_url.txt"), "w") as f:
                    f.write(webhook_url + "\n")
                with open(os.path.join(agent_dir, "cloudflare_tunnel_url.txt"), "w") as f:
                    f.write(cf_url + "\n")
            except Exception:
                logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, "_register/save_urls", exc_info=True)
        threading.Thread(target=_register, daemon=True, name="cloudflared-register").start()
    except Exception as e:
        logger.warning("cloudflared startup failed: %s", e)


# ============================================================================
# 6. Monitoring threads
# ============================================================================

def _cloudflared_watchdog():
    import time as _time
    _INTERVAL = 90
    _time.sleep(60)  # wait 60s after startup before first check
    while True:
        try:
            if not _is_cloudflared_alive():
                logger.warning("cloudflared died -- restarting...")
                _ensure_cloudflared()
        except Exception as e:
            logger.warning("cloudflared watchdog error: %s", e)
        _time.sleep(_INTERVAL)


def _paperclip_share_gateway_port() -> str:
    return str(os.environ.get("PAPERCLIP_SHARE_GATEWAY_PORT") or "5014").strip() or "5014"


def _paperclip_share_url_file() -> str:
    return os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".runtime", "osc_share_public_base_url.txt"))


def _paperclip_share_tunnel_pids_for_port(port: str) -> list[str]:
    try:
        pattern = f"cloudflared tunnel --url http://127.0.0.1:{str(port).strip()}"
        result = subprocess.run(["pgrep", "-f", pattern], capture_output=True, text=True, timeout=5)
        if result.returncode == 0:
            return [line.strip() for line in result.stdout.splitlines() if line.strip()]
    except Exception:
        logger.debug("silent-catch at %s:%s", __name__, "_paperclip_share_tunnel_pids_for_port", exc_info=True)
    return []


def _paperclip_share_gateway_health_ok(port: str) -> bool:
    try:
        with urllib.request.urlopen(f"http://127.0.0.1:{port}/health", timeout=5) as resp:
            return 200 <= int(resp.status) < 300
    except Exception:
        return False


def _paperclip_share_public_health_ok() -> bool:
    try:
        url_path = _paperclip_share_url_file()
        if not os.path.exists(url_path):
            return False
        base = open(url_path, encoding="utf-8").read().strip().rstrip("/")
        if not base:
            return False
        with urllib.request.urlopen(base + "/health", timeout=10) as resp:
            return 200 <= int(resp.status) < 300
    except Exception:
        return False


def _paperclip_share_public_base_is_managed_tunnel() -> bool:
    try:
        url_path = _paperclip_share_url_file()
        if not os.path.exists(url_path):
            return True
        base = open(url_path, encoding="utf-8").read().strip().rstrip("/")
        if not base:
            return True
        return ".trycloudflare.com" in base
    except Exception:
        return True


_PAPERCLIP_SHARE_LAUNCHD_LABELS = (
    "com.magi.paperclip-share-gateway",
    "com.magi.paperclip-share-tunnel",
)


def _paperclip_share_launchd_plists() -> dict[str, str]:
    launch_agents = os.path.expanduser("~/Library/LaunchAgents")
    return {
        label: os.path.join(launch_agents, f"{label}.plist")
        for label in _PAPERCLIP_SHARE_LAUNCHD_LABELS
    }


def _launchctl_label_loaded(label: str) -> bool:
    if sys.platform != "darwin":
        return False
    try:
        result = subprocess.run(["launchctl", "list", label], capture_output=True, text=True, timeout=5)
        return result.returncode == 0
    except Exception:
        logger.debug("silent-catch at %s:%s", __name__, "_launchctl_label_loaded", exc_info=True)
        return False


def _stop_unmanaged_paperclip_share_processes(port: str) -> None:
    """Stop fallback share-tunnel processes before launchd takes ownership."""
    patterns = [
        f"scripts/share_gateway.py --port {str(port).strip()}",
        f"cloudflared tunnel --url http://127.0.0.1:{str(port).strip()}",
    ]
    for pattern in patterns:
        try:
            subprocess.run(["pkill", "-f", pattern], capture_output=True, text=True, timeout=10)
        except Exception:
            logger.debug("silent-catch at %s:%s", __name__, "_stop_unmanaged_paperclip_share_processes", exc_info=True)


def _bootstrap_paperclip_share_launchd(port: str) -> bool:
    if sys.platform != "darwin":
        return False
    plists = _paperclip_share_launchd_plists()
    if not all(os.path.exists(path) for path in plists.values()):
        return False

    domain = f"gui/{os.getuid()}"
    loaded_before = [_launchctl_label_loaded(label) for label in _PAPERCLIP_SHARE_LAUNCHD_LABELS]
    if not all(loaded_before):
        _stop_unmanaged_paperclip_share_processes(port)

    for label, plist in plists.items():
        if _launchctl_label_loaded(label):
            continue
        try:
            subprocess.run(["launchctl", "bootstrap", domain, plist], capture_output=True, text=True, timeout=15)
        except Exception:
            logger.debug("silent-catch at %s:%s", __name__, f"_bootstrap_paperclip_share_launchd:{label}", exc_info=True)
        if not _launchctl_label_loaded(label):
            try:
                subprocess.run(["launchctl", "load", plist], capture_output=True, text=True, timeout=15)
            except Exception:
                logger.debug("silent-catch at %s:%s", __name__, f"_load_paperclip_share_launchd:{label}", exc_info=True)

    return all(_launchctl_label_loaded(label) for label in _PAPERCLIP_SHARE_LAUNCHD_LABELS)


def _paperclip_share_launchd_managed(port: str | None = None) -> bool:
    if sys.platform != "darwin":
        return False
    if all(_launchctl_label_loaded(label) for label in _PAPERCLIP_SHARE_LAUNCHD_LABELS):
        return True
    if port:
        return _bootstrap_paperclip_share_launchd(port)
    return False


def _kickstart_paperclip_share_launchd() -> None:
    domain = f"gui/{os.getuid()}"
    for label in _PAPERCLIP_SHARE_LAUNCHD_LABELS:
        try:
            subprocess.run(["launchctl", "kickstart", "-k", f"{domain}/{label}"], capture_output=True, text=True, timeout=15)
        except Exception:
            logger.debug("silent-catch at %s:%s", __name__, f"_kickstart_paperclip_share_launchd:{label}", exc_info=True)


def _ensure_paperclip_share_tunnel() -> None:
    if str(os.environ.get("PAPERCLIP_SHARE_TUNNEL_DISABLE") or "").strip().lower() in {"1", "true", "yes", "on"}:
        return
    with _SHARE_TUNNEL_LOCK:
        port = _paperclip_share_gateway_port()
        gateway_ok = _paperclip_share_gateway_health_ok(port)
        tunnel_ok = bool(_paperclip_share_tunnel_pids_for_port(port))
        public_ok = _paperclip_share_public_health_ok()
        if public_ok and not _paperclip_share_public_base_is_managed_tunnel():
            return
        if gateway_ok and tunnel_ok and public_ok:
            return

        if _paperclip_share_launchd_managed(port):
            logger.warning(
                "Paperclip share tunnel unhealthy but launchd-managed; kickstarting launchd jobs (gateway=%s tunnel=%s public=%s)",
                gateway_ok,
                tunnel_ok,
                public_ok,
            )
            _kickstart_paperclip_share_launchd()
            return

        root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
        script = os.path.join(root, "scripts", "start_paperclip_share_tunnel.sh")
        if not os.path.exists(script):
            logger.warning("Paperclip share tunnel script missing: %s", script)
            return
        logger.warning(
            "Paperclip share tunnel unhealthy; restarting (gateway=%s tunnel=%s public=%s)",
            gateway_ok,
            tunnel_ok,
            public_ok,
        )
        env = dict(os.environ)
        env.setdefault("MAGI_ROOT", root)
        subprocess.run(["bash", script], cwd=root, env=env, timeout=90, check=False)


def _paperclip_share_tunnel_watchdog():
    import time as _time
    interval = max(60, int(os.environ.get("PAPERCLIP_SHARE_TUNNEL_WATCHDOG_INTERVAL_SEC", "120") or "120"))
    _time.sleep(15)
    while True:
        try:
            _ensure_paperclip_share_tunnel()
        except Exception as e:
            logger.warning("Paperclip share tunnel watchdog error: %s", e)
        _time.sleep(interval)


def _preload_faiss():
    try:
        from skills.memory.mem_bridge import _get_faiss_index
        idx = _get_faiss_index()
        if idx:
            logger.info("FAISS index pre-loaded: %d vectors", getattr(idx, 'total', 0))
    except Exception as e:
        logger.warning("FAISS pre-load failed (non-fatal): %s", e)


def _warmup_omlx():
    try:
        import time as _t
        _t.sleep(5)  # let Ollama/oMLX finish startup
        from skills.bridge.http_pool import get_session
        from api.model_config import TEXT_PRIMARY_MODEL
        _model = os.environ.get("CASPER_LOCAL_MODEL", TEXT_PRIMARY_MODEL)
        _chat_url = os.environ.get("MAGI_OMLX_CHAT_URL", "http://127.0.0.1:11434")
        r = get_session().post(f"{_chat_url}/v1/chat/completions", json={
            "model": _model,
            "messages": [{"role": "user", "content": "ping"}],
            "max_tokens": 1, "temperature": 0,
        }, timeout=120)
        if r.status_code == 200:
            logger.info("Local LLM (%s) warmed up", _model)
        else:
            logger.warning("LLM warmup got %d", r.status_code)
    except Exception as e:
        logger.warning("LLM warmup failed (non-fatal): %s", e)


def _start_laf_gmail_monitor():
    """Background thread: scan Gmail for LAF emails every 300s."""
    try:
        _laf_paths = [
            os.path.join(os.path.dirname(__file__), '..', 'casper_ecosystem', 'law_firm_orchestrators'),
            os.path.join(os.path.dirname(__file__), '..', 'skills', 'legal'),
        ]
        for p in _laf_paths:
            if p not in sys.path:
                sys.path.insert(0, p)
        from laf_orchestrator import LAFOrchestrator
        laf_orch = LAFOrchestrator(dry_run=False)
        laf_orch.run_monitor()  # blocking loop (interval=300s)
    except Exception as e:
        logger.warning("LAF Gmail Monitor failed to start: %s", e)


# _start_filereview_email_monitor() removed 2026-04-20:
# file-review Gmail scan is integrated into the LAF Gmail monitor cycle
# (see laf_orchestrator.run_monitor() -> scan_for_file_review_emails()).
# The standalone function's log marker never emitted because run_startup_hooks
# did not spawn it; removal prevents menubar detection from relying on a
# string that was never logged.


# ============================================================================
# 7. Main startup entry point
# ============================================================================

def _cleanup_old_exports(days: int = 30) -> int:
    """Remove stale generated export files without importing api.server."""
    try:
        cutoff = time.time() - max(1, int(days)) * 86400
        cleaned = 0
        export_roots = {
            EXPORTS_DIR,
            os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "exports")),
        }
        for root in export_roots:
            if not root or not os.path.isdir(root):
                continue
            for dirpath, _dirnames, filenames in os.walk(root):
                for filename in filenames:
                    path = os.path.join(dirpath, filename)
                    try:
                        if os.path.getmtime(path) < cutoff:
                            os.remove(path)
                            cleaned += 1
                    except Exception:
                        logger.debug("silent-catch at %s:%s", __name__, "_cleanup_old_exports/file", exc_info=True)
        return cleaned
    except Exception:
        logger.debug("silent-catch at %s:%s", __name__, "_cleanup_old_exports", exc_info=True)
        return 0


def run_startup_hooks(app, orchestrator):
    """
    Run all startup hooks: FAISS preload, oMLX warmup, cloudflared tunnel,
    NAS mount guard, LAF Gmail monitor, and export cleanup.

    Called from server.py after all routes and helpers are registered.
    Parameters:
        app         - the Flask app instance
        orchestrator - the main orchestrator instance
    """
    global _STARTUP_HOOKS_DONE
    with _STARTUP_HOOKS_LOCK:
        if _STARTUP_HOOKS_DONE:
            logger.warning(
                "run_startup_hooks() called again in the same process — skipping "
                "(double-import or circular-import detected; LAF monitor already running)"
            )
            return
        _STARTUP_HOOKS_DONE = True

    _startup_enabled = str(
        os.environ.get("MAGI_DISABLE_SERVER_STARTUP_HOOKS", "0")
    ).strip().lower() not in {"1", "true", "yes", "on"}

    if not _startup_enabled:
        logger.info("Server startup hooks disabled by MAGI_DISABLE_SERVER_STARTUP_HOOKS")
        return

    # Cleanup old export files (>30 days). Keep this local to avoid importing
    # api.server while server.py is running as __main__, which double-initializes
    # routes, orchestrators, and startup hooks.
    _n_cleaned = _cleanup_old_exports(days=30)
    if _n_cleaned:
        logger.info("Startup: cleaned %d old exports", _n_cleaned)

    # Pre-load FAISS index in background
    threading.Thread(target=_preload_faiss, daemon=True, name="faiss-preload").start()

    # Warm up local LLM
    threading.Thread(target=_warmup_omlx, daemon=True, name="omlx-warmup").start()

    # Cloudflared tunnel
    try:
        _ensure_cloudflared()
    except Exception:
        logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, "run_startup_hooks/cloudflared", exc_info=True)

    # Cloudflared watchdog
    threading.Thread(target=_cloudflared_watchdog, daemon=True, name="cloudflared-watchdog").start()

    # Paperclip public share tunnel watchdog. This keeps the share-only tunnel
    # available and refreshes the base URL file when Cloudflare Quick Tunnel
    # rotates the temporary hostname.
    threading.Thread(
        target=_paperclip_share_tunnel_watchdog,
        daemon=True,
        name="paperclip-share-tunnel-watchdog",
    ).start()

    # NAS SMB auto-mount guard
    try:
        from api.nas_mount_guard import start_nas_mount_guard
        start_nas_mount_guard(interval=120)
    except Exception as e:
        logger.warning("NAS mount guard failed to start: %s", e)

    # LAF Gmail background monitor
    try:
        _laf_gmail_thread = threading.Thread(
            target=_start_laf_gmail_monitor,
            daemon=True,
            name="laf-gmail-monitor",
        )
        _laf_gmail_thread.start()
        logger.info("LAF Gmail Monitor background thread started")
    except Exception as e:
        logger.warning("LAF Gmail Monitor failed to start: %s", e)

    # 閱卷 Email 監控已整合進法扶 Gmail Monitor 的 poll cycle
    # （同一個信箱，每輪掃完法扶信件後順便掃閱卷信件，不另開 thread）
    logger.info("File Review Email Monitor: integrated into LAF Gmail Monitor cycle")
