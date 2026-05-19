"""
Runtime-facing Web/API routes extracted from server.py.

This module owns lightweight dashboard-supporting routes that depend on
runtime objects injected from the main server bootstrap:
  - process monitor page + APIs
  - vector memory dashboard APIs
  - OSC chat/poll helper APIs
  - legacy judgments JSON compatibility API
"""

from __future__ import annotations

import json
import html
import os
import re
import subprocess
import threading
import urllib.parse
import uuid
from collections import defaultdict
from datetime import datetime
from pathlib import Path
from typing import Any

from flask import Blueprint, jsonify, render_template, request
from flask_login import current_user, login_required
from werkzeug.utils import secure_filename


def _parse_etime_to_sec(raw: str) -> int:
    text = (raw or "").strip()
    if not text:
        return 0
    match = re.match(r"^(?:(\d+)-)?(?:(\d+):)?(\d+):(\d+)$", text)
    if not match:
        return 0
    dd = int(match.group(1) or 0)
    hh = int(match.group(2) or 0)
    mm = int(match.group(3) or 0)
    ss = int(match.group(4) or 0)
    return (dd * 86400) + (hh * 3600) + (mm * 60) + ss


def _process_monitor_markers(magi_root: Path) -> tuple[list[str], list[str], dict[str, str]]:
    worker_markers = [
        "skills/judgment-collector/action.py",
        "skills/file-review-orchestrator/action.py",
        "skills/transcript-downloader/action.py",
        "skills/laf-portal-automation/action.py",
        "skills/laf-orchestrator/action.py",
        "skills/laf-withdrawal-report/action.py",
        "skills/laf-refine-case/action.py",
        "skills/osc-orchestrator/action.py",
        "skills/osc-scan-folder/action.py",
        "skills/pdf-namer/action.py",
        "skills/crawler-targets/action.py",
        "skills/statutes-vdb/action.py",
        "skills/magi-autopilot/action.py",
    ]
    try:
        from daemon import REAPER_NEVER_KILL as daemon_never_kill

        core_markers = list(daemon_never_kill)
    except Exception:
        core_markers = [
            f"{magi_root}/daemon.py",
            "api/server.py",
            "api/discord_bot.py",
            "rpc-server",
        ]
    core_labels = {
        f"{magi_root}/daemon.py": "Daemon",
        "api/server.py": "API/LINE Webhook",
        "api/discord_bot.py": "Discord Bot",
        "rpc-server": "RPC Worker",
    }
    return worker_markers, core_markers, core_labels


def _chat_upload_dir(magi_root: Path) -> Path:
    path = magi_root / ".agent" / "chat_uploads"
    path.mkdir(parents=True, exist_ok=True)
    return path


def _magi_web_outputs_dir(magi_root: Path) -> Path:
    path = magi_root / "static" / "exports" / "magi_outputs"
    path.mkdir(parents=True, exist_ok=True)
    return path


def _extract_chat_upload_text(path: Path, filename: str) -> dict[str, Any]:
    try:
        from api.handlers.document_handler import extract_text_from_uploaded_file

        return extract_text_from_uploaded_file(str(path), filename=filename)
    except Exception as exc:
        return {"success": False, "text": "", "kind": "", "title": filename, "error": str(exc)}


def _extract_chat_upload_text_for_task(path: Path, filename: str, task: str = "") -> dict[str, Any]:
    suffix = Path(filename or "").suffix.lower()
    if task == "summary" and suffix == ".pdf":
        try:
            from skills.documents.pdf_bridge import _extract_text_pdftotext, _is_meaningful_text

            text, pages = _extract_text_pdftotext(str(path), max_pages=1000000)
            if _is_meaningful_text(text):
                return {
                    "success": True,
                    "text": str(text),
                    "kind": "pdf",
                    "title": filename,
                    "error": "",
                    "extractor": "pdftotext_fast",
                    "pages": pages,
                }
        except Exception:
            pass
    return _extract_chat_upload_text(path, filename)


def _safe_web_url(raw_url: str) -> str:
    text = str(raw_url or "").strip()
    try:
        parsed = urllib.parse.urlparse(text)
    except Exception:
        return ""
    if parsed.scheme.lower() not in {"http", "https", "mailto"}:
        return ""
    return text


def _artifact_download_url(path: Path, *, inline: bool = False) -> str:
    query = urllib.parse.urlencode({"path": str(path), "inline": "1" if inline else "0"})
    return f"/api/osc/files/content?{query}"


def _web_delivery_kind(instruction: str) -> tuple[str, str]:
    text = str(instruction or "").lower()
    if any(token in text for token in ("逐字稿", "轉錄", "transcript", "transcribe", "聽打")):
        return "transcript", "逐字稿"
    if any(token in text for token in ("翻譯", "translate", "translation", "譯成")):
        return "translation", "翻譯稿"
    if any(token in text for token in ("摘要", "summary", "summarize", "整理重點", "重點整理")):
        return "summary", "摘要報告"
    return "magi", "MAGI 處理結果"


def _should_create_web_delivery_artifacts(
    instruction: str,
    reply: str,
    *,
    source_filename: str = "",
) -> bool:
    text = f"{instruction}\n{source_filename}".lower()
    if not str(reply or "").strip():
        return False
    if any(token in text for token in ("逐字稿", "轉錄", "transcript", "transcribe", "聽打")):
        return True
    if any(token in text for token in ("翻譯", "translate", "translation", "譯成")):
        return True
    if any(token in text for token in ("摘要", "summary", "summarize", "整理重點", "重點整理")):
        return True
    return bool(source_filename and len(str(reply or "")) >= 900)


_CHAT_UPLOAD_AUDIO_EXTS = {".aac", ".aiff", ".flac", ".m4a", ".mp3", ".ogg", ".opus", ".wav", ".webm"}
_CHAT_UPLOAD_VIDEO_EXTS = {".m4v", ".mov", ".mp4", ".mpeg", ".mpg", ".webm"}


def _web_upload_requested_task(instruction: str, filename: str = "") -> str:
    text = f"{instruction}\n{filename}".lower().replace("＠", "@")
    suffix = Path(filename or "").suffix.lower()
    if suffix in _CHAT_UPLOAD_AUDIO_EXTS | _CHAT_UPLOAD_VIDEO_EXTS:
        return "transcript"
    if any(token in text for token in ("逐字稿", "轉錄", "transcript", "transcribe", "聽打")):
        return "transcript"
    if any(token in text for token in ("翻譯", "translate", "translation", "譯成")):
        return "translation"
    if any(token in text for token in ("摘要", "summary", "summarize", "整理重點", "重點整理")):
        return "summary"
    return ""


def _web_summary_length(instruction: str) -> str:
    text = str(instruction or "").lower().replace("＠", "@")
    if any(token in text for token in ("@heavy", "heavy", "詳細", "完整", "深度", "long", "detailed")):
        return "long"
    if any(token in text for token in ("簡短", "精簡", "簡要", "short", "brief")):
        return "short"
    return "medium"


def _web_heavy_opt_in(instruction: str) -> bool:
    text = str(instruction or "").lower().replace("＠", "@").lstrip()
    return text.startswith("@heavy ") or text.startswith("@重型 ") or any(
        token in text for token in (" @heavy ", "\n@heavy ", "使用 heavy", "重型模型", "深度模式")
    )


def _normalize_direct_reply(reply: str, *, task: str, source_text: str, source_filename: str, instruction: str) -> str:
    from api.handlers.output_quality_handler import (
        build_legal_document_summary_fallback,
        detect_output_quality_issue,
        format_quality_gate_failure,
    )

    output = str(reply or "").strip()
    issue = detect_output_quality_issue(task, output, source_chars=len(source_text or ""))
    if task == "summary" and issue:
        fallback = build_legal_document_summary_fallback(
            source_text,
            source_name=source_filename,
            instruction=instruction,
        )
        if fallback:
            return fallback
    if issue:
        return format_quality_gate_failure(task, issue)
    return output


def _run_direct_web_upload_text_task(
    orchestrator: Any,
    *,
    root: Path,
    target: Path,
    original_name: str,
    instruction: str,
    extracted: dict[str, Any],
    user_id: str,
) -> dict[str, Any] | None:
    task = _web_upload_requested_task(instruction, original_name)
    if task not in {"summary", "translation"}:
        return None

    source_text = str(extracted.get("text") or "").strip()
    if not source_text:
        return None

    suffix = Path(original_name or "").suffix.lower()
    # Tiny notes can stay on the conversational route; long files and PDFs need
    # a deterministic document route so they cannot be mistaken for case intake.
    if task == "summary" and suffix != ".pdf" and len(source_text) < 2000:
        return None

    if task == "summary":
        summary_length = _web_summary_length(instruction)
        heavy_opt_in = _web_heavy_opt_in(instruction)
        reply = ""
        try:
            extractive_threshold = int(os.environ.get("MAGI_WEB_UPLOAD_EXTRACTIVE_SUMMARY_THRESHOLD", "30000") or "30000")
        except Exception:
            extractive_threshold = 30000
        model_allowed = len(source_text) < max(12000, extractive_threshold)
        if suffix == ".pdf" and model_allowed:
            try:
                from skills.documents.pdf_bridge import summarize_pdf

                reply = summarize_pdf(str(target), summary_length=summary_length)
            except Exception:
                reply = ""
        if model_allowed and not str(reply or "").strip():
            try:
                from api.handlers.summary_handler import summarize_text_resilient

                result = summarize_text_resilient(source_text, summary_length=summary_length, heavy=heavy_opt_in)
                if isinstance(result, dict) and result.get("success"):
                    reply = str(result.get("text") or "").strip()
            except Exception:
                reply = ""
        reply = _normalize_direct_reply(
            reply,
            task=task,
            source_text=source_text,
            source_filename=original_name,
            instruction=instruction,
        )
        if not reply:
            return None
        if str(reply).lstrip().startswith("❌"):
            return {"task": task, "reply": reply, "artifacts": []}
        artifacts = _create_web_delivery_artifacts(
            root,
            instruction=instruction,
            reply=reply,
            source_filename=original_name,
        )
        return {"task": task, "reply": reply, "artifacts": artifacts}

    if task == "translation":
        try:
            from api.handlers.document_handler import (
                cap_translation_source_text,
                polish_translated_document_text,
                prepare_document_text_for_llm,
            )
            from api.handlers.translation_handler import translate_text_complete
            from api.handlers.output_quality_handler import run_output_quality_gate

            src_text = prepare_document_text_for_llm(source_text)
            src_text, was_capped = cap_translation_source_text(src_text)
            translator = getattr(orchestrator, "_translate_text_complete", None) or translate_text_complete
            result = translator(
                src_text,
                source_lang="auto",
                target_lang="繁體中文",
                heavy=_web_heavy_opt_in(instruction),
            )
            if not isinstance(result, dict) or not result.get("success"):
                err = str((result or {}).get("error") if isinstance(result, dict) else "translate_failed")
                return {"task": task, "reply": f"❌ 檔案翻譯失敗：{err[:260]}", "artifacts": []}
            translated = str(result.get("translated_text") or result.get("text") or "").strip()
            translated = polish_translated_document_text(translated) or translated
            gate = run_output_quality_gate(
                "translation",
                translated,
                source_chars=len(src_text or ""),
                source_text=src_text,
                source_name=original_name,
                instruction=instruction,
            )
            if not gate.get("ok"):
                return {
                    "task": task,
                    "reply": str(gate.get("message") or "❌ 檔案翻譯品質檢查未通過。"),
                    "artifacts": [],
                }
            notes = []
            if was_capped:
                notes.append("⚠️ 原文超過翻譯上限，本次僅處理上限內文字；請改用分段翻譯以取得全文。")
            failed = int(result.get("chunks_failed") or 0)
            if failed:
                notes.append(f"⚠️ 有 {failed} 個段落翻譯失敗，已保留原文位置供重跑。")
            term_glossary = str(result.get("term_glossary") or "").strip()
            if term_glossary:
                notes.append(term_glossary)
            bilingual_artifact, bilingual_reply = _create_bilingual_translation_artifact(
                orchestrator,
                source_text=src_text,
                translated_text=translated,
                result=result,
                title=original_name or "檔案翻譯",
                user_id=user_id,
            )
            if bilingual_artifact:
                notes.insert(0, "📄 已產出原文/翻譯雙語對照 Word 表格。")
            elif bilingual_reply:
                notes.append("⚠️ 雙語對照表格產生失敗，已改輸出純文字翻譯。")
            reply = "\n".join(notes + [translated]).strip()
            artifacts = []
            if bilingual_artifact:
                artifacts.append(bilingual_artifact)
            artifacts.extend(_create_web_delivery_artifacts(
                root,
                instruction=instruction,
                reply=reply,
                source_filename=original_name,
                include_docx=not bool(bilingual_artifact),
            ))
            return {"task": task, "reply": reply, "artifacts": artifacts}
        except Exception as exc:
            return {"task": task, "reply": f"❌ 檔案翻譯失敗：{str(exc)[:260]}", "artifacts": []}

    return None


def _clean_artifact_filename(value: str, fallback: str = "magi") -> str:
    text = str(value or "").strip()
    text = re.sub(r"[\\/:*?\"<>|]+", "_", text)
    text = re.sub(r"\s+", "_", text).strip("._ ")
    return text[:80] or fallback


def _artifact_size_label(size: int) -> str:
    value = float(max(0, int(size or 0)))
    for unit in ("B", "KB", "MB", "GB"):
        if value < 1024 or unit == "GB":
            return f"{value:.1f} {unit}" if unit != "B" else f"{int(value)} B"
        value /= 1024
    return f"{int(size)} B"


def _artifact_dict(path: Path, *, label: str, fmt: str) -> dict[str, Any]:
    try:
        size = path.stat().st_size
    except OSError:
        size = 0
    return {
        "label": label,
        "format": fmt,
        "filename": path.name,
        "path": str(path),
        "share_path": str(path),
        "download_url": _artifact_download_url(path),
        "preview_url": _artifact_download_url(path, inline=True),
        "size": size,
        "size_label": _artifact_size_label(size),
    }


def _path_from_export_reply(reply: str) -> Path | None:
    text = str(reply or "").strip()
    if not text:
        return None
    if "|||FILE_PATH|||" in text:
        candidate = text.split("|||FILE_PATH|||", 1)[1].strip().splitlines()[0].strip()
        p = Path(candidate)
        return p if p.exists() else None
    for line in reversed(text.splitlines()):
        candidate = line.strip()
        if not candidate or candidate.startswith("http"):
            continue
        p = Path(candidate)
        if p.exists():
            return p
    return None


def _create_bilingual_translation_artifact(
    orchestrator: Any,
    *,
    source_text: str,
    translated_text: str,
    result: dict[str, Any],
    title: str,
    user_id: str,
) -> tuple[dict[str, Any] | None, str]:
    export_reply = ""
    try:
        exporter = getattr(orchestrator, "_export_translation_docx", None)
        if not callable(exporter):
            from api.handlers.document_handler import export_translation_docx as exporter
        export_reply = str(
            exporter(
                source_text=source_text,
                translated_text=translated_text,
                source_chunks=result.get("source_chunks") or [],
                translated_chunks=result.get("translated_chunks") or [],
                term_glossary=str(result.get("term_glossary") or ""),
                title=title,
                subtitle="MAGI 原文/翻譯對照表",
                prefix="file_translate",
                user_id=user_id,
            )
            or ""
        ).strip()
    except Exception:
        return None, ""
    path = _path_from_export_reply(export_reply)
    if not path:
        return None, export_reply
    return _artifact_dict(path, label="翻譯雙語對照 Word", fmt="docx"), export_reply


def _reply_to_docx(path: Path, *, title: str, reply: str, instruction: str, source_filename: str = "") -> bool:
    try:
        from docx import Document
    except Exception:
        return False

    doc = Document()
    doc.add_heading(title, 0)
    meta = doc.add_paragraph()
    meta.add_run("產生時間：").bold = True
    meta.add_run(datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    if source_filename:
        p = doc.add_paragraph()
        p.add_run("來源檔案：").bold = True
        p.add_run(source_filename)
    if instruction:
        p = doc.add_paragraph()
        p.add_run("使用者指示：").bold = True
        p.add_run(instruction)
    doc.add_paragraph("")
    doc.add_heading("內容", level=1)

    for raw in str(reply or "").replace("\r\n", "\n").replace("\r", "\n").split("\n"):
        line = raw.strip()
        if not line:
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            level = min(3, max(1, len(heading.group(1))))
            doc.add_heading(heading.group(2).strip("*# "), level=level)
            continue
        bullet = re.match(r"^[-*•]\s+(.+)$", line)
        if bullet:
            doc.add_paragraph(bullet.group(1), style="List Bullet")
            continue
        ordered = re.match(r"^\d+[.)、]\s+(.+)$", line)
        if ordered:
            doc.add_paragraph(ordered.group(1), style="List Number")
            continue
        doc.add_paragraph(line.strip("*"))

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return True


def _create_web_delivery_artifacts(
    magi_root: Path,
    *,
    instruction: str,
    reply: str,
    source_filename: str = "",
    include_docx: bool = True,
) -> list[dict[str, Any]]:
    if not _should_create_web_delivery_artifacts(instruction, reply, source_filename=source_filename):
        return []

    kind, title_label = _web_delivery_kind(instruction)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    source_stem = _clean_artifact_filename(Path(source_filename).stem if source_filename else "", "web")
    base_name = f"{kind}_{stamp}_{source_stem}_{uuid.uuid4().hex[:8]}"
    output_dir = _magi_web_outputs_dir(magi_root)
    title = f"MAGI {title_label}"
    generated_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    body = str(reply or "").strip()
    instruction_text = str(instruction or "").strip()
    source_line = f"\n- 來源檔案：{source_filename}" if source_filename else ""
    md_text = (
        f"# {title}\n\n"
        f"- 產生時間：{generated_at}{source_line}\n"
        f"- 使用者指示：{instruction_text or '未提供'}\n\n"
        f"## 內容\n\n{body}\n"
    )

    artifacts: list[dict[str, Any]] = []
    md_path = output_dir / f"{base_name}.md"
    md_path.write_text(md_text, encoding="utf-8")
    artifacts.append(_artifact_dict(md_path, label=f"{title_label} Markdown", fmt="md"))

    txt_path = output_dir / f"{base_name}.txt"
    txt_path.write_text(body + "\n", encoding="utf-8")
    artifacts.append(_artifact_dict(txt_path, label=f"{title_label} 純文字", fmt="txt"))

    if include_docx:
        docx_path = output_dir / f"{base_name}.docx"
        if _reply_to_docx(docx_path, title=title, reply=body, instruction=instruction_text, source_filename=source_filename):
            artifacts.insert(0, _artifact_dict(docx_path, label=f"{title_label} Word", fmt="docx"))

    return artifacts


def _format_web_inline(text: str) -> str:
    """Render a small, safe Markdown subset for MAGI web chat replies."""
    raw = str(text or "")
    pieces: list[str] = []
    pos = 0
    link_re = re.compile(r"\[([^\]]{1,180})\]\(([^)\s]{1,600})\)")

    def _format_without_links(chunk: str) -> str:
        escaped = html.escape(chunk)
        escaped = re.sub(r"`([^`]{1,160})`", r"<code>\1</code>", escaped)
        escaped = re.sub(r"\*\*([^*]{1,220})\*\*", r"<strong>\1</strong>", escaped)
        escaped = re.sub(r"__([^_]{1,220})__", r"<strong>\1</strong>", escaped)
        return escaped

    for match in link_re.finditer(raw):
        pieces.append(_format_without_links(raw[pos:match.start()]))
        label = _format_without_links(match.group(1))
        url = _safe_web_url(match.group(2))
        if url:
            pieces.append(
                f'<a href="{html.escape(url, quote=True)}" target="_blank" rel="noopener noreferrer">{label}</a>'
            )
        else:
            pieces.append(label)
        pos = match.end()
    pieces.append(_format_without_links(raw[pos:]))
    return "".join(pieces)


def format_web_reply_html(reply: str) -> str:
    """
    Convert Discord/Telegram-style Markdown into readable, safe HTML for the web UI.

    The messaging platforms still receive the original text; this is only a display
    layer for /golem and other browser chat surfaces.
    """
    text = str(reply or "").replace("\r\n", "\n").replace("\r", "\n").strip()
    if not text:
        return '<div class="web-reply"><p>MAGI 沒有回傳內容。</p></div>'

    blocks: list[str] = []
    list_type: str | None = None
    in_code = False
    code_lines: list[str] = []

    def close_list() -> None:
        nonlocal list_type
        if list_type:
            blocks.append(f"</{list_type}>")
            list_type = None

    def open_list(kind: str) -> None:
        nonlocal list_type
        if list_type != kind:
            close_list()
            blocks.append(f"<{kind}>")
            list_type = kind

    for raw_line in text.split("\n"):
        line = raw_line.strip()
        if line.startswith("```"):
            if in_code:
                blocks.append(f"<pre><code>{html.escape('\\n'.join(code_lines))}</code></pre>")
                code_lines = []
                in_code = False
            else:
                close_list()
                in_code = True
                code_lines = []
            continue
        if in_code:
            code_lines.append(raw_line)
            continue
        if not line:
            close_list()
            continue
        if re.fullmatch(r"[━─=\-_*]{4,}", line):
            close_list()
            blocks.append("<hr>")
            continue
        if re.fullmatch(r"#{2,6}", line):
            close_list()
            blocks.append("<hr>")
            continue
        heading_line = line
        wrapped_heading = re.fullmatch(r"\*\*(#{1,6}\s*[^*]+?)\*\*", heading_line)
        if wrapped_heading:
            heading_line = wrapped_heading.group(1).strip()
        heading_line = re.sub(r"\*\*$", "", heading_line).strip()
        heading = re.match(r"^(#{1,6})\s*(.+)$", heading_line)
        if heading:
            close_list()
            title = heading.group(2).strip("# ").strip()
            if title:
                level = 3 if len(heading.group(1)) == 1 else 4
                blocks.append(f"<h{level}>{_format_web_inline(title)}</h{level}>")
                continue
        bold_heading = re.fullmatch(r"(?:[^\w\u4e00-\u9fff]{0,3}\s*)?\*\*([^*]{1,80})\*\*", line)
        if bold_heading:
            close_list()
            blocks.append(f"<h4>{_format_web_inline(bold_heading.group(1))}</h4>")
            continue
        unordered = re.match(r"^[-*•]\s+(.+)$", line)
        if unordered:
            open_list("ul")
            blocks.append(f"<li>{_format_web_inline(unordered.group(1))}</li>")
            continue
        ordered = re.match(r"^\d+[.)、]\s+(.+)$", line)
        if ordered:
            open_list("ol")
            blocks.append(f"<li>{_format_web_inline(ordered.group(1))}</li>")
            continue
        close_list()
        blocks.append(f"<p>{_format_web_inline(line)}</p>")

    if in_code:
        blocks.append(f"<pre><code>{html.escape('\\n'.join(code_lines))}</code></pre>")
    close_list()
    return f'<div class="web-reply">{"".join(blocks)}</div>'


def _collect_process_monitor(
    *,
    process_monitor_state_path: Path,
    magi_root: Path,
) -> dict[str, Any]:
    rows: list[dict[str, Any]] = []
    try:
        out = subprocess.run(
            ["ps", "-axo", "pid=,ppid=,etime=,command="],
            stdout=subprocess.PIPE,
            stderr=subprocess.DEVNULL,
            text=True,
            timeout=8,
        ).stdout or ""
        for raw in out.splitlines():
            line = (raw or "").strip()
            if not line:
                continue
            parts = line.split(None, 3)
            if len(parts) < 4:
                continue
            try:
                pid = int(parts[0])
                ppid = int(parts[1])
            except Exception:
                continue
            rows.append(
                {
                    "pid": pid,
                    "ppid": ppid,
                    "age_sec": _parse_etime_to_sec(parts[2]),
                    "age": parts[2],
                    "cmd": parts[3],
                }
            )
    except Exception as exc:
        return {
            "ok": False,
            "error": str(exc),
            "summary": {},
            "core": [],
            "workers": [],
            "orphans": [],
            "duplicates": [],
        }

    worker_markers, core_markers, core_labels = _process_monitor_markers(magi_root)
    core = []
    workers = []
    orphans = []
    grouped: dict[str, list[dict[str, Any]]] = defaultdict(list)

    for row in rows:
        cmd = str(row.get("cmd") or "")
        label = None
        for marker in core_markers:
            if marker in cmd:
                label = core_labels.get(marker, marker)
                break
        if label:
            entry = dict(row)
            entry["label"] = label
            core.append(entry)
        is_worker = any(marker in cmd for marker in worker_markers)
        if is_worker:
            workers.append(row)
            grouped[cmd].append(row)
            if int(row.get("ppid") or 0) == 1:
                orphans.append(row)

    duplicates = []
    for cmd, items in grouped.items():
        if len(items) <= 1:
            continue
        duplicates.append(
            {
                "count": len(items),
                "pids": [int(item["pid"]) for item in items],
                "cmd": cmd[:320],
            }
        )

    guardian_state: dict[str, Any] = {}
    try:
        if process_monitor_state_path.exists():
            guardian_state = json.loads(process_monitor_state_path.read_text(encoding="utf-8")) or {}
    except Exception:
        guardian_state = {}

    return {
        "ok": True,
        "ts": datetime.now().isoformat(timespec="seconds"),
        "summary": {
            "core_count": len(core),
            "worker_count": len(workers),
            "orphan_count": len(orphans),
            "duplicate_groups": len(duplicates),
        },
        "core": sorted(core, key=lambda x: (x.get("label", ""), x.get("pid", 0))),
        "workers": sorted(workers, key=lambda x: (x.get("age_sec", 0), x.get("pid", 0)), reverse=True),
        "orphans": sorted(orphans, key=lambda x: (x.get("age_sec", 0), x.get("pid", 0)), reverse=True),
        "duplicates": sorted(duplicates, key=lambda x: x.get("count", 0), reverse=True),
        "guardian_state": guardian_state,
    }


_MAGI_MODULE_COMMANDS = {
    "laf": {
        "label": "法扶",
        "commands": ("法扶指令", "二階段批次", "批次二階段", "二階段掃描", "報結掃描", "自動報結掃描", "批次報結", "結案掃描"),
    },
    "file_review": {
        "label": "閱卷",
        "commands": ("檢查閱卷信箱", "閱卷信箱", "閱卷郵件", "下載閱卷", "閱卷下載", "閱卷到期檢查", "閱卷到期", "閱卷期限"),
    },
    "transcript": {
        "label": "筆錄",
        "commands": ("下載筆錄", "筆錄下載", "調閱筆錄", "筆錄調閱", "筆錄同步", "同步筆錄", "筆錄全同步", "筆錄更名", "更名筆錄"),
    },
}


def _allowed_magi_module_command(module_key: str, command: str) -> tuple[bool, str]:
    meta = _MAGI_MODULE_COMMANDS.get(module_key)
    if not meta:
        return False, "unknown_module"
    text = str(command or "").strip()
    if not text:
        return False, "empty_command"
    for prefix in meta["commands"]:
        if text == prefix or text.startswith(prefix + " ") or text.startswith(prefix + "　"):
            return True, ""
    return False, "unsupported_command"


def _magi_module_runs_in_background(module_key: str, command: str) -> bool:
    text = str(command or "").strip()
    if module_key == "file_review":
        return True
    if module_key == "transcript":
        return True
    if module_key == "laf" and text != "法扶指令":
        return True
    return False


def create_web_runtime_blueprint(
    *,
    orchestrator: Any,
    logger: Any,
    web_notifications: dict[str, list[Any]],
    normalize_output_text=None,
    magi_root: str | Optional[Path] = None,
) -> Blueprint:
    bp = Blueprint("web_runtime", __name__)
    isolated_runtime = magi_root is not None
    root = Path(magi_root) if magi_root else Path(__file__).resolve().parents[2]
    agent_dir = root / ".agent"
    process_monitor_state_path = root / "static" / "process_guardian_state.json"
    guardian_control_path = root / "static" / "guardian_control.json"

    @bp.route("/ops/process-monitor")
    @login_required
    def process_monitor_page():
        return render_template("process_monitor.html", user=current_user)

    @bp.route("/api/memory/stats", methods=["GET"])
    @login_required
    def api_memory_stats():
        stats: dict[str, Any] = {"doc_count": 0, "source_count": 0, "last_ingest": None, "obsidian": {}, "faiss_size": 0}
        # Primary: get real document count from MariaDB
        if not isolated_runtime:
            try:
                from skills.memory.mem_bridge import _get_conn
                _conn = _get_conn()
                _cur = _conn.cursor()
                _cur.execute("SELECT COUNT(*) FROM documents")
                stats["doc_count"] = _cur.fetchone()[0]
                _cur.execute("SELECT COUNT(DISTINCT source) FROM documents")
                stats["source_count"] = _cur.fetchone()[0]
                _cur.execute("SELECT MAX(created_at) FROM documents")
                _last = _cur.fetchone()[0]
                if _last:
                    stats["last_ingest"] = str(_last)
                _conn.close()
            except Exception:
                pass
        if stats["doc_count"] == 0 and stats["last_ingest"] is None:
            # Fallback: read from doc_vector_index.json (attachment tracker only)
            try:
                idx_path = agent_dir / "doc_vector_index.json"
                if idx_path.exists():
                    idx = json.loads(idx_path.read_text(encoding="utf-8"))
                    entries = idx if isinstance(idx, list) else list(idx.values()) if isinstance(idx, dict) else []
                    stats["doc_count"] = len(entries)
                    dates = [entry.get("updated_at") or entry.get("created_at", "") for entry in entries if isinstance(entry, dict)]
                    dates = sorted([item for item in dates if item], reverse=True)
                    if dates:
                        stats["last_ingest"] = dates[0]
            except Exception as exc:
                stats["doc_index_error"] = str(exc)
        try:
            obs_cfg = agent_dir / "obsidian_vault_config.json"
            obs_idx = agent_dir / "obsidian_index.json"
            if obs_cfg.exists():
                cfg = json.loads(obs_cfg.read_text(encoding="utf-8"))
                stats["obsidian"]["vault_path"] = cfg.get("vault_path", "")
                stats["obsidian"]["vault_name"] = cfg.get("vault_name", "")
            if obs_idx.exists():
                oidx = json.loads(obs_idx.read_text(encoding="utf-8"))
                stats["obsidian"]["notes_indexed"] = len((oidx.get("notes") or {}))
                stats["obsidian"]["last_update"] = oidx.get("updated_at", "")
        except Exception as exc:
            stats["obsidian_error"] = str(exc)
        faiss_path = root / "skills" / "memory" / "index_cache" / "mem_index.faiss"
        if faiss_path.exists():
            stats["faiss_size"] = faiss_path.stat().st_size
        if not isolated_runtime:
            try:
                from skills.memory.mem_bridge import _get_faiss_index
                idx = _get_faiss_index()
                if idx:
                    stats["faiss_vector_count"] = idx.total
                    stats["faiss_index_type"] = idx.index_type
            except Exception:
                logger.debug("silent-catch in api_memory_stats index-sync", exc_info=True)
            try:
                from skills.memory.faiss_index import FAISSMemoryIndex
                idx = FAISSMemoryIndex.get_instance()
                stats["faiss_vector_count"] = idx.total
                stats["faiss_index_type"] = idx.index_type
            except Exception:
                pass
        return jsonify(stats)

    @bp.route("/api/memory/recall", methods=["POST"])
    @login_required
    def api_memory_recall():
        data = request.get_json(silent=True) or {}
        query = str(data.get("query", "")).strip()
        try:
            top_k = min(20, max(1, int(data.get("top_k", 5))))
        except (ValueError, TypeError):
            top_k = 5
        source_filter = str(data.get("source", "")).strip() or None
        if not query:
            return jsonify({"error": "請輸入搜尋關鍵字"}), 400
        try:
            from skills.memory.mem_bridge import recall

            results = recall(query, top_k=top_k, source_contains=source_filter)
            return jsonify({"memories": results or [], "query": query})
        except Exception as exc:
            logger.error("Memory recall error: %s", exc)
            return jsonify({"error": str(exc)}), 500

    @bp.route("/api/memory/remember", methods=["POST"])
    @login_required
    def api_memory_remember():
        data = request.get_json(silent=True) or {}
        content = str(data.get("content", "")).strip()
        source = str(data.get("source", "dashboard-manual")).strip() or "dashboard-manual"
        if not content:
            return jsonify({"error": "請輸入要記憶的內容"}), 400
        if len(content) > 50000:
            return jsonify({"error": "內容過長（上限 50,000 字元）"}), 400
        try:
            from skills.memory.mem_bridge import remember

            remember(content, source)
            return jsonify({"success": True, "message": f"已儲存 {len(content)} 字元至向量記憶庫"})
        except Exception as exc:
            logger.error("Memory remember error: %s", exc)
            return jsonify({"error": str(exc)}), 500

    @bp.route("/api/memory/obsidian-sync", methods=["POST"])
    @login_required
    def api_memory_obsidian_sync():
        def _run_ingest():
            try:
                from skills.obsidian.action import task_ingest

                task_ingest({})
            except Exception as exc:
                logger.error("Obsidian ingest error: %s", exc)

        thread = threading.Thread(target=_run_ingest, daemon=True)
        thread.start()
        return jsonify({"success": True, "message": "Obsidian 重新索引已啟動（背景執行中）"})

    @bp.route("/api/ops/process-monitor", methods=["GET"])
    @login_required
    def process_monitor_api():
        data = _collect_process_monitor(
            process_monitor_state_path=process_monitor_state_path,
            magi_root=root,
        )
        ctrl_enabled = True
        if guardian_control_path.exists():
            try:
                ctrl_enabled = json.loads(guardian_control_path.read_text(encoding="utf-8")).get("enabled", True)
            except Exception:
                logger.debug("silent-catch in process_monitor_api", exc_info=True)
        data["guardian_control_enabled"] = ctrl_enabled
        return jsonify(data), 200 if data.get("ok") else 500

    @bp.route("/api/ops/process-guardian/toggle", methods=["POST"])
    @login_required
    def process_guardian_toggle_api():
        try:
            ctrl = {"enabled": True}
            if guardian_control_path.exists():
                ctrl = json.loads(guardian_control_path.read_text(encoding="utf-8"))
            ctrl["enabled"] = not ctrl.get("enabled", True)
            guardian_control_path.write_text(json.dumps(ctrl, ensure_ascii=False, indent=2), encoding="utf-8")
            return jsonify({"ok": True, "enabled": ctrl["enabled"]})
        except Exception as exc:
            return jsonify({"ok": False, "error": str(exc)}), 500

    @bp.route("/api/osc/chat", methods=["POST"])
    @login_required
    def osc_chat_api():
        data = request.get_json(silent=True) or {}
        msg = (data.get("message") or "").strip()
        if not msg:
            return jsonify({"error": "Empty message"}), 400
        reply = orchestrator.process_message(
            user_id=str(current_user.id),
            message=msg,
            platform="WEB",
            role=current_user.role,
        )
        try:
            if normalize_output_text:
                reply = normalize_output_text(str(reply or ""), platform="WEB")
        except Exception:
            logger.debug("silent-catch in osc_chat_api", exc_info=True)
        artifacts = _create_web_delivery_artifacts(root, instruction=msg, reply=str(reply or ""))
        return jsonify({"reply": reply, "reply_html": format_web_reply_html(str(reply or "")), "artifacts": artifacts})

    @bp.route("/api/osc/magi-modules/run", methods=["POST"])
    @login_required
    def osc_magi_modules_run_api():
        data = request.get_json(silent=True) or {}
        module_key = str(data.get("module") or "").strip()
        command = str(data.get("command") or "").strip()
        ok, reason = _allowed_magi_module_command(module_key, command)
        if not ok:
            meta = _MAGI_MODULE_COMMANDS.get(module_key)
            return jsonify(
                {
                    "ok": False,
                    "error": reason,
                    "message": "此頁籤只接受法扶、閱卷、筆錄三模組的正式命令。",
                    "module": module_key,
                    "allowed_commands": list(meta["commands"]) if meta else [],
                }
            ), 400

        meta = _MAGI_MODULE_COMMANDS[module_key]
        if _magi_module_runs_in_background(module_key, command):
            user_id = str(current_user.id)
            role = current_user.role

            def _run_background_module() -> None:
                try:
                    orchestrator.process_message(
                        user_id=user_id,
                        message=command,
                        platform="WEB",
                        role=role,
                    )
                except Exception:
                    logger.exception("MAGI module background command failed: %s %s", module_key, command[:80])

            thread = threading.Thread(target=_run_background_module, daemon=True)
            thread.start()
            return jsonify(
                {
                    "ok": True,
                    "background": True,
                    "module": module_key,
                    "module_label": meta["label"],
                    "command": command,
                    "reply": f"已啟動{meta['label']}模組：{command}\\n背景作業完成後會由 MAGI 通知或寫入對應紀錄。",
                }
            ), 202

        reply = orchestrator.process_message(
            user_id=str(current_user.id),
            message=command,
            platform="WEB",
            role=current_user.role,
        )
        try:
            if normalize_output_text:
                reply = normalize_output_text(str(reply or ""), platform="WEB")
        except Exception:
            logger.debug("silent-catch in osc_magi_modules_run_api", exc_info=True)
        return jsonify(
            {
                "ok": True,
                "module": module_key,
                "module_label": meta["label"],
                "command": command,
                "reply": reply,
                "reply_html": format_web_reply_html(str(reply or "")),
            }
        )

    @bp.route("/api/osc/chat/upload", methods=["POST"])
    @login_required
    def osc_chat_upload_api():
        msg = (request.form.get("message") or "").strip()
        file = request.files.get("file")
        if not file or not file.filename:
            return jsonify({"error": "請先選擇檔案"}), 400

        max_mb = int(os.environ.get("MAGI_WEB_CHAT_UPLOAD_MAX_MB", "200") or "200")
        content_length = int(request.content_length or 0)
        if max_mb > 0 and content_length > max_mb * 1024 * 1024:
            return jsonify({"error": f"檔案過大，請選擇 {max_mb}MB 以下的檔案"}), 413

        original_name = Path(file.filename).name
        safe_name = secure_filename(original_name) or f"upload_{uuid.uuid4().hex}"
        stored_name = f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:8]}_{safe_name}"
        target = _chat_upload_dir(root) / stored_name
        try:
            file.save(target)
        except Exception as exc:
            return jsonify({"error": f"檔案儲存失敗：{exc}"}), 500

        upload_task = _web_upload_requested_task(msg, original_name)
        suffix = Path(original_name or "").suffix.lower()
        if upload_task == "transcript" and suffix in (_CHAT_UPLOAD_AUDIO_EXTS | _CHAT_UPLOAD_VIDEO_EXTS):
            user_instruction = msg or "請轉成逐字稿，保留時間戳記，並輸出可下載文字檔。"
            reply = orchestrator.process_message(
                user_id=str(current_user.id),
                message=user_instruction,
                platform="WEB",
                role=current_user.role,
                attachment={"type": "audio", "path": str(target), "filename": original_name},
            )
            try:
                if normalize_output_text:
                    reply = normalize_output_text(str(reply or ""), platform="WEB")
            except Exception:
                logger.debug("silent-catch in osc_chat_upload_api media", exc_info=True)
            try:
                from api.handlers.output_quality_handler import run_output_quality_gate

                gate = run_output_quality_gate(
                    "transcript",
                    str(reply or ""),
                    source_chars=0,
                    source_name=original_name,
                    instruction=user_instruction,
                )
                if not gate.get("ok"):
                    reply = str(gate.get("message") or "❌ 逐字稿品質檢查未通過。")
                    artifacts = []
                    return jsonify(
                        {
                            "reply": reply,
                            "reply_html": format_web_reply_html(str(reply or "")),
                            "artifacts": artifacts,
                            "filename": original_name,
                            "path": str(target),
                            "kind": "audio",
                            "chars": 0,
                            "truncated": False,
                            "task": upload_task,
                        }
                    )
            except Exception:
                logger.debug("silent-catch in osc_chat_upload_api transcript gate", exc_info=True)
            artifacts = _create_web_delivery_artifacts(
                root,
                instruction=user_instruction,
                reply=str(reply or ""),
                source_filename=original_name,
            )
            return jsonify(
                {
                    "reply": reply,
                    "reply_html": format_web_reply_html(str(reply or "")),
                    "artifacts": artifacts,
                    "filename": original_name,
                    "path": str(target),
                    "kind": "audio",
                    "chars": 0,
                    "truncated": False,
                    "task": upload_task,
                }
            )

        extracted = _extract_chat_upload_text_for_task(target, original_name, upload_task)
        if not extracted.get("success"):
            return jsonify(
                {
                    "error": f"檔案已接收，但無法讀取內容：{extracted.get('error') or 'extract_failed'}",
                    "filename": original_name,
                    "path": str(target),
                }
            ), 422

        text = str(extracted.get("text") or "").strip()
        user_instruction = msg or "請摘要這份檔案，必要時整理成可翻譯或分析的重點。"
        direct = _run_direct_web_upload_text_task(
            orchestrator,
            root=root,
            target=target,
            original_name=original_name,
            instruction=user_instruction,
            extracted=extracted,
            user_id=str(current_user.id),
        )
        if direct is not None:
            reply = str(direct.get("reply") or "")
            try:
                if normalize_output_text:
                    reply = normalize_output_text(reply, platform="WEB")
            except Exception:
                logger.debug("silent-catch in osc_chat_upload_api direct", exc_info=True)
            return jsonify(
                {
                    "reply": reply,
                    "reply_html": format_web_reply_html(reply),
                    "artifacts": direct.get("artifacts") or [],
                    "filename": original_name,
                    "path": str(target),
                    "kind": extracted.get("kind") or "",
                    "chars": len(text),
                    "truncated": False,
                    "task": direct.get("task") or upload_task,
                }
            )

        max_chars = int(os.environ.get("MAGI_WEB_CHAT_UPLOAD_TEXT_MAX_CHARS", "120000") or "120000")
        truncated = False
        if max_chars > 0 and len(text) > max_chars:
            text = text[:max_chars]
            truncated = True

        prompt = (
            f"{user_instruction}\n\n"
            f"[上傳檔案]\n"
            f"檔名：{original_name}\n"
            f"類型：{extracted.get('kind') or 'file'}\n"
            f"儲存位置：{target}\n"
            f"內容{'（因長度限制已截斷）' if truncated else ''}：\n"
            f"{text}"
        )
        reply = orchestrator.process_message(
            user_id=str(current_user.id),
            message=prompt,
            platform="WEB",
            role=current_user.role,
        )
        try:
            if normalize_output_text:
                reply = normalize_output_text(str(reply or ""), platform="WEB")
        except Exception:
            logger.debug("silent-catch in osc_chat_upload_api", exc_info=True)
        artifacts = _create_web_delivery_artifacts(
            root,
            instruction=user_instruction,
            reply=str(reply or ""),
            source_filename=original_name,
        )
        return jsonify(
            {
                "reply": reply,
                "reply_html": format_web_reply_html(str(reply or "")),
                "artifacts": artifacts,
                "filename": original_name,
                "path": str(target),
                "kind": extracted.get("kind") or "",
                "chars": len(text),
                "truncated": truncated,
            }
        )

    @bp.route("/api/osc/poll", methods=["GET"])
    @login_required
    def osc_poll_api():
        uid = str(current_user.id)
        messages = []
        if uid in web_notifications:
            messages = list(web_notifications[uid])
            web_notifications[uid].clear()
        return jsonify({"messages": messages})

    @bp.route("/api/osc/judgments_legacy", methods=["GET"])
    @login_required
    def osc_judgments_api():
        try:
            json_path = root / "skills" / "judgment-collector" / "judgments.json"
            if json_path.exists():
                return jsonify(json.loads(json_path.read_text(encoding="utf-8")))
            return jsonify([])
        except Exception as exc:
            logger.error("Error serving judgments: %s", exc)
            return jsonify([])

    return bp
