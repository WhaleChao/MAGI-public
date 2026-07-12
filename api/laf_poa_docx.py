from __future__ import annotations

import datetime as _dt
import html
import os
import re
import tempfile
import zipfile
from pathlib import Path
from typing import Any

from api.laf_branch_profiles import (
    DEFAULT_LAWYER_NAME,
    PUBLIC_PLACEHOLDERS,
    get_law_firm_profile,
    normalize_branch_label,
    resolve_laf_branch_profile,
)


FALSE_VALUES = {"0", "false", "no", "off", "disabled"}
TEMPLATE_DIR = Path(__file__).resolve().parents[1] / "templates" / "laf_poa"
TEMPLATE_FILENAMES = {
    "general": "general.docx",
    "indigenous_center": "indigenous_center.docx",
}
NON_POA_FILENAME_MARKERS = ("回執", "收件", "郵件", "掛號", "委任狀卷")


def laf_poa_docx_enabled() -> bool:
    return os.environ.get("MAGI_LAF_POA_DOCX", "1").strip().lower() not in FALSE_VALUES


def laf_poa_docx_templates_enabled() -> bool:
    return os.environ.get("MAGI_LAF_POA_DOCX_TEMPLATES", "1").strip().lower() not in FALSE_VALUES


def laf_poa_static_templates_enabled() -> bool:
    # Keep curated Word templates available as an explicit compatibility mode
    # when exact PDF-backed layout is disabled.
    return os.environ.get("MAGI_LAF_POA_STATIC_TEMPLATES", "1").strip().lower() not in FALSE_VALUES


def laf_poa_pdf_render_fallback_enabled() -> bool:
    return os.environ.get("MAGI_LAF_POA_ALLOW_PDF_RENDER_FALLBACK", "0").strip().lower() not in FALSE_VALUES


def laf_poa_exact_pdf_layout_enabled() -> bool:
    # The default companion should visually match the official PDF.  Editable
    # Word controls are overlaid only where MAGI has reliable values.
    return os.environ.get("MAGI_LAF_POA_EXACT_PDF_LAYOUT", "1").strip().lower() not in FALSE_VALUES


def is_laf_power_of_attorney_pdf(path: str | os.PathLike[str]) -> bool:
    p = Path(path)
    if p.suffix.lower() != ".pdf" or "委任狀" not in p.name or p.name.startswith("._"):
        return False
    return not any(marker in p.stem for marker in NON_POA_FILENAME_MARKERS)


def laf_poa_docx_path(pdf_path: str | os.PathLike[str]) -> Path:
    p = Path(pdf_path)
    return p.with_name(f"{p.stem}（可填寫版）.docx")


def laf_poa_template_docx_path(pdf_path: str | os.PathLike[str]) -> Path:
    p = Path(pdf_path)
    return p.with_name(f"{p.stem}（範本）.docx")


def laf_poa_case_docx_path(pdf_path: str | os.PathLike[str]) -> Path:
    return laf_poa_docx_path(pdf_path)


def _text(value: Any) -> str:
    return str(value or "").strip()


VALID_EAST_ASIA_FONTS = {
    "標楷體",
    "微軟正黑體",
    "新細明體",
    "PMingLiU",
    "MingLiU",
    "MingLiU-ExtB",
    "Noto Sans TC",
    "Noto Sans CJK TC",
    "Noto Sans CJK TC Bold",
    "Microsoft JhengHei",
    "Microsoft JhengHei UI",
}
LAF_POA_QUALITY_PLACEHOLDERS = (
    "請填律師姓名",
    "請填事務所地址",
    "請填電話",
    "請填法院",
    "請填法院/檢署/委員會",
    "請填法院案號",
    "請填案由",
    "請填當事人",
)


def _collect_docx_xml_text_and_fonts(docx_path: str | os.PathLike[str]) -> tuple[str, set[str]]:
    """Read document and related XML payload plus declared East Asia fonts."""
    with zipfile.ZipFile(docx_path) as zf:
        text = zf.read("word/document.xml").decode("utf-8", "ignore")
        fonts: set[str] = set(re.findall(r'w:eastAsia="([^"]+)"', text))
        for name in zf.namelist():
            if not name.endswith(".xml") or name == "word/document.xml":
                continue
            xml = zf.read(name).decode("utf-8", "ignore")
            fonts.update(re.findall(r'w:eastAsia="([^"]+)"', xml))
        return text, fonts


def _docx_plain_text(xml: str) -> str:
    return re.sub(r"<[^>]+>", "", xml)


def _normalize_quality_value(value: str) -> str:
    text = _text(value).replace(" ", "")
    if not text or text in PUBLIC_PLACEHOLDERS:
        return ""
    return text


def validate_laf_poa_docx_quality(
    docx_path: str | os.PathLike[str],
    *,
    expected_fields: dict[str, str] | None = None,
    check_expected_values: bool = True,
) -> dict[str, Any]:
    """Validate output quality for generated 委任狀 DOCX.

    Returns:
      {"ok": bool, "issues": [ ... ]}.
    """

    path = Path(docx_path)
    if not path.exists():
        return {"ok": False, "path": str(path), "issues": [{"code": "missing_file", "field": "docx_path", "message": "docx 文件不存在"}]}

    try:
        document_xml, fonts = _collect_docx_xml_text_and_fonts(path)
    except Exception as exc:
        return {
            "ok": False,
            "path": str(path),
            "issues": [{"code": "read_failed", "field": "docx_path", "message": f"讀取 docx 失敗: {exc!s}"}],
        }

    issues: list[dict[str, str]] = []
    plain = _docx_plain_text(document_xml)
    expected = expected_fields or {}
    if not fonts or not (set(fonts) & VALID_EAST_ASIA_FONTS):
        issues.append({
            "code": "font_violation",
            "field": "east_asia_font",
            "message": "缺少可接受的 East Asia 字型（建議標楷體）",
        })

    public_placeholder_tokens = set(PUBLIC_PLACEHOLDERS) - {"事務所地址"}
    for token in set(LAF_POA_QUALITY_PLACEHOLDERS) | public_placeholder_tokens:
        if token in plain:
            issues.append({"code": "placeholder_leftover", "field": "template_placeholders", "message": f"發現未填欄位留白：{token}"})

    if check_expected_values:
        checks: list[tuple[str, str, str, bool]] = [
            ("client_name", "CLIENT_NAME", "當事人", True),
            ("case_reason", "CASE_REASON", "案由", True),
            ("lawyer_name", "LAWYER_NAME", "律師姓名", True),
            ("lawyer_address", "LAW_FIRM_ADDRESS_LINE", "律師地址", True),
            ("lawyer_phone", "LAW_FIRM_PHONE", "律師電話", True),
            ("court", "COURT_NAME", "法院", False),
        ]
        for key, payload_key, field_label, required in checks:
            value = _normalize_quality_value(expected.get(key, "") or expected.get(payload_key, ""))
            if not value:
                if required:
                    issues.append({"code": f"missing_field", "field": key, "message": f"缺少{field_label}資料"})
                continue
            if value not in plain:
                issues.append({"code": "value_mismatch", "field": key, "message": f"{field_label}與文件內容不符", "expected": value})

    if "{{" in document_xml or "}}" in document_xml:
        issues.append({"code": "unreplaced_placeholder", "field": "template_placeholders", "message": "發現未清除的 template placeholder"})

    return {"ok": not issues, "path": str(path), "issues": issues}


def _path_needs_timeout(path: str | os.PathLike[str]) -> bool:
    p = str(path or "")
    return (
        p.startswith("/Volumes/")
        or "/.magi_mounts/" in p
        or "/Library/CloudStorage/" in p
        or "/SynologyDrive/" in p
    )


def _stat_quick(path: str | os.PathLike[str], timeout: float = 1.5):
    p = str(path or "")
    if not _path_needs_timeout(p):
        return Path(p).stat()
    import threading

    result: dict[str, Any] = {"stat": None, "error": None}

    def _run() -> None:
        try:
            result["stat"] = Path(p).stat()
        except Exception as exc:
            result["error"] = exc

    thread = threading.Thread(target=_run, daemon=True)
    thread.start()
    thread.join(timeout=timeout)
    if thread.is_alive():
        raise TimeoutError(f"stat timeout: {p}")
    if result["error"] is not None:
        raise result["error"]
    return result["stat"]


def _exists_quick(path: str | os.PathLike[str]) -> bool:
    try:
        _stat_quick(path)
        return True
    except Exception:
        return False


def _is_exact_pdf_backed_docx(path: str | os.PathLike[str]) -> bool:
    try:
        with zipfile.ZipFile(path) as zf:
            xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
    except Exception:
        return False
    return "<wp:anchor" in xml and 'behindDoc="1"' in xml and "magi_laf_poa_overlay_" in xml


def _remove_stale_laf_poa_template(path: Path) -> None:
    try:
        if path.exists():
            path.unlink()
    except OSError:
        pass


def _metadata_first(data: dict[str, Any], *keys: str) -> str:
    for key in keys:
        value = _text(data.get(key))
        if value:
            return value
    return ""


def normalize_laf_branch_label(branch: str) -> str:
    return normalize_branch_label(branch)


def _laf_branch_phone(data: dict[str, Any], branch_label: str) -> str:
    explicit = _metadata_first(
        data,
        "branch_phone",
        "laf_branch_phone",
        "office_phone",
        "division_phone",
        "phone",
    )
    if explicit:
        return explicit
    profile = resolve_laf_branch_profile(branch_label)
    return profile.phone if profile and profile.phone else "待確認"


def _default_lawyer_name(case_type: str, case_reason: str) -> str:
    return DEFAULT_LAWYER_NAME


def _normalize_lawyer_name(value: str, case_type: str, case_reason: str) -> str:
    # 法扶委任狀受任人固定使用法扶委任狀設定；不得受案件承辦律師欄位干擾。
    text = _text(value)
    if text and text not in {"受任律師", "事務所名稱"}:
        if "律師" not in text and re.fullmatch(r"[\u4e00-\u9fff]{2,4}", text):
            text = f"{text}律師"
        return text
    return get_law_firm_profile().lawyer_name or DEFAULT_LAWYER_NAME


def _roc_date_parts(data: dict[str, Any]) -> tuple[str, str, str]:
    explicit_roc = (
        _metadata_first(data, "roc_year", "roc_y"),
        _metadata_first(data, "roc_month", "roc_m"),
        _metadata_first(data, "roc_day", "roc_d"),
    )
    if all(explicit_roc):
        return explicit_roc

    raw = _metadata_first(data, "document_date", "poa_date", "date", "download_date")
    parsed: _dt.date | None = None
    if raw:
        for fmt in ("%Y-%m-%d", "%Y/%m/%d", "%Y%m%d"):
            try:
                parsed = _dt.datetime.strptime(raw, fmt).date()
                break
            except ValueError:
                pass
    if parsed is None:
        parsed = _dt.date.today()
    return str(parsed.year - 1911), str(parsed.month), str(parsed.day)


def _roc_date_from_filename(path: Path) -> dict[str, str]:
    # Official LAF files usually end with a ROC date, e.g.
    # 委任狀_1150421-W-004_1150529.pdf.
    matches = list(re.finditer(r"(?:^|_)(1\d{2})(\d{2})(\d{2})(?:\D|$)", path.stem))
    if not matches:
        return {}
    match = matches[-1]
    return {
        "roc_year": str(int(match.group(1))),
        "roc_month": str(int(match.group(2))),
        "roc_day": str(int(match.group(3))),
    }


def _laf_case_number_from_filename(path: Path) -> str:
    match = re.search(r"委任狀[_-]([0-9]{6,7}-[A-Z]-\d{3})(?:[_-]|\D|$)", path.stem)
    return match.group(1) if match else ""


def _case_folder_metadata_from_path(path: Path) -> dict[str, str]:
    """Infer stable case fields from MAGI/OSC case folder names.

    Official LAF POA PDFs intentionally leave the case reason blank.  The case
    folder is therefore a safe fallback for draft Word companions, while PDF
    fields still remain authoritative for party name, LAF number, branch, and
    branch phone.
    """

    known_types = {"刑事", "民事", "行政", "非訟", "消費者債務清理", "無償案件"}
    data: dict[str, str] = {}
    for parent in path.parents:
        name = parent.name
        match = re.match(r"^(20\d{2}-\d{4})-(.+?)-([^-]+)-(.+)$", name)
        if not match:
            continue
        data["case_number"] = match.group(1)
        data["client_name"] = match.group(2).strip()
        stage_or_type = match.group(3).strip()
        reason = match.group(4).strip()
        data["case_reason"] = reason
        if stage_or_type and stage_or_type != "消費者債務清理":
            data["stage"] = stage_or_type
        for ancestor in parent.parents:
            if ancestor.name in known_types:
                data["case_type"] = ancestor.name
                break
        if stage_or_type == "消費者債務清理" or data.get("case_type") == "消費者債務清理":
            data["case_type"] = "消費者債務清理"
            data.setdefault("stage", "")
        return {k: v for k, v in data.items() if v}
    return data


def _stage_marks(data: dict[str, Any]) -> str:
    stage = _metadata_first(data, "stage", "trial_stage", "instance", "審級")
    if "偵" in stage:
        return "□第　　審　■偵 查 中　□"
    if "二" in stage or "2" in stage:
        return "□第 一 審　■第 二 審　□偵 查 中　□"
    if "三" in stage or "3" in stage:
        return "□第 一 審　□第 二 審　■第 三 審　□偵 查 中　□"
    if "一" in stage or "1" in stage:
        return "■第 一 審　□第 二 審　□偵 查 中　□"
    return "□第　　審　□偵 查 中　□"


def _role_marks(data: dict[str, Any], case_type: str) -> str:
    role = _metadata_first(data, "poa_role", "role", "case_role")
    joined = f"{case_type} {role}"
    if "刑" in joined or "辯" in joined:
        return "□代 理 人　□告訴代理人　■辯 護 人　□輔 佐 人"
    if "告訴" in joined:
        return "□代 理 人　■告訴代理人　□辯 護 人　□輔 佐 人"
    return "■代 理 人　□告訴代理人　□辯 護 人　□輔 佐 人"


def _court_line(data: dict[str, Any]) -> str:
    court = _metadata_first(data, "court", "court_name", "法院", "prosecutor_office")
    if not court:
        return "□　　　　　　　法 院　□　　　　　　檢察署　□轉呈　□　　　　　委員會"
    if "檢" in court:
        return f"□ 法 院　■ {court}　□轉呈　□ 委員會"
    if "委員會" in court:
        return f"□ 法 院　□ 檢察署　□轉呈　■ {court}"
    return f"■ {court}　□ 檢察署　□轉呈　□ 委員會"


def _read_pdf_text(path: Path) -> str:
    try:
        import fitz  # type: ignore

        with fitz.open(str(path)) as pdf:
            return "\n".join(page.get_text("text") for page in pdf)
    except Exception:
        return ""


def _extract_application_contact_metadata(text: str) -> dict[str, str]:
    lines = [_compact_pdf_line(line) for line in text.splitlines()]
    data: dict[str, str] = {}
    mobile = ""
    mobile_match = re.search(r"\b(09\d{8})\b", text)
    if mobile_match:
        mobile = mobile_match.group(1)

    address = ""
    for idx, line in enumerate(lines):
        if line != "通訊地址":
            continue
        for candidate in lines[idx + 1 : idx + 5]:
            if not candidate or candidate.startswith("<") or candidate == "[ ]":
                continue
            if re.search(r"\[\d{3}\].+", candidate) or any(token in candidate for token in ("縣", "市", "鄉", "鎮", "區", "路", "街")):
                address = re.sub(r"^\[(\d{3})\]\s*", r"\1", candidate)
                break
        if address:
            break

    if not address:
        address_match = re.search(r"\[(\d{3})\]\s*([^\n\r<]{6,100})", text)
        if address_match:
            address = f"{address_match.group(1)}{_compact_pdf_line(address_match.group(2))}"

    if address and mobile:
        data["client_address_phone"] = f"{address}\n手機：{mobile}"
    elif address:
        data["client_address_phone"] = address
    elif mobile:
        data["client_address_phone"] = f"手機：{mobile}"
    return data


def _extract_case_overview_metadata(text: str) -> dict[str, str]:
    data: dict[str, str] = {}
    match = re.search(r"本案管轄法院[（(]或機關[）)][:：]\s*([^\n\r]+)", text)
    if not match:
        return data
    line = _compact_pdf_line(match.group(1))
    before_case_no = re.split(r"\s*案號[:：]", line, maxsplit=1)[0]
    court_candidates = [
        _compact_pdf_line(part)
        for part in re.split(r"[,，、;；]", before_case_no)
        if _compact_pdf_line(part)
    ]
    for candidate in reversed(court_candidates):
        if any(token in candidate for token in ("法院", "檢察署", "委員會")):
            data["court_name"] = candidate
            break

    case_no_match = re.search(r"案號[:：]\s*([^股別\n\r]+)", line)
    if case_no_match:
        case_no = _compact_pdf_line(case_no_match.group(1))
        if case_no and case_no not in {"不清楚", "待確認"} and any(ch.isdigit() for ch in case_no):
            data["court_case_number"] = case_no
    return data


def _extract_laf_poa_sibling_metadata(path: Path) -> dict[str, str]:
    """Extract fields from other PDFs downloaded with the LAF POA packet."""

    data: dict[str, str] = {}
    parent = path.parent
    if not parent.exists():
        return data

    for pdf in sorted(parent.glob("法律扶助申請書*.pdf")):
        extracted = _extract_application_contact_metadata(_read_pdf_text(pdf))
        for key, value in extracted.items():
            data.setdefault(key, value)

    for pdf in sorted(parent.glob("案件概述單*.pdf")):
        extracted = _extract_case_overview_metadata(_read_pdf_text(pdf))
        for key, value in extracted.items():
            data.setdefault(key, value)
    return data


def _compact_pdf_line(line: str) -> str:
    return re.sub(r"\s+", " ", line).strip()


def _next_meaningful_line(lines: list[str], start: int) -> str:
    for line in lines[start:]:
        value = _compact_pdf_line(line)
        if value:
            return value
    return ""


def _extract_laf_poa_pdf_metadata(path: Path) -> dict[str, str]:
    """Extract official fields from the downloaded LAF POA PDF itself."""

    data: dict[str, str] = {}
    try:
        import fitz  # type: ignore

        with fitz.open(str(path)) as pdf:
            text = "\n".join(page.get_text("text") for page in pdf)
            first_page_words = pdf[0].get_text("words") if len(pdf) else []
    except Exception:
        return data

    normalized_text = re.sub(r"[ \t\u3000]+", " ", text)
    lines = [_compact_pdf_line(line) for line in text.splitlines()]
    if "刑事委任狀" in normalized_text:
        data["poa_layout"] = "criminal"

        def words_in_rect(x0: float, y0: float, x1: float, y1: float) -> str:
            items: list[tuple[float, float, str]] = []
            for word in first_page_words:
                wx0, wy0, wx1, wy1, value, *_ = word
                if wx0 >= x0 and wy0 >= y0 and wx1 <= x1 and wy1 <= y1 and _text(value):
                    items.append((wy0, wx0, _text(value)))
            return _compact_pdf_line(" ".join(value for _y, _x, value in sorted(items)))

        client_name = words_in_rect(70, 135, 245, 250)
        if client_name and len(client_name) <= 20:
            data["client_name"] = client_name
        lawyer_name = words_in_rect(80, 285, 235, 335)
        lawyer_match = re.search(r"([\u4e00-\u9fff]{2,4}律師)", lawyer_name)
        if lawyer_match:
            data["lawyer_name"] = lawyer_match.group(1)
        office_text = words_in_rect(245, 280, 560, 385)
        if office_text:
            office_match = re.search(r"([\u4e00-\u9fffA-Za-z0-9　 ]{2,30}(?:事務所|法律事務所))", office_text)
            if office_match:
                data["law_firm_office_name"] = _compact_pdf_line(office_match.group(1))
            address_match = re.search(r"地址[:：]\s*([^電話傳真]+)", office_text)
            if address_match:
                data["law_firm_address_line"] = _compact_pdf_line(address_match.group(1))
            phone_match = re.search(r"電話[:：]\s*([^傳真]+)", office_text)
            if phone_match:
                data["law_firm_phone"] = _compact_pdf_line(phone_match.group(1))
            fax_match = re.search(r"傳真[:：]\s*(.+)$", office_text)
            if fax_match:
                data["law_firm_fax"] = _compact_pdf_line(fax_match.group(1))
        court_text = words_in_rect(65, 520, 265, 555)
        court_match = re.search(r"(臺灣|台灣).+?(?:法院|檢察署)", court_text)
        if court_match:
            data["court_name"] = court_match.group(0)
    elif "財團法人法律扶助基金會專用委任狀" in normalized_text:
        data["poa_layout"] = "laf_foundation"

    laf_match = re.search(r"本會申請編號[:：]\s*([0-9]{6,7}-[A-Z]-\d{3})", normalized_text)
    if laf_match:
        data["laf_case_number"] = laf_match.group(1)
    else:
        filename_laf_number = _laf_case_number_from_filename(path)
        if filename_laf_number:
            data["laf_case_number"] = filename_laf_number

    branch_match = re.search(r"本事件經本會\s*(.+?)\s*審核准予扶助", normalized_text, re.S)
    if branch_match:
        data["branch"] = _compact_pdf_line(branch_match.group(1))
    phone_match = re.search(r"逕致電分會\(([^)]+)\)", normalized_text)
    if phone_match:
        data["branch_phone"] = _compact_pdf_line(phone_match.group(1))

    if "受原住民族委員會委託辦理原住民法律扶助專用委任狀" in normalized_text:
        data["branch"] = data.get("branch") or "原住民族法律服務中心"

    court_case_match = re.search(r"案號[:：]\s*([^\n\r]+)", normalized_text)
    if court_case_match:
        court_case = _compact_pdf_line(court_case_match.group(1))
        if any(ch.isdigit() for ch in court_case):
            data["court_case_number"] = court_case

    id_match = re.search(r"\b([A-Z][12]\d{8})\b", normalized_text)
    if id_match:
        data["client_id"] = id_match.group(1)

    lawyer_match = re.search(r"([\u4e00-\u9fff]{2,4}律師)", normalized_text)
    if lawyer_match:
        data["lawyer_name"] = lawyer_match.group(1)

    reason_match = re.search(r"為\s*([^\n\r]{1,40}?)\s*事[（(]案[）)]件", normalized_text)
    if reason_match:
        reason = _compact_pdf_line(reason_match.group(1))
        if reason and " " not in reason:
            data["case_reason"] = reason

    for idx, line in enumerate(lines):
        if line == "姓名" and not data.get("client_name"):
            name = _next_meaningful_line(lines, idx + 1)
            if name and "律師" not in name and len(name) <= 20:
                data["client_name"] = name
        elif line == "出生年月日" and not data.get("client_birthday"):
            birthday = _next_meaningful_line(lines, idx + 1)
            if re.search(r"\d+\s*年\s*\d+\s*月\s*\d+\s*日", birthday):
                data["client_birthday"] = birthday.replace(" ", "")

    data.update({k: v for k, v in _roc_date_from_filename(path).items() if v})
    return data


def _normalize_case_metadata(case_metadata: dict[str, Any] | None) -> dict[str, str]:
    data = case_metadata or {}
    case_type = _metadata_first(data, "case_type", "type", "category", "case_category")
    case_reason = _metadata_first(data, "case_reason", "reason", "cause", "案由")
    court_name = _metadata_first(data, "court", "court_name", "法院", "prosecutor_office")
    branch_label = normalize_laf_branch_label(_metadata_first(data, "branch", "laf_branch", "division"))
    branch_phone = _laf_branch_phone(data, branch_label)
    roc_year, roc_month, roc_day = _roc_date_parts(data)
    # Only official POA/PDF lawyer fields can set the POA recipient.  Generic
    # OSC case lawyer fields may be the internal assigned lawyer and must not
    # silently change the LAF form's受任人.
    lawyer_name = _normalize_lawyer_name(
        _metadata_first(data, "poa_lawyer_name", "official_lawyer_name", "pdf_lawyer_name"),
        case_type,
        case_reason,
    )
    return {
        "poa_layout": _metadata_first(data, "poa_layout", "layout"),
        "client_name": _metadata_first(data, "client_name", "name", "party_name", "當事人"),
        "client_birthday": _metadata_first(data, "client_birthday", "birthday", "birth_date", "birth"),
        "client_id": _metadata_first(data, "client_id", "id_number", "identity_number", "national_id"),
        "client_address_phone": _metadata_first(data, "client_address_phone", "address_phone", "address"),
        "laf_case_number": _metadata_first(data, "laf_case_number", "legal_aid_number", "applyno", "case_number"),
        "court_case_number": _metadata_first(data, "court_case_number", "court_case_no", "court_no", "court_number"),
        "branch": branch_label,
        "branch_phone": branch_phone,
        "case_reason": case_reason,
        "case_type": case_type,
        "lawyer_name": lawyer_name,
        "law_firm_office_name": _metadata_first(data, "law_firm_office_name", "office_name", "firm_name"),
        "law_firm_address_line": _metadata_first(data, "law_firm_address_line", "firm_address", "office_address"),
        "law_firm_phone": _metadata_first(data, "law_firm_phone", "firm_phone", "office_phone"),
        "law_firm_fax": _metadata_first(data, "law_firm_fax", "firm_fax", "office_fax"),
        "court_name": court_name,
        "court_line": _court_line(data),
        "stage_marks": _stage_marks(data),
        "role_marks": _role_marks(data, case_type),
        "roc_year": roc_year,
        "roc_month": roc_month,
        "roc_day": roc_day,
    }


def _laf_poa_template_key(case_metadata: dict[str, str] | None = None) -> str:
    data = case_metadata or {}
    key = "indigenous_center" if "原住民族" in _text(data.get("branch")) else "general"
    return key


def _laf_poa_template_title(template_key: str) -> str:
    if template_key == "indigenous_center":
        return "受原住民族委員會委託辦理原住民法律扶助專用委任狀"
    return "財團法人法律扶助基金會專用委任狀"


def select_laf_poa_template(
    case_metadata: dict[str, str] | None = None,
) -> tuple[str, Path] | None:
    if not laf_poa_docx_templates_enabled() or not laf_poa_static_templates_enabled():
        return None
    if _text((case_metadata or {}).get("poa_layout")) == "criminal":
        return None
    key = _laf_poa_template_key(case_metadata)
    template_path = TEMPLATE_DIR / TEMPLATE_FILENAMES[key]
    if not template_path.exists():
        return None
    return key, template_path


def _template_values(metadata: dict[str, str]) -> dict[str, str]:
    firm = get_law_firm_profile()
    court_case_year, court_case_kind, court_case_number, court_case_stock = (
        _court_case_number_parts(metadata["court_case_number"]) or ("", "", "", "")
    )
    court_target = metadata.get("court_name", "")
    court_name_placeholder = "" if any(token in court_target for token in ("檢", "委員會")) else court_target
    return {
        "LAF_CASE_NUMBER": metadata["laf_case_number"],
        "COURT_CASE_NUMBER": metadata["court_case_number"],
        "COURT_CASE_YEAR": court_case_year,
        "COURT_CASE_KIND": court_case_kind,
        "COURT_CASE_NUMBER_ONLY": court_case_number,
        "COURT_CASE_STOCK": court_case_stock,
        "CLIENT_NAME": metadata["client_name"],
        "CLIENT_BIRTHDAY": metadata["client_birthday"],
        "CLIENT_ID": metadata["client_id"],
        "CLIENT_ADDRESS_PHONE": metadata["client_address_phone"],
        "CASE_TYPE": metadata["case_type"],
        "LAWYER_NAME": metadata["lawyer_name"] or firm.lawyer_name,
        "LAW_FIRM_OFFICE_NAME": metadata.get("law_firm_office_name", "") or firm.office_name,
        "LAW_FIRM_ADDRESS_LINE": metadata.get("law_firm_address_line", "") or firm.address_line,
        "LAW_FIRM_PHONE": metadata.get("law_firm_phone", "") or firm.phone,
        "LAW_FIRM_FAX": metadata.get("law_firm_fax", "") or firm.fax,
        "LAW_FIRM_MOBILE": firm.mobile,
        "CASE_REASON": metadata["case_reason"],
        "COURT_NAME": court_name_placeholder,
        "COURT_TARGET": court_target,
        "COURT_LINE": metadata["court_line"],
        "STAGE_MARKS": metadata["stage_marks"],
        "ROLE_MARKS": metadata["role_marks"],
        "BRANCH_LABEL": metadata["branch"] or "待確認分會",
        "BRANCH_PHONE": metadata["branch_phone"],
        "ROC_YEAR": metadata["roc_year"],
        "ROC_MONTH": metadata["roc_month"],
        "ROC_DAY": metadata["roc_day"],
    }


def _quality_expected_values(metadata: dict[str, str], values: dict[str, str]) -> dict[str, str]:
    """Build expected text set for quality validation in output docs."""
    return {
        "client_name": metadata.get("client_name", ""),
        "case_reason": metadata.get("case_reason", ""),
        "lawyer_name": values.get("LAWYER_NAME", ""),
        "lawyer_address": values.get("LAW_FIRM_ADDRESS_LINE", ""),
        "lawyer_phone": values.get("LAW_FIRM_PHONE", ""),
        "court": metadata.get("court_name", "") or _strip_court_target(values.get("COURT_LINE", "")),
    }


def _apply_quality_gate(
    result: dict[str, Any],
    docx_path: Path,
    metadata: dict[str, str],
    template_values: dict[str, str],
    *,
    check_expected_values: bool = True,
) -> None:
    validation = validate_laf_poa_docx_quality(
        docx_path,
        expected_fields=_quality_expected_values(metadata, template_values),
        check_expected_values=check_expected_values,
    )
    result["quality"] = validation
    if not validation["ok"]:
        result["warnings"].append("quality_validation_failed")


def _strip_court_target(court_line: str) -> str:
    value = _text(court_line)
    for token in ("■", "□", "法 院", "法  院", "法院", "檢察署", "轉呈", "委員會"):
        value = value.replace(token, " ")
    return re.sub(r"\s+", "", value).strip()


def _court_target_from_line(court_line: str) -> str:
    value = _text(court_line)
    if "■" not in value:
        return ""
    selected = value.split("■", 1)[1]
    selected = re.split(r"\s*□", selected, maxsplit=1)[0]
    return re.sub(r"\s+", "", selected).strip()


def _insert_text_after_first_wt(xml: str, label: str, text: str) -> str:
    if not text:
        return xml
    token = f"<w:t>{label}</w:t>"
    index = xml.find(token)
    if index < 0:
        return xml
    insert_at = index + len(token)
    safe_text = html.escape(text)
    run = f'</w:r><w:r><w:rPr><w:sz w:val="24"/></w:rPr><w:t>{safe_text}</w:t></w:r><w:r>'
    return f"{xml[:insert_at]}{run}{xml[insert_at:]}"


def _court_case_number_parts(court_case_number: str) -> tuple[str, str, str, str] | None:
    value = _text(court_case_number)
    match = re.match(
        r"^(?P<year>\d{2,3})\s*年度\s*(?P<kind>.+?)\s*字\s*第\s*(?P<number>.+?)\s*號(?:\s*(?P<stock>.*?)\s*股)?$",
        value,
    )
    if not match:
        return None
    return (
        _text(match.group("year")),
        _text(match.group("kind")),
        _text(match.group("number")),
        _text(match.group("stock")),
    )


def _fill_court_case_number(xml: str, court_case_number: str) -> str:
    value = _text(court_case_number)
    if not value:
        return xml
    else:
        parts = _court_case_number_parts(value) or ("", "", value, "")
    year, kind, number, stock = (html.escape(part) for part in parts)

    # General LAF template: the blanks are tabs in the same run.
    general_pattern = (
        "<w:t>案號：</w:t><w:tab/><w:t>年度</w:t><w:tab/><w:t>字第</w:t>"
        "<w:tab/><w:t>號</w:t><w:tab/><w:t>股</w:t>"
    )
    if general_pattern in xml:
        general_replacement = (
            f"<w:t>案號：</w:t><w:t>{year}</w:t><w:tab/><w:t>年度</w:t>"
            f"<w:t>{kind}</w:t><w:tab/><w:t>字第</w:t>"
            f"<w:t>{number}</w:t><w:tab/><w:t>號</w:t>"
            f"<w:t>{stock}</w:t><w:tab/><w:t>股</w:t>"
        )
        return xml.replace(general_pattern, general_replacement, 1)

    label = "<w:t>案號：</w:t>"
    label_index = xml.find(label)
    if label_index < 0:
        return xml
    paragraph_start = xml.rfind("<w:p", 0, label_index)
    paragraph_end = xml.find("</w:p>", label_index)
    if paragraph_start < 0 or paragraph_end < 0:
        return _insert_text_after_first_wt(xml, "案號：", value or "請填法院案號")
    paragraph_end += len("</w:p>")
    paragraph = xml[paragraph_start:paragraph_end]
    replacements = iter((year, kind, number, stock))

    def _run_text(text: str) -> str:
        if not text:
            return ""
        return (
            '<w:r><w:rPr><w:rFonts w:ascii="標楷體" w:eastAsia="標楷體" '
            'w:hAnsi="標楷體" w:cs="標楷體"/><w:sz w:val="20"/></w:rPr>'
            f"<w:t>{text}</w:t></w:r>"
        )

    if "<w:t>年度</w:t>" in paragraph and "<w:t>字第</w:t>" in paragraph and "<w:t>號</w:t>" in paragraph:
        filled = paragraph
        for marker, replacement in (
            ("案號：", year),
            ("年度", kind),
            ("字第", number),
            ("號", stock),
        ):
            if not replacement:
                continue
            filled = re.sub(
                fr"(<w:t>{re.escape(marker)}</w:t></w:r>)(<w:r><w:tab/></w:r>)",
                lambda match, text=replacement: f"{match.group(1)}{_run_text(text)}{match.group(2)}",
                filled,
                count=1,
            )
        if filled != paragraph:
            return f"{xml[:paragraph_start]}{filled}{xml[paragraph_end:]}"

    def _replace_blank(match: re.Match[str]) -> str:
        try:
            replacement = next(replacements)
        except StopIteration:
            return match.group(0)
        if not replacement:
            return match.group(0)
        return f"{match.group(1)}{replacement}{match.group(3)}"

    filled = re.sub(
        r'(<w:t xml:space="preserve">)([\s\u3000]+)(</w:t>)',
        _replace_blank,
        paragraph,
        count=4,
    )
    if filled == paragraph and value:
        return _insert_text_after_first_wt(xml, "案號：", value)
    return f"{xml[:paragraph_start]}{filled}{xml[paragraph_end:]}"


def _replace_last_underlined_blank_before(xml: str, marker: str, text: str) -> str:
    if not text:
        return xml
    marker_index = xml.find(marker)
    if marker_index < 0:
        return xml
    prefix = xml[:marker_index]
    suffix = xml[marker_index:]
    blank_pattern = re.compile(
        r'(<w:r\b(?:(?!</w:r>).)*?<w:u w:val="single"/>(?:(?!</w:r>).)*?<w:t(?: xml:space="preserve")?>)'
        r"([\s\u3000]*)"
        r"(</w:t>(?:(?!</w:r>).)*?</w:r>)",
        re.S,
    )
    matches = list(blank_pattern.finditer(prefix))
    if not matches:
        return xml
    match = matches[-1]
    safe_text = html.escape(text)
    prefix = f"{prefix[:match.start()]}{match.group(1)}{safe_text}{match.group(3)}{prefix[match.end():]}"
    return prefix + suffix


def _fill_static_word_form_gaps(xml: str, values: dict[str, str]) -> str:
    """Fill visible blanks that are part of the official Word layout itself.

    The downloaded LAF Word forms leave the court and court-case-number areas as
    literal blank runs rather than placeholders.  Filling those blanks here keeps
    MAGI's output as a native Word form while still making every needed field
    discoverable and editable.
    """

    court_case_number = _text(values.get("COURT_CASE_NUMBER"))
    court_target = _text(values.get("COURT_TARGET")) or _text(values.get("COURT_NAME")) or _strip_court_target(values.get("COURT_LINE", ""))

    xml = _fill_court_case_number(xml, court_case_number)
    xml = _fill_static_court_target(xml, values, court_target)
    return _mark_static_word_choices(xml, values)


def _replace_first(xml: str, old: str, new: str) -> str:
    return xml.replace(old, new, 1) if old in xml else xml


def _fill_static_court_target(xml: str, values: dict[str, str], court_target: str) -> str:
    court_line = _text(values.get("COURT_LINE"))
    target = _text(court_target)
    wants_prosecutor = "檢" in target or ("■" in court_line and "檢" in court_line)
    wants_committee = "委員會" in target or ("■" in court_line and "委員會" in court_line)

    if wants_prosecutor:
        replacement = f"☒{target}" if target else "☒　　　　　　　　　檢察署"
        return _replace_first(xml, "□　　　　　　　　　檢察署", replacement)
    if wants_committee:
        replacement = f"☒{target}" if target else "☒　　　　　　　　　委員會"
        return _replace_first(xml, "□　　　　　　　　　委員會", replacement)
    if target:
        xml = _replace_first(xml, f"□{target}　　　　　　　法　院", f"☒{target}")
        if target not in xml:
            xml = _replace_first(xml, "□　　　　　　　法　院", f"☒{target}")
            xml = _replace_last_underlined_blank_before(xml, "<w:t>法</w:t>", target)
            xml = _replace_last_underlined_blank_before(xml, "<w:t>法</w:t><w:tab/><w:t>院</w:t>", target)
            xml = _replace_last_underlined_blank_before(xml, "<w:t>法  院</w:t>", target)
    return xml


def _mark_static_word_choices(xml: str, values: dict[str, str]) -> str:
    """Mark checkbox glyphs in the native Word template."""

    case_type = _text(values.get("CASE_TYPE"))
    if case_type:
        if "刑" in case_type:
            xml = _replace_first(xml, "□　刑事", "☒　刑事")
        elif "訴願" in case_type:
            xml = _replace_first(xml, "□　訴願", "☒　訴願")
        elif "行政" in case_type:
            xml = _replace_first(xml, "□　行政訴訟", "☒　行政訴訟")
        else:
            xml = _replace_first(xml, "□　民事", "☒　民事")

    stage_marks = _text(values.get("STAGE_MARKS"))
    if "■偵" in stage_marks:
        xml = _replace_first(xml, "□偵 查 中", "☒偵 查 中")
    elif "■第 一 審" in stage_marks:
        xml = _replace_first(xml, "□第　　審", "☒第 一 審")
    elif "■第 二 審" in stage_marks:
        xml = _replace_first(xml, "□第　　審", "☒第 二 審")
    elif "■第 三 審" in stage_marks:
        xml = _replace_first(xml, "□第　　審", "☒第 三 審")

    role_marks = _text(values.get("ROLE_MARKS"))
    if "■代 理 人" in role_marks:
        xml = _replace_first(xml, "□代　理　人", "☒代　理　人")
    if "■告訴代理人" in role_marks:
        xml = _replace_first(xml, "□告訴代理人", "☒告訴代理人")
    if "■辯 護 人" in role_marks:
        xml = _replace_first(xml, "□辯 護 人", "☒辯 護 人")
    if "■輔 佐 人" in role_marks:
        xml = _replace_first(xml, "□輔 佐 人", "☒輔 佐 人")
    return xml


def _docx_text_fragment(value: str) -> str:
    text = html.escape(_text(value))
    if "\n" not in text:
        return text
    return "</w:t><w:br/><w:t>".join(text.splitlines())


def _law_firm_contact_fragment(values: dict[str, str]) -> str:
    address = _text(values.get("LAW_FIRM_ADDRESS_LINE"))
    phone = _text(values.get("LAW_FIRM_PHONE"))
    if address.startswith("970") and len(address) > 3:
        address_lines = ["970", address[3:]]
    else:
        address_lines = [address] if address else []
    if phone:
        address_lines.append(f"電話：{phone}")
    return _docx_text_fragment("\n".join(line for line in address_lines if line))


def _replace_law_firm_contact_placeholders(xml: str, values: dict[str, str]) -> str:
    fragment = _law_firm_contact_fragment(values)
    if not fragment:
        return xml
    for token in (
        "{{LAW_FIRM_ADDRESS_LINE}}　電話：{{LAW_FIRM_PHONE}}",
        "{{LAW_FIRM_ADDRESS_LINE}}電話：{{LAW_FIRM_PHONE}}",
    ):
        xml = xml.replace(token, fragment)
    return xml


def _relax_single_line_form_fields(xml: str, values: dict[str, str]) -> str:
    for key in ("CLIENT_ID",):
        value = html.escape(_text(values.get(key)))
        if not value:
            continue
        pattern = re.compile(
            r"(<w:p\b(?:(?!</w:p>).)*?<w:t>"
            + re.escape(value)
            + r"</w:t>(?:(?!</w:p>).)*?</w:p>)",
            re.S,
        )

        def _relax(match: re.Match[str]) -> str:
            paragraph = match.group(1)
            if "<w:ind" in paragraph:
                paragraph = re.sub(r"<w:ind\b[^>]*/>", '<w:ind w:left="0" w:right="0"/>', paragraph)
            else:
                paragraph = paragraph.replace("<w:pPr>", '<w:pPr><w:ind w:left="0" w:right="0"/>', 1)
            if key == "CLIENT_ID":
                paragraph = re.sub(r'<w:sz w:val="28"/>', '<w:sz w:val="24"/>', paragraph)
            return paragraph

        xml = pattern.sub(_relax, xml)
    return xml


def _force_biau_kai_xml(xml: str) -> str:
    font = "標楷體"
    rfonts = f'<w:rFonts w:ascii="{font}" w:eastAsia="{font}" w:hAnsi="{font}" w:cs="{font}"/>'

    def _normalize_rfonts_tag(match: re.Match[str]) -> str:
        tag = match.group(0)
        for attr in ("w:ascii", "w:eastAsia", "w:hAnsi", "w:cs"):
            if re.search(fr'{attr}="[^"]*"', tag):
                tag = re.sub(fr'{attr}="[^"]*"', f'{attr}="{font}"', tag)
            else:
                tag = tag[:-2] + f' {attr}="{font}"/>'
        return tag

    xml = re.sub(r"<w:rFonts\b[^>]*/>", _normalize_rfonts_tag, xml)
    xml = xml.replace("<w:rPr/>", f"<w:rPr>{rfonts}</w:rPr>")
    xml = re.sub(r"<w:rPr>(?!<w:rFonts\b)", f"<w:rPr>{rfonts}", xml)
    return xml


def _replace_docx_placeholders(
    template_path: Path,
    output_path: Path,
    values: dict[str, str],
    *,
    fill_static_gaps: bool = True,
) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    tmp_path = output_path.with_suffix(output_path.suffix + ".tmp")
    with zipfile.ZipFile(template_path, "r") as src, zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as dst:
        for item in src.infolist():
            payload = src.read(item.filename)
            if item.filename.endswith(".xml"):
                xml = payload.decode("utf-8")
                xml = _replace_law_firm_contact_placeholders(xml, values)
                for key, value in values.items():
                    xml = xml.replace(f"{{{{{key}}}}}", _docx_text_fragment(_text(value)))
                if item.filename == "word/document.xml" and fill_static_gaps:
                    xml = _fill_static_word_form_gaps(xml, values)
                    xml = _relax_single_line_form_fields(xml, values)
                if item.filename.startswith("word/"):
                    xml = _force_biau_kai_xml(xml)
                payload = xml.encode("utf-8")
            dst.writestr(item, payload)
    try:
        from docx import Document  # type: ignore

        Document(str(tmp_path))
    except Exception:
        tmp_path.unlink(missing_ok=True)
        raise
    os.replace(tmp_path, output_path)


def _create_from_laf_poa_template(
    template_key: str,
    template_path: Path,
    template_target: Path,
    target: Path,
    metadata: dict[str, str],
) -> dict[str, str]:
    _replace_docx_placeholders(template_path, template_target, {}, fill_static_gaps=False)

    values = _template_values(metadata)
    _replace_docx_placeholders(template_path, target, values)
    values["template_key"] = template_key
    return values


def ensure_laf_poa_docx_companion(
    pdf_path: str | os.PathLike[str],
    *,
    overwrite: bool = False,
    case_metadata: dict[str, Any] | None = None,
) -> dict[str, Any]:
    """Create a Word companion for LAF power-of-attorney PDFs.

    The default output uses the curated native Word template and fills official
    case fields directly.  PDF-backed reconstruction is available only as an
    explicit compatibility mode for layouts that do not have a template.
    """

    source = Path(pdf_path)
    template_target = laf_poa_template_docx_path(source)
    target = laf_poa_case_docx_path(source)
    result: dict[str, Any] = {
        "ok": False,
        "status": "",
        "pdf_path": str(source),
        "template_docx_path": str(template_target),
        "docx_path": str(target),
        "pages": 0,
        "error": "",
        "filled_fields": {},
        "template_key": "",
        "warnings": [],
    }

    if not laf_poa_docx_enabled():
        result.update(ok=True, status="disabled")
        return result
    if not is_laf_power_of_attorney_pdf(source):
        result.update(ok=True, status="not_poa_pdf")
        return result
    if not _exists_quick(source):
        result.update(status="missing_pdf", error="pdf_not_found")
        return result

    pdf_metadata = _extract_laf_poa_pdf_metadata(source)
    if pdf_metadata.get("lawyer_name"):
        pdf_metadata["poa_lawyer_name"] = pdf_metadata["lawyer_name"]
    if not pdf_metadata.get("laf_case_number"):
        filename_laf_number = _laf_case_number_from_filename(source)
        if filename_laf_number:
            pdf_metadata["laf_case_number"] = filename_laf_number
    combined_metadata: dict[str, Any] = {}
    combined_metadata.update(_case_folder_metadata_from_path(source))
    combined_metadata.update(_extract_laf_poa_sibling_metadata(source))
    combined_metadata.update(case_metadata or {})
    # The official POA PDF is authoritative for its own case number, party, and
    # branch footer. DB/email metadata only fills fields the PDF leaves blank.
    combined_metadata.update({k: v for k, v in pdf_metadata.items() if v})
    normalized_metadata = _normalize_case_metadata(combined_metadata)
    template_key = _laf_poa_template_key(normalized_metadata)
    result["pdf_extracted_fields"] = pdf_metadata
    exact_pdf_layout = laf_poa_exact_pdf_layout_enabled()
    if exact_pdf_layout:
        _remove_stale_laf_poa_template(template_target)

    existing_ready = _exists_quick(target)
    if existing_ready and not overwrite:
        try:
            target_stat = _stat_quick(target)
            source_stat = _stat_quick(source)
            if target_stat.st_mtime >= source_stat.st_mtime and target_stat.st_size > 0:
                target_is_exact = _is_exact_pdf_backed_docx(target)
                if exact_pdf_layout and target_is_exact:
                    result.update(ok=True, status="exists")
                    return result
                if not exact_pdf_layout and not target_is_exact:
                    result.update(ok=True, status="exists")
                    return result
                result["warnings"].append("stale_non_exact_docx_rebuilt" if exact_pdf_layout else "stale_exact_docx_rebuilt")
        except OSError:
            pass

    selected_template = None if exact_pdf_layout else select_laf_poa_template(normalized_metadata)
    fallback_allowed = exact_pdf_layout or (selected_template is None and laf_poa_pdf_render_fallback_enabled())

    if selected_template is not None:
        template_key, template_path = selected_template
        try:
            filled_fields = _create_from_laf_poa_template(
                template_key,
                template_path,
                template_target,
                target,
                normalized_metadata,
            )
            _apply_quality_gate(result, target, normalized_metadata, filled_fields)
            result.update(
                ok=True,
                status="created",
                pages=1,
                filled_fields={k: v for k, v in normalized_metadata.items() if v},
                template_key=template_key,
                branch_label=normalized_metadata.get("branch", ""),
                branch_phone=normalized_metadata.get("branch_phone", ""),
                template_values={k: v for k, v in filled_fields.items() if v},
            )
            if normalized_metadata.get("branch_phone") == "待確認":
                result["warnings"].append("missing_branch_phone")
            return result
        except Exception as exc:
            result["warnings"].append(f"template_failed:{exc}")
            result.update(status="template_failed", error=str(exc))
            return result

    if selected_template is None and not fallback_allowed:
        result.update(
            status="template_unavailable",
            error="laf_poa_docx_template_required",
            template_key=template_key,
        )
        return result

    try:
        import fitz  # type: ignore
        from docx import Document  # type: ignore
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT  # type: ignore
        from docx.enum.section import WD_SECTION_START  # type: ignore
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
        from docx.enum.table import WD_ROW_HEIGHT_RULE  # type: ignore
        from docx.oxml import OxmlElement, parse_xml  # type: ignore
        from docx.oxml.ns import nsdecls, qn  # type: ignore
        from docx.shared import Mm, Pt  # type: ignore

        def emu(points: float) -> int:
            return int(round(points * 12700))

        def clear_paragraph(paragraph) -> None:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1

        def set_table_borders_none(table) -> None:
            tbl = table._tbl
            tbl_pr = tbl.tblPr
            borders = tbl_pr.find(qn("w:tblBorders"))
            if borders is None:
                borders = OxmlElement("w:tblBorders")
                tbl_pr.append(borders)
            for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
                edge = borders.find(qn(f"w:{name}"))
                if edge is None:
                    edge = OxmlElement(f"w:{name}")
                    borders.append(edge)
                edge.set(qn("w:val"), "nil")

        def set_cell_margins_zero(cell) -> None:
            tc_pr = cell._tc.get_or_add_tcPr()
            margins = tc_pr.first_child_found_in("w:tcMar")
            if margins is None:
                margins = OxmlElement("w:tcMar")
                tc_pr.append(margins)
            for name in ("top", "start", "left", "bottom", "end", "right"):
                margin = margins.find(qn(f"w:{name}"))
                if margin is None:
                    margin = OxmlElement(f"w:{name}")
                    margins.append(margin)
                margin.set(qn("w:w"), "0")
                margin.set(qn("w:type"), "dxa")

        def set_table_fixed_geometry(table, width_pt: float, col_width_pt: float, cols: int) -> None:
            tbl = table._tbl
            tbl_pr = tbl.tblPr
            layout = tbl_pr.find(qn("w:tblLayout"))
            if layout is None:
                layout = OxmlElement("w:tblLayout")
                tbl_pr.append(layout)
            layout.set(qn("w:type"), "fixed")
            tbl_w = tbl_pr.find(qn("w:tblW"))
            if tbl_w is None:
                tbl_w = OxmlElement("w:tblW")
                tbl_pr.append(tbl_w)
            tbl_w.set(qn("w:type"), "dxa")
            tbl_w.set(qn("w:w"), str(int(round(width_pt * 20))))
            grid = tbl.tblGrid
            if grid is None:
                grid = OxmlElement("w:tblGrid")
                tbl.insert(0, grid)
            for child in list(grid):
                grid.remove(child)
            for _ in range(cols):
                grid_col = OxmlElement("w:gridCol")
                grid_col.set(qn("w:w"), str(int(round(col_width_pt * 20))))
                grid.append(grid_col)

        overlay_counter = 0

        def field_tag(name: str) -> str:
            tag = re.sub(r"[^A-Za-z0-9_.:-]+", "_", name).strip("_")
            return tag or "field"

        def _overlay_grid_paragraph(
            host,
            *,
            x_pt: float,
            y_pt: float,
            width_pt: float,
            height_pt: float,
            align: str = "center",
        ):
            if not isinstance(host, dict):
                return host
            table = host["table"]
            rows = int(host["rows"])
            cols = int(host["cols"])
            row_height = float(host["row_height"])
            col_width = float(host["col_width"])
            row_idx = max(0, min(rows - 1, int(y_pt / row_height)))
            col_idx = max(0, min(cols - 1, int(x_pt / col_width)))
            end_row = max(row_idx, min(rows - 1, int((y_pt + max(height_pt, row_height) - 0.1) / row_height)))
            end_col = max(col_idx, min(cols - 1, int((x_pt + max(width_pt, col_width) - 0.1) / col_width)))
            cell = table.cell(row_idx, col_idx)
            try:
                if end_row != row_idx or end_col != col_idx:
                    cell = cell.merge(table.cell(end_row, end_col))
            except Exception:
                pass
            paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
            if paragraph._p.xpath("./w:r|./w:sdt"):
                paragraph = cell.add_paragraph()
            clear_paragraph(paragraph)
            paragraph.alignment = {
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
            }.get(align, WD_ALIGN_PARAGRAPH.CENTER)
            return paragraph

        def _append_overlay_marker(paragraph) -> None:
            nonlocal overlay_counter
            paragraph._p.append(
                parse_xml(
                    f'<w:bookmarkStart {nsdecls("w")} w:id="{overlay_counter}" '
                    f'w:name="magi_laf_poa_overlay_{overlay_counter}"/>'
                )
            )
            paragraph._p.append(
                parse_xml(f'<w:bookmarkEnd {nsdecls("w")} w:id="{overlay_counter}"/>')
            )

        def _inline_text_control_xml(
            text: str,
            *,
            font_pt: float = 10.5,
            align: str = "center",
            color: str = "000000",
            italic: bool = False,
            field_name: str = "text",
        ):
            clean = html.escape(_text(text), quote=True)
            if not clean:
                return None
            lines = clean.splitlines() or [clean]
            run_xml = ""
            for index, line in enumerate(lines):
                if index:
                    run_xml += "<w:br/>"
                run_xml += f'<w:t xml:space="preserve">{line}</w:t>'
            italic_xml = "<w:i/>" if italic else ""
            size = max(1, int(round(font_pt * 2)))
            align = align if align in {"left", "center", "right"} else "center"
            alias = html.escape(field_name.replace("_", " "), quote=True)
            tag = html.escape(f"magi_laf_poa.{field_tag(field_name)}", quote=True)
            return parse_xml(
                f"""
                <w:sdt {nsdecls("w")}>
                  <w:sdtPr>
                    <w:alias w:val="{alias}"/>
                    <w:tag w:val="{tag}"/>
                    <w:text/>
                  </w:sdtPr>
                  <w:sdtContent>
                    <w:r>
                      <w:rPr>
                        <w:rFonts w:ascii="標楷體" w:eastAsia="標楷體" w:hAnsi="標楷體" w:cs="標楷體"/>
                        <w:color w:val="{color}"/>
                        <w:sz w:val="{size}"/>
                        {italic_xml}
                      </w:rPr>
                      {run_xml}
                    </w:r>
                  </w:sdtContent>
                </w:sdt>
                """
            )

        def add_overlay_textbox(host, text: str, **kwargs) -> None:
            nonlocal overlay_counter
            paragraph = _overlay_grid_paragraph(
                host,
                x_pt=float(kwargs.pop("x_pt")),
                y_pt=float(kwargs.pop("y_pt")),
                width_pt=float(kwargs.pop("width_pt")),
                height_pt=float(kwargs.pop("height_pt")),
                align=str(kwargs.get("align", "center")),
            )
            overlay_counter += 1
            element = _inline_text_control_xml(text, **kwargs)
            if element is not None:
                _append_overlay_marker(paragraph)
                paragraph._p.append(element)

        def _inline_checkbox_control_xml(
            field_name: str,
            *,
            checked: bool = False,
            size_pt: float = 11,
            color: str = "000000",
        ):
            char = "☒" if checked else "☐"
            checked_value = "1" if checked else "0"
            size = max(1, int(round(size_pt * 2)))
            alias = html.escape(field_name.replace("_", " "), quote=True)
            tag = html.escape(f"magi_laf_poa.{field_tag(field_name)}", quote=True)
            return parse_xml(
                f"""
                <w:sdt {nsdecls("w")} xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml">
                  <w:sdtPr>
                    <w:alias w:val="{alias}"/>
                    <w:tag w:val="{tag}"/>
                    <w14:checkbox>
                      <w14:checked w14:val="{checked_value}"/>
                      <w14:checkedState w14:val="2612" w14:font="MS Gothic"/>
                      <w14:uncheckedState w14:val="2610" w14:font="MS Gothic"/>
                    </w14:checkbox>
                  </w:sdtPr>
                  <w:sdtContent>
                    <w:r>
                      <w:rPr>
                        <w:rFonts w:ascii="MS Gothic" w:eastAsia="MS Gothic" w:hAnsi="MS Gothic"/>
                        <w:color w:val="{html.escape(color, quote=True)}"/>
                        <w:sz w:val="{size}"/>
                      </w:rPr>
                      <w:t>{char}</w:t>
                    </w:r>
                  </w:sdtContent>
                </w:sdt>
                """
            )

        def add_overlay_checkbox(host, field_name: str, **kwargs) -> None:
            nonlocal overlay_counter
            size_pt = float(kwargs.get("size_pt", 11))
            visible = bool(kwargs.pop("visible", True))
            paragraph = _overlay_grid_paragraph(
                host,
                x_pt=float(kwargs.pop("x_pt")),
                y_pt=float(kwargs.pop("y_pt")),
                width_pt=size_pt,
                height_pt=size_pt,
                align="center",
            )
            overlay_counter += 1
            _append_overlay_marker(paragraph)
            if visible:
                paragraph._p.append(_inline_checkbox_control_xml(field_name, **kwargs))
            else:
                kwargs["size_pt"] = 1
                kwargs["color"] = "FFFFFF"
                paragraph._p.append(_inline_checkbox_control_xml(field_name, **kwargs))

        def _blank_pdf_overlay(fill_values) -> bool:
            return str(fill_values).strip().lower() in {"blank", "pdf"}

        def _value_or_placeholder(value: str, placeholder: str, fill_values) -> tuple[str, bool]:
            if _blank_pdf_overlay(fill_values):
                return "\u200b", True
            if fill_values and _text(value):
                return _text(value), False
            if fill_values:
                return "\u200b", True
            return placeholder, True

        def add_fillable_overlay_fields(paragraph, values: dict[str, str], *, fill_values) -> None:
            """Place editable Word text boxes over the official PDF form fields."""
            checkbox_visible = not _blank_pdf_overlay(fill_values)

            lawyer_name, lawyer_placeholder = _value_or_placeholder(
                values.get("LAWYER_NAME", ""),
                "請填律師姓名",
                fill_values,
            )
            add_overlay_textbox(
                paragraph,
                lawyer_name,
                field_name="lawyer_name",
                x_pt=412,
                y_pt=184,
                width_pt=118,
                height_pt=22,
                font_pt=12,
                color="666666" if lawyer_placeholder else "000000",
                italic=lawyer_placeholder,
            )

            address_parts = []
            if _text(values.get("LAW_FIRM_ADDRESS_LINE", "")):
                address_parts.append(_text(values["LAW_FIRM_ADDRESS_LINE"]))
            if _text(values.get("LAW_FIRM_PHONE", "")):
                address_parts.append(f"電話：{_text(values['LAW_FIRM_PHONE'])}")
            firm_text, firm_placeholder = _value_or_placeholder(
                "\n".join(address_parts),
                "請填事務所地址\n請填電話",
                fill_values,
            )
            add_overlay_textbox(
                paragraph,
                firm_text,
                field_name="law_firm_contact",
                x_pt=414,
                y_pt=286,
                width_pt=126,
                height_pt=60,
                font_pt=8.5,
                align="left",
                color="666666" if firm_placeholder else "000000",
                italic=firm_placeholder,
            )

            laf_case_number, laf_case_placeholder = _value_or_placeholder(
                values.get("LAF_CASE_NUMBER", ""),
                "請填法扶案號",
                fill_values,
            )
            add_overlay_textbox(
                paragraph,
                laf_case_number,
                field_name="laf_case_number",
                x_pt=118,
                y_pt=100,
                width_pt=142,
                height_pt=20,
                font_pt=9.5,
                align="left",
                color="666666" if laf_case_placeholder else "000000",
                italic=laf_case_placeholder,
            )

            client_name, client_placeholder = _value_or_placeholder(
                values.get("CLIENT_NAME", ""),
                "請填當事人",
                fill_values,
            )
            add_overlay_textbox(
                paragraph,
                client_name,
                field_name="client_name",
                x_pt=118,
                y_pt=184,
                width_pt=96,
                height_pt=22,
                font_pt=12,
                color="666666" if client_placeholder else "000000",
                italic=client_placeholder,
            )

            birthday, birthday_placeholder = _value_or_placeholder(
                values.get("CLIENT_BIRTHDAY", ""),
                "請填生日",
                fill_values,
            )
            add_overlay_textbox(
                paragraph,
                birthday,
                field_name="client_birthday",
                x_pt=222,
                y_pt=184,
                width_pt=92,
                height_pt=22,
                font_pt=10,
                color="666666" if birthday_placeholder else "000000",
                italic=birthday_placeholder,
            )

            client_id, client_id_placeholder = _value_or_placeholder(
                values.get("CLIENT_ID", ""),
                "請填身分證",
                fill_values,
            )
            add_overlay_textbox(
                paragraph,
                client_id,
                field_name="client_id",
                x_pt=322,
                y_pt=184,
                width_pt=82,
                height_pt=22,
                font_pt=10,
                color="666666" if client_id_placeholder else "000000",
                italic=client_id_placeholder,
            )

            address_phone, address_phone_placeholder = _value_or_placeholder(
                values.get("CLIENT_ADDRESS_PHONE", ""),
                "請填地址電話",
                fill_values,
            )
            add_overlay_textbox(
                paragraph,
                address_phone,
                field_name="client_address_phone",
                x_pt=118,
                y_pt=286,
                width_pt=282,
                height_pt=60,
                font_pt=9.5,
                align="left",
                color="666666" if address_phone_placeholder else "000000",
                italic=address_phone_placeholder,
            )

            court_case_number, court_case_placeholder = _value_or_placeholder(
                values.get("COURT_CASE_NUMBER", ""),
                "請填法院案號",
                fill_values,
            )
            add_overlay_textbox(
                paragraph,
                court_case_number,
                field_name="court_case_number",
                x_pt=388,
                y_pt=100,
                width_pt=145,
                height_pt=20,
                font_pt=8.5 if not court_case_placeholder else 9.5,
                color="666666" if court_case_placeholder else "000000",
                italic=court_case_placeholder,
            )

            case_reason, case_reason_placeholder = _value_or_placeholder(
                values.get("CASE_REASON", ""),
                "請填案由",
                fill_values,
            )
            add_overlay_textbox(
                paragraph,
                case_reason,
                field_name="case_reason",
                x_pt=156,
                y_pt=433,
                width_pt=212,
                height_pt=22,
                font_pt=12,
                color="666666" if case_reason_placeholder else "000000",
                italic=case_reason_placeholder,
            )

            court_line = _text(values.get("COURT_LINE", ""))
            court_name = _court_target_from_line(court_line) or _text(values.get("COURT_TARGET", "")) or _text(values.get("COURT_NAME", ""))
            court_text, court_placeholder = _value_or_placeholder(
                court_name,
                "請填法院/檢署/委員會",
                fill_values,
            )
            add_overlay_textbox(
                paragraph,
                court_text,
                field_name="court_name",
                x_pt=104,
                y_pt=625,
                width_pt=198,
                height_pt=20,
                font_pt=11 if not court_placeholder else 9.5,
                align="left",
                color="666666" if court_placeholder else "000000",
                italic=court_placeholder,
            )

            court_line_source = _text(values.get("COURT_LINE", ""))
            add_overlay_checkbox(paragraph, "court_kind_court", x_pt=82, y_pt=625, checked="■" in court_line_source and "法院" in court_line_source, visible=checkbox_visible)
            add_overlay_checkbox(paragraph, "court_kind_prosecutor", x_pt=164, y_pt=625, checked="■" in court_line_source and "檢" in court_line_source, visible=checkbox_visible)
            add_overlay_checkbox(paragraph, "court_kind_transfer", x_pt=266, y_pt=625, checked="■轉呈" in court_line_source, visible=checkbox_visible)
            add_overlay_checkbox(paragraph, "court_kind_committee", x_pt=324, y_pt=625, checked="■" in court_line_source and "委員會" in court_line_source, visible=checkbox_visible)

            stage_marks = _text(values.get("STAGE_MARKS", ""))
            add_overlay_checkbox(paragraph, "stage_first", x_pt=118, y_pt=478, checked="■第 一 審" in stage_marks, visible=checkbox_visible)
            add_overlay_checkbox(paragraph, "stage_second", x_pt=204, y_pt=478, checked="■第 二 審" in stage_marks, visible=checkbox_visible)
            add_overlay_checkbox(paragraph, "stage_third", x_pt=290, y_pt=478, checked="■第 三 審" in stage_marks, visible=checkbox_visible)
            add_overlay_checkbox(paragraph, "stage_investigation", x_pt=376, y_pt=478, checked="■偵" in stage_marks, visible=checkbox_visible)

            role_marks = _text(values.get("ROLE_MARKS", ""))
            add_overlay_checkbox(paragraph, "role_agent", x_pt=118, y_pt=514, checked="■代 理 人" in role_marks, visible=checkbox_visible)
            add_overlay_checkbox(paragraph, "role_complainant_agent", x_pt=204, y_pt=514, checked="■告訴代理人" in role_marks, visible=checkbox_visible)
            add_overlay_checkbox(paragraph, "role_defender", x_pt=318, y_pt=514, checked="■辯 護 人" in role_marks, visible=checkbox_visible)
            add_overlay_checkbox(paragraph, "role_assistant", x_pt=430, y_pt=514, checked="■輔 佐 人" in role_marks, visible=checkbox_visible)

            year, year_placeholder = _value_or_placeholder(values.get("ROC_YEAR", ""), "年", fill_values)
            month, month_placeholder = _value_or_placeholder(values.get("ROC_MONTH", ""), "月", fill_values)
            day, day_placeholder = _value_or_placeholder(values.get("ROC_DAY", ""), "日", fill_values)
            for text, field_name, x_pt, placeholder in (
                (year, "roc_year", 407, year_placeholder),
                (month, "roc_month", 471, month_placeholder),
                (day, "roc_day", 527, day_placeholder),
            ):
                add_overlay_textbox(
                    paragraph,
                    text,
                    field_name=field_name,
                    x_pt=x_pt,
                    y_pt=808,
                    width_pt=32,
                    height_pt=18,
                    font_pt=11,
                    color="666666" if placeholder else "000000",
                    italic=placeholder,
                )

        def add_laf_foundation_overlay_fields(paragraph, values: dict[str, str], *, fill_values) -> None:
            """Place controls for the current LAF foundation official POA layout."""

            checkbox_visible = not _blank_pdf_overlay(fill_values)
            case_type = _text(values.get("CASE_TYPE", ""))
            role_marks = _text(values.get("ROLE_MARKS", ""))
            stage_marks = _text(values.get("STAGE_MARKS", ""))
            court_line_source = _text(values.get("COURT_LINE", ""))

            add_overlay_checkbox(
                paragraph,
                "case_type_civil",
                x_pt=67,
                y_pt=51,
                checked=not any(key in case_type for key in ("刑", "行政", "訴願")),
                visible=checkbox_visible,
            )
            add_overlay_checkbox(paragraph, "case_type_criminal", x_pt=67, y_pt=67, checked="刑" in case_type, visible=checkbox_visible)
            add_overlay_checkbox(
                paragraph,
                "case_type_administrative",
                x_pt=67,
                y_pt=83,
                checked="行政" in case_type and "訴願" not in case_type,
                visible=checkbox_visible,
            )
            add_overlay_checkbox(paragraph, "case_type_petition", x_pt=67, y_pt=99, checked="訴願" in case_type, visible=checkbox_visible)

            laf_case_number, laf_case_placeholder = _value_or_placeholder(
                values.get("LAF_CASE_NUMBER", ""),
                "請填法扶案號",
                fill_values,
            )
            add_overlay_textbox(
                paragraph,
                laf_case_number,
                field_name="laf_case_number",
                x_pt=365,
                y_pt=66,
                width_pt=120,
                height_pt=14,
                font_pt=9,
                align="left",
                color="666666" if laf_case_placeholder else "000000",
                italic=laf_case_placeholder,
            )

            court_case_number, court_case_placeholder = _value_or_placeholder(
                values.get("COURT_CASE_NUMBER", ""),
                "請填法院案號",
                fill_values,
            )
            add_overlay_textbox(
                paragraph,
                court_case_number,
                field_name="court_case_number",
                x_pt=350,
                y_pt=84,
                width_pt=190,
                height_pt=15,
                font_pt=8.5,
                color="666666" if court_case_placeholder else "000000",
                italic=court_case_placeholder,
            )

            client_name, client_placeholder = _value_or_placeholder(
                values.get("CLIENT_NAME", ""),
                "請填當事人",
                fill_values,
            )
            add_overlay_textbox(
                paragraph,
                client_name,
                field_name="client_name",
                x_pt=220,
                y_pt=140,
                width_pt=110,
                height_pt=17,
                font_pt=10.5,
                color="666666" if client_placeholder else "000000",
                italic=client_placeholder,
            )

            birthday, birthday_placeholder = _value_or_placeholder(
                values.get("CLIENT_BIRTHDAY", ""),
                "請填生日",
                fill_values,
            )
            add_overlay_textbox(
                paragraph,
                birthday,
                field_name="client_birthday",
                x_pt=210,
                y_pt=161,
                width_pt=126,
                height_pt=17,
                font_pt=10,
                color="666666" if birthday_placeholder else "000000",
                italic=birthday_placeholder,
            )

            client_id, client_id_placeholder = _value_or_placeholder(
                values.get("CLIENT_ID", ""),
                "請填身分證",
                fill_values,
            )
            add_overlay_textbox(
                paragraph,
                client_id,
                field_name="client_id",
                x_pt=215,
                y_pt=184,
                width_pt=126,
                height_pt=17,
                font_pt=10,
                color="666666" if client_id_placeholder else "000000",
                italic=client_id_placeholder,
            )

            lawyer_name, lawyer_placeholder = _value_or_placeholder(
                values.get("LAWYER_NAME", ""),
                "請填律師姓名",
                fill_values,
            )
            add_overlay_textbox(
                paragraph,
                lawyer_name,
                field_name="lawyer_name",
                x_pt=425,
                y_pt=182,
                width_pt=110,
                height_pt=18,
                font_pt=10.5,
                color="666666" if lawyer_placeholder else "000000",
                italic=lawyer_placeholder,
            )

            address_phone, address_phone_placeholder = _value_or_placeholder(
                values.get("CLIENT_ADDRESS_PHONE", ""),
                "請填地址電話",
                fill_values,
            )
            add_overlay_textbox(
                paragraph,
                address_phone,
                field_name="client_address_phone",
                x_pt=150,
                y_pt=210,
                width_pt=178,
                height_pt=48,
                font_pt=8.5,
                align="left",
                color="666666" if address_phone_placeholder else "000000",
                italic=address_phone_placeholder,
            )

            address_parts = []
            if _text(values.get("LAW_FIRM_ADDRESS_LINE", "")):
                address_parts.append(_text(values["LAW_FIRM_ADDRESS_LINE"]))
            if _text(values.get("LAW_FIRM_PHONE", "")):
                address_parts.append(f"電話：{_text(values['LAW_FIRM_PHONE'])}")
            firm_text, firm_placeholder = _value_or_placeholder(
                "\n".join(address_parts),
                "請填事務所地址\n請填電話",
                fill_values,
            )
            add_overlay_textbox(
                paragraph,
                firm_text,
                field_name="law_firm_contact",
                x_pt=375,
                y_pt=292,
                width_pt=170,
                height_pt=64,
                font_pt=8.5,
                align="left",
                color="666666" if firm_placeholder else "000000",
                italic=firm_placeholder,
            )

            case_reason, case_reason_placeholder = _value_or_placeholder(
                values.get("CASE_REASON", ""),
                "請填案由",
                fill_values,
            )
            add_overlay_textbox(
                paragraph,
                case_reason,
                field_name="case_reason",
                x_pt=83,
                y_pt=405,
                width_pt=280,
                height_pt=18,
                font_pt=10.5,
                align="left",
                color="666666" if case_reason_placeholder else "000000",
                italic=case_reason_placeholder,
            )

            add_overlay_checkbox(paragraph, "stage_first", x_pt=62, y_pt=473, checked="■第 一 審" in stage_marks, visible=checkbox_visible)
            add_overlay_checkbox(paragraph, "stage_investigation", x_pt=62, y_pt=487, checked="■偵" in stage_marks, visible=checkbox_visible)
            add_overlay_checkbox(paragraph, "stage_second", x_pt=62, y_pt=501, checked="■第 二 審" in stage_marks, visible=checkbox_visible)
            add_overlay_checkbox(paragraph, "stage_third", x_pt=62, y_pt=515, checked="■第 三 審" in stage_marks, visible=checkbox_visible)

            add_overlay_checkbox(paragraph, "role_agent", x_pt=156, y_pt=466, checked="■代 理 人" in role_marks, visible=checkbox_visible)
            add_overlay_checkbox(paragraph, "role_complainant_agent", x_pt=156, y_pt=480, checked="■告訴代理人" in role_marks, visible=checkbox_visible)
            add_overlay_checkbox(paragraph, "role_defender", x_pt=156, y_pt=494, checked="■辯 護 人" in role_marks, visible=checkbox_visible)
            add_overlay_checkbox(paragraph, "role_assistant", x_pt=156, y_pt=508, checked="■輔 佐 人" in role_marks, visible=checkbox_visible)

            add_overlay_checkbox(paragraph, "right_settlement_yes", x_pt=256, y_pt=454, checked=False, visible=checkbox_visible)
            add_overlay_checkbox(paragraph, "right_settlement_no", x_pt=286, y_pt=454, checked=False, visible=checkbox_visible)
            add_overlay_checkbox(paragraph, "right_execution_yes", x_pt=256, y_pt=498, checked=False, visible=checkbox_visible)
            add_overlay_checkbox(paragraph, "right_execution_no", x_pt=286, y_pt=498, checked=False, visible=checkbox_visible)

            court_name = ""
            court_name = _court_target_from_line(court_line_source) or _text(values.get("COURT_TARGET", "")) or _text(values.get("COURT_NAME", ""))
            court_text, court_placeholder = _value_or_placeholder(
                court_name,
                "請填法院/檢署/委員會",
                fill_values,
            )
            if "檢" in court_line_source:
                court_y = 607
            elif "委員會" in court_line_source:
                court_y = 637
            else:
                court_y = 577
            add_overlay_checkbox(paragraph, "court_kind_court", x_pt=62, y_pt=578, checked="■" in court_line_source and "法院" in court_line_source, visible=checkbox_visible)
            add_overlay_checkbox(paragraph, "court_kind_prosecutor", x_pt=62, y_pt=608, checked="■" in court_line_source and "檢" in court_line_source, visible=checkbox_visible)
            add_overlay_checkbox(paragraph, "court_kind_transfer", x_pt=370, y_pt=608, checked="■轉呈" in court_line_source, visible=checkbox_visible)
            add_overlay_checkbox(paragraph, "court_kind_committee", x_pt=62, y_pt=638, checked="■" in court_line_source and "委員會" in court_line_source, visible=checkbox_visible)
            add_overlay_textbox(
                paragraph,
                court_text,
                field_name="court_name",
                x_pt=78,
                y_pt=court_y,
                width_pt=118,
                height_pt=18,
                font_pt=10.5,
                align="left",
                color="666666" if court_placeholder else "000000",
                italic=court_placeholder,
            )

            signer_client, signer_client_placeholder = _value_or_placeholder(values.get("CLIENT_NAME", ""), "請填委任人", fill_values)
            signer_lawyer, signer_lawyer_placeholder = _value_or_placeholder(values.get("LAWYER_NAME", ""), "請填受任人", fill_values)
            add_overlay_textbox(
                paragraph,
                signer_client,
                field_name="client_signature_name",
                x_pt=340,
                y_pt=684,
                width_pt=110,
                height_pt=18,
                font_pt=10,
                color="666666" if signer_client_placeholder else "000000",
                italic=signer_client_placeholder,
            )
            add_overlay_textbox(
                paragraph,
                signer_lawyer,
                field_name="lawyer_signature_name",
                x_pt=340,
                y_pt=718,
                width_pt=110,
                height_pt=18,
                font_pt=10,
                color="666666" if signer_lawyer_placeholder else "000000",
                italic=signer_lawyer_placeholder,
            )

            year, year_placeholder = _value_or_placeholder(values.get("ROC_YEAR", ""), "年", fill_values)
            month, month_placeholder = _value_or_placeholder(values.get("ROC_MONTH", ""), "月", fill_values)
            day, day_placeholder = _value_or_placeholder(values.get("ROC_DAY", ""), "日", fill_values)
            for text, field_name, x_pt, placeholder in (
                (year, "roc_year", 350, year_placeholder),
                (month, "roc_month", 410, month_placeholder),
                (day, "roc_day", 470, day_placeholder),
            ):
                add_overlay_textbox(
                    paragraph,
                    text,
                    field_name=field_name,
                    x_pt=x_pt,
                    y_pt=786,
                    width_pt=42,
                    height_pt=18,
                    font_pt=10.5,
                    color="666666" if placeholder else "000000",
                    italic=placeholder,
                )

        def add_criminal_overlay_fields(paragraph, values: dict[str, str], *, fill_values) -> None:
            """Place editable controls over the common criminal POA PDF layout."""

            def value(key: str, placeholder: str) -> tuple[str, bool]:
                return _value_or_placeholder(values.get(key, ""), placeholder, fill_values)

            court_case_number, court_case_placeholder = value("COURT_CASE_NUMBER", "請填案號")
            add_overlay_textbox(
                paragraph,
                court_case_number,
                field_name="criminal_case_number",
                x_pt=382,
                y_pt=55,
                width_pt=86,
                height_pt=18,
                font_pt=10,
                color="666666" if court_case_placeholder else "000000",
                italic=court_case_placeholder,
            )
            stock, stock_placeholder = value("COURT_CASE_STOCK", "股別")
            add_overlay_textbox(
                paragraph,
                stock,
                field_name="criminal_stock",
                x_pt=526,
                y_pt=55,
                width_pt=42,
                height_pt=18,
                font_pt=10,
                color="666666" if stock_placeholder else "000000",
                italic=stock_placeholder,
            )
            client_name, client_placeholder = value("CLIENT_NAME", "請填委任人")
            add_overlay_textbox(
                paragraph,
                client_name,
                field_name="criminal_client_name",
                x_pt=90,
                y_pt=181,
                width_pt=144,
                height_pt=22,
                font_pt=12,
                color="666666" if client_placeholder else "000000",
                italic=client_placeholder,
            )
            for field_name, placeholder, x_pt, width_pt in (
                ("criminal_client_age", "年齡", 248, 30),
                ("criminal_client_occupation", "職業", 277, 30),
                ("criminal_client_registry", "籍貫", 304, 30),
            ):
                text, is_placeholder = _value_or_placeholder("", placeholder, fill_values)
                add_overlay_textbox(
                    paragraph,
                    text,
                    field_name=field_name,
                    x_pt=x_pt,
                    y_pt=181,
                    width_pt=width_pt,
                    height_pt=22,
                    font_pt=9,
                    color="666666",
                    italic=is_placeholder,
                )
            address, address_placeholder = value("CLIENT_ADDRESS_PHONE", "請填住居所")
            add_overlay_textbox(
                paragraph,
                address,
                field_name="criminal_client_address",
                x_pt=340,
                y_pt=180,
                width_pt=200,
                height_pt=44,
                font_pt=9.5,
                color="666666" if address_placeholder else "000000",
                italic=address_placeholder,
            )
            lawyer_name, lawyer_placeholder = value("LAWYER_NAME", "請填受任律師")
            add_overlay_textbox(
                paragraph,
                lawyer_name,
                field_name="criminal_lawyer_name",
                x_pt=88,
                y_pt=306,
                width_pt=146,
                height_pt=22,
                font_pt=12,
                color="666666" if lawyer_placeholder else "000000",
                italic=lawyer_placeholder,
            )
            office_lines = [
                _text(values.get("LAW_FIRM_OFFICE_NAME", "")),
                f"地址：{_text(values.get('LAW_FIRM_ADDRESS_LINE', ''))}" if _text(values.get("LAW_FIRM_ADDRESS_LINE", "")) else "",
                f"電話：{_text(values.get('LAW_FIRM_PHONE', ''))}" if _text(values.get("LAW_FIRM_PHONE", "")) else "",
                f"傳真：{_text(values.get('LAW_FIRM_FAX', ''))}" if _text(values.get("LAW_FIRM_FAX", "")) else "",
            ]
            office_text, office_placeholder = _value_or_placeholder(
                "\n".join(line for line in office_lines if line),
                "請填事務所資料",
                fill_values,
            )
            add_overlay_textbox(
                paragraph,
                office_text,
                field_name="criminal_law_firm_contact",
                x_pt=262,
                y_pt=290,
                width_pt=282,
                height_pt=86,
                font_pt=10,
                align="left",
                color="666666" if office_placeholder else "000000",
                italic=office_placeholder,
            )
            case_reason, reason_placeholder = value("CASE_REASON", "請填案由")
            add_overlay_textbox(
                paragraph,
                case_reason,
                field_name="criminal_case_reason",
                x_pt=86,
                y_pt=432,
                width_pt=68,
                height_pt=20,
                font_pt=11,
                color="666666" if reason_placeholder else "000000",
                italic=reason_placeholder,
            )
            court_name = _text(values.get("COURT_NAME", "")) or _strip_court_target(values.get("COURT_LINE", ""))
            court_text, court_placeholder = _value_or_placeholder(court_name, "請填法院/檢署", fill_values)
            add_overlay_textbox(
                paragraph,
                court_text,
                field_name="criminal_court_name",
                x_pt=68,
                y_pt=529,
                width_pt=184,
                height_pt=20,
                font_pt=11,
                align="left",
                color="666666" if court_placeholder else "000000",
                italic=court_placeholder,
            )
            signer_client, signer_client_placeholder = value("CLIENT_NAME", "請填委任人")
            add_overlay_textbox(
                paragraph,
                "\u200b" if _blank_pdf_overlay(fill_values) else f"委任人：{signer_client}",
                field_name="criminal_client_signature_name",
                x_pt=334,
                y_pt=638,
                width_pt=150,
                height_pt=20,
                font_pt=11,
                color="666666" if signer_client_placeholder else "000000",
                italic=signer_client_placeholder,
            )
            signer_lawyer, signer_lawyer_placeholder = value("LAWYER_NAME", "請填受任人")
            add_overlay_textbox(
                paragraph,
                "\u200b" if _blank_pdf_overlay(fill_values) else f"受任人：{signer_lawyer}",
                field_name="criminal_lawyer_signature_name",
                x_pt=334,
                y_pt=701,
                width_pt=170,
                height_pt=20,
                font_pt=11,
                color="666666" if signer_lawyer_placeholder else "000000",
                italic=signer_lawyer_placeholder,
            )
            year, year_placeholder = _value_or_placeholder(values.get("ROC_YEAR", ""), "年", fill_values)
            month, month_placeholder = _value_or_placeholder(values.get("ROC_MONTH", ""), "月", fill_values)
            day, day_placeholder = _value_or_placeholder(values.get("ROC_DAY", ""), "日", fill_values)
            for text, field_name, x_pt, placeholder in (
                (year, "criminal_roc_year", 278, year_placeholder),
                (month, "criminal_roc_month", 374, month_placeholder),
                (day, "criminal_roc_day", 470, day_placeholder),
            ):
                add_overlay_textbox(
                    paragraph,
                    text,
                    field_name=field_name,
                    x_pt=x_pt,
                    y_pt=750,
                    width_pt=42,
                    height_pt=18,
                    font_pt=11,
                    color="666666" if placeholder else "000000",
                    italic=placeholder,
                )

        def inline_to_page_background(inline, width_pt: float, height_pt: float) -> None:
            inline.tag = qn("wp:anchor")
            for key, value in {
                "distT": "0",
                "distB": "0",
                "distL": "0",
                "distR": "0",
                "simplePos": "0",
                "relativeHeight": "251659264",
                "behindDoc": "1",
                "locked": "0",
                "layoutInCell": "1",
                "allowOverlap": "1",
            }.items():
                inline.set(key, value)

            extent = inline.find(qn("wp:extent"))
            if extent is not None:
                extent.set("cx", str(emu(width_pt)))
                extent.set("cy", str(emu(height_pt)))

            simple_pos = OxmlElement("wp:simplePos")
            simple_pos.set("x", "0")
            simple_pos.set("y", "0")
            pos_h = OxmlElement("wp:positionH")
            pos_h.set("relativeFrom", "page")
            h_offset = OxmlElement("wp:posOffset")
            h_offset.text = "0"
            pos_h.append(h_offset)
            pos_v = OxmlElement("wp:positionV")
            pos_v.set("relativeFrom", "page")
            v_offset = OxmlElement("wp:posOffset")
            v_offset.text = "0"
            pos_v.append(v_offset)

            inline.insert(0, pos_v)
            inline.insert(0, pos_h)
            inline.insert(0, simple_pos)

            extent = inline.find(qn("wp:extent"))
            extent_index = list(inline).index(extent) if extent is not None else 2
            effect = OxmlElement("wp:effectExtent")
            for key in ("l", "t", "r", "b"):
                effect.set(key, "0")
            wrap = OxmlElement("wp:wrapNone")
            inline.insert(extent_index + 1, effect)
            inline.insert(extent_index + 2, wrap)

        def add_typing_page(
            doc,
            section,
            image_path: Path,
            width_pt: float,
            height_pt: float,
            *,
            overlay_values: dict[str, str] | None = None,
            fill_values=False,
            poa_layout: str = "",
        ) -> None:
            section.page_width = Pt(width_pt)
            section.page_height = Pt(height_pt)
            section.top_margin = Mm(0)
            section.bottom_margin = Mm(0)
            section.left_margin = Mm(0)
            section.right_margin = Mm(0)
            section.header_distance = Mm(0)
            section.footer_distance = Mm(0)

            rows, cols = 84, 60
            table = doc.add_table(rows=rows, cols=cols)
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            table.autofit = False
            set_table_borders_none(table)

            col_width = width_pt / cols
            row_height = max(1.0, (height_pt - 2.0) / rows)
            set_table_fixed_geometry(table, width_pt, col_width, cols)
            for row in table.rows:
                row.height = Pt(row_height)
                row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
                for cell in row.cells:
                    cell.width = Pt(col_width)
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                    set_cell_margins_zero(cell)
                    for paragraph in cell.paragraphs:
                        clear_paragraph(paragraph)

            paragraph = table.cell(0, 0).paragraphs[0]
            inline_shape = paragraph.add_run().add_picture(
                str(image_path),
                width=Pt(width_pt),
                height=Pt(height_pt),
            )
            inline_to_page_background(inline_shape._inline, width_pt, height_pt)
            if overlay_values is not None:
                overlay_host = {
                    "table": table,
                    "rows": rows,
                    "cols": cols,
                    "row_height": row_height,
                    "col_width": col_width,
                }
                if poa_layout == "criminal":
                    add_criminal_overlay_fields(overlay_host, overlay_values, fill_values=fill_values)
                elif poa_layout == "laf_foundation":
                    add_laf_foundation_overlay_fields(overlay_host, overlay_values, fill_values=fill_values)
                else:
                    add_fillable_overlay_fields(overlay_host, overlay_values, fill_values=fill_values)

        def build_template_document(*, fill_values) -> tuple[Any, int]:
            doc = Document()
            first_section = doc.sections[0]
            values = _template_values(normalized_metadata)
            with fitz.open(str(source)) as pdf, tempfile.TemporaryDirectory(prefix="magi_laf_poa_") as tmpdir:
                if len(pdf) == 0:
                    return doc, 0

                page_count = 0
                for idx, page in enumerate(pdf):
                    section = first_section if idx == 0 else doc.add_section(WD_SECTION_START.NEW_PAGE)
                    rect = page.rect

                    try:
                        scale = float(os.environ.get("MAGI_LAF_POA_RENDER_SCALE", "2") or "2")
                    except Exception:
                        scale = 2.0
                    scale = max(1.0, min(scale, 2.5))
                    pix = page.get_pixmap(matrix=fitz.Matrix(scale, scale), alpha=False)
                    img_path = Path(tmpdir) / f"page_{idx + 1}.png"
                    pix.save(str(img_path))
                    add_typing_page(
                        doc,
                        section,
                        img_path,
                        float(rect.width),
                        float(rect.height),
                        overlay_values=values if idx == 0 else None,
                        fill_values=fill_values,
                        poa_layout=normalized_metadata.get("poa_layout", ""),
                    )
                    page_count += 1
            return doc, page_count

        def add_case_metadata_page(doc) -> None:
            if not any(normalized_metadata.values()):
                return
            values = _template_values(normalized_metadata)
            section = doc.add_section(WD_SECTION_START.NEW_PAGE)
            section.top_margin = Mm(18)
            section.bottom_margin = Mm(18)
            section.left_margin = Mm(20)
            section.right_margin = Mm(20)
            title = doc.add_paragraph()
            title_run = title.add_run(_laf_poa_template_title(template_key))
            title_run.bold = True
            title_run.font.size = Pt(14)
            info = [
                ("分會", values["BRANCH_LABEL"]),
                ("分會電話", values["BRANCH_PHONE"]),
                ("法扶案號", values["LAF_CASE_NUMBER"]),
                ("法院案號", values["COURT_CASE_NUMBER"]),
                ("當事人", values["CLIENT_NAME"]),
                ("出生年月日", values["CLIENT_BIRTHDAY"]),
                ("身分證字號", values["CLIENT_ID"]),
                ("案件類型", normalized_metadata["case_type"]),
                ("案由", values["CASE_REASON"]),
                ("法院/檢署", values["COURT_LINE"]),
                ("審級", values["STAGE_MARKS"]),
                ("身分", values["ROLE_MARKS"]),
                ("律師", values["LAWYER_NAME"]),
                ("事務所", values["LAW_FIRM_OFFICE_NAME"]),
                ("事務所地址", values["LAW_FIRM_ADDRESS_LINE"]),
                ("事務所電話", values["LAW_FIRM_PHONE"]),
                ("事務所傳真", values["LAW_FIRM_FAX"]),
                ("律師手機", values["LAW_FIRM_MOBILE"]),
                ("填寫日期", f"{values['ROC_YEAR']} 年 {values['ROC_MONTH']} 月 {values['ROC_DAY']} 日"),
            ]
            table = doc.add_table(rows=0, cols=2)
            table.autofit = True
            for label, value in info:
                if not value:
                    continue
                cells = table.add_row().cells
                cells[0].text = label
                cells[1].text = value

        target.parent.mkdir(parents=True, exist_ok=True)
        from docx import Document as _Document  # type: ignore
        if exact_pdf_layout:
            _remove_stale_laf_poa_template(template_target)
            page_count = 0
        else:
            template_doc, page_count = build_template_document(fill_values=False)
            if page_count == 0:
                result.update(status="empty_pdf", error="pdf_has_no_pages")
                return result

            tmp_template = template_target.with_suffix(template_target.suffix + ".tmp")
            template_doc.save(str(tmp_template))
            _Document(str(tmp_template))
            os.replace(tmp_template, template_target)

        case_doc, case_page_count = build_template_document(fill_values="blank" if exact_pdf_layout else True)
        if case_page_count == 0:
            result.update(status="empty_pdf", error="pdf_has_no_pages")
            return result
        if page_count == 0:
            page_count = case_page_count
        if not exact_pdf_layout:
            add_case_metadata_page(case_doc)
        tmp_docx = target.with_suffix(target.suffix + ".tmp")
        case_doc.save(str(tmp_docx))
        _Document(str(tmp_docx))
        os.replace(tmp_docx, target)
        _apply_quality_gate(
            result,
            target,
            normalized_metadata,
            _template_values(normalized_metadata),
            check_expected_values=not exact_pdf_layout,
        )

        result.update(
            ok=True,
            status="created",
            pages=page_count,
            filled_fields={k: v for k, v in normalized_metadata.items() if v},
            template_key=template_key,
            branch_label=normalized_metadata.get("branch", ""),
            branch_phone=normalized_metadata.get("branch_phone", ""),
            template_values={k: v for k, v in _template_values(normalized_metadata).items() if v},
        )
        if normalized_metadata.get("branch_phone") == "待確認":
            result["warnings"].append("missing_branch_phone")
        return result
    except Exception as exc:
        result.update(status="failed", error=str(exc))
        return result
