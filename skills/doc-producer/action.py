#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
doc-producer/action.py
書狀製作技能：DOCX→PDF 轉換、正本/副本/繕本標記、PDF 合併
"""

import argparse
import json
import logging
import os
import subprocess
import sys
import tempfile
from pathlib import Path
from typing import Dict, List, Optional

_MAGI_ROOT = Path(__file__).resolve().parents[2]
if str(_MAGI_ROOT) not in sys.path:
    sys.path.insert(0, str(_MAGI_ROOT))

logger = logging.getLogger("doc-producer")

# ── LibreOffice discovery ──

_SOFFICE_CANDIDATES = [
    os.environ.get("MAGI_SOFFICE_PATH", ""),
    "/Applications/LibreOffice.app/Contents/MacOS/soffice",
    "/opt/homebrew/bin/soffice",
    "/usr/bin/soffice",
    "/usr/local/bin/soffice",
]


def _find_soffice():
    # type: () -> Optional[str]
    for p in _SOFFICE_CANDIDATES:
        if p and os.path.isfile(p) and os.access(p, os.X_OK):
            return p
    return None


# ── DOCX → PDF conversion ──

def convert_docx_to_pdf(input_path, output_path=None):
    # type: (str, Optional[str]) -> Dict
    """Convert a DOCX file to PDF using LibreOffice headless."""
    input_path = os.path.abspath(input_path)
    if not os.path.isfile(input_path):
        return {"success": False, "output": "", "error": "輸入檔案不存在: %s" % input_path}

    ext = os.path.splitext(input_path)[1].lower()
    if ext not in (".docx", ".doc", ".odt", ".rtf"):
        return {"success": False, "output": "", "error": "不支援的檔案格式: %s" % ext}

    if output_path:
        outdir = os.path.dirname(os.path.abspath(output_path))
    else:
        outdir = os.path.dirname(input_path)
        output_path = os.path.splitext(input_path)[0] + ".pdf"

    os.makedirs(outdir, exist_ok=True)

    soffice = _find_soffice()
    if soffice:
        return _convert_via_libreoffice(soffice, input_path, outdir, output_path)

    # Fallback: try docx2pdf (requires MS Word on macOS)
    return _convert_via_docx2pdf(input_path, output_path)


def _convert_via_libreoffice(soffice, input_path, outdir, output_path):
    # type: (str, str, str, str) -> Dict
    """Use LibreOffice headless to convert."""
    # Use a temp profile to avoid locking issues with concurrent conversions
    with tempfile.TemporaryDirectory(prefix="lo_profile_") as profile_dir:
        cmd = [
            soffice,
            "--headless",
            "--norestore",
            "-env:UserInstallation=file://%s" % profile_dir,
            "--convert-to", "pdf",
            "--outdir", outdir,
            input_path,
        ]
        try:
            proc = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=120,
            )
        except subprocess.TimeoutExpired:
            return {"success": False, "output": "", "error": "LibreOffice 轉換逾時 (120s)"}
        except Exception as e:
            return {"success": False, "output": "", "error": "LibreOffice 執行失敗: %s" % str(e)}

    # LibreOffice outputs to outdir with the same basename but .pdf
    lo_output = os.path.join(outdir, os.path.splitext(os.path.basename(input_path))[0] + ".pdf")

    if not os.path.isfile(lo_output):
        stderr_snippet = (proc.stderr or "")[:500]
        return {
            "success": False,
            "output": "",
            "error": "LibreOffice 轉換完成但找不到輸出檔案。stderr: %s" % stderr_snippet,
        }

    # If desired output_path differs from lo_output, rename
    if os.path.abspath(lo_output) != os.path.abspath(output_path):
        os.rename(lo_output, output_path)

    return {"success": True, "output": output_path, "error": ""}


def _convert_via_docx2pdf(input_path, output_path):
    # type: (str, str) -> Dict
    """Fallback: use docx2pdf package."""
    try:
        from docx2pdf import convert as _d2p_convert
        _d2p_convert(input_path, output_path)
        if os.path.isfile(output_path):
            return {"success": True, "output": output_path, "error": ""}
        return {"success": False, "output": "", "error": "docx2pdf 轉換完成但找不到輸出檔案"}
    except ImportError:
        return {
            "success": False,
            "output": "",
            "error": "找不到 LibreOffice 也找不到 docx2pdf，無法轉換 DOCX→PDF",
        }
    except Exception as e:
        return {"success": False, "output": "", "error": "docx2pdf 轉換失敗: %s" % str(e)}


# ── PDF copy-type marking ──


def _add_stamp_image_to_last_page(doc, stamp_image, stamp_center=None):
    # type: (object, Optional[str], Optional[dict]) -> None
    """Place the configured lawyer stamp image near the signature area."""
    if not stamp_image or not os.path.isfile(stamp_image) or len(doc) < 1:
        return
    import fitz

    stamp_path = stamp_image
    temp_stamp = None
    try:
        try:
            from PIL import Image

            img = Image.open(stamp_image).convert("RGBA")
            pixels = []
            for r, g, b, a in img.getdata():
                if r > 245 and g > 245 and b > 245:
                    pixels.append((r, g, b, 0))
                else:
                    pixels.append((r, g, b, int(a * 0.62)))
            img.putdata(pixels)
            fd, temp_stamp = tempfile.mkstemp(prefix="magi_stamp_", suffix=".png")
            os.close(fd)
            img.save(temp_stamp)
            stamp_path = temp_stamp
        except Exception:
            stamp_path = stamp_image

        page = doc[-1]
        width = 82.0
        try:
            from PIL import Image

            with Image.open(stamp_image) as raw:
                aspect = raw.height / raw.width if raw.width else 1.0
        except Exception:
            aspect = 1.0
        height = max(48.0, width * aspect)
        center_x = page.rect.width - 144
        center_y = page.rect.height - 144
        if isinstance(stamp_center, dict):
            try:
                center_x = float(stamp_center.get("x"))
                center_y = float(stamp_center.get("y"))
            except Exception:
                center_x = page.rect.width - 144
                center_y = page.rect.height - 144
        rect = fitz.Rect(
            center_x - width / 2,
            center_y - height / 2,
            center_x + width / 2,
            center_y + height / 2,
        )
        page.insert_image(rect, filename=stamp_path, overlay=True)
    finally:
        if temp_stamp:
            try:
                os.unlink(temp_stamp)
            except OSError:
                pass


def mark_copy_type(input_pdf, output_pdf=None, copy_type="正本", add_poa=False, add_sent_to_opponent=False, stamp_image=None, stamp_center=None):
    # type: (str, Optional[str], str, bool, bool, Optional[str], Optional[dict]) -> Dict
    """Add copy-type label (正本/副本/繕本) to top-right corner of first page."""
    import fitz

    input_pdf = os.path.abspath(input_pdf)
    if not os.path.isfile(input_pdf):
        return {"success": False, "output": "", "error": "輸入 PDF 不存在: %s" % input_pdf}

    valid_types = ("正本", "副本", "繕本", "留底")
    if copy_type not in valid_types:
        return {
            "success": False,
            "output": "",
            "error": "copy_type 必須為 %s 之一，收到: %s" % (str(valid_types), copy_type),
        }

    if not output_pdf:
        base, ext = os.path.splitext(input_pdf)
        output_pdf = "%s_%s%s" % (base, copy_type, ext)

    try:
        doc = fitz.open(input_pdf)
    except Exception as e:
        return {"success": False, "output": "", "error": "無法開啟 PDF: %s" % str(e)}

    try:
        page = doc[0]
        page_width = page.rect.width

        # Position: top-right area
        margin_right = 30
        font_main = "china-ss"
        font_size_main = 16
        font_size_sub = 12
        line_gap = 6

        # Calculate text width for right-alignment
        x_start = page_width - margin_right
        y_pos = 40.0

        # Main label (copy_type)
        _tw_main = fitz.get_text_length(copy_type, fontname=font_main, fontsize=font_size_main)
        page.insert_text(
            (x_start - _tw_main, y_pos),
            copy_type,
            fontname=font_main,
            fontsize=font_size_main,
            color=(0, 0, 0),
        )
        y_pos += font_size_main + line_gap

        # Sub-label: 附委任狀
        if add_poa:
            sub_text = "附委任狀"
            _tw_sub = fitz.get_text_length(sub_text, fontname=font_main, fontsize=font_size_sub)
            page.insert_text(
                (x_start - _tw_sub, y_pos),
                sub_text,
                fontname=font_main,
                fontsize=font_size_sub,
                color=(0, 0, 0),
            )
            y_pos += font_size_sub + line_gap

        # Sub-label: 繕本已送對造
        if add_sent_to_opponent:
            sub_text2 = "繕本已送對造"
            _tw_sub2 = fitz.get_text_length(sub_text2, fontname=font_main, fontsize=font_size_sub)
            page.insert_text(
                (x_start - _tw_sub2, y_pos),
                sub_text2,
                fontname=font_main,
                fontsize=font_size_sub,
                color=(0, 0, 0),
            )

        _add_stamp_image_to_last_page(doc, stamp_image, stamp_center=stamp_center)

        os.makedirs(os.path.dirname(os.path.abspath(output_pdf)), exist_ok=True)
        doc.save(output_pdf)
        doc.close()

        return {"success": True, "output": output_pdf, "error": ""}

    except Exception as e:
        doc.close()
        return {"success": False, "output": "", "error": "標記失敗: %s" % str(e)}


# ── PDF merging ──

def merge_pdfs(input_paths, output_path):
    # type: (List[str], str) -> Dict
    """Merge multiple PDFs into one."""
    import fitz

    if not input_paths:
        return {"success": False, "output": "", "page_count": 0, "error": "未提供任何輸入檔案"}

    # Validate all inputs exist
    for p in input_paths:
        if not os.path.isfile(p):
            return {"success": False, "output": "", "page_count": 0, "error": "檔案不存在: %s" % p}

    output_path = os.path.abspath(output_path)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    try:
        merged = fitz.open()
        total_pages = 0

        for pdf_path in input_paths:
            src = fitz.open(pdf_path)
            merged.insert_pdf(src)
            total_pages += len(src)
            src.close()

        merged.save(output_path)
        merged.close()

        return {
            "success": True,
            "output": output_path,
            "page_count": total_pages,
            "error": "",
        }

    except Exception as e:
        return {"success": False, "output": "", "page_count": 0, "error": "合併失敗: %s" % str(e)}


# ── Full produce pipeline ──

def produce(input_path, copy_type="正本", add_poa=False, add_sent_to_opponent=False, merge_with=None, output_dir=None, stamp_image=None, stamp_center=None):
    # type: (str, str, bool, bool, Optional[List[str]], Optional[str], Optional[str], Optional[dict]) -> Dict
    """Full pipeline: convert DOCX→PDF → mark copy type → optionally merge."""
    input_path = os.path.abspath(input_path)
    if not os.path.isfile(input_path):
        return {"success": False, "outputs": {}, "error": "輸入檔案不存在: %s" % input_path}

    if output_dir:
        os.makedirs(output_dir, exist_ok=True)

    outputs = {}  # type: Dict[str, str]
    ext = os.path.splitext(input_path)[1].lower()

    # Step 1: Convert if DOCX
    if ext in (".docx", ".doc", ".odt", ".rtf"):
        if output_dir:
            pdf_name = os.path.splitext(os.path.basename(input_path))[0] + ".pdf"
            pdf_path = os.path.join(output_dir, pdf_name)
        else:
            pdf_path = None  # defaults to same dir
        result = convert_docx_to_pdf(input_path, pdf_path)
        if not result["success"]:
            return {"success": False, "outputs": {}, "error": "轉檔失敗: %s" % result["error"]}
        pdf_path = result["output"]
        outputs["pdf"] = pdf_path
    elif ext == ".pdf":
        pdf_path = input_path
        outputs["pdf"] = pdf_path
    else:
        return {"success": False, "outputs": {}, "error": "不支援的輸入格式: %s" % ext}

    # Step 2: Mark copy type
    if output_dir:
        marked_name = "%s_%s.pdf" % (os.path.splitext(os.path.basename(pdf_path))[0], copy_type)
        marked_path = os.path.join(output_dir, marked_name)
    else:
        marked_path = None  # defaults to auto-naming

    result = mark_copy_type(
        pdf_path,
        output_pdf=marked_path,
        copy_type=copy_type,
        add_poa=add_poa,
        add_sent_to_opponent=add_sent_to_opponent,
        stamp_image=stamp_image,
        stamp_center=stamp_center,
    )
    if not result["success"]:
        return {"success": False, "outputs": outputs, "error": "標記失敗: %s" % result["error"]}
    outputs["marked"] = result["output"]

    # Step 3: Merge (optional)
    if merge_with:
        all_pdfs = [result["output"]] + [os.path.abspath(p) for p in merge_with]
        if output_dir:
            merged_name = "%s_%s_合併.pdf" % (os.path.splitext(os.path.basename(input_path))[0], copy_type)
            merged_path = os.path.join(output_dir, merged_name)
        else:
            base = os.path.splitext(result["output"])[0]
            merged_path = "%s_合併.pdf" % base

        merge_result = merge_pdfs(all_pdfs, merged_path)
        if not merge_result["success"]:
            return {"success": False, "outputs": outputs, "error": "合併失敗: %s" % merge_result["error"]}
        outputs["merged"] = merge_result["output"]

    return {"success": True, "outputs": outputs, "error": ""}


# ── Self test ──

def _self_test():
    # type: () -> Dict
    """Create a test DOCX, convert, mark, merge, verify outputs."""
    import fitz

    with tempfile.TemporaryDirectory(prefix="doc_producer_test_") as tmpdir:
        results = {}  # type: Dict[str, object]

        # -- Test 1: Create a minimal DOCX for conversion test --
        test_docx = os.path.join(tmpdir, "test_document.docx")
        soffice = _find_soffice()
        has_lo = soffice is not None
        results["libreoffice_found"] = has_lo

        if has_lo:
            # Create a simple DOCX using python-docx if available, otherwise skip
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument()
                doc.add_heading("測試書狀", level=1)
                doc.add_paragraph("本件為自動測試用之書狀範本。")
                doc.save(test_docx)
                results["docx_created"] = True
            except ImportError:
                # No python-docx, create a minimal ODT-like file via LibreOffice
                # Just test with a real PDF instead
                results["docx_created"] = False
                results["note"] = "python-docx 未安裝，跳過 DOCX 轉換測試"

            if results.get("docx_created"):
                conv = convert_docx_to_pdf(test_docx)
                results["convert"] = conv
                if not conv["success"]:
                    return {"success": False, "results": results, "error": "轉檔測試失敗: %s" % conv["error"]}
        else:
            results["note"] = "LibreOffice 未安裝，跳過 DOCX 轉換測試"

        # -- Test 2: Create a test PDF and mark it --
        test_pdf = os.path.join(tmpdir, "test_mark.pdf")
        doc = fitz.open()
        page = doc.new_page(width=595, height=842)  # A4
        page.insert_text((72, 100), "測試書狀內容", fontname="china-ss", fontsize=14)
        doc.save(test_pdf)
        doc.close()

        mark_result = mark_copy_type(test_pdf, copy_type="正本", add_poa=True, add_sent_to_opponent=True)
        results["mark"] = mark_result
        if not mark_result["success"]:
            return {"success": False, "results": results, "error": "標記測試失敗: %s" % mark_result["error"]}

        # Verify marked file exists and is valid PDF
        marked_path = mark_result["output"]
        if not os.path.isfile(marked_path):
            return {"success": False, "results": results, "error": "標記後的 PDF 不存在"}

        marked_doc = fitz.open(marked_path)
        marked_text = marked_doc[0].get_text()
        marked_doc.close()
        results["mark_text_check"] = {
            "has_copy_type": "正本" in marked_text,
            "has_poa": "附委任狀" in marked_text,
            "has_sent_to_opponent": "繕本已送對造" in marked_text,
        }

        # -- Test 3: Merge test --
        test_pdf2 = os.path.join(tmpdir, "test_merge_extra.pdf")
        doc2 = fitz.open()
        page2 = doc2.new_page(width=595, height=842)
        page2.insert_text((72, 100), "附件內容", fontname="china-ss", fontsize=14)
        doc2.save(test_pdf2)
        doc2.close()

        merge_output = os.path.join(tmpdir, "merged.pdf")
        merge_result = merge_pdfs([marked_path, test_pdf2], merge_output)
        results["merge"] = merge_result
        if not merge_result["success"]:
            return {"success": False, "results": results, "error": "合併測試失敗: %s" % merge_result["error"]}

        if merge_result["page_count"] != 2:
            return {
                "success": False,
                "results": results,
                "error": "合併後頁數不正確，預期 2，實際 %d" % merge_result["page_count"],
            }

        # -- Test 4: Full produce pipeline (with existing PDF) --
        produce_result = produce(
            test_pdf,
            copy_type="副本",
            add_poa=False,
            add_sent_to_opponent=True,
            merge_with=[test_pdf2],
            output_dir=tmpdir,
        )
        results["produce"] = produce_result
        if not produce_result["success"]:
            return {"success": False, "results": results, "error": "produce 流程測試失敗: %s" % produce_result["error"]}

        return {"success": True, "results": results, "error": ""}


# ── CLI entry ──

def main():
    parser = argparse.ArgumentParser(description="doc-producer skill")
    parser.add_argument("--task", required=True, help="Task: convert/mark/merge/produce/self_test")
    parser.add_argument("payload", nargs="?", default="{}", help="JSON payload")
    args = parser.parse_args()

    task = args.task.strip()

    # self_test shortcut
    if task == "self_test":
        result = _self_test()
        print(json.dumps(result, ensure_ascii=False, indent=2))
        return

    # Parse payload (can be embedded in --task or as positional arg)
    payload_str = args.payload
    if " " in task:
        parts = task.split(" ", 1)
        task = parts[0]
        payload_str = parts[1]

    try:
        payload = json.loads(payload_str)
    except json.JSONDecodeError as e:
        print(json.dumps({"success": False, "error": "JSON 解析失敗: %s" % str(e)}, ensure_ascii=False))
        return

    if task == "convert":
        result = convert_docx_to_pdf(
            payload.get("input", ""),
            payload.get("output"),
        )
    elif task == "mark":
        result = mark_copy_type(
            payload.get("input", ""),
            output_pdf=payload.get("output"),
            copy_type=payload.get("copy_type", "正本"),
            add_poa=payload.get("add_poa", False),
            add_sent_to_opponent=payload.get("add_sent_to_opponent", False),
            stamp_image=payload.get("stamp_image"),
            stamp_center=payload.get("stamp_center"),
        )
    elif task == "merge":
        result = merge_pdfs(
            payload.get("inputs", []),
            payload.get("output", ""),
        )
    elif task == "produce":
        result = produce(
            payload.get("input", ""),
            copy_type=payload.get("copy_type", "正本"),
            add_poa=payload.get("add_poa", False),
            add_sent_to_opponent=payload.get("add_sent_to_opponent", False),
            merge_with=payload.get("merge_with"),
            output_dir=payload.get("output_dir"),
            stamp_image=payload.get("stamp_image"),
            stamp_center=payload.get("stamp_center"),
        )
    else:
        result = {"success": False, "error": "未知 task: %s" % task}

    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
