# -*- coding: utf-8 -*-
"""
export_docx.py
==============
將翻譯、逐字稿、摘要等結構化內容輸出成 docx 表格，
並儲存至 /static/exports，方便透過 LINE/DC 傳送連結或檔案。

支援三種表格模式：
  1. bilingual  — 雙語對照表（翻譯用）
  2. transcript — 逐字稿表格（發言人｜時間｜內容）
  3. summary    — 摘要表格（段落｜摘要｜原文節錄）

V2/source 優先使用 docx-js；sealed V3 使用已由 runtime snapshot
雜湊綁定的 python-docx，避免依賴主機全域 Node 模組。
"""

from __future__ import annotations

import hashlib
import io
import json
import logging
import os
import re
import stat
import subprocess
import tempfile
import time
import uuid
import zipfile
from pathlib import Path
from typing import Any, Dict, List, Optional

_MAGI_ROOT = Path(__file__).resolve().parents[2]
_DEFAULT_EXPORTS_DIR = _MAGI_ROOT / "static" / "exports"
_EXPORTS_DIR = os.environ.get("MAGI_EXPORTS_DIR", str(_DEFAULT_EXPORTS_DIR))

# Reuse public URL logic from export_text
try:
    from skills.ops.export_text import _load_public_base_url
except Exception:
    def _load_public_base_url() -> str:
        return ""


def _find_node() -> str:
    """Find node binary."""
    for p in ["/opt/homebrew/bin/node", "/usr/local/bin/node", "/usr/bin/node"]:
        if os.path.exists(p):
            return p
    return "node"


def _find_node_path() -> str:
    """Find NODE_PATH for docx module."""
    candidates = [
        str(_MAGI_ROOT / "node_modules"),
        str(Path(__file__).resolve().parent / "node_modules"),
        "/opt/homebrew/lib/node_modules",
        "/usr/local/lib/node_modules",
        "/usr/lib/node_modules",
    ]
    for c in candidates:
        docx_path = os.path.join(c, "docx")
        if os.path.isdir(docx_path):
            return c
    return ""


def _sealed_release_context() -> bool:
    """Return whether this module is executing from a sealed V3 release."""
    return bool(
        str(os.environ.get("MAGI_V3_RELEASE_ID") or "").strip()
        or str(os.environ.get("MAGI_V3_DEPLOYMENT_MODE") or "").strip()
        or str(os.environ.get("MAGI_V3_RELEASE_MANIFEST") or "").strip()
        or (_MAGI_ROOT / "release-manifest.json").is_file()
        or (_MAGI_ROOT / "RELEASE_COMPLETE.json").is_file()
    )


def _node_backend() -> tuple[str, str] | None:
    """Resolve the legacy Node backend only in an unsealed source/V2 runtime.

    V3 releases intentionally exclude ``node_modules`` and do not declare a
    manifest-bound Node runtime.  Falling back directly prevents a sealed
    candidate from silently importing Homebrew/global Node or docx modules.
    """
    if _sealed_release_context():
        return None
    return _find_node(), _find_node_path()


def _exports_dir() -> Path:
    return Path(os.environ.get("MAGI_EXPORTS_DIR", str(_DEFAULT_EXPORTS_DIR))).expanduser().resolve(strict=False)


def _is_relative_to(path: Path, root: Path) -> bool:
    try:
        path.relative_to(root)
        return True
    except ValueError:
        return False


def _safe_generated_filename(prefix: str) -> str:
    safe_prefix = "".join(ch if (ch.isalnum() or ch in {"_", "-"}) else "_" for ch in str(prefix or "export"))
    safe_prefix = safe_prefix.strip("._-") or "export"
    stamp = time.strftime("%Y%m%d_%H%M%S")
    token = uuid.uuid4().hex[:8]
    return f"{safe_prefix}_{stamp}_{token}.docx"


def _resolve_export_docx_path(filename: str) -> tuple[Path, str]:
    name = str(filename or "").strip()
    if not name:
        raise ValueError("empty filename")
    candidate = Path(name)
    if candidate.is_absolute() or candidate.name != name or "\\" in name or ".." in candidate.parts:
        raise ValueError("filename must be a plain .docx basename")
    if candidate.suffix.lower() != ".docx":
        raise ValueError("filename must end with .docx")
    exports_dir = _exports_dir()
    exports_dir.mkdir(parents=True, exist_ok=True)
    out_path = (exports_dir / candidate.name).resolve(strict=False)
    if not _is_relative_to(out_path, exports_dir):
        raise ValueError("resolved DOCX path escapes exports directory")
    return out_path, candidate.name


def _validate_docx_file(path: Path, contract: dict) -> dict:
    flags = os.O_RDONLY | getattr(os, "O_NOFOLLOW", 0)
    try:
        descriptor = os.open(path, flags)
    except OSError as exc:
        return {"ok": False, "error": f"docx file open failed:{type(exc).__name__}"}
    try:
        return _validate_docx_descriptor(descriptor, contract)
    finally:
        os.close(descriptor)


_XML_CONTROL_RE = re.compile(r"[\x00-\x08\x0B\x0C\x0E-\x1F]")
# V3 is a macOS runtime; use the system Traditional-Chinese family so Word,
# Pages, and Quick Look do not depend on an optional Microsoft font install.
_FALLBACK_FONT_CJK = "PingFang TC"
_FALLBACK_ACCENT = "1F4E79"
_FALLBACK_HEADER_FILL = "1F4E79"
_FALLBACK_ROW_EVEN = "FFFFFF"
_FALLBACK_ROW_ODD = "F0F4F8"
_FALLBACK_BORDER = "D0D5DD"


class _ExportDirectoryChangedError(ValueError):
    """Raised when an anchored exports directory no longer owns its path."""


def _xml_text(value: Any) -> str:
    return _XML_CONTROL_RE.sub("", str(value or ""))


def _normalized_docx_text(value: Any) -> str:
    return re.sub(r"\s+", " ", _xml_text(value)).strip()


def _docx_contract(data: dict) -> dict:
    mode = str(data.get("mode") or "")
    collection_key = {
        "bilingual": "pages",
        "transcript": "segments",
        "summary": "sections",
    }.get(mode)
    if collection_key is None:
        raise ValueError(f"unsupported DOCX mode:{mode}")
    rows = data.get(collection_key)
    if not isinstance(rows, list) or not rows:
        raise ValueError(f"{collection_key} must be a non-empty list")

    required: list[str] = []
    expected_cells: list[list[str]] = []

    def require(value: Any) -> None:
        normalized = _normalized_docx_text(value)
        if normalized and normalized not in required:
            required.append(normalized)

    require(data.get("title"))
    if mode == "bilingual":
        require(data.get("subtitle"))
        require(data.get("header_text") or data.get("title") or "MAGI 文件")
        hide_page = data.get("hide_page_column") is True
        labels = data.get("col_labels") if isinstance(data.get("col_labels"), dict) else {}
        headers = [labels.get("col2") or "原文", labels.get("col3") or "翻譯"]
        if not hide_page:
            headers.insert(0, labels.get("col1") or "頁碼")
        expected_cells.append([_normalized_docx_text(value) for value in headers])
        for header in headers:
            require(header)
        for index, item in enumerate(rows):
            if not isinstance(item, dict):
                raise ValueError("bilingual rows must be objects")
            if not hide_page:
                require(item.get("page") or index + 1)
            require(item.get("source"))
            require(item.get("target"))
            values = [item.get("source"), item.get("target")]
            if not hide_page:
                values.insert(0, item.get("page") or index + 1)
            expected_cells.append([_normalized_docx_text(value) for value in values])
        columns = 2 if hide_page else 3
    elif mode == "transcript":
        require(data.get("case_info") or data.get("title") or "MAGI 文件")
        transcript_headers = ("發言人", "時間", "內容")
        expected_cells.append([_normalized_docx_text(value) for value in transcript_headers])
        for header in transcript_headers:
            require(header)
        for item in rows:
            if not isinstance(item, dict):
                raise ValueError("transcript rows must be objects")
            require(item.get("speaker"))
            require(item.get("time"))
            require(item.get("content"))
            expected_cells.append(
                [
                    _normalized_docx_text(item.get("speaker")),
                    _normalized_docx_text(item.get("time")),
                    _normalized_docx_text(item.get("content")),
                ]
            )
        columns = 3
    else:
        require(data.get("title") or "MAGI 文件")
        summary_headers = ("#", "段落", "摘要", "原文節錄")
        expected_cells.append([_normalized_docx_text(value) for value in summary_headers])
        for header in summary_headers:
            require(header)
        for index, item in enumerate(rows):
            if not isinstance(item, dict):
                raise ValueError("summary rows must be objects")
            require(index + 1)
            require(item.get("heading"))
            require(item.get("summary"))
            require(item.get("excerpt"))
            expected_cells.append(
                [
                    str(index + 1),
                    _normalized_docx_text(item.get("heading")),
                    _normalized_docx_text(item.get("summary")),
                    _normalized_docx_text(item.get("excerpt")),
                ]
            )
        columns = 4
    return {
        "mode": mode,
        "rows": len(rows) + 1,
        "columns": columns,
        "required_text": required,
        "cells": expected_cells,
    }


def _validate_docx_payload(payload: bytes, contract: dict) -> dict:
    if len(payload) < 512:
        return {"ok": False, "error": f"docx file too small:{len(payload)}"}
    try:
        with zipfile.ZipFile(io.BytesIO(payload)) as archive:
            names = set(archive.namelist())
            for required in ("[Content_Types].xml", "word/document.xml"):
                if required not in names:
                    return {"ok": False, "error": f"docx missing zip entry:{required}"}
            bad = archive.testzip()
            if bad:
                return {"ok": False, "error": f"docx corrupt zip entry:{bad}"}
    except (OSError, zipfile.BadZipFile) as exc:
        return {"ok": False, "error": f"docx bad zip:{type(exc).__name__}"}
    try:
        from docx import Document

        document = Document(io.BytesIO(payload))
    except Exception as exc:
        return {"ok": False, "error": f"docx readback failed:{type(exc).__name__}"}

    if len(document.tables) != 1:
        return {"ok": False, "error": f"docx table count mismatch:{len(document.tables)}"}
    table = document.tables[0]
    expected_rows = int(contract["rows"])
    expected_columns = int(contract["columns"])
    if len(table.rows) != expected_rows:
        return {
            "ok": False,
            "error": f"docx table row mismatch:{len(table.rows)}!={expected_rows}",
        }
    if any(len(row.cells) != expected_columns for row in table.rows):
        return {"ok": False, "error": f"docx table column mismatch:expected {expected_columns}"}
    for row_index, (row, expected) in enumerate(zip(table.rows, contract["cells"])):
        actual = [_normalized_docx_text(cell.text) for cell in row.cells]
        if actual != expected:
            return {"ok": False, "error": f"docx table content mismatch:row {row_index}"}

    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    parts.extend(
        cell.text
        for row in table.rows
        for cell in row.cells
        if cell.text
    )
    for section in document.sections:
        parts.extend(
            paragraph.text
            for paragraph in section.header.paragraphs
            if paragraph.text
        )
    visible = _normalized_docx_text("\n".join(parts))
    missing = [text for text in contract["required_text"] if text not in visible]
    if missing:
        return {
            "ok": False,
            "error": f"docx required content missing:{len(missing)}",
        }
    return {
        "ok": True,
        "mode": contract["mode"],
        "table_rows": len(table.rows),
        "table_columns": expected_columns,
        "content_items": len(contract["required_text"]),
        "sha256": hashlib.sha256(payload).hexdigest(),
        "size": len(payload),
    }


def _validate_docx_descriptor(descriptor: int, contract: dict) -> dict:
    metadata = os.fstat(descriptor)
    if not stat.S_ISREG(metadata.st_mode):
        return {"ok": False, "error": "docx path is not a regular file"}
    os.lseek(descriptor, 0, os.SEEK_SET)
    chunks: list[bytes] = []
    while True:
        chunk = os.read(descriptor, 1024 * 1024)
        if not chunk:
            break
        chunks.append(chunk)
    after = os.fstat(descriptor)
    if (
        after.st_dev != metadata.st_dev
        or after.st_ino != metadata.st_ino
        or after.st_size != metadata.st_size
        or after.st_mtime_ns != metadata.st_mtime_ns
    ):
        return {"ok": False, "error": "docx changed during validation"}
    return _validate_docx_payload(b"".join(chunks), contract)


def _set_run_font(
    run: Any,
    *,
    size: float,
    bold: bool = False,
    italic: bool = False,
    color: str = "000000",
) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    run.font.name = _FALLBACK_FONT_CJK
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{key}"), _FALLBACK_FONT_CJK)


def _set_cell_box(cell: Any, *, fill: str, header: bool = False) -> None:
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    cell.vertical_alignment = (
        WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if header
        else WD_CELL_VERTICAL_ALIGNMENT.TOP
    )
    tc_pr = cell._tc.get_or_add_tcPr()
    for tag in ("w:shd", "w:tcBorders", "w:tcMar"):
        for existing in tc_pr.findall(qn(tag)):
            tc_pr.remove(existing)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")
    tc_pr.append(shading)

    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "8" if header else "4")
        border.set(qn("w:color"), _FALLBACK_ACCENT if header else _FALLBACK_BORDER)
        borders.append(border)
    tc_pr.append(borders)

    margins = OxmlElement("w:tcMar")
    for edge, width in (("top", 100), ("bottom", 100), ("start", 140), ("end", 140)):
        margin = OxmlElement(f"w:{edge}")
        margin.set(qn("w:w"), str(width))
        margin.set(qn("w:type"), "dxa")
        margins.append(margin)
    tc_pr.append(margins)


def _set_table_geometry(table: Any, widths: list[int]) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    table.autofit = False
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    for tag in ("w:tblW", "w:tblInd", "w:tblLayout"):
        for existing in tbl_pr.findall(qn(tag)):
            tbl_pr.remove(existing)
    table_width = OxmlElement("w:tblW")
    table_width.set(qn("w:w"), str(total))
    table_width.set(qn("w:type"), "dxa")
    tbl_pr.append(table_width)
    table_indent = OxmlElement("w:tblInd")
    table_indent.set(qn("w:w"), "0")
    table_indent.set(qn("w:type"), "dxa")
    tbl_pr.append(table_indent)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)

    grid = table._tbl.tblGrid
    for column in list(grid):
        grid.remove(column)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.get_or_add_tcW()
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def _repeat_table_header(row: Any) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tr_pr = row._tr.get_or_add_trPr()
    for existing in tr_pr.findall(qn("w:tblHeader")):
        tr_pr.remove(existing)
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def _fill_cell(
    cell: Any,
    value: Any,
    *,
    fill: str,
    size: float = 10,
    bold: bool = False,
    color: str = "000000",
    centered: bool = False,
    header: bool = False,
) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    _set_cell_box(cell, fill=fill, header=header)
    lines = [line.strip() for line in _xml_text(value).splitlines() if line.strip()] or [""]
    cell.text = ""
    for index, line in enumerate(lines):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
        )
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.25
        run = paragraph.add_run(line)
        _set_run_font(run, size=size, bold=bold, color=color)


def _add_title(document: Any, title: Any, *, size: float = 16) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    text = _xml_text(title).strip()
    if not text:
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(text)
    _set_run_font(run, size=size, bold=True, color=_FALLBACK_ACCENT)


def _add_subtitle(document: Any, subtitle: Any) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    text = _xml_text(subtitle).strip()
    if not text:
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(15)
    run = paragraph.add_run(text)
    _set_run_font(run, size=11, italic=True, color="666666")


def _set_page_number_footer(section: Any) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prefix = paragraph.add_run("- ")
    _set_run_font(prefix, size=8, color="999999")
    run = paragraph.add_run()
    _set_run_font(run, size=8, color="999999")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instruction, separate, text, end):
        run._r.append(element)
    suffix = paragraph.add_run(" -")
    _set_run_font(suffix, size=8, color="999999")


def _build_python_docx(data: dict) -> Any:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Inches, Mm, Pt

    mode = str(data.get("mode") or "")
    collection_key = {
        "bilingual": "pages",
        "transcript": "segments",
        "summary": "sections",
    }.get(mode)
    if collection_key is None:
        raise ValueError(f"unsupported DOCX mode:{mode}")
    rows_data = data.get(collection_key)
    if not isinstance(rows_data, list) or not rows_data:
        raise ValueError(f"{collection_key} must be a non-empty list")

    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = _FALLBACK_FONT_CJK
    normal.font.size = Pt(10)
    normal_r_pr = normal._element.get_or_add_rPr()
    normal_r_fonts = normal_r_pr.get_or_add_rFonts()
    normal_r_fonts.set(qn("w:eastAsia"), _FALLBACK_FONT_CJK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.25

    section = document.sections[0]
    landscape = mode in {"bilingual", "summary"}
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Mm(297)
        section.page_height = Mm(210)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7875)
        section.right_margin = Inches(0.7875)
    else:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    header_text = data.get("header_text") if mode == "bilingual" else None
    if mode == "transcript":
        header_text = data.get("case_info")
    if not header_text:
        header_text = data.get("title") or "MAGI 文件"
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.add_run(_xml_text(header_text))
    _set_run_font(header_run, size=8, italic=True, color="999999")
    _set_page_number_footer(section)

    _add_title(document, data.get("title"), size=16 if landscape else 14)
    if mode == "bilingual":
        _add_subtitle(document, data.get("subtitle"))

    if mode == "bilingual":
        hide_page = data.get("hide_page_column") is True
        labels = data.get("col_labels") if isinstance(data.get("col_labels"), dict) else {}
        widths = [7200, 7200] if hide_page else [660, 6870, 6870]
        headers = [labels.get("col2") or "原文", labels.get("col3") or "翻譯"]
        if not hide_page:
            headers.insert(0, labels.get("col1") or "頁碼")
        table = document.add_table(rows=1, cols=len(widths))
        for cell, label in zip(table.rows[0].cells, headers):
            _fill_cell(
                cell,
                label,
                fill=_FALLBACK_HEADER_FILL,
                size=11,
                bold=True,
                color="FFFFFF",
                centered=True,
                header=True,
            )
        for index, item in enumerate(rows_data):
            row = table.add_row()
            fill = _FALLBACK_ROW_EVEN if index % 2 == 0 else _FALLBACK_ROW_ODD
            values = [_xml_text((item or {}).get("source")), _xml_text((item or {}).get("target"))]
            if not hide_page:
                values.insert(0, _xml_text((item or {}).get("page") or index + 1))
            for column, (cell, value) in enumerate(zip(row.cells, values)):
                _fill_cell(
                    cell,
                    value,
                    fill=fill,
                    size=10,
                    bold=column == 0 and not hide_page,
                    color="555555" if column == 0 and not hide_page else "000000",
                    centered=column == 0 and not hide_page,
                )
    elif mode == "transcript":
        widths = [1200, 1000, 6826]
        table = document.add_table(rows=1, cols=3)
        for cell, label in zip(table.rows[0].cells, ("發言人", "時間", "內容")):
            _fill_cell(
                cell,
                label,
                fill=_FALLBACK_HEADER_FILL,
                size=11,
                bold=True,
                color="FFFFFF",
                centered=True,
                header=True,
            )
        speaker_fills = {
            "法官": "DBEAFE",
            "審判長": "DBEAFE",
            "被告": "FEF3C7",
            "辯護人": "EDE9FE",
            "檢察官": "FCE7F3",
            "證人": "D1FAE5",
            "告訴人": "FEF9C3",
            "告訴代理人": "FEF9C3",
        }
        for index, item in enumerate(rows_data):
            row = table.add_row()
            speaker = _xml_text((item or {}).get("speaker"))
            fill = speaker_fills.get(
                speaker,
                _FALLBACK_ROW_EVEN if index % 2 == 0 else _FALLBACK_ROW_ODD,
            )
            values = (speaker, _xml_text((item or {}).get("time")), _xml_text((item or {}).get("content")))
            for column, (cell, value) in enumerate(zip(row.cells, values)):
                _fill_cell(
                    cell,
                    value,
                    fill=fill,
                    size=9 if column == 1 else 10,
                    bold=column == 0,
                    color="888888" if column == 1 else "000000",
                )
    else:
        widths = [550, 2000, 5925, 5925]
        table = document.add_table(rows=1, cols=4)
        for cell, label in zip(table.rows[0].cells, ("#", "段落", "摘要", "原文節錄")):
            _fill_cell(
                cell,
                label,
                fill=_FALLBACK_HEADER_FILL,
                size=11,
                bold=True,
                color="FFFFFF",
                centered=True,
                header=True,
            )
        for index, item in enumerate(rows_data):
            row = table.add_row()
            fill = _FALLBACK_ROW_EVEN if index % 2 == 0 else _FALLBACK_ROW_ODD
            values = (
                str(index + 1),
                _xml_text((item or {}).get("heading")),
                _xml_text((item or {}).get("summary")),
                _xml_text((item or {}).get("excerpt")),
            )
            for column, (cell, value) in enumerate(zip(row.cells, values)):
                _fill_cell(
                    cell,
                    value,
                    fill=fill,
                    size=9 if column == 3 else 10,
                    bold=column in {0, 1},
                    centered=column == 0,
                )

    _repeat_table_header(table.rows[0])
    _set_table_geometry(table, widths)
    return document


def _validated_fallback_path(data: dict, out_path: str, filename: str) -> Path:
    expected, normalized_name = _resolve_export_docx_path(filename)
    provided = Path(str(out_path or "")).expanduser().resolve(strict=False)
    payload_out = Path(str(data.get("out_path") or "")).expanduser().resolve(strict=False)
    payload_root = Path(str(data.get("exports_dir") or "")).expanduser().resolve(strict=False)
    if normalized_name != filename or expected.name != filename:
        raise ValueError("fallback output basename does not match filename")
    if provided != expected or payload_out != expected or payload_root != _exports_dir():
        raise ValueError("fallback output path contract mismatch")
    if expected.exists() and not expected.is_file():
        raise ValueError("fallback output path is not a file")
    return expected


def _export_directory_identity(descriptor: int) -> tuple[int, int]:
    opened = os.fstat(descriptor)
    if not stat.S_ISDIR(opened.st_mode):
        raise _ExportDirectoryChangedError("exports directory descriptor is not a directory")
    return opened.st_dev, opened.st_ino


def _assert_export_directory_path_identity(
    parent: Path,
    expected_identity: tuple[int, int],
) -> None:
    try:
        current = os.stat(parent, follow_symlinks=False)
        resolved = parent.resolve(strict=True)
    except (OSError, RuntimeError) as exc:
        raise _ExportDirectoryChangedError("exports directory identity is unavailable") from exc
    if (
        not stat.S_ISDIR(current.st_mode)
        or (current.st_dev, current.st_ino) != expected_identity
        or resolved != parent
        or parent != _exports_dir()
    ):
        raise _ExportDirectoryChangedError("exports directory changed during DOCX export")


def _assert_export_directory_identity(
    parent: Path,
    descriptor: int,
    *,
    expected_identity: tuple[int, int] | None = None,
) -> None:
    opened = os.fstat(descriptor)
    try:
        current = os.stat(parent, follow_symlinks=False)
        resolved = parent.resolve(strict=True)
    except (OSError, RuntimeError) as exc:
        raise _ExportDirectoryChangedError("exports directory identity is unavailable") from exc
    if (
        not stat.S_ISDIR(opened.st_mode)
        or not stat.S_ISDIR(current.st_mode)
        or opened.st_dev != current.st_dev
        or opened.st_ino != current.st_ino
        or (
            expected_identity is not None
            and (opened.st_dev, opened.st_ino) != expected_identity
        )
        or resolved != parent
        or parent != _exports_dir()
    ):
        raise _ExportDirectoryChangedError("exports directory changed during DOCX export")


def _open_export_directory(
    parent: Path,
    *,
    expected_identity: tuple[int, int] | None = None,
) -> int:
    flags = os.O_RDONLY | getattr(os, "O_DIRECTORY", 0) | getattr(os, "O_NOFOLLOW", 0)
    descriptor = os.open(parent, flags)
    try:
        _assert_export_directory_identity(
            parent,
            descriptor,
            expected_identity=expected_identity,
        )
    except Exception:
        os.close(descriptor)
        raise
    return descriptor


def _create_docx_candidate(descriptor: int, *, prefix: str) -> tuple[str, int]:
    flags = (
        os.O_RDWR
        | os.O_CREAT
        | os.O_EXCL
        | getattr(os, "O_NOFOLLOW", 0)
    )
    for _ in range(8):
        name = f".{prefix}-{uuid.uuid4().hex}.docx"
        try:
            return name, os.open(name, flags, 0o600, dir_fd=descriptor)
        except FileExistsError:
            continue
    raise FileExistsError("could not allocate a secure DOCX candidate")


def _unlink_docx_candidate(descriptor: int, name: str | None) -> None:
    if not name:
        return
    try:
        os.unlink(name, dir_fd=descriptor)
    except FileNotFoundError:
        pass


def _validate_docx_entry(descriptor: int, name: str, contract: dict) -> dict:
    flags = os.O_RDONLY | getattr(os, "O_NOFOLLOW", 0)
    try:
        candidate = os.open(name, flags, dir_fd=descriptor)
    except OSError as exc:
        return {"ok": False, "error": f"docx candidate open failed:{type(exc).__name__}"}
    try:
        return _validate_docx_descriptor(candidate, contract)
    finally:
        os.close(candidate)


def _copy_validated_docx_to_candidate(
    source: Path,
    directory_descriptor: int,
    contract: dict,
    *,
    prefix: str,
) -> str:
    source_flags = os.O_RDONLY | getattr(os, "O_NOFOLLOW", 0)
    source_descriptor = os.open(source, source_flags)
    candidate_name: str | None = None
    candidate_descriptor: int | None = None
    try:
        validation = _validate_docx_descriptor(source_descriptor, contract)
        if not validation.get("ok"):
            raise ValueError(str(validation.get("error") or "staged DOCX validation failed"))
        os.lseek(source_descriptor, 0, os.SEEK_SET)
        candidate_name, candidate_descriptor = _create_docx_candidate(
            directory_descriptor,
            prefix=prefix,
        )
        while True:
            chunk = os.read(source_descriptor, 1024 * 1024)
            if not chunk:
                break
            view = memoryview(chunk)
            while view:
                written = os.write(candidate_descriptor, view)
                view = view[written:]
        os.fsync(candidate_descriptor)
        os.close(candidate_descriptor)
        candidate_descriptor = None
        return candidate_name
    except Exception:
        if candidate_descriptor is not None:
            os.close(candidate_descriptor)
        _unlink_docx_candidate(directory_descriptor, candidate_name)
        raise
    finally:
        os.close(source_descriptor)


def _publish_docx_candidate(
    descriptor: int,
    parent: Path,
    candidate_name: str,
    target_name: str,
    contract: dict,
) -> dict:
    candidate_validation = _validate_docx_entry(descriptor, candidate_name, contract)
    if not candidate_validation.get("ok"):
        raise ValueError(str(candidate_validation.get("error") or "DOCX candidate validation failed"))
    _assert_export_directory_identity(parent, descriptor)
    try:
        target_metadata = os.stat(target_name, dir_fd=descriptor, follow_symlinks=False)
    except FileNotFoundError:
        target_metadata = None
    if target_metadata is not None and stat.S_ISDIR(target_metadata.st_mode):
        raise ValueError("DOCX target is a directory")
    os.replace(
        candidate_name,
        target_name,
        src_dir_fd=descriptor,
        dst_dir_fd=descriptor,
    )
    os.fsync(descriptor)
    _assert_export_directory_identity(parent, descriptor)
    published = _validate_docx_entry(descriptor, target_name, contract)
    if not published.get("ok"):
        raise ValueError(str(published.get("error") or "published DOCX validation failed"))
    if published.get("sha256") != candidate_validation.get("sha256"):
        raise ValueError("published DOCX digest mismatch")
    return published


def _success_result(
    path: Path,
    filename: str,
    validation: dict,
    *,
    generator: str,
    fallback_code: str = "",
) -> dict:
    base = _load_public_base_url()
    result = {
        "success": True,
        "path": str(path),
        "filename": filename,
        "url": (base.rstrip("/") + f"/static/exports/{filename}") if base else "",
        "format": "docx",
        "validation": validation,
        "generator": generator,
    }
    if fallback_code:
        result["fallback_code"] = fallback_code
    return result


def _diagnostic_evidence(*values: Any) -> str:
    material = "\x1f".join(str(value or "")[:1000] for value in values)
    return hashlib.sha256(material.encode("utf-8", errors="replace")).hexdigest()[:16]


def _run_python_docx_fallback(
    data: dict,
    out_path: str,
    filename: str,
    *,
    node_error: str,
    anchored_target: Path | None = None,
    expected_directory_identity: tuple[int, int] | None = None,
) -> dict:
    directory_descriptor: int | None = None
    candidate_name: str | None = None
    try:
        if anchored_target is None:
            target = _validated_fallback_path(data, out_path, filename)
        else:
            target = anchored_target
            if Path(str(out_path)).expanduser() != target or target.name != filename:
                raise ValueError("fallback target changed after initial validation")
        contract = _docx_contract(data)
        directory_descriptor = _open_export_directory(
            target.parent,
            expected_identity=expected_directory_identity,
        )
        document = _build_python_docx(data)
        candidate_name, candidate_descriptor = _create_docx_candidate(
            directory_descriptor,
            prefix="magi-docx-fallback",
        )
        with os.fdopen(candidate_descriptor, "w+b") as candidate:
            document.save(candidate)
            candidate.flush()
            os.fsync(candidate.fileno())
        validation = _publish_docx_candidate(
            directory_descriptor,
            target.parent,
            candidate_name,
            target.name,
            contract,
        )
        candidate_name = None
        evidence = _diagnostic_evidence(node_error)
        logging.getLogger(__name__).info(
            "DOCX Node backend unavailable; python-docx fallback succeeded "
            "(evidence=%s)",
            evidence,
        )
        return _success_result(
            target,
            filename,
            validation,
            generator="python-docx-fallback",
            fallback_code="node_backend_unavailable",
        )
    except _ExportDirectoryChangedError as exc:
        evidence = _diagnostic_evidence(node_error, type(exc).__name__, exc)
        logging.getLogger(__name__).warning(
            "DOCX export directory changed before fallback publication "
            "(evidence=%s)",
            evidence,
        )
        return {
            "success": False,
            "error": "DOCX export directory changed",
            "error_code": "docx_export_directory_changed",
            "evidence": evidence,
        }
    except Exception as exc:
        evidence = _diagnostic_evidence(node_error, type(exc).__name__, exc)
        logging.getLogger(__name__).warning(
            "DOCX export failed after Node fallback "
            "(node_error_type=%s, fallback_error_type=%s, evidence=%s)",
            "backend-unavailable",
            type(exc).__name__,
            evidence,
        )
        return {
            "success": False,
            "error": "DOCX export failed",
            "error_code": "docx_export_failed",
            "evidence": evidence,
        }
    finally:
        if directory_descriptor is not None:
            _unlink_docx_candidate(directory_descriptor, candidate_name)
            os.close(directory_descriptor)


def export_bilingual_docx(
    pages: List[Dict[str, Any]],
    *,
    title: str = "",
    subtitle: str = "",
    header_text: str = "",
    prefix: str = "translate",
    filename: str = "",
    col_labels: Optional[Dict[str, str]] = None,
    hide_page_column: bool = False,
) -> dict:
    """
    產生雙語對照 docx 表格。

    pages: [{"page": 1, "source": "English text...", "target": "中文翻譯..."}]
    title: 文件標題
    subtitle: 副標題
    header_text: 頁首文字
    prefix: 檔名前綴
    filename: 指定檔名（不含路徑），若空則自動產生
    col_labels: 自訂表頭 {"col1": "段落", "col2": "原文", "col3": "摘要"}

    Returns: {"success": True, "path": "...", "filename": "...", "url": "..."}
    """
    if not pages:
        return {"success": False, "error": "empty pages"}

    if not filename:
        filename = _safe_generated_filename(prefix)

    try:
        out_path, filename = _resolve_export_docx_path(filename)
    except ValueError as exc:
        return {"success": False, "error": str(exc)}

    # Write page data to temp JSON
    data = {
        "mode": "bilingual",
        "title": title or "",
        "subtitle": subtitle or "",
        "header_text": header_text or "",
        "pages": pages,
        "out_path": str(out_path),
        "filename": filename,
        "exports_dir": str(_exports_dir()),
        "hide_page_column": bool(hide_page_column),
    }
    if col_labels:
        data["col_labels"] = col_labels

    return _run_docx_generator(data, str(out_path), filename)


def export_transcript_docx(
    segments: List[Dict[str, Any]],
    *,
    title: str = "",
    case_info: str = "",
    prefix: str = "transcript",
    filename: str = "",
) -> dict:
    """
    產生逐字稿 docx 表格。

    segments: [{"speaker": "法官", "time": "10:30", "content": "..."}]
    title: 文件標題（如「114年度訴字第123號 審理程序筆錄」）
    case_info: 案件資訊（頁首）
    prefix: 檔名前綴
    filename: 指定檔名

    Returns: {"success": True, "path": "...", "filename": "...", "url": "..."}
    """
    if not segments:
        return {"success": False, "error": "empty segments"}

    if not filename:
        filename = _safe_generated_filename(prefix)

    try:
        out_path, filename = _resolve_export_docx_path(filename)
    except ValueError as exc:
        return {"success": False, "error": str(exc)}

    data = {
        "mode": "transcript",
        "title": title or "",
        "case_info": case_info or "",
        "segments": segments,
        "out_path": str(out_path),
        "filename": filename,
        "exports_dir": str(_exports_dir()),
    }

    return _run_docx_generator(data, str(out_path), filename)


def export_summary_docx(
    sections: List[Dict[str, Any]],
    *,
    title: str = "",
    prefix: str = "summary",
    filename: str = "",
) -> dict:
    """
    產生摘要 docx 表格。

    sections: [{"heading": "第一部分", "summary": "摘要...", "excerpt": "原文節錄..."}]
    title: 文件標題
    prefix: 檔名前綴
    filename: 指定檔名

    Returns: {"success": True, "path": "...", "filename": "...", "url": "..."}
    """
    if not sections:
        return {"success": False, "error": "empty sections"}

    if not filename:
        filename = _safe_generated_filename(prefix)

    try:
        out_path, filename = _resolve_export_docx_path(filename)
    except ValueError as exc:
        return {"success": False, "error": str(exc)}

    data = {
        "mode": "summary",
        "title": title or "",
        "sections": sections,
        "out_path": str(out_path),
        "filename": filename,
        "exports_dir": str(_exports_dir()),
    }

    return _run_docx_generator(data, str(out_path), filename)


def _run_docx_generator(data: dict, out_path: str, filename: str) -> dict:
    """Prefer Node.js and fall back to an atomic python-docx implementation."""
    try:
        target = _validated_fallback_path(data, out_path, filename)
        contract = _docx_contract(data)
    except ValueError as exc:
        return {"success": False, "error": str(exc)}

    backend = _node_backend()
    if backend is None:
        return _run_python_docx_fallback(
            data,
            str(target),
            filename,
            node_error="sealed_release_node_backend_unbound",
            anchored_target=target,
        )
    node, node_path = backend
    directory_descriptor: int | None = None
    initial_directory_identity: tuple[int, int] | None = None
    candidate_name: str | None = None
    node_error = "node_backend_unavailable"
    directory_identity_failed = False
    try:
        directory_descriptor = _open_export_directory(target.parent)
        initial_directory_identity = _export_directory_identity(directory_descriptor)
        staging_parent = str(os.environ.get("TMPDIR") or "").strip() or None
        with tempfile.TemporaryDirectory(
            prefix="magi-docx-node-",
            dir=staging_parent,
        ) as staging_name:
            staging = Path(staging_name)
            os.chmod(staging, 0o700)
            node_output = staging / "node-output.docx"
            json_path = staging / "payload.json"
            node_data = dict(data)
            node_data["out_path"] = str(node_output)
            node_data["filename"] = node_output.name
            node_data["exports_dir"] = str(staging)
            with json_path.open("w", encoding="utf-8") as payload_file:
                json.dump(node_data, payload_file, ensure_ascii=False)

            script_path = os.path.join(os.path.dirname(__file__), "_docx_table_gen.js")
            env = os.environ.copy()
            if node_path:
                env["NODE_PATH"] = node_path

            try:
                cp = subprocess.run(
                    [node, script_path, str(json_path)],
                    capture_output=True,
                    text=True,
                    timeout=30,
                    env=env,
                )
            except subprocess.TimeoutExpired:
                node_error = "docx generator timeout"
            except Exception as exc:
                node_error = f"{type(exc).__name__}: {exc}"
            else:
                if cp.returncode == 0:
                    try:
                        candidate_name = _copy_validated_docx_to_candidate(
                            node_output,
                            directory_descriptor,
                            contract,
                            prefix="magi-docx-node",
                        )
                        validation = _publish_docx_candidate(
                            directory_descriptor,
                            target.parent,
                            candidate_name,
                            target.name,
                            contract,
                        )
                    except _ExportDirectoryChangedError as exc:
                        directory_identity_failed = True
                        node_error = f"node DOCX export directory changed:{type(exc).__name__}"
                    except Exception as exc:
                        node_error = f"node DOCX candidate rejected:{type(exc).__name__}"
                    else:
                        candidate_name = None
                        return _success_result(
                            target,
                            filename,
                            validation,
                            generator="node-docx",
                        )
                else:
                    node_error = (
                        f"docx generator failed (rc={cp.returncode}): "
                        f"{(cp.stderr or '')[:300]}"
                    )
    except _ExportDirectoryChangedError as exc:
        directory_identity_failed = True
        node_error = f"node DOCX export directory changed:{type(exc).__name__}"
    except Exception as exc:
        node_error = f"node DOCX staging failed:{type(exc).__name__}"
    finally:
        if directory_descriptor is not None:
            _unlink_docx_candidate(directory_descriptor, candidate_name)
            os.close(directory_descriptor)
    if not directory_identity_failed and initial_directory_identity is not None:
        try:
            _assert_export_directory_path_identity(
                target.parent,
                initial_directory_identity,
            )
        except _ExportDirectoryChangedError:
            directory_identity_failed = True
    if directory_identity_failed or target.parent != _exports_dir():
        return {
            "success": False,
            "error": "DOCX export directory changed",
            "error_code": "docx_export_directory_changed",
        }
    return _run_python_docx_fallback(
        data,
        str(target),
        filename,
        node_error=node_error,
        anchored_target=target,
        expected_directory_identity=initial_directory_identity,
    )
