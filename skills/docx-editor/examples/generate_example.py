"""
generate_example.py — Phase 4: generate_docx 用法示範

執行方式：
    python skills/docx-editor/examples/generate_example.py
"""

import os
import sys

# Ensure skill is importable
_SKILL_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, _SKILL_DIR)
sys.path.insert(0, os.path.join(_SKILL_DIR, "lib"))

from lib.generator import generate_docx, GenerateDocxRequest, SectionSpec, TableSpec


def main():
    req = GenerateDocxRequest(
        title="結案報告書",
        sections=[
            SectionSpec(
                heading="事實摘要",
                level=1,
                content="本案當事人甲方（原告）於民國114年3月提起損害賠償訴訟。\n\n雙方爭點在於系爭合約之解釋。",
            ),
            SectionSpec(
                heading="法律分析",
                level=1,
            ),
            SectionSpec(
                heading="損害賠償依據",
                level=2,
                content="依民法第184條，故意或過失，不法侵害他人之權利者，負損害賠償責任。",
            ),
            SectionSpec(
                heading="費用明細",
                level=1,
                table=TableSpec(
                    headers=["項目", "金額（元）", "說明"],
                    rows=[
                        ["律師費", "50,000", "第一審代理費"],
                        ["書狀費", "2,000", "書狀繕打費"],
                        ["合計", "52,000", ""],
                    ],
                ),
            ),
            SectionSpec(
                heading="結語",
                level=1,
                content="爰請鈞院依法裁判。",
                page_break=True,
            ),
        ],
    )

    docx_bytes = generate_docx(req)
    out_path = "/tmp/generate_example_output.docx"
    with open(out_path, "wb") as f:
        f.write(docx_bytes)

    print(f"✅ 產出 {len(docx_bytes)} bytes → {out_path}")

    # Verify round-trip
    from docx import Document
    import io
    doc = Document(io.BytesIO(docx_bytes))
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "結案報告書" in all_text
    assert "費用明細" in all_text
    assert "52,000" in "\n".join(
        cell.text for tbl in doc.tables for row in tbl.rows for cell in row.cells
    )
    print("✅ Round-trip 驗證通過")


if __name__ == "__main__":
    main()
