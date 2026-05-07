"""
generator.py — 從 sections list 產出 .docx bytes

Phase 4 of docx-editor skill.

Usage:
    from lib.generator import generate_docx, GenerateDocxRequest, SectionSpec, TableSpec
"""

import io
import warnings
from dataclasses import dataclass, field
from typing import List, Optional


@dataclass
class TableSpec:
    """表格規格：headers + rows。每 row 的欄數必須等於 len(headers)。"""
    headers: List[str]
    rows: List[List[str]]


@dataclass
class SectionSpec:
    """文件章節規格。"""
    heading: Optional[str] = None
    level: int = 1                  # heading 層級 1 / 2 / 3
    content: Optional[str] = None  # 段落文字，多段以 \n\n 分隔
    table: Optional[TableSpec] = None
    page_break: bool = False        # True → 此 section 從新頁開始


@dataclass
class GenerateDocxRequest:
    """產生文件的完整請求。"""
    title: str
    sections: List[SectionSpec]
    landscape: bool = False
    author: str = "MAGI"


def generate_docx(req: GenerateDocxRequest) -> bytes:
    """從 sections 產出 .docx bytes。

    規則（仿 Mike SYSTEM_PROMPT DOCX GENERATION 段）：
    - heading 層級：必須遞增不跳級（1→2→3→2→3 OK；1→3 拒絕回 ValueError）
    - 編號：所有 numbered list 從 1 開始（python-docx 預設）
    - heading text 含 numbered prefix（例如「1. 」）只警告不拒絕
    - landscape=True → 設 page orientation=landscape + width/height swap
    - 表格 headers/rows 列數必須一致

    Returns:
        .docx bytes（可直接寫入檔案或回傳 HTTP response）

    Raises:
        ValueError: heading 層級跳級、表格欄數不一致
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.section import WD_ORIENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError as e:
        raise ImportError(f"python-docx not installed: {e}")

    doc = Document()

    # --- 文件標題 ---
    if req.title:
        doc.add_heading(req.title, level=0)

    # --- 頁面設定（landscape）---
    if req.landscape:
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        # swap width / height
        new_width = section.page_height
        new_height = section.page_width
        section.page_width = new_width
        section.page_height = new_height

    # --- 追蹤 heading 層級（防止跳級）---
    _prev_heading_level = 0  # 0 = 尚未出現任何 heading

    for sec in req.sections:
        # Page break 在 section 開始前插入
        if sec.page_break:
            doc.add_page_break()

        # Heading
        if sec.heading is not None:
            level = max(1, min(3, sec.level))  # 限制 1-3

            # 跳級檢查：只有從非 0 的前一層跳超過 1 級才報錯
            if _prev_heading_level > 0 and level > _prev_heading_level + 1:
                raise ValueError(
                    f"Heading level jump: previous={_prev_heading_level}, current={level} "
                    f"(heading: {sec.heading!r}). Level must increase by at most 1."
                )

            # 帶數字前綴的 heading 警告
            import re
            if re.match(r"^\d+[.、。\s]", sec.heading):
                warnings.warn(
                    f"Heading contains numbering prefix: {sec.heading!r}. "
                    "Consider using auto-numbering instead.",
                    UserWarning,
                    stacklevel=2,
                )

            doc.add_heading(sec.heading, level=level)
            _prev_heading_level = level

        # Content（多段以 \n\n 分隔）
        if sec.content:
            paragraphs = sec.content.split("\n\n")
            for para_text in paragraphs:
                para_text = para_text.strip()
                if para_text:
                    doc.add_paragraph(para_text)

        # Table
        if sec.table is not None:
            _add_table(doc, sec.table)

    # --- 輸出 bytes ---
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_table(doc, table_spec: TableSpec) -> None:
    """在文件中加入表格。"""
    if not table_spec.headers:
        return

    n_cols = len(table_spec.headers)

    # 驗證每 row 欄數一致
    for i, row in enumerate(table_spec.rows):
        if len(row) != n_cols:
            raise ValueError(
                f"Table row {i} has {len(row)} columns, expected {n_cols} "
                f"(matching headers: {table_spec.headers})"
            )

    n_rows = len(table_spec.rows) + 1  # +1 for header row
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"

    # Header row
    header_row = table.rows[0]
    for j, header in enumerate(table_spec.headers):
        cell = header_row.cells[j]
        cell.text = header
        # Bold header
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    # Data rows
    for i, row_data in enumerate(table_spec.rows):
        data_row = table.rows[i + 1]
        for j, cell_text in enumerate(row_data):
            data_row.cells[j].text = str(cell_text)
