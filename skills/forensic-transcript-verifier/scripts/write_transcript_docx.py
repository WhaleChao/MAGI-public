#!/usr/bin/env python3
"""Write the court transcript DOCX with the hash-bound V3 Python runtime."""

from __future__ import annotations

import json
import os
import re
import sys
from pathlib import Path
from typing import Any, Mapping

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


FONT = "Noto Sans CJK TC"
CONTROL_CHARACTERS = re.compile(r"[\x00-\x08\x0B\x0C\x0E-\x1F]")
UNCERTAIN = re.compile(r"【[^】]*(?:聽辨|發話者|未定)[^】]*】")


def clean(value: Any) -> str:
    return CONTROL_CHARACTERS.sub("", str(value or ""))


def set_run_font(run: Any, *, size: float, bold: bool = False, color: str | None = None) -> None:
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_run(
    paragraph: Any,
    text: Any,
    *,
    size: float = 11,
    bold: bool = False,
    color: str | None = None,
) -> Any:
    run = paragraph.add_run(clean(text))
    set_run_font(run, size=size, bold=bold, color=color)
    return run


def set_cell_shading(cell: Any, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_text(cell: Any, value: Any, *, header: bool = False) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    add_run(paragraph, value, size=9.5, bold=header)
    if header:
        set_cell_shading(cell, "D9EAF7")


def add_page_number(paragraph: Any) -> None:
    add_run(paragraph, "— ", size=8, color="6B7280")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    set_run_font(run, size=8, color="6B7280")
    run._r.extend((begin, instruction, separate, value, end))
    add_run(paragraph, " —", size=8, color="6B7280")


def write_document(task: Mapping[str, Any]) -> Path:
    output = Path(clean(task.get("output_path"))).expanduser().resolve()
    if output.suffix.lower() != ".docx":
        raise ValueError("output_path must be .docx")
    output.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.9)
    section.right_margin = Cm(2.54)
    section.bottom_margin = Cm(1.9)
    section.left_margin = Cm(2.54)

    normal = document.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(11)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    zoom = document.settings.element.find(qn("w:zoom"))
    if zoom is not None:
        # python-docx emits bestFit without the percent attribute required by
        # the OOXML validator used by MAGI's court-grade DOCX gate.
        zoom.set(qn("w:percent"), "100")

    title = clean(task.get("title") or "訊問影音完整譯文（重新勘驗校正版）")
    header_text = clean(task.get("header") or title or "法院影音勘驗譯文")
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(header, header_text, size=8, color="6B7280")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(footer)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(13)
    add_run(paragraph, title, size=17, bold=True)

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(7)
    add_run(paragraph, task.get("case_info"), size=10, color="4B5563")

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(13)
    add_run(
        paragraph,
        "說明：本譯文由 MAGI 依影音、時間戳 ASR、畫面序列及人工校正節文進行兩次獨立勘驗。"
        "人工校正節文於其範圍內逐字保留；不確定內容以【】明示。",
        size=10,
    )

    turns = task.get("turns") if isinstance(task.get("turns"), list) else []
    for raw in turns:
        turn = raw if isinstance(raw, Mapping) else {}
        heading = document.add_paragraph()
        heading.paragraph_format.keep_with_next = True
        heading.paragraph_format.space_before = Pt(7.5)
        heading.paragraph_format.space_after = Pt(3)
        add_run(heading, f"[{clean(turn.get('display'))}] ", size=10, bold=True, color="1F4E79")
        add_run(heading, turn.get("speaker") or "【發話者未定】", size=10, bold=True)

        body = document.add_paragraph()
        body.paragraph_format.space_after = Pt(4)
        body.paragraph_format.line_spacing = 1.5
        text = clean(turn.get("text"))
        add_run(body, f"「{text}」", size=11, color="9C0006" if UNCERTAIN.search(text) else "111827")

    document.add_page_break()
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(9)
    add_run(paragraph, "未決事項與人工最終確認清單", size=14, bold=True)

    unresolved = task.get("unresolved") if isinstance(task.get("unresolved"), list) else []
    if unresolved:
        table = document.add_table(rows=1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        widths = (Cm(2.5), Cm(2.3), Cm(4.2), Cm(6.0))
        labels = ("時間", "發話者", "未決內容", "原因／所需確認")
        for index, label in enumerate(labels):
            table.columns[index].width = widths[index]
            set_cell_text(table.rows[0].cells[index], label, header=True)
        for raw in unresolved:
            row = raw if isinstance(raw, Mapping) else {}
            cells = table.add_row().cells
            values = (row.get("time"), row.get("speaker"), row.get("content"), row.get("reason"))
            for index, value in enumerate(values):
                cells[index].width = widths[index]
                set_cell_text(cells[index], value)
    else:
        paragraph = document.add_paragraph()
        add_run(paragraph, "本次兩輪技術複核未留下未決項目；法院送件前仍須由承辦人確認。", size=10)

    temporary = output.with_name(f"{output.name}.{os.getpid()}.tmp")
    try:
        document.save(temporary)
        os.replace(temporary, output)
    finally:
        temporary.unlink(missing_ok=True)
    return output


def main() -> int:
    if len(sys.argv) != 2:
        raise ValueError("usage: write_transcript_docx.py task.json")
    task_path = Path(sys.argv[1]).expanduser().resolve(strict=True)
    payload = json.loads(task_path.read_text(encoding="utf-8"))
    if not isinstance(payload, dict):
        raise ValueError("task JSON must be an object")
    output = write_document(payload)
    print(
        json.dumps(
            {"success": True, "output": str(output), "turns": len(payload.get("turns") or [])},
            ensure_ascii=False,
            sort_keys=True,
        )
    )
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except Exception as exc:
        print(str(exc), file=sys.stderr)
        raise SystemExit(1)
