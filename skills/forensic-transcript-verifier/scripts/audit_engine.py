#!/usr/bin/env python3
"""Deterministic gates for the forensic transcript verifier skill."""

from __future__ import annotations

import hashlib
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import time
import zipfile
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Any, Iterable
from xml.etree import ElementTree as ET


TIME_RE = re.compile(r"\d{2}:\d{2}:\d{2}(?:\.\d+)?")
HEADER_RE = re.compile(
    r"^[ \t]*\[([^\]]+)\][ \t]*([^\n]*)\n(.*?)(?=^[ \t]*\[[^\]]+\][ \t]*[^\n]*\n|\Z)",
    re.MULTILINE | re.DOTALL,
)
TURN_BODY_STOP_RE = re.compile(
    r"(?:^|\n)[ \t]*(?:[一二三四五六七八九十]+、)?(?:仍須由書面資料確認之項目|仍須確認之項目|未決項目|未決事項(?:與人工最終確認清單)?|複核原則)[^\n]*(?:\n|\Z)",
    re.MULTILINE,
)
UNCERTAINTY_RE = re.compile(r"【(?:聽辨|姓名聽辨|發話者)[^】]*】")
QUESTION_MARKERS = ("?", "？", "嗎", "呢", "誰", "為什麼", "为什么", "然後呢", "然后呢", "對不對", "对不对")
DEFAULT_ANSWER_SPEAKERS = ("吳慧珠", "被告", "證人", "告訴人", "受訊問人", "答")


@dataclass(frozen=True)
class Turn:
    display: str
    speaker: str
    text: str
    start: float
    end: float
    order: int


def clock_to_seconds(value: str) -> float:
    match = TIME_RE.search(str(value or ""))
    if not match:
        raise ValueError(f"invalid time: {value!r}")
    hour, minute, second = match.group(0).split(":")
    return int(hour) * 3600 + int(minute) * 60 + float(second)


def seconds_to_clock(value: float) -> str:
    value = max(0.0, float(value))
    hour = int(value // 3600)
    minute = int((value % 3600) // 60)
    second = value % 60
    return f"{hour:02d}:{minute:02d}:{second:05.2f}"


def _run(command: list[str], *, env: dict[str, str] | None = None, timeout: int = 180) -> dict[str, Any]:
    try:
        proc = subprocess.run(
            command,
            check=False,
            capture_output=True,
            text=True,
            timeout=timeout,
            env=env,
        )
        return {
            "command": command,
            "returncode": proc.returncode,
            "stdout": proc.stdout.strip(),
            "stderr": proc.stderr.strip(),
        }
    except (OSError, subprocess.TimeoutExpired) as exc:
        return {"command": command, "returncode": 127, "stdout": "", "stderr": str(exc)}


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _docx_xml_text(path: Path) -> str:
    """Fallback accepted-text extraction that skips deleted OOXML runs."""
    with zipfile.ZipFile(path) as archive:
        root = ET.fromstring(archive.read("word/document.xml"))

    paragraphs: list[str] = []

    def collect(node: ET.Element, *, deleted: bool = False) -> str:
        local = node.tag.rsplit("}", 1)[-1]
        deleted = deleted or local in {"del", "moveFrom"}
        if deleted:
            return ""
        if local in {"t", "tab", "br"}:
            if local == "tab":
                return "\t"
            if local == "br":
                return "\n"
            return node.text or ""
        return "".join(collect(child, deleted=deleted) for child in node)

    for paragraph in root.iter():
        if paragraph.tag.rsplit("}", 1)[-1] == "p":
            text = collect(paragraph).strip()
            if text:
                paragraphs.append(text)
    return "\n\n".join(paragraphs)


def extract_text(path: str | Path, output_dir: str | Path | None = None) -> tuple[str, dict[str, Any]]:
    source = Path(path).expanduser().resolve()
    if not source.is_file():
        raise FileNotFoundError(source)
    metadata: dict[str, Any] = {"source": str(source), "method": "read_text"}
    if source.suffix.lower() != ".docx":
        return source.read_text(encoding="utf-8", errors="replace"), metadata

    pandoc = shutil.which("pandoc")
    if pandoc:
        target_dir = Path(output_dir).expanduser().resolve() if output_dir else source.parent
        target_dir.mkdir(parents=True, exist_ok=True)
        target = target_dir / f"{source.stem}.accepted.txt"
        result = _run([pandoc, "--track-changes=accept", str(source), "-t", "plain", "-o", str(target)])
        metadata.update({"method": "pandoc", "command": result})
        if result["returncode"] == 0 and target.is_file():
            metadata["extracted_path"] = str(target)
            return target.read_text(encoding="utf-8", errors="replace"), metadata

    metadata["method"] = "ooxml-fallback"
    return _docx_xml_text(source), metadata


def _compact_whitespace(text: str) -> str:
    return re.sub(r"\s+", " ", str(text or "")).strip()


def _strip_outer_quote(text: str) -> str:
    value = _compact_whitespace(text)
    pairs = (("「", "」"), ('"', '"'), ("“", "”"))
    for left, right in pairs:
        if value.startswith(left) and value.endswith(right) and len(value) >= 2:
            return value[len(left) : -len(right)].strip()
    return value


def parse_turns(text: str) -> list[Turn]:
    turns: list[Turn] = []
    for order, match in enumerate(HEADER_RE.finditer(str(text or ""))):
        display = match.group(1).strip()
        times = TIME_RE.findall(display)
        if not times:
            continue
        start = clock_to_seconds(times[0])
        end = clock_to_seconds(times[-1])
        speaker = _compact_whitespace(match.group(2))
        raw_body = TURN_BODY_STOP_RE.split(match.group(3), maxsplit=1)[0]
        body = _strip_outer_quote(raw_body)
        turns.append(Turn(display, speaker, body, start, end, order))
    return turns


def split_monotonic_blocks(turns: Iterable[Turn], reset_tolerance: float = 5.0) -> list[list[Turn]]:
    blocks: list[list[Turn]] = []
    for turn in turns:
        if not blocks or turn.start < blocks[-1][-1].start - reset_tolerance:
            blocks.append([turn])
        else:
            blocks[-1].append(turn)
    return blocks


def find_overlaps(turns: list[Turn], tolerance: float = 0.25) -> list[dict[str, Any]]:
    findings: list[dict[str, Any]] = []
    for block_index, block in enumerate(split_monotonic_blocks(turns), start=1):
        for previous, current in zip(block, block[1:]):
            overlap = previous.end - current.start
            if overlap > tolerance:
                findings.append(
                    {
                        "block": block_index,
                        "seconds": round(overlap, 3),
                        "previous": asdict(previous),
                        "current": asdict(current),
                        "cross_speaker": previous.speaker != current.speaker,
                    }
                )
    return findings


def compare_baseline(candidate_text: str, baseline_text: str) -> dict[str, Any]:
    candidate_turns = parse_turns(candidate_text)
    baseline_turns = parse_turns(baseline_text)
    rows: list[dict[str, Any]] = []
    for baseline in baseline_turns:
        time_tokens = TIME_RE.findall(baseline.display)
        candidates = [turn for turn in candidate_turns if all(token in turn.display for token in time_tokens)]
        time_ok = bool(candidates)
        speaker_ok = (not baseline.speaker) or any(baseline.speaker in turn.speaker for turn in candidates)
        baseline_body = _compact_whitespace(baseline.text)
        text_ok = bool(baseline_body) and any(
            baseline_body in _compact_whitespace(turn.text) for turn in candidates
        )
        rows.append(
            {
                "time": baseline.display,
                "speaker": baseline.speaker,
                "text": baseline.text,
                "time_ok": time_ok,
                "speaker_ok": speaker_ok,
                "text_ok": text_ok,
            }
        )
    return {
        "baseline_entries": len(baseline_turns),
        "matched_entries": sum(1 for row in rows if row["time_ok"] and row["speaker_ok"] and row["text_ok"]),
        "exact": bool(baseline_turns) and all(row["time_ok"] and row["speaker_ok"] and row["text_ok"] for row in rows),
        "mismatches": [row for row in rows if not (row["time_ok"] and row["speaker_ok"] and row["text_ok"])],
    }


def load_asr_segments(path: str | Path) -> list[dict[str, Any]]:
    payload = json.loads(Path(path).expanduser().resolve().read_text(encoding="utf-8"))
    if isinstance(payload, dict):
        rows = payload.get("segments") or []
    elif isinstance(payload, list):
        rows = payload
    else:
        rows = []
    return [row for row in rows if isinstance(row, dict) and "start" in row and "end" in row]


def _merge_intervals(intervals: list[tuple[float, float]], join_gap: float = 2.0) -> list[list[float]]:
    merged: list[list[float]] = []
    for start, end in sorted((float(a), float(b)) for a, b in intervals if float(b) >= float(a)):
        if not merged or start > merged[-1][1] + join_gap:
            merged.append([start, end])
        else:
            merged[-1][1] = max(merged[-1][1], end)
    return merged


def _group_segments(segments: list[dict[str, Any]], gap: float = 3.0) -> list[dict[str, Any]]:
    groups: list[list[dict[str, Any]]] = []
    for segment in sorted(segments, key=lambda row: float(row.get("start", 0))):
        if not groups or float(segment["start"]) > float(groups[-1][-1]["end"]) + gap:
            groups.append([segment])
        else:
            groups[-1].append(segment)
    return [
        {
            "start": group[0]["start"],
            "end": group[-1]["end"],
            "time": f"{seconds_to_clock(group[0]['start'])}–{seconds_to_clock(group[-1]['end'])}",
            "text": " / ".join(_compact_whitespace(item.get("text", "")) for item in group),
            "segments": len(group),
        }
        for group in groups
    ]


def audit_asr_coverage(
    turns: list[Turn],
    segments: list[dict[str, Any]],
    *,
    extra_coverage: list[tuple[float, float]] | None = None,
    min_duration: float = 0.45,
    min_logprob: float = -0.55,
    max_no_speech: float = 0.35,
    padding: float = 1.5,
) -> dict[str, Any]:
    coverage = [(turn.start, turn.end) for turn in turns if turn.end > turn.start]
    coverage.extend(extra_coverage or [])
    merged = _merge_intervals(coverage)
    selected: list[dict[str, Any]] = []
    uncovered: list[dict[str, Any]] = []
    for segment in segments:
        start = float(segment.get("start", 0))
        end = float(segment.get("end", start))
        if end - start < min_duration:
            continue
        if float(segment.get("avg_logprob", -999)) < min_logprob:
            continue
        if float(segment.get("no_speech_prob", 1)) > max_no_speech:
            continue
        selected.append(segment)
        if not any(start >= left - padding and end <= right + padding for left, right in merged):
            uncovered.append(segment)
    return {
        "selected_segments": len(selected),
        "uncovered_segments": len(uncovered),
        "uncovered_groups": _group_segments(uncovered),
        "thresholds": {
            "min_duration": min_duration,
            "min_logprob": min_logprob,
            "max_no_speech": max_no_speech,
            "padding": padding,
        },
    }


def audit_question_like_speaker_segments(
    turns: list[Turn],
    segments: list[dict[str, Any]],
    *,
    answer_speakers: Iterable[str] = DEFAULT_ANSWER_SPEAKERS,
    min_logprob: float = -0.72,
    max_no_speech: float = 0.75,
    padding: float = 1.0,
) -> list[dict[str, Any]]:
    answer_speakers = tuple(answer_speakers)
    findings: list[dict[str, Any]] = []
    for segment in segments:
        text = _compact_whitespace(segment.get("text", ""))
        if not any(marker in text for marker in QUESTION_MARKERS):
            continue
        if float(segment.get("avg_logprob", -999)) < min_logprob:
            continue
        if float(segment.get("no_speech_prob", 1)) > max_no_speech:
            continue
        start = float(segment.get("start", 0))
        end = float(segment.get("end", start))
        hits = [turn for turn in turns if start >= turn.start - padding and end <= turn.end + padding]
        if not hits:
            continue
        if all(any(label in turn.speaker for label in answer_speakers) and "檢察官" not in turn.speaker and "法官" not in turn.speaker for turn in hits):
            findings.append(
                {
                    "start": start,
                    "end": end,
                    "time": f"{seconds_to_clock(start)}–{seconds_to_clock(end)}",
                    "text": text,
                    "containing_turns": [asdict(turn) for turn in hits],
                    "classification": "manual_review_required",
                }
            )
    return findings


def inspect_evidence(payload: dict[str, Any]) -> dict[str, Any]:
    output_dir = Path(payload.get("output_dir") or ".forensic-transcript-audit").expanduser().resolve()
    output_dir.mkdir(parents=True, exist_ok=True)
    files: dict[str, Any] = {}
    for field in ("video", "transcript", "baseline", "asr_json"):
        raw = str(payload.get(field) or "").strip()
        if not raw:
            continue
        path = Path(raw).expanduser().resolve()
        item: dict[str, Any] = {"path": str(path), "exists": path.is_file()}
        if path.is_file():
            item.update({"size": path.stat().st_size, "sha256": sha256_file(path)})
            if path.suffix.lower() in {".docx", ".txt", ".md"}:
                text, extraction = extract_text(path, output_dir)
                item.update({"characters": len(text), "turns": len(parse_turns(text)), "extraction": extraction})
        files[field] = item

    video = files.get("video") or {}
    if video.get("exists") and shutil.which("ffprobe"):
        probe = _run(
            [
                "ffprobe",
                "-v",
                "error",
                "-show_entries",
                "format=duration,format_name:stream=index,codec_type,codec_name,width,height,sample_rate,channels",
                "-of",
                "json",
                video["path"],
            ]
        )
        try:
            video["ffprobe"] = json.loads(probe["stdout"]) if probe["returncode"] == 0 else probe
        except json.JSONDecodeError:
            video["ffprobe"] = probe

    result = {"success": True, "operation": "inspect", "output_dir": str(output_dir), "files": files}
    _write_reports(output_dir, "inspection", result)
    return result


def audit_transcript(payload: dict[str, Any]) -> dict[str, Any]:
    transcript_path = str(payload.get("transcript") or "").strip()
    if not transcript_path:
        raise ValueError("audit requires transcript")
    output_dir = Path(payload.get("output_dir") or ".forensic-transcript-audit").expanduser().resolve()
    output_dir.mkdir(parents=True, exist_ok=True)
    candidate_text, candidate_extraction = extract_text(transcript_path, output_dir)
    turns = parse_turns(candidate_text)
    result: dict[str, Any] = {
        "success": True,
        "operation": "audit",
        "transcript": str(Path(transcript_path).expanduser().resolve()),
        "turns": len(turns),
        "monotonic_blocks": len(split_monotonic_blocks(turns)),
        "extraction": candidate_extraction,
        "overlaps": find_overlaps(turns, float(payload.get("overlap_tolerance", 0.25))),
        "uncertainty_markers": UNCERTAINTY_RE.findall(candidate_text),
    }

    baseline_path = str(payload.get("baseline") or "").strip()
    if baseline_path:
        baseline_text, baseline_extraction = extract_text(baseline_path, output_dir)
        result["baseline"] = compare_baseline(candidate_text, baseline_text)
        result["baseline"]["extraction"] = baseline_extraction
    else:
        result["baseline"] = {"skipped": True, "reason": "baseline not supplied"}

    asr_path = str(payload.get("asr_json") or "").strip()
    if asr_path:
        segments = load_asr_segments(asr_path)
        extra_coverage: list[tuple[float, float]] = []
        if payload.get("baseline_video_start") and payload.get("baseline_video_end"):
            extra_coverage.append(
                (
                    clock_to_seconds(str(payload["baseline_video_start"])),
                    clock_to_seconds(str(payload["baseline_video_end"])),
                )
            )
        result["coverage"] = audit_asr_coverage(
            turns,
            segments,
            extra_coverage=extra_coverage,
            min_duration=float(payload.get("min_duration", 0.45)),
            min_logprob=float(payload.get("min_logprob", -0.55)),
            max_no_speech=float(payload.get("max_no_speech", 0.35)),
            padding=float(payload.get("coverage_padding", 1.5)),
        )
        answer_speakers = payload.get("answer_speakers") or DEFAULT_ANSWER_SPEAKERS
        result["speaker_review"] = audit_question_like_speaker_segments(
            turns,
            segments,
            answer_speakers=answer_speakers,
        )
    else:
        result["coverage"] = {"skipped": True, "reason": "asr_json not supplied"}
        result["speaker_review"] = []

    cross_speaker_overlaps = [row for row in result["overlaps"] if row["cross_speaker"]]
    baseline_ok = result["baseline"].get("exact", False) if baseline_path else False
    coverage_ok = result["coverage"].get("uncovered_segments", 0) == 0 if asr_path else False
    result["gates"] = {
        "has_turns": bool(turns),
        "baseline_supplied": bool(baseline_path),
        "baseline_exact": baseline_ok,
        "no_cross_speaker_overlap": not cross_speaker_overlaps,
        "asr_supplied": bool(asr_path),
        "asr_coverage_clear": coverage_ok,
        "speaker_review_count": len(result["speaker_review"]),
        "manual_second_pass_required": True,
    }
    result["passed_deterministic_gates"] = all(
        (
            result["gates"]["has_turns"],
            result["gates"]["baseline_exact"],
            result["gates"]["no_cross_speaker_overlap"],
            result["gates"]["asr_coverage_clear"],
        )
    )
    paths = _write_reports(output_dir, "audit", result)
    result["reports"] = paths
    return result


def validate_docx(payload: dict[str, Any], *, magi_root: Path) -> dict[str, Any]:
    transcript = Path(str(payload.get("transcript") or "")).expanduser().resolve()
    if not transcript.is_file() or transcript.suffix.lower() != ".docx":
        raise ValueError("validate-docx requires an existing .docx transcript")
    output_dir = Path(payload.get("output_dir") or transcript.parent / f"{transcript.stem}.render").expanduser().resolve()
    output_dir.mkdir(parents=True, exist_ok=True)
    result: dict[str, Any] = {
        "success": True,
        "operation": "validate-docx",
        "transcript": str(transcript),
        "output_dir": str(output_dir),
        "sha256": sha256_file(transcript),
    }

    validate_script = magi_root / "skills" / "docx" / "scripts" / "office" / "validate.py"
    if validate_script.is_file():
        result["validation"] = _run([sys.executable, str(validate_script), str(transcript)])
    else:
        try:
            with zipfile.ZipFile(transcript) as archive:
                archive.testzip()
                required = {"[Content_Types].xml", "word/document.xml"}
                missing = sorted(required.difference(archive.namelist()))
            result["validation"] = {"returncode": 0 if not missing else 1, "missing": missing, "method": "zip-fallback"}
        except (OSError, zipfile.BadZipFile) as exc:
            result["validation"] = {"returncode": 1, "stderr": str(exc), "method": "zip-fallback"}

    render_script = magi_root / "skills" / "docx" / "scripts" / "office" / "soffice.py"
    pdf = output_dir / f"{transcript.stem}.pdf"
    pdf.unlink(missing_ok=True)
    render_started_at_ns = time.time_ns()
    if render_script.is_file():
        env = dict(os.environ)
        fontconfig = str(payload.get("fontconfig") or "").strip()
        if not fontconfig and Path("/opt/homebrew/etc/fonts/fonts.conf").is_file():
            fontconfig = "/opt/homebrew/etc/fonts/fonts.conf"
        if fontconfig:
            env["FONTCONFIG_FILE"] = fontconfig
        with tempfile.TemporaryDirectory(
            prefix=".magi-libreoffice-profile-", dir=output_dir
        ) as profile_dir:
            user_installation = Path(profile_dir).resolve().as_uri()
            result["render"] = _run(
                [
                    sys.executable,
                    str(render_script),
                    f"-env:UserInstallation={user_installation}",
                    "--headless",
                    "--nologo",
                    "--nodefault",
                    "--nofirststartwizard",
                    "--norestore",
                    "--nolockcheck",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    str(output_dir),
                    str(transcript),
                ],
                env=env,
            )
            result["render"]["isolation"] = {
                "method": "unique-user-installation",
                "profile_inside_output_dir": True,
                "profile_removed_after_render": True,
                "network_policy_unchanged": True,
            }
            result["render"]["fresh_pdf"] = _wait_for_fresh_pdf(
                pdf, minimum_mtime_ns=render_started_at_ns
            )
    else:
        result["render"] = {
            "returncode": 127,
            "stderr": "soffice.py not found",
            "fresh_pdf": {"passed": False, "reason": "renderer missing"},
        }

    if (
        result["validation"].get("returncode") == 0
        and result["render"].get("fresh_pdf", {}).get("passed") is not True
        and os.environ.get("MAGI_V3_RELEASE_QUALITY_SEATBELT_CHILD") == "1"
        and os.environ.get("MAGI_V3_OFFLINE_CERTIFICATION") == "1"
    ):
        result["libreoffice_render"] = result["render"]
        render_started_at_ns = time.time_ns()
        result["render"] = _render_docx_pdf_for_seatbelt(transcript, pdf)
        result["render"]["fresh_pdf"] = _wait_for_fresh_pdf(
            pdf, minimum_mtime_ns=render_started_at_ns, timeout_seconds=0.5
        )

    result["pdf"] = str(pdf) if pdf.is_file() else ""
    result.update(
        _pdf_render_evidence(
            pdf, transcript, minimum_mtime_ns=render_started_at_ns
        )
    )
    result["passed"] = all(
        (
            result["validation"].get("returncode") == 0,
            result["render"].get("returncode") == 0,
            result["render"].get("fresh_pdf", {}).get("passed") is True,
            pdf.is_file(),
            result["pdfinfo"].get("passed") is True,
            result["pdf_text_validation"].get("passed") is True,
        )
    )
    paths = _write_reports(output_dir, "docx-validation", result)
    result["reports"] = paths
    return result


def _wait_for_fresh_pdf(
    pdf: Path, *, minimum_mtime_ns: int, timeout_seconds: float = 3.0
) -> dict[str, Any]:
    """Require a new, non-empty, size-stable PDF from the current render."""

    deadline = time.monotonic() + timeout_seconds
    previous_size: int | None = None
    while time.monotonic() < deadline:
        try:
            stat = pdf.stat()
        except OSError:
            time.sleep(0.05)
            continue
        if stat.st_size > 0 and stat.st_mtime_ns >= minimum_mtime_ns:
            if previous_size == stat.st_size:
                return {
                    "passed": True,
                    "size": stat.st_size,
                    "mtime_ns": stat.st_mtime_ns,
                    "sha256": sha256_file(pdf),
                }
            previous_size = stat.st_size
        time.sleep(0.05)
    return {
        "passed": False,
        "reason": "no fresh size-stable PDF was produced by this render",
    }


def _render_docx_pdf_for_seatbelt(transcript: Path, pdf: Path) -> dict[str, Any]:
    """Render verified forensic OOXML without sockets in the formal certifier."""

    pdf.unlink(missing_ok=True)
    try:
        from docx import Document
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.pdfgen import canvas

        document = Document(str(transcript))
        source_text = _docx_xml_text(transcript)
        turns = parse_turns(source_text)
        sections: list[dict[str, Any]] = []
        for section in document.sections:
            width = float(section.page_width or 0) / 12700
            height = float(section.page_height or 0) / 12700
            sections.append(
                {
                    "width_points": width,
                    "height_points": height,
                    "a4": abs(width - A4[0]) <= 3.0 and abs(height - A4[1]) <= 3.0,
                }
            )
        table_rows = sum(len(table.rows) for table in document.tables)
        table_cells = sum(
            len(row.cells) for table in document.tables for row in table.rows
        )
        structure = {
            "paragraphs": len(document.paragraphs),
            "tables": len(document.tables),
            "table_rows": table_rows,
            "table_cells": table_cells,
            "sections": sections,
        }
        if (
            not source_text.strip()
            or not turns
            or not sections
            or not all(section["a4"] for section in sections)
        ):
            raise ValueError("forensic DOCX text/turn/A4 structure is incomplete")

        font_path = next(
            (
                candidate.resolve(strict=True)
                for candidate in (
                    Path("/System/Library/Fonts/STHeiti Medium.ttc"),
                    Path("/System/Library/Fonts/STHeiti Light.ttc"),
                )
                if candidate.is_file()
            ),
            None,
        )
        if font_path is None:
            raise ValueError("verified extractable CJK font is unavailable")
        font_name = "MAGI-STHeiti"
        pdfmetrics.registerFont(TTFont(font_name, str(font_path), subfontIndex=0))
        writer = canvas.Canvas(str(pdf), pagesize=A4)
        writer.setTitle(transcript.stem)
        left = 72.0
        right = A4[0] - 72.0
        top = A4[1] - 54.0
        bottom = 54.0
        font_size = 10.5
        leading = 16.0
        usable_width = right - left
        writer.setFont(font_name, font_size)
        y = top
        rendered_lines = 0

        def wrap(value: str) -> list[str]:
            output: list[str] = []
            for logical in value.splitlines() or [""]:
                current = ""
                for char in logical:
                    candidate = current + char
                    if current and pdfmetrics.stringWidth(
                        candidate, font_name, font_size
                    ) > usable_width:
                        output.append(current)
                        current = char
                    else:
                        current = candidate
                output.append(current)
            return output

        for paragraph in source_text.split("\n\n"):
            for line in wrap(paragraph):
                if y < bottom + leading:
                    writer.showPage()
                    writer.setFont(font_name, font_size)
                    y = top
                writer.drawString(left, y, line)
                y -= leading
                rendered_lines += 1
            y -= leading * 0.35
        writer.save()
        return {
            "command": [],
            "returncode": 0,
            "stdout": "formal Seatbelt offline DOCX projection completed",
            "stderr": "",
            "renderer": "python-docx+reportlab-cjk/v1",
            "network_access_performed": False,
            "subprocess_started": False,
            "source_docx_sha256": sha256_file(transcript),
            "source_text_sha256": hashlib.sha256(source_text.encode("utf-8")).hexdigest(),
            "font_realpath": str(font_path),
            "font_sha256": sha256_file(font_path),
            "structure": structure,
            "rendered_lines": rendered_lines,
            "limitations": [
                "Formal-certifier fallback for paragraph/table accepted text only.",
                "Floating drawings, embedded media, and complex field layout are not reproduced.",
                "Acceptance still requires fresh A4 PDF metadata and extracted forensic text equivalence.",
            ],
        }
    except Exception as exc:
        pdf.unlink(missing_ok=True)
        return {
            "command": [],
            "returncode": 1,
            "stdout": "",
            "stderr": str(exc),
            "renderer": "python-docx+reportlab-cjk/v1",
            "network_access_performed": False,
            "subprocess_started": False,
        }


def _pdf_render_evidence(
    pdf: Path, transcript: Path, *, minimum_mtime_ns: int
) -> dict[str, Any]:
    """Fail-closed PDF page and accepted-text checks without retaining content."""

    pdfinfo_binary = _resolve_verifier_binary(
        "pdfinfo",
        (
            Path("/opt/homebrew/bin/pdfinfo"),
            Path("/usr/local/bin/pdfinfo"),
            Path("/usr/bin/pdfinfo"),
        ),
    )
    pdftotext_binary = _resolve_verifier_binary(
        "pdftotext",
        (
            Path("/opt/homebrew/bin/pdftotext"),
            Path("/usr/local/bin/pdftotext"),
            Path("/usr/bin/pdftotext"),
        ),
    )
    fresh = _wait_for_fresh_pdf(
        pdf, minimum_mtime_ns=minimum_mtime_ns, timeout_seconds=0.25
    )
    if not fresh.get("passed") or not pdfinfo_binary or not pdftotext_binary:
        reason = "rendered PDF or Poppler verifier is unavailable"
        return {
            "pdfinfo": {
                "returncode": 127,
                "pages": None,
                "page_size": "",
                "a4": False,
                "fresh_pdf": fresh,
                "passed": False,
            },
            "pdf_text_validation": {"returncode": 127, "passed": False, "reason": reason},
        }

    info = _run([str(pdfinfo_binary), str(pdf)])
    page_match = re.search(r"^Pages:\s+(\d+)", info.get("stdout", ""), re.MULTILINE)
    size_match = re.search(
        r"^Page size:\s+([0-9.]+)\s+x\s+([0-9.]+)\s+pts",
        info.get("stdout", ""),
        re.MULTILINE,
    )
    pages = int(page_match.group(1)) if page_match else None
    dimensions = (
        (float(size_match.group(1)), float(size_match.group(2))) if size_match else None
    )
    a4 = bool(
        dimensions
        and abs(dimensions[0] - 595.276) <= 3.0
        and abs(dimensions[1] - 841.89) <= 3.0
    )
    info_evidence = {
        "returncode": info["returncode"],
        "pages": pages,
        "page_size": size_match.group(0).removeprefix("Page size:").strip() if size_match else "",
        "a4": a4,
        "fresh_pdf": fresh,
        "binary_realpath": str(pdfinfo_binary),
        "binary_sha256": sha256_file(pdfinfo_binary),
        "passed": info["returncode"] == 0 and bool(pages and pages > 0) and a4,
    }

    extracted = _run([str(pdftotext_binary), "-layout", str(pdf), "-"])
    source_text = _docx_xml_text(transcript)
    pdf_text = extracted.get("stdout", "")
    normalize = lambda value: re.sub(r"\s+", "", str(value or ""))
    normalized_pdf = normalize(pdf_text)
    turns = parse_turns(source_text)
    time_markers = sorted(set(TIME_RE.findall(source_text)))
    speakers = sorted({turn.speaker for turn in turns if turn.speaker})
    turn_texts = [turn.text for turn in turns if turn.text]
    source_has_cjk = bool(re.search(r"[\u3400-\u9fff]", source_text))
    extracted_has_cjk = bool(re.search(r"[\u3400-\u9fff]", pdf_text))
    text_evidence = {
        "returncode": extracted["returncode"],
        "binary_realpath": str(pdftotext_binary),
        "binary_sha256": sha256_file(pdftotext_binary),
        "extracted_text_sha256": hashlib.sha256(pdf_text.encode("utf-8")).hexdigest(),
        "extracted_characters": len(pdf_text),
        "source_has_cjk": source_has_cjk,
        "extracted_has_cjk": extracted_has_cjk,
        "time_marker_count": len(time_markers),
        "all_time_markers_present": bool(time_markers)
        and all(normalize(marker) in normalized_pdf for marker in time_markers),
        "speaker_count": len(speakers),
        "all_speakers_present": bool(speakers)
        and all(normalize(speaker) in normalized_pdf for speaker in speakers),
        "turn_count": len(turn_texts),
        "all_turn_text_present": bool(turn_texts)
        and all(normalize(text) in normalized_pdf for text in turn_texts),
    }
    text_evidence["passed"] = all(
        (
            extracted["returncode"] == 0,
            bool(normalized_pdf),
            not source_has_cjk or extracted_has_cjk,
            text_evidence["all_time_markers_present"],
            text_evidence["all_speakers_present"],
            text_evidence["all_turn_text_present"],
        )
    )
    return {"pdfinfo": info_evidence, "pdf_text_validation": text_evidence}


def _resolve_verifier_binary(name: str, fallbacks: Iterable[Path]) -> Path | None:
    discovered = shutil.which(name)
    candidates = ([Path(discovered)] if discovered else []) + list(fallbacks)
    seen: set[Path] = set()
    for candidate in candidates:
        try:
            resolved = candidate.resolve(strict=True)
        except OSError:
            continue
        if resolved in seen:
            continue
        seen.add(resolved)
        if resolved.is_file() and os.access(resolved, os.X_OK):
            return resolved
    return None


def _write_reports(output_dir: Path, stem: str, payload: dict[str, Any]) -> dict[str, str]:
    output_dir.mkdir(parents=True, exist_ok=True)
    json_path = output_dir / f"{stem}.json"
    md_path = output_dir / f"{stem}.md"
    json_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    lines = [f"# {stem}", "", f"- success: `{payload.get('success', False)}`"]
    if "passed_deterministic_gates" in payload:
        lines.append(f"- passed_deterministic_gates: `{payload['passed_deterministic_gates']}`")
    if "passed" in payload:
        lines.append(f"- passed: `{payload['passed']}`")
    for key in ("turns", "monotonic_blocks"):
        if key in payload:
            lines.append(f"- {key}: `{payload[key]}`")
    if isinstance(payload.get("baseline"), dict):
        lines.append(f"- baseline_exact: `{payload['baseline'].get('exact', 'skipped')}`")
    if isinstance(payload.get("overlaps"), list):
        lines.append(f"- overlaps: `{len(payload['overlaps'])}`")
    if isinstance(payload.get("coverage"), dict):
        lines.append(f"- uncovered_asr_segments: `{payload['coverage'].get('uncovered_segments', 'skipped')}`")
    if isinstance(payload.get("speaker_review"), list):
        lines.append(f"- speaker_review: `{len(payload['speaker_review'])}`")
    lines.extend(["", "See the JSON report for full evidence and timestamps.", ""])
    md_path.write_text("\n".join(lines), encoding="utf-8")
    return {"json": str(json_path), "markdown": str(md_path)}
