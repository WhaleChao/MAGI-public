#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
brief-gen/action.py

書狀輔助產生技能
- 依案件類型產生書狀草稿架構
- 支援多種書狀類型（起訴狀、答辯狀、上訴狀等）
- 可匯出為 Word 文件
"""

from __future__ import annotations
import logging

import argparse
import json
import os
import re
import subprocess
import sys
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

_MAGI_ROOT = Path(os.environ.get("MAGI_ROOT", str(Path(__file__).resolve().parents[2])))
if str(_MAGI_ROOT) not in sys.path:
    sys.path.insert(0, str(_MAGI_ROOT))

CASE_BASE = Path(os.environ.get(
    "MAGI_CASE_BASE",
    "/Users/ai/Library/CloudStorage/SynologyDrive-homes/01_案件",
))

# ── 書狀範本定義 ──────────────────────────────────────────────────

BRIEF_TEMPLATES: Dict[str, Dict[str, Any]] = {
    "complaint": {
        "name": "民事起訴狀",
        "sections": [
            {"title": "案號", "hint": "（由法院分案後填入）"},
            {"title": "原告", "hint": "姓名、身分證字號、住址、送達代收人及地址"},
            {"title": "被告", "hint": "姓名、身分證字號、住址"},
            {"title": "訴訟標的金額或價額", "hint": "新臺幣 _____ 元"},
            {"title": "訴之聲明", "hint": "一、被告應給付原告新臺幣___元，及自起訴狀繕本送達翌日起至清償日止，按年息百分之五計算之利息。\n二、訴訟費用由被告負擔。\n三、原告願供擔保，請准宣告假執行。"},
            {"title": "事實及理由", "subsections": [
                "壹、事實經過",
                "貳、法律依據",
                "參、損害計算",
            ]},
            {"title": "證據清單", "hint": "附件一：___\n附件二：___"},
            {"title": "附屬文件", "hint": "一、本狀繕本 __ 份\n二、證物影本 __ 份"},
        ],
    },
    "answer": {
        "name": "民事答辯狀",
        "sections": [
            {"title": "案號"},
            {"title": "被告（答辯人）", "hint": "姓名、身分證字號、住址"},
            {"title": "原告", "hint": "姓名"},
            {"title": "答辯聲明", "hint": "一、原告之訴駁回。\n二、訴訟費用由原告負擔。\n三、如受不利判決，被告願供擔保，請准宣告免為假執行。"},
            {"title": "答辯理由", "subsections": [
                "壹、原告主張不實之處",
                "貳、被告之抗辯",
                "參、法律見解",
            ]},
            {"title": "證據清單"},
        ],
    },
    "appeal": {
        "name": "上訴狀",
        "sections": [
            {"title": "案號", "hint": "原審案號：___"},
            {"title": "上訴人", "hint": "姓名、身分證字號、住址"},
            {"title": "被上訴人", "hint": "姓名"},
            {"title": "上訴聲明", "hint": "一、原判決廢棄。\n二、___（依案件類型填寫）"},
            {"title": "上訴理由", "subsections": [
                "壹、原判決違背法令之處",
                "貳、原判決認定事實錯誤之處",
                "參、補充證據",
            ]},
            {"title": "證據清單"},
        ],
    },
    "motion": {
        "name": "聲請狀",
        "sections": [
            {"title": "案號"},
            {"title": "聲請人", "hint": "姓名、身分證字號、住址"},
            {"title": "相對人", "hint": "姓名"},
            {"title": "聲請事項", "hint": "請准予___"},
            {"title": "聲請理由", "subsections": [
                "壹、事實經過",
                "貳、聲請之依據",
                "參、釋明事項",
            ]},
            {"title": "附件"},
        ],
    },
    "closing": {
        "name": "辯論意旨狀",
        "sections": [
            {"title": "案號"},
            {"title": "原告／被告"},
            {"title": "辯論要旨", "subsections": [
                "壹、本案爭點整理",
                "貳、就各爭點之主張",
                "參、證據評價",
                "肆、結論",
            ]},
            {"title": "聲明"},
        ],
    },
    "statement": {
        "name": "準備書狀（陳報狀）",
        "sections": [
            {"title": "案號"},
            {"title": "陳報人"},
            {"title": "陳報事項", "subsections": [
                "壹、就鈞院諭知事項之說明",
                "貳、補充事實及理由",
                "參、補提證據",
            ]},
        ],
    },
    "labor": {
        "name": "勞動調解聲請狀",
        "sections": [
            {"title": "案號"},
            {"title": "聲請人（勞工）", "hint": "姓名、身分證字號、住址、電話"},
            {"title": "相對人（雇主）", "hint": "公司名稱、統一編號、地址、法定代理人"},
            {"title": "聲請調解事項", "hint": "一、相對人應給付聲請人新臺幣___元。\n二、相對人應開立非自願離職證明。"},
            {"title": "事實及理由", "subsections": [
                "壹、勞動關係說明（到職日、職稱、薪資）",
                "貳、爭議事實",
                "參、法律依據（勞基法條文）",
                "肆、請求金額計算",
            ]},
            {"title": "證據清單", "hint": "附件一：勞動契約\n附件二：薪資單\n附件三：出勤紀錄"},
        ],
    },
}

# ── 指令實作 ──────────────────────────────────────────────────────


def _cmd_template(text: str) -> str:
    """列出可用書狀範本。"""
    if text:
        text_lower = text.strip().lower()
        # Match specific template
        for key, tmpl in BRIEF_TEMPLATES.items():
            if text_lower in key or text_lower in tmpl["name"]:
                return _render_template_detail(key, tmpl)

    lines = ["📝 可用書狀範本", ""]
    for key, tmpl in BRIEF_TEMPLATES.items():
        sections = tmpl.get("sections", [])
        lines.append(f"  {key:<12} {tmpl['name']}（{len(sections)} 個段落）")
    lines.append("")
    lines.append("使用 --task template --text \"類型\" 查看詳細結構。")
    lines.append("使用 --task draft --text \"案件描述\" 產生草稿。")
    return "\n".join(lines)


def _render_template_detail(key: str, tmpl: Dict[str, Any]) -> str:
    """Render detailed template structure."""
    lines = [f"📝 書狀範本：{tmpl['name']}（{key}）", ""]
    for i, sec in enumerate(tmpl.get("sections", []), 1):
        title = sec.get("title", "")
        hint = sec.get("hint", "")
        subsections = sec.get("subsections", [])
        lines.append(f"{i}. {title}")
        if hint:
            for h in hint.split("\n"):
                lines.append(f"   {h}")
        if subsections:
            for j, sub in enumerate(subsections, 1):
                lines.append(f"   {j}) {sub}")
    return "\n".join(lines)


def _detect_brief_type(text: str) -> str:
    """Auto-detect brief type from text description."""
    text_lower = text.lower()
    type_keywords = {
        "complaint": ["起訴", "提告", "告"],
        "answer": ["答辯", "被告"],
        "appeal": ["上訴", "抗告"],
        "motion": ["聲請", "假扣押", "假處分", "保全"],
        "closing": ["辯論意旨", "最後言詞", "辯論"],
        "statement": ["準備書狀", "陳報", "補充"],
        "labor": ["勞動", "勞資", "勞調", "勞基法", "資遣", "加班費"],
    }
    for brief_type, keywords in type_keywords.items():
        for kw in keywords:
            if kw in text_lower:
                return brief_type
    return "complaint"  # default


def _cmd_draft(text: str) -> str:
    """產生書狀草稿架構。"""
    if not text:
        return "⚠️ 請描述案件情況，例如：\n--task draft --text \"勞資爭議 雇主未付加班費 請求給付加班費及資遣費\""

    # Detect brief type
    brief_type = _detect_brief_type(text)
    tmpl = BRIEF_TEMPLATES.get(brief_type, BRIEF_TEMPLATES["complaint"])

    # Extract case info from text
    case_no = _extract_case_number(text)

    lines = [
        f"📝 書狀草稿 — {tmpl['name']}",
        f"案件描述：{text[:100]}",
        f"自動判斷書狀類型：{brief_type}（{tmpl['name']}）",
        "",
        "=" * 50,
        f"　　　　　　{tmpl['name']}",
        "=" * 50,
        "",
    ]

    for sec in tmpl.get("sections", []):
        title = sec.get("title", "")
        hint = sec.get("hint", "")
        subsections = sec.get("subsections", [])

        # Auto-fill case number if available
        if title == "案號" and case_no:
            lines.append(f"{title}：{case_no}")
        elif hint:
            lines.append(f"{title}：")
            for h in hint.split("\n"):
                lines.append(f"  {h}")
        elif subsections:
            lines.append(f"{title}：")
            for sub in subsections:
                lines.append(f"  {sub}")
                lines.append(f"    （待填寫）")
                lines.append("")
        else:
            lines.append(f"{title}：（待填寫）")
        lines.append("")

    # Query related statutes
    case_keywords = _extract_legal_keywords(text)
    if case_keywords:
        lines.append("【建議引用法條】")
        statutes = _query_statutes(case_keywords)
        if statutes:
            for sl in statutes.splitlines()[:8]:
                lines.append(f"  {sl}")
        else:
            lines.append("  （法條查詢中或無結果）")
        lines.append("")

    # Query related judgments
    if case_keywords:
        lines.append("【參考判決見解】")
        judgments = _query_judgments(case_keywords)
        if judgments:
            for jl in judgments.splitlines()[:8]:
                lines.append(f"  {jl}")
        else:
            lines.append("  （見解查詢中或無結果）")
        lines.append("")

    lines.append("=" * 50)
    lines.append("提示：此為草稿架構，請律師審閱修改後使用。")
    lines.append("使用 --task export --text \"案號\" --mode docx 可匯出為 Word 文件。")
    return "\n".join(lines)


def _cmd_enrich(text: str) -> str:
    """從案件資料自動充實書狀內容。"""
    if not text:
        return "⚠️ 請指定案號。"

    case_no = _extract_case_number(text) or text.strip()
    folder = _find_case_folder(case_no)
    if not folder:
        return f"⚠️ 找不到案號「{case_no}」的資料夾。"

    cats = _scan_case_folder(folder)
    lines = [
        f"📝 書狀充實資料 — {case_no}",
        f"資料夾：{folder.name}",
        "",
    ]

    # Summarize existing briefs
    if cats["briefs"]:
        lines.append("【既有書狀】")
        for b in cats["briefs"]:
            lines.append(f"  - {b}")
        lines.append("")

    # Summarize transcripts for key points
    if cats["transcripts"]:
        lines.append("【筆錄摘要】")
        for t in cats["transcripts"]:
            lines.append(f"  - {t}")
        lines.append("  （建議使用 transcript-indexer 搜尋特定爭點）")
        lines.append("")

    # Evidence inventory
    if cats["evidence"]:
        lines.append("【證據清單（可直接引用）】")
        for i, e in enumerate(cats["evidence"], 1):
            lines.append(f"  附件{_chinese_num(i)}：{e}")
        lines.append("")

    if not any(cats.values()):
        lines.append("⚠️ 案件資料夾中沒有找到文件。")

    return "\n".join(lines)


def _cmd_export(text: str, mode: str = "docx") -> str:
    """匯出書狀為 Word 文件。"""
    if not text:
        return "⚠️ 請指定案號或書狀內容。"

    case_no = _extract_case_number(text) or text.strip()
    brief_type = _detect_brief_type(text)
    tmpl = BRIEF_TEMPLATES.get(brief_type, BRIEF_TEMPLATES["complaint"])

    # Generate docx via docx skill
    try:
        docx_skill = str(_MAGI_ROOT / "skills" / "docx" / "action.py")
        py = os.environ.get("MAGI_SKILL_PYTHON", "") or "python3"
        if not os.path.exists(docx_skill):
            return "⚠️ docx skill 不存在，無法匯出 Word 文件。"

        # Build document structure
        paragraphs = []
        paragraphs.append({"text": tmpl["name"], "style": "Title"})
        paragraphs.append({"text": f"案號：{case_no or '（待填）'}", "style": "Normal"})
        paragraphs.append({"text": f"日期：{datetime.now().strftime('%Y年%m月%d日')}", "style": "Normal"})
        paragraphs.append({"text": "", "style": "Normal"})

        for sec in tmpl.get("sections", []):
            title = sec.get("title", "")
            hint = sec.get("hint", "")
            subsections = sec.get("subsections", [])
            paragraphs.append({"text": title, "style": "Heading 1"})
            if hint:
                paragraphs.append({"text": hint, "style": "Normal"})
            if subsections:
                for sub in subsections:
                    paragraphs.append({"text": sub, "style": "Heading 2"})
                    paragraphs.append({"text": "（待填寫）", "style": "Normal"})

        # Determine output path
        folder = _find_case_folder(case_no) if case_no else None
        if folder:
            output_path = str(folder / f"{tmpl['name']}_{datetime.now().strftime('%Y%m%d')}.docx")
        else:
            export_dir = _MAGI_ROOT / "static" / "exports"
            export_dir.mkdir(parents=True, exist_ok=True)
            output_path = str(export_dir / f"{tmpl['name']}_{case_no or 'draft'}_{datetime.now().strftime('%Y%m%d')}.docx")

        payload = json.dumps({
            "paragraphs": paragraphs,
            "output_path": output_path,
            "font": "標楷體",
            "font_size": 14,
        }, ensure_ascii=False)

        result = subprocess.run(
            [py, docx_skill, "--task", "create", "--text", payload],
            capture_output=True, timeout=30, text=True,
            cwd=str(_MAGI_ROOT),
        )
        if result.returncode == 0 and os.path.exists(output_path):
            return f"✅ 書狀已匯出：{output_path}"

        # Fallback: try direct python-docx
        return _export_docx_direct(tmpl, case_no, output_path)
    except Exception as e:
        return f"⚠️ 匯出失敗：{type(e).__name__}: {e}"


def _export_docx_direct(tmpl: Dict[str, Any], case_no: str, output_path: str) -> str:
    """Direct export using python-docx."""
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Page margins
        for section in doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        # Title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(tmpl["name"])
        run.font.size = Pt(18)
        run.bold = True

        doc.add_paragraph(f"案號：{case_no or '（待填）'}")
        doc.add_paragraph(f"日期：{datetime.now().strftime('%Y年%m月%d日')}")
        doc.add_paragraph("")

        for sec in tmpl.get("sections", []):
            title = sec.get("title", "")
            hint = sec.get("hint", "")
            subsections = sec.get("subsections", [])

            heading = doc.add_heading(title, level=1)
            if hint:
                doc.add_paragraph(hint)
            if subsections:
                for sub in subsections:
                    doc.add_heading(sub, level=2)
                    doc.add_paragraph("（待填寫）")

        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        return f"✅ 書狀已匯出：{output_path}"
    except ImportError:
        return "⚠️ python-docx 未安裝，無法匯出 Word 文件。請執行 pip install python-docx"
    except Exception as e:
        return f"⚠️ 匯出失敗：{type(e).__name__}: {e}"


# ── 輔助函式 ──────────────────────────────────────────────────────


def _extract_case_number(text: str) -> str:
    """Extract Taiwan court case number from text."""
    m = re.search(r"(\d{2,3})\s*年度?\s*[\u4e00-\u9fff]{1,6}\s*字?\s*第?\s*(\d+)\s*號?", text)
    if m:
        return m.group(0)
    m = re.search(r"\d{2,3}年[\u4e00-\u9fff\d]+號", text)
    if m:
        return m.group(0)
    return ""


def _find_case_folder(case_no: str) -> Optional[Path]:
    """Find case folder on NAS by case number."""
    if not case_no or not CASE_BASE.exists():
        return None
    try:
        for entry in CASE_BASE.iterdir():
            if entry.is_dir() and case_no in entry.name:
                return entry
    except Exception:
        logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 455, exc_info=True)
    nums = re.findall(r"\d+", case_no)
    if len(nums) >= 2:
        try:
            for entry in CASE_BASE.iterdir():
                if entry.is_dir() and nums[0] in entry.name and nums[-1] in entry.name:
                    return entry
        except Exception:
            logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 463, exc_info=True)
    return None


def _scan_case_folder(folder: Path) -> Dict[str, List[str]]:
    """Scan case folder and categorize files."""
    categories: Dict[str, List[str]] = {
        "briefs": [], "transcripts": [], "evidence": [],
        "rulings": [], "other": [],
    }
    if not folder or not folder.exists():
        return categories
    try:
        for f in sorted(folder.rglob("*")):
            if not f.is_file():
                continue
            name = f.name.lower()
            rel = str(f.relative_to(folder))
            if any(kw in name for kw in ["書狀", "起訴", "答辯", "聲請", "準備", "辯論", "陳報", "上訴"]):
                categories["briefs"].append(rel)
            elif any(kw in name for kw in ["筆錄", "言詞", "準備程序", "調解"]):
                categories["transcripts"].append(rel)
            elif any(kw in name for kw in ["證據", "附件", "證物"]):
                categories["evidence"].append(rel)
            elif any(kw in name for kw in ["裁定", "判決"]):
                categories["rulings"].append(rel)
            elif f.suffix.lower() in {".pdf", ".docx", ".doc", ".jpg", ".png"}:
                categories["other"].append(rel)
    except Exception:
        logging.getLogger(__name__).debug("silent-catch at %s:%s", __name__, 492, exc_info=True)
    return categories


def _extract_legal_keywords(text: str) -> str:
    """Extract legal keywords for statute/judgment lookup."""
    keywords = []
    kw_map = {
        "資遣": "勞動基準法 資遣費",
        "加班": "勞動基準法 加班費 延長工時",
        "職災": "勞動基準法 職業災害補償",
        "離婚": "民法親屬編 離婚",
        "扶養": "民法親屬編 扶養",
        "繼承": "民法繼承編",
        "侵權": "民法 侵權行為",
        "契約": "民法債編 契約",
        "借貸": "民法 消費借貸",
        "租賃": "民法 租賃",
        "買賣": "民法 買賣",
        "保險": "保險法",
        "公司": "公司法",
        "票據": "票據法",
        "智財": "著作權法 專利法 商標法",
        "勞動": "勞動基準法 勞動事件法",
        "消費": "消費者保護法",
        "交通事故": "民法 侵權行為 強制汽車責任保險法",
    }
    for kw, laws in kw_map.items():
        if kw in text:
            keywords.append(laws)
    return " ".join(keywords) if keywords else text[:50]


def _query_statutes(keywords: str) -> str:
    """Query statutes-vdb skill."""
    try:
        skill_path = str(_MAGI_ROOT / "skills" / "statutes-vdb" / "action.py")
        py = os.environ.get("MAGI_SKILL_PYTHON", "") or "python3"
        if not os.path.exists(skill_path):
            return ""
        result = subprocess.run(
            [py, skill_path, "--task", "search", "--text", keywords],
            capture_output=True, timeout=30, text=True,
            cwd=str(_MAGI_ROOT),
        )
        return result.stdout.strip()[:2000] if result.returncode == 0 else ""
    except Exception:
        return ""


def _query_judgments(keywords: str) -> str:
    """Query judgment-collector skill."""
    try:
        skill_path = str(_MAGI_ROOT / "skills" / "judgment-collector" / "action.py")
        py = os.environ.get("MAGI_SKILL_PYTHON", "") or "python3"
        if not os.path.exists(skill_path):
            return ""
        result = subprocess.run(
            [py, skill_path, "--task", "search", "--text", keywords],
            capture_output=True, timeout=30, text=True,
            cwd=str(_MAGI_ROOT),
        )
        return result.stdout.strip()[:2000] if result.returncode == 0 else ""
    except Exception:
        return ""


def _chinese_num(n: int) -> str:
    """Convert number to Chinese numeral for attachment numbering."""
    nums = ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十",
            "十一", "十二", "十三", "十四", "十五", "十六", "十七", "十八", "十九", "二十"]
    if 1 <= n <= len(nums):
        return nums[n - 1]
    return str(n)


# ── 主程式 ────────────────────────────────────────────────────────


def main() -> int:
    ap = argparse.ArgumentParser(description="MAGI 書狀輔助產生")
    ap.add_argument("--task", default="template", help="draft|template|enrich|export")
    ap.add_argument("--text", default="", help="案件描述或案號")
    ap.add_argument("--mode", default="docx", help="匯出格式（docx）")
    args = ap.parse_args()

    task = str(args.task or "template").strip().lower()
    text = str(args.text or "").strip()
    mode = str(args.mode or "docx").strip().lower()

    if task in {"template", "list", "範本", "清單"}:
        print(_cmd_template(text))
        return 0
    if task in {"draft", "write", "草稿", "擬狀"}:
        print(_cmd_draft(text))
        return 0
    if task in {"enrich", "充實", "補充"}:
        print(_cmd_enrich(text))
        return 0
    if task in {"export", "匯出", "docx", "word"}:
        print(_cmd_export(text, mode=mode))
        return 0

    print("⚠️ 不支援的 task，請使用：template|draft|enrich|export")
    return 2


if __name__ == "__main__":
    raise SystemExit(main())
