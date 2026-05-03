# -*- coding: utf-8 -*-
"""
docx_builder.py — M5 模組 1：產生補件書狀 docx

公開 API:
    build_supplement_docx(case_meta, extracted, matched, *, procedure, brief_seq,
                          output_path, template_path=None) -> dict
"""
from __future__ import annotations

import logging
import os
import re
import unicodedata
from datetime import datetime
from pathlib import Path
from typing import Optional

from .exceptions import SupplementError

logger = logging.getLogger("docx_builder")

# ── 預設模板路徑解析 ──────────────────────────────────────────────────────────

def _default_template_path() -> str:
    """解析預設模板路徑。優先用 MAGI_ROOT env，否則從 __file__ 往上三層。"""
    magi_root = os.environ.get("MAGI_ROOT", "")
    if magi_root:
        p = Path(magi_root) / "data" / "templates" / "D_supplement.docx"
    else:
        # src/supplement_core/docx_builder.py → MAGI_v2/
        p = Path(__file__).parent.parent.parent / "data" / "templates" / "D_supplement.docx"
    return str(p)


# ── 中文數字 ──────────────────────────────────────────────────────────────────

_CN_DIGITS = ["", "一", "二", "三", "四", "五", "六", "七", "八", "九"]


def _int_to_chinese(n: int) -> str:
    """將 int 轉為中文數字（支援 1~99）。"""
    if not (1 <= n <= 99):
        raise ValueError(f"n={n} 超出範圍 [1, 99]")
    if n < 10:
        return _CN_DIGITS[n]
    tens, ones = n // 10, n % 10
    prefix = "十" if tens == 1 else _CN_DIGITS[tens] + "十"
    return prefix if ones == 0 else prefix + _CN_DIGITS[ones]


# ── 段落替換（ port 自 05_E.py apply_inputs_to_doc / replace_in_paragraph）──

def _replace_in_paragraph(para, values: dict) -> None:
    """就地替換段落中的佔位符；樣式繼承首 run。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    original_text = "".join(run.text for run in para.runs)
    new_text = original_text
    changed = False
    for key in sorted(values, key=len, reverse=True):
        val = values[key]
        if val is not None and key in new_text:
            new_text = new_text.replace(key, val)
            changed = True
    if not changed:
        return

    base_run = para.runs[0] if para.runs else para.add_run("")
    while para.runs:
        para.runs[0]._element.getparent().remove(para.runs[0]._element)

    new_run = para.add_run(new_text)
    try:
        if base_run.font and base_run.font.name:
            new_run.font.name = base_run.font.name
            rPr = new_run._element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.append(rFonts)
            rFonts.set(qn("w:eastAsia"), base_run.font.name)
        new_run.font.size = base_run.font.size
        new_run.bold = base_run.font.bold
        new_run.italic = base_run.font.italic
        new_run.underline = base_run.font.underline
    except Exception as exc:
        logger.debug("樣式繼承失敗：%s", exc)


def _apply_fields_to_doc(doc, filled: dict, deleted: list[str]) -> None:
    """
    將 filled 填入 doc，值為 None（即「刪除此項」）的欄位整段刪除。
    ported from 05_E.py apply_inputs_to_doc()
    """
    from docx.oxml import OxmlElement

    # paragraph_hint_map：幫助用提示文字找到要刪的段落
    paragraph_hint_map = {
        "D4": "全戶戶籍謄本", "D5": "家族系統表", "D6": "綜合所得稅各類所得清單",
        "D7": "年金及健保資料", "D8": "勞保資料", "D9": "存摺影本或交易明細表",
        "D10": "集保公司資料", "D11": "壽險查詢結果", "D12": "社會補助或津貼",
        "D13": "財產變動情形", "D14": "公司營運情形", "D15": "財產及收入狀況說明書",
    }

    # 讓 B1/B2/B3、C1/C2/C3 等值為 None 的也列入 hint（用 key 本身比對）
    sorted_keys = sorted(filled.keys(), key=len, reverse=True)
    para_indices_to_delete: set[int] = set()

    for i, para in enumerate(doc.paragraphs):
        full_text = "".join(run.text for run in para.runs).strip()
        deleted_this = False
        for key in sorted_keys:
            if filled[key] is None:
                hint = paragraph_hint_map.get(key, "")
                if key in full_text or (hint and hint in full_text):
                    para_indices_to_delete.add(i)
                    deleted_this = True
                    break
        if not deleted_this:
            _replace_in_paragraph(para, {k: v for k, v in filled.items() if v is not None})

    # 從後往前刪除段落
    for idx in sorted(para_indices_to_delete, reverse=True):
        doc.paragraphs[idx]._element.getparent().remove(doc.paragraphs[idx]._element)

    # 表格處理
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    should_delete = False
                    for key in sorted_keys:
                        if filled[key] is None and key in para.text:
                            para.clear()
                            should_delete = True
                            break
                    if not should_delete:
                        _replace_in_paragraph(para, {k: v for k, v in filled.items() if v is not None})


def _apply_proof_table(doc, items: list[dict], matched: list[dict], label_to_proof: dict[str, str]) -> None:
    """替換模板末段 F/H 證據表，並清掉未使用列的欄位代碼。"""
    proof_values: dict[str, str] = {}
    proof_idx = 1
    for idx, item in enumerate(items[:15]):
        m = matched[idx] if idx < len(matched) else None
        if not m or m.get("status") != "have":
            continue
        label = label_to_proof.get(f"D{idx + 1}", "")
        if not label:
            continue

        selected = m.get("selected", "")
        filename = Path(selected).name if selected else ""
        category = item.get("category", f"補正資料{idx + 1}")
        proof_values[f"F{proof_idx}"] = label
        proof_values[f"H{proof_idx}"] = filename or f"{category}文件乙份。"
        proof_idx += 1

    for n in range(proof_idx, 16):
        proof_values[f"F{n}"] = ""
        proof_values[f"H{n}"] = ""

    all_paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paragraphs.extend(cell.paragraphs)

    for para in all_paragraphs:
        _replace_in_paragraph(para, proof_values)


# ── 補件 lead-in 段落插入 ─────────────────────────────────────────────────────

def _insert_lead_in(doc, extracted: dict, matched: list[dict], label_to_proof: dict) -> str:
    """在「謹就消費者債務清理事件」段落前或文件開頭插入補件說明段落。

    Returns: lead-in 文字（debug 用）
    """
    case_meta_e = extracted.get("case_meta", {})
    ruling_date = case_meta_e.get("ruling_date", "")
    case_no = case_meta_e.get("case_no", "")
    items = extracted.get("items", [])

    # 建立 lead-in 文字
    lines = []
    intro = f"緣鈞院 {ruling_date} {case_no} 民事裁定（下稱系爭裁定）所列補正事項，聲請人已備齊下列文件，謹依命提送，並逐項說明如下：".strip()
    lines.append(intro)
    lines.append("")

    cn_nums = ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十",
               "十一", "十二", "十三", "十四", "十五"]

    for idx, item in enumerate(items):
        proof_label = label_to_proof.get(f"D{idx+1}", f"【聲證{cn_nums[idx] if idx < len(cn_nums) else str(idx+1)}】")
        m = matched[idx] if idx < len(matched) else None
        attach_name = ""
        if m:
            cands = m.get("candidates", [])
            if m.get("selected"):
                attach_name = Path(m["selected"]).name
            elif cands:
                attach_name = cands[0].get("filename", "")

        item_line = f"{cn_nums[idx] if idx < len(cn_nums) else str(idx+1)}、{item.get('category', '')}：{item.get('quote', '')}"
        lines.append(item_line)
        attach_line = f"   {proof_label}{('　' + attach_name) if attach_name else ''}"
        lines.append(attach_line)
        lines.append("")

    lead_in_text = "\n".join(lines)

    # 尋找合適的插入點（paragraph 1 之前，即「謹就...陳報事：」前）
    anchor_para = None
    for para in doc.paragraphs:
        if "謹就消費者債務清理" in para.text or "陳報事：" in para.text:
            anchor_para = para
            break

    if anchor_para is None and doc.paragraphs:
        anchor_para = doc.paragraphs[0]

    if anchor_para is not None:
        new_para = anchor_para.insert_paragraph_before(lead_in_text)
        new_para.style = anchor_para.style
    else:
        # 文件沒有段落，直接加
        doc.add_paragraph(lead_in_text)

    return lead_in_text


# ── 主 API ────────────────────────────────────────────────────────────────────

def build_supplement_docx(
    case_meta: dict,
    extracted: dict,
    matched: list[dict],
    *,
    procedure: str,
    brief_seq: int,
    output_path: str,
    template_path: Optional[str] = None,
) -> dict:
    """產生補件書狀 docx 並寫到 output_path。

    Returns:
        {
            "output_path": str,
            "filled_fields": dict[str, str],
            "deleted_fields": list[str],
            "main_content": str,
            "warnings": list[str],
        }
    """
    from docx import Document

    warnings: list[str] = []

    # ── 1. 模板 ──────────────────────────────────────────────────────────────
    if template_path is None:
        template_path = _default_template_path()
    template_path = unicodedata.normalize("NFC", str(template_path))
    if not os.path.isfile(template_path):
        raise SupplementError(f"模板不存在：{template_path}")

    doc = Document(template_path)

    # ── 2. 欄位映射 ──────────────────────────────────────────────────────────
    today = datetime.today()
    roc_year = today.year - 1911
    g1 = f"中華民國{roc_year}年{today.month:02d}月{today.day:02d}日"

    brief_seq_cn = _int_to_chinese(brief_seq)

    # extracted.case_meta 由 LLM 填充（可能比 case_meta 更詳細）
    e_meta = extracted.get("case_meta", {})
    case_no = e_meta.get("case_no") or case_meta.get("case_no") or case_meta.get("case_year_seq", "")
    court = e_meta.get("court") or case_meta.get("court") or ""
    parties = case_meta.get("parties", [])
    party_name = parties[0] if parties else ""

    items = extracted.get("items", [])

    # 聲證順序（D1~D15 中哪些有值）
    proof_fields_order = [f"D{i}" for i in range(1, 16)]
    cn_nums_list = ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十",
                    "十一", "十二", "十三", "十四", "十五"]

    filled: dict[str, str | None] = {
        "A1": brief_seq_cn,
        "A2": case_no,
        "A3": "",
        "A4": party_name,
        "E1": court,
        "G1": g1,
        "B1": None,  # 刪除此項
        "B2": None,
        "B3": None,
        "C1": None,
        "C2": None,
        "C3": None,
    }

    # D1~D15
    if len(items) > 15:
        warnings.append(f"items 共 {len(items)} 筆，超過 15 項，第 16 筆起略去。")

    proof_targets: list[str] = []
    for idx in range(1, 16):
        key = f"D{idx}"
        if idx <= len(items):
            item = items[idx - 1]
            m = matched[idx - 1] if idx - 1 < len(matched) else None
            status = (m.get("status", "missing") if m else "missing")
            attach_status = "已備齊【聲證O】" if status == "have" else "向相關機關申請中"
            period = item.get("period", "")
            period_str = f"（{period}）" if period else ""
            filled[key] = f"{item['category']}{period_str}{attach_status}"
            if "【聲證O】" in filled[key]:
                proof_targets.append(key)
        else:
            filled[key] = None  # 刪除此項

    # 聲證自動編號（【聲證O】→【聲證一】等）
    ordered_proof_targets = [k for k in proof_fields_order if k in proof_targets]
    label_to_proof: dict[str, str] = {
        k: f"【聲證{cn_nums_list[i] if i < len(cn_nums_list) else str(i+1)}】"
        for i, k in enumerate(ordered_proof_targets)
    }

    for key in list(filled.keys()):
        if isinstance(filled[key], str) and "【聲證O】" in filled[key]:
            filled[key] = filled[key].replace("【聲證O】", label_to_proof.get(key, "【聲證O】"))

    # ── 3. 套用欄位到文件 ────────────────────────────────────────────────────
    deleted_fields = [k for k, v in filled.items() if v is None]
    _apply_fields_to_doc(doc, filled, deleted_fields)

    # ── 4. 插入補件 lead-in 段落 ─────────────────────────────────────────────
    main_content = _insert_lead_in(doc, extracted, matched, label_to_proof)
    _apply_proof_table(doc, items, matched, label_to_proof)

    # ── 5. 存檔 ──────────────────────────────────────────────────────────────
    output_path = unicodedata.normalize("NFC", str(output_path))
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    logger.info("docx 已儲存：%s", output_path)

    # 回傳 filled_fields（只取有值的，None 不列）
    filled_fields_display = {k: v for k, v in filled.items() if v is not None}

    return {
        "output_path": output_path,
        "filled_fields": filled_fields_display,
        "deleted_fields": deleted_fields,
        "main_content": main_content,
        "warnings": warnings,
    }
