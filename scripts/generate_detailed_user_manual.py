from __future__ import annotations

import re
import shutil
from datetime import date
from pathlib import Path
from textwrap import wrap

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    Image as PdfImage,
    PageBreak,
    Paragraph,
    Preformatted,
    SimpleDocTemplate,
    Spacer,
    Table as PdfTable,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "USER_GUIDE.md"
GUIDES = ROOT / "docs" / "guides"
ASSETS = GUIDES / "assets"

VERSION = date(2026, 5, 19)
BASE_NAME = f"MAGI_一般使用者超詳細操作手冊_{VERSION.isoformat()}"
MD_OUT = GUIDES / f"{BASE_NAME}.md"
DOCX_OUT = GUIDES / f"{BASE_NAME}.docx"
PDF_OUT = GUIDES / f"{BASE_NAME}.pdf"

# Backward-compatible file names used by README and smoke tests.
VISUAL_DOCX_OUT = GUIDES / f"MAGI_一般使用者圖文操作手冊_{VERSION.isoformat()}.docx"
VISUAL_PDF_OUT = GUIDES / f"MAGI_一般使用者圖文操作手冊_{VERSION.isoformat()}.pdf"

FONT_PATH = Path("/System/Library/Fonts/STHeiti Medium.ttc")
FONT_EAST_ASIA = "Microsoft JhengHei"
FONT_LATIN = "Arial"

INK = "172033"
MUTED = "52627A"
BLUE = "0B74D1"
GREEN = "15803D"
AMBER = "B45309"
RED = "B91C1C"
PURPLE = "6D28D9"
PANEL = "F8FAFC"
BORDER = "D8E2EF"


def pil_font(size: int) -> ImageFont.FreeTypeFont:
    if FONT_PATH.exists():
        return ImageFont.truetype(str(FONT_PATH), size)
    return ImageFont.load_default()


def draw_round_rect(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], fill: str, outline: str, radius: int = 22, width: int = 2) -> None:
    draw.rounded_rectangle(box, radius=radius, fill=f"#{fill}", outline=f"#{outline}", width=width)


def draw_wrapped(draw: ImageDraw.ImageDraw, text: str, xy: tuple[int, int], width_chars: int, size: int, fill: str, gap: int = 8) -> int:
    x, y = xy
    for line in wrap(text, width=width_chars):
        draw.text((x, y), line, font=pil_font(size), fill=f"#{fill}")
        y += size + gap
    return y


def save_img(path: Path, img: Image.Image) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path, "PNG")
    return path


def make_cover_image() -> Path:
    path = ASSETS / "manual_cover_detailed.png"
    img = Image.new("RGB", (1600, 900), "#EEF6FF")
    draw = ImageDraw.Draw(img)
    for y in range(900):
        ratio = y / 899
        r = int(238 * (1 - ratio) + 248 * ratio)
        g = int(246 * (1 - ratio) + 250 * ratio)
        b = int(255 * (1 - ratio) + 252 * ratio)
        draw.line([(0, y), (1600, y)], fill=(r, g, b))

    draw_round_rect(draw, (90, 75, 1510, 825), "FFFFFF", BORDER, radius=34)
    draw.text((140, 135), "MAGI", font=pil_font(72), fill=f"#{BLUE}")
    draw.text((142, 218), "一般使用者超詳細操作手冊", font=pil_font(64), fill=f"#{INK}")
    draw_wrapped(
        draw,
        "從第一次安裝、登入、案件、檔案、PDF、OCR、摘要、翻譯、逐字稿、書狀、法扶、閱卷、筆錄、法律資料、通知、健康檢查到疑難排除的完整指南。",
        (145, 320),
        width_chars=34,
        size=34,
        fill=MUTED,
        gap=10,
    )
    badges = [("外部使用者", BLUE), ("公開版 / 私有版", GREEN), ("繁體中文", PURPLE), ("2026-05-19", AMBER)]
    x = 145
    for label, color in badges:
        draw_round_rect(draw, (x, 610, x + 245, 684), color, color, radius=18)
        bbox = draw.textbbox((0, 0), label, font=pil_font(29))
        draw.text((x + (245 - (bbox[2] - bbox[0])) / 2, 630), label, font=pil_font(29), fill="#FFFFFF")
        x += 270

    draw_round_rect(draw, (1045, 145, 1435, 580), "111827", "334155", radius=26)
    draw_round_rect(draw, (1072, 185, 1408, 545), "F8FAFC", "CBD5E1", radius=18)
    draw.text((1105, 230), "健康狀態", font=pil_font(30), fill=f"#{INK}")
    rows = [("資料庫", "正常"), ("模型", "E4B / 26B"), ("OCR", "正常"), ("NAS", "已掛載")]
    y = 292
    for label, value in rows:
        draw_round_rect(draw, (1100, y, 1380, y + 48), "FFFFFF", BORDER, radius=12)
        draw.text((1118, y + 10), label, font=pil_font(23), fill=f"#{MUTED}")
        draw.text((1282, y + 10), value, font=pil_font(23), fill=f"#{GREEN}")
        y += 62
    return save_img(path, img)


def make_module_map() -> Path:
    path = ASSETS / "manual_module_map_detailed.png"
    img = Image.new("RGB", (1600, 920), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((80, 58), "MAGI 功能地圖", font=pil_font(54), fill=f"#{INK}")
    draw.text((80, 125), "每個功能都有明確資料來源；工具查不到就回報查不到，不由模型硬猜。", font=pil_font(28), fill=f"#{MUTED}")
    center = (535, 205, 1065, 330)
    draw_round_rect(draw, center, "E0F2FE", "7DD3FC", radius=32)
    draw.text((720, 230), "MAGI", font=pil_font(58), fill=f"#{BLUE}")
    draw.text((620, 292), "案件與文件資料核心", font=pil_font(30), fill=f"#{INK}")

    modules = [
        ("案件", "案件卡片、狀態、資料夾、人工結案優先", 80, 410, BLUE),
        ("檔案 / OCR", "預覽、下載、分享、PDF 命名、待辦建立", 555, 410, GREEN),
        ("AI 交付", "摘要、翻譯、逐字稿、@heavy、品質閘門", 1030, 410, PURPLE),
        ("法扶 / 閱卷 / 筆錄", "開辦、進度回報、活動計數、去重", 80, 650, AMBER),
        ("法律資料", "法條、判決、實務見解、通譯實證研究", 555, 650, "2563EB"),
        ("維運", "健康頁、模型切換、NAS、磁碟、通知分流", 1030, 650, RED),
    ]
    for title, desc, x, y, color in modules:
        draw_round_rect(draw, (x, y, x + 430, y + 165), PANEL, color, radius=28)
        draw.text((x + 28, y + 28), title, font=pil_font(36), fill=f"#{color}")
        draw_wrapped(draw, desc, (x + 28, y + 82), width_chars=18, size=25, fill=MUTED, gap=5)

    draw_round_rect(draw, (80, 850, 1520, 895), "F8FAFC", BORDER, radius=16)
    draw.text((108, 860), "使用原則：正式送出、刪除、批次搬移、公開分享、法律引用與金額計算，都要由使用者確認。", font=pil_font(24), fill=f"#{INK}")
    return save_img(path, img)


def make_quality_map() -> Path:
    path = ASSETS / "manual_quality_gate_detailed.png"
    img = Image.new("RGB", (1600, 860), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((80, 58), "高品質交付流程", font=pil_font(54), fill=f"#{INK}")
    steps = [
        ("1", "讀取來源", "檔案、案件、日曆或法律資料"),
        ("2", "工具處理", "OCR、搜尋、摘要、翻譯、分類"),
        ("3", "品質閘門", "檢查漏譯、幻覺、格式、專有名詞"),
        ("4", "不通過重跑", "回到工具或要求使用者補資料"),
        ("5", "可交付成品", "保留來源、路徑、檔案與人工確認點"),
    ]
    colors_ = [BLUE, GREEN, AMBER, RED, PURPLE]
    x = 80
    for idx, (num, title, desc) in enumerate(steps):
        y = 260
        draw_round_rect(draw, (x, y, x + 255, y + 300), PANEL, colors_[idx], radius=28)
        draw.ellipse((x + 28, y + 28, x + 88, y + 88), fill=f"#{colors_[idx]}")
        draw.text((x + 50, y + 40), num, font=pil_font(30), fill="#FFFFFF")
        draw.text((x + 28, y + 120), title, font=pil_font(32), fill=f"#{INK}")
        draw_wrapped(draw, desc, (x + 28, y + 178), width_chars=10, size=24, fill=MUTED, gap=6)
        if idx < len(steps) - 1:
            draw.line((x + 268, y + 150, x + 298, y + 150), fill=f"#{MUTED}", width=4)
            draw.polygon([(x + 298, y + 150), (x + 282, y + 138), (x + 282, y + 162)], fill=f"#{MUTED}")
        x += 300
    draw_round_rect(draw, (80, 650, 1520, 775), "FFF7ED", "FDBA74", radius=22)
    draw_wrapped(
        draw,
        "重點：@heavy 不是只把模型變大，而是要求更嚴格的讀取、對照、專有名詞保留、格式檢查與必要時重新處理。法律文件、翻譯與可交付摘要都應使用這套流程。",
        (112, 678),
        width_chars=45,
        size=28,
        fill=INK,
        gap=8,
    )
    return save_img(path, img)


def make_daily_workflow() -> Path:
    path = ASSETS / "manual_daily_workflow_detailed.png"
    img = Image.new("RGB", (1600, 860), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((80, 58), "日常使用流程", font=pil_font(54), fill=f"#{INK}")
    draw.text((80, 125), "每天先看健康與待辦，再處理案件；正式送出前保留人工確認。", font=pil_font(28), fill=f"#{MUTED}")
    steps = [
        ("1", "健康檢查", "看 /health：DB、模型、OCR、NAS、外網"),
        ("2", "業務概覽", "分別看 OSC 建立待辦與行事曆事件"),
        ("3", "案件處理", "開資料夾、預覽檔案、確認狀態"),
        ("4", "AI 交付", "摘要、翻譯、逐字稿、書狀、分類"),
        ("5", "更新紀錄", "寫回待辦、行事曆、學習修正、通知"),
    ]
    x = 82
    for i, (num, title, desc) in enumerate(steps):
        y = 265
        draw_round_rect(draw, (x, y, x + 245, y + 280), PANEL, [BLUE, GREEN, AMBER, PURPLE, RED][i], radius=28)
        draw.ellipse((x + 28, y + 28, x + 88, y + 88), fill=f"#{[BLUE, GREEN, AMBER, PURPLE, RED][i]}")
        draw.text((x + 50, y + 40), num, font=pil_font(30), fill="#FFFFFF")
        draw.text((x + 28, y + 118), title, font=pil_font(32), fill=f"#{INK}")
        draw_wrapped(draw, desc, (x + 28, y + 174), width_chars=11, size=23, fill=MUTED, gap=6)
        if i < len(steps) - 1:
            draw.line((x + 258, y + 138, x + 292, y + 138), fill=f"#{MUTED}", width=4)
            draw.polygon([(x + 292, y + 138), (x + 276, y + 126), (x + 276, y + 150)], fill=f"#{MUTED}")
        x += 300
    draw_round_rect(draw, (80, 665, 1520, 760), "F8FAFC", BORDER, radius=20)
    draw_wrapped(
        draw,
        "提醒：問行程要查行事曆；問案件要查案件資料庫；問判決要查法律資料。工具調用錯誤時，應回報並修正路由，不應讓模型猜。",
        (112, 690),
        width_chars=50,
        size=26,
        fill=INK,
        gap=7,
    )
    return save_img(path, img)


def make_todo_split() -> Path:
    path = ASSETS / "manual_todo_split_detailed.png"
    img = Image.new("RGB", (1600, 920), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((80, 58), "待辦與行事曆要分開顯示", font=pil_font(54), fill=f"#{INK}")
    draw.text((80, 125), "文件期限、系統建立待辦與真實行程混在一起時，使用者會不知道該先處理什麼。", font=pil_font(28), fill=f"#{MUTED}")
    panels = [
        ("OSC 建立待辦", "來源：PDF、案件文件、系統規則", 80, BLUE, [("05/20", "2025-0084", "補正期限"), ("05/21", "2026-0042", "陳報資料"), ("05/24", "2025-0121", "消債待補")]),
        ("行事曆事件", "來源：Google 日曆、OSC 自動行程", 835, GREEN, [("05/20 10:00", "2025-0032", "調解"), ("05/21 14:30", "2026-0033", "開庭"), ("05/22 16:00", "-", "電話會議")]),
    ]
    for title, subtitle, x, color, rows in panels:
        draw_round_rect(draw, (x, 225, x + 685, 790), PANEL, BORDER, radius=26)
        draw.text((x + 35, 260), title, font=pil_font(38), fill=f"#{color}")
        draw.text((x + 35, 312), subtitle, font=pil_font(24), fill=f"#{MUTED}")
        y = 385
        widths = [150, 170, 275]
        for label, w, xx in zip(["日期 / 時間", "案件", "描述"], widths, [x + 35, x + 200, x + 385]):
            draw_round_rect(draw, (xx, y, xx + w - 12, y + 48), "E0F2FE", "BAE6FD", radius=10)
            draw.text((xx + 14, y + 12), label, font=pil_font(22), fill=f"#{INK}")
        y += 62
        for row in rows:
            for cell, w, xx in zip(row, widths, [x + 35, x + 200, x + 385]):
                draw_round_rect(draw, (xx, y, xx + w - 12, y + 58), "FFFFFF", BORDER, radius=10)
                draw_wrapped(draw, cell, (xx + 13, y + 11), width_chars=max(5, w // 24), size=20, fill=INK, gap=2)
            y += 72
    draw_round_rect(draw, (80, 825, 1520, 875), "FFF7ED", "FDBA74", radius=16)
    draw.text((110, 838), "原則：有具體時間的是行事曆事件；文件期限與案件規則產生的是 OSC 建立待辦。兩者可互相連結，但不混成同一欄。", font=pil_font(24), fill=f"#{INK}")
    return save_img(path, img)


def make_file_flow() -> Path:
    path = ASSETS / "manual_file_flow_detailed.png"
    img = Image.new("RGB", (1600, 880), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((80, 58), "檔案、NAS 與分享連結", font=pil_font(54), fill=f"#{INK}")
    draw.text((80, 125), "網頁版用可讀取路徑預覽與分享；資料庫仍保留 Windows/NAS 慣用路徑，避免破壞使用者原本工作流。", font=pil_font(28), fill=f"#{MUTED}")
    steps = [
        ("DB 路徑", "Z:\\lumi... 或 Y:\\...", BLUE),
        ("路徑轉換", "本機 /Volumes 或 CloudStorage", GREEN),
        ("檔案服務", "預覽、下載、分享連結", PURPLE),
        ("人工確認", "正式外傳前確認權限", AMBER),
    ]
    x = 120
    y = 290
    for idx, (title, desc, color) in enumerate(steps):
        draw_round_rect(draw, (x, y, x + 300, y + 210), PANEL, color, radius=28)
        draw.text((x + 35, y + 42), title, font=pil_font(36), fill=f"#{color}")
        draw_wrapped(draw, desc, (x + 35, y + 105), width_chars=13, size=25, fill=MUTED, gap=6)
        if idx < len(steps) - 1:
            draw.line((x + 315, y + 105, x + 370, y + 105), fill=f"#{MUTED}", width=5)
            draw.polygon([(x + 370, y + 105), (x + 350, y + 90), (x + 350, y + 120)], fill=f"#{MUTED}")
        x += 370
    draw_round_rect(draw, (120, 595, 1480, 760), "F8FAFC", BORDER, radius=22)
    draw_wrapped(
        draw,
        "若看到 Resource deadlock avoided、找不到檔案、外網連不上，先檢查檔案是否正在同步、NAS 是否掛載、分享通道是否健康，再重試預覽與下載；不要把資料庫路徑改成本機臨時路徑。",
        (152, 630),
        width_chars=49,
        size=27,
        fill=INK,
        gap=8,
    )
    return save_img(path, img)


def clean_inline(text: str) -> str:
    text = re.sub(r"!\[([^\]]*)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1（\2）", text)
    text = re.sub(r"(\*\*|__)(.*?)\1", r"\2", text)
    text = re.sub(r"(\*|_)(.*?)\1", r"\2", text)
    text = text.replace("`", "")
    return text.strip()


def parse_markdown_tables(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        raw = lines[i].strip().strip("|")
        cells = [clean_inline(cell.strip()) for cell in raw.split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells):
            rows.append(cells)
        i += 1
    return rows, i


def parse_blocks(markdown: str) -> list[dict]:
    lines = markdown.splitlines()
    blocks: list[dict] = []
    paragraph: list[str] = []
    code: list[str] | None = None

    def flush_para() -> None:
        nonlocal paragraph
        if paragraph:
            blocks.append({"type": "p", "text": clean_inline(" ".join(paragraph))})
            paragraph = []

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if code is not None:
            if stripped.startswith("```"):
                blocks.append({"type": "code", "text": "\n".join(code).rstrip()})
                code = None
            else:
                code.append(line.rstrip())
            i += 1
            continue

        if stripped.startswith("```"):
            flush_para()
            code = []
            i += 1
            continue
        if not stripped:
            flush_para()
            i += 1
            continue
        if stripped.startswith("|"):
            flush_para()
            rows, next_i = parse_markdown_tables(lines, i)
            if rows:
                blocks.append({"type": "table", "rows": rows})
            i = next_i
            continue
        heading = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading:
            flush_para()
            blocks.append({"type": "heading", "level": len(heading.group(1)), "text": clean_inline(heading.group(2))})
            i += 1
            continue
        if re.match(r"^[-*]\s+", stripped):
            flush_para()
            blocks.append({"type": "bullet", "text": clean_inline(re.sub(r"^[-*]\s+", "", stripped))})
            i += 1
            continue
        if re.match(r"^\d+[.)]\s+", stripped):
            flush_para()
            blocks.append({"type": "number", "text": clean_inline(re.sub(r"^\d+[.)]\s+", "", stripped))})
            i += 1
            continue
        if re.fullmatch(r"-{3,}", stripped):
            flush_para()
            blocks.append({"type": "rule"})
            i += 1
            continue
        paragraph.append(stripped)
        i += 1
    flush_para()
    return blocks


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_run_east_asia(run, font_name: str = FONT_EAST_ASIA) -> None:
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = FONT_LATIN
    set_run_east_asia(run)
    run.font.size = Pt(9)
    paragraph.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def configure_doc_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.9)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    normal = doc.styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.18
    normal.paragraph_format.space_after = Pt(5)
    for name, size, color in [
        ("Title", 24, BLUE),
        ("Heading 1", 18, BLUE),
        ("Heading 2", 14, INK),
        ("Heading 3", 12, INK),
    ]:
        style = doc.styles[name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)


def add_doc_footer(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("MAGI 一般使用者超詳細操作手冊")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(MUTED)
        run.font.name = FONT_LATIN
        set_run_east_asia(run)


def add_doc_paragraph(doc: Document, text: str, style: str | None = None) -> None:
    para = doc.add_paragraph(style=style)
    run = para.add_run(text)
    run.font.name = FONT_LATIN
    set_run_east_asia(run)
    if style == "Intense Quote":
        run.font.color.rgb = RGBColor.from_string(INK)


def add_doc_code(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F1F5F9")
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(text or " ")
    run.font.name = "Courier New"
    set_run_east_asia(run, "Microsoft JhengHei")
    run.font.size = Pt(8.5)
    doc.add_paragraph()


def add_doc_table(doc: Document, rows: list[list[str]]) -> None:
    max_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(max_cols):
            cell = table.cell(r_idx, c_idx)
            set_cell_text(cell, row[c_idx] if c_idx < len(row) else "", bold=(r_idx == 0))
            if r_idx == 0:
                set_cell_shading(cell, "E0F2FE")
    doc.add_paragraph()


def build_docx(blocks: list[dict], images: list[Path], out: Path) -> None:
    doc = Document()
    configure_doc_styles(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("MAGI 一般使用者超詳細操作手冊")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor.from_string(BLUE)
    run.font.name = FONT_LATIN
    set_run_east_asia(run)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run("公開版與私有版適用｜一般使用者、事務協作者、導入顧問")
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = RGBColor.from_string(MUTED)
    set_run_east_asia(sub_run)
    doc.add_picture(str(images[0]), width=Inches(6.7))
    doc.add_paragraph("本文件以 docs/USER_GUIDE.md 的 37 章內容為基礎，加入圖示導覽、完整功能索引、常用命令碼與故障排除。", style="Intense Quote")

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    doc.add_heading("圖示導覽", level=1)
    for image in images[1:]:
        doc.add_picture(str(image), width=Inches(6.7))
        doc.add_paragraph()

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    doc.add_heading("完整手冊", level=1)
    for block in blocks:
        kind = block["type"]
        if kind == "heading":
            level = min(max(block["level"], 1), 3)
            doc.add_heading(block["text"], level=level)
        elif kind == "p":
            add_doc_paragraph(doc, block["text"])
        elif kind == "bullet":
            add_doc_paragraph(doc, block["text"], style="List Bullet")
        elif kind == "number":
            add_doc_paragraph(doc, block["text"], style="List Number")
        elif kind == "code":
            add_doc_code(doc, block["text"])
        elif kind == "table":
            add_doc_table(doc, block["rows"])
        elif kind == "rule":
            doc.add_paragraph("—" * 24)

    add_doc_footer(doc)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)


def register_pdf_font() -> str:
    if FONT_PATH.exists():
        try:
            pdfmetrics.registerFont(TTFont("STHeitiMAGI", str(FONT_PATH)))
            return "STHeitiMAGI"
        except Exception:
            pass
    return "Helvetica"


def pdf_styles() -> dict[str, ParagraphStyle]:
    font_name = register_pdf_font()
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle("MAGITitle", parent=base["Title"], fontName=font_name, fontSize=24, leading=31, textColor=colors.HexColor("#" + BLUE), spaceAfter=14),
        "h1": ParagraphStyle("MAGIH1", parent=base["Heading1"], fontName=font_name, fontSize=18, leading=24, textColor=colors.HexColor("#" + BLUE), spaceBefore=12, spaceAfter=8),
        "h2": ParagraphStyle("MAGIH2", parent=base["Heading2"], fontName=font_name, fontSize=14, leading=19, textColor=colors.HexColor("#" + INK), spaceBefore=10, spaceAfter=6),
        "h3": ParagraphStyle("MAGIH3", parent=base["Heading3"], fontName=font_name, fontSize=12, leading=17, textColor=colors.HexColor("#" + INK), spaceBefore=8, spaceAfter=4),
        "body": ParagraphStyle("MAGIBody", parent=base["BodyText"], fontName=font_name, fontSize=9.5, leading=14, textColor=colors.HexColor("#" + INK), spaceAfter=5),
        "bullet": ParagraphStyle("MAGIBullet", parent=base["BodyText"], fontName=font_name, fontSize=9.3, leading=13.5, leftIndent=13, bulletIndent=0, textColor=colors.HexColor("#" + INK), spaceAfter=4),
        "code": ParagraphStyle("MAGICode", parent=base["Code"], fontName=font_name, fontSize=7.7, leading=10.5, backColor=colors.HexColor("#F1F5F9"), textColor=colors.HexColor("#" + INK), leftIndent=5, rightIndent=5, spaceBefore=4, spaceAfter=7),
        "caption": ParagraphStyle("MAGICaption", parent=base["BodyText"], fontName=font_name, fontSize=8.2, leading=11, textColor=colors.HexColor("#" + MUTED), alignment=1, spaceAfter=8),
    }


def pdf_para(text: str, style: ParagraphStyle) -> Paragraph:
    escaped = (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )
    return Paragraph(escaped, style)


def add_pdf_table(story: list, rows: list[list[str]], styles: dict[str, ParagraphStyle], usable_width: float) -> None:
    if not rows:
        return
    max_cols = max(len(row) for row in rows)
    if max_cols > 6:
        for row in rows:
            story.append(pdf_para("｜".join(row), styles["body"]))
        return
    col_width = usable_width / max_cols
    data = []
    for row in rows:
        cells = []
        for i in range(max_cols):
            text = row[i] if i < len(row) else ""
            cells.append(pdf_para(text, styles["body"]))
        data.append(cells)
    table = PdfTable(data, colWidths=[col_width] * max_cols, repeatRows=1, hAlign="LEFT")
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E0F2FE")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#" + INK)),
                ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E1")),
                ("INNERGRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#E2E8F0")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    story.append(table)
    story.append(Spacer(1, 0.15 * cm))


def build_pdf(blocks: list[dict], images: list[Path], out: Path) -> None:
    styles = pdf_styles()
    doc = SimpleDocTemplate(
        str(out),
        pagesize=A4,
        rightMargin=1.7 * cm,
        leftMargin=1.7 * cm,
        topMargin=1.5 * cm,
        bottomMargin=1.5 * cm,
        title="MAGI 一般使用者超詳細操作手冊",
        author="MAGI",
    )
    usable_width = A4[0] - doc.leftMargin - doc.rightMargin
    story: list = [pdf_para("MAGI 一般使用者超詳細操作手冊", styles["title"])]
    story.append(pdf_para("公開版與私有版適用｜一般使用者、事務協作者、導入顧問", styles["body"]))
    story.append(PdfImage(str(images[0]), width=usable_width, height=usable_width * 900 / 1600))
    story.append(PageBreak())
    story.append(pdf_para("圖示導覽", styles["h1"]))
    for idx, image in enumerate(images[1:], 1):
        story.append(PdfImage(str(image), width=usable_width, height=usable_width * Image.open(image).height / Image.open(image).width))
        story.append(Spacer(1, 0.25 * cm))
        if idx < len(images) - 1:
            story.append(PageBreak())
    story.append(PageBreak())
    story.append(pdf_para("完整手冊", styles["h1"]))
    for block in blocks:
        kind = block["type"]
        if kind == "heading":
            level = block["level"]
            story.append(pdf_para(block["text"], styles["h1" if level <= 2 else "h2" if level == 3 else "h3"]))
        elif kind == "p":
            story.append(pdf_para(block["text"], styles["body"]))
        elif kind == "bullet":
            story.append(Paragraph(f"• {block['text']}", styles["bullet"]))
        elif kind == "number":
            story.append(Paragraph(f"• {block['text']}", styles["bullet"]))
        elif kind == "code":
            story.append(Preformatted(block["text"] or " ", styles["code"], maxLineLength=88))
        elif kind == "table":
            add_pdf_table(story, block["rows"], styles, usable_width)
        elif kind == "rule":
            story.append(Spacer(1, 0.25 * cm))
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.build(story)


def main() -> None:
    markdown = SOURCE.read_text(encoding="utf-8")
    generated_note = (
        f"> 本檔由 `{SOURCE.relative_to(ROOT)}` 於 {VERSION.isoformat()} 產生，"
        "供外部使用者閱讀與匯出 DOCX/PDF。正式內容請以本檔與 `docs/USER_GUIDE.md` 同步維護。\n\n"
    )
    MD_OUT.write_text(generated_note + markdown, encoding="utf-8")
    blocks = parse_blocks(markdown)
    images = [
        make_cover_image(),
        make_module_map(),
        make_quality_map(),
        make_daily_workflow(),
        make_todo_split(),
        make_file_flow(),
    ]
    build_docx(blocks, images, DOCX_OUT)
    build_pdf(blocks, images, PDF_OUT)
    shutil.copyfile(DOCX_OUT, VISUAL_DOCX_OUT)
    shutil.copyfile(PDF_OUT, VISUAL_PDF_OUT)
    print(f"wrote {MD_OUT}")
    print(f"wrote {DOCX_OUT}")
    print(f"wrote {PDF_OUT}")
    print(f"updated {VISUAL_DOCX_OUT}")
    print(f"updated {VISUAL_PDF_OUT}")


if __name__ == "__main__":
    main()
