#!/usr/bin/env python3
"""Render, but never install or start, a release-bound V3 launchd deployment."""

from __future__ import annotations

import argparse
import hashlib
import json
import os
import plistlib
import re
import resource
import signal
import sqlite3
import stat
import subprocess
import sys
import tempfile
import uuid
from dataclasses import dataclass, replace
from datetime import datetime, timezone
from pathlib import Path, PurePosixPath
from typing import Any, Mapping

if __package__ in {None, ""}:
    sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from scripts.v3_cron_snapshot import CronSnapshotBlocked, render_snapshot
from scripts.v3_python_runtime_snapshot import PythonRuntimeBlocked, build_runtime_manifest
from scripts.v3_static_external_staging import (
    RECEIPT_NAME as STATIC_EXTERNAL_RECEIPT_NAME,
    RELEASE_BINDING_RECEIPT_NAME as STATIC_EXTERNAL_RELEASE_RECEIPT_NAME,
    StaticExternalStagingError,
    render_static_external_release_binding,
    verify_static_external_payload,
)
from magi_v3.errors import ConfigurationError
from magi_v3.external_inputs import (
    NAMED_MUTABLE_STATE_BINDINGS,
    named_mutable_state_paths,
)
from magi_v3.service_manifest import load_service_manifest

EXPECTED_LABELS = {
    "gateway": "com.magi.v3.gateway",
    "control": "com.magi.v3.control",
    "supervisor": "com.magi.v3.supervisor",
}
# The public HTTP gateway performs bounded, user-initiated work.  Rendering it
# as a launchd Background process makes its spawned workers inherit background
# CPU policy and can exhaust their wall-clock budget despite staying within
# the CPU/RSS contract.  Control and supervisor remain background services.
PROCESS_TYPE_BY_ROLE = {
    "gateway": "Interactive",
    "control": "Background",
    "supervisor": "Background",
}
RELEASE_ID_RE = re.compile(r"^[A-Za-z0-9][A-Za-z0-9._-]{0,127}$")
SHA256_RE = re.compile(r"^[0-9a-f]{64}$")
MODULE_RE = re.compile(r"^[A-Za-z_][A-Za-z0-9_]*(?:\.[A-Za-z_][A-Za-z0-9_]*)*$")
RELEASE_MANIFEST_NAME = "release-manifest.json"
RELEASE_MARKER_NAME = "RELEASE_COMPLETE.json"
OWNERSHIP_MANIFEST_NAME = "ownership/ownership-manifest.json"
DEPLOY_MANIFEST_NAME = "deploy-manifest.json"
COMPLETION_MARKER_NAME = "DEPLOY_PREPARED.json"
DEPLOYMENT_MODES = frozenset({"production", "isolated_live_validation"})
SERVICE_MANIFEST_NAMES = {
    "production": "config/v3_service_manifest.json",
    "isolated_live_validation": "config/v3_live_validation_service_manifest.json",
}
CRON_DISPATCH_POLICY_NAME = "config/v3_schedule_dispatch_policy.json"
VALIDATION_ENV_BYTES = b"MAGI_V3_VALIDATION_FIXTURE=1\n"
VALIDATION_CRON_JOB = {
    "id": "v3_live_validation_inert",
    "enabled": False,
    "cron": "0 0 31 2 *",
    "command": "@MAGI validation_inert",
}
VALIDATION_LAF_CONFIG = {
    "gmail": {},
    "laf": {
        "auto_create_case": False,
        "auto_portal_draft": False,
        "base_url": "http://127.0.0.1:17002",
        "check_interval": 3600,
        "headless": True,
        "portal_env": "test",
        "test_base_url": "http://127.0.0.1:17002",
    },
    "product_runtime": {
        "laf": {
            "compare_base_url": "",
            "portal_env": "test",
            "prod_base_url": "http://127.0.0.1:17002",
            "test_base_url": "http://127.0.0.1:17002",
        }
    },
}
VALIDATION_LAF_CONFIG_BYTES = (
    json.dumps(VALIDATION_LAF_CONFIG, ensure_ascii=False, sort_keys=True, indent=2) + "\n"
).encode("utf-8")
VALIDATION_GOOGLE_CREDENTIALS_BYTES = b'{"installed":{"client_id":"inert.invalid"}}\n'
VALIDATION_GOOGLE_CALENDAR_TOKEN_BYTES = b'{"token":"inert","valid":false}\n'
VALIDATION_LAF_GMAIL_TOKEN_BYTES = b"MAGI_V3_INERT_LAF_GMAIL_TOKEN\n"
REQUIRED_DOCX_MODULE = "docx"
REQUIRED_DOCX_MINIMUM_VERSION = "1.0"
RUNTIME_MODULE_PROBE_TIMEOUT_SECONDS = 10
RUNTIME_MODULE_PROBE_MAX_OUTPUT_BYTES = 4096
_RUNTIME_MODULE_PROBE = r"""
import contextlib
import base64
import hashlib
import importlib.metadata
import json
import os
from io import BytesIO
from pathlib import Path, PurePosixPath

try:
    with open(os.devnull, "w", encoding="utf-8") as sink:
        with contextlib.redirect_stdout(sink), contextlib.redirect_stderr(sink):
            import docx
            from docx import Document
            from packaging.utils import canonicalize_name
            from packaging.version import Version

            origin = Path(str(getattr(docx, "__file__", "") or "")).resolve(
                strict=True
            )
            distributions = [
                candidate
                for candidate in importlib.metadata.distributions()
                if canonicalize_name(str(candidate.metadata.get("Name", "")))
                == "python-docx"
            ]
            if len(distributions) != 1:
                raise RuntimeError("python-docx distribution is ambiguous")
            distribution_metadata = distributions[0]
            distribution = canonicalize_name(
                str(distribution_metadata.metadata["Name"])
            )
            distribution_version = Version(distribution_metadata.version)
            module_version = Version(
                str(getattr(docx, "__version__", "") or "").strip()
            )
            if module_version != distribution_version:
                raise RuntimeError("python-docx module/distribution version mismatch")
            version = str(distribution_version)
            version_at_least_minimum = distribution_version >= Version("1.0")
            distribution_files = tuple(distribution_metadata.files or ())
            metadata_entries = [
                entry
                for entry in distribution_files
                if PurePosixPath(str(entry)).name == "METADATA"
                and PurePosixPath(str(entry)).parent.name.endswith(".dist-info")
            ]
            if len(metadata_entries) != 1:
                raise RuntimeError("python-docx METADATA is ambiguous or missing")
            metadata_entry = metadata_entries[0]
            metadata_entry_path = PurePosixPath(str(metadata_entry))
            if metadata_entry_path.is_absolute() or ".." in metadata_entry_path.parts:
                raise RuntimeError("python-docx METADATA path is invalid")
            record_entries = [
                entry
                for entry in distribution_files
                if PurePosixPath(str(entry)).name == "RECORD"
                and PurePosixPath(str(entry)).parent == metadata_entry_path.parent
            ]
            if len(record_entries) != 1:
                raise RuntimeError("python-docx RECORD is ambiguous or missing")
            record_entry = record_entries[0]
            record_entry_path = PurePosixPath(str(record_entry))
            if record_entry_path.is_absolute() or ".." in record_entry_path.parts:
                raise RuntimeError("python-docx RECORD path is invalid")
            metadata_origin = Path(
                distribution_metadata.locate_file(metadata_entry)
            ).resolve(strict=True)
            record_origin = Path(
                distribution_metadata.locate_file(record_entry)
            ).resolve(strict=True)
            owned_module_entries = []
            for entry in distribution_files:
                try:
                    located = Path(distribution_metadata.locate_file(entry)).resolve(
                        strict=True
                    )
                except OSError:
                    continue
                if located == origin:
                    owned_module_entries.append(entry)
            if len(owned_module_entries) != 1:
                raise RuntimeError("python-docx package ownership is ambiguous or missing")
            owned_module_entry = PurePosixPath(str(owned_module_entries[0]))
            if (
                owned_module_entry.is_absolute()
                or ".." in owned_module_entry.parts
                or owned_module_entry.as_posix() != "docx/__init__.py"
            ):
                raise RuntimeError("python-docx package ownership path is invalid")
            owned_module_record = owned_module_entries[0]
            module_digest = hashlib.sha256(origin.read_bytes()).hexdigest()
            expected_record_digest = base64.urlsafe_b64encode(
                bytes.fromhex(module_digest)
            ).decode("ascii").rstrip("=")
            record_hash = getattr(owned_module_record, "hash", None)
            record_size = getattr(owned_module_record, "size", None)
            if (
                record_hash is None
                or getattr(record_hash, "mode", None) != "sha256"
                or getattr(record_hash, "value", None) != expected_record_digest
                or isinstance(record_size, bool)
                or not isinstance(record_size, int)
                or record_size != origin.stat().st_size
            ):
                raise RuntimeError("python-docx RECORD module binding is invalid")
            document = Document()
            document.add_paragraph("MAGI V3 python-docx preflight")
            buffer = BytesIO()
            document.save(buffer)
            reopened = Document(BytesIO(buffer.getvalue()))
            roundtrip_succeeded = (
                len(reopened.paragraphs) == 1
                and reopened.paragraphs[0].text == "MAGI V3 python-docx preflight"
            )
    version = str(version).strip()
    digest = hashlib.sha256(origin.read_bytes()).hexdigest()
    metadata_digest = hashlib.sha256(metadata_origin.read_bytes()).hexdigest()
    record_digest = hashlib.sha256(record_origin.read_bytes()).hexdigest()
    payload = {
        "ok": True,
        "module": "docx",
        "distribution": distribution,
        "distribution_count": len(distributions),
        "distribution_metadata_origin": str(metadata_origin),
        "distribution_metadata_sha256": metadata_digest,
        "distribution_record_origin": str(record_origin),
        "distribution_record_sha256": record_digest,
        "distribution_module_entry": owned_module_entry.as_posix(),
        "distribution_module_origin": str(origin),
        "distribution_module_record_sha256": module_digest,
        "distribution_module_record_size": record_size,
        "distribution_version_matches_module": True,
        "version": version,
        "version_at_least_minimum": version_at_least_minimum,
        "module_sha256": digest,
        "origin": str(origin),
        "roundtrip_succeeded": roundtrip_succeeded,
    }
except BaseException as exc:
    payload = {"ok": False, "error_type": type(exc).__name__}
print(json.dumps(payload, sort_keys=True, separators=(",", ":")))
"""


class DeployPrepareBlocked(ValueError):
    """The requested deployment cannot truthfully be rendered as ready."""


@dataclass(frozen=True, slots=True)
class ReleaseIdentity:
    release_id: str
    manifest_sha256: str
    manifest_path: Path
    file_sha256: Mapping[str, str]
    file_mode: Mapping[str, int]


@dataclass(frozen=True, slots=True)
class RoleDefinition:
    role: str
    label: str
    entrypoint_module: str
    ports: tuple[int, ...]
    ownership_domains: tuple[str, ...]


@dataclass(frozen=True, slots=True)
class ExternalRuntimeInputs:
    env_file: Path
    env_file_sha256: str
    cron_jobs_file: Path
    cron_jobs_sha256: str
    website_root: Path
    website_admin_sha256: str
    python_runtime: Path
    python_runtime_realpath: Path
    python_runtime_sha256: str
    chromedriver_path: Path | None
    chromedriver_sha256: str | None
    chromedriver_mode: str | None
    cron_jobs_source_file: Path
    cron_jobs_source_sha256: str
    cron_snapshot_evidence: Mapping[str, Any]
    laf_config_file: Path
    laf_config_sha256: str
    laf_config_mode: str
    google_credentials_file: Path
    google_credentials_sha256: str
    google_credentials_mode: str
    google_calendar_token_source_file: Path
    google_calendar_token_source_sha256: str
    laf_gmail_token_source_file: Path
    laf_gmail_token_source_sha256: str
    file_review_token_source_file: Path
    file_review_token_source_sha256: str
    gmail_compose_token_source_file: Path | None
    gmail_compose_token_source_sha256: str | None
    accounting_credentials_file: Path
    accounting_credentials_sha256: str
    accounting_credentials_mode: str
    accounting_sheets_token_source_file: Path
    accounting_sheets_token_source_sha256: str
    drive_sync_token_source_file: Path
    drive_sync_token_source_sha256: str
    drive_sync_write_token_source_file: Path
    drive_sync_write_token_source_sha256: str
    nas_ocr_queue_db_file: Path
    nas_ocr_queue_db_mode: str


@dataclass(frozen=True, slots=True)
class StaticExternalEvidence:
    receipt: Path
    receipt_sha256: str
    receipt_bytes: bytes
    payload_receipt: Path
    payload_receipt_sha256: str
    source_snapshot_sha256: str
    target_snapshot_sha256: str


def _sha256(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def _sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _sandbox_literal(path: Path) -> str:
    return json.dumps(str(path))


def _probe_process_exec_paths(python_runtime: Path) -> tuple[Path, ...]:
    declared = Path(os.path.abspath(python_runtime.expanduser()))
    realpath = declared.resolve(strict=True)
    allowed = {declared, realpath}
    framework_payload = (
        realpath.parent.parent
        / "Resources"
        / "Python.app"
        / "Contents"
        / "MacOS"
        / "Python"
    )
    if framework_payload.is_file():
        allowed.add(framework_payload.resolve(strict=True))
    return tuple(sorted(allowed, key=str))


def _probe_seatbelt_profile(
    python_runtime: Path,
    *,
    stdout_path: Path,
    stderr_path: Path,
) -> str:
    exec_filters = " ".join(
        f"(literal {_sandbox_literal(path)})"
        for path in _probe_process_exec_paths(python_runtime)
    )
    return "".join(
        (
            "(version 1)",
            "(allow default)",
            "(deny network*)",
            "(deny process-fork)",
            f"(deny process-exec (require-not (require-any {exec_filters})))",
            "(deny file-write*)",
            '(allow file-write* (literal "/dev/null"))',
            f"(allow file-write* (literal {_sandbox_literal(stdout_path)}))",
            f"(allow file-write* (literal {_sandbox_literal(stderr_path)}))",
        )
    )


def _limit_probe_output_files() -> None:
    resource.setrlimit(
        resource.RLIMIT_FSIZE,
        (
            RUNTIME_MODULE_PROBE_MAX_OUTPUT_BYTES,
            RUNTIME_MODULE_PROBE_MAX_OUTPUT_BYTES,
        ),
    )
    resource.setrlimit(resource.RLIMIT_CORE, (0, 0))


def _terminate_probe_process_group(process: subprocess.Popen[bytes]) -> None:
    try:
        os.killpg(process.pid, signal.SIGTERM)
    except ProcessLookupError:
        pass
    try:
        process.wait(timeout=0.25)
    except subprocess.TimeoutExpired:
        try:
            os.killpg(process.pid, signal.SIGKILL)
        except ProcessLookupError:
            pass
        process.wait(timeout=1)


def _run_bounded_runtime_probe(python_runtime: Path) -> bytes:
    """Run the transient probe under Seatbelt with bounded, ephemeral output."""

    sandbox_exec = Path("/usr/bin/sandbox-exec")
    if (
        sys.platform != "darwin"
        or sandbox_exec.is_symlink()
        or not sandbox_exec.is_file()
        or not os.access(sandbox_exec, os.X_OK)
    ):
        raise DeployPrepareBlocked(
            "external Python runtime python-docx probe requires macOS Seatbelt"
        )
    try:
        with tempfile.TemporaryDirectory(prefix="magi-v3-docx-preflight-") as temporary:
            temporary_root = Path(temporary).resolve(strict=True)
            stdout_path = temporary_root / "stdout"
            stderr_path = temporary_root / "stderr"
            stdout_descriptor = os.open(
                stdout_path,
                os.O_RDWR | os.O_CREAT | os.O_EXCL,
                0o600,
            )
            stderr_descriptor = os.open(
                stderr_path,
                os.O_RDWR | os.O_CREAT | os.O_EXCL,
                0o600,
            )
            with os.fdopen(stdout_descriptor, "w+b") as stdout_handle, os.fdopen(
                stderr_descriptor, "w+b"
            ) as stderr_handle:
                profile = _probe_seatbelt_profile(
                    python_runtime,
                    stdout_path=stdout_path,
                    stderr_path=stderr_path,
                )
                command = [
                    str(sandbox_exec),
                    "-p",
                    profile,
                    str(python_runtime),
                    "-B",
                    "-X",
                    "pycache_prefix=/dev/null",
                    "-I",
                    "-c",
                    _RUNTIME_MODULE_PROBE,
                ]
                process = subprocess.Popen(
                    command,
                    stdin=subprocess.DEVNULL,
                    stdout=stdout_handle,
                    stderr=stderr_handle,
                    cwd="/",
                    env={
                        "LANG": "C",
                        "LC_ALL": "C",
                        "PATH": "/usr/bin:/bin",
                        "PYTHONDONTWRITEBYTECODE": "1",
                        "PYTHONPYCACHEPREFIX": "/dev/null",
                        "PYTHONNOUSERSITE": "1",
                        "TMPDIR": str(temporary_root),
                    },
                    start_new_session=True,
                    preexec_fn=_limit_probe_output_files,
                )
                try:
                    returncode = process.wait(
                        timeout=RUNTIME_MODULE_PROBE_TIMEOUT_SECONDS
                    )
                except subprocess.TimeoutExpired as exc:
                    _terminate_probe_process_group(process)
                    raise DeployPrepareBlocked(
                        "external Python runtime python-docx import probe could not "
                        "complete before timeout"
                    ) from exc
                stdout_size = os.fstat(stdout_handle.fileno()).st_size
                stderr_size = os.fstat(stderr_handle.fileno()).st_size
                if (
                    stdout_size >= RUNTIME_MODULE_PROBE_MAX_OUTPUT_BYTES
                    or stderr_size >= RUNTIME_MODULE_PROBE_MAX_OUTPUT_BYTES
                ):
                    raise DeployPrepareBlocked(
                        "external Python runtime python-docx import probe exceeded its output limit"
                    )
                if returncode != 0:
                    raise DeployPrepareBlocked(
                        "external Python runtime python-docx import probe failed"
                    )
                stdout_handle.seek(0)
                return stdout_handle.read(RUNTIME_MODULE_PROBE_MAX_OUTPUT_BYTES)
    except DeployPrepareBlocked:
        raise
    except (OSError, UnicodeError, ValueError, subprocess.SubprocessError) as exc:
        raise DeployPrepareBlocked(
            "external Python runtime python-docx import probe could not complete"
        ) from exc


def _probe_python_docx_runtime(
    python_runtime: Path,
    runtime_manifest: Mapping[str, Any],
) -> dict[str, Any]:
    """Prove the selected, snapshotted runtime can import the core DOCX module."""

    try:
        payload = json.loads(_run_bounded_runtime_probe(python_runtime).decode("utf-8"))
    except (UnicodeError, json.JSONDecodeError) as exc:
        raise DeployPrepareBlocked(
            "external Python runtime python-docx import probe returned invalid evidence"
        ) from exc
    if not isinstance(payload, dict) or payload.get("ok") is not True:
        raise DeployPrepareBlocked(
            "external Python runtime is missing the required python-docx module"
        )
    version = payload.get("version")
    module_sha256 = payload.get("module_sha256")
    origin_text = payload.get("origin")
    metadata_origin_text = payload.get("distribution_metadata_origin")
    metadata_sha256 = payload.get("distribution_metadata_sha256")
    record_origin_text = payload.get("distribution_record_origin")
    record_sha256 = payload.get("distribution_record_sha256")
    distribution_module_entry_text = payload.get("distribution_module_entry")
    distribution_module_origin_text = payload.get("distribution_module_origin")
    distribution_module_record_sha256 = payload.get(
        "distribution_module_record_sha256"
    )
    distribution_module_record_size = payload.get("distribution_module_record_size")
    if (
        payload.get("module") != REQUIRED_DOCX_MODULE
        or payload.get("distribution") != "python-docx"
        or payload.get("distribution_count") != 1
        or not isinstance(version, str)
        or len(version) > 64
        or not isinstance(module_sha256, str)
        or SHA256_RE.fullmatch(module_sha256) is None
        or not isinstance(origin_text, str)
        or len(origin_text) > 2048
        or not isinstance(metadata_origin_text, str)
        or len(metadata_origin_text) > 2048
        or not isinstance(metadata_sha256, str)
        or SHA256_RE.fullmatch(metadata_sha256) is None
        or not isinstance(record_origin_text, str)
        or len(record_origin_text) > 2048
        or not isinstance(record_sha256, str)
        or SHA256_RE.fullmatch(record_sha256) is None
        or not isinstance(distribution_module_entry_text, str)
        or len(distribution_module_entry_text) > 256
        or not isinstance(distribution_module_origin_text, str)
        or len(distribution_module_origin_text) > 2048
        or distribution_module_record_sha256 != module_sha256
        or isinstance(distribution_module_record_size, bool)
        or not isinstance(distribution_module_record_size, int)
        or distribution_module_record_size < 0
        or payload.get("distribution_version_matches_module") is not True
        or payload.get("roundtrip_succeeded") is not True
    ):
        raise DeployPrepareBlocked(
            "external Python runtime python-docx import evidence is invalid"
        )
    if payload.get("version_at_least_minimum") is not True:
        raise DeployPrepareBlocked(
            "external Python runtime python-docx version is below 1.0"
        )
    declared = Path(os.path.abspath(python_runtime.expanduser()))
    try:
        realpath = declared.resolve(strict=True)
    except OSError as exc:
        raise DeployPrepareBlocked(
            "external Python runtime python-docx executable binding is invalid"
        ) from exc
    if (
        runtime_manifest.get("python_runtime") != str(declared)
        or runtime_manifest.get("python_runtime_realpath") != str(realpath)
    ):
        raise DeployPrepareBlocked(
            "external Python runtime python-docx executable is not bound by its snapshot"
        )
    def canonical_probe_file(raw: str, *, description: str) -> Path:
        try:
            path = Path(raw)
            if (
                not path.is_absolute()
                or path.is_symlink()
                or not path.is_file()
                or path.resolve(strict=True) != path
            ):
                raise DeployPrepareBlocked(
                    f"external Python runtime python-docx {description} is not canonical"
                )
            return path
        except OSError as exc:
            raise DeployPrepareBlocked(
                f"external Python runtime python-docx {description} is not canonical"
            ) from exc

    origin = canonical_probe_file(origin_text, description="origin")
    metadata_origin = canonical_probe_file(
        metadata_origin_text,
        description="METADATA origin",
    )
    record_origin = canonical_probe_file(
        record_origin_text,
        description="RECORD origin",
    )
    distribution_module_origin = canonical_probe_file(
        distribution_module_origin_text,
        description="distribution-owned module origin",
    )
    distribution_module_entry = PurePosixPath(distribution_module_entry_text)
    if (
        distribution_module_origin != origin
        or distribution_module_entry.is_absolute()
        or ".." in distribution_module_entry.parts
        or distribution_module_entry.as_posix() != "docx/__init__.py"
        or not metadata_origin.parent.name.endswith(".dist-info")
        or metadata_origin.name != "METADATA"
        or record_origin.parent != metadata_origin.parent
        or record_origin.name != "RECORD"
        or distribution_module_record_size != origin.stat().st_size
    ):
        raise DeployPrepareBlocked(
            "external Python runtime python-docx distribution ownership is invalid"
        )
    try:
        located_module_origin = (
            metadata_origin.parent.parent
            / Path(*distribution_module_entry.parts)
        ).resolve(strict=True)
    except OSError as exc:
        raise DeployPrepareBlocked(
            "external Python runtime python-docx distribution ownership is invalid"
        ) from exc
    if located_module_origin != origin:
        raise DeployPrepareBlocked(
            "external Python runtime python-docx distribution is orphaned from its module"
        )

    manifest_roots: list[tuple[str, Path, list[Any]]] = []
    for root_key, rows_key in (
        ("runtime_root", "files"),
        ("base_runtime_root", "base_files"),
    ):
        root_text = runtime_manifest.get(root_key)
        rows = runtime_manifest.get(rows_key)
        if not isinstance(root_text, str) or not isinstance(rows, list):
            raise DeployPrepareBlocked(
                "external Python runtime snapshot origin inventory is invalid"
            )
        root = Path(root_text)
        if (
            not root.is_absolute()
            or root.is_symlink()
            or not root.is_dir()
            or root.resolve(strict=True) != root
        ):
            raise DeployPrepareBlocked(
                "external Python runtime snapshot origin root is not canonical"
            )
        manifest_roots.append((root_key, root, rows))

    def exact_manifest_binding(path: Path, digest: str) -> tuple[str, str] | None:
        for root_key, root, rows in manifest_roots:
            try:
                relative = path.relative_to(root).as_posix()
            except ValueError:
                continue
            matches = [
                row
                for row in rows
                if isinstance(row, dict) and row.get("path") == relative
            ]
            if len(matches) != 1:
                return None
            row = matches[0]
            if row.get("kind") != "file" or row.get("sha256") != digest:
                return None
            return root_key, relative
        return None

    module_binding = exact_manifest_binding(origin, module_sha256)
    metadata_binding = exact_manifest_binding(metadata_origin, metadata_sha256)
    record_binding = exact_manifest_binding(record_origin, record_sha256)
    if module_binding is None:
        raise DeployPrepareBlocked(
            "external Python runtime python-docx origin is outside or not hash-bound by its snapshot"
        )
    if metadata_binding is None:
        raise DeployPrepareBlocked(
            "external Python runtime python-docx METADATA is outside or not hash-bound by its snapshot"
        )
    if record_binding is None:
        raise DeployPrepareBlocked(
            "external Python runtime python-docx RECORD is outside or not hash-bound by its snapshot"
        )
    if not (
        module_binding[0] == metadata_binding[0] == record_binding[0]
    ):
        raise DeployPrepareBlocked(
            "external Python runtime python-docx module and METADATA roots differ"
        )
    tree_sha256 = runtime_manifest.get("tree_sha256")
    if not isinstance(tree_sha256, str) or SHA256_RE.fullmatch(tree_sha256) is None:
        raise DeployPrepareBlocked(
            "external Python runtime snapshot tree binding is invalid"
        )
    evidence: dict[str, Any] = {
        "module": REQUIRED_DOCX_MODULE,
        "distribution": "python-docx",
        "minimum_version": REQUIRED_DOCX_MINIMUM_VERSION,
        "version": version,
        "module_sha256": module_sha256,
        "distribution_metadata_sha256": metadata_sha256,
        "distribution_record_sha256": record_sha256,
        "runtime_tree_sha256": tree_sha256,
        "distribution_unambiguous": True,
        "distribution_metadata_manifest_bound": True,
        "distribution_record_manifest_bound": True,
        "distribution_module_owned": True,
        "distribution_module_record_bound": True,
        "distribution_version_matches_module": True,
        "import_succeeded": True,
        "module_manifest_bound": True,
        "roundtrip_succeeded": True,
    }
    evidence["evidence_sha256"] = _sha256(
        json.dumps(
            evidence,
            sort_keys=True,
            separators=(",", ":"),
        ).encode("utf-8")
    )
    return evidence


def _load_object(path: Path, *, description: str) -> dict[str, Any]:
    try:
        value = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        raise DeployPrepareBlocked(f"{description} is unreadable: {exc}") from exc
    if not isinstance(value, dict):
        raise DeployPrepareBlocked(f"{description} must be a JSON object")
    return value


def _inside(path: Path, parent: Path) -> bool:
    try:
        path.relative_to(parent)
    except ValueError:
        return False
    return True


def _application_support_root() -> Path:
    return (Path.home() / "Library" / "Application Support").resolve()


def _canonical_runtime_root() -> Path:
    return _application_support_root() / "MAGI" / "runtime" / "MAGI_v3"


def _canonical_installed_release_root(release_id: str) -> Path:
    return _application_support_root() / "MAGI" / "releases" / release_id


def _canonical_nas_ocr_queue_db() -> Path:
    return (Path.home() / ".magi_nas_ocr_queue.db").resolve(strict=False)


def _validated_absolute(path: Path, *, name: str) -> Path:
    expanded = path.expanduser()
    if not expanded.is_absolute():
        raise DeployPrepareBlocked(f"{name} must be absolute")
    return expanded.resolve(strict=False)


def _validate_external_runtime_inputs(
    env_file: Path | None,
    cron_jobs_file: Path | None,
    website_root: Path | None,
    python_runtime: Path | None,
    laf_config_file: Path | None,
    google_credentials_file: Path | None,
    google_calendar_token_file: Path | None,
    laf_gmail_token_file: Path | None,
    file_review_token_file: Path | None,
    gmail_compose_token_file: Path | None,
    accounting_credentials_file: Path | None,
    accounting_sheets_token_file: Path | None,
    drive_sync_token_file: Path | None,
    drive_sync_write_token_file: Path | None,
    nas_ocr_queue_db_path: Path | None,
    chromedriver_path: Path | None,
) -> ExternalRuntimeInputs:
    raw_env = env_file or Path(os.environ.get("MAGI_ENV_FILE", ""))
    raw_cron = cron_jobs_file or Path(os.environ.get("MAGI_CRON_JOBS_FILE", ""))
    raw_website = website_root or Path(os.environ.get("MAGI_WEBSITE_ROOT", ""))
    raw_python = python_runtime or Path(os.environ.get("MAGI_V3_PYTHON_RUNTIME", ""))
    raw_laf_config = laf_config_file or Path(os.environ.get("MAGI_LAF_CONFIG_FILE", ""))
    raw_google_credentials = google_credentials_file or Path(
        os.environ.get("MAGI_GOOGLE_CREDENTIALS_PATH", "")
    )
    raw_google_calendar_token = google_calendar_token_file or Path(
        os.environ.get("MAGI_GOOGLE_CALENDAR_TOKEN_PATH", "")
    )
    raw_laf_gmail_token = laf_gmail_token_file or Path(
        os.environ.get("MAGI_LAF_GMAIL_TOKEN_PATH", "")
    )
    raw_file_review_token = file_review_token_file or Path(
        os.environ.get("MAGI_FILE_REVIEW_TOKEN_PATH", "")
    )
    compose_env = os.environ.get("MAGI_GMAIL_COMPOSE_TOKEN_PATH", "").strip()
    raw_gmail_compose_token = gmail_compose_token_file or (
        Path(compose_env) if compose_env else None
    )
    raw_accounting_credentials = accounting_credentials_file or Path(
        os.environ.get("MAGI_ACCOUNTING_GOOGLE_CREDENTIALS_PATH", "")
    )
    raw_accounting_token = accounting_sheets_token_file or Path(
        os.environ.get("MAGI_ACCOUNTING_GOOGLE_SHEETS_TOKEN", "")
    )
    raw_drive_sync_token = drive_sync_token_file or Path(
        os.environ.get("MAGI_DRIVE_SYNC_TOKEN", "")
        or str(raw_accounting_token)
    )
    raw_drive_sync_write_token = drive_sync_write_token_file or Path(
        os.environ.get("MAGI_DRIVE_SYNC_WRITE_TOKEN", "")
    )
    raw_ocr_queue = nas_ocr_queue_db_path or Path(
        os.environ.get("MAGI_NAS_OCR_QUEUE_DB_PATH", "")
    )
    chromedriver_env = os.environ.get("MAGI_CHROMEDRIVER_PATH", "").strip()
    raw_chromedriver = chromedriver_path or (
        Path(chromedriver_env) if chromedriver_env else None
    )
    if not str(raw_env) or not raw_env.expanduser().is_absolute() or raw_env.expanduser().is_symlink():
        raise DeployPrepareBlocked("external MAGI env file must be an absolute non-symlink path")
    if not str(raw_website) or not raw_website.expanduser().is_absolute() or raw_website.expanduser().is_symlink():
        raise DeployPrepareBlocked("external website root must be an absolute non-symlink path")
    if not str(raw_cron) or not raw_cron.expanduser().is_absolute() or raw_cron.expanduser().is_symlink():
        raise DeployPrepareBlocked("external cron jobs file must be an absolute non-symlink path")
    if not str(raw_python) or not raw_python.expanduser().is_absolute():
        raise DeployPrepareBlocked("external Python runtime must be an absolute path")
    if (
        not str(raw_laf_config)
        or not raw_laf_config.expanduser().is_absolute()
        or raw_laf_config.expanduser().is_symlink()
    ):
        raise DeployPrepareBlocked("external LAF config must be an absolute non-symlink path")
    for name, raw in (
        ("Google credentials", raw_google_credentials),
        ("Google Calendar token", raw_google_calendar_token),
        ("LAF Gmail token", raw_laf_gmail_token),
        ("FileReview token", raw_file_review_token),
        ("accounting credentials", raw_accounting_credentials),
        ("accounting Sheets token", raw_accounting_token),
        ("Drive sync token", raw_drive_sync_token),
        ("Drive sync write token", raw_drive_sync_write_token),
        ("NAS OCR queue database", raw_ocr_queue),
    ):
        if not str(raw) or not raw.expanduser().is_absolute() or raw.expanduser().is_symlink():
            raise DeployPrepareBlocked(f"external {name} must be an absolute non-symlink path")
    if raw_chromedriver is not None and (
        not raw_chromedriver.expanduser().is_absolute()
        or raw_chromedriver.expanduser().is_symlink()
    ):
        raise DeployPrepareBlocked(
            "external ChromeDriver must be an absolute non-symlink path"
        )
    try:
        secret = raw_env.expanduser().resolve(strict=True)
        cron_jobs = raw_cron.expanduser().resolve(strict=True)
        website = raw_website.expanduser().resolve(strict=True)
        python_declared = Path(os.path.abspath(raw_python.expanduser()))
        python_realpath = python_declared.resolve(strict=True)
        laf_config = raw_laf_config.expanduser().resolve(strict=True)
        google_credentials = raw_google_credentials.expanduser().resolve(strict=True)
        google_calendar_token = raw_google_calendar_token.expanduser().resolve(strict=True)
        laf_gmail_token = raw_laf_gmail_token.expanduser().resolve(strict=True)
        file_review_token = raw_file_review_token.expanduser().resolve(strict=True)
        gmail_compose_token = (
            raw_gmail_compose_token.expanduser().resolve(strict=True)
            if raw_gmail_compose_token is not None
            else None
        )
        accounting_credentials = raw_accounting_credentials.expanduser().resolve(strict=True)
        accounting_token = raw_accounting_token.expanduser().resolve(strict=True)
        drive_sync_token = raw_drive_sync_token.expanduser().resolve(strict=True)
        drive_sync_write_token = raw_drive_sync_write_token.expanduser().resolve(strict=True)
        ocr_queue = raw_ocr_queue.expanduser().resolve(strict=True)
        chromedriver = (
            raw_chromedriver.expanduser().resolve(strict=True)
            if raw_chromedriver is not None
            else None
        )
    except OSError as exc:
        raise DeployPrepareBlocked(f"external runtime input is missing: {exc}") from exc
    if (
        not secret.is_file()
        or not cron_jobs.is_file()
        or not website.is_dir()
        or not laf_config.is_file()
        or not google_credentials.is_file()
        or not google_calendar_token.is_file()
        or not laf_gmail_token.is_file()
        or not file_review_token.is_file()
        or (gmail_compose_token is not None and not gmail_compose_token.is_file())
        or not accounting_credentials.is_file()
        or not accounting_token.is_file()
        or not drive_sync_token.is_file()
        or not drive_sync_write_token.is_file()
        or not ocr_queue.is_file()
        or (
            chromedriver is not None
            and (not chromedriver.is_file() or not os.access(chromedriver, os.X_OK))
        )
    ):
        raise DeployPrepareBlocked("external runtime inputs have invalid types")
    for name, resolved, raw in (
        ("LAF config", laf_config, raw_laf_config),
        ("Google credentials", google_credentials, raw_google_credentials),
        ("Google Calendar token", google_calendar_token, raw_google_calendar_token),
        ("LAF Gmail token", laf_gmail_token, raw_laf_gmail_token),
        ("FileReview token", file_review_token, raw_file_review_token),
        ("accounting credentials", accounting_credentials, raw_accounting_credentials),
        ("accounting Sheets token", accounting_token, raw_accounting_token),
        ("Drive sync token", drive_sync_token, raw_drive_sync_token),
        ("Drive sync write token", drive_sync_write_token, raw_drive_sync_write_token),
        ("NAS OCR queue database", ocr_queue, raw_ocr_queue),
    ):
        if resolved != raw.expanduser():
            raise DeployPrepareBlocked(f"external {name} must be canonical and non-symlinked")
    if (
        chromedriver is not None
        and raw_chromedriver is not None
        and chromedriver != raw_chromedriver.expanduser()
    ):
        raise DeployPrepareBlocked(
            "external ChromeDriver must be canonical and non-symlinked"
        )
    if (
        gmail_compose_token is not None
        and raw_gmail_compose_token is not None
        and gmail_compose_token != raw_gmail_compose_token.expanduser()
    ):
        raise DeployPrepareBlocked(
            "external Gmail compose token must be canonical and non-symlinked"
        )
    if not python_declared.is_file() or not os.access(python_declared, os.X_OK):
        raise DeployPrepareBlocked("external Python runtime must be an executable regular file")
    if python_realpath.is_symlink() or not python_realpath.is_file():
        raise DeployPrepareBlocked("external Python runtime must resolve to a regular file")
    if stat.S_IMODE(secret.stat().st_mode) & 0o077:
        raise DeployPrepareBlocked("external MAGI env file permissions must be 0600 or stricter")
    for name, path in (
        ("LAF config", laf_config),
        ("Google credentials", google_credentials),
        ("accounting credentials", accounting_credentials),
    ):
        if stat.S_IMODE(path.stat().st_mode) != 0o600:
            raise DeployPrepareBlocked(f"external {name} permissions must be exactly 0600")
    ocr_metadata = ocr_queue.lstat()
    if (
        ocr_metadata.st_uid != os.getuid()
        or ocr_metadata.st_nlink != 1
        or stat.S_IMODE(ocr_metadata.st_mode) not in {0o600, 0o640, 0o644}
    ):
        raise DeployPrepareBlocked("external NAS OCR queue owner or mode is unsafe")
    try:
        with sqlite3.connect(f"file:{ocr_queue}?mode=ro", uri=True) as connection:
            quick_check = connection.execute("PRAGMA quick_check").fetchone()
    except sqlite3.Error as exc:
        raise DeployPrepareBlocked(f"external NAS OCR queue quick_check failed: {exc}") from exc
    if quick_check != ("ok",):
        raise DeployPrepareBlocked("external NAS OCR queue quick_check did not return ok")
    admin = website / "admin" / "admin_server.py"
    if admin.is_symlink() or not admin.is_file():
        raise DeployPrepareBlocked("external Website Admin source is missing or unsafe")
    try:
        cron_payload = json.loads(cron_jobs.read_text(encoding="utf-8"))
    except (OSError, UnicodeError, json.JSONDecodeError) as exc:
        raise DeployPrepareBlocked(f"external cron jobs file is unreadable: {exc}") from exc
    if not isinstance(cron_payload, list):
        raise DeployPrepareBlocked("external cron jobs file must contain a JSON list")
    cron_ids: list[str] = []
    for index, row in enumerate(cron_payload):
        if not isinstance(row, dict) or not isinstance(row.get("id"), str) or not row["id"].strip():
            raise DeployPrepareBlocked(f"external cron job {index} must have a non-empty string id")
        cron_ids.append(row["id"].strip())
    if len(cron_ids) != len(set(cron_ids)):
        raise DeployPrepareBlocked("external cron job ids must be unique")
    try:
        laf_config_payload = json.loads(laf_config.read_text(encoding="utf-8"))
    except (OSError, UnicodeError, json.JSONDecodeError) as exc:
        raise DeployPrepareBlocked(f"external LAF config is unreadable: {exc}") from exc
    if not isinstance(laf_config_payload, dict):
        raise DeployPrepareBlocked("external LAF config must contain a JSON object")
    return ExternalRuntimeInputs(
        env_file=secret,
        env_file_sha256=_sha256_file(secret),
        cron_jobs_file=cron_jobs,
        cron_jobs_sha256=_sha256_file(cron_jobs),
        website_root=website,
        website_admin_sha256=_sha256_file(admin),
        python_runtime=python_declared,
        python_runtime_realpath=python_realpath,
        python_runtime_sha256=_sha256_file(python_realpath),
        chromedriver_path=chromedriver,
        chromedriver_sha256=(
            _sha256_file(chromedriver) if chromedriver is not None else None
        ),
        chromedriver_mode=(
            f"{stat.S_IMODE(chromedriver.stat().st_mode):04o}"
            if chromedriver is not None
            else None
        ),
        cron_jobs_source_file=cron_jobs,
        cron_jobs_source_sha256=_sha256_file(cron_jobs),
        cron_snapshot_evidence={},
        laf_config_file=laf_config,
        laf_config_sha256=_sha256_file(laf_config),
        laf_config_mode=f"{stat.S_IMODE(laf_config.stat().st_mode):04o}",
        google_credentials_file=google_credentials,
        google_credentials_sha256=_sha256_file(google_credentials),
        google_credentials_mode=f"{stat.S_IMODE(google_credentials.stat().st_mode):04o}",
        google_calendar_token_source_file=google_calendar_token,
        google_calendar_token_source_sha256=_sha256_file(google_calendar_token),
        laf_gmail_token_source_file=laf_gmail_token,
        laf_gmail_token_source_sha256=_sha256_file(laf_gmail_token),
        file_review_token_source_file=file_review_token,
        file_review_token_source_sha256=_sha256_file(file_review_token),
        gmail_compose_token_source_file=gmail_compose_token,
        gmail_compose_token_source_sha256=(
            _sha256_file(gmail_compose_token) if gmail_compose_token is not None else None
        ),
        accounting_credentials_file=accounting_credentials,
        accounting_credentials_sha256=_sha256_file(accounting_credentials),
        accounting_credentials_mode="0600",
        accounting_sheets_token_source_file=accounting_token,
        accounting_sheets_token_source_sha256=_sha256_file(accounting_token),
        drive_sync_token_source_file=drive_sync_token,
        drive_sync_token_source_sha256=_sha256_file(drive_sync_token),
        drive_sync_write_token_source_file=drive_sync_write_token,
        drive_sync_write_token_source_sha256=_sha256_file(drive_sync_write_token),
        nas_ocr_queue_db_file=ocr_queue,
        nas_ocr_queue_db_mode=f"{stat.S_IMODE(ocr_metadata.st_mode):04o}",
    )


def _validate_release(release_root: Path) -> tuple[Path, ReleaseIdentity]:
    if release_root.is_symlink():
        raise DeployPrepareBlocked("release root must not be a symlink")
    try:
        root = release_root.expanduser().resolve(strict=True)
    except OSError as exc:
        raise DeployPrepareBlocked(f"release root is missing: {exc}") from exc
    if not root.is_dir():
        raise DeployPrepareBlocked("release root must be a directory")

    manifest_path = root / RELEASE_MANIFEST_NAME
    marker_path = root / RELEASE_MARKER_NAME
    manifest = _load_object(manifest_path, description="release manifest")
    marker = _load_object(marker_path, description="release completion marker")
    if manifest.get("schema_version") != 1 or manifest.get("immutable") is not True:
        raise DeployPrepareBlocked("release manifest is not an immutable schema version 1 release")
    if marker.get("schema_version") != 1:
        raise DeployPrepareBlocked("release completion marker schema_version must equal 1")
    release_id = manifest.get("release_id")
    marker_release_id = marker.get("release_id")
    marker_digest = marker.get("manifest_sha256")
    release_sha256 = manifest.get("release_sha256")
    source_snapshot_sha256 = manifest.get("source_snapshot_sha256")
    if not isinstance(release_id, str) or not RELEASE_ID_RE.fullmatch(release_id):
        raise DeployPrepareBlocked("release manifest has an invalid release_id")
    if marker_release_id != release_id:
        raise DeployPrepareBlocked("release completion marker release_id mismatch")
    if marker.get("manifest") != RELEASE_MANIFEST_NAME:
        raise DeployPrepareBlocked("release completion marker names the wrong manifest")
    if (
        not isinstance(release_sha256, str)
        or not SHA256_RE.fullmatch(release_sha256)
        or source_snapshot_sha256 != release_sha256
        or marker.get("release_sha256") != release_sha256
        or marker.get("source_snapshot_sha256") != release_sha256
    ):
        raise DeployPrepareBlocked("release content snapshot identity is invalid or inconsistent")
    if not isinstance(marker_digest, str) or not SHA256_RE.fullmatch(marker_digest):
        raise DeployPrepareBlocked("release completion marker has an invalid manifest hash")
    try:
        manifest_bytes = manifest_path.read_bytes()
    except OSError as exc:
        raise DeployPrepareBlocked(f"release manifest is unreadable: {exc}") from exc
    if _sha256(manifest_bytes) != marker_digest:
        raise DeployPrepareBlocked("release manifest hash does not match completion marker")
    rows = manifest.get("files")
    if not isinstance(rows, list) or not rows:
        raise DeployPrepareBlocked("release manifest files must be a non-empty list")
    file_sha256: dict[str, str] = {}
    file_mode: dict[str, int] = {}
    snapshot_rows: list[dict[str, Any]] = []
    for index, row in enumerate(rows):
        if not isinstance(row, dict):
            raise DeployPrepareBlocked(f"release manifest file {index} must be an object")
        relative = row.get("path")
        digest = row.get("sha256")
        size = row.get("size")
        mode_text = row.get("mode")
        if (
            not isinstance(relative, str)
            or not relative
            or Path(relative).is_absolute()
            or ".." in Path(relative).parts
            or not isinstance(digest, str)
            or not SHA256_RE.fullmatch(digest)
            or not isinstance(size, int)
            or isinstance(size, bool)
            or size < 0
            or not isinstance(mode_text, str)
            or not re.fullmatch(r"0[0-7]{3}", mode_text)
            or relative in file_sha256
        ):
            raise DeployPrepareBlocked(f"release manifest file {index} is invalid")
        file_sha256[relative] = digest
        file_mode[relative] = int(mode_text, 8)
        snapshot_rows.append(
            {"path": relative, "sha256": digest, "size": size, "mode": mode_text}
        )
    snapshot_rows.sort(key=lambda row: row["path"])
    computed_snapshot = _sha256(
        json.dumps(snapshot_rows, sort_keys=True, separators=(",", ":")).encode("utf-8")
    )
    if computed_snapshot != release_sha256:
        raise DeployPrepareBlocked("release content snapshot hash does not match manifest files")
    actual_files: set[str] = set()
    for directory, directory_names, file_names in os.walk(root, followlinks=False):
        base = Path(directory)
        for name in directory_names:
            candidate = base / name
            if candidate.is_symlink():
                raise DeployPrepareBlocked(
                    f"release contains a symlinked directory: {candidate.relative_to(root).as_posix()}"
                )
        for name in file_names:
            candidate = base / name
            relative = candidate.relative_to(root).as_posix()
            if relative in {RELEASE_MANIFEST_NAME, RELEASE_MARKER_NAME}:
                continue
            if candidate.is_symlink() or not candidate.is_file():
                raise DeployPrepareBlocked(f"release member is unsafe: {relative}")
            actual_files.add(relative)
    if actual_files != set(file_sha256):
        missing = sorted(set(file_sha256) - actual_files)
        extra = sorted(actual_files - set(file_sha256))
        raise DeployPrepareBlocked(
            f"release contents differ from manifest: missing={missing[:3]}, extra={extra[:3]}"
        )
    for relative, expected_digest in file_sha256.items():
        candidate = root / relative
        try:
            candidate.resolve(strict=True).relative_to(root)
        except (OSError, ValueError) as exc:
            raise DeployPrepareBlocked(f"release member escapes root: {relative}") from exc
        if _sha256_file(candidate) != expected_digest:
            raise DeployPrepareBlocked(f"release member hash mismatch: {relative}")
        row = next(item for item in snapshot_rows if item["path"] == relative)
        file_stat = candidate.stat()
        if file_stat.st_size != row["size"] or stat.S_IMODE(file_stat.st_mode) != int(row["mode"], 8):
            raise DeployPrepareBlocked(f"release member metadata mismatch: {relative}")
    return root, ReleaseIdentity(
        release_id,
        marker_digest,
        manifest_path,
        file_sha256,
        file_mode,
    )


def _validate_installed_release_immutability(
    installed_root: Path,
    identity: ReleaseIdentity,
) -> None:
    try:
        root_stat = installed_root.lstat()
    except OSError as exc:
        raise DeployPrepareBlocked("installed release root is unreadable") from exc
    if (
        not stat.S_ISDIR(root_stat.st_mode)
        or stat.S_IMODE(root_stat.st_mode) != 0o555
    ):
        raise DeployPrepareBlocked(
            "installed release root must be an immutable 0555 directory"
        )
    expected_files = set(identity.file_sha256) | {
        RELEASE_MANIFEST_NAME,
        RELEASE_MARKER_NAME,
    }
    for directory, directory_names, file_names in os.walk(
        installed_root,
        followlinks=False,
    ):
        base = Path(directory)
        base_stat = base.lstat()
        if (
            stat.S_ISLNK(base_stat.st_mode)
            or not stat.S_ISDIR(base_stat.st_mode)
            or stat.S_IMODE(base_stat.st_mode) != 0o555
        ):
            raise DeployPrepareBlocked(
                "installed release contains a mutable or unsafe directory"
            )
        for name in directory_names:
            if (base / name).is_symlink():
                raise DeployPrepareBlocked(
                    "installed release contains a symlinked directory"
                )
        for name in file_names:
            candidate = base / name
            relative = candidate.relative_to(installed_root).as_posix()
            metadata = candidate.lstat()
            expected_mode = identity.file_mode.get(relative, 0o444)
            if relative not in expected_files:
                raise DeployPrepareBlocked(
                    "installed release contains an unmanifested file"
                )
            if relative in {RELEASE_MANIFEST_NAME, RELEASE_MARKER_NAME}:
                expected_mode = 0o444
            if (
                stat.S_ISLNK(metadata.st_mode)
                or not stat.S_ISREG(metadata.st_mode)
                or metadata.st_nlink != 1
                or expected_mode not in {0o444, 0o555}
                or stat.S_IMODE(metadata.st_mode) != expected_mode
            ):
                raise DeployPrepareBlocked(
                    f"installed release member is not immutable: {relative}"
                )


def _validate_production_release_root(
    source_root: Path,
    source_identity: ReleaseIdentity,
    installed_release_root: Path | None,
) -> tuple[Path, ReleaseIdentity]:
    """Bind production only to the canonical installed copy of the release.

    A candidate/evidence tree is useful for validation, but launchd cannot safely
    execute it from Desktop or another operator-owned staging domain.  The
    installed tree is independently revalidated and must have exactly the same
    immutable identity as the candidate.
    """

    canonical = _canonical_installed_release_root(source_identity.release_id)
    requested = installed_release_root or canonical
    installed, installed_identity = _validate_release(requested)
    _validate_installed_release_immutability(installed, installed_identity)
    try:
        canonical_realpath = canonical.resolve(strict=True)
    except OSError as exc:
        raise DeployPrepareBlocked(
            f"canonical installed release is missing: {canonical}"
        ) from exc
    if installed != canonical_realpath:
        raise DeployPrepareBlocked(
            "production installed release must resolve to the canonical "
            f"Application Support release root: {canonical_realpath}"
        )
    if (
        installed_identity.release_id != source_identity.release_id
        or installed_identity.manifest_sha256 != source_identity.manifest_sha256
        or dict(installed_identity.file_sha256) != dict(source_identity.file_sha256)
        or dict(installed_identity.file_mode) != dict(source_identity.file_mode)
    ):
        raise DeployPrepareBlocked(
            "installed release immutable identity does not match the validated candidate"
        )
    return installed, installed_identity


def _installed_release_member(
    source_member: Path,
    *,
    source_root: Path,
    installed_root: Path,
    installed_identity: ReleaseIdentity,
    description: str,
) -> Path:
    try:
        relative = source_member.relative_to(source_root)
    except ValueError as exc:  # Defensive: source members are already validated.
        raise DeployPrepareBlocked(f"{description} escapes candidate release") from exc
    installed_member = installed_root / relative
    expected = installed_identity.file_sha256.get(relative.as_posix())
    if (
        expected is None
        or installed_member.is_symlink()
        or not installed_member.is_file()
        or _sha256_file(installed_member) != expected
    ):
        raise DeployPrepareBlocked(
            f"installed {description} is missing or not bound by the installed release"
        )
    return installed_member


def load_roles(path: Path) -> tuple[RoleDefinition, ...]:
    payload = _load_object(path, description="launchagent role config")
    if payload.get("schema_version") != 1:
        raise DeployPrepareBlocked("launchagent role config schema_version must equal 1")
    rows = payload.get("roles")
    if not isinstance(rows, list) or len(rows) != len(EXPECTED_LABELS):
        raise DeployPrepareBlocked("launchagent role config must define exactly three roles")
    roles: list[RoleDefinition] = []
    seen_ports: set[int] = set()
    seen_domains: set[str] = set()
    for index, row in enumerate(rows):
        if not isinstance(row, dict):
            raise DeployPrepareBlocked(f"launchagent role {index} must be an object")
        role = row.get("role")
        label = row.get("label")
        module = row.get("entrypoint_module")
        if role not in EXPECTED_LABELS or label != EXPECTED_LABELS[role]:
            raise DeployPrepareBlocked(f"launchagent role {index} has an invalid fixed label")
        if not isinstance(module, str) or not MODULE_RE.fullmatch(module):
            raise DeployPrepareBlocked(f"launchagent role {role} has an invalid entrypoint module")
        raw_ports = row.get("ports")
        if not isinstance(raw_ports, list) or any(
            not isinstance(port, int) or isinstance(port, bool) or not 1 <= port <= 65535
            for port in raw_ports
        ):
            raise DeployPrepareBlocked(f"launchagent role {role} has invalid ports")
        ports = tuple(raw_ports)
        if len(ports) != len(set(ports)) or seen_ports.intersection(ports):
            raise DeployPrepareBlocked("launchagent ports must be globally unique")
        seen_ports.update(ports)
        raw_domains = row.get("ownership_domains")
        if not isinstance(raw_domains, list) or not raw_domains or any(
            not isinstance(domain, str) or not domain.strip() for domain in raw_domains
        ):
            raise DeployPrepareBlocked(f"launchagent role {role} has invalid ownership domains")
        domains = tuple(raw_domains)
        if len(domains) != len(set(domains)) or seen_domains.intersection(domains):
            raise DeployPrepareBlocked("ownership domains must be globally unique")
        seen_domains.update(domains)
        roles.append(RoleDefinition(role, label, module, ports, domains))
    if {role.role for role in roles} != set(EXPECTED_LABELS):
        raise DeployPrepareBlocked("launchagent role config is incomplete or duplicated")
    return tuple(sorted(roles, key=lambda role: role.role))


def _validate_executable(
    path: Path,
    release_root: Path,
    identity: ReleaseIdentity,
) -> Path:
    if path.expanduser().is_symlink():
        raise DeployPrepareBlocked("required executable must not be a symlink")
    executable = _validated_absolute(path, name="python executable")
    try:
        relative = executable.relative_to(release_root)
    except ValueError as exc:
        raise DeployPrepareBlocked("required executable must be inside the immutable release root") from exc
    try:
        mode = executable.stat().st_mode
    except OSError as exc:
        raise DeployPrepareBlocked(f"required executable is missing: {executable}") from exc
    if not stat.S_ISREG(mode) or not os.access(executable, os.X_OK):
        raise DeployPrepareBlocked(f"required executable is not executable: {executable}")
    expected_digest = identity.file_sha256.get(relative.as_posix())
    if expected_digest is None:
        raise DeployPrepareBlocked(
            f"required executable is absent from release manifest: {relative.as_posix()}"
        )
    if _sha256_file(executable) != expected_digest:
        raise DeployPrepareBlocked(
            f"required executable hash does not match release manifest: {relative.as_posix()}"
        )
    return executable


def _validate_release_data_file(
    path: Path,
    release_root: Path,
    identity: ReleaseIdentity,
    *,
    description: str,
) -> Path:
    if path.expanduser().is_symlink():
        raise DeployPrepareBlocked(f"{description} must not be a symlink")
    candidate = path.expanduser().resolve(strict=False)
    try:
        relative = candidate.relative_to(release_root)
    except ValueError as exc:
        raise DeployPrepareBlocked(f"{description} must be inside the immutable release root") from exc
    if not candidate.is_file():
        raise DeployPrepareBlocked(f"{description} is missing: {candidate}")
    expected_digest = identity.file_sha256.get(relative.as_posix())
    if expected_digest is None or _sha256_file(candidate) != expected_digest:
        raise DeployPrepareBlocked(f"{description} is not hash-bound by the release manifest")
    return candidate


def _validate_service_manifest(
    path: Path,
    release_root: Path,
    identity: ReleaseIdentity,
    *,
    deployment_mode: str,
) -> tuple[Path, str, tuple[Path, ...]]:
    candidate = _validate_release_data_file(
        path,
        release_root,
        identity,
        description="service manifest",
    )
    payload = _load_object(candidate, description="service manifest")
    if (
        payload.get("schema_version") != 1
        or payload.get("release_mode") != "single_active_replacement"
        or payload.get("deployment_mode") != deployment_mode
    ):
        raise DeployPrepareBlocked("service manifest does not match the requested deployment mode")
    try:
        parsed = load_service_manifest(candidate)
    except ConfigurationError as exc:
        raise DeployPrepareBlocked(f"service manifest topology is invalid: {exc}") from exc
    if parsed.deployment_mode != deployment_mode:
        raise DeployPrepareBlocked("service manifest parsed mode mismatch")
    process_entrypoints: list[Path] = []
    for service in parsed.services:
        if service.kind != "process":
            continue
        process_entrypoints.append(
            _validate_release_data_file(
                release_root / service.argv[1],
                release_root,
                identity,
                description=f"service entrypoint {service.service_id}",
            )
        )
    return candidate, _sha256_file(candidate), tuple(process_entrypoints)


def _validate_production_cron_policy_binding(
    release_root: Path,
    identity: ReleaseIdentity,
    external_inputs: ExternalRuntimeInputs,
) -> None:
    """Reject a production deployment whose cron source is not release-approved."""

    policy_path = _validate_release_data_file(
        release_root / CRON_DISPATCH_POLICY_NAME,
        release_root,
        identity,
        description="cron dispatch policy",
    )
    policy = _load_object(policy_path, description="cron dispatch policy")
    expected_source_sha256 = policy.get("cron_jobs_sha256")
    if (
        not isinstance(expected_source_sha256, str)
        or SHA256_RE.fullmatch(expected_source_sha256) is None
    ):
        raise DeployPrepareBlocked(
            "cron dispatch policy cron_jobs_sha256 must be lowercase SHA-256"
        )
    if expected_source_sha256 != external_inputs.cron_jobs_source_sha256:
        raise DeployPrepareBlocked(
            "production cron source SHA-256 does not match the release cron dispatch policy"
        )


def _validate_production_external_independence(
    external_inputs: ExternalRuntimeInputs,
) -> None:
    v2_roots = (
        (Path.home() / "Library/Application Support/MAGI/runtime/MAGI_v2").resolve(
            strict=False
        ),
        (Path.home() / "Desktop/MAGI_v2").resolve(strict=False),
    )
    for name, path in (
        ("environment", external_inputs.env_file),
        ("website", external_inputs.website_root),
        ("runtime config", external_inputs.laf_config_file),
        ("Google credentials", external_inputs.google_credentials_file),
        ("accounting credentials", external_inputs.accounting_credentials_file),
    ):
        if any(path == root or path.is_relative_to(root) for root in v2_roots):
            raise DeployPrepareBlocked(
                f"production {name} must be staged outside the V2 rollback tree"
            )
    if external_inputs.nas_ocr_queue_db_file != _canonical_nas_ocr_queue_db():
        raise DeployPrepareBlocked(
            "production NAS OCR queue must use the canonical ~/.magi_nas_ocr_queue.db"
        )


def _validate_production_static_external(
    static_external_receipt: Path | None,
    *,
    identity: ReleaseIdentity,
    runtime_root: Path,
    deployment_root: Path,
    external_inputs: ExternalRuntimeInputs,
) -> StaticExternalEvidence:
    raw = static_external_receipt or Path(
        os.environ.get("MAGI_V3_STATIC_EXTERNAL_RECEIPT", "")
    )
    if not str(raw) or not raw.expanduser().is_absolute() or raw.expanduser().is_symlink():
        raise DeployPrepareBlocked(
            "production deployment requires an absolute non-symlink static external receipt"
        )
    expected_root = runtime_root / "shared" / "external"
    expected_receipt = expected_root / STATIC_EXTERNAL_RECEIPT_NAME
    if raw.expanduser() != expected_receipt:
        raise DeployPrepareBlocked(
            "production static external receipt must use the canonical V3 shared target"
        )
    try:
        report = verify_static_external_payload(
            target_root=expected_root,
        )
    except StaticExternalStagingError as exc:
        raise DeployPrepareBlocked(f"production static external receipt is invalid: {exc}") from exc
    summaries = {
        row["logical_id"]: row for row in report.get("logical_inputs", [])
        if isinstance(row, dict) and isinstance(row.get("logical_id"), str)
    }
    required_summaries = {
        "environment",
        "website",
        "runtime_config",
        "google_credentials",
        "accounting_credentials",
    }
    if set(summaries) != required_summaries:
        raise DeployPrepareBlocked("production static external receipt logical inputs mismatch")
    required_paths = {
        "environment": external_inputs.env_file,
        "website": external_inputs.website_root,
        "runtime_config": external_inputs.laf_config_file,
        "google_credentials": external_inputs.google_credentials_file,
        "accounting_credentials": external_inputs.accounting_credentials_file,
    }
    expected_paths = {
        "environment": expected_root / ".env",
        "website": expected_root / "website",
        "runtime_config": expected_root / "config.json",
        "google_credentials": expected_root / "google-credentials.json",
        "accounting_credentials": expected_root / "accounting-credentials.json",
    }
    if any(required_paths[key] != expected_paths[key] for key in required_summaries):
        raise DeployPrepareBlocked(
            "production five static runtime inputs must exactly use the staged V3 target"
        )
    expected_content_hashes = {
        "environment": external_inputs.env_file_sha256,
        "runtime_config": external_inputs.laf_config_sha256,
        "google_credentials": external_inputs.google_credentials_sha256,
        "accounting_credentials": external_inputs.accounting_credentials_sha256,
    }
    if any(
        summaries[key].get("content_sha256") != expected_hash
        for key, expected_hash in expected_content_hashes.items()
    ):
        raise DeployPrepareBlocked("production static external file hash binding mismatch")
    payload_receipt = expected_receipt.resolve(strict=True)
    release_receipt = (
        deployment_root / "runtime-inputs" / STATIC_EXTERNAL_RELEASE_RECEIPT_NAME
    )
    try:
        receipt_bytes, binding = render_static_external_release_binding(
            identity.manifest_path,
            expected_release_manifest_sha256=identity.manifest_sha256,
            target_root=expected_root,
            binding_receipt=release_receipt,
        )
    except StaticExternalStagingError as exc:
        raise DeployPrepareBlocked(
            f"production static external release binding is invalid: {exc}"
        ) from exc
    return StaticExternalEvidence(
        receipt=release_receipt,
        receipt_sha256=str(binding["binding_receipt_sha256"]),
        receipt_bytes=receipt_bytes,
        payload_receipt=payload_receipt,
        payload_receipt_sha256=str(report["receipt_sha256"]),
        source_snapshot_sha256=str(report["source_snapshot_sha256"]),
        target_snapshot_sha256=str(report["target_snapshot_sha256"]),
    )


def _validate_isolated_validation_inputs(
    validation_input_root: Path | None,
    external_inputs: ExternalRuntimeInputs,
    *,
    release_root: Path,
    staging: Path,
    publish: Path,
    runtime: Path,
) -> Path:
    if validation_input_root is None:
        raise DeployPrepareBlocked("isolated live validation requires --validation-input-root")
    raw = validation_input_root.expanduser()
    if not raw.is_absolute() or raw.is_symlink():
        raise DeployPrepareBlocked("validation input root must be an absolute non-symlink directory")
    try:
        root = raw.resolve(strict=True)
    except OSError as exc:
        raise DeployPrepareBlocked(f"validation input root is missing: {exc}") from exc
    if not root.is_dir():
        raise DeployPrepareBlocked("validation input root must be a directory")
    if any(
        _inside(root, domain) or _inside(domain, root)
        for domain in (release_root, staging, publish, runtime, _application_support_root())
    ):
        raise DeployPrepareBlocked("validation input root overlaps a release or live runtime domain")
    for name, path in (
        ("environment", external_inputs.env_file),
        ("cron", external_inputs.cron_jobs_source_file),
        ("website", external_inputs.website_root),
        ("LAF config", external_inputs.laf_config_file),
        ("Google credentials", external_inputs.google_credentials_file),
        ("Google Calendar token", external_inputs.google_calendar_token_source_file),
        ("LAF Gmail token", external_inputs.laf_gmail_token_source_file),
        ("FileReview token", external_inputs.file_review_token_source_file),
        ("accounting credentials", external_inputs.accounting_credentials_file),
        ("accounting Sheets token", external_inputs.accounting_sheets_token_source_file),
        ("Drive sync token", external_inputs.drive_sync_token_source_file),
        ("Drive sync write token", external_inputs.drive_sync_write_token_source_file),
        ("NAS OCR queue database", external_inputs.nas_ocr_queue_db_file),
    ):
        if not _inside(path, root):
            raise DeployPrepareBlocked(f"validation {name} input must be inside validation input root")
    if external_inputs.env_file.read_bytes() != VALIDATION_ENV_BYTES:
        raise DeployPrepareBlocked("validation environment file must contain only the inert fixture marker")
    try:
        cron_payload = json.loads(external_inputs.cron_jobs_source_file.read_text(encoding="utf-8"))
    except (OSError, UnicodeError, json.JSONDecodeError) as exc:
        raise DeployPrepareBlocked(f"validation cron input is unreadable: {exc}") from exc
    if cron_payload != [VALIDATION_CRON_JOB]:
        raise DeployPrepareBlocked("validation cron input must contain only the disabled inert job")
    if external_inputs.laf_config_file.read_bytes() != VALIDATION_LAF_CONFIG_BYTES:
        raise DeployPrepareBlocked(
            "validation LAF config must contain only the inert local fixture"
        )
    if (
        external_inputs.google_credentials_file.read_bytes() != VALIDATION_GOOGLE_CREDENTIALS_BYTES
        or external_inputs.google_calendar_token_source_file.read_bytes()
        != VALIDATION_GOOGLE_CALENDAR_TOKEN_BYTES
        or external_inputs.laf_gmail_token_source_file.read_bytes()
        != VALIDATION_LAF_GMAIL_TOKEN_BYTES
        or external_inputs.file_review_token_source_file.read_bytes()
        != VALIDATION_LAF_GMAIL_TOKEN_BYTES
    ):
        raise DeployPrepareBlocked("validation credential inputs must be inert fixtures")
    if (
        external_inputs.accounting_credentials_file.read_bytes()
        != VALIDATION_GOOGLE_CREDENTIALS_BYTES
        or external_inputs.accounting_sheets_token_source_file.read_bytes()
        != VALIDATION_GOOGLE_CALENDAR_TOKEN_BYTES
        or external_inputs.drive_sync_token_source_file.read_bytes()
        != VALIDATION_GOOGLE_CALENDAR_TOKEN_BYTES
        or external_inputs.drive_sync_write_token_source_file.read_bytes()
        != VALIDATION_GOOGLE_CALENDAR_TOKEN_BYTES
    ):
        raise DeployPrepareBlocked("validation accounting/Drive inputs must be inert fixtures")
    fixture = external_inputs.website_root / "data" / "live-validation-document.txt"
    if fixture.is_symlink() or not fixture.is_file() or fixture.stat().st_size > 1024 * 1024:
        raise DeployPrepareBlocked("validation website document fixture is missing or unsafe")
    for directory, directory_names, file_names in os.walk(root, followlinks=False):
        base = Path(directory)
        for name in (*directory_names, *file_names):
            member = base / name
            if member.is_symlink():
                raise DeployPrepareBlocked("validation input tree must not contain symlinks")
    return root


def _validate_case_filesystem_bindings(
    deployment_mode: str,
    case_root: Path | None,
    archive_root: Path | None,
    path_mappings: tuple[tuple[str, str], ...] | None,
) -> tuple[Path | None, Path | None, tuple[tuple[str, str], ...]]:
    if deployment_mode == "isolated_live_validation":
        if any(value is not None for value in (case_root, archive_root, path_mappings)):
            raise DeployPrepareBlocked("isolated live validation cannot bind production case roots")
        return None, None, ()
    raw_case = case_root or Path(os.environ.get("MAGI_V3_CASE_ROOT", ""))
    raw_archive = archive_root or Path(os.environ.get("MAGI_V3_ARCHIVE_ROOT", ""))
    if (
        not str(raw_case)
        or not raw_case.expanduser().is_absolute()
        or raw_case.expanduser().is_symlink()
        or not str(raw_archive)
        or not raw_archive.expanduser().is_absolute()
        or raw_archive.expanduser().is_symlink()
    ):
        raise DeployPrepareBlocked("production case and archive roots must be absolute non-symlink paths")
    try:
        active = raw_case.expanduser().resolve(strict=True)
        archive = raw_archive.expanduser().resolve(strict=True)
    except OSError as exc:
        raise DeployPrepareBlocked(f"production case filesystem root is unavailable: {exc}") from exc
    if not active.is_dir() or not archive.is_dir() or active == archive:
        raise DeployPrepareBlocked("production case and archive roots must be distinct directories")
    mappings = path_mappings
    if mappings is None:
        raw = os.environ.get("MAGI_V3_PATH_MAPPINGS_JSON", "").strip()
        try:
            decoded = json.loads(raw) if raw else None
        except json.JSONDecodeError as exc:
            raise DeployPrepareBlocked("MAGI_V3_PATH_MAPPINGS_JSON is invalid") from exc
        if not isinstance(decoded, list):
            raise DeployPrepareBlocked("production path mappings must be a JSON list")
        mappings = tuple(tuple(row) for row in decoded if isinstance(row, list))
        if len(mappings) != len(decoded):
            raise DeployPrepareBlocked("production path mappings contain an invalid row")
    normalized: list[tuple[str, str]] = []
    for index, row in enumerate(mappings):
        if (
            not isinstance(row, tuple)
            or len(row) != 2
            or any(not isinstance(item, str) or not item.strip() or "\x00" in item for item in row)
        ):
            raise DeployPrepareBlocked(f"production path mapping {index} is invalid")
        source = Path(row[0]).expanduser()
        if not source.is_absolute() or source.is_symlink():
            raise DeployPrepareBlocked(f"production path mapping {index} source is unsafe")
        try:
            source = source.resolve(strict=True)
        except OSError as exc:
            raise DeployPrepareBlocked(f"production path mapping {index} source is unavailable") from exc
        target = row[1].strip().replace("/", "\\").rstrip("\\")
        if not (re.fullmatch(r"[A-Za-z]:", target) or target.startswith("\\\\")):
            raise DeployPrepareBlocked(
                f"production path mapping {index} target must be a drive or UNC root"
            )
        normalized.append((str(source), target))
    if not normalized or len({row[0].lower() for row in normalized}) != len(normalized):
        raise DeployPrepareBlocked("production path mappings are missing or duplicated")
    if not all(
        any(_inside(root, Path(source)) for source, _target in normalized)
        for root in (active, archive)
    ):
        raise DeployPrepareBlocked("production case roots are not covered by canonical path mappings")
    return active, archive, tuple(normalized)


def _entrypoint_path(
    release_root: Path,
    module: str,
    identity: ReleaseIdentity,
) -> Path:
    relative = Path(*module.split(".")).with_suffix(".py")
    candidate = release_root / relative
    if candidate.is_symlink() or not candidate.is_file():
        raise DeployPrepareBlocked(f"required entrypoint is missing: {relative.as_posix()}")
    try:
        candidate.resolve(strict=True).relative_to(release_root)
    except (OSError, ValueError) as exc:
        raise DeployPrepareBlocked(f"required entrypoint escapes release: {relative.as_posix()}") from exc
    expected_digest = identity.file_sha256.get(relative.as_posix())
    if expected_digest is None:
        raise DeployPrepareBlocked(
            f"required entrypoint is absent from release manifest: {relative.as_posix()}"
        )
    if _sha256_file(candidate) != expected_digest:
        raise DeployPrepareBlocked(
            f"required entrypoint hash does not match release manifest: {relative.as_posix()}"
        )
    return candidate


def _safe_paths(
    release_root: Path,
    staging_dir: Path,
    runtime_root: Path,
    publish_dir: Path | None = None,
) -> tuple[Path, Path, Path]:
    staging = _validated_absolute(staging_dir, name="staging directory")
    runtime = _validated_absolute(runtime_root, name="runtime root")
    publish = (
        _validated_absolute(publish_dir, name="publish directory")
        if publish_dir is not None
        else staging
    )
    if staging.exists() or staging.is_symlink():
        raise DeployPrepareBlocked("staging directory must not already exist")
    if not staging.parent.is_dir():
        raise DeployPrepareBlocked("staging parent must already exist")
    if publish != staging:
        if publish.exists() or publish.is_symlink():
            raise DeployPrepareBlocked("publish directory must not already exist")
        if publish.parent != staging.parent:
            raise DeployPrepareBlocked(
                "publish directory must share the staging parent for atomic publication"
            )
    application_support = _application_support_root()
    if _inside(staging, application_support) or _inside(publish, application_support):
        raise DeployPrepareBlocked(
            "staging directory and publish directory must not target Application Support"
        )
    canonical_runtime = _canonical_runtime_root().expanduser().resolve(strict=False)
    if runtime != canonical_runtime:
        raise DeployPrepareBlocked(
            f"runtime root must equal canonical V3 runtime: {canonical_runtime}"
        )
    if (
        _inside(staging, release_root)
        or _inside(publish, release_root)
        or _inside(staging, runtime)
        or _inside(publish, runtime)
        or _inside(runtime, release_root)
        or _inside(release_root, runtime)
    ):
        raise DeployPrepareBlocked(
            "release, staging, publish, and runtime domains must not overlap"
        )
    return staging, publish, runtime


def _role_binding(
    role: RoleDefinition,
    *,
    release_root: Path,
    executable: Path,
    runtime_root: Path,
    identity: ReleaseIdentity,
    executable_sha256: str,
    external_inputs: ExternalRuntimeInputs,
    python_runtime_manifest: Path,
    python_runtime_manifest_sha256: str,
    python_runtime_tree_sha256: str,
    deployment_mode: str,
    service_manifest: Path,
    service_manifest_sha256: str,
    case_root: Path | None,
    archive_root: Path | None,
    path_mappings: tuple[tuple[str, str], ...],
    static_external: StaticExternalEvidence | None,
) -> dict[str, Any]:
    state_dir = runtime_root / "state" / role.role
    log_dir = runtime_root / "logs" / role.role
    pid_file = runtime_root / "pids" / f"{role.role}.pid"
    ownership_path = runtime_root / OWNERSHIP_MANIFEST_NAME
    # The externally managed Python process replaces the release-contained
    # shell launcher. Keep one no-space, release-resolving token in argv so
    # the fail-closed listener probe can classify the process after it binds a
    # production port. The installer creates this exact symlink atomically.
    owner_token = Path.home() / f".magi-{identity.release_id}-owner"
    invocation = (
        "import runpy,sys; sys.argv=sys.argv[:1]; "
        f"runpy.run_module({role.entrypoint_module!r}, run_name='__main__')"
    )
    arguments = [str(executable), "-c", invocation, str(owner_token)]
    secrets_root = runtime_root / "shared" / "secrets"
    return {
        "role": role.role,
        "label": role.label,
        "ProgramArguments": arguments,
        "owner_token": str(owner_token),
        "executable_sha256": executable_sha256,
        "WorkingDirectory": str(release_root),
        "state_dir": str(state_dir),
        "log_dir": str(log_dir),
        "pid_file": str(pid_file),
        "ports": list(role.ports),
        "ownership_domains": list(role.ownership_domains),
        "ownership_manifest": str(ownership_path),
        "runtime_root": str(runtime_root),
        "release_id": identity.release_id,
        "release_manifest": str(identity.manifest_path),
        "release_manifest_sha256": identity.manifest_sha256,
        "deployment_mode": deployment_mode,
        "service_manifest": str(service_manifest),
        "service_manifest_sha256": service_manifest_sha256,
        "case_root": str(case_root) if case_root is not None else None,
        "archive_root": str(archive_root) if archive_root is not None else None,
        "path_mappings": [list(row) for row in path_mappings],
        "env_file": str(external_inputs.env_file),
        "env_file_sha256": external_inputs.env_file_sha256,
        "cron_jobs_file": str(external_inputs.cron_jobs_file),
        "cron_jobs_sha256": external_inputs.cron_jobs_sha256,
        "cron_jobs_source_sha256": external_inputs.cron_jobs_source_sha256,
        "website_root": str(external_inputs.website_root),
        "website_admin_sha256": external_inputs.website_admin_sha256,
        "laf_config_file": str(external_inputs.laf_config_file),
        "laf_config_sha256": external_inputs.laf_config_sha256,
        "laf_config_mode": external_inputs.laf_config_mode,
        "google_credentials_file": str(external_inputs.google_credentials_file),
        "google_credentials_sha256": external_inputs.google_credentials_sha256,
        "google_credentials_mode": external_inputs.google_credentials_mode,
        "google_calendar_token_source_file": str(external_inputs.google_calendar_token_source_file),
        "google_calendar_token_source_sha256": external_inputs.google_calendar_token_source_sha256,
        "google_calendar_token_file": str(secrets_root / "google_calendar_token.json"),
        "laf_gmail_token_source_file": str(external_inputs.laf_gmail_token_source_file),
        "laf_gmail_token_source_sha256": external_inputs.laf_gmail_token_source_sha256,
        "laf_gmail_token_file": str(secrets_root / "laf_gmail_token.pickle"),
        "file_review_token_source_file": str(external_inputs.file_review_token_source_file),
        "file_review_token_source_sha256": external_inputs.file_review_token_source_sha256,
        "file_review_token_file": str(secrets_root / "filereview_token.pickle"),
        "gmail_compose_token_source_file": (
            str(external_inputs.gmail_compose_token_source_file)
            if external_inputs.gmail_compose_token_source_file is not None
            else None
        ),
        "gmail_compose_token_source_sha256": external_inputs.gmail_compose_token_source_sha256,
        "gmail_compose_token_file": str(secrets_root / "gmail_compose_token.json"),
        "optional_degraded_inputs": (
            [] if external_inputs.gmail_compose_token_source_file is not None else ["gmail_compose_token"]
        ),
        "accounting_credentials_file": str(external_inputs.accounting_credentials_file),
        "accounting_credentials_sha256": external_inputs.accounting_credentials_sha256,
        "accounting_credentials_mode": external_inputs.accounting_credentials_mode,
        "accounting_sheets_token_source_file": str(
            external_inputs.accounting_sheets_token_source_file
        ),
        "accounting_sheets_token_source_sha256": (
            external_inputs.accounting_sheets_token_source_sha256
        ),
        "accounting_sheets_token_file": str(secrets_root / "accounting_sheets_token.json"),
        "drive_sync_token_source_file": str(external_inputs.drive_sync_token_source_file),
        "drive_sync_token_source_sha256": external_inputs.drive_sync_token_source_sha256,
        "drive_sync_token_file": str(secrets_root / "drive_sync_token.json"),
        "drive_sync_write_token_source_file": str(
            external_inputs.drive_sync_write_token_source_file
        ),
        "drive_sync_write_token_source_sha256": (
            external_inputs.drive_sync_write_token_source_sha256
        ),
        "drive_sync_write_token_file": str(secrets_root / "drive_sync_write_token.json"),
        "nas_ocr_queue_db_file": str(external_inputs.nas_ocr_queue_db_file),
        "nas_ocr_queue_db_mode": external_inputs.nas_ocr_queue_db_mode,
        "python_runtime": str(external_inputs.python_runtime),
        "python_runtime_realpath": str(external_inputs.python_runtime_realpath),
        "python_runtime_sha256": external_inputs.python_runtime_sha256,
        "chromedriver_path": (
            str(external_inputs.chromedriver_path)
            if external_inputs.chromedriver_path is not None
            else None
        ),
        "chromedriver_sha256": external_inputs.chromedriver_sha256,
        "chromedriver_mode": external_inputs.chromedriver_mode,
        "python_runtime_manifest": str(python_runtime_manifest),
        "python_runtime_manifest_sha256": python_runtime_manifest_sha256,
        "python_runtime_tree_sha256": python_runtime_tree_sha256,
        "static_external_receipt": (
            str(static_external.receipt) if static_external is not None else None
        ),
        "static_external_receipt_sha256": (
            static_external.receipt_sha256 if static_external is not None else None
        ),
        "static_external_target_snapshot_sha256": (
            static_external.target_snapshot_sha256 if static_external is not None else None
        ),
        **named_mutable_state_paths(runtime_root),
    }


def _plist(binding: Mapping[str, Any]) -> bytes:
    log_dir = Path(binding["log_dir"])
    runtime_root = Path(binding["runtime_root"])
    shared_root = runtime_root / "shared"
    validation_mode = binding["deployment_mode"] == "isolated_live_validation"
    environment = {
        "MAGI_V3_ROLE": binding["role"],
        "MAGI_V3_EXECUTABLE_PATH": binding["ProgramArguments"][0],
        "MAGI_V3_EXECUTABLE_SHA256": binding["executable_sha256"],
        "MAGI_V3_RELEASE_ID": binding["release_id"],
        "MAGI_V3_RELEASE_MANIFEST": binding["release_manifest"],
        "MAGI_V3_RELEASE_MANIFEST_SHA256": binding["release_manifest_sha256"],
        "MAGI_V3_DEPLOYMENT_MODE": binding["deployment_mode"],
        "MAGI_V3_EXTERNAL_INPUT_CONTRACT": "1",
        "MAGI_V3_SERVICE_MANIFEST": binding["service_manifest"],
        "MAGI_V3_SERVICE_MANIFEST_SHA256": binding["service_manifest_sha256"],
        "MAGI_V3_LIVE_VALIDATION": "1" if validation_mode else "0",
        "MAGI_V3_EXTERNAL_WRITES_ENABLED": "0" if validation_mode else "1",
        "MAGI_V3_NOTIFICATIONS_ENABLED": "0" if validation_mode else "1",
        "MAGI_V3_SCHEDULER_ENABLED": "0" if validation_mode else "1",
        "MAGI_V3_STATE_DIR": binding["state_dir"],
        "MAGI_V3_SHARED_STATE_DIR": str(shared_root),
        "MAGI_SHARED_STATE_DIR": str(shared_root),
        "MAGI_V3_LOG_DIR": binding["log_dir"],
        "MAGI_V3_PID_FILE": binding["pid_file"],
        "MAGI_V3_PORTS": ",".join(str(port) for port in binding["ports"]),
        "MAGI_V3_OWNERSHIP_DOMAINS": ",".join(binding["ownership_domains"]),
        "MAGI_V3_OWNERSHIP_MANIFEST": binding["ownership_manifest"],
        "MAGI_V3_OWNERSHIP_MANIFEST_SHA256": binding["ownership_manifest_sha256"],
        "MAGI_V3_ACTIVE_RELEASE_MARKER": str(
            Path.home()
            / "Library"
            / "Application Support"
            / "MAGI"
            / "runtime"
            / "active-release.json"
        ),
        "MAGI_V3_REQUIRE_ACTIVE_MARKER": "0" if validation_mode else "1",
        "MAGI_ENV_FILE": binding["env_file"],
        "MAGI_ENV_FILE_SHA256": binding["env_file_sha256"],
        "MAGI_CRON_JOBS_FILE": binding["cron_jobs_file"],
        "MAGI_CRON_JOBS_SHA256": binding["cron_jobs_sha256"],
        "MAGI_CRON_JOBS_SOURCE_SHA256": binding["cron_jobs_source_sha256"],
        "MAGI_WEBSITE_ROOT": binding["website_root"],
        "MAGI_WEBSITE_ADMIN_SHA256": binding["website_admin_sha256"],
        "MAGI_LAF_CONFIG_FILE": binding["laf_config_file"],
        "MAGI_LAF_CONFIG_SHA256": binding["laf_config_sha256"],
        "MAGI_CONFIG_PATH": binding["laf_config_file"],
        "MAGI_CONFIG_SHA256": binding["laf_config_sha256"],
        "MAGI_CONFIG_MODE": binding["laf_config_mode"],
        "MAGI_GOOGLE_CREDENTIALS_PATH": binding["google_credentials_file"],
        "MAGI_GOOGLE_CREDENTIALS_SHA256": binding["google_credentials_sha256"],
        "MAGI_GOOGLE_CREDENTIALS_MODE": binding["google_credentials_mode"],
        "MAGI_GMAIL_CREDENTIALS_PATH": binding["google_credentials_file"],
        "MAGI_GOOGLE_CALENDAR_TOKEN_PATH": binding["google_calendar_token_file"],
        "MAGI_LAF_GMAIL_TOKEN_PATH": binding["laf_gmail_token_file"],
        "MAGI_FILE_REVIEW_TOKEN_PATH": binding["file_review_token_file"],
        "MAGI_GMAIL_COMPOSE_TOKEN_PATH": binding["gmail_compose_token_file"],
        "MAGI_V3_OPTIONAL_DEGRADED_INPUTS": ",".join(binding["optional_degraded_inputs"]),
        "MAGI_ACCOUNTING_GOOGLE_CREDENTIALS_PATH": binding["accounting_credentials_file"],
        "MAGI_ACCOUNTING_GOOGLE_CREDENTIALS_SHA256": binding[
            "accounting_credentials_sha256"
        ],
        "MAGI_ACCOUNTING_GOOGLE_CREDENTIALS_MODE": binding["accounting_credentials_mode"],
        "MAGI_ACCOUNTING_GOOGLE_SHEETS_TOKEN": binding["accounting_sheets_token_file"],
        "MAGI_DRIVE_SYNC_CREDENTIALS_PATH": binding["google_credentials_file"],
        "MAGI_DRIVE_SYNC_TOKEN": binding["drive_sync_token_file"],
        "MAGI_DRIVE_SYNC_WRITE_TOKEN": binding["drive_sync_write_token_file"],
        "MAGI_NAS_OCR_QUEUE_DB_PATH": binding["nas_ocr_queue_db_file"],
        "MAGI_V3_PYTHON_RUNTIME": binding["python_runtime"],
        "MAGI_V3_PYTHON_RUNTIME_REALPATH": binding["python_runtime_realpath"],
        "MAGI_V3_PYTHON_RUNTIME_SHA256": binding["python_runtime_sha256"],
        "MAGI_V3_PYTHON_RUNTIME_MANIFEST": binding["python_runtime_manifest"],
        "MAGI_V3_PYTHON_RUNTIME_MANIFEST_SHA256": binding["python_runtime_manifest_sha256"],
        "MAGI_V3_PYTHON_RUNTIME_TREE_SHA256": binding["python_runtime_tree_sha256"],
        "PYTHONDONTWRITEBYTECODE": "1",
        "PYTHONPYCACHEPREFIX": "/dev/null",
        "MAGI_ROOT": binding["WorkingDirectory"],
        "MAGI_ROOT_DIR": binding["WorkingDirectory"],
        "MAGI_PUBLIC_SOURCE_ROOT_DIR": binding["WorkingDirectory"],
        "OSC_CONFIG_PATH": binding["laf_config_file"],
        "MAGI_ORCH_DIR": str(Path(binding["WorkingDirectory"]) / "casper_ecosystem" / "law_firm_orchestrators"),
        "MAGI_CODE_DIR": str(Path(binding["WorkingDirectory"]) / "casper_ecosystem" / "law_firm_orchestrators"),
        "MAGI_JSON_DIR": str(Path(binding["laf_config_file"]).parent),
        "MAGI_SKILL_PYTHON": binding["python_runtime"],
        "MAGI_RUNTIME_DIR": str(shared_root / "runtime"),
        "MAGI_AGENT_DIR": str(shared_root / "agent"),
        "MAGI_DATA_DIR": str(shared_root / "agent"),
        "MAGI_EXPORTS_DIR": str(shared_root / "exports"),
        "MAGI_GCAL_DUP_AUDIT_OUTPUT_DIR": str(shared_root / "exports" / "gcal_dedup"),
        "MAGI_METRICS_DIR": str(shared_root / "metrics"),
        "MAGI_AUTOPILOT_RUNS_DIR": str(shared_root / "autopilot-runs"),
        "MAGI_MUTABLE_STATIC_DIR": str(shared_root / "static"),
        "MAGI_LOG_DIR": str(log_dir),
        "MAGI_FILE_REVIEW_STATE_DIR": str(shared_root / "file-review"),
        "MAGI_FILE_REVIEW_BG_JOB_DIR": str(shared_root / "file-review" / "bg-jobs"),
        "MAGI_EEFILE_DOWNLOAD_FOLDER": str(shared_root / "file-review" / "downloads"),
        "MAGI_TRAINING_LOCK_PATH": str(
            Path.home() / "Library" / "Application Support" / "MAGI" / "training.lock"
        ),
        "MAGI_BACKGROUND_LOCK_DIR": str(shared_root / "runtime" / "locks"),
        "MAGI_LAF_GMAIL_STATE_PATH": str(shared_root / "static" / "laf_gmail_monitor_state.json"),
        "MAGI_LAF_GMAIL_MONITOR_STATE": str(shared_root / "static" / "laf_gmail_monitor_state.json"),
        "MAGI_LAF_GMAIL_PENDING_PATH": str(shared_root / "runtime" / "laf_gmail_dispatch_pending.json"),
        "MAGI_FILE_REVIEW_EMAIL_MONITOR_STATE": str(
            shared_root / "static" / "file_review_email_monitor_state.json"
        ),
        "MAGI_FILE_REVIEW_PENDING_PATH": str(
            shared_root / "agent" / "file-review" / "review_submit_pending.json"
        ),
        "MAGI_BRAIN_SQLITE_PATH": str(shared_root / "agent" / "magi_brain.db"),
        "MAGI_CLOUDFLARED_LOG_PATH": str(log_dir / "cloudflared.log"),
        "MAGI_DAEMON_LOG_PATH": str(shared_root / "agent" / "daemon.log"),
        "MAGI_PENDING_ENV_UPDATE_FILE": str(shared_root / "runtime" / "pending-config" / "env_updates.json"),
        "MAGI_OSC_FILE_SHARE_STORE": str(shared_root / "runtime" / "osc_file_shares.json"),
        "MAGI_OSC_FILE_SHARE_PUBLIC_BASE_FILE": str(
            shared_root / "runtime" / "osc_share_public_base_url.txt"
        ),
        "MAGI_OSC_PREVIEW_CACHE_DIR": str(shared_root / "runtime" / "cache" / "paperclip-preview"),
        "MAGI_OSC_PREVIEW_CACHE_MAX_BYTES": str(1024 * 1024 * 1024),
        "MAGI_OSC_UPLOAD_CACHE_DIR": str(shared_root / "runtime" / "cache" / "paperclip-uploads"),
        "MAGI_WORLDMONITOR_REPORT_DIR": str(shared_root / "static" / "worldmonitor_reports"),
        "MAGI_AGENT_STATUS_PUBLIC_PATH": str(
            shared_root / "static" / "agent_status_public_latest.json"
        ),
        "MAGI_SAAS_AUDIT_PATH": str(shared_root / "runtime" / "saas_audit_events.jsonl"),
        "MAGI_RATE_LIMIT_DB_PATH": str(shared_root / "runtime" / "rate_limit.sqlite3"),
        "MAGI_GIBBERISH_LOG": str(shared_root / "static" / "gibberish_samples.jsonl"),
        "MAGI_TAIWAN_LEGAL_MCP_ROOT": str(shared_root / "runtime" / "mcp-taiwan-legal-db"),
        "MAGI_TAIWAN_LEGAL_MCP_CACHE": str(
            shared_root / "runtime" / "taiwan_legal_mcp" / "cache.sqlite3"
        ),
        "MAGI_PDF_NAMER_STATE_DIR": str(shared_root / "pdf-namer"),
        "MAGI_SKILL_OVERLAY_DIR": str(shared_root / "skill-overlays"),
        "MAGI_SKILL_RUNTIME_SITE_PACKAGES": str(
            shared_root / "skill-overlays" / ".runtime-site-packages"
        ),
        "MAGI_SKILL_EVENTS_FILE": str(
            shared_root / "skill-overlays" / ".logs" / "skill_runtime_events.jsonl"
        ),
        "MAGI_SKILL_USAGE_TRACKER_FILE": str(
            shared_root / "skill-overlays" / ".logs" / "skill_usage_events.jsonl"
        ),
        "MAGI_SKILL_INTERVIEW_HISTORY_FILE": str(
            shared_root / "skill-overlays" / ".logs" / "skill_interview_history.jsonl"
        ),
        "MAGI_IRON_DOME_STATE_DIR": str(shared_root / "skill-overlays" / ".iron-dome"),
        "MAGI_IRON_DOME_DYNAMIC_RULES_PATH": str(
            shared_root / "skill-overlays" / ".iron-dome" / "dynamic_rules.json"
        ),
        "MAGI_IRON_DOME_PATTERNS_CACHE_FILE": str(
            shared_root / "skill-overlays" / ".iron-dome" / "patterns_cache.json"
        ),
        "MAGI_IRON_DOME_UPSTREAM_STATE_FILE": str(
            shared_root / "skill-overlays" / ".iron-dome" / "upstream_last.json"
        ),
        "MAGI_AUTORESEARCH_RUNS_DIR": str(shared_root / "autoresearch-runs"),
        "JUDICIAL_CACHE_DIR": str(shared_root / "runtime" / "cache" / "judicial_web_search"),
        "MAGI_LAW_CACHE_DIR": str(shared_root / "runtime" / "cache" / "laws"),
        "MAGI_LAW_VDB_STATE_PATH": str(shared_root / "agent" / "_statutes_vdb_state.json"),
        "FAISS_INDEX_DIR": str(shared_root / "memory" / "index_cache"),
        "MAGI_USE_RUNTIME_DIR": "1",
        "MAGI_CRON_DEFINITIONS_IMMUTABLE": "1",
    }
    environment.update(
        {
            env_name: binding[binding_name]
            for env_name, (binding_name, _relative) in NAMED_MUTABLE_STATE_BINDINGS.items()
        }
    )
    if binding.get("chromedriver_path") is not None:
        environment.update(
            {
                "MAGI_CHROMEDRIVER_PATH": binding["chromedriver_path"],
                "MAGI_CHROMEDRIVER_SHA256": binding["chromedriver_sha256"],
            }
        )
    if binding.get("static_external_receipt") is not None:
        environment.update(
            {
                "MAGI_V3_STATIC_EXTERNAL_RECEIPT": binding["static_external_receipt"],
                "MAGI_V3_STATIC_EXTERNAL_RECEIPT_SHA256": binding[
                    "static_external_receipt_sha256"
                ],
                "MAGI_V3_STATIC_EXTERNAL_TARGET_SNAPSHOT_SHA256": binding[
                    "static_external_target_snapshot_sha256"
                ],
            }
        )
    if not validation_mode:
        environment.update(
            {
                "MAGI_V3_CASE_ROOT": binding["case_root"],
                "MAGI_V3_ARCHIVE_ROOT": binding["archive_root"],
                "MAGI_V3_PATH_MAPPINGS_JSON": json.dumps(
                    binding["path_mappings"],
                    ensure_ascii=False,
                    separators=(",", ":"),
                ),
            }
        )
    payload = {
        "Label": binding["label"],
        "ProgramArguments": binding["ProgramArguments"],
        "WorkingDirectory": binding["WorkingDirectory"],
        "EnvironmentVariables": environment,
        "StandardOutPath": str(log_dir / "stdout.log"),
        "StandardErrorPath": str(log_dir / "stderr.log"),
        "RunAtLoad": True,
        "KeepAlive": {"SuccessfulExit": False},
        "ThrottleInterval": 10,
        "ProcessType": PROCESS_TYPE_BY_ROLE[binding["role"]],
    }
    return plistlib.dumps(payload, fmt=plistlib.FMT_XML, sort_keys=True)


_PRODUCTION_RELEASE_ENV_SUFFIXES = {
    "MAGI_V3_RELEASE_MANIFEST": Path(RELEASE_MANIFEST_NAME),
    "MAGI_ROOT": Path("."),
    "MAGI_ROOT_DIR": Path("."),
    "MAGI_PUBLIC_SOURCE_ROOT_DIR": Path("."),
    "MAGI_ORCH_DIR": Path("casper_ecosystem/law_firm_orchestrators"),
    "MAGI_CODE_DIR": Path("casper_ecosystem/law_firm_orchestrators"),
}


def _validate_no_candidate_release_reference(
    value: Any,
    *,
    candidate_root: Path,
    working_directory: Path,
    location: str,
) -> None:
    """Recursively reject candidate-root paths hidden in launch metadata."""

    if isinstance(value, Mapping):
        for key, child in value.items():
            _validate_no_candidate_release_reference(
                key,
                candidate_root=candidate_root,
                working_directory=working_directory,
                location=f"{location}.<key>",
            )
            _validate_no_candidate_release_reference(
                child,
                candidate_root=candidate_root,
                working_directory=working_directory,
                location=f"{location}.{key}",
            )
        return
    if isinstance(value, (list, tuple)):
        for index, child in enumerate(value):
            _validate_no_candidate_release_reference(
                child,
                candidate_root=candidate_root,
                working_directory=working_directory,
                location=f"{location}[{index}]",
            )
        return
    if not isinstance(value, str) or not value:
        return
    candidate_text = str(candidate_root)
    references_candidate = candidate_text in value
    try:
        possible_path = Path(value).expanduser()
        if possible_path.is_absolute():
            resolved = possible_path.resolve(strict=False)
            references_candidate = references_candidate or (
                resolved == candidate_root or _inside(resolved, candidate_root)
            )
        elif ".." in possible_path.parts:
            resolved = (working_directory / possible_path).resolve(strict=False)
            references_candidate = references_candidate or (
                resolved == candidate_root or _inside(resolved, candidate_root)
            )
    except (OSError, RuntimeError, ValueError):
        # Non-path strings remain legal; exact candidate text was checked above.
        pass
    if references_candidate:
        raise DeployPrepareBlocked(
            f"production launchagent metadata retained candidate release path: {location}"
        )


def _validate_production_rendered_binding(
    binding: Mapping[str, Any],
    plist_bytes: bytes,
    *,
    installed_root: Path,
    candidate_root: Path,
    installed_executable: Path,
    installed_service_manifest: Path,
) -> None:
    """Reject a production plist if any immutable binding retained staging."""

    try:
        payload = plistlib.loads(plist_bytes)
        arguments = payload["ProgramArguments"]
        working_directory = Path(payload["WorkingDirectory"])
        environment = payload["EnvironmentVariables"]
    except (KeyError, TypeError, ValueError, plistlib.InvalidFileException) as exc:
        raise DeployPrepareBlocked("rendered production launchagent is invalid") from exc
    if (
        not isinstance(arguments, list)
        or not arguments
        or Path(arguments[0]) != installed_executable
    ):
        raise DeployPrepareBlocked(
            "production launchagent executable is not bound to the installed release"
        )
    if working_directory != installed_root:
        raise DeployPrepareBlocked(
            "production launchagent working directory is not the installed release"
        )
    for name, suffix in _PRODUCTION_RELEASE_ENV_SUFFIXES.items():
        expected = installed_root if suffix == Path(".") else installed_root / suffix
        if environment.get(name) != str(expected):
            raise DeployPrepareBlocked(
                f"production launchagent {name} is not bound to the installed release"
            )
    if environment.get("MAGI_V3_EXECUTABLE_PATH") != str(installed_executable):
        raise DeployPrepareBlocked(
            "production launchagent executable environment is not bound to the installed release"
        )
    if environment.get("MAGI_V3_SERVICE_MANIFEST") != str(installed_service_manifest):
        raise DeployPrepareBlocked(
            "production launchagent service manifest environment is not bound to the installed release"
        )
    if candidate_root != installed_root:
        _validate_no_candidate_release_reference(
            payload,
            candidate_root=candidate_root,
            working_directory=installed_root,
            location="plist",
        )
        _validate_no_candidate_release_reference(
            binding,
            candidate_root=candidate_root,
            working_directory=installed_root,
            location="binding",
        )
    if binding.get("release_manifest") != str(installed_root / RELEASE_MANIFEST_NAME):
        raise DeployPrepareBlocked(
            "production role metadata retained a non-installed release manifest"
        )
    if binding.get("service_manifest") != str(installed_service_manifest):
        raise DeployPrepareBlocked(
            "production role metadata retained a non-installed service manifest"
        )


def _encoded_json(payload: Mapping[str, Any]) -> bytes:
    return (json.dumps(payload, ensure_ascii=False, indent=2, sort_keys=True) + "\n").encode()


def _fsync_directory(path: Path) -> None:
    descriptor = os.open(path, os.O_RDONLY)
    try:
        os.fsync(descriptor)
    finally:
        os.close(descriptor)


def _write_atomic_exclusive(path: Path, data: bytes, *, mode: int = 0o644) -> dict[str, Any]:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = path.parent / f".{path.name}.tmp-{os.getpid()}-{uuid.uuid4().hex}"
    try:
        descriptor = os.open(temporary, os.O_WRONLY | os.O_CREAT | os.O_EXCL, mode)
        try:
            with os.fdopen(descriptor, "wb", closefd=False) as handle:
                handle.write(data)
                handle.flush()
                os.fsync(handle.fileno())
        finally:
            os.close(descriptor)
        os.link(temporary, path)
        temporary.unlink()
        _fsync_directory(path.parent)
    finally:
        if temporary.exists():
            temporary.unlink()
    persisted = path.read_bytes()
    if persisted != data:
        raise DeployPrepareBlocked(f"atomic write verification failed: {path.name}")
    return {"sha256": _sha256(persisted), "size": len(persisted)}


def _validate_shapely_sealed_runtime(python_runtime: Path) -> None:
    """Probe Shapely in the exact bound runtime, never the prepare process."""
    probe = (
        "import json,shapely; from pathlib import Path; "
        "from shapely import constrained_delaunay_triangles as c; "
        "print(json.dumps({'version':shapely.__version__,'origin':str(Path(shapely.__file__).resolve()),'api':callable(c)}))"
    )
    try:
        completed = subprocess.run(
            [
                str(python_runtime),
                "-B",
                "-X",
                "pycache_prefix=/dev/null",
                "-c",
                probe,
            ],
            text=True,
            capture_output=True,
            timeout=8,
        )
        payload = json.loads(completed.stdout or "{}") if completed.returncode == 0 else {}
        version = tuple(int(x) for x in str(payload.get("version", "")).split(".")[:2])
        origin = Path(str(payload.get("origin", ""))).resolve()
        root = python_runtime.parent.parent.resolve()
    except Exception as exc:
        raise DeployPrepareBlocked("sealed Shapely runtime probe failed") from exc
    if not ((2, 0) <= version < (3, 0)) or not payload.get("api") or "site-packages" not in origin.parts or root not in origin.parents:
        raise DeployPrepareBlocked("sealed Shapely runtime contract failed")


def prepare_deployment(
    release_root: Path,
    staging_dir: Path,
    runtime_root: Path,
    python_executable: Path,
    *,
    roles_config: Path | None = None,
    env_file: Path | None = None,
    cron_jobs_file: Path | None = None,
    website_root: Path | None = None,
    python_runtime: Path | None = None,
    laf_config_file: Path | None = None,
    google_credentials_file: Path | None = None,
    google_calendar_token_file: Path | None = None,
    laf_gmail_token_file: Path | None = None,
    file_review_token_file: Path | None = None,
    gmail_compose_token_file: Path | None = None,
    accounting_credentials_file: Path | None = None,
    accounting_sheets_token_file: Path | None = None,
    drive_sync_token_file: Path | None = None,
    drive_sync_write_token_file: Path | None = None,
    nas_ocr_queue_db_path: Path | None = None,
    chromedriver_path: Path | None = None,
    static_external_receipt: Path | None = None,
    deployment_mode: str = "production",
    service_manifest: Path | None = None,
    validation_input_root: Path | None = None,
    case_root: Path | None = None,
    archive_root: Path | None = None,
    path_mappings: tuple[tuple[str, str], ...] | None = None,
    publish_dir: Path | None = None,
    installed_release_root: Path | None = None,
    now: datetime | None = None,
) -> dict[str, Any]:
    """Generate a hashed, non-installed deployment and optionally publish it atomically.

    When ``publish_dir`` is supplied, files are written and verified below
    ``staging_dir`` but every persisted absolute binding names ``publish_dir``.
    The completed staging tree is then atomically renamed to the publish path.
    This prevents a caller-side staging rename from invalidating launchd input
    paths after their hashes have already been sealed.  Production launchagents
    are always rebound to the separately validated canonical installed release;
    ``release_root`` remains only the candidate identity being prepared.
    """

    if deployment_mode not in DEPLOYMENT_MODES:
        raise DeployPrepareBlocked("deployment mode must be production or isolated_live_validation")
    if deployment_mode == "production" and validation_input_root is not None:
        raise DeployPrepareBlocked("production deployment cannot accept a validation input root")
    if deployment_mode != "production" and installed_release_root is not None:
        raise DeployPrepareBlocked(
            "installed release root is only valid for production deployment"
        )
    root, identity = _validate_release(release_root)
    config_path = roles_config or root / "config" / "v3_launchagent_roles.json"
    config_path = _validate_release_data_file(
        config_path,
        root,
        identity,
        description="launchagent role config",
    )
    roles = load_roles(config_path)
    (
        service_manifest_path,
        service_manifest_sha256,
        service_process_entrypoints,
    ) = _validate_service_manifest(
        service_manifest or root / SERVICE_MANIFEST_NAMES[deployment_mode],
        root,
        identity,
        deployment_mode=deployment_mode,
    )
    executable = _validate_executable(python_executable, root, identity)
    executable_sha256 = _sha256_file(executable)
    binding_root = root
    binding_identity = identity
    binding_service_manifest = service_manifest_path
    binding_executable = executable
    if deployment_mode == "production":
        binding_root, binding_identity = _validate_production_release_root(
            root,
            identity,
            installed_release_root,
        )
        binding_executable = _installed_release_member(
            executable,
            source_root=root,
            installed_root=binding_root,
            installed_identity=binding_identity,
            description="executable",
        )
        binding_service_manifest = _installed_release_member(
            service_manifest_path,
            source_root=root,
            installed_root=binding_root,
            installed_identity=binding_identity,
            description="service manifest",
        )
    if deployment_mode == "isolated_live_validation":
        accounting_credentials_file = accounting_credentials_file or google_credentials_file
        accounting_sheets_token_file = (
            accounting_sheets_token_file or google_calendar_token_file
        )
        drive_sync_token_file = drive_sync_token_file or accounting_sheets_token_file
        drive_sync_write_token_file = (
            drive_sync_write_token_file or accounting_sheets_token_file
        )
    external_inputs = _validate_external_runtime_inputs(
        env_file,
        cron_jobs_file,
        website_root,
        python_runtime,
        laf_config_file,
        google_credentials_file,
        google_calendar_token_file,
        laf_gmail_token_file,
        file_review_token_file,
        gmail_compose_token_file,
        accounting_credentials_file,
        accounting_sheets_token_file,
        drive_sync_token_file,
        drive_sync_write_token_file,
        nas_ocr_queue_db_path,
        chromedriver_path,
    )
    if deployment_mode == "production":
        _validate_shapely_sealed_runtime(external_inputs.python_runtime)
    if deployment_mode == "production":
        _validate_production_cron_policy_binding(root, identity, external_inputs)
        _validate_production_external_independence(external_inputs)
    active_case_root, closed_case_root, canonical_path_mappings = (
        _validate_case_filesystem_bindings(
            deployment_mode,
            case_root,
            archive_root,
            path_mappings,
        )
    )
    staging, publish, runtime = _safe_paths(
        root,
        staging_dir,
        runtime_root,
        publish_dir,
    )
    static_external: StaticExternalEvidence | None = None
    if deployment_mode == "production":
        static_external = _validate_production_static_external(
            static_external_receipt,
            identity=identity,
            runtime_root=runtime,
            deployment_root=publish,
            external_inputs=external_inputs,
        )
    validated_validation_root: Path | None = None
    if deployment_mode == "isolated_live_validation":
        validated_validation_root = _validate_isolated_validation_inputs(
            validation_input_root,
            external_inputs,
            release_root=root,
            staging=staging,
            publish=publish,
            runtime=runtime,
        )
    cron_relative = Path("runtime-inputs") / "cron_jobs.v3.json"
    cron_artifact = staging / cron_relative
    cron_binding = publish / cron_relative
    try:
        cron_snapshot_bytes, cron_snapshot_evidence = render_snapshot(
            source=external_inputs.cron_jobs_source_file,
            # Production launch agents execute from the canonical immutable
            # installed release.  Cron commands must use that same binding;
            # otherwise a running worker remains coupled to the caller-owned
            # evidence/staging tree even after the release is installed.
            release_root=binding_root if deployment_mode == "production" else root,
            runtime_root=runtime,
            python_runtime=external_inputs.python_runtime,
        )
    except CronSnapshotBlocked as exc:
        raise DeployPrepareBlocked(f"release-bound cron snapshot is blocked: {exc}") from exc
    if cron_snapshot_evidence.get("source_sha256") != external_inputs.cron_jobs_source_sha256:
        raise DeployPrepareBlocked("release-bound cron snapshot source binding drifted")
    external_inputs = replace(
        external_inputs,
        cron_jobs_file=cron_binding,
        cron_jobs_sha256=str(cron_snapshot_evidence["snapshot_sha256"]),
        cron_snapshot_evidence=dict(cron_snapshot_evidence),
    )
    python_runtime_relative = Path("runtime-inputs") / "python-runtime-manifest.json"
    python_runtime_manifest = staging / python_runtime_relative
    python_runtime_manifest_binding = publish / python_runtime_relative
    try:
        python_runtime_manifest_bytes, python_runtime_evidence = build_runtime_manifest(
            external_inputs.python_runtime
        )
    except PythonRuntimeBlocked as exc:
        raise DeployPrepareBlocked(f"external Python runtime snapshot is blocked: {exc}") from exc
    try:
        python_runtime_manifest_payload = json.loads(
            python_runtime_manifest_bytes.decode("utf-8")
        )
    except (UnicodeError, json.JSONDecodeError) as exc:
        raise DeployPrepareBlocked(
            "external Python runtime snapshot returned invalid JSON"
        ) from exc
    if not isinstance(python_runtime_manifest_payload, dict):
        raise DeployPrepareBlocked(
            "external Python runtime snapshot must contain a JSON object"
        )
    python_runtime_module_preflight = _probe_python_docx_runtime(
        external_inputs.python_runtime,
        python_runtime_manifest_payload,
    )
    bindings: list[dict[str, Any]] = []
    entrypoints: list[Path] = []
    for role in roles:
        entrypoints.append(_entrypoint_path(root, role.entrypoint_module, identity))
        bindings.append(
            _role_binding(
                role,
                release_root=binding_root,
                executable=binding_executable,
                runtime_root=runtime,
                identity=binding_identity,
                executable_sha256=executable_sha256,
                external_inputs=external_inputs,
                python_runtime_manifest=python_runtime_manifest_binding,
                python_runtime_manifest_sha256=str(python_runtime_evidence["manifest_sha256"]),
                python_runtime_tree_sha256=str(python_runtime_evidence["tree_sha256"]),
                deployment_mode=deployment_mode,
                service_manifest=binding_service_manifest,
                service_manifest_sha256=service_manifest_sha256,
                case_root=active_case_root,
                archive_root=closed_case_root,
                path_mappings=canonical_path_mappings,
                static_external=static_external,
            )
        )

    generated_at = (now or datetime.now(timezone.utc)).astimezone(timezone.utc).isoformat()
    ownership = {
        "schema_version": 1,
        "status": "prepared_not_installed",
        "release_id": identity.release_id,
        "release_manifest": str(binding_identity.manifest_path),
        "release_manifest_sha256": identity.manifest_sha256,
        "deployment_mode": deployment_mode,
        "service_manifest": str(binding_service_manifest),
        "service_manifest_sha256": service_manifest_sha256,
        "validation_input_root": (
            str(validated_validation_root) if validated_validation_root is not None else None
        ),
        "runtime_root": str(runtime),
        "static_external_receipt": (
            str(static_external.receipt) if static_external is not None else None
        ),
        "static_external_receipt_sha256": (
            static_external.receipt_sha256 if static_external is not None else None
        ),
        "static_external_source_snapshot_sha256": (
            static_external.source_snapshot_sha256 if static_external is not None else None
        ),
        "static_external_target_snapshot_sha256": (
            static_external.target_snapshot_sha256 if static_external is not None else None
        ),
        "external_inputs": {
            "env_file": str(external_inputs.env_file),
            "env_file_sha256": external_inputs.env_file_sha256,
            "cron_jobs_file": str(external_inputs.cron_jobs_file),
            "cron_jobs_sha256": external_inputs.cron_jobs_sha256,
            "cron_jobs_source_file": str(external_inputs.cron_jobs_source_file),
            "cron_jobs_source_sha256": external_inputs.cron_jobs_source_sha256,
            "cron_snapshot_evidence": dict(external_inputs.cron_snapshot_evidence),
            "website_root": str(external_inputs.website_root),
            "website_admin_sha256": external_inputs.website_admin_sha256,
            "laf_config_file": str(external_inputs.laf_config_file),
            "laf_config_sha256": external_inputs.laf_config_sha256,
            "laf_config_mode": external_inputs.laf_config_mode,
            "google_credentials_file": str(external_inputs.google_credentials_file),
            "google_credentials_sha256": external_inputs.google_credentials_sha256,
            "google_credentials_mode": external_inputs.google_credentials_mode,
            "google_calendar_token_source_file": str(external_inputs.google_calendar_token_source_file),
            "google_calendar_token_source_sha256": external_inputs.google_calendar_token_source_sha256,
            "google_calendar_token_file": bindings[0]["google_calendar_token_file"],
            "laf_gmail_token_source_file": str(external_inputs.laf_gmail_token_source_file),
            "laf_gmail_token_source_sha256": external_inputs.laf_gmail_token_source_sha256,
            "laf_gmail_token_file": bindings[0]["laf_gmail_token_file"],
            "file_review_token_source_file": str(external_inputs.file_review_token_source_file),
            "file_review_token_source_sha256": external_inputs.file_review_token_source_sha256,
            "file_review_token_file": bindings[0]["file_review_token_file"],
            "gmail_compose_token_source_file": (
                str(external_inputs.gmail_compose_token_source_file)
                if external_inputs.gmail_compose_token_source_file is not None
                else None
            ),
            "gmail_compose_token_source_sha256": external_inputs.gmail_compose_token_source_sha256,
            "gmail_compose_token_file": bindings[0]["gmail_compose_token_file"],
            "optional_degraded_inputs": bindings[0]["optional_degraded_inputs"],
            "accounting_credentials_file": str(external_inputs.accounting_credentials_file),
            "accounting_credentials_sha256": external_inputs.accounting_credentials_sha256,
            "accounting_credentials_mode": external_inputs.accounting_credentials_mode,
            "accounting_sheets_token_source_file": str(
                external_inputs.accounting_sheets_token_source_file
            ),
            "accounting_sheets_token_source_sha256": (
                external_inputs.accounting_sheets_token_source_sha256
            ),
            "accounting_sheets_token_file": bindings[0]["accounting_sheets_token_file"],
            "drive_sync_token_source_file": str(external_inputs.drive_sync_token_source_file),
            "drive_sync_token_source_sha256": external_inputs.drive_sync_token_source_sha256,
            "drive_sync_token_file": bindings[0]["drive_sync_token_file"],
            "drive_sync_write_token_source_file": str(
                external_inputs.drive_sync_write_token_source_file
            ),
            "drive_sync_write_token_source_sha256": (
                external_inputs.drive_sync_write_token_source_sha256
            ),
            "drive_sync_write_token_file": bindings[0]["drive_sync_write_token_file"],
            "nas_ocr_queue_db_file": str(external_inputs.nas_ocr_queue_db_file),
            "nas_ocr_queue_db_mode": external_inputs.nas_ocr_queue_db_mode,
            "python_runtime": str(external_inputs.python_runtime),
            "python_runtime_realpath": str(external_inputs.python_runtime_realpath),
            "python_runtime_sha256": external_inputs.python_runtime_sha256,
            "chromedriver_path": (
                str(external_inputs.chromedriver_path)
                if external_inputs.chromedriver_path is not None
                else None
            ),
            "chromedriver_sha256": external_inputs.chromedriver_sha256,
            "chromedriver_mode": external_inputs.chromedriver_mode,
            "python_runtime_manifest": str(python_runtime_manifest_binding),
            "python_runtime_manifest_sha256": python_runtime_evidence["manifest_sha256"],
            "python_runtime_tree_sha256": python_runtime_evidence["tree_sha256"],
            "python_runtime_file_count": python_runtime_evidence["file_count"],
            "python_runtime_directory_count": python_runtime_evidence["directory_count"],
            "python_runtime_module_preflight": python_runtime_module_preflight,
            "case_root": str(active_case_root) if active_case_root is not None else None,
            "archive_root": str(closed_case_root) if closed_case_root is not None else None,
            "path_mappings": [list(row) for row in canonical_path_mappings],
            "static_external_receipt": (
                str(static_external.receipt) if static_external is not None else None
            ),
            "static_external_receipt_sha256": (
                static_external.receipt_sha256 if static_external is not None else None
            ),
            "static_external_source_snapshot_sha256": (
                static_external.source_snapshot_sha256 if static_external is not None else None
            ),
            "static_external_target_snapshot_sha256": (
                static_external.target_snapshot_sha256 if static_external is not None else None
            ),
            **named_mutable_state_paths(runtime),
        },
        # Keep the installed ownership document independent of the plist-only
        # hash binding.  Embedding its own digest would create a circular hash.
        "roles": [dict(binding) for binding in bindings],
        "uninstall_labels": [binding["label"] for binding in bindings],
    }
    ownership_bytes = _encoded_json(ownership)
    ownership_sha256 = _sha256(ownership_bytes)
    for binding in bindings:
        binding["ownership_manifest_sha256"] = ownership_sha256
    rendered_plists: list[tuple[str, bytes]] = []
    for binding in bindings:
        relative = f"launchagents/{binding['label']}.plist"
        plist_bytes = _plist(binding)
        if deployment_mode == "production":
            _validate_production_rendered_binding(
                binding,
                plist_bytes,
                installed_root=binding_root,
                candidate_root=root,
                installed_executable=binding_executable,
                installed_service_manifest=binding_service_manifest,
            )
        rendered_plists.append((relative, plist_bytes))
    staging.mkdir(mode=0o700)
    artifacts: list[dict[str, Any]] = []
    cron_relative = cron_relative.as_posix()
    cron_info = _write_atomic_exclusive(
        staging / cron_relative,
        cron_snapshot_bytes,
        mode=0o600,
    )
    artifacts.append({"path": cron_relative, **cron_info})
    python_runtime_relative = python_runtime_relative.as_posix()
    python_runtime_info = _write_atomic_exclusive(
        staging / python_runtime_relative,
        python_runtime_manifest_bytes,
        mode=0o600,
    )
    artifacts.append({"path": python_runtime_relative, **python_runtime_info})
    if static_external is not None:
        static_receipt_relative = (
            Path("runtime-inputs") / STATIC_EXTERNAL_RELEASE_RECEIPT_NAME
        ).as_posix()
        static_receipt_info = _write_atomic_exclusive(
            staging / static_receipt_relative,
            static_external.receipt_bytes,
            mode=0o600,
        )
        artifacts.append({"path": static_receipt_relative, **static_receipt_info})
    ownership_info = _write_atomic_exclusive(
        staging / OWNERSHIP_MANIFEST_NAME,
        ownership_bytes,
    )
    artifacts.append({"path": OWNERSHIP_MANIFEST_NAME, **ownership_info})
    for relative, plist_bytes in rendered_plists:
        info = _write_atomic_exclusive(staging / relative, plist_bytes)
        artifacts.append({"path": relative, **info})
    deploy_manifest = {
        "schema_version": 1,
        "status": "prepared_not_installed",
        "generated_at": generated_at,
        "mutation_performed": False,
        "release_id": identity.release_id,
        "release_manifest": str(binding_identity.manifest_path),
        "release_manifest_sha256": identity.manifest_sha256,
        "deployment_mode": deployment_mode,
        "service_manifest": str(binding_service_manifest),
        "service_manifest_sha256": service_manifest_sha256,
        "validation_input_root": (
            str(validated_validation_root) if validated_validation_root is not None else None
        ),
        "runtime_root": str(runtime),
        "static_external_receipt": ownership["static_external_receipt"],
        "static_external_receipt_sha256": ownership[
            "static_external_receipt_sha256"
        ],
        "static_external_source_snapshot_sha256": ownership[
            "static_external_source_snapshot_sha256"
        ],
        "static_external_target_snapshot_sha256": ownership[
            "static_external_target_snapshot_sha256"
        ],
        "ownership_manifest": str(runtime / OWNERSHIP_MANIFEST_NAME),
        "ownership_manifest_sha256": ownership_sha256,
        "external_inputs": ownership["external_inputs"],
        "roles": bindings,
        "artifacts": artifacts,
    }
    deploy_info = _write_atomic_exclusive(
        staging / DEPLOY_MANIFEST_NAME,
        _encoded_json(deploy_manifest),
    )
    if _sha256_file(identity.manifest_path) != identity.manifest_sha256:
        raise DeployPrepareBlocked("release manifest changed while rendering deployment")
    if deployment_mode == "production":
        final_binding_root, final_binding_identity = _validate_production_release_root(
            root,
            identity,
            binding_root,
        )
        if (
            final_binding_root != binding_root
            or final_binding_identity.manifest_sha256 != binding_identity.manifest_sha256
        ):
            raise DeployPrepareBlocked("installed release changed while rendering deployment")
    if _sha256_file(service_manifest_path) != service_manifest_sha256:
        raise DeployPrepareBlocked("service manifest changed while rendering deployment")
    for service_entrypoint in service_process_entrypoints:
        relative = service_entrypoint.relative_to(root).as_posix()
        if _sha256_file(service_entrypoint) != identity.file_sha256[relative]:
            raise DeployPrepareBlocked(
                f"service entrypoint changed while rendering deployment: {relative}"
            )
    if _sha256_file(executable) != executable_sha256:
        raise DeployPrepareBlocked("required executable changed while rendering deployment")
    if external_inputs.python_runtime.resolve(strict=True) != external_inputs.python_runtime_realpath:
        raise DeployPrepareBlocked("external Python runtime symlink target changed while rendering deployment")
    if _sha256_file(external_inputs.python_runtime_realpath) != external_inputs.python_runtime_sha256:
        raise DeployPrepareBlocked("external Python runtime changed while rendering deployment")
    if external_inputs.chromedriver_path is not None and (
        not external_inputs.chromedriver_path.is_file()
        or not os.access(external_inputs.chromedriver_path, os.X_OK)
        or _sha256_file(external_inputs.chromedriver_path)
        != external_inputs.chromedriver_sha256
        or f"{stat.S_IMODE(external_inputs.chromedriver_path.stat().st_mode):04o}"
        != external_inputs.chromedriver_mode
    ):
        raise DeployPrepareBlocked(
            "external ChromeDriver changed while rendering deployment"
        )
    if _sha256_file(python_runtime_manifest) != python_runtime_evidence["manifest_sha256"]:
        raise DeployPrepareBlocked("external Python runtime manifest changed while rendering deployment")
    if _sha256_file(external_inputs.cron_jobs_source_file) != external_inputs.cron_jobs_source_sha256:
        raise DeployPrepareBlocked("external cron jobs source changed while rendering deployment")
    if _sha256_file(external_inputs.laf_config_file) != external_inputs.laf_config_sha256:
        raise DeployPrepareBlocked("external LAF config changed while rendering deployment")
    if (
        static_external is not None
        and _sha256_file(static_external.payload_receipt)
        != static_external.payload_receipt_sha256
    ):
        raise DeployPrepareBlocked(
            "production static external payload receipt changed while rendering deployment"
        )
    if static_external is not None and _sha256_file(
        staging / "runtime-inputs" / STATIC_EXTERNAL_RELEASE_RECEIPT_NAME
    ) != static_external.receipt_sha256:
        raise DeployPrepareBlocked(
            "production static external release receipt changed while rendering deployment"
        )
    for name, path, expected in (
        (
            "Google credentials",
            external_inputs.google_credentials_file,
            external_inputs.google_credentials_sha256,
        ),
        (
            "Google Calendar token",
            external_inputs.google_calendar_token_source_file,
            external_inputs.google_calendar_token_source_sha256,
        ),
        (
            "LAF Gmail token",
            external_inputs.laf_gmail_token_source_file,
            external_inputs.laf_gmail_token_source_sha256,
        ),
        (
            "FileReview token",
            external_inputs.file_review_token_source_file,
            external_inputs.file_review_token_source_sha256,
        ),
        (
            "accounting credentials",
            external_inputs.accounting_credentials_file,
            external_inputs.accounting_credentials_sha256,
        ),
        (
            "accounting Sheets token",
            external_inputs.accounting_sheets_token_source_file,
            external_inputs.accounting_sheets_token_source_sha256,
        ),
        (
            "Drive sync token",
            external_inputs.drive_sync_token_source_file,
            external_inputs.drive_sync_token_source_sha256,
        ),
        (
            "Drive sync write token",
            external_inputs.drive_sync_write_token_source_file,
            external_inputs.drive_sync_write_token_source_sha256,
        ),
    ):
        if _sha256_file(path) != expected:
            raise DeployPrepareBlocked(f"external {name} changed while rendering deployment")
    if _sha256_file(cron_artifact) != external_inputs.cron_jobs_sha256:
        raise DeployPrepareBlocked("release-bound cron snapshot changed while rendering deployment")
    for description, path in (
        ("active case root", active_case_root),
        ("closed case root", closed_case_root),
    ):
        if path is not None and (path.is_symlink() or path.resolve(strict=True) != path):
            raise DeployPrepareBlocked(f"{description} changed while rendering deployment")
    for entrypoint in entrypoints:
        relative = entrypoint.relative_to(root).as_posix()
        if _sha256_file(entrypoint) != identity.file_sha256[relative]:
            raise DeployPrepareBlocked(
                f"required entrypoint changed while rendering deployment: {relative}"
            )
    try:
        final_python_runtime_manifest_bytes, final_python_runtime_evidence = (
            build_runtime_manifest(external_inputs.python_runtime)
        )
    except PythonRuntimeBlocked as exc:
        raise DeployPrepareBlocked(
            f"external Python runtime final snapshot is blocked: {exc}"
        ) from exc
    if (
        final_python_runtime_manifest_bytes != python_runtime_manifest_bytes
        or final_python_runtime_evidence.get("manifest_sha256")
        != python_runtime_evidence["manifest_sha256"]
        or final_python_runtime_evidence.get("tree_sha256")
        != python_runtime_evidence["tree_sha256"]
    ):
        raise DeployPrepareBlocked(
            "external Python runtime tree changed after python-docx preflight"
        )
    marker = {
        "schema_version": 1,
        "status": "prepared_not_installed",
        "ready_to_install": True,
        "mutation_performed": False,
        "release_id": identity.release_id,
        "deployment_mode": deployment_mode,
        "release_manifest_sha256": identity.manifest_sha256,
        "ownership_manifest_sha256": ownership_sha256,
        "manifest": DEPLOY_MANIFEST_NAME,
        "manifest_sha256": deploy_info["sha256"],
    }
    _write_atomic_exclusive(staging / COMPLETION_MARKER_NAME, _encoded_json(marker))
    _fsync_directory(staging)
    if publish != staging:
        os.replace(staging, publish)
        _fsync_directory(publish.parent)
    return marker


def _blocked(exc: BaseException) -> dict[str, Any]:
    return {
        "schema_version": 1,
        "status": "blocked",
        "ready_to_install": False,
        "mutation_performed": False,
        "reason": str(exc),
    }


def main(argv: list[str] | None = None) -> int:
    root = Path(__file__).resolve().parents[1]
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--release-root", type=Path, required=True)
    parser.add_argument("--staging-dir", type=Path, required=True)
    parser.add_argument(
        "--publish-dir",
        type=Path,
        help="final sibling path atomically published after staging verification",
    )
    parser.add_argument("--runtime-root", type=Path, required=True)
    parser.add_argument("--python-executable", type=Path, required=True)
    parser.add_argument(
        "--installed-release-root",
        type=Path,
        help=(
            "canonical immutable Application Support release used by production "
            "launchagents"
        ),
    )
    parser.add_argument("--env-file", type=Path)
    parser.add_argument("--cron-jobs-file", type=Path)
    parser.add_argument("--website-root", type=Path)
    parser.add_argument("--python-runtime", type=Path)
    parser.add_argument("--laf-config-file", type=Path)
    parser.add_argument("--google-credentials-file", type=Path)
    parser.add_argument("--google-calendar-token-file", type=Path)
    parser.add_argument("--laf-gmail-token-file", type=Path)
    parser.add_argument("--file-review-token-file", type=Path)
    parser.add_argument("--gmail-compose-token-file", type=Path)
    parser.add_argument("--accounting-credentials-file", type=Path)
    parser.add_argument("--accounting-sheets-token-file", type=Path)
    parser.add_argument("--drive-sync-token-file", type=Path)
    parser.add_argument("--drive-sync-write-token-file", type=Path)
    parser.add_argument("--nas-ocr-queue-db-path", type=Path)
    parser.add_argument("--chromedriver-path", type=Path)
    parser.add_argument("--static-external-receipt", type=Path)
    parser.add_argument(
        "--deployment-mode",
        choices=sorted(DEPLOYMENT_MODES),
        default="production",
    )
    parser.add_argument("--service-manifest", type=Path)
    parser.add_argument("--validation-input-root", type=Path)
    parser.add_argument("--case-root", type=Path)
    parser.add_argument("--archive-root", type=Path)
    parser.add_argument(
        "--path-mapping",
        action="append",
        default=[],
        metavar="LOCAL=CANONICAL",
    )
    parser.add_argument(
        "--roles-config",
        type=Path,
    )
    args = parser.parse_args(argv)
    parsed_mappings: tuple[tuple[str, str], ...] | None = None
    if args.path_mapping:
        if any("=" not in row for row in args.path_mapping):
            print(
                json.dumps(
                    _blocked(
                        DeployPrepareBlocked("--path-mapping must use LOCAL=CANONICAL")
                    ),
                    ensure_ascii=False,
                    sort_keys=True,
                )
            )
            return 2
        parsed_mappings = tuple(tuple(row.split("=", 1)) for row in args.path_mapping)
    try:
        result = prepare_deployment(
            args.release_root,
            args.staging_dir,
            args.runtime_root,
            args.python_executable,
            roles_config=args.roles_config,
            env_file=args.env_file,
            cron_jobs_file=args.cron_jobs_file,
            website_root=args.website_root,
            python_runtime=args.python_runtime,
            laf_config_file=args.laf_config_file,
            google_credentials_file=args.google_credentials_file,
            google_calendar_token_file=args.google_calendar_token_file,
            laf_gmail_token_file=args.laf_gmail_token_file,
            file_review_token_file=args.file_review_token_file,
            gmail_compose_token_file=args.gmail_compose_token_file,
            accounting_credentials_file=args.accounting_credentials_file,
            accounting_sheets_token_file=args.accounting_sheets_token_file,
            drive_sync_token_file=args.drive_sync_token_file,
            drive_sync_write_token_file=args.drive_sync_write_token_file,
            nas_ocr_queue_db_path=args.nas_ocr_queue_db_path,
            chromedriver_path=args.chromedriver_path,
            static_external_receipt=args.static_external_receipt,
            deployment_mode=args.deployment_mode,
            service_manifest=args.service_manifest,
            validation_input_root=args.validation_input_root,
            case_root=args.case_root,
            archive_root=args.archive_root,
            path_mappings=parsed_mappings,
            publish_dir=args.publish_dir,
            installed_release_root=args.installed_release_root,
        )
    except (OSError, DeployPrepareBlocked) as exc:
        print(json.dumps(_blocked(exc), ensure_ascii=False, sort_keys=True))
        return 2
    print(json.dumps(result, ensure_ascii=False, sort_keys=True))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
