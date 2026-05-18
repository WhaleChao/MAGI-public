from __future__ import annotations

import tempfile
from pathlib import Path
from textwrap import wrap

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs/guides/MAGI_一般使用者圖文操作手冊_2026-05-19.docx"
FONT_PATH = Path("/System/Library/Fonts/STHeiti Medium.ttc")
FONT_LATIN = "Arial"
FONT_EAST_ASIA = "Microsoft JhengHei"

INK = "1A2738"
MUTED = "5D6B82"
BLUE = "0EA5E9"
GREEN = "16A34A"
AMBER = "D97706"
RED = "DC2626"
PURPLE = "7C3AED"
PANEL = "F8FAFC"
BORDER = "D8E2EF"


def font(size: int) -> ImageFont.FreeTypeFont:
    if FONT_PATH.exists():
        return ImageFont.truetype(str(FONT_PATH), size)
    return ImageFont.load_default()


def hex_to_rgb(value: str) -> tuple[int, int, int]:
    value = value.lstrip("#")
    return tuple(int(value[i : i + 2], 16) for i in (0, 2, 4))


def draw_round_rect(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int, int, int],
    fill: str,
    outline: str | None = None,
    width: int = 2,
    radius: int = 24,
) -> None:
    draw.rounded_rectangle(xy, radius=radius, fill="#" + fill, outline="#" + outline if outline else None, width=width)


def draw_text_box(
    draw: ImageDraw.ImageDraw,
    text: str,
    xy: tuple[int, int],
    max_chars: int,
    size: int = 34,
    fill: str = INK,
    line_gap: int = 10,
) -> None:
    x, y = xy
    for line in wrap(text, width=max_chars):
        draw.text((x, y), line, font=font(size), fill="#" + fill)
        y += size + line_gap


def save_image(path: Path, image: Image.Image) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    image.save(path, "PNG")
    return path


def make_cover(path: Path) -> Path:
    img = Image.new("RGB", (1600, 900), "#EEF6FF")
    draw = ImageDraw.Draw(img)
    for y in range(900):
        ratio = y / 899
        r = int(239 * (1 - ratio) + 248 * ratio)
        g = int(246 * (1 - ratio) + 250 * ratio)
        b = int(255 * (1 - ratio) + 252 * ratio)
        draw.line([(0, y), (1600, y)], fill=(r, g, b))

    draw_round_rect(draw, (80, 70, 1520, 830), fill="FFFFFF", outline=BORDER, radius=36)
    draw_round_rect(draw, (116, 112, 430, 204), fill="E0F2FE", outline="BAE6FD", radius=24)
    draw.text((148, 136), "MAGI", font=font(54), fill="#" + BLUE)
    draw.text((148, 200), "法律事務 AI 工作系統", font=font(28), fill="#" + MUTED)
    draw.text((118, 300), "一般使用者圖文操作手冊", font=font(70), fill="#" + INK)
    draw_text_box(
        draw,
        "從登入、案件、檔案、PDF、摘要、翻譯、逐字稿、法扶、閱卷、法律資料到系統健康檢查，一份給一般使用者真正看得懂的操作指南。",
        (122, 410),
        max_chars=30,
        size=34,
        fill=MUTED,
    )
    badges = [("公開版", BLUE), ("私有版", GREEN), ("商用部署", PURPLE), ("2026-05-19", AMBER)]
    x = 122
    for label, color in badges:
        draw_round_rect(draw, (x, 635, x + 210, 705), fill=color, outline=color, radius=18)
        draw.text((x + 38, 653), label, font=font(30), fill="#FFFFFF")
        x += 235

    # Safe mock workstation.
    draw_round_rect(draw, (1030, 190, 1438, 650), fill="111827", outline="334155", radius=28)
    draw_round_rect(draw, (1060, 230, 1408, 610), fill="F8FAFC", outline="CBD5E1", radius=18)
    for i, color in enumerate([BLUE, GREEN, AMBER]):
        draw.ellipse((1084 + i * 34, 250, 1106 + i * 34, 272), fill="#" + color)
    draw_round_rect(draw, (1090, 300, 1378, 350), fill="E0F2FE", outline="BAE6FD", radius=14)
    draw.text((1110, 311), "系統健康：正常", font=font(24), fill="#" + INK)
    rows = [("案件", "185"), ("待辦", "24"), ("文件", "3,420"), ("模型", "day")]
    y = 380
    for label, value in rows:
        draw_round_rect(draw, (1090, y, 1378, y + 42), fill="FFFFFF", outline=BORDER, radius=12)
        draw.text((1110, y + 8), label, font=font(22), fill="#" + MUTED)
        draw.text((1305, y + 8), value, font=font(22), fill="#" + INK)
        y += 58
    return save_image(path, img)


def make_module_map(path: Path) -> Path:
    img = Image.new("RGB", (1500, 860), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((70, 55), "MAGI 功能地圖", font=font(52), fill="#" + INK)
    draw.text((70, 120), "每個入口都回到同一套案件、檔案、待辦與知識資料，不重複建資料。", font=font(28), fill="#" + MUTED)
    draw_round_rect(draw, (585, 320, 915, 510), fill="E0F2FE", outline="7DD3FC", radius=34)
    draw.text((655, 365), "MAGI", font=font(58), fill="#" + BLUE)
    draw.text((626, 440), "案件與知識總控", font=font(28), fill="#" + INK)
    modules = [
        ("案件管理", "案件卡片、狀態、資料夾", 130, 245, BLUE),
        ("檔案與 OCR", "預覽、下載、命名、分享", 1040, 245, GREEN),
        ("摘要翻譯逐字稿", "品質閘門與 DOCX 交付", 120, 575, PURPLE),
        ("法扶閱卷筆錄", "開辦、回報、結案、去重", 1030, 575, AMBER),
        ("法律資料", "判決、法條、實務見解", 575, 95, "2563EB"),
        ("健康維運", "模型、DB、NAS、外網", 575, 650, RED),
    ]
    for title, desc, x, y, color in modules:
        draw.line((750, 415, x + 155, y + 72), fill="#" + BORDER, width=5)
        draw_round_rect(draw, (x, y, x + 320, y + 145), fill="F8FAFC", outline=color, radius=26)
        draw.text((x + 28, y + 24), title, font=font(34), fill="#" + color)
        draw_text_box(draw, desc, (x + 28, y + 76), max_chars=14, size=24, fill=MUTED, line_gap=4)
    return save_image(path, img)


def make_workflow(path: Path) -> Path:
    img = Image.new("RGB", (1500, 740), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((70, 55), "一天的使用流程", font=font(52), fill="#" + INK)
    steps = [
        ("1", "看健康狀態", "確認 DB、模型、OCR、NAS"),
        ("2", "查今日工作", "行事曆事件與 OSC 待辦分開看"),
        ("3", "處理案件", "開資料夾、看檔案、更新狀態"),
        ("4", "交給 MAGI", "摘要、翻譯、逐字稿、命名"),
        ("5", "人工確認", "正式送出、分享、結案前確認"),
        ("6", "留下紀錄", "待辦、行事曆、學習修正"),
    ]
    x = 70
    y = 210
    colors = [BLUE, GREEN, PURPLE, AMBER, RED, "2563EB"]
    for i, (num, title, desc) in enumerate(steps):
        draw_round_rect(draw, (x, y, x + 210, y + 260), fill="F8FAFC", outline=colors[i], radius=26)
        draw.ellipse((x + 22, y + 22, x + 78, y + 78), fill="#" + colors[i])
        draw.text((x + 42, y + 32), num, font=font(28), fill="#FFFFFF")
        draw_text_box(draw, title, (x + 22, y + 100), max_chars=8, size=30, fill=INK)
        draw_text_box(draw, desc, (x + 22, y + 155), max_chars=9, size=22, fill=MUTED, line_gap=5)
        if i < len(steps) - 1:
            draw.line((x + 220, y + 130, x + 242, y + 130), fill="#" + MUTED, width=4)
            draw.polygon([(x + 242, y + 130), (x + 228, y + 120), (x + 228, y + 140)], fill="#" + MUTED)
        x += 235
    return save_image(path, img)


def make_dashboard_mock(path: Path) -> Path:
    img = Image.new("RGB", (1500, 900), "#EFF4FA")
    draw = ImageDraw.Draw(img)
    draw_round_rect(draw, (60, 50, 1440, 850), fill="FFFFFF", outline=BORDER, radius=32)
    draw_round_rect(draw, (60, 50, 300, 850), fill="172033", outline="172033", radius=32)
    draw.text((95, 95), "Paperclip", font=font(38), fill="#" + BLUE)
    menu = ["業務概覽", "案件列表", "行事曆", "待辦事項", "書狀", "法扶", "實務見解", "系統設定"]
    y = 170
    for item in menu:
        fill = "0EA5E9" if item == "業務概覽" else "22304A"
        draw_round_rect(draw, (88, y, 270, y + 46), fill=fill, outline=fill, radius=10)
        draw.text((108, y + 9), item, font=font(22), fill="#FFFFFF")
        y += 62
    draw.text((340, 92), "業務概覽", font=font(46), fill="#" + INK)
    draw_round_rect(draw, (340, 152, 1020, 210), fill="F8FAFC", outline=BORDER, radius=14)
    draw.text((365, 167), "搜尋案件、當事人、法院案號...", font=font(24), fill="#" + MUTED)
    cards = [
        ("今日行程", "3 件", BLUE),
        ("OSC 待辦", "8 件", AMBER),
        ("檔案待處理", "5 份", PURPLE),
        ("系統健康", "正常", GREEN),
    ]
    x = 340
    for title, value, color in cards:
        draw_round_rect(draw, (x, 245, x + 235, 360), fill="FFFFFF", outline=BORDER, radius=18)
        draw.text((x + 24, 270), title, font=font(24), fill="#" + MUTED)
        draw.text((x + 24, 310), value, font=font(34), fill="#" + color)
        x += 260
    draw_round_rect(draw, (340, 410, 880, 765), fill="FFFFFF", outline=BORDER, radius=18)
    draw.text((370, 438), "OSC 建立待辦", font=font(28), fill="#" + INK)
    draw.text((370, 475), "由案件文件或系統規則建立", font=font(20), fill="#" + MUTED)
    draw_round_rect(draw, (915, 410, 1378, 765), fill="FFFFFF", outline=BORDER, radius=18)
    draw.text((945, 438), "行事曆事件", font=font(28), fill="#" + INK)
    draw.text((945, 475), "來自 Google Calendar 或匯入行程", font=font(20), fill="#" + MUTED)
    todo_rows = [
        ("2026-05-20", "補正", "2026-0001"),
        ("2026-05-21", "陳報", "2026-0002"),
        ("2026-05-22", "繳費確認", "2026-0003"),
    ]
    y = 525
    for date, kind, case_id in todo_rows:
        draw_round_rect(draw, (370, y, 850, y + 50), fill="F8FAFC", outline=BORDER, radius=10)
        draw.text((390, y + 12), date, font=font(20), fill="#" + MUTED)
        draw.text((540, y + 12), case_id, font=font(20), fill="#" + INK)
        draw.text((685, y + 12), kind, font=font(20), fill="#" + AMBER)
        y += 70
    cal_rows = [
        ("05/20 10:00", "開庭"),
        ("05/21 14:30", "與當事人會議"),
        ("05/22 09:00", "電話確認"),
    ]
    y = 525
    for date, kind in cal_rows:
        draw_round_rect(draw, (945, y, 1348, y + 50), fill="F8FAFC", outline=BORDER, radius=10)
        draw.text((965, y + 12), date, font=font(20), fill="#" + MUTED)
        draw.text((1135, y + 12), kind, font=font(20), fill="#" + GREEN)
        y += 70
    return save_image(path, img)


def make_quality_gate(path: Path) -> Path:
    img = Image.new("RGB", (1500, 720), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((70, 55), "摘要、翻譯、逐字稿品質閘門", font=font(48), fill="#" + INK)
    lanes = [
        ("輸入", "PDF、音訊、網址、案件資料", BLUE),
        ("取用工具", "OCR、搜尋、法律資料、模型", PURPLE),
        ("品質檢查", "不漏段、不亂譯、有來源", AMBER),
        ("交付", "文字、DOCX、表格、連結", GREEN),
    ]
    x = 90
    y = 225
    for i, (title, desc, color) in enumerate(lanes):
        draw_round_rect(draw, (x, y, x + 280, y + 250), fill="F8FAFC", outline=color, radius=28)
        draw.text((x + 34, y + 36), title, font=font(36), fill="#" + color)
        draw_text_box(draw, desc, (x + 34, y + 98), max_chars=12, size=25, fill=MUTED)
        if i < len(lanes) - 1:
            draw.line((x + 295, y + 125, x + 345, y + 125), fill="#" + MUTED, width=5)
            draw.polygon([(x + 345, y + 125), (x + 325, y + 112), (x + 325, y + 138)], fill="#" + MUTED)
        x += 360
    draw_round_rect(draw, (245, 545, 1255, 630), fill="FFF7ED", outline="FDBA74", radius=22)
    draw.text((282, 570), "品質不通過時：MAGI 應重跑、降級為抽取式摘要，或明確說明失敗原因，不交付殘缺內容。", font=font(27), fill="#" + INK)
    return save_image(path, img)


def make_health_cards(path: Path) -> Path:
    img = Image.new("RGB", (1500, 800), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((70, 55), "健康狀態怎麼看", font=font(52), fill="#" + INK)
    cards = [
        ("主狀態", "正常", GREEN),
        ("資料庫", "MariaDB", GREEN),
        ("推論服務", "day profile", GREEN),
        ("OCR", "macOS Vision", GREEN),
        ("向量資料庫", "索引正常", GREEN),
        ("NAS 掛載", "homes / lumi", GREEN),
        ("外網通道", "可連線", GREEN),
        ("日常稽核", "最近檢查", GREEN),
    ]
    x = 90
    y = 180
    for idx, (title, value, color) in enumerate(cards):
        draw_round_rect(draw, (x, y, x + 300, y + 135), fill="F8FAFC", outline=BORDER, radius=22)
        draw.text((x + 26, y + 24), title, font=font(28), fill="#" + INK)
        draw_round_rect(draw, (x + 200, y + 24, x + 270, y + 58), fill="DCFCE7", outline="BBF7D0", radius=16)
        draw.text((x + 214, y + 29), "正常", font=font(17), fill="#" + GREEN)
        draw.text((x + 26, y + 78), value, font=font(24), fill="#" + MUTED)
        x += 345
        if (idx + 1) % 4 == 0:
            x = 90
            y += 175
    draw_round_rect(draw, (90, 580, 1410, 680), fill="FEF2F2", outline="FECACA", radius=22)
    draw.text((124, 608), "遇到黃燈：先看警示原因，再處理檔案、NAS、模型或排程。遇到紅燈：暫停正式送出與批次搬檔。", font=font(28), fill="#" + RED)
    return save_image(path, img)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_border(cell, color: str = BORDER) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_run_font(run, size: int = 10, bold: bool = False, color: str = INK) -> None:
    run.font.name = FONT_LATIN
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(*hex_to_rgb(color))
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.65)
    section.right_margin = Cm(1.65)
    section.start_type = WD_SECTION_START.NEW_PAGE

    for style_name, size in [("Normal", 10), ("Title", 24), ("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12)]:
        style = doc.styles[style_name]
        style.font.name = FONT_LATIN
        style.font.size = Pt(size)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
        if style_name.startswith("Heading") or style_name == "Title":
            style.font.bold = True
            style.font.color.rgb = RGBColor(*hex_to_rgb(INK))

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("MAGI 一般使用者圖文操作手冊  |  第 ")
    set_run_font(r, 9, color=MUTED)
    add_page_number(footer)
    r = footer.add_run(" 頁")
    set_run_font(r, 9, color=MUTED)


def p(doc: Document, text: str = "", size: int = 10, bold: bool = False, color: str = INK, before: int = 0, after: int = 5):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after = Pt(after)
    para.paragraph_format.line_spacing = 1.25
    run = para.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return para


def heading(doc: Document, text: str, level: int = 1) -> None:
    para = doc.add_heading(text, level=level)
    para.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    para.paragraph_format.space_after = Pt(5)
    for run in para.runs:
        set_run_font(run, size=18 if level == 1 else 14 if level == 2 else 12, bold=True, color=INK)


def callout(doc: Document, title: str, body: str, fill: str = "EFF6FF", accent: str = BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, accent)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(2)
    r = para.add_run(title)
    set_run_font(r, 11, True, accent)
    para = cell.add_paragraph()
    para.paragraph_format.line_spacing = 1.25
    r = para.add_run(body)
    set_run_font(r, 10, False, INK)
    doc.add_paragraph()


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    hdr = t.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_shading(hdr[i], "E0F2FE")
        set_cell_border(hdr[i])
        r = hdr[i].paragraphs[0].add_run(header)
        set_run_font(r, 9, True, INK)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            set_cell_shading(cells[i], "FFFFFF")
            set_cell_border(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            para = cells[i].paragraphs[0]
            para.paragraph_format.space_after = Pt(1)
            r = para.add_run(value)
            set_run_font(r, 9, False, INK)
    doc.add_paragraph()


def command_block(doc: Document, commands: list[str]) -> None:
    t = doc.add_table(rows=1, cols=1)
    cell = t.cell(0, 0)
    set_cell_shading(cell, "111827")
    set_cell_border(cell, "111827")
    para = cell.paragraphs[0]
    para.paragraph_format.line_spacing = 1.1
    for idx, command in enumerate(commands):
        if idx:
            para.add_run("\n")
        r = para.add_run(command)
        r.font.name = "Menlo"
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(226, 232, 240)
    doc.add_paragraph()


def image(doc: Document, path: Path, caption: str) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(str(path), width=Inches(6.6))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    set_run_font(r, 9, False, MUTED)


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.space_after = Pt(2)
        r = para.add_run(item)
        set_run_font(r, 10, False, INK)


def build_manual() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="magi_visual_manual_") as td:
        temp = Path(td)
        assets = {
            "cover": make_cover(temp / "cover.png"),
            "map": make_module_map(temp / "module_map.png"),
            "workflow": make_workflow(temp / "workflow.png"),
            "dashboard": make_dashboard_mock(temp / "dashboard.png"),
            "quality": make_quality_gate(temp / "quality.png"),
            "health": make_health_cards(temp / "health.png"),
        }

        doc = Document()
        configure_doc(doc)
        image(doc, assets["cover"], "封面示意：MAGI 是案件、檔案、待辦、知識與健康狀態的共同入口。")
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("MAGI 一般使用者圖文操作手冊")
        set_run_font(run, 22, True, INK)
        p(doc, "版本：2026-05-19", size=11, bold=True, color=MUTED)
        p(doc, "適用對象：一般使用者、事務協作者、導入顧問、外部客戶", size=11, color=MUTED)
        p(doc, "適用版本：MAGI 公開版與私有版", size=11, color=MUTED)
        callout(
            doc,
            "閱讀方式",
            "第一天只看第 1 至第 4 章即可開始使用。需要法扶、閱卷、翻譯、法律研究或維運時，再跳到對應章節。",
            fill="F0FDF4",
            accent=GREEN,
        )
        doc.add_page_break()

        heading(doc, "目錄與導覽", 1)
        table(
            doc,
            ["章節", "你會學到什麼", "適合誰先看"],
            [
                ["1. 快速上手", "入口、登入、健康燈號、第一句話怎麼問", "所有人"],
                ["2. 功能地圖", "案件、檔案、AI、法扶、法律資料如何串起來", "所有人"],
                ["3. 業務概覽", "OSC 待辦與行事曆事件分開閱讀", "每日使用者"],
                ["4. 對話與工具", "怎麼問 MAGI，什麼時候用 @heavy", "每日使用者"],
                ["5. 案件與檔案", "案件卡片、資料夾、預覽、下載、分享", "案件承辦者"],
                ["6. PDF 與 OCR", "命名、信封頁排除、從 PDF 建立待辦", "案件承辦者"],
                ["7. 交付品質", "摘要、翻譯、逐字稿、DOCX 產出與品質閘門", "要交付文件者"],
                ["8. 法扶、閱卷、筆錄", "開辦、進度、結案、去重與通知", "私有版使用者"],
                ["9. 法律資料", "判決、法條、實務見解、通譯實證分類", "法律研究者"],
                ["10. 維運與安裝", "健康頁、商用檢查、公開版安裝", "管理者"],
            ],
        )

        heading(doc, "1. 快速上手", 1)
        image(doc, assets["workflow"], "圖 1：建議每日使用流程。")
        callout(
            doc,
            "第一件事：看健康狀態",
            "登入後先看 /health 或 NERV。若主狀態、資料庫、模型、OCR、NAS 或外網顯示需檢查，先處理警示，再做正式送出、批次搬檔或結案。",
            fill="EFF6FF",
            accent=BLUE,
        )
        table(
            doc,
            ["你想做的事", "直接這樣問", "MAGI 應該做什麼"],
            [
                ["查今日行程", "今天有什麼行程？", "查行事曆，不查天氣。"],
                ["查案件", "查 2025-0033 的案件狀態", "查案件資料庫與資料夾。"],
                ["摘要文件", "請摘要這份 PDF", "讀檔、OCR、列出事實與結論。"],
                ["高品質翻譯", "@heavy 翻譯，保留專有名詞原文", "逐段翻譯、產出對照與術語表。"],
                ["建立待辦", "從這份法院通知建立待辦", "解析期限、對應案件、建立 OSC 待辦。"],
                ["分享檔案", "建立這份檔案的分享連結", "建立可複製連結並確認可開啟。"],
            ],
        )

        heading(doc, "2. MAGI 功能地圖", 1)
        image(doc, assets["map"], "圖 2：MAGI 的主要功能模組。")
        bullets(
            doc,
            [
                "案件資料是核心。案件、檔案、行事曆、法扶、閱卷與書狀都應回到同一個案件編號。",
                "同名不同案不可混用。MAGI 應用案件編號、法院案號、法扶案號、案由與資料夾一起判斷。",
                "公開版不包含私有案件資料、私有 NAS 路徑、金鑰、token、cookie 或私有法律資料來源。",
                "私有版可啟用法扶、閱卷、筆錄、內部通訊頻道、NAS 與私有實務見解來源。",
            ],
        )

        heading(doc, "3. 業務概覽與畫面閱讀", 1)
        image(doc, assets["dashboard"], "圖 3：安全示意畫面。實際畫面會依版本與權限不同。")
        table(
            doc,
            ["區塊", "用途", "閱讀方式"],
            [
                ["OSC 建立待辦", "由 PDF、案件文件或系統規則建立的期限工作", "看日期、案件編號、類型與描述。"],
                ["行事曆事件", "由 Google Calendar 或匯入行程來的開庭、會議、電話", "看完整時間，不與 OSC 待辦混在同一框。"],
                ["案件卡片", "看案件狀態、案由、法院案號與資料夾", "法扶案件應標示法扶/進行中或法扶/已結案。"],
                ["健康狀態", "顯示模型、DB、OCR、NAS、外網與排程", "黃燈先排查，紅燈停止正式批次作業。"],
            ],
        )

        heading(doc, "4. 和 MAGI 對話", 1)
        p(doc, "MAGI 可以用白話中文操作。你不用記複雜指令，但要把案件、檔案、期限或輸出格式講清楚。", size=10)
        command_block(
            doc,
            [
                "查王小明案件最近有什麼待辦",
                "請摘要這份 PDF，列出事實、爭點、法院理由、結論與下一步",
                "@heavy 請翻譯這份英文文獻，產出中英對照表，專有名詞後保留原文",
                "用「最高法院 通譯」抓取判決，分類是否涉及通譯品質或未使用通譯",
            ],
        )
        callout(
            doc,
            "工具調用原則",
            "問行程就查行事曆；問案件就查案件資料庫；問判決就查法律資料；查不到就說查不到，不用模型猜。",
            fill="FFF7ED",
            accent=AMBER,
        )

        heading(doc, "5. 案件管理與檔案", 1)
        table(
            doc,
            ["功能", "怎麼操作", "注意事項"],
            [
                ["查案件", "用案件編號、當事人、法院案號、法扶案號搜尋。", "同名多案時要補法院案號或案由。"],
                ["開資料夾", "在案件卡片點資料夾，或說「打開 2026-0001 資料夾」。", "NAS 未掛載會造成外網無法列檔。"],
                ["預覽檔案", "在檔案列表點預覽。", "預覽失敗要查檔案是否存在、是否被占用、外網通道是否正常。"],
                ["下載檔案", "在檔案列表點下載。", "下載連結應確認可用，不應只開本機 Finder。"],
                ["分享連結", "點分享或請 MAGI 建立分享連結。", "分享前確認檔案內容與收件對象。"],
                ["結案", "案件卡片應有「結案」按鈕。", "人工狀態優先，掃描流程不得把已結案改回進行中。"],
            ],
        )

        heading(doc, "6. PDF、OCR、命名與待辦", 1)
        table(
            doc,
            ["情境", "MAGI 應做的事", "使用者怎麼說"],
            [
                ["法院通知", "排除信封頁、辨識法院、案號、期限、文件類型。", "請把這份法院通知命名並建立待辦。"],
                ["程序裁定", "辨識是否為結案依據或程序文件。", "請判斷這份裁定是否可作結案文件。"],
                ["判決書", "命名、摘要、入庫、建立可引用段落。", "請摘要判決理由並保留頁碼。"],
                ["對方書狀", "歸檔到對方歷次書狀並摘要主張。", "請整理對方書狀的主張與反駁方向。"],
                ["掃描不清", "標示 OCR 品質不足，請人工複核。", "請 OCR 並標示不確定文字。"],
            ],
        )
        command_block(
            doc,
            [
                "請批次整理這個資料夾的法院文件",
                "從這份 PDF 建立 OSC 待辦和 Google 行事曆事件",
                "請把掃描 PDF OCR 後輸出可搜尋文字",
            ],
        )

        heading(doc, "7. 摘要、翻譯、逐字稿與品質閘門", 1)
        image(doc, assets["quality"], "圖 4：正式交付前的品質閘門。")
        table(
            doc,
            ["輸出", "合格標準", "失敗時應怎麼做"],
            [
                ["摘要", "有文件類型、事實、爭點、理由、結論、期限與頁碼。", "重跑或改為抽取式摘要，不交付空泛文字。"],
                ["翻譯", "不漏段，專有名詞後保留原文，法律用語符合臺灣慣用語。", "用 @heavy 重跑，附術語表與中英對照表。"],
                ["逐字稿", "有可讀段落、時間資訊、說話者或不確定標示。", "音質差時提示人工複核。"],
                ["DOCX", "版面清楚、表格不擠壓、標題層級一致。", "重新產出，不用殘缺檔交付。"],
            ],
        )

        heading(doc, "8. 書狀、範本與學習修正", 1)
        bullets(
            doc,
            [
                "書狀草稿應先讀案件資料、檔案、範本與既有完稿，不應直接空想。",
                "範本資料夾是特殊資料，不顯示一般案件狀態，也不套用結案流程。",
                "使用者修改後，MAGI 可記錄差異，但只套用在同案由、同文件類型。",
                "正式書狀送出前，使用者仍需確認格式、引用、附件與事實。",
            ],
        )

        heading(doc, "9. 法扶、閱卷與筆錄", 1)
        table(
            doc,
            ["模組", "主要功能", "重要原則"],
            [
                ["法扶", "新派案、開辦、應備事項、進度回報、結案。", "同名不同案不可混搬；開辦不應有暫存舊流程混淆。"],
                ["消債應備事項", "依 OSC 邏輯列出當事人待補資料。", "所得清單每年 5 月後按可申請年度自動推進。"],
                ["進度回報", "逾期案件完整列出，回覆已回報後冷卻 60 天。", "只把進度回報送到正確頻道。"],
                ["閱卷", "檢查可下載、到院閱卷、下載與歸檔。", "已下載或已歸檔不重複通知；只有繳費單不算閱卷次數。"],
                ["筆錄", "下載、命名、歸檔與通知。", "通知只列有下載到檔案的案件，不列一長串 0 份。"],
            ],
        )

        heading(doc, "10. 法律資料與通譯實證研究", 1)
        p(doc, "法律資料功能包含法條、判決、知識庫、實務見解與指定關鍵字裁判抓取。查不到時應明確說查不到，不以模型補編。", size=10)
        command_block(
            doc,
            [
                "查民法第184條",
                "搜尋最高法院關於通譯的判決",
                "用「最高法院 通譯」抓取 TXT 與 PDF，並分類法院如何處理通譯爭議",
                "整理成表格，欄位包含裁判字號、通譯原文、分類、判決結果、來源連結",
            ],
        )
        table(
            doc,
            ["分類方向", "說明"],
            [
                ["單純帶過", "判決只提到通譯或翻譯，但不是爭點。"],
                ["通譯品質爭議", "當事人主張通譯錯誤、不完整或影響陳述。"],
                ["未使用通譯", "爭點是未提供通譯或語言理解不足。"],
                ["翻譯證據", "涉及外文證據、筆錄、譯文或翻譯可靠性。"],
                ["法院結果", "上訴駁回、發回、更正、撤銷或其他結果。"],
            ],
        )

        heading(doc, "11. 帳務、通知與對外資料", 1)
        table(
            doc,
            ["功能", "用途", "注意事項"],
            [
                ["帳務匯入", "按月匯入收入支出，排除非本人標識資料。", "薪資、勞健保、固定支出要去重。"],
                ["通知分流", "系統、法扶、閱卷、筆錄、股市與研究報告送到正確頻道。", "系統巡檢不應丟到業務頻道。"],
                ["對外資料", "產生可複製文字、分享連結、DOCX 或 PDF。", "分享前確認個資與附件。"],
                ["市場與情報", "晨報、研究報告、全球情報摘要。", "報告不得截斷，非投資建議應標示。"],
            ],
        )

        heading(doc, "12. 系統健康、安裝與商用守門", 1)
        image(doc, assets["health"], "圖 5：健康頁應該讓一般使用者看得懂。")
        command_block(
            doc,
            [
                "python3 scripts/customer_install_wizard.py --public --yes",
                "python3 scripts/magi_doctor.py --json",
                "python3 scripts/public_release_audit.py --public-isolation --strict --json",
                "./venv/bin/python scripts/ops/run_test_suite.py --suite smoke62",
                "./venv/bin/python scripts/ops/commercial_readiness_live.py --strict-public",
            ],
        )
        callout(
            doc,
            "商用部署最低標準",
            "公開版要通過 public isolation、乾淨公開版安裝檢查、secret audit、ci、smoke62 與商用 readiness。私有版要確認模型、DB、NAS、OCR、外網、通知、備份與資料清理都正常。",
            fill="F0FDF4",
            accent=GREEN,
        )

        heading(doc, "13. 常見問題速查", 1)
        table(
            doc,
            ["問題", "使用者可以怎麼說", "MAGI 應怎麼處理"],
            [
                ["檔案預覽失敗", "檢查這份檔案為何無法預覽。", "查路徑、NAS、檔案鎖定、分享服務與外網。"],
                ["行事曆沒新增", "檢查高弘軒案件 PDF 待辦是否已進 Google 日曆。", "查 PDF 掃描、OSC 待辦、Calendar 寫入與去重。"],
                ["翻譯品質差", "用 @heavy 重跑，專有名詞後保留原文。", "重翻、對照、術語表、品質閘門。"],
                ["摘要太短", "請依頁面重做摘要並列可引用段落。", "改用抽取式或分段摘要。"],
                ["法扶狀態錯", "重新比對法扶案號、案件編號與資料夾。", "以人工狀態與明確案號優先。"],
                ["NAS 未掛載", "檢查 NAS 掛載並重新掛回正確名稱。", "掛載守門、路徑轉換、健康頁回報。"],
                ["公版推送前", "執行公開版隔離檢查。", "阻擋金鑰、私有路徑、私有來源與案件資料。"],
            ],
        )

        heading(doc, "14. 完整功能索引", 1)
        table(
            doc,
            ["類別", "功能", "常見產出"],
            [
                ["案件", "案件查詢、新增、狀態、結案、資料夾", "案件卡片、資料夾、狀態紀錄"],
                ["檔案", "預覽、下載、分享、OCR、命名", "分享連結、可搜尋文字、建議檔名"],
                ["文件", "摘要、翻譯、逐字稿、DOCX", "摘要報告、中英對照、逐字稿"],
                ["書狀", "草稿、範本、校對、學習修正", "書狀草稿、修正紀錄"],
                ["行事曆", "查詢、建立、提醒、冷卻", "行程、提醒、完成狀態"],
                ["法扶", "開辦、應備事項、進度、結案", "回報草稿、附件清單、結案狀態"],
                ["閱卷筆錄", "檢查、下載、歸檔、去重", "下載清單、歸檔結果"],
                ["法律資料", "法條、判決、實務見解、通譯研究", "來源表格、TXT、PDF、XLSX"],
                ["帳務", "匯入、去重、固定支出", "月報、成本摘要"],
                ["維運", "健康頁、安裝、清理、備份、商用 gate", "檢查報告、smoke 結果"],
            ],
        )

        heading(doc, "15. 命令範例總表", 1)
        p(
            doc,
            "下表是一般使用者最常用的白話命令。MAGI 應依照語意自動選擇正確工具；如果資訊不足，應要求補充案件編號、當事人、法院案號或檔案，不應猜測。",
            size=10,
        )
        table(
            doc,
            ["場景", "建議說法", "成功時應看到"],
            [
                ["查行程", "今天有什麼行程？", "依時間排列的行事曆事件，含案件與地點。"],
                ["查待辦", "列出本週 OSC 建立待辦。", "只列系統建立的待辦，不混入一般日曆行程。"],
                ["查案件", "查 2026-0001 的案件狀態。", "案件卡片、狀態、法院案號、資料夾入口。"],
                ["開資料夾", "打開 2026-0001 資料夾。", "可預覽、下載、分享，不只開本機 Finder。"],
                ["新增案件", "新增案件：王小明，民事，一審，損害賠償。", "系統自動產生案件編號並建立資料夾。"],
                ["結案", "把 2026-0001 標示為結案。", "狀態更新，後續掃描不得自動改回進行中。"],
                ["PDF 命名", "請幫這份 PDF 命名並歸檔。", "含日期、法院、案號、文件類型的檔名。"],
                ["PDF 待辦", "從這份法院通知建立待辦。", "OSC 待辦與必要的 Google Calendar 事件。"],
                ["OCR", "請 OCR 這份掃描檔，並標示不確定文字。", "可搜尋文字與 OCR 品質提醒。"],
                ["摘要", "摘要這份裁定，列出事實、爭點、理由、結論。", "結構化摘要與可引用段落。"],
                ["翻譯", "@heavy 翻譯這份 PDF，專有名詞後保留原文。", "中英對照、術語表、完整譯文。"],
                ["逐字稿", "請轉逐字稿，並整理決議與待辦。", "逐字稿、摘要、待辦清單。"],
                ["書狀", "依這件案件資料草擬陳報狀。", "書狀草稿、引用來源、待補資料。"],
                ["範本", "開啟書狀範本資料夾。", "範本清單、預覽、下載、分享連結。"],
                ["法扶狀態", "查 1150421-W-004 法扶狀態。", "開辦、進度、附件、結案狀態。"],
                ["消債應備", "產生這件消債待補資料文字。", "可複製給當事人的待補清單。"],
                ["進度回報", "羅伊辰已回報。", "冷卻 60 天並建立下次提醒。"],
                ["閱卷", "檢查這件是否有新閱卷資料。", "只列真正新資料，已歸檔不重複通知。"],
                ["筆錄", "下載這件的新筆錄。", "新筆錄檔案與歸檔位置。"],
                ["法律資料", "查民法第184條與相關判決。", "法條、裁判來源、引用片段。"],
                ["通譯研究", "用最高法院與通譯抓判決並分類。", "TXT、PDF、表格、分類與原文摘錄。"],
                ["帳務", "匯入這個月帳務，排除非本人項目。", "月報與去重後支出。"],
                ["通知", "檢查通知是否送錯頻道。", "通知分流檢查結果。"],
                ["健康", "MAGI 系統狀態。", "主狀態、DB、模型、OCR、NAS、外網。"],
                ["外網", "檢查外網為什麼連不上。", "本機健康、通道、憑證與服務狀態。"],
                ["公版檢查", "執行公開版隔離檢查。", "0 errors / 0 warnings 或待修項。"],
                ["商用檢查", "跑完整 smoke62 與 commercial readiness。", "通過/失敗摘要與報告路徑。"],
            ],
        )

        heading(doc, "16. 交付前檢查清單", 1)
        table(
            doc,
            ["檢查項目", "合格標準", "不合格時"],
            [
                ["手冊與文件", "README、公開版、私有版、DOCX 手冊連結一致。", "更新連結並重新產生 DOCX。"],
                ["公版隔離", "public_release_audit 為 0 errors / 0 warnings。", "移除金鑰、私有路徑、私有資料來源。"],
                ["安裝精靈", "乾淨 clone 可跑 customer_install_wizard dry-run。", "修正缺檔、依賴或設定範例。"],
                ["完整 smoke", "smoke62 全部通過。", "先修失敗項，不以口頭說明替代。"],
                ["健康頁", "主狀態、DB、模型、OCR、NAS、外網正常。", "暫停正式作業並修復。"],
                ["AI 品質", "摘要、翻譯、逐字稿不漏段、不亂譯、有來源。", "重跑或阻斷交付。"],
                ["工具調用", "行程、案件、法律資料、檔案各走正確工具。", "修正路由或工具健康。"],
                ["檔案服務", "預覽、下載、分享連結可用。", "查 NAS、外網、分享服務與檔案鎖定。"],
                ["資料安全", "DOCX、README、範例設定不含真實個資與金鑰。", "清除後重跑 audit。"],
            ],
        )

        doc.save(OUT)


if __name__ == "__main__":
    build_manual()
    print(f"created: {OUT}")
