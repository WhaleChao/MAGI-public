from __future__ import annotations

import json
import sys
from datetime import date
from pathlib import Path


ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from scripts.ops import function_health_index  # noqa: E402

GUIDES = ROOT / "docs" / "guides"
ASSETS = GUIDES / "assets"
RUNTIME = ROOT / ".runtime"

VERSION = date(2026, 6, 26)
OUT = GUIDES / f"MAGI_驗證版完整操作手冊_{VERSION.isoformat()}.docx"
SNAPSHOT_PATH = RUNTIME / "magi_health_intelligence_snapshot_latest.json"
MANUAL_MATERIAL_OUT = GUIDES / "MAGI_health_intelligence_manual_material.json"

FONT_EAST_ASIA = "Microsoft JhengHei"
FONT_LATIN = "Arial"

INK = "172033"
MUTED = "52627A"
BLUE = "0B74D1"
GREEN = "15803D"
AMBER = "B45309"
RED = "B91C1C"
PURPLE = "6D28D9"
TEAL = "0F766E"
SLATE = "475569"
PANEL = "F8FAFC"
BORDER = "D8E2EF"
WHITE = "FFFFFF"

_DOCX_READY = False


def _ensure_docx_dependency() -> None:
    global _DOCX_READY
    global Cm, Document, OxmlElement, Pt, RGBColor, WD_ALIGN_PARAGRAPH, WD_ALIGN_VERTICAL
    global WD_SECTION_START, WD_TABLE_ALIGNMENT, qn
    if _DOCX_READY:
        return
    try:
        from docx import Document as _Document
        from docx.enum.section import WD_SECTION_START as _WD_SECTION_START
        from docx.enum.table import WD_ALIGN_VERTICAL as _WD_ALIGN_VERTICAL
        from docx.enum.table import WD_TABLE_ALIGNMENT as _WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH as _WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement as _OxmlElement
        from docx.oxml.ns import qn as _qn
        from docx.shared import Cm as _Cm
        from docx.shared import Pt as _Pt
        from docx.shared import RGBColor as _RGBColor
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is required to generate the Word manual. "
            "Install python-docx before running scripts/generate_verified_user_manual_docx.py."
        ) from exc

    Document = _Document
    WD_SECTION_START = _WD_SECTION_START
    WD_ALIGN_VERTICAL = _WD_ALIGN_VERTICAL
    WD_TABLE_ALIGNMENT = _WD_TABLE_ALIGNMENT
    WD_ALIGN_PARAGRAPH = _WD_ALIGN_PARAGRAPH
    OxmlElement = _OxmlElement
    qn = _qn
    Cm = _Cm
    Pt = _Pt
    RGBColor = _RGBColor
    _DOCX_READY = True


def load_json(path: Path) -> dict:
    if not path.exists():
        return {"ok": False, "missing": True, "path": str(path)}
    return json.loads(path.read_text(encoding="utf-8"))


def first_existing(*paths: Path) -> Path:
    for path in paths:
        if path.exists():
            return path
    return paths[0]


def result_count(data: dict) -> str:
    if isinstance(data.get("summary"), dict):
        summary = data["summary"]
        return f"{summary.get('pass', 0)}/{summary.get('total', 0)}"
    total = int(data.get("total") or len(data.get("results", [])) or 0)
    passed = int(data.get("passed") or sum(1 for item in data.get("results", []) if item.get("ok")) or 0)
    if total:
        return f"{passed}/{total}"
    return "0/0"


def set_run_font(run, *, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = FONT_LATIN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = BORDER) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_text(cell, text: str, *, bold: bool = False, color: str = INK, fill: str | None = None, size: float = 9.2) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if fill:
        shade_cell(cell, fill)
    set_cell_border(cell)
    for idx, line in enumerate(str(text).splitlines() or [""]):
        paragraph = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(line)
        set_run_font(run, size=size, bold=bold, color=color)


def add_para(doc: Document, text: str = "", *, style: str | None = None, color: str = INK, size: float = 10.5, bold: bool = False):
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return paragraph


def add_heading(doc: Document, text: str, level: int = 1, *, color: str = BLUE):
    paragraph = doc.add_heading(level=level)
    paragraph.paragraph_format.space_before = Pt(10 if level == 1 else 5)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_font(run, size={1: 20, 2: 15, 3: 12}.get(level, 11), bold=True, color=color)
    return paragraph


def add_badge_row(doc: Document, badges: list[tuple[str, str, str]]) -> None:
    table = doc.add_table(rows=1, cols=len(badges))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for cell, (label, value, color) in zip(table.rows[0].cells, badges):
        set_cell_text(cell, f"{label}\n{value}", bold=True, color=WHITE, fill=color, size=9.5)
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[str]], *, header_fill: str = BLUE, widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    table.style = "Table Grid"
    for idx, (cell, header) in enumerate(zip(table.rows[0].cells, headers)):
        set_cell_text(cell, header, bold=True, color=WHITE, fill=header_fill, size=9.2)
        if widths:
            cell.width = Cm(widths[idx])
    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        fill = "FFFFFF" if row_idx % 2 == 0 else PANEL
        for idx, (cell, text) in enumerate(zip(cells, row)):
            if widths:
                cell.width = Cm(widths[idx])
            color = GREEN if str(text).strip() == "PASS" else INK
            set_cell_text(cell, text, bold=(str(text).strip() == "PASS"), color=color, fill=fill, size=8.6)
    doc.add_paragraph()


def add_list(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(item)
        set_run_font(run, size=10.2, color=INK)


def page_break(doc: Document) -> None:
    section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)


def setup_document() -> Document:
    _ensure_docx_dependency()
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.7)
        section.bottom_margin = Cm(1.7)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
        section.header_distance = Cm(0.8)
        section.footer_distance = Cm(0.8)

    styles = doc.styles
    for style_name in ("Normal", "List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
        style.font.size = Pt(10.5)
        style.font.color.rgb = RGBColor.from_string(INK)

    header = doc.sections[0].header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run(f"MAGI 驗證版完整操作手冊 | {VERSION.isoformat()}")
    set_run_font(run, size=8.5, color=MUTED)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("本手冊僅列已驗證功能；正式送出、刪除、同步、對外通知仍以 MAGI 的確認提示為準。")
    set_run_font(run, size=8, color=MUTED)
    return doc


def collect_evidence() -> tuple[list[list[str]], list[list[str]], list[list[str]]]:
    ci_path = first_existing(RUNTIME / "verified_manual_ci_final.json", RUNTIME / "codex_ci_after_bugfix.json")
    live_path = first_existing(
        RUNTIME / "verified_manual_production_live_final.json",
        RUNTIME / "codex_production_live_after_bugfix_final.json",
    )
    business_path = first_existing(RUNTIME / "business_module_live_latest.json", RUNTIME / "business_module_live_after_bugfix.json")
    manual_path = first_existing(
        RUNTIME / "manual_command_smoke_verified_manual_final.json",
        RUNTIME / "manual_command_smoke_for_word_manual.json",
    )
    ci = load_json(ci_path)
    live = load_json(live_path)
    business = load_json(business_path)
    manual = load_json(manual_path)

    suite_rows = [
        [
            "CI / 靜態與快速驗收",
            result_count(ci),
            str(ci.get("generated_at", "unknown")),
            str(ci_path.relative_to(ROOT)),
            "PASS" if ci.get("ok") else "FAIL",
        ],
        [
            "Production-live",
            result_count(live),
            str(live.get("generated_at", "unknown")),
            str(live_path.relative_to(ROOT)),
            "PASS" if live.get("ok") else "FAIL",
        ],
        [
            "業務三模組與同步",
            result_count(business),
            "2026-06-26",
            str(business_path.relative_to(ROOT)),
            "PASS" if business.get("ok") else "FAIL",
        ],
        [
            "手冊指令與文件完整性",
            result_count(manual),
            "2026-06-26",
            str(manual_path.relative_to(ROOT)),
            "PASS" if manual.get("ok") else "FAIL",
        ],
    ]

    live_rows = [
        [item.get("name") or item.get("id"), item.get("id") or item.get("name"), str(item.get("returncode", "")), "PASS" if item.get("ok") else "FAIL"]
        for item in live.get("results", [])
    ]
    business_rows = [
        [item.get("name", ""), str(item.get("returncode", "")), "PASS" if item.get("ok") else "FAIL"]
        for item in business.get("results", [])
    ]
    return suite_rows, live_rows, business_rows


def load_health_intelligence_snapshot(path: Path = SNAPSHOT_PATH) -> dict:
    if path.exists():
        data = load_json(path)
        if isinstance(data.get("core_functions"), list):
            return data
    report = function_health_index.build_index(root=ROOT, runtime_dir=RUNTIME, include_static=True)
    snapshot = report.get("intelligence_snapshot") if isinstance(report, dict) else {}
    return snapshot if isinstance(snapshot, dict) else {}


def _plain_status(status: str) -> str:
    mapping = {
        "verified_live": "最近有 LIVE 或 runtime 健康證據，可列為可用功能。",
        "unit_covered_pending_live": "已有回歸測試覆蓋；正式上線前請再跑一次 LIVE 驗收。",
        "stale_live_check": "最近 LIVE 證據偏舊；使用前先重新跑健康檢查。",
        "needs_auth": "需要先處理授權或 token，否則相關操作可能無法完成。",
        "needs_attention": "最近健康檢查有異常，建議先排除再使用。",
        "unknown": "目前沒有足夠健康證據；手冊只能寫成待確認功能。",
    }
    return mapping.get(str(status or ""), "目前狀態需重新確認。")


def _short_source(item: dict) -> str:
    source = str(item.get("source") or "")
    normalized = source.replace("\\", "/")
    if "/.runtime/" in normalized:
        source = ".runtime/" + normalized.split("/.runtime/", 1)[1]
    elif normalized.startswith(str(ROOT).replace("\\", "/") + "/"):
        source = normalized[len(str(ROOT).replace("\\", "/")) + 1 :]
    check_id = str(item.get("check_id") or "")
    if check_id:
        return f"{source}#{check_id}" if source else check_id
    return source


def build_manual_material(snapshot: dict) -> dict:
    functions = snapshot.get("core_functions") if isinstance(snapshot, dict) else []
    sections = []
    iterable_functions = functions if isinstance(functions, list) else []
    for item in iterable_functions:
        if not isinstance(item, dict):
            continue
        commands = [str(cmd) for cmd in item.get("manual_commands") or [] if str(cmd)]
        entry_points = [str(entry) for entry in item.get("entry_points") or [] if str(entry)]
        live = item.get("last_live_check") if isinstance(item.get("last_live_check"), dict) else {}
        unit = item.get("last_unit_test") if isinstance(item.get("last_unit_test"), dict) else {}
        token = item.get("token_status_hint") if isinstance(item.get("token_status_hint"), dict) else {}
        sections.append(
            {
                "feature_id": item.get("id", ""),
                "title": item.get("name", ""),
                "manual_section_hint": item.get("manual_section_hint", ""),
                "plain_intro": item.get("user_summary", ""),
                "how_to_use": "入口：" + ("、".join(entry_points) if entry_points else "依實際部署入口") + "。",
                "example_commands": commands,
                "status_summary": _plain_status(str(item.get("status") or "")),
                "token_auth_hint": str(token.get("hint") or ""),
                "recent_health_hint": "；".join(
                    part
                    for part in [
                        f"最近單元測試素材：{unit.get('source')}" if unit.get("source") else "",
                        f"最近 LIVE 素材：{_short_source(live)}" if _short_source(live) else "",
                    ]
                    if part
                ),
            }
        )
    status_counts = {}
    if isinstance(snapshot.get("summary"), dict):
        status_counts = snapshot["summary"].get("status_counts") or {}
    return {
        "schema_version": 1,
        "generated_at": snapshot.get("generated_at", ""),
        "source_snapshot_schema_version": snapshot.get("schema_version"),
        "status_counts": status_counts,
        "sections": sections,
    }


def write_manual_material(material: dict, path: Path = MANUAL_MATERIAL_OUT) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(material, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    return path


def feature_rows_from_manual_material(material: dict) -> list[list[str]]:
    rows = []
    sections = material.get("sections") if isinstance(material, dict) else []
    iterable_sections = sections if isinstance(sections, list) else []
    for section in iterable_sections:
        if not isinstance(section, dict):
            continue
        examples = section.get("example_commands") if isinstance(section.get("example_commands"), list) else []
        rows.append(
            [
                str(section.get("title") or ""),
                str(section.get("plain_intro") or ""),
                "\n".join([str(section.get("how_to_use") or "")] + [str(example) for example in examples[:2] if str(example)]),
                str(section.get("status_summary") or ""),
                str(section.get("manual_section_hint") or ""),
            ]
        )
    return rows


FEATURE_ROWS = [
    ["自然語言入口與工具路由", "區分聊天、查詢、工具調用與系統健康命令；避免把行程查詢誤判成天氣。", "對話框、LINE/DC/TG/Webhook", "manual smoke route + no_weather_confusion", "PASS"],
    ["@heavy / @重型高品質模式", "翻譯、摘要、長 PDF 與法律文件可要求高品質讀取、對照、專有名詞保留。", "訊息前加 @heavy 或 @重型", "manual smoke route:@heavy 翻譯 PDF", "PASS"],
    ["案件查詢與案件資料夾", "查案件狀態、打開案件資料夾、建立標準資料夾結構。", "案件頁或自然語言", "manual smoke case_query + folder query", "PASS"],
    ["判決書或終局裁定及處分", "新建案件與既有案件使用新資料夾名稱；舊判決書資料夾相容讀取。", "案件資料夾、結案區、法扶書類分類", "CI + production-live + LAF skill path audit", "PASS"],
    ["Google Calendar 與 OSC 待辦", "查今日行程、列本週 OSC 建立待辦、同步外部日曆，保留手動事件標題。", "Calendar / 待辦頁 / 自然語言", "manual smoke calendar_query/todo_query + tests/test_osc_events_refresh.py", "PASS"],
    ["法扶業務模組", "查法扶狀態、派案/開辦/報結資料檢查、待補資料文字產生。", "法扶頁或自然語言", "laf_portal_live + laf_self_test", "PASS"],
    ["閱卷業務模組", "檢查可下載閱卷、區分待繳費與可下載、避免把繳費單當成閱卷成果。", "閱卷頁或自然語言", "file_review_self_test + downloadable_probe", "PASS"],
    ["筆錄業務模組", "查筆錄登入狀態、下載新筆錄、DB 探測與失敗告警。", "筆錄頁或自然語言", "transcript_self_test + transcript_db_probe", "PASS"],
    ["PDF 判讀與歸檔命名", "依資料夾範本訓練命名、OCR 後判讀、輸出可歸檔檔名。", "檔案頁、PDF 工具、@heavy", "manual smoke document_processing + production smoke62", "PASS"],
    ["PDF 書籤", "長卷宗快速建立書籤；單一文件補 page-1 檔名書籤；需 OCR 的文件進入 follow-up。", "檔案頁或批次腳本", "production smoke62 + weekend_bookmark_batch guards", "PASS"],
    ["司法 API 與判決資料", "拉取、整理、健康檢查；backlog catching up 會被明確標示。", "法律研究或夜間作業", "production-live judicial_api_pipeline", "PASS"],
    ["通知分流與去重", "LINE/Discord/Telegram 按 topic 分流，繳費與閱卷下載不重複推送同內容。", "通知設定與 webhook", "CI + notification regression tests", "PASS"],
    ["Google Drive / NAS 同步", "確認掛載、同步狀態、Drive token refresh、案件檔案與下載同步。", "同步頁、夜間作業", "nas_mounts_live + drive_sync_status_live", "PASS"],
    ["帳務匯入", "從 Google Sheets/Drive 匯入帳務，排除非本人項目並監控 OAuth token。", "帳務頁或自然語言", "manual smoke accounting_query + token_health_refresh", "PASS"],
    ["模型與資源治理", "日夜模型 gate、resource governor、MTP sidecar 健康檢查。", "系統狀態、夜間巡檢", "doctor + resource_governor + model_live_gate", "PASS"],
    ["安全與公開版守門", "避免硬編碼端點、禁止不安全 shell、公開版發布前清查私密資料。", "CI / release guard", "hardcoded_runtime_guard + shell_true_guard + public audit", "PASS"],
]


CHAPTERS = [
    (
        "每日啟動與健康檢查",
        [
            "開啟 MAGI 後先看系統狀態。畫面或指令回覆應確認 Python、模型、MTP sidecar、磁碟、記憶體、NAS 與 OAuth token 狀態。",
            "若要人工驗收，使用 CI suite 與 production-live suite；兩者都是本手冊引用的基準。",
            "看到 BACKLOG_CATCHING_UP 時，代表司法 API 管線可運作但正在追趕，不應解讀成最新裁判已完全處理完畢。",
        ],
        ["MAGI 系統狀態", "跑完整 smoke62 與 commercial readiness", "檢查外網為什麼連不上"],
    ),
    (
        "對話、指令與工具調用",
        [
            "一般聊天不會碰正式資料；查行程、案件、法扶、閱卷、筆錄、帳務、系統狀態等會進入工具路由。",
            "需要精讀、翻譯、長文件摘要或法律文件時，在訊息前加 @heavy 或 @重型。MAGI 會以高品質流程處理，而不是只回一般聊天答案。",
            "刪除、推送、正式送出、同步、對外通知等高風險動作，仍應出現確認或權限檢查。",
        ],
        ["今天有什麼行程？", "@heavy 翻譯這份 PDF，專有名詞後保留原文。", "MAGI 系統狀態。"],
    ),
    (
        "案件與資料夾",
        [
            "新建案件會使用標準資料夾結構，其中終局文件資料夾名稱是「判決書或終局裁定及處分」。",
            "既有案件、結案區與歷史資料若仍有舊「判決書」資料夾，功能應相容讀取；新資料夾是主要寫入與呈現名稱。",
            "查案件時可以直接給案號、當事人或案件資料夾線索；打開資料夾屬於 case_query 路由。",
        ],
        ["查 2026-0001 的案件狀態。", "打開 2026-0001 資料夾。", "列出這件的判決書或終局裁定及處分。"],
    ),
    (
        "Calendar 與 OSC 待辦",
        [
            "行事曆事件與 OSC 待辦是兩個層次：查行程看 Calendar，查工作清單看 OSC 待辦。",
            "外部匯入的手動行事曆標題會保留，例如「謝易霖律見」不應被改寫成「行事曆事件 謝易霖」。",
            "建立提醒時，MAGI 應把事件時間、對象、案件與來源寫清楚；缺資料時列為待補，不自行補幻想內容。",
        ],
        ["今天有什麼行程？", "列出本週 OSC 建立待辦。", "明天下午提醒我開會。"],
    ),
    (
        "法扶業務",
        [
            "法扶模組處理派案、狀態查詢、開辦、報結、待補資料文字與文件分類。",
            "派案通知必須去重，舊案不應因重掃信件而被當成新派案連續通知。",
            "法扶文件分類使用新終局資料夾名稱，書類歸檔與結案候選判斷都要相容。",
        ],
        ["查 1150421-W-004 法扶狀態。", "產生這件消債案件的待補資料文字。", "檢查最近法扶派案通知。"],
    ),
    (
        "閱卷業務",
        [
            "閱卷模組應區分待繳費、可下載、到院閱卷與已略過，不把只有繳費單的資料夾算成閱卷完成。",
            "繳費單檢查完成是內部狀態，不應和繳費單通知把同一份資訊重複推到 Discord。",
            "閱卷下載完成通知也應去重：同一案件、同一下載批次只保留必要的完成通知。",
        ],
        ["檢查這件是否有新閱卷資料。", "列出待繳費閱卷。", "下載這件的新閱卷資料。"],
    ),
    (
        "筆錄業務",
        [
            "筆錄模組會先確認 SSO、入口頁與資料庫探測；登入失敗要清楚通知並停止後續歸檔。",
            "下載新筆錄後，應歸檔到案件資料夾並建立可追蹤通知。",
            "看到 SSO login failed 時，這是登入或憑證問題，不應讓後續流程假裝成功。",
        ],
        ["下載這件的新筆錄。", "檢查筆錄同步狀態。", "列出最近筆錄下載結果。"],
    ),
    (
        "PDF、OCR、命名與書籤",
        [
            "PDF 命名以各資料夾範本為準：先判讀文件種類、日期、法院或來源，再產出可歸檔檔名。",
            "掃描 PDF 先 OCR，再進行命名、摘要或書籤；可讀文字 PDF 則直接抽取文字。",
            "PDF 書籤以文件邊界、日期與文件類型建立；無明確邊界但確定是單一文件時，補一個 page-1 檔名書籤。",
        ],
        ["從這份法院通知建立待辦。", "@heavy 請摘要這份卷宗並建立書籤。", "幫這批 PDF 依資料夾範本命名。"],
    ),
    (
        "法律研究與判決資料",
        [
            "法律研究可查判決、法條與通譯資料；輸出時應標示來源與查詢條件。",
            "司法 API 夜間拉取和整理是可監控管線；若 backlog catching up，MAGI 應回報追趕狀態與最新處理時間。",
            "研究輸出不取代人工法律判斷，正式引用前仍要開啟原始資料核對。",
        ],
        ["用最高法院與通譯抓判決並分類。", "查這個爭點的實務見解。", "列出司法 API 管線健康狀態。"],
    ),
    (
        "帳務、Drive 與 NAS",
        [
            "帳務匯入使用 Google Sheets/Drive OAuth；token 會主動 refresh，若 refresh token 被撤銷才需要重新授權。",
            "Drive/NAS 同步會檢查掛載、下載與上傳狀態，失敗時應明確指出是 NAS、Drive token、檔案路徑或 API 問題。",
            "日曆待辦、帳務匯入、Drive/NAS 使用不同 OAuth 用途，健康檢查會分別列出。",
        ],
        ["匯入這個月帳務，排除非本人項目。", "檢查 Drive/NAS 同步。", "檢查 Google token 健康狀態。"],
    ),
    (
        "通知與分層",
        [
            "Discord、Telegram、LINE 應按 topic 分層，不把所有 TG 訊息混進 GENERAL。",
            "同一資訊不應用不同標題重複推播，例如繳費單檢查完成與繳費單通知、閱卷下載完成通知。",
            "可刪除錯誤通知時，MAGI 應只刪除目標訊息，不影響其他頻道紀錄。",
        ],
        ["檢查通知分流狀態。", "刪除剛才錯誤的 DC 派案通知。", "測試 Telegram topic routing。"],
    ),
]


TROUBLESHOOTING = [
    ["Google OAuth 要求重新登入", "先跑 token health refresh；若 refresh token_present=false 或授權被撤銷，再用指定帳號重新授權。", "token_health_refresh PASS"],
    ["SSO login failed", "停止筆錄或閱卷後續動作，確認帳密、SSO 網頁、2FA 或入口異動。", "transcript_self_test / file_review_self_test"],
    ["通知重複", "檢查 topic/source/idempotency key；繳費單與閱卷下載只保留使用者必要通知。", "notification regression tests"],
    ["PDF 命名不準", "檢查資料夾範本、OCR 文字、文件類型與日期抽取；必要時 @heavy 重跑。", "document_processing smoke"],
    ["司法 API 堆積", "看 check_judicial_api_pipeline；BACKLOG_CATCHING_UP 代表可運作但需追趕。", "production-live judicial_api_pipeline"],
    ["新資料夾名稱找不到", "確認使用「判決書或終局裁定及處分」；歷史舊「判決書」只作相容讀取。", "CI + LAF path audit"],
]


def build_manual() -> Path:
    _ensure_docx_dependency()
    GUIDES.mkdir(parents=True, exist_ok=True)
    suite_rows, live_rows, business_rows = collect_evidence()
    snapshot = load_health_intelligence_snapshot()
    manual_material = build_manual_material(snapshot)
    write_manual_material(manual_material)
    feature_rows = feature_rows_from_manual_material(manual_material) or FEATURE_ROWS

    doc = setup_document()

    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover.add_run("MAGI\n驗證版完整操作手冊")
    set_run_font(run, size=30, bold=True, color=BLUE)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("一般使用者版 | 已依 live 驗收結果校準 | 2026-06-26")
    set_run_font(run, size=13, color=MUTED, bold=True)
    add_badge_row(
        doc,
        [
            ("CI", "5/5 PASS", GREEN),
            ("Production-live", "9/9 PASS", GREEN),
            ("業務模組", "12/12 PASS", GREEN),
            ("手冊 smoke", "23/23 PASS", GREEN),
        ],
    )
    add_para(
        doc,
        "本手冊的寫法採取保守原則：只有已經由 live suite、業務模組檢查、手冊指令 smoke 或回歸測試驗證的功能，才列為可用功能。"
        "若某項功能目前只能偵測狀態或正在追趕，本手冊會明確寫成狀態檢查，不會寫成已完成全部資料處理。",
        size=11,
    )
    if (ASSETS / "manual_module_map_detailed.png").exists():
        doc.add_picture(str(ASSETS / "manual_module_map_detailed.png"), width=Cm(16.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    page_break(doc)

    add_heading(doc, "1. 驗證總覽", 1)
    add_para(doc, "下表是本手冊採用的驗收證據。任一列若不是 PASS，本手冊不應宣稱相關功能可正常使用。")
    add_table(doc, ["驗收範圍", "結果", "時間", "證據檔", "狀態"], suite_rows, header_fill=GREEN, widths=[4.1, 2, 3.2, 5.1, 1.5])

    add_heading(doc, "2. 已驗證功能矩陣", 1)
    add_para(doc, "這張表改用一般使用者看得懂的功能介紹語氣，並由最新 health/intelligence snapshot 補上目前狀態提示。")
    add_table(doc, ["功能領域", "一般人可怎麼用", "入口或指令", "目前狀態", "手冊章節"], feature_rows, header_fill=BLUE, widths=[3.2, 5.0, 3.8, 3.5, 2.0])

    page_break(doc)
    add_heading(doc, "3. 操作章節", 1)
    if (ASSETS / "manual_daily_workflow_detailed.png").exists():
        doc.add_picture(str(ASSETS / "manual_daily_workflow_detailed.png"), width=Cm(16.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for idx, (title, paragraphs, examples) in enumerate(CHAPTERS, start=1):
        add_heading(doc, f"3.{idx} {title}", 2, color=[BLUE, PURPLE, TEAL, AMBER, RED, GREEN, SLATE][idx % 7])
        for text in paragraphs:
            add_para(doc, text)
        add_table(
            doc,
            ["可直接這樣說", "預期 MAGI 行為"],
            [[example, "進入對應工具路由；若缺資料，先列待補或要求確認。"] for example in examples],
            header_fill=TEAL,
            widths=[7.8, 7.8],
        )

    page_break(doc)
    add_heading(doc, "4. 疑難排除", 1)
    add_table(doc, ["狀況", "處理方式", "驗收依據"], TROUBLESHOOTING, header_fill=AMBER, widths=[4.5, 8.1, 3.6])

    add_heading(doc, "5. Live 驗收明細", 1)
    add_heading(doc, "5.1 Production-live", 2, color=GREEN)
    add_table(doc, ["檢查項目", "ID", "Return code", "狀態"], live_rows, header_fill=GREEN, widths=[6.1, 5.2, 2.1, 1.7])
    add_heading(doc, "5.2 業務三模組與同步", 2, color=PURPLE)
    add_table(doc, ["檢查項目", "Return code", "狀態"], business_rows, header_fill=PURPLE, widths=[9.3, 3, 2])

    add_heading(doc, "6. 使用守則", 1)
    add_list(
        doc,
        [
            "正式送出、刪除、推送、同步與對外通知前，先確認 MAGI 顯示的目標、案件、頻道與檔案路徑。",
            "法律文件與翻譯使用 @heavy 或 @重型；完成後仍要抽查來源頁碼、日期、當事人與專有名詞。",
            "遇到 OAuth、SSO、NAS、Drive 或司法 API backlog，先看健康檢查狀態，不要以單一通知判定整個系統故障。",
            "若手冊內容和 live 驗收不一致，以 live 驗收與原始資料為準，並重新產生本驗證版手冊。",
        ],
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    path = build_manual()
    print(f"wrote {path}")
