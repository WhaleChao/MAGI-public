from __future__ import annotations

import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
GUIDES = ROOT / "docs" / "guides"
ASSETS = GUIDES / "assets"
RUNTIME = ROOT / ".runtime"
VERSION = date(2026, 6, 26)
OUT = GUIDES / f"MAGI_一般使用者完整操作手冊_{VERSION.isoformat()}.docx"

FONT_EAST_ASIA = "Microsoft JhengHei"
FONT_LATIN = "Arial"
INK = "172033"
MUTED = "52627A"
BLUE = "0B74D1"
GREEN = "15803D"
AMBER = "B45309"
RED = "B91C1C"
PURPLE = "6D28D9"
TEAL = "0F766E"
SLATE = "475569"
PANEL = "F8FAFC"
WHITE = "FFFFFF"
BORDER = "D8E2EF"


def set_run_font(run, *, size: float = 10.5, bold: bool = False, color: str = INK) -> None:
    run.font.name = FONT_LATIN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = BORDER) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        item = borders.find(qn(f"w:{edge}"))
        if item is None:
            item = OxmlElement(f"w:{edge}")
            borders.append(item)
        item.set(qn("w:val"), "single")
        item.set(qn("w:sz"), "6")
        item.set(qn("w:space"), "0")
        item.set(qn("w:color"), color)


def cell_text(cell, text: str, *, fill: str | None = None, color: str = INK, bold: bool = False, size: float = 9.2) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(cell)
    if fill:
        shade_cell(cell, fill)
    for idx, line in enumerate(str(text).splitlines() or [""]):
        paragraph = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(line)
        set_run_font(run, size=size, bold=bold, color=color)


def add_para(doc: Document, text: str, *, size: float = 10.5, color: str = INK, bold: bool = False, style: str | None = None):
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.2
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)
    return paragraph


def add_heading(doc: Document, text: str, level: int = 1, *, color: str = BLUE):
    paragraph = doc.add_heading(level=level)
    paragraph.paragraph_format.space_before = Pt(12 if level == 1 else 7)
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(text)
    set_run_font(run, size={1: 20, 2: 15, 3: 12}.get(level, 11), bold=True, color=color)
    return paragraph


def add_table(doc: Document, headers: list[str], rows: list[list[str]], *, header_fill: str = BLUE, widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        if widths:
            cell.width = Cm(widths[idx])
        cell_text(cell, header, fill=header_fill, color=WHITE, bold=True, size=9.2)
    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        fill = WHITE if row_idx % 2 == 0 else PANEL
        for idx, value in enumerate(row):
            if widths:
                cells[idx].width = Cm(widths[idx])
            cell_text(cells[idx], value, fill=fill, size=8.9)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(item)
        set_run_font(run, size=10.2)


def page_break(doc: Document) -> None:
    section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)


def load_json(path: Path) -> dict:
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return {}


def first_existing(*paths: Path) -> Path:
    for path in paths:
        if path.exists():
            return path
    return paths[0]


def result_count(data: dict) -> str:
    if isinstance(data.get("summary"), dict):
        summary = data["summary"]
        return f"{summary.get('pass', 0)}/{summary.get('total', 0)}"
    total = int(data.get("total") or len(data.get("results", [])) or 0)
    passed = int(data.get("passed") or sum(1 for item in data.get("results", []) if item.get("ok")) or 0)
    return f"{passed}/{total}" if total else "未記錄"


def setup_doc() -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.7)
        section.bottom_margin = Cm(1.7)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
        section.header_distance = Cm(0.8)
        section.footer_distance = Cm(0.8)
    for style_name in ("Normal", "List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
        style.font.size = Pt(10.5)
    header = doc.sections[0].header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(header.add_run(f"MAGI 一般使用者完整操作手冊 | {VERSION.isoformat()}"), size=8.5, color=MUTED)
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(footer.add_run("先用白話交辦；正式送出、刪除、上傳、同步與對外通知前，一律看清楚 MAGI 的確認內容。"), size=8, color=MUTED)
    return doc


QUICK_START = [
    ["我想知道今天要做什麼", "今天有什麼行程？\n列出本週 OSC 建立待辦。", "MAGI 會把行事曆和待辦分開列，不把行程誤當天氣。"],
    ["我想查案件", "查 2026-0001 的案件狀態。\n打開 2026-0001 資料夾。", "MAGI 會找案件資料、資料夾與相關文件。"],
    ["我有一份 PDF 要處理", "請幫我命名這份 PDF。\n從這份法院通知建立待辦。", "MAGI 會先讀文件，必要時 OCR，再提出檔名、摘要或待辦。"],
    ["我需要高品質讀文件", "@heavy 翻譯這份 PDF，保留專有名詞原文。", "@heavy / @重型 代表要更仔細讀取、對照與檢查。"],
    ["我想知道系統是否健康", "MAGI 系統狀態。\n檢查 Drive/NAS 同步。", "MAGI 會回報模型、同步、token、業務模組與警示。"],
]


WORKFLOWS = [
    (
        "案件管理",
        "用來查案件、開資料夾、看文件是否歸位，也能建立案件工作清單。",
        [
            "新建案件時，MAGI 會建立標準案件資料夾。",
            "終局文件資料夾名稱使用「判決書或終局裁定及處分」。舊資料夾「判決書」仍可相容讀取。",
            "查案件可用案號、當事人或案件名稱；不確定時先叫 MAGI 列候選案件。",
        ],
        ["查王惠薰案件狀態。", "打開 2026-0058 案件資料夾。", "列出這件的判決書或終局裁定及處分。"],
    ),
    (
        "行事曆與待辦",
        "行事曆是時間事件，OSC 待辦是工作清單。兩者會同步，但用途不同。",
        [
            "外部匯入的行事曆標題會保留，例如「謝易霖律見」不應被改名。",
            "建立提醒時，請說清日期、時間、對象和案件。",
            "不確定來源的期限會先列為待確認，不會直接當正式期限。",
        ],
        ["今天有什麼行程？", "明天下午提醒我開會。", "列出本週 OSC 建立待辦。"],
    ),
    (
        "文件、PDF、OCR 與命名",
        "適合處理法院通知、判決、對造書狀、閱卷資料、掃描卷宗與長 PDF。",
        [
            "PDF 有文字層時直接讀文字；掃描檔會先 OCR。",
            "檔名以案件資料夾內的命名範本為準。",
            "長卷宗可建立書籤；單一文件會用 page-1 檔名書籤補基本導覽。",
        ],
        ["請幫我命名這份 PDF。", "幫這份卷宗建立書籤。", "@heavy 摘要這份判決，列爭點與可引用段落。"],
    ),
    (
        "委任狀、收據與契約",
        "用來產生常用法律文件。委任狀可輸出可編輯 Word 與 PDF。",
        [
            "法扶委任狀的可填寫 Word 會以 PDF 作底圖並加上可編輯欄位，盡量貼近原 PDF。",
            "一般 OSC 委任狀、收據、委任契約會先產 DOCX，再由 DOCX 轉 PDF，避免兩版內容不同。",
            "下載或送出前，仍要人工確認姓名、案號、法院、案由、日期與律師資料。",
        ],
        ["幫我做民事委任狀。", "製作刑事辯護人委任狀 115年度偵字第1234號。", "產生這件的律師酬金收據。"],
    ),
    (
        "法扶",
        "用來追蹤派案、開辦、報結、待補資料與法扶文件。",
        [
            "派案通知會去重，舊案不應因重掃信件重複通知。",
            "開辦、報結、疑義與進度回報屬正式動作，送出前要看確認內容。",
            "法扶資料夾內的委任狀、開辦通知、接案通知與終局文件會依規則分類。",
        ],
        ["查 1150421-W-004 法扶狀態。", "產生這件消債案件的待補資料文字。", "檢查最近法扶派案通知。"],
    ),
    (
        "閱卷",
        "用來查閱卷狀態、下載卷證、處理繳費單與到院閱卷。",
        [
            "繳費單和可下載卷證會分開處理。",
            "只有繳費單的資料夾，不算閱卷下載完成。",
            "同一批繳費或下載通知會去重，不應重複推播同一內容。",
        ],
        ["檢查這件是否有新閱卷資料。", "列出待繳費閱卷。", "下載這件的新閱卷資料。"],
    ),
    (
        "筆錄",
        "用來檢查筆錄入口、下載新筆錄、歸檔到案件資料夾。",
        [
            "SSO 登入失敗時，MAGI 會停止後續流程並清楚報錯。",
            "下載成功後會留下檔案與通知紀錄。",
            "筆錄、閱卷與法扶是三個業務模組，健康檢查會分別驗證。",
        ],
        ["下載這件的新筆錄。", "檢查筆錄同步狀態。", "列出最近筆錄下載結果。"],
    ),
    (
        "帳務、Drive 與 NAS",
        "用來同步案件檔案、Google Drive、NAS 與帳務表。",
        [
            "Google OAuth token 會主動 refresh；只有 refresh token 被撤銷才需重新授權。",
            "Drive/NAS 同步會回報掛載、下載、上傳與 token 狀態。",
            "帳務匯入會排除非本人項目，匯入前後都可查狀態。",
        ],
        ["匯入這個月帳務，排除非本人項目。", "檢查 Drive/NAS 同步。", "檢查 Google token 健康狀態。"],
    ),
    (
        "通知",
        "MAGI 會把重要結果推到 LINE、Discord 或 Telegram。",
        [
            "Telegram 和 Discord 會依 topic 分層，不應全部混在 GENERAL。",
            "繳費單、閱卷下載、派案等通知會去重。",
            "如果誤發，刪除錯誤通知時應只刪目標訊息。",
        ],
        ["檢查通知分流狀態。", "刪除剛才錯誤的 DC 派案通知。", "測試 Telegram topic routing。"],
    ),
]


TROUBLESHOOTING = [
    ["MAGI 回答像聊天，沒有調工具", "把任務說具體：查案件、下載筆錄、命名 PDF、建立待辦。需要精讀時加 @heavy。"],
    ["Google 授權又過期", "先叫 MAGI 檢查 token 健康狀態；若 refresh token 被撤銷，才需要重新登入。"],
    ["筆錄同步失敗 SSO login failed", "這是入口或憑證問題，MAGI 應停止流程；不要把後續歸檔當成功。"],
    ["PDF 命名不準", "確認 PDF 是否 OCR 成功、是否放在正確資料夾、該資料夾是否有命名範本。"],
    ["通知重複", "回報是哪一類通知和案件，MAGI 會查 topic/source 去重鍵。"],
    ["找不到判決書資料夾", "新名稱是「判決書或終局裁定及處分」；舊「判決書」只作相容讀取。"],
]


def validation_rows() -> list[list[str]]:
    sources = [
        (
            "CI",
            first_existing(
                RUNTIME / "user_manual_poa_ci_precommit.json",
                RUNTIME / "user_manual_poa_ci_postcommit.json",
                RUNTIME / "user_manual_poa_ci_final.json",
                RUNTIME / "verified_manual_ci_postcommit.json",
                RUNTIME / "verified_manual_ci_final.json",
            ),
        ),
        (
            "Production-live",
            first_existing(
                RUNTIME / "user_manual_poa_production_live_precommit.json",
                RUNTIME / "user_manual_poa_production_live_postcommit.json",
                RUNTIME / "user_manual_poa_production_live_final.json",
                RUNTIME / "verified_manual_production_live_postcommit.json",
                RUNTIME / "verified_manual_production_live_final.json",
            ),
        ),
        (
            "手冊指令 smoke",
            first_existing(
                RUNTIME / "manual_command_smoke_user_manual_postfix.json",
                RUNTIME / "manual_command_smoke_user_manual_final.json",
                RUNTIME / "manual_command_smoke_verified_manual_final.json",
                RUNTIME / "manual_command_smoke_for_word_manual.json",
            ),
        ),
    ]
    rows: list[list[str]] = []
    for label, path in sources:
        data = load_json(path)
        rows.append([label, result_count(data), "PASS" if data.get("ok") else "未確認", str(path.relative_to(ROOT)) if path.exists() else str(path)])
    return rows


def build() -> Path:
    GUIDES.mkdir(parents=True, exist_ok=True)
    doc = setup_doc()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("MAGI\n一般使用者完整操作手冊")
    set_run_font(run, size=30, bold=True, color=BLUE)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(subtitle.add_run("用白話交辦法律事務工作 | 2026-06-26"), size=13, bold=True, color=MUTED)
    add_para(
        doc,
        "MAGI 是法律事務工作的 AI 助理。你不需要記很多指令；多數時候只要用中文說明想查什麼、要處理哪個案件、要 MAGI 幫你讀哪份文件即可。"
        "這份手冊先教你怎麼用，再在最後附上驗收摘要。",
        size=11,
    )
    if (ASSETS / "manual_module_map_detailed.png").exists():
        doc.add_picture(str(ASSETS / "manual_module_map_detailed.png"), width=Cm(16.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    page_break(doc)

    add_heading(doc, "1. 先看這裡", 1)
    add_para(doc, "最簡單的用法是：把 MAGI 當成懂案件資料夾、文件、行事曆和業務入口的助理。你可以先問它今天要做什麼，也可以直接交辦一件事。")
    add_table(doc, ["你想做什麼", "可以這樣說", "MAGI 應該怎麼做"], QUICK_START, header_fill=GREEN, widths=[3.7, 5.7, 6.4])

    add_heading(doc, "2. 跟 MAGI 說話的方式", 1)
    add_table(
        doc,
        ["情境", "建議說法", "提醒"],
        [
            ["聊天或詢問功能", "你可以做什麼？\n這件事你能幫我處理嗎？", "這類問題不應啟動正式工具。"],
            ["查資料", "查某案狀態。\n列出本週待辦。", "MAGI 會進入查詢工具，不會只閒聊。"],
            ["處理文件", "摘要這份 PDF。\n從法院通知建立待辦。", "檔案要能被 MAGI 讀到；掃描檔需要 OCR。"],
            ["高品質工作", "@heavy 翻譯這份 PDF。\n@重型 摘要這份判決。", "適合法律長文、翻譯、重要摘要。"],
            ["正式動作", "送出法扶回報。\n刪除錯誤通知。", "正式送出、刪除、同步前要確認目標。"],
        ],
        header_fill=BLUE,
        widths=[3.6, 5.8, 6.4],
    )

    add_heading(doc, "3. 一天的建議流程", 1)
    if (ASSETS / "manual_daily_workflow_detailed.png").exists():
        doc.add_picture(str(ASSETS / "manual_daily_workflow_detailed.png"), width=Cm(16.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_bullets(
        doc,
        [
            "早上先查今日行程、本週待辦和系統狀態。",
            "處理案件前，先打開案件資料夾，確認文件是否在正確分類。",
            "遇到 PDF、掃描檔、判決或長文，先讓 MAGI 命名、OCR、摘要或建立書籤。",
            "法扶、閱卷、筆錄、帳務、同步和通知都屬於業務功能，完成後看結果摘要。",
            "正式送出前，人工確認當事人、案號、日期、法院、金額、頻道和檔案路徑。",
        ],
    )

    page_break(doc)
    add_heading(doc, "4. 常用功能", 1)
    for idx, (title_text, intro, bullets, examples) in enumerate(WORKFLOWS, start=1):
        add_heading(doc, f"4.{idx} {title_text}", 2, color=[BLUE, TEAL, GREEN, PURPLE, AMBER, RED, SLATE][idx % 7])
        add_para(doc, intro)
        add_bullets(doc, bullets)
        add_table(
            doc,
            ["直接這樣說", "MAGI 的預期行為"],
            [[example, "進入對應功能；缺資料時列出待補或請你確認。"] for example in examples],
            header_fill=TEAL,
            widths=[7.2, 8.4],
        )

    page_break(doc)
    add_heading(doc, "5. 常見問題", 1)
    add_table(doc, ["狀況", "處理方式"], TROUBLESHOOTING, header_fill=AMBER, widths=[5.2, 10.2])

    add_heading(doc, "6. 使用安全線", 1)
    add_bullets(
        doc,
        [
            "MAGI 可以幫你整理、查詢、草擬、下載、同步，但正式送出與刪除仍應由你確認。",
            "法律文件輸出後，務必核對姓名、案號、法院、日期、金額、對造、來源頁碼。",
            "如果 MAGI 說資料不足，先補資料；不要要求它憑空猜。",
            "如果系統狀態顯示 token、SSO、NAS 或 API 異常，先處理異常再交辦業務動作。",
        ],
    )

    add_heading(doc, "附錄：本版功能驗收摘要", 1, color=SLATE)
    add_para(doc, "以下只是背書，不是操作主體。本手冊前述功能已用 smoke、CI 或 live suite 驗證；正式環境仍以當下健康檢查為準。", color=MUTED)
    add_table(doc, ["驗收項目", "結果", "狀態", "證據檔"], validation_rows(), header_fill=SLATE, widths=[4, 2.3, 2.4, 7])

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(f"wrote {build()}")
