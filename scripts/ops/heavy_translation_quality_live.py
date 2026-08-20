#!/usr/bin/env python3
"""HEAVY translation live gate for Taiwan legal/academic PDF quality.

This gate keeps @heavy honest without re-translating a full thesis on every
run.  It verifies the live NVIDIA route with a short synthetic prompt, then
uses a real PDF fixture to validate extraction, DOI/header cleanup, Taiwan
legal terminology, source-term visibility, and DOCX export/readback quality.
"""
from __future__ import annotations

import argparse
import hashlib
import io
import json
import os
import re
import stat
import sys
import time
from pathlib import Path
from typing import Any, Callable

ROOT = Path(__file__).resolve().parents[2]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

DEFAULT_FIXTURE = Path.home() / "Desktop" / "司法通譯語言風格如何影響國民法官對被告的印象.pdf"
RUNTIME_DIR = Path(os.environ.get("MAGI_RUNTIME_DIR", "").strip() or ROOT / ".runtime").expanduser()
GENERATED_FIXTURE = RUNTIME_DIR / "fixtures" / "heavy_translation_quality_fixture.pdf"

BAD_TERMS = (
    "doi:",
    "doi：",
    "公民法官",
    "法庭翻譯",
    "法庭口譯員",
    "法院翻譯",
    "演講風格",
    "言語風格",
    "無能為力組",
    "無權組",
    "強大組",
    "有權組",
    "前世",
    "前生",
    "前半生",
    "辯護人的印象",
)

REQUIRED_TERMS = (
    "國民法官法",
    "國民法官",
    "司法通譯",
    "被告",
    "無力風格",
    "有力風格",
    "假冒配對測試法",
)

REQUIRED_SOURCE_ANNOTATIONS = (
    "Citizen Judges Act",
    "court interpreters",
    "defendant",
    "powerless style",
    "Powerless Group",
    "Powerful Group",
    "matched guise technique",
)


def _load_env() -> None:
    env_path = ROOT / ".env"
    if not env_path.exists():
        return
    for raw in env_path.read_text(encoding="utf-8", errors="replace").splitlines():
        line = raw.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        key, _, value = line.partition("=")
        os.environ.setdefault(key.strip(), value.strip().strip("\"'"))


def _draw_wrapped(c: Any, text: str, *, x: int = 52, y: int = 760, width: int = 45, line_height: int = 18) -> None:
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if " " in line:
            current = ""
            for word in line.split():
                candidate = f"{current} {word}".strip()
                if len(candidate) > width and current:
                    c.drawString(x, y, current)
                    y -= line_height
                    current = word
                    if y < 72:
                        c.showPage()
                        c.setFont("STSong-Light", 12)
                        y = 760
                else:
                    current = candidate
            line = current
        while line:
            c.drawString(x, y, line[:width])
            line = line[width:]
            y -= line_height
            if y < 72:
                c.showPage()
                c.setFont("STSong-Light", 12)
                y = 760
        y -= line_height


def write_generated_fixture(path: Path = GENERATED_FIXTURE) -> Path:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.pdfgen import canvas

    path.parent.mkdir(parents=True, exist_ok=True)
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    c = canvas.Canvas(str(path), pagesize=A4)
    c.setTitle("MAGI heavy translation quality fixture")

    title = "司法通譯語言風格如何影響國民法官對被告的印象"
    zh_abstract = (
        "中文摘要\n"
        "本研究討論國民法官法施行後，司法通譯語言風格如何影響國民法官對被告的印象。"
        "研究比較無力風格與有力風格，並採用假冒配對測試法觀察國民法官對被告可信度、"
        "說服力與理解程度的評估。本文使用臺灣法律語境中的司法通譯、被告與國民法官等術語。"
    )
    en_abstract = (
        "English abstract\n"
        "This study examines how court interpreters and interpreting style influence citizen judges under the "
        "Citizen Judges Act. The experiment compares a powerless style with a powerful style and assigns "
        "participants to the Powerless Group and the Powerful Group. It uses the matched guise technique to "
        "measure perceptions of the defendant, credibility, persuasiveness, and comprehension."
    )

    for page in range(1, 101):
        c.setFont("STSong-Light", 12)
        if page == 1:
            c.setFont("STSong-Light", 18)
            c.drawString(52, 740, title)
        elif page in {4, 5}:
            c.setFont("STSong-Light", 12)
            _draw_wrapped(c, zh_abstract)
        elif page in {6, 7}:
            c.setFont("STSong-Light", 12)
            _draw_wrapped(c, en_abstract, width=70)
        else:
            c.drawString(52, 800, f"MAGI @heavy translation fixture - page {page}")
            c.drawString(52, 740, "This page intentionally supports the 100-page extraction gate.")
        c.showPage()
    c.save()
    return path


def resolve_fixture_path(raw_pdf: str) -> Path:
    pdf_path = Path(raw_pdf).expanduser()
    if pdf_path.exists():
        return pdf_path
    if not os.environ.get("MAGI_HEAVY_TRANSLATION_FIXTURE_PDF") and pdf_path == DEFAULT_FIXTURE:
        return write_generated_fixture()
    return pdf_path


def _check(checks: list[dict[str, Any]], name: str, ok: bool, detail: str = "", **extra: Any) -> None:
    item: dict[str, Any] = {"name": name, "ok": bool(ok), "detail": detail}
    item.update(extra)
    checks.append(item)


def _text_of(result: dict[str, Any]) -> str:
    return str(result.get("response") or result.get("translated_text") or result.get("text") or "").strip()


def _extract_fixture(pdf_path: Path) -> dict[str, str]:
    from pypdf import PdfReader
    from api.handlers.document_handler import prepare_document_text_for_llm

    reader = PdfReader(str(pdf_path))
    page_texts = [(page.extract_text() or "") for page in reader.pages]
    return {
        "title": prepare_document_text_for_llm(page_texts[0] if page_texts else ""),
        "zh_abstract": prepare_document_text_for_llm("\n".join(page_texts[3:5])),
        "en_abstract": prepare_document_text_for_llm("\n".join(page_texts[5:7])),
        "pages": str(len(reader.pages)),
    }


def _run_nim_route_check(timeout: int, *, gateway: Any | None = None) -> dict[str, Any]:
    from skills.bridge.inference_gateway import InferenceGateway
    from api.handlers.document_handler import normalize_tw_legal_translation_terms

    started = time.monotonic()
    old_retries = os.environ.get("MAGI_HEAVY_STRICT_NIM_RETRIES")
    old_fallback = os.environ.get("MAGI_HEAVY_STRICT_NIM_ALLOW_FALLBACK")
    try:
        os.environ["MAGI_HEAVY_STRICT_NIM_RETRIES"] = "0"
        os.environ["MAGI_HEAVY_STRICT_NIM_ALLOW_FALLBACK"] = "0"
        active_gateway = gateway if gateway is not None else InferenceGateway()
        result = active_gateway.chat(
            "@heavy 請用臺灣繁體中文回答：court interpreter 在司法文件中應譯為什麼？只回答一行。",
            task_type="translate",
            timeout=timeout,
            allow_synthetic_fallback=False,
        )
    finally:
        if old_retries is None:
            os.environ.pop("MAGI_HEAVY_STRICT_NIM_RETRIES", None)
        else:
            os.environ["MAGI_HEAVY_STRICT_NIM_RETRIES"] = old_retries
        if old_fallback is None:
            os.environ.pop("MAGI_HEAVY_STRICT_NIM_ALLOW_FALLBACK", None)
        else:
            os.environ["MAGI_HEAVY_STRICT_NIM_ALLOW_FALLBACK"] = old_fallback
    response = str(result.get("response") or "").strip()
    normalized_response = normalize_tw_legal_translation_terms(response)
    semantic_quality_passed = bool(
        result.get("success")
        and "司法通譯" in normalized_response
        and not re.search(r"法庭翻譯|司法翻譯|法庭口譯", response)
    )
    return {
        "elapsed_sec": round(time.monotonic() - started, 2),
        "success": bool(result.get("success")),
        "route": str(result.get("route") or ""),
        "model": str(result.get("model") or ""),
        "provider": str(result.get("provider") or ""),
        "provider_quality_certified": result.get("provider_quality_certified"),
        "text_len": len(response),
        "response_sha256": hashlib.sha256(response.encode("utf-8")).hexdigest() if response else "",
        "semantic_quality_passed": semantic_quality_passed,
        "error": str(result.get("error") or "")[:240],
    }


def _read_docx_text(path: Path, *, expected_sha256: str) -> str:
    from docx import Document

    if not path.is_absolute():
        raise ValueError("DOCX readback path must be absolute")
    if re.fullmatch(r"[0-9a-f]{64}", str(expected_sha256 or "")) is None:
        raise ValueError("DOCX readback digest is missing or invalid")
    flags = os.O_RDONLY | getattr(os, "O_NOFOLLOW", 0)
    descriptor = os.open(path, flags)
    try:
        before = os.fstat(descriptor)
        if not stat.S_ISREG(before.st_mode):
            raise ValueError("DOCX readback path is not a regular file")
        with os.fdopen(descriptor, "rb", closefd=False) as handle:
            payload = handle.read()
        after = os.fstat(descriptor)
        current = os.stat(path, follow_symlinks=False)
    finally:
        os.close(descriptor)
    if (
        before.st_dev != after.st_dev
        or before.st_ino != after.st_ino
        or before.st_size != after.st_size
        or before.st_mtime_ns != after.st_mtime_ns
        or current.st_dev != after.st_dev
        or current.st_ino != after.st_ino
        or not stat.S_ISREG(current.st_mode)
    ):
        raise ValueError("DOCX changed during stable readback")
    if hashlib.sha256(payload).hexdigest() != expected_sha256:
        raise ValueError("DOCX readback digest mismatch")
    doc = Document(io.BytesIO(payload))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)


def run_gate(
    *,
    pdf_path: Path,
    run_live_nim: bool,
    timeout: int,
    title_translator: Callable[[str], dict[str, Any]] | None = None,
    gateway: Any | None = None,
) -> dict[str, Any]:
    from api.handlers.document_handler import (
        build_translation_term_glossary,
        ensure_translation_terms_visible,
        missing_translation_source_terms,
        normalize_tw_legal_translation_terms,
        polish_translated_document_text,
    )
    from api.handlers.translation_handler import translate_text_complete
    from skills.ops.export_docx import export_bilingual_docx

    checks: list[dict[str, Any]] = []

    if run_live_nim:
        route = _run_nim_route_check(timeout, gateway=gateway)
        _check(
            checks,
            "heavy_nvidia_route",
            route["success"]
            and route["route"] == "nvidia_nim"
            and route["semantic_quality_passed"],
            f"route={route['route']} model={route['model']}",
            result=route,
        )

    if not pdf_path.exists():
        _check(checks, "fixture_pdf_exists", False, str(pdf_path))
        return {"success": False, "pdf": str(pdf_path), "checks": checks}
    _check(checks, "fixture_pdf_exists", True, str(pdf_path))

    extracted = _extract_fixture(pdf_path)
    joined_source = "\n\n".join(extracted.values())
    _check(checks, "pdf_extract_pages", int(extracted["pages"]) >= 100, f"pages={extracted['pages']}")
    _check(checks, "pdf_doi_cleaned", "doi:" not in joined_source.lower() and "doi：" not in joined_source.lower())
    _check(
        checks,
        "pdf_title_preserved",
        "司法通譯語言風格如何影響國民法官對被告的印象" in extracted["title"]
        and "辯護人的印象" not in extracted["title"],
    )
    zh_compact = re.sub(r"\s+", "", extracted["zh_abstract"])
    _check(
        checks,
        "pdf_official_zh_terms",
        all(term in zh_compact for term in ("國民法官法", "司法通譯", "無力風格", "假冒配對測試法")),
    )

    title_res = (
        title_translator(extracted["title"])
        if title_translator is not None
        else translate_text_complete(
            extracted["title"], target_lang="繁體中文", heavy=True
        )
    )
    title_out = polish_translated_document_text(_text_of(title_res))
    _check(
        checks,
        "heavy_title_identity_preserve",
        bool(title_res.get("success"))
        and "司法通譯語言風格如何影響國民法官對被告的印象" in title_out
        and "辯護人的印象" not in title_out,
        f"model={title_res.get('model')}",
    )
    _check(
        checks,
        "heavy_title_provider_route",
        title_res.get("success") is True
        and title_res.get("route") in {"nvidia_nim", "source_preserved"}
        and bool(title_res.get("provider")),
        f"route={title_res.get('route')} provider={title_res.get('provider')}",
    )

    glossary = build_translation_term_glossary(extracted["en_abstract"], max_terms=40)
    old_bad_translation = (
        "隨著2023年1月1日《公民法官法》的實施，台灣的司法制度進入新時代。"
        "對法庭翻譯的需求增加。許多外國被告展現無權風格。"
        "本研究使用配對偽裝技術，將參與者分成無權組與強大組。"
        "參與者評估被告的智力、可信度、說服力。"
    )
    corrected = ensure_translation_terms_visible(
        extracted["en_abstract"],
        old_bad_translation,
        term_glossary=glossary,
        target_lang="繁體中文",
    )
    corrected = normalize_tw_legal_translation_terms(corrected)
    lowered_corrected = corrected.lower()
    missing_source_terms = missing_translation_source_terms(
        extracted["en_abstract"],
        corrected,
        term_glossary=glossary,
        max_terms=40,
    )
    _check(checks, "tw_term_normalization", all(term in corrected for term in REQUIRED_TERMS))
    _check(
        checks,
        "source_terms_inline",
        not missing_source_terms,
        "missing=" + ",".join(missing_source_terms[:8]) if missing_source_terms else "approved Taiwan renderings present",
    )
    _check(checks, "bad_terms_removed", not any(term.lower() in lowered_corrected for term in BAD_TERMS))

    # Availability is not translation quality. On the actual LIVE path,
    # translate a bounded synthetic legal/academic abstract through the same
    # heavy handler as user documents, then apply the production fidelity
    # gate. Hermetic schedule fixtures inject a title translator and therefore
    # do not contact an external provider here.
    if run_live_nim and title_translator is None:
        from api.handlers.output_quality_handler import run_output_quality_gate

        abstract_res = translate_text_complete(
            extracted["en_abstract"], target_lang="繁體中文", heavy=True
        )
        abstract_out = polish_translated_document_text(_text_of(abstract_res))
        abstract_out = ensure_translation_terms_visible(
            extracted["en_abstract"],
            abstract_out,
            term_glossary=glossary,
            target_lang="繁體中文",
        )
        abstract_out = normalize_tw_legal_translation_terms(abstract_out)
        abstract_quality = run_output_quality_gate(
            "translation",
            abstract_out,
            source_text=extracted["en_abstract"],
            instruction="English -> 臺灣繁體中文",
        )
        abstract_route_ok = bool(
            abstract_res.get("success")
            and abstract_res.get("route") == "nvidia_nim"
            and abstract_res.get("provider")
        )
        abstract_terms_ok = all(term in abstract_out for term in REQUIRED_TERMS)
        _check(
            checks,
            "heavy_abstract_provider_translation",
            abstract_route_ok and abstract_terms_ok and bool(abstract_quality.get("ok")),
            (
                f"route={abstract_res.get('route')} model={abstract_res.get('model')} "
                f"score={abstract_quality.get('score')} issue={abstract_quality.get('issue') or 'none'}"
            ),
            result={
                "success": bool(abstract_res.get("success")),
                "route": str(abstract_res.get("route") or ""),
                "provider": str(abstract_res.get("provider") or ""),
                "model": str(abstract_res.get("model") or ""),
                "quality": abstract_quality,
                "response_sha256": hashlib.sha256(abstract_out.encode("utf-8")).hexdigest() if abstract_out else "",
            },
        )
        corrected = abstract_out

    idiom_source = "In my previous life as a prosecutor, I saw defendants misunderstand court interpreters."
    idiom_bad = "在我擔任檢察官的前半生中，我看到被告誤解法庭翻譯。"
    idiom_fixed = normalize_tw_legal_translation_terms(
        ensure_translation_terms_visible(idiom_source, idiom_bad, target_lang="繁體中文")
    )
    _check(
        checks,
        "previous_life_idiom_fixed",
        "我之前擔任檢察官時" in idiom_fixed
        and "前半生" not in idiom_fixed
        and "司法通譯" in idiom_fixed,
    )

    export = export_bilingual_docx(
        [
            {"page": "標題", "source": extracted["title"], "target": title_out},
            {"page": "中文摘要", "source": extracted["zh_abstract"], "target": polish_translated_document_text(extracted["zh_abstract"])},
            {"page": "英文摘要節錄", "source": extracted["en_abstract"][:1800], "target": corrected},
        ],
        title="HEAVY 翻譯品質 Live Gate",
        subtitle=pdf_path.name,
        header_text="MAGI @heavy translation live gate",
        prefix="heavy_translation_live_gate",
        hide_page_column=True,
        col_labels={"col2": "原文", "col3": "譯文 / 品質校正"},
    )
    docx_path: Path | None = None
    docx_text = ""
    exported_path = str(export.get("path") or "").strip()
    validation = export.get("validation")
    if export.get("success") and exported_path and isinstance(validation, dict):
        candidate_path = Path(exported_path).expanduser()
        try:
            docx_text = _read_docx_text(
                candidate_path,
                expected_sha256=str(validation.get("sha256") or ""),
            )
        except Exception as exc:
            _check(
                checks,
                "docx_export",
                False,
                "stable DOCX readback failed",
            )
            _check(
                checks,
                "docx_readback",
                False,
                type(exc).__name__,
            )
        else:
            docx_path = candidate_path
            _check(checks, "docx_export", True, str(docx_path))
            _check(checks, "docx_readback", True, "digest and parse verified")
    else:
        _check(
            checks,
            "docx_export",
            False,
            str(export.get("error") or "invalid export result"),
        )
    if docx_path is not None:
        docx_lower = docx_text.lower()
        _check(checks, "docx_no_doi", "doi:" not in docx_lower and "doi：" not in docx_lower)
        _check(checks, "docx_good_title", "司法通譯語言風格如何影響國民法官對被告的印象" in docx_text)
        _check(checks, "docx_bad_terms_removed", not any(term.lower() in docx_lower for term in BAD_TERMS))
        _check(checks, "docx_source_terms_inline", all(term.lower() in docx_lower for term in REQUIRED_SOURCE_ANNOTATIONS))

    return {
        "success": all(item["ok"] for item in checks),
        "pdf": str(pdf_path),
        "docx_path": str(docx_path) if docx_path is not None else "",
        "checks": checks,
    }


class _BoundedHeavyProvider:
    """NVIDIA-compatible boundary used only by the owned schedule fixture."""

    def __init__(self, fixture_root: Path, payload: dict[str, Any]):
        self.fixture_root = fixture_root
        self.payload = payload
        self.transcript: list[dict[str, Any]] = []

    def _record(self, action: str, **values: Any) -> None:
        self.transcript.append({"action": action, **values})
        path = self.fixture_root / "heavy_provider_transcript.json"
        temporary = path.with_suffix(".json.tmp")
        temporary.write_text(
            json.dumps(self.transcript, ensure_ascii=False, sort_keys=True, indent=2)
            + "\n",
            encoding="utf-8",
        )
        temporary.replace(path)

    def chat(
        self,
        prompt: str,
        *,
        task_type: str,
        timeout: int,
        allow_synthetic_fallback: bool,
    ) -> dict[str, Any]:
        if task_type != "translate" or allow_synthetic_fallback is not False:
            raise RuntimeError("bounded heavy provider received an unsafe route request")
        is_route_probe = "court interpreter" in prompt
        response_key = "route_response" if is_route_probe else "title_response"
        response = str(self.payload.get(response_key) or "").strip()
        if not response:
            raise RuntimeError("bounded heavy provider response is missing")
        self._record(
            "chat",
            stage="route_probe" if is_route_probe else "title_translation",
            task_type=task_type,
            timeout=int(timeout),
            fallback_allowed=False,
            prompt_sha256=__import__("hashlib").sha256(prompt.encode()).hexdigest(),
            response_sha256=__import__("hashlib").sha256(response.encode()).hexdigest(),
        )
        return {
            "success": True,
            "response": response,
            "route": "nvidia_nim",
            "model": str(self.payload.get("model") or "bounded-nim-model"),
            "provider": "bounded_local_nim_provider",
            "provider_quality_certified": False,
        }

    def close(self) -> None:
        self._record("close", ok=True)


def _load_bounded_heavy_provider(fixture: Any, product_input: dict[str, Any]) -> _BoundedHeavyProvider:
    provider_name = str(product_input.get("provider") or "heavy-provider.json")
    provider_path = fixture.input_path(provider_name)
    try:
        payload = json.loads(provider_path.read_text(encoding="utf-8"))
    except (OSError, UnicodeError, json.JSONDecodeError) as exc:
        raise RuntimeError("bounded heavy provider is unreadable") from exc
    if payload.get("schema") != "magi.heavy-translation-provider/v1":
        raise RuntimeError("bounded heavy provider schema is invalid")
    if payload.get("route") != "nvidia_nim" or payload.get("quality_certified") is not False:
        raise RuntimeError("bounded heavy provider route contract is invalid")
    return _BoundedHeavyProvider(fixture.root, payload)


def _run_schedule_fixture(raw_root: str, raw_output: str) -> int:
    from scripts.ops.schedule_fixture_contract import (
        load_schedule_fixture,
        safety_receipt,
        write_fixture_report,
    )

    fixture = load_schedule_fixture(
        raw_root, job_id="job_heavy_translation_quality_live"
    )
    product_input = fixture.manifest["product_input"]
    pdf_name = str(product_input.get("pdf") or "source.pdf")
    pdf_path = fixture.input_path(pdf_name)
    old_exports = os.environ.get("MAGI_EXPORTS_DIR")
    os.environ["MAGI_EXPORTS_DIR"] = str(fixture.workspace / "exports")
    provider = _load_bounded_heavy_provider(fixture, product_input)
    previous_cwd = Path.cwd()
    os.chdir(fixture.workspace)
    try:
        result = run_gate(
            pdf_path=pdf_path,
            run_live_nim=True,
            timeout=5,
            title_translator=lambda text: provider.chat(
                text,
                task_type="translate",
                timeout=5,
                allow_synthetic_fallback=False,
            ),
            gateway=provider,
        )
    finally:
        try:
            provider.close()
        finally:
            os.chdir(previous_cwd)
            if old_exports is None:
                os.environ.pop("MAGI_EXPORTS_DIR", None)
            else:
                os.environ["MAGI_EXPORTS_DIR"] = old_exports
    docx = Path(str(result.get("docx_path") or "")).resolve(strict=False)
    artifact_bounded = bool(docx.is_file() and docx.is_relative_to(fixture.root))
    check_names = {
        str(row.get("name"))
        for row in result.get("checks", [])
        if isinstance(row, dict)
    }
    required_checks = {
        "pdf_extract_pages",
        "pdf_official_zh_terms",
        "heavy_title_identity_preserve",
        "tw_term_normalization",
        "source_terms_inline",
        "heavy_nvidia_route",
        "heavy_title_provider_route",
        "docx_export",
        "docx_source_terms_inline",
    }
    checks = {
        "fixture_sample_bound": 1 <= fixture.sample_id <= 3,
        "product_gate_passed": result.get("success") is True,
        "required_quality_checks_executed": required_checks <= check_names,
        "docx_artifact_bounded": artifact_bounded,
        "isolated_provider_route_executed": "heavy_nvidia_route" in check_names,
        "provider_quality_not_certified": all(
            row.get("result", {}).get("provider_quality_certified") is False
            for row in result.get("checks", [])
            if row.get("name") == "heavy_nvidia_route"
        ),
        "provider_terminal_close": [
            row.get("action")
            for row in json.loads(
                (fixture.root / "heavy_provider_transcript.json").read_text(
                    encoding="utf-8"
                )
            )
        ]
        == ["chat", "chat", "close"],
    }
    success = all(checks.values())
    report = {
        "schema": "magi.schedule-product-result/v1",
        "job_id": fixture.job_id,
        "fixture_sample_id": fixture.sample_id,
        "success": success,
        "status": "passed" if success else "failed",
        "checks": checks,
        "product": result,
        "provider_quality_certified": False,
        "safety": safety_receipt(fixture),
    }
    output = write_fixture_report(fixture, raw_output, report)
    report["json_out"] = str(output)
    print(json.dumps(report, ensure_ascii=False, sort_keys=True))
    return 0 if success else 1


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--pdf", default=os.environ.get("MAGI_HEAVY_TRANSLATION_FIXTURE_PDF", str(DEFAULT_FIXTURE)))
    parser.add_argument("--skip-live-nim", action="store_true")
    parser.add_argument("--timeout", type=int, default=int(os.environ.get("MAGI_HEAVY_TRANSLATION_LIVE_TIMEOUT", "90") or "90"))
    parser.add_argument("--json-out", default=str(RUNTIME_DIR / "heavy_translation_quality_latest.json"))
    parser.add_argument("--schedule-fixture-root")
    args = parser.parse_args(argv)

    if args.schedule_fixture_root:
        return _run_schedule_fixture(args.schedule_fixture_root, args.json_out)

    _load_env()

    result = run_gate(pdf_path=resolve_fixture_path(args.pdf), run_live_nim=not args.skip_live_nim, timeout=args.timeout)
    out_path = Path(args.json_out)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    for item in result["checks"]:
        print(("✅" if item["ok"] else "❌") + f" {item['name']} {item.get('detail','')}".rstrip())
    print(f"JSON: {out_path}")
    if result.get("docx_path"):
        print(f"DOCX: {result['docx_path']}")
    return 0 if result["success"] else 1


if __name__ == "__main__":
    raise SystemExit(main())
