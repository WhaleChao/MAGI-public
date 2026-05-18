from __future__ import annotations

import argparse
import json
import shutil
import subprocess
import sys
import tempfile
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Any

ROOT = Path(__file__).resolve().parents[2]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))


@dataclass
class Check:
    name: str
    ok: bool
    detail: str
    expected: str = ""
    actual: str = ""


ROUTE_CASES: list[tuple[str, str, str]] = [
    ("今天有什麼行程？", "calendar_query", "required"),
    ("列出本週 OSC 建立待辦。", "todo_query", "required"),
    ("查 2026-0001 的案件狀態。", "case_query", "required"),
    ("打開 2026-0001 資料夾。", "case_query", "required"),
    ("從這份法院通知建立待辦。", "document_processing", "required"),
    ("@heavy 翻譯這份 PDF，專有名詞後保留原文。", "document_processing", "required"),
    ("請轉逐字稿，並整理決議與待辦。", "document_processing", "required"),
    ("檢查這件是否有新閱卷資料。", "file_review_query", "required"),
    ("下載這件的新筆錄。", "transcript_query", "required"),
    ("用最高法院與通譯抓判決並分類。", "judgment_query", "required"),
    ("查 1150421-W-004 法扶狀態。", "laf_query", "required"),
    ("產生這件消債案件的待補資料文字。", "laf_query", "required"),
    ("匯入這個月帳務，排除非本人項目。", "accounting_query", "required"),
    ("MAGI 系統狀態。", "system_health", "required"),
    ("檢查外網為什麼連不上。", "system_health", "required"),
    ("跑完整 smoke62 與 commercial readiness。", "system_health", "required"),
]


def run_json(cmd: list[str], *, timeout: int = 120, cwd: Path | None = None) -> tuple[bool, str]:
    proc = subprocess.run(
        cmd,
        cwd=cwd or ROOT,
        text=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        timeout=timeout,
        check=False,
    )
    return proc.returncode == 0, (proc.stdout or "")[-1200:]


def route_checks() -> list[Check]:
    from api.tools.policies import classify_tool_requirement
    from skills.engine.realtime_data_gateway import classify_realtime_query

    checks: list[Check] = []
    for prompt, expected_tool, expected_level in ROUTE_CASES:
        req = classify_tool_requirement(prompt)
        ok = req.tool_hint == expected_tool and req.level == expected_level
        checks.append(
            Check(
                name=f"route:{prompt}",
                ok=ok,
                expected=f"{expected_level}/{expected_tool}",
                actual=f"{req.level}/{req.tool_hint}",
                detail=req.reason,
            )
        )

    # Regression guard: reminder/schedule wording must not fall into weather.
    for prompt in ["今天有什麼行程？", "明天下午提醒我開會", "列出本週 OSC 建立待辦。"]:
        actual = classify_realtime_query(prompt)
        ok = actual != "weather"
        checks.append(Check(name=f"no_weather_confusion:{prompt}", ok=ok, expected="not weather", actual=str(actual), detail=""))
    return checks


def docx_pdf_checks() -> list[Check]:
    checks: list[Check] = []
    docx = ROOT / "docs/guides/MAGI_一般使用者圖文操作手冊_2026-05-19.docx"
    pdf = ROOT / "docs/guides/MAGI_一般使用者圖文操作手冊_2026-05-19.pdf"

    if not docx.exists():
        checks.append(Check("docx_exists", False, "missing", expected=str(docx)))
    else:
        from docx import Document

        d = Document(docx)
        ok = len(d.inline_shapes) >= 6 and len(d.tables) >= 10
        checks.append(
            Check(
                "docx_visual_integrity",
                ok,
                f"images={len(d.inline_shapes)} tables={len(d.tables)} paragraphs={len(d.paragraphs)}",
                expected="images>=6 tables>=10",
                actual=f"images={len(d.inline_shapes)} tables={len(d.tables)}",
            )
        )

    if not pdf.exists():
        checks.append(Check("pdf_exists", False, "missing", expected=str(pdf)))
    else:
        import fitz

        doc = fitz.open(pdf)
        page_details: list[str] = []
        ok = doc.page_count >= 5
        image_pages = 0
        for idx, page in enumerate(doc, 1):
            text_len = len(page.get_text())
            images = len(page.get_images(full=True))
            pix = page.get_pixmap(matrix=fitz.Matrix(0.25, 0.25), alpha=False)
            # A very small set of unique samples means a broken/blank page.
            samples = pix.samples[:: max(1, len(pix.samples) // 4000)]
            unique = len(set(samples))
            page_ok = (text_len > 10 or images > 0) and unique > 8
            ok = ok and page_ok
            if images:
                image_pages += 1
            page_details.append(f"p{idx}:text={text_len},images={images},colors={unique}")
        ok = ok and image_pages >= 5
        checks.append(
            Check(
                "pdf_visual_integrity",
                ok,
                "; ".join(page_details),
                expected="pages>=5 image_pages>=5 text-or-image/page nonblank",
                actual=f"pages={doc.page_count} image_pages={image_pages}",
            )
        )
    return checks


def safe_cli_checks() -> list[Check]:
    checks: list[Check] = []
    py = sys.executable
    ok, tail = run_json([py, "scripts/generate_visual_user_manual_docx.py"], timeout=120)
    checks.append(Check("manual_generator", ok, tail, expected="regenerate DOCX/PDF manual", actual="exit=0" if ok else "exit!=0"))

    tmp = Path(tempfile.mkdtemp(prefix="magi_manual_command_clean_"))
    try:
        clone = tmp / "repo"
        proc = subprocess.run(
            ["git", "clone", "--local", "--no-hardlinks", "--quiet", str(ROOT), str(clone)],
            cwd=ROOT,
            text=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT,
            timeout=180,
            check=False,
        )
        if proc.returncode != 0:
            checks.append(Check("clean_clone", False, (proc.stdout or "")[-1200:], expected="clone committed checkout", actual="exit!=0"))
        else:
            ok, tail = run_json(
                [
                    py,
                    "scripts/customer_install_wizard.py",
                    "--public",
                    "--no-live",
                    "--skip-readiness",
                    "--no-optional",
                    "--json",
                    "--output",
                    ".runtime/manual_command_customer_install_latest.json",
                ],
                timeout=180,
                cwd=clone,
            )
            checks.append(Check("customer_install_wizard_dry_run", ok, tail, expected="public dry-run install wizard on clean checkout", actual="exit=0" if ok else "exit!=0"))
    finally:
        shutil.rmtree(tmp, ignore_errors=True)
    return checks


def main() -> int:
    parser = argparse.ArgumentParser(description="Validate user-manual command examples without mutating production data.")
    parser.add_argument("--json-out", default=".runtime/manual_command_smoke_latest.json")
    parser.add_argument("--skip-cli", action="store_true", help="Skip safe CLI executions; still validate routing and artifacts.")
    args = parser.parse_args()

    checks = route_checks()
    if not args.skip_cli:
        checks.extend(safe_cli_checks())
    checks.extend(docx_pdf_checks())

    payload: dict[str, Any] = {
        "ok": all(c.ok for c in checks),
        "summary": {
            "total": len(checks),
            "pass": sum(1 for c in checks if c.ok),
            "fail": sum(1 for c in checks if not c.ok),
        },
        "checks": [asdict(c) for c in checks],
    }
    out = ROOT / args.json_out
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(payload, ensure_ascii=False, indent=2))
    return 0 if payload["ok"] else 1


if __name__ == "__main__":
    raise SystemExit(main())
