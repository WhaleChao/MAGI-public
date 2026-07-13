#!/usr/bin/env python3
"""Smoke-test MAGI translation DOCX delivery gates."""
from __future__ import annotations

import json
import sys
import tempfile
from pathlib import Path

ROOT = Path(__file__).resolve().parents[2]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))


def _make_docx(path: Path, rows: list[tuple[str, str]], *, title: str = "MAGI test") -> None:
    from docx import Document

    doc = Document()
    doc.add_heading(title, level=1)
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "原文"
    table.rows[0].cells[1].text = "翻譯"
    for source, target in rows:
        row = table.add_row()
        row.cells[0].text = source
        row.cells[1].text = target
    doc.save(str(path))


def main() -> int:
    from api.handlers.document_handler import validate_translation_docx

    checks: list[dict[str, object]] = []
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)

        ok_path = root / "ok.docx"
        ok_source = (
            "Court interpreters and the Citizen Judges Act require careful handling of "
            "defendant impressions and matched guise technique evidence. "
        ) * 12
        ok_target = (
            "司法通譯（court interpreters）與國民法官法（Citizen Judges Act）需要謹慎處理"
            "被告（defendant）印象及假冒配對測試法（matched guise technique）證據。"
        ) * 12
        _make_docx(ok_path, [(ok_source, ok_target)] * 3)
        gate = validate_translation_docx(str(ok_path), source_text=ok_source, translated_text=ok_target, source_name="ok")
        checks.append({"name": "accept_clean_docx", "ok": bool(gate.get("ok")), "gate": gate})

        tiny_path = root / "tiny.docx"
        long_source = "This is a long legal academic source about court interpreters. " * 1200
        _make_docx(tiny_path, [("short", "短")])
        gate = validate_translation_docx(str(tiny_path), source_text=long_source, translated_text="短", source_name="tiny")
        checks.append({"name": "reject_tiny_docx_for_long_source", "ok": not bool(gate.get("ok")), "gate": gate})

        bad_path = root / "bad.docx"
        bad_source = "In my previous life as a prosecutor, I saw court interpreters. " * 20
        bad_target = "我前世是檢察官時，看過法庭翻譯。" * 20
        _make_docx(bad_path, [(bad_source, bad_target)] * 2)
        gate = validate_translation_docx(str(bad_path), source_text=bad_source, translated_text=bad_target, source_name="bad")
        checks.append({"name": "reject_bad_terms_and_idiom", "ok": not bool(gate.get("ok")), "gate": gate})

    result = {"success": all(bool(c["ok"]) for c in checks), "checks": checks}
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0 if result["success"] else 1


if __name__ == "__main__":
    raise SystemExit(main())
