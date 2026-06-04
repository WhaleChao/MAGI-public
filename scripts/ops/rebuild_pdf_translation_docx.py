#!/usr/bin/env python3
"""Rebuild a full bilingual translation DOCX from a PDF.

This is a conservative delivery path for long legal/academic PDFs:
- clean repeated DOI/page headers
- keep Chinese pages as authoritative source text
- translate English-heavy pages by sentence-safe chunks through the stable
  non-LLM GTX path
- run Taiwan legal term normalization and source-term visibility rules
- export a full bilingual DOCX and read it back for quality checks
"""
from __future__ import annotations

import argparse
import json
import os
import re
import sys
from pathlib import Path
from typing import Any

ROOT = Path(__file__).resolve().parents[2]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))


BAD_TERMS = (
    "doi:",
    "doi：",
    "公民法官",
    "法庭翻譯",
    "法庭口譯員",
    "法院翻譯",
    "演講風格",
    "言語風格",
    "無權組",
    "無權力組",
    "強大組",
    "有權力組",
    "辯護人的印象",
    "前世",
    "前生",
    "前半生",
)


def _split_sentence_safe(text: str, limit: int = 950) -> list[str]:
    s = str(text or "").strip()
    if not s:
        return []
    rough: list[str] = []
    for para in re.split(r"\n{2,}", s):
        para = re.sub(r"[ \t]+", " ", para).strip()
        if not para:
            continue
        parts = re.split(r"(?<=[.!?])\s+(?=(?:[A-Z0-9(\"'“‘]))", para)
        rough.extend([p.strip() for p in parts if p.strip()])
    out: list[str] = []
    cur = ""
    for part in rough:
        if len(part) > limit:
            if cur:
                out.append(cur)
                cur = ""
            for i in range(0, len(part), limit):
                out.append(part[i : i + limit].strip())
            continue
        cand = f"{cur} {part}".strip()
        if cur and len(cand) > limit:
            out.append(cur)
            cur = part
        else:
            cur = cand
    if cur:
        out.append(cur)
    return out


def _is_english_heavy(text: str) -> bool:
    cjk = len(re.findall(r"[\u4e00-\u9fff]", text))
    alpha = len(re.findall(r"[A-Za-z]", text))
    return alpha > max(200, cjk * 1.2)


def _translate_english_text(text: str) -> str:
    from api.handlers.document_handler import (
        build_translation_term_glossary,
        ensure_translation_terms_visible,
        normalize_tw_legal_translation_terms,
        polish_translated_document_text,
    )
    from skills.translator.action import _translate_via_google_gtx

    chunks = _split_sentence_safe(text)
    translated: list[str] = []
    glossary = build_translation_term_glossary(text, max_terms=80)
    for chunk in chunks:
        raw = _translate_via_google_gtx(chunk, target_lang="繁體中文", timeout_sec=12).strip()
        if not raw:
            raw = chunk
        raw = normalize_tw_legal_translation_terms(raw)
        raw = ensure_translation_terms_visible(chunk, raw, term_glossary=glossary, target_lang="繁體中文")
        raw = polish_translated_document_text(raw) or raw
        translated.append(raw)
    return "\n\n".join(t for t in translated if t.strip()).strip()


def _read_docx_text(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)


def rebuild(pdf: Path, *, prefix: str) -> dict[str, Any]:
    from pypdf import PdfReader

    from api.handlers.document_handler import (
        normalize_tw_legal_translation_terms,
        polish_translated_document_text,
        prepare_document_text_for_llm,
        validate_translation_docx,
    )
    from skills.ops.export_docx import export_bilingual_docx

    reader = PdfReader(str(pdf))
    rows: list[dict[str, Any]] = []
    translated_pages = 0
    preserved_pages = 0
    for idx, page in enumerate(reader.pages, 1):
        source = prepare_document_text_for_llm(page.extract_text() or "")
        if not source.strip():
            continue
        if _is_english_heavy(source):
            target = _translate_english_text(source)
            translated_pages += 1
        else:
            target = polish_translated_document_text(source) or source
            preserved_pages += 1
        target = normalize_tw_legal_translation_terms(target)
        target = polish_translated_document_text(target) or target
        rows.append({"page": idx, "source": source, "target": target})

    export = export_bilingual_docx(
        rows,
        title="MAGI HEAVY 完整翻譯交付檔",
        subtitle=pdf.name,
        header_text="MAGI full bilingual translation",
        prefix=prefix,
        hide_page_column=False,
        col_labels={"col1": "頁", "col2": "原文", "col3": "繁體中文譯文 / 原文保留"},
    )
    path = Path(str(export.get("path") or "")) if export.get("success") else Path()
    text = _read_docx_text(path) if path.exists() else ""
    lower = text.lower()
    bad = [term for term in BAD_TERMS if term.lower() in lower]
    source_joined = "\n\n".join(str(row.get("source") or "") for row in rows)
    target_joined = "\n\n".join(str(row.get("target") or "") for row in rows)
    gate = validate_translation_docx(
        str(path),
        source_text=source_joined,
        translated_text=target_joined,
        source_name=pdf.name,
    ) if path.exists() else {"ok": False, "issues": ["missing_output"]}
    return {
        "success": bool(export.get("success")) and path.exists() and not bad and bool(gate.get("ok")),
        "path": str(path) if path else "",
        "pages": len(reader.pages),
        "rows": len(rows),
        "translated_pages": translated_pages,
        "preserved_pages": preserved_pages,
        "chars": len(text),
        "size": path.stat().st_size if path.exists() else 0,
        "bad_terms": bad,
        "quality_gate": gate,
    }


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("pdf")
    ap.add_argument("--prefix", default="file_translate_heavy_full_rebuilt")
    ap.add_argument("--json-out", default="")
    args = ap.parse_args()
    result = rebuild(Path(args.pdf), prefix=args.prefix)
    print(json.dumps(result, ensure_ascii=False, indent=2))
    if args.json_out:
        Path(args.json_out).write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    return 0 if result.get("success") else 1


if __name__ == "__main__":
    raise SystemExit(main())
