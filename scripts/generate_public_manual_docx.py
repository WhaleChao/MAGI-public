from __future__ import annotations

import re
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
import os
_MAGI_ROOT = os.path.abspath(os.path.dirname(os.path.dirname(__file__)))


ROOT = Path(_MAGI_ROOT)
SOURCE = ROOT / "MAGI_對外版使用說明.md"
OUTPUT = ROOT / "MAGI_對外版使用說明.docx"


HEADING_RE = re.compile(r"^(#{1,3})\s+(.*)$")
NUMBER_RE = re.compile(r"^(\d+)\.\s+(.*)$")
HORIZONTAL_RULES = {"---", "***"}
FONT_LATIN = "Arial"
FONT_EAST_ASIA = "Microsoft JhengHei"


def set_run_font(run, size_pt: int, bold: bool = False, color: RGBColor | None = None) -> None:
    font = run.font
    font.name = FONT_LATIN
    font.size = Pt(size_pt)
    font.bold = bold
    if color is not None:
        font.color.rgb = color
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.header_distance = Cm(1)
    section.footer_distance = Cm(1)
    section.start_type = WD_SECTION_START.NEW_PAGE

    normal = doc.styles["Normal"]
    normal.font.name = FONT_LATIN
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)

    for style_name, size in [("Title", 20), ("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
        style = doc.styles[style_name]
        style.font.name = FONT_LATIN
        style.font.size = Pt(size)
        style.font.bold = True
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label = footer.add_run("第 ")
    set_run_font(label, 9)
    add_page_number(footer)
    suffix = footer.add_run(" 頁")
    set_run_font(suffix, 9)


def add_styled_text(paragraph, text: str, size_pt: int = 11) -> None:
    left, sep, right = text.partition("：")
    if sep and right.strip():
        run = paragraph.add_run(left + "：")
        set_run_font(run, size_pt, bold=True, color=RGBColor(35, 64, 94))
        run = paragraph.add_run(right.strip())
        set_run_font(run, size_pt)
        return

    run = paragraph.add_run(text)
    set_run_font(run, size_pt)


def flush_buffer(doc: Document, buffer: list[str]) -> None:
    if not buffer:
        return
    text = " ".join(part.strip() for part in buffer if part.strip())
    if not text:
        buffer.clear()
        return

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.35
    add_styled_text(paragraph, text)
    buffer.clear()


def build_docx(source_path: Path, output_path: Path) -> None:
    doc = Document()
    configure_document(doc)

    lines = source_path.read_text(encoding="utf-8").splitlines()
    buffer: list[str] = []

    for raw_line in lines:
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            flush_buffer(doc, buffer)
            continue

        if stripped in HORIZONTAL_RULES:
            flush_buffer(doc, buffer)
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(2)
            continue

        heading_match = HEADING_RE.match(stripped)
        if heading_match:
            flush_buffer(doc, buffer)
            level = len(heading_match.group(1)) - 1
            text = heading_match.group(2).strip()
            heading = doc.add_heading(text, level=level)
            heading.paragraph_format.space_before = Pt(10 if level else 0)
            heading.paragraph_format.space_after = Pt(4)
            continue

        number_match = NUMBER_RE.match(stripped)
        if number_match:
            flush_buffer(doc, buffer)
            paragraph = doc.add_paragraph(style="List Number")
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.line_spacing = 1.25
            add_styled_text(paragraph, number_match.group(2).strip())
            continue

        if stripped.startswith("- "):
            flush_buffer(doc, buffer)
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.line_spacing = 1.25
            add_styled_text(paragraph, stripped[2:].strip())
            continue

        buffer.append(stripped)

    flush_buffer(doc, buffer)
    doc.save(str(output_path))
    patch_settings_xml(output_path)


def patch_settings_xml(output_path: Path) -> None:
    with tempfile.TemporaryDirectory() as tmpdir:
        temp_zip = Path(tmpdir) / "patched.docx"
        with zipfile.ZipFile(output_path, "r") as src, zipfile.ZipFile(
            temp_zip, "w", compression=zipfile.ZIP_DEFLATED
        ) as dst:
            for info in src.infolist():
                data = src.read(info.filename)
                if info.filename == "word/settings.xml":
                    xml = data.decode("utf-8")
                    xml = xml.replace(
                        '<w:zoom w:val="bestFit"/>',
                        '<w:zoom w:val="bestFit" w:percent="100"/>',
                    )
                    data = xml.encode("utf-8")
                dst.writestr(info, data)
        temp_zip.replace(output_path)


if __name__ == "__main__":
    build_docx(SOURCE, OUTPUT)
    print(f"created: {OUTPUT}")
