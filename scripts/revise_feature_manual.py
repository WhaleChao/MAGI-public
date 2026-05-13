#!/usr/bin/env python3
from __future__ import annotations

import argparse
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_DOC = ROOT / "MAGI_功能介紹手冊.docx"


PARAGRAPH_REPLACEMENTS = {
    "MAGI（Multi-Agent Governance Infrastructure，多代理治理基礎設施）是一套專為台灣法律事務所打造的 AI 作業平台。它不是一個聊天機器人，而是一位能夠理解法律工作流程、自動執行繁瑣任務的「數位助理團隊」。": "MAGI（Multi-Agent Governance Infrastructure，多代理治理基礎設施）是一套專為台灣法律事務所打造的 AI 作業平台。它不是單一聊天機器人，而是一套可依部署模組啟用的「數位助理工作台」，能理解法律工作流程並協助執行繁瑣任務。",
    "透過 LINE、Discord、Telegram 隨時接受您的指令": "透過已啟用的 LINE、Discord、Telegram 或網頁入口接受您的指令",
    "MAGI 已經安裝在事務所的工作站上，您只需要透過 LINE 就可以開始使用。以下是最常用的三個場景：": "MAGI 安裝完成後，可透過已啟用的通訊管道或網頁入口使用。以下以 LINE 為例，示範三個常用場景：",
    "情境：王律師需要為三件案件申請閱卷": "情境：王律師需要為一件案件申請閱卷",
    "在 LINE 上傳送：「批次閱卷 113訴123、113訴456、114訴789」": "在 LINE 上傳送：「申請閱卷 台北 113訴123 民事」",
    "MAGI 自動登入系統，依序處理三件申請": "MAGI 自動登入系統並開始處理申請",
    "LINE 回報：「三件閱卷申請已完成，共下載 47 份文件」": "LINE 回報：「閱卷申請已完成，後續可下載卷宗」",
    "效果：原本需要 30-40 分鐘的工作，現在只需 3 分鐘回覆驗證碼。": "效果：原本大量的重複登入與填表工作，現在可收斂為少量人工確認步驟。",
    "MAGI 自動掃描 Apple 行事曆 / Outlook，找到明天的庭期": "MAGI 可讀取已串接的 Apple 行事曆 / Outlook，找到明天的庭期",
    "產生的草稿可以直接匯出為 Word 文件，方便律師進一步編輯修改。MAGI 還會自動查詢相關法條和判決，供您參考引用。": "產生的草稿可以直接匯出為 Word 文件，方便律師進一步編輯修改。必要時也可搭配相關法條與判決查詢結果，作為引用參考。",
    "MAGI 可以透過三種通訊管道與您互動：": "MAGI 可依部署啟用下列通訊管道：",
    "MAGI 能從自然語言指令中自動產生新的技能。如果您經常重複某個工作流程，MAGI 可以將它自動化為一個新技能，日後一句話就能完成。": "MAGI 具備技能生成與安全檢查流程。對於重複性工作，可在管理員審核後包裝成新技能，供後續重複使用。",
    "MAGI 的核心服務由「程序守護 daemon」管理，會自動監控並重啟意外終止的元件，確保 24 小時不間斷運行。": "MAGI 的核心服務由「程序守護 daemon」管理，會自動監控並重啟意外終止的元件，支援長時間持續運行。",
    "所有資料和 AI 模型都在事務所本地電腦上運行，不上傳雲端": "核心資料處理與檔案儲存可在事務所本地環境運行；若啟用外部模型、政府網站或研究來源，僅在該功能執行時連線對應服務",
    "所有通訊使用 Tailscale VPN 加密隧道": "管理維運與內網存取可透過 Tailscale VPN 加密隧道",
    "外部存取透過 Cloudflare Tunnel 加密": "如啟用公開入口，可透過 Cloudflare Tunnel 保護對外連線",
    "A：不會。MAGI 所有的 AI 模型和資料處理都在事務所本地電腦上運行。案件資料不會上傳到任何雲端服務。唯一的外部連線是司法院等政府網站（執行閱卷等操作時）和通訊管道（LINE/Discord）。": "A：核心案件檔案與本地推理可留在事務所環境中。若啟用司法院或法扶網站串接、通訊 Bot、外部研究來源或其他模型路由，系統只會在執行該功能時與對應服務通訊，不會把案件資料無限制地上傳到未知的第三方服務。",
    "A：可以。MAGI 透過 LINE 管理員身分識別不同使用者。管理員可以執行所有操作，一般使用者則有讀取和查詢的權限。": "A：可以。MAGI 可依通道帳號與系統角色區分不同使用者。管理員可執行完整操作，一般使用者則依授權使用查詢或指定流程。",
    "Q：如何新增追蹤的案件？": "Q：如何新增追蹤的股票？",
    "A：直接在 LINE 上告訴 MAGI 案號即可。例如：「追蹤案件 114年度訴字第123號」，MAGI 會自動建立追蹤並定期更新進度。": "A：直接在 LINE 上告訴 MAGI 您要追蹤的標的即可，例如：「追蹤股票 台積電 AAPL」。MAGI 會更新追蹤清單，並在後續晨報中納入分析。",
    "A：支援。MAGI 同時支援 macOS（Apple Silicon）和 Windows（NVIDIA/CPU）。Windows 版使用 Ollama 作為推理引擎，功能完全相同。": "A：支援。MAGI 可部署於 macOS（Apple Silicon）與 Windows（NVIDIA/CPU）環境；Windows 版通常使用 Ollama 作為推理引擎，但實際可用功能仍取決於當地模型、憑證、通道與外部整合設定。",
}


CELL_REPLACEMENTS = {
    "完全本地運行": "核心資料本地部署",
    "所有資料、AI 模型都在事務所自己的電腦上運行，案件資料不會上傳到雲端，保障當事人隱私。": "核心資料處理與案件檔案可留在事務所環境；若啟用外部網站、通訊或模型路由，僅在該功能執行時與對應服務連線。",
    "不需要學習複雜的操作介面。透過 LINE 傳訊息，像跟助理說話一樣，MAGI 就會幫您完成工作。": "不需要學習複雜操作介面。透過已啟用的 LINE、Discord、Telegram 或網頁入口，用自然語句即可觸發常用流程。",
    "24 小時不間斷": "長時間守護運行",
    "MAGI 每天凌晨會自動進行系統維護、資料同步、健康檢查，確保隨時可用。": "MAGI 會透過 daemon、夜間巡檢與健康檢查維持服務可用性。",
    "自動登入，支援 SSO 單一登入": "自動登入既有網站流程（含既有 SSO）",
    "自動追蹤繳費期限、LINE 提醒": "自動追蹤繳費期限，並透過已啟用通道提醒",
    "自動歸檔到 NAS 案件資料夾": "依已設定的案件根目錄自動歸檔",
    "「追蹤消費者保護法第51條的新判決」": "「找最高法院關於消費者保護法第51條的判決」",
    "依案件類型自動選擇格式，填入當事人資料、案由": "依案由提供起訴狀草稿架構與段落範本",
    "分析對造起訴狀，自動標記需要回應的爭點": "依案件描述與既有資料整理答辯架構",
    "整理原審爭點、判決理由，對應上訴理由架構": "提供上訴理由架構與常見檢核段落",
    "支援假扣押、假處分等各類聲請": "支援常見聲請類型的草稿模板",
    "整理證據清單、爭點整理、法條引用": "整理證據清單、爭點與法條提示",
    "依進度自動產生陳報內容": "依案件進度提供陳報模板",
    "上傳書狀草稿即自動審核": "可對草稿進行法律用語與文體檢查",
    "Word (.docx)": "Word (.docx)",
    "建立、編輯、套用台灣法律文書格式範本": "模板填充、法律文件生成、內容抽取",
    "建立簡報、提取內容": "投影片內容抽取、增修與結構操作工具",
    "建立試算表、公式驗證、資料分析": "報表匯出、公式重算、驗證與資料處理",
    "最主要的互動管道，支援圖片、檔案傳送": "主要對話入口之一，需設定 LINE Bot 與 webhook",
    "支援多頻道、豐富的格式化訊息": "可選通道，需設定 Discord Bot",
    "輕量、快速、支援 Markdown": "可選通道，需設定 Telegram Bot",
    "使用本地 AI 模型進行中英文翻譯，資料不外流": "可優先使用本地翻譯流程；若部署開啟其他路由則依系統設定執行",
    "「閱卷 113訴123」": "「申請閱卷 台北 114訴123 民事」",
    "「批次閱卷 113訴123、114訴456」": "「下載閱卷 114年度訴字第123號」",
    "批次申請多件閱卷": "下載指定案號卷宗",
    "「草擬起訴狀 [案由]」": "「幫我做委任狀」",
    "產生書狀草稿": "產生委任狀文件",
    "「搜尋法條 [描述]」": "「法扶回報 疑義 1140728-K-002 原因 資力不合標準」",
    "語意搜尋相關法條": "處理法扶疑義回報",
    "「法扶報結 [案號]」": "「法扶回報 結案 [案號] 原因 [說明]」",
    "處理法扶報結": "處理法扶結案或報結",
}


def _replace_paragraphs(doc: Document) -> int:
    changed = 0
    for paragraph in doc.paragraphs:
        original = paragraph.text.strip()
        replacement = PARAGRAPH_REPLACEMENTS.get(original)
        if replacement and paragraph.text != replacement:
            paragraph.text = replacement
            changed += 1
    return changed


def _replace_cells(doc: Document) -> int:
    changed = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                original = cell.text.strip()
                replacement = CELL_REPLACEMENTS.get(original)
                if replacement and cell.text != replacement:
                    cell.text = replacement
                    changed += 1
    return changed


def main() -> int:
    parser = argparse.ArgumentParser(description="Revise MAGI feature manual claims to match verified product behavior.")
    parser.add_argument("--doc", default=str(DEFAULT_DOC), help="Path to the source DOCX")
    parser.add_argument("--output", default="", help="Optional explicit output path")
    args = parser.parse_args()

    source = Path(args.doc).resolve()
    if not source.exists():
        raise FileNotFoundError(f"Manual not found: {source}")

    date_tag = datetime.now().strftime("%Y%m%d")
    backup = source.with_name(f"{source.stem}_{date_tag}_pre_audit_backup{source.suffix}")
    audited_copy = Path(args.output).resolve() if args.output else source.with_name(f"{source.stem}_對外核實版{source.suffix}")

    if not backup.exists():
        shutil.copy2(source, backup)

    doc = Document(str(source))
    para_changed = _replace_paragraphs(doc)
    cell_changed = _replace_cells(doc)

    doc.save(str(source))
    doc.save(str(audited_copy))

    print(f"source={source}")
    print(f"backup={backup}")
    print(f"audited_copy={audited_copy}")
    print(f"paragraph_changes={para_changed}")
    print(f"cell_changes={cell_changed}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
