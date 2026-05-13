"""
Tests for skills/docx-editor/lib/tracked_edits.py
"""

import io
import os
import sys
import zipfile

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "skills", "docx-editor"))

import docx as python_docx
from lib.tracked_edits import apply_tracked_edits, EditInput

FIXTURES_DIR = os.path.join(os.path.dirname(__file__), "fixtures", "docx_editor")


def load_fixture(name: str) -> bytes:
    path = os.path.join(FIXTURES_DIR, name)
    with open(path, "rb") as f:
        return f.read()


def get_xml_str(docx_bytes: bytes) -> str:
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
        return zf.read("word/document.xml").decode("utf-8")


def get_zip_entry_names(docx_bytes: bytes) -> set:
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
        return set(zf.namelist())


class TestApplyTrackedEdits:
    def test_apply_single_edit_simple_docx(self):
        """happy path: 單一 edit 套用成功"""
        docx_bytes = load_fixture("simple.docx")
        edits = [EditInput(
            find="Hello World",
            replace="Hello MAGI",
            context_before="",
            context_after="",
        )]
        result = apply_tracked_edits(docx_bytes, edits, author="TestAuthor")

        assert len(result.changes) == 1
        assert len(result.errors) == 0
        assert result.changes[0].deleted_text == "World"
        assert result.changes[0].inserted_text == "MAGI"

        xml_str = get_xml_str(result.bytes)
        assert "w:del" in xml_str
        assert "w:ins" in xml_str
        assert "World" in xml_str  # in <w:delText>
        assert "MAGI" in xml_str   # in <w:ins>

    def test_apply_multiple_edits_independent(self):
        """多 edits 互不干擾"""
        docx_bytes = load_fixture("multi_paragraph.docx")
        edits = [
            EditInput(
                find="First",
                replace="1st",
                context_before="",
                context_after=" paragraph",
            ),
            EditInput(
                find="Second",
                replace="2nd",
                context_before="",
                context_after=" paragraph",
            ),
        ]
        result = apply_tracked_edits(docx_bytes, edits)

        assert len(result.changes) == 2
        assert len(result.errors) == 0

        xml_str = get_xml_str(result.bytes)
        # collapse_diff minimizes range: "First"→"1st" keeps "st" suffix, only "Fir"→"1"
        # "Second"→"2nd" keeps "nd" suffix, only "Seco"→"2"
        assert "w:del" in xml_str
        assert "w:ins" in xml_str
        # The inserted portions should appear
        assert result.changes[0].inserted_text in ("1st", "1")
        assert result.changes[1].inserted_text in ("2nd", "2")

    def test_apply_edit_in_table_does_not_break_table(self):
        """表格內 edit 不破壞表格結構"""
        docx_bytes = load_fixture("with_table.docx")
        # Edit the text after the table
        edits = [EditInput(
            find="please review",
            replace="kindly review",
            context_before="",
            context_after=" above.",
        )]
        result = apply_tracked_edits(docx_bytes, edits)

        assert len(result.errors) == 0
        assert len(result.changes) == 1

        # Verify table content still exists
        xml_str = get_xml_str(result.bytes)
        assert "Item A" in xml_str
        assert "100" in xml_str

        # Verify docx is readable
        doc = python_docx.Document(io.BytesIO(result.bytes))
        table_text = doc.tables[0].cell(1, 0).text
        assert "Item A" in table_text

    def test_id_uniqueness_with_existing_changes(self):
        """既有 tracked changes + 新 edits 後 ID 不衝突"""
        docx_bytes = load_fixture("with_existing_changes.docx")
        # The fixture has w:id=5 and w:id=3
        edits = [EditInput(
            find="Original text",
            replace="Modified text",
            context_before="",
            context_after=" here.",
        )]
        result = apply_tracked_edits(docx_bytes, edits)

        assert len(result.changes) >= 1

        xml_str = get_xml_str(result.bytes)
        # Extract all w:id values
        import re
        ids = re.findall(r'w:id="(\d+)"', xml_str)
        id_list = [int(i) for i in ids]
        # IDs must be unique
        assert len(id_list) == len(set(id_list)), f"Duplicate IDs found: {id_list}"
        # New IDs must be > 5 (max existing)
        if result.changes[0].del_id:
            assert int(result.changes[0].del_id) > 5
        if result.changes[0].ins_id:
            assert int(result.changes[0].ins_id) > 5

    def test_apply_with_anchor_not_found_returns_error(self):
        """失敗 edit 進 errors，其他成功"""
        docx_bytes = load_fixture("simple.docx")
        edits = [
            EditInput(find="Hello World", replace="Hello MAGI", context_before="", context_after=""),
            EditInput(find="不存在的文字", replace="新文字", context_before="", context_after=""),
        ]
        result = apply_tracked_edits(docx_bytes, edits)

        assert len(result.changes) == 1
        assert len(result.errors) == 1
        assert result.errors[0].index == 1
        assert "not_found" in result.errors[0].reason.lower() or "could not locate" in result.errors[0].reason.lower()

    def test_apply_all_failed_returns_original_bytes(self):
        """全失敗時 bytes 等於原 bytes"""
        docx_bytes = load_fixture("simple.docx")
        edits = [
            EditInput(find="不存在A", replace="新A", context_before="", context_after=""),
            EditInput(find="不存在B", replace="新B", context_before="", context_after=""),
        ]
        result = apply_tracked_edits(docx_bytes, edits)

        assert len(result.changes) == 0
        assert len(result.errors) == 2
        assert result.bytes == docx_bytes

    def test_invalid_docx_bytes_raises_valueerror(self):
        """非 ZIP 輸入 raise ValueError"""
        import pytest
        with pytest.raises(ValueError, match="not a valid docx"):
            apply_tracked_edits(b"not a zip file", [])

    def test_round_trip_with_python_docx(self):
        """套 edit → python-docx 讀回不報錯 + paragraph 數不變"""
        docx_bytes = load_fixture("multi_paragraph.docx")
        original_doc = python_docx.Document(io.BytesIO(docx_bytes))
        original_para_count = len(original_doc.paragraphs)

        edits = [EditInput(
            find="defendant shall pay",
            replace="defendant must pay",
            context_before="important clause: the ",
            context_after=".",
        )]
        result = apply_tracked_edits(docx_bytes, edits)

        assert len(result.errors) == 0

        # Should be readable with python-docx
        new_doc = python_docx.Document(io.BytesIO(result.bytes))
        assert len(new_doc.paragraphs) == original_para_count

    def test_zip_entries_unchanged_except_document_xml(self):
        """套 edit 前後 ZIP entry set 一致（除 document.xml 該變）"""
        docx_bytes = load_fixture("simple.docx")
        edits = [EditInput(
            find="Hello World",
            replace="Hello MAGI",
            context_before="",
            context_after="",
        )]
        result = apply_tracked_edits(docx_bytes, edits)

        original_entries = get_zip_entry_names(docx_bytes)
        new_entries = get_zip_entry_names(result.bytes)
        assert original_entries == new_entries

        # Other entries should be byte-identical
        with zipfile.ZipFile(io.BytesIO(docx_bytes)) as orig_zf:
            with zipfile.ZipFile(io.BytesIO(result.bytes)) as new_zf:
                for name in original_entries:
                    if name == "word/document.xml":
                        continue
                    orig_data = orig_zf.read(name)
                    new_data = new_zf.read(name)
                    assert orig_data == new_data, f"Entry {name} changed unexpectedly"

    def test_multi_run_edit_preserves_rpr(self):
        """對 multi_run.docx 編輯後 rPr 數量 ≥ 編輯前"""
        docx_bytes = load_fixture("multi_run.docx")

        # Count rPr elements before
        orig_xml = get_xml_str(docx_bytes)
        orig_rpr_count = orig_xml.count("<w:rPr")

        edits = [EditInput(
            find="brown fox",
            replace="red fox",
            context_before="The quick ",
            context_after=" jumps",
        )]
        result = apply_tracked_edits(docx_bytes, edits)

        assert len(result.errors) == 0

        new_xml = get_xml_str(result.bytes)
        new_rpr_count = new_xml.count("<w:rPr")
        # After splitting a run, we may have additional rPr copies
        assert new_rpr_count >= orig_rpr_count

    def test_pure_deletion(self):
        """replace = '' → 純刪除"""
        docx_bytes = load_fixture("simple.docx")
        edits = [EditInput(
            find=" World",
            replace="",
            context_before="Hello",
            context_after="",
        )]
        result = apply_tracked_edits(docx_bytes, edits)

        assert len(result.errors) == 0
        assert len(result.changes) == 1
        assert result.changes[0].deleted_text == " World"
        assert result.changes[0].inserted_text == ""
        assert result.changes[0].ins_id is None  # No insertion

        xml_str = get_xml_str(result.bytes)
        assert "w:del" in xml_str
        assert "w:ins" not in xml_str  # No insertion element

    def test_author_in_output(self):
        """author 出現在 XML tracked changes 中"""
        docx_bytes = load_fixture("simple.docx")
        edits = [EditInput(find="World", replace="MAGI", context_before="Hello ", context_after="")]
        result = apply_tracked_edits(docx_bytes, edits, author="TestLawyer")

        xml_str = get_xml_str(result.bytes)
        assert "TestLawyer" in xml_str
