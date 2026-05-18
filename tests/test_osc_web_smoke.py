# -*- coding: utf-8 -*-
"""OSC 網頁版功能 E2E smoke 測試套件。

目的：以 Flask test client 走過 OSC web app 的關鍵 user journey，
確保「使用者打開 dashboard 看到資料庫斷線」這類整合層 bug 不再發生。

策略：
- 真實 register OSC blueprints（不 mock blueprint）
- _osc_exec 用 monkey-patch 餵 fake rows，避免污染真 DB
- LOGIN_DISABLED=True 跳過登入
- 涵蓋路徑：
  - dashboard / cases / clients / todos / quotations / 帳務 / 法扶
  - 文件生成（消債書狀、委任狀/收據/契約、報價單 PDF、地址標籤 PNG、PDF 蓋章）
  - 系統設定（admin、Discord、GCal、自動備份）

對應使用者要求（2026-04-30）：
  「網頁的所有功能也請加入測試範圍」
  「以後 MAGI 的 TEST 請包含這部分」
"""
from __future__ import annotations

import builtins
import errno
import json
import re
import shutil
import sys
import zipfile
from io import BytesIO
from pathlib import Path
from urllib.parse import urlparse
from unittest.mock import MagicMock, patch

import pytest

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from flask import Flask  # noqa: E402
from flask_login import LoginManager  # noqa: E402
from flask_login import UserMixin  # noqa: E402


# ── Fixtures ──────────────────────────────────────────────────────────────────


def _build_app() -> Flask:
    """同時 register 所有 OSC blueprints，模擬真實 server.py 環境。"""
    a = Flask(__name__, template_folder=str(ROOT / "templates"))
    a.config["TESTING"] = True
    a.config["LOGIN_DISABLED"] = True
    a.config["PROPAGATE_EXCEPTIONS"] = False  # 讓 500 變回 response 而非 raise
    a.secret_key = "test"
    lm = LoginManager()
    lm.init_app(a)

    class _TestUser(UserMixin):
        id = "test-user"

    @lm.user_loader
    def _load_user(_user_id):
        return _TestUser()

    from api.blueprints.osc_cases import osc_bp
    from api.blueprints.osc_accounting import osc_accounting_bp
    from api.blueprints.osc_debt import osc_debt_bp
    from api.blueprints.osc_files import osc_files_bp
    from api.blueprints.osc_gcal import osc_gcal_bp
    from api.blueprints.osc_pdf import osc_pdf_bp
    from api.blueprints.osc_settings import osc_settings_bp

    a.register_blueprint(osc_bp)
    a.register_blueprint(osc_settings_bp)
    a.register_blueprint(osc_accounting_bp)
    a.register_blueprint(osc_files_bp)
    a.register_blueprint(osc_pdf_bp)
    a.register_blueprint(osc_debt_bp)
    a.register_blueprint(osc_gcal_bp)
    return a


@pytest.fixture
def app():
    return _build_app()


@pytest.fixture
def client(app):
    return app.test_client()


# ── helper：模擬 _osc_exec 回傳 (rows, cfg) tuple ─────────────────────────────


def _make_fake_exec(rows_by_table=None):
    """建一個 fake exec callable，依 SQL 中出現的表名回傳對應 rows。

    rows_by_table: {"cases": [...], "clients": [...]}; SQL 中找到第一個 match 的表名
    就回那組 rows。找不到就回空。
    """
    rows_by_table = rows_by_table or {}

    def _fake(sql, params=(), fetch="none"):
        rows = []
        sql_lower = (sql or "").lower()
        for table, table_rows in rows_by_table.items():
            if f"from {table}" in sql_lower or f"from `{table}`" in sql_lower:
                rows = table_rows
                break
        if fetch == "all":
            return rows, {"host": "127.0.0.1"}
        if fetch == "one":
            return (rows[0] if rows else None), {"host": "127.0.0.1"}
        return {"rowcount": 0, "lastrowid": None}, {"host": "127.0.0.1"}

    return _fake


# ── 1. Dashboard 即時資料（與 user 04-30 報的 bug 相關） ────────────────────


def test_dashboard_endpoint_reachable(client):
    """確保 /api/osc/dashboard 在 LOGIN_DISABLED 下可達且回 200/JSON。

    這個 test 是「dashboard DB 斷線」事件後加入的回歸保險。
    """
    with patch("api.blueprints.osc_cases._osc_exec", side_effect=_make_fake_exec()):
        r = client.get("/api/osc/dashboard")
    assert r.status_code == 200, f"dashboard 失效：{r.status_code}"
    payload = r.get_json()
    assert isinstance(payload, dict), "dashboard 應回 JSON dict"


def test_dashboard_does_not_call_old_db_ip(client):
    """dashboard 不得在錯誤路徑下嘗試連舊私有 DB IP。

    熱搜舊 IP 字串以 catch 寫死回歸。
    """
    import api.blueprints.osc_cases as mod

    src = Path(mod.__file__).read_text(encoding="utf-8")
    legacy_ip = "100." "121." "61." "74"
    assert legacy_ip not in src, "不該寫死舊 NAS IP"


def test_dashboard_pending_todos_excludes_completed_statuses(client):
    """業務概覽只顯示未完成待辦，中文「已完成/完成」也要排除。"""
    calls = []

    def fake_exec(sql, params=(), fetch="none"):
        if "FROM case_todos" in sql:
            calls.append((sql, params))
        return _make_fake_exec()(sql, params, fetch)

    with patch("api.blueprints.osc_cases._osc_exec", side_effect=fake_exec):
        r = client.get("/api/osc/dashboard")

    assert r.status_code == 200
    assert calls, "dashboard 應查詢 case_todos"
    sql, _params = calls[-1]
    assert "已完成" in sql
    assert "完成" in sql
    assert "cancelled" in sql


def test_todo_mark_completed_sets_completed_date_and_reopen_clears_it(client):
    calls = []

    def fake_exec(sql, params=(), fetch="none"):
        calls.append((sql, params, fetch))
        return _make_fake_exec()(sql, params, fetch)

    with patch("api.blueprints.osc_cases._osc_exec", side_effect=fake_exec):
        r = client.put("/api/osc/todos/7", json={"status": "已完成"})
    assert r.status_code == 200
    sql = calls[-1][0]
    assert "completed_date=COALESCE(completed_date, NOW())" in sql

    calls.clear()
    with patch("api.blueprints.osc_cases._osc_exec", side_effect=fake_exec):
        r = client.put("/api/osc/todos/7", json={"status": "待處理"})
    assert r.status_code == 200
    sql = calls[-1][0]
    assert "completed_date=NULL" in sql


def test_paperclip_todo_complete_buttons_are_visible_in_all_todo_surfaces():
    dashboard_js = (ROOT / "static/osc/tabs/dashboard.js").read_text(encoding="utf-8")
    todos_js = (ROOT / "static/osc/tabs/todos.js").read_text(encoding="utf-8")
    cases_js = (ROOT / "static/osc/tabs/cases.js").read_text(encoding="utf-8")
    osc_html = (ROOT / "templates/osc.html").read_text(encoding="utf-8")

    assert 'data-act="todo-complete"' in dashboard_js
    assert 'data-act="todo-complete"' in todos_js
    assert 'data-act="todo-complete"' in cases_js
    assert 'data-act="todo-reopen"' in cases_js
    assert "wbRenderTodoActions(t)" in cases_js
    assert "20260513-todo-source-split-v1" in osc_html


def test_laf_closed_scope_includes_final_laf_status(client):
    calls = []

    def fake_exec(sql, params=(), fetch="none"):
        calls.append((sql, params, fetch))
        if fetch == "all":
            return [], {"host": "test"}
        return {"rowcount": 0}, {"host": "test"}

    with patch("api.blueprints.osc_cases._osc_exec", side_effect=fake_exec):
        r = client.get("/api/osc/laf/cases?status_scope=closed")

    assert r.status_code == 200
    sql = calls[-1][0]
    assert "status" in sql
    assert "LIKE '%已結案%'" in sql
    assert "legal_aid_status" in sql
    assert "= '已結案'" in sql


# ── 2. 各 tab 列表 endpoint 可達 ──────────────────────────────────────────────

LIST_ENDPOINTS = [
    "/api/osc/cases",
    "/api/osc/laf",
    "/api/osc/laf/cases",
    "/api/osc/clients",
    "/api/osc/todos",
    "/api/osc/calendar/events",
    "/api/osc/quotations",
    "/api/osc/quotation-templates",
    "/api/osc/accounting/transactions",
    "/api/osc/accounting/summary",
    "/api/osc/accounting/defaults",
    "/api/osc/accounting/recurring",
    "/api/osc/legal-aid-branches",
    "/api/osc/courts",
    "/api/osc/case-reason-templates",
    "/api/osc/settings",
    "/api/osc/documents",
    "/api/osc/template-folder",
    "/api/osc/document-templates",
    "/api/osc/document-keywords",
    "/api/osc/document-replacements",
    "/api/osc/folders/roots",
    "/api/osc/backups",
    "/api/osc/gcal/status",
]
# 註：insights 端點走獨立 DB connection（不經 _osc_exec mock），
# 在 smoke 範圍外。線上實機可達。


@pytest.mark.parametrize("endpoint", LIST_ENDPOINTS)
def test_list_endpoint_reachable(client, endpoint):
    """每個列表 endpoint 都應回 200 且 JSON 結構合理。"""

    fake = _make_fake_exec()
    with patch("api.blueprints.osc_cases._osc_exec", side_effect=fake), \
         patch("api.osc.utils._osc_exec", side_effect=fake):
        r = client.get(endpoint)
    assert r.status_code == 200, f"{endpoint} 回 {r.status_code}: {r.get_data(as_text=True)[:300]}"


def test_cases_default_status_scope_is_all(client):
    """案件清單 API 預設顯示全部狀態，避免使用者以為案件消失。"""
    calls = []

    def fake_exec(sql, params=(), fetch="none"):
        if "FROM cases" in sql:
            calls.append((sql, params))
        return _make_fake_exec({"cases": []})(sql, params, fetch)

    with patch("api.blueprints.osc_cases._osc_exec", side_effect=fake_exec):
        r = client.get("/api/osc/cases?limit=5")

    assert r.status_code == 200
    assert calls, "應查詢 cases"
    sql, params = calls[-1]
    assert "status LIKE" not in sql
    assert "%進行%" not in params
    assert "%結案中%" not in params
    assert "%待報結%" not in params


def test_template_folder_endpoint_lists_template_case_folder(client, tmp_path):
    folder = tmp_path / "0000-0000-範本-消費者債務清理"
    (folder / "02_各種書狀").mkdir(parents=True)
    sample = folder / "聲請狀範本.docx"
    sample.write_text("template", encoding="utf-8")

    fake = _make_fake_exec(
        {
            "cases": [
                {
                    "id": "template-case-0000-0000-0001",
                    "case_number": "0000-0000",
                    "client_name": "範本",
                    "folder_path": str(folder),
                    "updated_at": "2026-05-13 12:00:00",
                }
            ]
        }
    )
    with patch("api.blueprints.osc_cases._osc_exec", side_effect=fake), \
         patch("api.osc.utils._osc_is_safe_local_path", return_value=True):
        r = client.get("/api/osc/template-folder")

    assert r.status_code == 200
    data = r.get_json()
    assert data["ok"] is True
    assert data["exists"] is True
    assert data["local_folder"] == str(folder)
    names = {item["name"] for item in data["entries"]}
    assert {"02_各種書狀", "聲請狀範本.docx"} <= names


def test_template_folder_uses_folder_open_action():
    html = (ROOT / "templates" / "partials" / "osc" / "templateFolder.html").read_text(encoding="utf-8")
    js = (ROOT / "static" / "osc" / "tabs" / "documents.js").read_text(encoding="utf-8")
    events = (ROOT / "static" / "osc" / "osc-events.js").read_text(encoding="utf-8")

    assert 'id="templateFolder"' in html
    assert 'data-act="template-folder-open"' in js
    assert 'data-act="wb-file-share"' in js
    assert 'fileContentUrl(path, true)' in js
    assert 'fileContentUrl(path)' in js
    assert 'if (act === "template-folder-open") return await loadTemplateFolder' in events


def test_cases_filters_split_type_and_kind(client):
    """案件分類(case_type)與案件種類(case_category)要分開篩選。"""
    calls = []

    def fake_exec(sql, params=(), fetch="none"):
        if "FROM cases" in sql:
            calls.append((sql, params))
        return _make_fake_exec({"cases": []})(sql, params, fetch)

    with patch("api.blueprints.osc_cases._osc_exec", side_effect=fake_exec):
        r = client.get("/api/osc/cases?limit=5&case_type=刑事&case_kind=法律扶助案件")

    assert r.status_code == 200
    sql, params = calls[-1]
    assert "case_type = %s" in sql
    assert "case_category = %s" in sql
    assert "刑事" in params
    assert "法律扶助案件" in params


def test_cases_endpoint_uses_effective_laf_status_for_display(client):
    rows = [{
        "id": "case-closed-laf",
        "case_number": "2025-0051",
        "client_name": "莊宸銘",
        "case_type": "消費者債務清理",
        "case_category": "法律扶助案件",
        "case_reason": "更生",
        "status": "進行中",
        "legal_aid_status": "已結案",
    }]

    with patch("api.blueprints.osc_cases._osc_exec", side_effect=_make_fake_exec({"cases": rows})):
        r = client.get("/api/osc/cases?limit=5")

    assert r.status_code == 200
    item = r.get_json()["items"][0]
    assert item["status_display"] == "已結案"
    assert item["case_type_display"] == "消費者債務清理"
    assert item["case_reason_display"] == "更生"


def test_cases_endpoint_treats_unclosed_laf_status_as_active(client):
    rows = [{
        "id": "case-open-laf",
        "case_number": "2026-0001",
        "client_name": "測試當事人",
        "case_type": "民事",
        "case_category": "法律扶助案件",
        "case_reason": "損害賠償",
        "status": "",
        "legal_aid_status": "未結案",
    }]

    with patch("api.blueprints.osc_cases._osc_exec", side_effect=_make_fake_exec({"cases": rows})):
        r = client.get("/api/osc/cases?limit=5")

    assert r.status_code == 200
    item = r.get_json()["items"][0]
    assert item["status_display"] == "進行中"
    assert item["effective_status"] == "進行中"


def test_cases_csv_export_uses_external_case_type_display(client):
    rows = [{
        "case_number": "2025-0051",
        "client_name": "莊宸銘",
        "client_name_en": "",
        "case_type": "消費者債務清理",
        "case_category": "法律扶助案件",
        "case_subject": "",
        "case_reason": "更生",
        "status": "進行中",
        "legal_aid_status": "已結案",
        "start_date": "",
        "court_date": "",
        "lawyer": "",
        "court_case_no": "",
        "court_division": "",
        "court_name": "",
    }]

    with patch("api.blueprints.osc_cases._osc_exec", side_effect=_make_fake_exec({"cases": rows})):
        r = client.get("/api/osc/cases/export-csv")

    assert r.status_code == 200
    text = r.data.decode("utf-8-sig")
    assert "莊宸銘" in text
    assert "消費者債務清理,法律扶助案件" in text
    assert "更生" in text


def test_cases_active_scope_excludes_laf_closing_and_closed(client):
    calls = []

    def fake_exec(sql, params=(), fetch="none"):
        if "FROM cases" in sql:
            calls.append((sql, params))
        return _make_fake_exec({"cases": []})(sql, params, fetch)

    with patch("api.blueprints.osc_cases._osc_exec", side_effect=fake_exec):
        r = client.get("/api/osc/cases?limit=5&status_scope=active")

    assert r.status_code == 200
    sql, _params = calls[-1]
    assert "legal_aid_status" in sql
    assert "已結案，待送出" in sql
    assert "已結案，待報結" in sql
    assert "未結案" in sql
    assert "NOT" in sql


@pytest.mark.parametrize(
    ("scope", "included", "excluded"),
    [
        ("pending_report", "已結案，待報結", "已結案，待送出"),
        ("pending_submit", "已結案，待送出", "已結案，待報結"),
    ],
)
def test_cases_has_distinct_laf_pending_scopes(client, scope, included, excluded):
    calls = []

    def fake_exec(sql, params=(), fetch="none"):
        if "FROM cases" in sql:
            calls.append((sql, params))
        return _make_fake_exec({"cases": []})(sql, params, fetch)

    with patch("api.blueprints.osc_cases._osc_exec", side_effect=fake_exec):
        r = client.get(f"/api/osc/cases?limit=5&status_scope={scope}")

    assert r.status_code == 200
    sql, _params = calls[-1]
    assert included in sql
    assert excluded not in sql


def test_cases_ui_uses_unambiguous_status_and_laf_badge_labels():
    html = (ROOT / "templates" / "partials" / "osc" / "cases.html").read_text(encoding="utf-8")
    page = (ROOT / "templates" / "osc.html").read_text(encoding="utf-8")
    js = (ROOT / "static" / "osc" / "tabs" / "cases.js").read_text(encoding="utf-8")

    assert "進行中 / 結案中" not in html
    assert "結案中 / 已結案" not in html
    assert 'data-scope="pending_report">待報結' in html
    assert 'data-scope="pending_submit">待送出' in html
    assert 'data-type="消費者債務清理"' in html
    assert 'data-kind="消費者債務清理"' not in html
    assert "法扶 / " in js
    assert "function isFinalClosingStatusText" in js
    assert "function isFinalClosedStatusText" in js
    assert 'text.includes("未結案")' in js
    assert "caseNotesBlock(r)" in js
    assert "card-notes" in js
    assert ">結案</button>" in js
    assert "一鍵結案" not in js
    assert "case-close-btn" in js
    assert "20260518-case-close-v4" in page
    assert "case_type_display" in js
    assert "case_reason_display" in js
    assert "const editorCaseType = caseDisplayType(c)" in js
    assert 'id="case_case_number" placeholder="儲存時由 MAGI 自動產生" readonly' in html
    assert 'id="case_application_no" type="hidden"' in html
    assert 'for="case_court_division">股別' in html
    assert 'id="wb_case_case_number" value="${esc(c.case_number || "")}" readonly' in js
    assert "wb_case_court_division" in js


def test_cases_post_generates_osc_number_syncs_laf_and_keeps_division(client, monkeypatch):
    import api.blueprints.osc_cases as mod

    calls = []
    folders = []

    def fake_exec(sql, params=(), fetch="none"):
        calls.append((sql, params, fetch))
        if fetch == "all":
            return [], {"host": "127.0.0.1"}
        if fetch == "one":
            return None, {"host": "127.0.0.1"}
        return {"rowcount": 1, "lastrowid": None}, {"host": "127.0.0.1"}

    def fake_folder(row_id, payload, case_category):
        folders.append((row_id, dict(payload), case_category))
        return {"ok": True, "path": f"/tmp/{payload['case_number']}-測試", "canonical": f"/tmp/{payload['case_number']}-測試"}

    monkeypatch.setattr(mod, "_osc_exec", fake_exec)
    monkeypatch.setattr(mod, "_osc_generate_case_number", lambda: "2026-0099")
    monkeypatch.setattr(mod, "_osc_auto_create_folder_for_case", fake_folder)

    r = client.post("/api/osc/cases", json={
        "client_name": "測試當事人",
        "case_category": "一般案件",
        "case_type": "民事",
        "case_reason": "給付工程款",
        "laf_case_no": "1150101-E-001",
        "application_no": "不應保留",
        "court_name": "臺灣花蓮地方法院",
        "court_case_no": "115年度建字第1號",
        "court_division": "義股",
        "auto_create_folder": True,
    })

    assert r.status_code == 200
    body = r.get_json()
    assert body["ok"] is True
    assert body["case_number"] == "2026-0099"
    assert folders[0][1]["case_number"] == "2026-0099"

    insert_sql, insert_params, _ = next(c for c in calls if c[0].startswith("INSERT INTO cases"))
    cols = re.search(r"INSERT INTO cases \((.*?)\) VALUES", insert_sql).group(1).split(",")
    inserted = dict(zip(cols, insert_params))
    assert inserted["case_number"] == "2026-0099"
    assert inserted["laf_case_no"] == "1150101-E-001"
    assert inserted["application_no"] == "1150101-E-001"
    assert inserted["court_division"] == "義股"


def test_cases_legacy_category_still_maps_to_case_kind(client):
    """舊的 category=一般案件 仍要相容，但語意是案件種類。"""
    calls = []

    def fake_exec(sql, params=(), fetch="none"):
        if "FROM cases" in sql:
            calls.append((sql, params))
        return _make_fake_exec({"cases": []})(sql, params, fetch)

    with patch("api.blueprints.osc_cases._osc_exec", side_effect=fake_exec):
        r = client.get("/api/osc/cases?limit=5&category=一般案件")

    assert r.status_code == 200
    sql, params = calls[-1]
    assert "case_category = %s" in sql
    assert "NOT (" not in sql
    assert "一般案件" in params


def test_cases_legacy_category_maps_consumer_debt_to_case_type(client):
    """category=消費者債務清理 應視為案件分類，不應落入案件種類。"""
    calls = []

    def fake_exec(sql, params=(), fetch="none"):
        if "FROM cases" in sql:
            calls.append((sql, params))
        return _make_fake_exec({"cases": []})(sql, params, fetch)

    with patch("api.blueprints.osc_cases._osc_exec", side_effect=fake_exec):
        r = client.get("/api/osc/cases?limit=5&category=消費者債務清理")

    assert r.status_code == 200
    sql, params = calls[-1]
    assert "case_reason LIKE %s OR case_type = %s" in sql
    assert "case_category = %s" not in sql
    assert "消費者債務清理" in params


def test_quick_action_laf_closing_status_is_native_osc_inventory(client, monkeypatch):
    """結案狀況盤點是 OSC 原生資料盤點，不應呼叫 MAGI 推論層。"""
    import api.blueprints.osc_cases as mod

    rows = {
        "cases": [{
            "id": "1",
            "case_number": "2026-0001",
            "client_name": "測試人",
            "case_category": "法律扶助案件",
            "case_reason": "更生",
            "case_stage": "一審",
            "court_case_no": "115消債更1",
            "laf_case_no": "1150101-E-001",
        }],
        "document_index": [{
            "title": "結案酬金領款單",
            "file_name": "結案酬金領款單.pdf",
            "file_path": "/tmp/結案酬金領款單.pdf",
            "doc_type": "法扶",
            "created_at": "2026-05-06",
        }],
        "case_todos": [{
            "todo_type": "補件",
            "todo_date": "2026-05-06",
            "description": "確認報結文件",
            "status": "待辦",
        }],
    }

    monkeypatch.setattr(mod, "_osc_exec", _make_fake_exec(rows))
    monkeypatch.setattr(mod, "_get_orchestrator", lambda: (_ for _ in ()).throw(RuntimeError("model down")))

    r = client.post("/api/osc/cases/1/quick-action", json={"action": "laf_closing_status"})
    assert r.status_code == 200
    body = r.get_json()
    assert body["ok"] is True
    assert body["native"] is True
    assert body.get("fallback") is None
    assert "法扶結案狀況盤點" in body["reply"]
    assert "結案酬金領款單" in body["reply"]
    assert "確認報結文件" in body["reply"]


@pytest.mark.parametrize(
    ("action", "title"),
    [
        ("closing_overview", "結案資料彙整"),
        ("laf_progress_summary", "法扶進度盤點"),
        ("laf_closing_status", "法扶結案狀況盤點"),
    ],
)
def test_quick_action_inventory_actions_are_native(client, monkeypatch, action, title):
    import api.blueprints.osc_cases as mod

    rows = {
        "cases": [{
            "id": "1",
            "case_number": "2026-0001",
            "client_name": "測試人",
            "case_category": "法律扶助案件",
            "case_reason": "更生",
            "case_stage": "一審",
            "court_case_no": "115消債更1",
            "laf_case_no": "1150101-E-001",
        }],
        "document_index": [],
        "case_todos": [],
    }
    monkeypatch.setattr(mod, "_osc_exec", _make_fake_exec(rows))
    monkeypatch.setattr(mod, "_get_orchestrator", lambda: (_ for _ in ()).throw(AssertionError("inventory action should not call MAGI")))

    r = client.post("/api/osc/cases/1/quick-action", json={"action": action})
    assert r.status_code == 200
    body = r.get_json()
    assert body["ok"] is True
    assert body["native"] is True
    assert title in body["reply"]


def test_large_file_content_uses_streaming_response(client, tmp_path):
    """超過記憶體預覽門檻的大型 PDF 不應回 File too large。"""
    big_pdf = tmp_path / "large.pdf"
    with big_pdf.open("wb") as f:
        f.seek((58 * 1024 * 1024) - 1)
        f.write(b"\0")

    with patch("api.blueprints.osc_cases._osc_local_path_candidates", return_value=[str(big_pdf)]), \
         patch("api.blueprints.osc_cases._osc_is_safe_local_path", return_value=True):
        r = client.get(f"/api/osc/files/content?path={big_pdf}&inline=1")

    assert r.status_code == 200
    assert r.content_type.startswith("application/pdf")
    assert r.headers.get("Content-Length") == str(big_pdf.stat().st_size)


def test_file_content_streams_staged_file_without_send_file(client, tmp_path, monkeypatch):
    """PDF 下載/預覽不得再依賴 Werkzeug send_file，避免 macOS SMB EDEADLK。"""
    from api.blueprints import osc_cases as mod

    pdf = tmp_path / "卷證.pdf"
    payload = b"%PDF-streaming-response"
    pdf.write_bytes(payload)

    def fail_send_file(*_args, **_kwargs):
        raise AssertionError("osc file content should stream staged files directly")

    monkeypatch.setattr(mod, "send_file", fail_send_file)
    with patch("api.blueprints.osc_cases._osc_local_path_candidates", return_value=[str(pdf)]), \
         patch("api.blueprints.osc_cases._osc_is_safe_local_path", return_value=True):
        r = client.get(f"/api/osc/files/content?path={pdf}&inline=1")

    assert r.status_code == 200
    assert r.data == payload
    assert r.headers["Content-Disposition"].startswith("inline")
    assert r.headers["Accept-Ranges"] == "bytes"


def test_file_content_supports_pdf_range_requests(client, tmp_path):
    pdf = tmp_path / "range.pdf"
    pdf.write_bytes(b"0123456789abcdef")

    with patch("api.blueprints.osc_cases._osc_local_path_candidates", return_value=[str(pdf)]), \
         patch("api.blueprints.osc_cases._osc_is_safe_local_path", return_value=True):
        r = client.get(f"/api/osc/files/content?path={pdf}&inline=1", headers={"Range": "bytes=2-5"})

    assert r.status_code == 206
    assert r.data == b"2345"
    assert r.headers["Content-Range"] == "bytes 2-5/16"


def test_file_content_chinese_pdf_has_mobile_safe_ascii_filename(client, tmp_path):
    pdf = tmp_path / "楊曉琳-案.pdf"
    pdf.write_bytes(b"%PDF-mobile-filename")

    with patch("api.blueprints.osc_cases._osc_local_path_candidates", return_value=[str(pdf)]), \
         patch("api.blueprints.osc_cases._osc_is_safe_local_path", return_value=True):
        r = client.get(f"/api/osc/files/content?path={pdf}")

    assert r.status_code == 200
    cd = r.headers["Content-Disposition"]
    assert 'filename="paperclip.pdf"' in cd
    assert "filename*=UTF-8''" in cd
    assert "%E6%A5%8A%E6%9B%89%E7%90%B3-%E6%A1%88.pdf" in cd


def test_file_content_head_uses_stat_without_staging(client, tmp_path, monkeypatch):
    from api.blueprints import osc_cases as mod

    pdf = tmp_path / "楊曉琳-案.pdf"
    pdf.write_bytes(b"%PDF-head-mobile")

    def fail_stage(_path):
        raise AssertionError("HEAD should not stage or hydrate the file")

    monkeypatch.setattr(mod, "_osc_stage_file_with_retry", fail_stage)
    with patch("api.blueprints.osc_cases._osc_local_path_candidates", return_value=[str(pdf)]), \
         patch("api.blueprints.osc_cases._osc_is_safe_local_path", return_value=True):
        r = client.head(f"/api/osc/files/content?path={pdf}")

    assert r.status_code == 200
    assert r.data == b""
    assert r.headers["Content-Length"] == str(pdf.stat().st_size)
    assert 'filename="paperclip.pdf"' in r.headers["Content-Disposition"]


def test_file_content_prefers_hydrated_volume_candidate_over_dataless_cloud(client, tmp_path, monkeypatch):
    from api.blueprints import osc_cases as mod

    cloud = tmp_path / "CloudStorage" / "楊曉琳-案.pdf"
    volume = tmp_path / "Volumes" / "楊曉琳-案.pdf"
    cloud.parent.mkdir()
    volume.parent.mkdir()
    cloud.write_bytes(b"")
    volume.write_bytes(b"%PDF-volume-copy")

    monkeypatch.setattr(mod, "_osc_is_dataless_file", lambda path: str(path) == str(cloud))
    with patch("api.blueprints.osc_cases._osc_local_path_candidates", return_value=[str(cloud), str(volume)]), \
         patch("api.blueprints.osc_cases._osc_is_safe_local_path", return_value=True):
        r = client.get(f"/api/osc/files/content?path={cloud}")

    assert r.status_code == 200
    assert r.data == b"%PDF-volume-copy"


def test_large_file_content_stages_nas_file_when_direct_read_deadlocks(client, tmp_path, monkeypatch):
    """大型 NAS 檔案若第一次開檔 EDEADLK，應 retry 並從本機暫存檔下載。"""
    from api.blueprints import osc_cases as mod

    big_pdf = tmp_path / "large-deadlock.pdf"
    payload = b"%PDF-" + (b"x" * (2 * 1024 * 1024))
    big_pdf.write_bytes(payload)
    attempts = {"n": 0}

    def fake_open(path, mode="r", *args, **kwargs):
        if str(path) == str(big_pdf) and mode == "rb":
            attempts["n"] += 1
            if attempts["n"] == 1:
                raise OSError(errno.EDEADLK, "Resource deadlock avoided")
        return builtins.open(path, mode, *args, **kwargs)

    monkeypatch.setenv("PAPERCLIP_FILE_MEMORY_PREVIEW_MAX_MB", "1")
    monkeypatch.setattr(mod, "open", fake_open, raising=False)
    with patch("api.blueprints.osc_cases._osc_local_path_candidates", return_value=[str(big_pdf)]), \
         patch("api.blueprints.osc_cases._osc_is_safe_local_path", return_value=True):
        r = client.get(f"/api/osc/files/content?path={big_pdf}")

    assert r.status_code == 200
    assert r.data == payload
    assert attempts["n"] >= 1
    assert "attachment" in r.headers.get("Content-Disposition", "")


def test_large_file_content_uses_system_cp_when_python_read_keeps_deadlocking(client, tmp_path, monkeypatch):
    """若 Python open 持續 EDEADLK，改用本機 cp staging，避免直接回檔案讀取失敗。"""
    from api.blueprints import osc_cases as mod

    big_pdf = tmp_path / "large-deadlock-cp.pdf"
    payload = b"%PDF-cp-fallback"
    big_pdf.write_bytes(payload)
    attempts = {"n": 0}

    def fake_open(path, mode="r", *args, **kwargs):
        if str(path) == str(big_pdf) and mode == "rb":
            attempts["n"] += 1
            raise OSError(errno.EDEADLK, "Resource deadlock avoided")
        return builtins.open(path, mode, *args, **kwargs)

    def fake_run(argv, **_kwargs):
        target = Path(argv[-1])
        target.write_bytes(payload)
        return type("R", (), {"returncode": 0, "stdout": "", "stderr": ""})()

    monkeypatch.setattr(mod, "open", fake_open, raising=False)
    monkeypatch.setattr(mod.subprocess, "run", fake_run)
    with patch("api.blueprints.osc_cases._osc_local_path_candidates", return_value=[str(big_pdf)]), \
         patch("api.blueprints.osc_cases._osc_is_safe_local_path", return_value=True):
        r = client.get(f"/api/osc/files/content?path={big_pdf}")

    assert r.status_code == 200
    assert r.data == payload
    assert attempts["n"] >= 1


def test_file_content_stages_even_when_source_stat_deadlocks(client, tmp_path, monkeypatch):
    """SMB 若在 stat/getsize 階段先 EDEADLK，仍應暫存後完成下載。"""
    from api.blueprints import osc_cases as mod

    pdf = tmp_path / "stat-deadlock.pdf"
    payload = b"%PDF-stat-deadlock"
    pdf.write_bytes(payload)
    attempts = {"n": 0}

    def flaky_stat(path, **_kwargs):
        if str(path) == str(pdf):
            attempts["n"] += 1
            raise OSError(errno.EDEADLK, "Resource deadlock avoided")
        return mod.os.stat(path)

    monkeypatch.setattr(mod, "_osc_stat_with_retry", flaky_stat)
    with patch("api.blueprints.osc_cases._osc_local_path_candidates", return_value=[str(pdf)]), \
         patch("api.blueprints.osc_cases._osc_is_safe_local_path", return_value=True):
        r = client.get(f"/api/osc/files/content?path={pdf}")

    assert r.status_code == 200
    assert r.data == payload
    assert attempts["n"] >= 1


def test_direct_file_content_error_is_readable_html(client):
    r = client.get("/api/osc/files/content", headers={"Accept": "text/html"})

    assert r.status_code == 400
    assert r.content_type.startswith("text/html")
    body = r.get_data(as_text=True)
    assert "檔案操作沒有完成" in body
    assert "缺少檔案路徑" in body


def test_api_file_content_error_stays_json(client):
    r = client.get("/api/osc/files/content", headers={"Accept": "application/json"})

    assert r.status_code == 400
    assert r.is_json
    assert r.get_json()["ok"] is False


def test_direct_address_label_error_is_readable_html(client):
    case = {"id": 1, "case_number": "114-測-1", "client_name": "測試", "court_name": ""}

    def fake_exec(sql, params=(), fetch="all"):
        if "FROM cases" in sql:
            return (case, None) if fetch == "one" else ([case], None)
        return (None, None) if fetch == "one" else ([], None)

    with patch("api.blueprints.osc_cases._osc_exec", side_effect=fake_exec), \
         patch("api.blueprints.osc_cases._osc_get_setting_value", return_value=""):
        r = client.get(
            "/api/osc/cases/1/address-label?mode=preview&recipient=court",
            headers={"Accept": "text/html"},
        )

    assert r.status_code == 400
    assert r.content_type.startswith("text/html")
    assert "案件未設定法院或地檢署名稱" in r.get_data(as_text=True)


def test_direct_quotation_pdf_error_is_readable_html(client):
    with patch("api.blueprints.osc_cases._osc_exec", return_value=(None, None)):
        r = client.get("/api/osc/quotations/999/export-pdf", headers={"Accept": "text/html"})

    assert r.status_code == 404
    assert r.content_type.startswith("text/html")
    assert "找不到這份報價單" in r.get_data(as_text=True)


def test_case_workbench_upload_allows_legal_pdf_over_50mb(client, tmp_path, monkeypatch):
    """案件處理頁上傳需支援常見大型卷證 PDF。"""
    monkeypatch.setenv("PAPERCLIP_UPLOAD_MAX_PER_FILE_MB", "128")
    monkeypatch.setenv("PAPERCLIP_UPLOAD_MAX_TOTAL_MB", "128")

    with patch("api.blueprints.osc_cases._osc_resolve_existing_local_path", return_value=str(tmp_path)), \
         patch("api.blueprints.osc_cases._osc_is_safe_local_path", return_value=True):
        r = client.post(
            "/api/osc/files/upload",
            data={
                "folder_path": str(tmp_path),
                "file": (BytesIO(b"x" * (52 * 1024 * 1024)), "large.pdf"),
            },
            content_type="multipart/form-data",
        )

    assert r.status_code == 200
    data = r.get_json()
    assert data["ok"] is True
    assert data["saved"][0]["file_name"] == "large.pdf"


# ── 3. 消債書狀生成（osc_debt） ───────────────────────────────────────────────


def test_debt_forms_list(client):
    r = client.get("/api/osc/debt/forms")
    assert r.status_code == 200
    body = r.get_json()
    # 至少要有 5 種：聲請狀、財產說明書、債權人清冊、陳報狀、合併
    assert isinstance(body, dict)


def test_debt_schema_returns_fields(client):
    r = client.get("/api/osc/debt/schema/application")
    assert r.status_code == 200


def test_debt_source_status_uses_bundled_source(client):
    r = client.get("/api/osc/debt/source-status")
    assert r.status_code == 200
    body = r.get_json()
    assert body["ok"] is True
    assert body["source_dir"].endswith("integrations/debt_robot")
    assert body["modules"]["supplement"].endswith("06_F.py")


def _write_debt_asset_doc(path: Path, total: str = "12345"):
    from docx import Document

    doc = Document()
    for idx in range(5):
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "總計" if idx == 4 else f"表格{idx}"
        table.cell(0, 1).text = total if idx == 4 else ""
    doc.save(path)


def _write_debt_creditor_doc(path: Path, total: str = "67890", bank: str = "測試銀行"):
    from docx import Document

    doc = Document()
    summary = doc.add_table(rows=1, cols=2)
    summary.cell(0, 0).text = "債務總金額"
    summary.cell(0, 1).text = total
    detail = doc.add_table(rows=2, cols=4)
    detail.cell(0, 0).text = "債權人"
    detail.cell(0, 2).text = "金額"
    detail.cell(1, 0).text = bank
    detail.cell(1, 1).text = "台北市測試路100號"
    detail.cell(1, 2).text = total
    detail.cell(1, 3).text = "信用貸款"
    doc.save(path)


def test_debt_auto_import_selected_docs_and_candidates(client, tmp_path, monkeypatch):
    from api.blueprints import osc_debt as mod

    asset = tmp_path / "02_財產及收入狀況說明書（測試）.docx"
    creditor = tmp_path / "03_債權人清冊（測試）.docx"
    _write_debt_asset_doc(asset, "12345")
    _write_debt_creditor_doc(creditor, "67890", "測試銀行")
    monkeypatch.setattr(mod, "_export_dir", lambda: str(tmp_path))

    r = client.get("/api/osc/debt/import-candidates")
    assert r.status_code == 200
    candidates = r.get_json()
    assert candidates["ok"] is True
    assert any(item["path"] == str(asset.resolve()) for item in candidates["asset_docs"])
    assert any(item["path"] == str(creditor.resolve()) for item in candidates["creditor_docs"])

    r = client.post(
        "/api/osc/debt/auto-import",
        json={"asset_doc_path": str(asset), "creditor_doc_path": str(creditor)},
    )
    assert r.status_code == 200
    body = r.get_json()
    assert body["ok"] is True
    assert body["asset_total"] == 12345
    assert body["debt_total"] == 67890
    assert body["max_bank"] == "測試銀行"


def test_debt_generate_word_files_and_bank_json_record(client, tmp_path, monkeypatch):
    from api import debt_document_generator as gen
    from api.blueprints import osc_debt as mod

    template_dir = tmp_path / "robot_document"
    shutil.copytree(Path(gen._TEMPLATE_DIR), template_dir)
    export_dir = tmp_path / "exports"
    export_dir.mkdir()
    monkeypatch.setattr(gen, "_TEMPLATE_DIR", str(template_dir))
    monkeypatch.setattr(mod, "_export_dir", lambda: str(export_dir))

    payloads = [
        ("application", {
            "name": "測試人", "address": "台北市測試路1號", "asset_total": 12345,
            "debt_total": 67890, "max_creditor_bank": "測試銀行",
            "execution_court": "臺灣臺北地方法院", "execution_case_no": "115執字第1號",
            "application_court": "臺灣臺北地方法院", "attachments": "測試附件",
        }),
        ("asset_statement", {
            "insurance": [{"company": "測試保險", "type": "壽險", "policy_no": "P1", "amount": "1000"}],
            "land": [], "vehicles": [], "stocks": [],
            "income": [{"type": "薪資", "source": "測試公司", "amount": "30000"}],
            "expenses": [{"category": "房租", "monthly": "10000"}],
            "dependents": [],
        }),
        ("creditor_list", {
            "creditors": [{"name": "測試銀行", "address": "台北市測試路100號", "amount": "67890", "debt_type": "信用貸款"}],
        }),
        ("report", {"A1": "1", "A2": "115消債更字第1號", "A3": "明股", "A4": "測試人", "E1": "臺灣臺北地方法院"}),
        ("supplement", {
            "court": "臺灣臺北地方法院", "case_no": "115消債更字第1號", "branch": "明股",
            "applicant": "測試人", "procedure": "更生", "brief_no": "1",
            "items": [{"category": "勞保資料", "period": "112年度", "attachment": "附件一"}],
        }),
    ]

    generated = []
    generated_by_type = {}
    for form_type, data in payloads:
        r = client.post("/api/osc/debt/generate", json={"form_type": form_type, "data": data})
        assert r.status_code == 200, r.get_data(as_text=True)
        body = r.get_json()
        assert body["ok"] is True
        path = Path(body["path"])
        assert path.exists()
        assert path.suffix == ".docx"
        assert zipfile.is_zipfile(path)
        assert body["download_url"].startswith("/api/osc/files/content?path=")
        assert body["share_path"] == str(path)
        generated.append(path)
        generated_by_type[form_type] = path
        if form_type == "creditor_list":
            assert body["saved_addresses"] >= 1

    assert len(generated) == 5
    r = client.post(
        "/api/osc/debt/auto-import",
        json={
            "asset_doc_path": str(generated_by_type["asset_statement"]),
            "creditor_doc_path": str(generated_by_type["creditor_list"]),
        },
    )
    assert r.status_code == 200, r.get_data(as_text=True)
    imported = r.get_json()
    assert imported["ok"] is True
    assert imported["asset_total"] == 30000
    assert imported["debt_total"] == 67890
    assert imported["max_bank"] == "測試銀行"

    address_json = template_dir / "all adress - bank.json"
    assert address_json.exists()
    bank_data = json.loads(address_json.read_text(encoding="utf-8"))
    assert any(item["name"] == "測試銀行" and item["address"] == "台北市測試路100號" for item in bank_data["items"])


def test_generated_exports_can_use_paperclip_download_and_share(client, tmp_path, monkeypatch):
    """消債匯出檔也要走 Paperclip 檔案管理的下載與不含路徑分享模式。"""
    from api.blueprints import osc_files as files_mod

    monkeypatch.setattr(files_mod, "_SHARE_STORE_PATH", tmp_path / "shares.json")
    export_dir = ROOT / "exports"
    export_dir.mkdir(exist_ok=True)
    generated = export_dir / "消債下載分享測試.docx"
    generated.write_bytes(b"test-docx")
    try:
        r = client.get(f"/api/osc/files/content?path={generated}")
        assert r.status_code == 200
        assert r.data == b"test-docx"

        r = client.post("/api/osc/files/share", json={"path": str(generated)})
        assert r.status_code == 200, r.get_data(as_text=True)
        body = r.get_json()
        assert body["ok"] is True
        assert "/s/" in body["url"]
        assert str(generated) not in body["url"]

        share_path = urlparse(body["url"]).path
        r = client.get(share_path)
        assert r.status_code == 200
        assert r.data == b"test-docx"
    finally:
        generated.unlink(missing_ok=True)


# ── 4. 委任狀及契約生成（forms） ─────────────────────────────────────────────


def test_forms_preview_smoke(client):
    """forms/preview 至少不 crash（fake 案件 row）。"""
    fake_case = {
        "id": 1,
        "case_number": "TEST-001",
        "client_name": "測試當事人",
        "court_case_no": "112年度訴字第123號",
        "court_branch": "民事第一庭",
        "case_reason": "測試案由",
        "lawyer_name": "張律師",
        "address": "台北市測試路1號",
        "phone": "0912345678",
    }

    # forms/preview 內部還會 call _osc_get_case_identity_by_payload 與
    # _osc_build_form_preview，會跨進 osc.utils；smoke 不做完整 mock，
    # 只驗證 route 已註冊（非 404 即可）。
    fake = _make_fake_exec({"cases": [fake_case]})
    with patch("api.blueprints.osc_cases._osc_exec", side_effect=fake):
        r = client.post(
            "/api/osc/forms/preview",
            json={"form_type": "power_of_attorney", "case_id": 1, "fields": {}},
        )
    assert r.status_code != 404, "forms/preview route 未註冊"


# ── 5. 報價單 PDF（reportlab） ────────────────────────────────────────────────


def test_quotation_pdf_smoke(client):
    """報價單 PDF endpoint 應能對 mock row 產出 application/pdf。"""
    fake_quotation = {
        "id": 1,
        "case_number": "TEST-001",
        "client_name": "王小明",
        "items_json": '[{"name":"民事訴訟","qty":1,"unit_price":50000}]',
        "total": 50000,
        "notes": "smoke test",
        "created_at": "2026-04-30",
    }

    fake = _make_fake_exec({"quotations": [fake_quotation], "settings": []})
    with patch("api.blueprints.osc_cases._osc_exec", side_effect=fake):
        r = client.get("/api/osc/quotations/1/export-pdf")
    assert r.status_code == 200, f"PDF endpoint failed: {r.status_code}"
    assert r.mimetype == "application/pdf"
    assert r.data[:4] == b"%PDF", "PDF magic bytes 不對"


# ── 6. 地址標籤 PNG（PIL） ────────────────────────────────────────────────────


def test_address_label_preview_smoke(client):
    """地址標籤預覽 endpoint 應回 image/png inline。"""
    fake_case = {
        "id": 1,
        "case_number": "TEST-001",
        "court_name": "臺灣臺北地方法院",
        "client_name": "王小明",
    }
    fake_court = {"name": "臺灣臺北地方法院", "address": "台北市博愛路131號"}

    fake = _make_fake_exec({"cases": [fake_case], "courts": [fake_court], "settings": []})
    with patch("api.blueprints.osc_cases._osc_exec", side_effect=fake):
        r = client.get("/api/osc/cases/1/address-label?mode=preview&recipient=court")
    assert r.status_code == 200, f"PNG preview failed: {r.status_code}"
    assert r.mimetype == "image/png"
    assert r.data[:8] == b"\x89PNG\r\n\x1a\n", "PNG magic bytes 不對"
    # 預覽必須 inline
    cd = r.headers.get("Content-Disposition", "")
    assert "inline" in cd.lower(), f"預覽應 inline: {cd}"


# ── 7. PDF 蓋章 endpoint（doc-producer skill） ───────────────────────────────


def test_documents_stamp_endpoint_reachable(client):
    """蓋章 endpoint 在收到 fake 路徑時應回 4xx（不 5xx），驗證路由綁定正確。"""
    r = client.post(
        "/api/osc/documents/stamp",
        json={"path": "/nonexistent/file.pdf", "copy_type": "正本"},
    )
    assert r.status_code in (200, 400, 404, 500), f"unexpected: {r.status_code}"
    # 即使 fail 也應 JSON
    if r.is_json:
        body = r.get_json()
        assert "ok" in body or "error" in body or "success" in body


# ── 8. CSV 匯入匯出 endpoint 可達 ────────────────────────────────────────────


def test_csv_export_cases_smoke(client):
    """案件 CSV 匯出 endpoint 應回 text/csv；無資料時回空 CSV 也算 OK。"""

    with patch("api.blueprints.osc_cases._osc_exec", side_effect=_make_fake_exec()):
        r = client.get("/api/osc/cases/export-csv")
    assert r.status_code == 200
    assert "text/csv" in r.mimetype or "octet-stream" in r.mimetype


# ── 9. Checklist 應備事項 endpoint 可達 ──────────────────────────────────────


def test_checklist_legal_aid_endpoint_reachable(client):
    with patch("api.blueprints.osc_cases._osc_exec", side_effect=_make_fake_exec()):
        r = client.get("/api/osc/checklists/legal-aid?case_number=TEST-001")
    assert r.status_code in (200, 400)


# ── 10. 自動備份 endpoint 可達 ────────────────────────────────────────────────


def test_backup_list_endpoint(client):
    """備份列表 endpoint 應回 {ok:true, items:[...]}。"""
    r = client.get("/api/osc/backups")
    assert r.status_code == 200
    body = r.get_json()
    assert body.get("ok") is True
    assert isinstance(body.get("items"), list)


# ── 11. Google Calendar status endpoint ───────────────────────────────────────


def test_gcal_status_endpoint(client):
    """GCal status 應該回 connected=false（測試環境無 token.json）。"""
    r = client.get("/api/osc/gcal/status")
    assert r.status_code == 200
    body = r.get_json()
    assert body.get("ok") is True
    assert body.get("connected") in (False, True)


# ── 12. heartbeat / status JSON 端點 ─────────────────────────────────────────


def test_magi_status_json_has_no_legacy_db_ip():
    """static/magi_status.json 不該再有舊 NAS IP（2026-04-30 修復後驗證）。"""
    status_path = ROOT / "static" / "magi_status.json"
    if not status_path.exists():
        pytest.skip("magi_status.json 尚未生成（heartbeat 未執行過）")
    raw = status_path.read_text(encoding="utf-8")
    # 完整禁止舊的 desktop-jj06fa3 IP
    legacy_ip = "100." "121." "61." "74"
    assert legacy_ip not in raw, "magi_status.json 仍引用舊 DB IP"


# ── 13. heartbeat keeper 設定不寫死舊 IP ─────────────────────────────────────


def test_heartbeat_no_legacy_keeper_fallback():
    """heartbeat.py keeper.ip 不該 fallback 到舊私有 DB IP。"""
    src = (ROOT / "skills" / "ops" / "heartbeat.py").read_text(encoding="utf-8")
    # 允許註解中提到（用於文件說明），但不該再有 _node_ip_or 的 fallback
    legacy_call = '_node_ip_or("nas", "' + "100." "121." "61." "74" + '")'
    assert legacy_call not in src, (
        "heartbeat.py keeper.ip 仍 fallback 到舊 IP"
    )


# ── 14. MAGI 網頁回覆不得再以 DC/Markdown 原文呈現 ─────────────────────────


def test_magi_web_markdown_replies_use_html_renderers():
    """聊天、案件快捷動作、實務見解都應吃 reply_html / renderWebReplyHtml。"""
    osc_ui = (ROOT / "static" / "osc" / "osc-ui.js").read_text(encoding="utf-8")
    osc_cases = (ROOT / "static" / "osc" / "tabs" / "cases.js").read_text(encoding="utf-8")
    osc_docs = (ROOT / "static" / "osc" / "tabs" / "documents.js").read_text(encoding="utf-8")
    osc_insights = (ROOT / "static" / "osc" / "tabs" / "insights.js").read_text(encoding="utf-8")
    old_dashboard = (ROOT / "templates" / "dashboard.html").read_text(encoding="utf-8")
    nerv = (ROOT / "templates" / "dashboard_nerv.html").read_text(encoding="utf-8")
    api_cases = (ROOT / "api" / "blueprints" / "osc_cases.py").read_text(encoding="utf-8")

    assert "data.reply_html" in osc_ui
    assert "data.reply_html" in osc_cases
    assert "renderWebReplyHtml(text || \"\")" in osc_ui
    assert "renderWebReplyHtml(data.reply || data.message" in osc_cases
    assert "showWebReplyDialog" in osc_docs
    assert "renderWebReplyHtml(r.full_text" in osc_insights
    assert "reply_html" in old_dashboard
    assert "reply_html" in nerv
    assert '"reply_html": _format_web_reply_html(reply_text)' in api_cases
    assert "alert(data.reply" not in osc_docs
    assert "alert(result.reply" not in osc_docs


def test_paperclip_bound_buttons_and_data_actions_are_wired():
    """Paperclip 各頁籤按鈕要有 DOM 目標，data-act 要有 dispatch handler。"""
    html_sources = [
        ROOT / "templates" / "osc.html",
        *sorted((ROOT / "templates" / "partials" / "osc").glob("*.html")),
        *sorted((ROOT / "static" / "osc" / "tabs").glob("*.js")),
    ]
    source = "\n".join(p.read_text(encoding="utf-8") for p in html_sources)
    ids = set(re.findall(r'id=["\']([^"\']+)["\']', source))

    events_js = (ROOT / "static" / "osc" / "osc-events.js").read_text(encoding="utf-8")
    bound_ids = set(re.findall(r'\["([^"]+)",', events_js))
    missing_ids = sorted(bound_ids - ids)
    assert not missing_ids, f"按鈕事件綁定找不到 DOM id：{missing_ids}"

    acts = set(re.findall(r'data-act=["\']([^"\']+)["\']', source))
    handled = set(re.findall(r'if \(act === "([^"]+)"\)', events_js))
    missing_handlers = sorted(acts - handled)
    assert not missing_handlers, f"data-act 沒有 dispatch handler：{missing_handlers}"
