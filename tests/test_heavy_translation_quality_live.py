from __future__ import annotations

import hashlib
import os

import pytest
from docx import Document
from pypdf import PdfReader

from scripts.ops import heavy_translation_quality_live as gate


def test_live_route_probe_requires_correct_taiwan_term() -> None:
    class Gateway:
        def __init__(self, response: str):
            self.response = response

        def chat(self, *_args, **_kwargs):
            return {
                "success": True,
                "response": self.response,
                "route": "nvidia_nim",
                "model": "synthetic-model",
                "provider": "synthetic-provider",
            }

    good = gate._run_nim_route_check(5, gateway=Gateway("司法通譯"))
    bad = gate._run_nim_route_check(5, gateway=Gateway("法庭翻譯"))

    assert good["semantic_quality_passed"] is True
    assert bad["semantic_quality_passed"] is False
    assert "response" not in good


def test_generated_heavy_translation_fixture_is_extractable(tmp_path):
    path = gate.write_generated_fixture(tmp_path / "fixture.pdf")

    reader = PdfReader(str(path))
    texts = [(page.extract_text() or "") for page in reader.pages]

    assert len(reader.pages) == 100
    assert "司法通譯語言風格如何影響國民法官對被告的印象" in texts[0]
    assert "國民法官法" in "".join(texts[3:5])
    assert "court interpreters" in "".join(texts[5:7])


def test_gate_accepts_verified_taiwan_renderings_without_forcing_english_into_target(
    tmp_path,
):
    fixture = gate.write_generated_fixture(tmp_path / "fixture.pdf")
    result = gate.run_gate(
        pdf_path=fixture,
        run_live_nim=False,
        timeout=1,
        title_translator=lambda _text: {
            "success": True,
            "response": "司法通譯語言風格如何影響國民法官對被告的印象",
            "route": "nvidia_nim",
            "provider": "synthetic-test-provider",
            "model": "synthetic-test-model",
        },
    )

    checks = {item["name"]: item for item in result["checks"]}
    assert checks["tw_term_normalization"]["ok"] is True
    assert checks["source_terms_inline"]["ok"] is True
    assert checks["source_terms_inline"]["detail"] == "approved Taiwan renderings present"


def test_run_gate_does_not_read_directory_when_docx_export_fails(
    monkeypatch,
    tmp_path,
):
    fixture = tmp_path / "fixture.pdf"
    fixture.write_bytes(b"synthetic fixture marker")
    monkeypatch.setattr(
        gate,
        "_extract_fixture",
        lambda _path: {
            "pages": "100",
            "title": "司法通譯語言風格如何影響國民法官對被告的印象",
            "zh_abstract": "國民法官法 司法通譯 無力風格 假冒配對測試法 被告 國民法官 有力風格",
            "en_abstract": (
                "Citizen Judges Act court interpreters defendant powerless style "
                "Powerless Group Powerful Group matched guise technique"
            ),
        },
    )
    monkeypatch.setattr(
        gate,
        "_read_docx_text",
        lambda _path: pytest.fail("directory/non-file DOCX must never be read"),
    )
    from skills.ops import export_docx

    monkeypatch.setattr(
        export_docx,
        "export_bilingual_docx",
        lambda *args, **kwargs: {"success": False, "error": "node module missing"},
    )

    result = gate.run_gate(
        pdf_path=fixture,
        run_live_nim=False,
        timeout=1,
        title_translator=lambda _text: {
            "success": True,
            "response": "司法通譯語言風格如何影響國民法官對被告的印象",
            "route": "nvidia_nim",
            "provider": "synthetic-test-provider",
            "model": "synthetic-test-model",
        },
    )

    assert result["docx_path"] == ""
    assert next(
        check for check in result["checks"] if check["name"] == "docx_export"
    )["ok"] is False


@pytest.mark.parametrize("failure_mode", ["missing", "swapped", "unreadable"])
def test_run_gate_fails_closed_when_exported_docx_cannot_be_stably_read(
    monkeypatch,
    tmp_path,
    failure_mode,
):
    fixture = tmp_path / "fixture.pdf"
    fixture.write_bytes(b"synthetic fixture marker")
    monkeypatch.setattr(
        gate,
        "_extract_fixture",
        lambda _path: {
            "pages": "100",
            "title": "司法通譯語言風格如何影響國民法官對被告的印象",
            "zh_abstract": "國民法官法 司法通譯 無力風格 假冒配對測試法 被告 國民法官 有力風格",
            "en_abstract": (
                "Citizen Judges Act court interpreters defendant powerless style "
                "Powerless Group Powerful Group matched guise technique"
            ),
        },
    )
    target = (tmp_path / "exported.docx").resolve()
    document = Document()
    document.add_paragraph("original export")
    document.save(target)
    expected_sha256 = hashlib.sha256(target.read_bytes()).hexdigest()
    if failure_mode == "missing":
        target.unlink()
    elif failure_mode == "swapped":
        replacement = tmp_path / "replacement.docx"
        replacement_document = Document()
        replacement_document.add_paragraph("replacement export")
        replacement_document.save(replacement)
        os.replace(replacement, target)
    else:
        target.write_bytes(b"not a docx package")
        expected_sha256 = hashlib.sha256(target.read_bytes()).hexdigest()

    from skills.ops import export_docx

    monkeypatch.setattr(
        export_docx,
        "export_bilingual_docx",
        lambda *args, **kwargs: {
            "success": True,
            "path": str(target),
            "validation": {"sha256": expected_sha256},
        },
    )
    result = gate.run_gate(
        pdf_path=fixture,
        run_live_nim=False,
        timeout=1,
        title_translator=lambda _text: {
            "success": True,
            "response": "司法通譯語言風格如何影響國民法官對被告的印象",
            "route": "nvidia_nim",
            "provider": "synthetic-test-provider",
            "model": "synthetic-test-model",
        },
    )

    checks = {check["name"]: check for check in result["checks"]}
    assert result["success"] is False
    assert result["docx_path"] == ""
    assert checks["docx_export"]["ok"] is False
    assert checks["docx_readback"]["ok"] is False
