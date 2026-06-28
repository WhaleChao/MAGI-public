from pathlib import Path
from unittest.mock import patch

from docx import Document
from flask import Flask
from flask_login import LoginManager, UserMixin


ROOT = Path(__file__).resolve().parent.parent


def _build_app() -> Flask:
    app = Flask(__name__, template_folder=str(ROOT / "templates"))
    app.config["TESTING"] = True
    app.config["LOGIN_DISABLED"] = True
    app.secret_key = "test"
    login = LoginManager()
    login.init_app(app)

    class _User(UserMixin):
        id = "test-user"

    @login.user_loader
    def _load_user(_user_id):
        return _User()

    from api.blueprints.osc_cases import osc_bp

    app.register_blueprint(osc_bp)
    return app


def _make_source_docx(path: Path) -> None:
    doc = Document()
    header = doc.sections[0].header.paragraphs[0]
    header.add_run("臺灣")
    header.add_run("臺北")
    header.add_run("地方法院")
    p = doc.add_paragraph()
    p.add_run("案號：113年度")
    p.add_run("訴字第1號")
    p.add_run(" 股別：義股")
    doc.add_paragraph("內部案號：SRC-001")
    doc.add_paragraph("案由：損害賠償")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "原告"
    table.cell(0, 1).text = "王小明"
    table.cell(1, 0).text = "被告"
    table.cell(1, 1).text = "李大華"
    doc.sections[0].footer.paragraphs[0].text = "此致 臺灣臺北地方法院"
    doc.save(path)


def _doc_text(path: Path) -> str:
    doc = Document(path)
    parts = []

    def collect(container):
        parts.extend(p.text for p in container.paragraphs)
        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    collect(cell)

    collect(doc)
    for section in doc.sections:
        collect(section.header)
        collect(section.footer)
    return "\n".join(parts)


def test_reuse_document_api_creates_target_case_word_doc_and_logs(tmp_path):
    app = _build_app()
    client = app.test_client()
    source_dir = tmp_path / "source"
    target_case_dir = tmp_path / "target-case"
    source_dir.mkdir()
    target_case_dir.mkdir()
    source = source_dir / "舊案民事準備書狀.doc"
    converted_source = source_dir / "converted-source.docx"
    source.write_bytes(b"legacy-doc")
    _make_source_docx(converted_source)

    calls = []
    source_case = {
        "id": "1",
        "case_number": "SRC-001",
        "client_name": "王小明",
        "opponent_name": "李大華",
        "court_name": "臺灣臺北地方法院",
        "court_case_no": "113年度訴字第1號",
        "court_case_number": "113年度訴字第1號",
        "court_division": "義股",
        "case_reason": "損害賠償",
        "folder_path": str(source_dir),
    }
    target_case = {
        "id": "2",
        "case_number": "TGT-002",
        "client_name": "陳新明",
        "opponent_name": "林新華",
        "court_name": "臺灣新北地方法院",
        "court_case_no": "115年度重訴字第9號",
        "court_case_number": "115年度重訴字第9號",
        "court_division": "忠股",
        "case_reason": "返還借款",
        "folder_path": str(target_case_dir),
    }

    def fake_exec(sql, params=(), fetch="none"):
        calls.append((sql, params, fetch))
        lowered = " ".join(str(sql).lower().split())
        if "from cases" in lowered and fetch == "one":
            if params and str(params[0]) == "2":
                return target_case, {}
            if params and str(params[0]) == "SRC-001":
                return source_case, {}
        if "from opponents" in lowered and fetch == "all":
            return [], {}
        if fetch == "all":
            return [], {}
        if fetch == "one":
            return None, {}
        return {"rowcount": 1, "lastrowid": 99}, {}

    def resolve_existing(path, prefer_dir=None):
        p = Path(str(path))
        return str(p) if p.exists() else ""

    with patch("api.blueprints.osc_cases._osc_exec", side_effect=fake_exec), patch(
        "api.blueprints.osc_cases._osc_resolve_existing_local_path", side_effect=resolve_existing
    ), patch(
        "api.osc.document_reuse._convert_doc_to_docx", return_value=converted_source
    ):
        response = client.post(
            "/api/osc/drafts/reuse-document",
            json={
                "case_id": "2",
                "case_number": "115年度重訴字第9號",
                "division": "忠股",
                "court_name": "臺灣新北地方法院",
                "reason": "返還借款",
                "plaintiff": "陳新明",
                "defendant": "林新華",
                "source_path": str(source),
                "source_case_number": "SRC-001",
                "suggested_filename": "新案準備書狀.docx",
            },
        )

    assert response.status_code == 200, response.get_data(as_text=True)
    payload = response.get_json()
    assert payload["ok"] is True
    result = payload["result"]
    output = Path(result["output_path"])
    assert output.exists()
    assert output.parent == target_case_dir / "04_我方歷次書狀"
    text = _doc_text(output)
    for expected in ("115年度重訴字第9號", "忠股", "陳新明", "林新華", "臺灣新北地方法院", "返還借款"):
        assert expected in text
    assert "113年度訴字第1號" not in text
    assert "王小明" not in text
    assert any("insert into case_documents" in " ".join(sql.lower().split()) for sql, _, _ in calls)
    assert any("insert into document_replacements" in " ".join(sql.lower().split()) for sql, _, _ in calls)


def test_documents_api_own_pleading_word_scope_excludes_poa_word_files(tmp_path):
    app = _build_app()
    client = app.test_client()
    case_root = tmp_path / "case"
    pleading_dir = case_root / "04_我方歷次書狀"
    poa_dir = case_root / "01_委任狀"
    pleading_dir.mkdir(parents=True)
    poa_dir.mkdir(parents=True)

    def fake_exec(sql, params=(), fetch="none"):
        lowered = " ".join(str(sql).lower().split())
        if "from document_index" in lowered:
            return [
                {
                    "id": 1,
                    "case_number": "SRC-001",
                    "file_name": "民事準備書狀.docx",
                    "file_path": str(pleading_dir / "民事準備書狀.docx"),
                    "subfolder_name": "04_我方歷次書狀",
                    "reason": "損害賠償",
                    "party": "王小明",
                    "modified_date": None,
                },
                {
                    "id": 2,
                    "case_number": "SRC-001",
                    "file_name": "委任狀（可填寫版）.docx",
                    "file_path": str(poa_dir / "委任狀（可填寫版）.docx"),
                    "subfolder_name": "01_委任狀",
                    "reason": "",
                    "party": "王小明",
                    "modified_date": None,
                },
                {
                    "id": 3,
                    "case_number": "SRC-001",
                    "file_name": "民事準備書狀.pdf",
                    "file_path": str(pleading_dir / "民事準備書狀.pdf"),
                    "subfolder_name": "04_我方歷次書狀",
                    "reason": "",
                    "party": "",
                    "modified_date": None,
                },
            ], {}
        if "from case_documents" in lowered:
            return [
                {
                    "id": 4,
                    "case_id": "SRC-001",
                    "case_number_ref": "SRC-001",
                    "document_type": "04_我方歷次書狀",
                    "file_name": "舊式民事聲請狀.doc",
                    "file_path": str(pleading_dir / "舊式民事聲請狀.doc"),
                    "description": "沿用舊書狀",
                    "upload_date": None,
                },
                {
                    "id": 5,
                    "case_id": "SRC-001",
                    "case_number_ref": "SRC-001",
                    "document_type": "01_委任狀",
                    "file_name": "委任狀.docx",
                    "file_path": str(poa_dir / "委任狀.docx"),
                    "description": "",
                    "upload_date": None,
                },
            ], {}
        return [], {}

    with patch("api.blueprints.osc_cases._osc_exec", side_effect=fake_exec):
        response = client.get("/api/osc/documents?kind=own_pleading_word&reuse_scope=own_pleading_word&limit=20")

    assert response.status_code == 200, response.get_data(as_text=True)
    payload = response.get_json()
    assert payload["ok"] is True
    assert [item["file_name"] for item in payload["items"]] == [
        "民事準備書狀.docx",
        "舊式民事聲請狀.doc",
    ]
    assert {item["kind_label"] for item in payload["items"]} == {"書狀"}


def test_document_reuse_ui_is_separate_from_document_index_and_wired_to_api():
    drafts_html = (ROOT / "templates" / "partials" / "osc" / "drafts.html").read_text(encoding="utf-8")
    reuse_html = (ROOT / "templates" / "partials" / "osc" / "documentReuse.html").read_text(encoding="utf-8")
    documents_html = (ROOT / "templates" / "partials" / "osc" / "documents.html").read_text(encoding="utf-8")
    osc_html = (ROOT / "templates" / "osc.html").read_text(encoding="utf-8")
    js = (ROOT / "static" / "osc" / "tabs" / "drafts.js").read_text(encoding="utf-8")
    events = (ROOT / "static" / "osc" / "osc-events.js").read_text(encoding="utf-8")
    state = (ROOT / "static" / "osc" / "osc-state.js").read_text(encoding="utf-8")

    assert 'id="documentReuse"' in reuse_html
    assert "沿用舊書狀" in osc_html
    assert "1. 新案件資料" in reuse_html
    assert "2. 舊書狀底稿" in reuse_html
    assert "3. 檢查後產生" in reuse_html
    assert "書狀名稱" in reuse_html
    assert "填上方書狀名稱後會自動搜尋" in reuse_html
    assert 'id="reuseDocsBody"' in reuse_html
    assert 'id="reusePreview"' in reuse_html
    assert "產生新書狀" in reuse_html
    assert 'id="draftReuseDocsBody"' not in drafts_html
    assert "我方歷次書狀 Word 索引 / 沿用舊書狀" not in drafts_html
    assert 'data-tab="documentReuse"' in osc_html
    assert 'data-tab="documents"' in osc_html
    assert osc_html.index('data-tab="documentReuse"') < osc_html.index('data-tab="documents"')
    assert "文件總索引" in osc_html
    assert "文件總索引" in documents_html
    assert "/api/osc/drafts/reuse-document" in js
    assert "reuse_scope=own_pleading_word" in js
    assert "documentReuseIsWord" in js
    assert "renderDocumentReusePreview" in js
    assert "syncDocumentReuseDocNameSearch" in js
    assert "showDocumentReuseWarning" in js
    assert "選為底稿" in js
    assert "reuseDocType.addEventListener(\"input\"" in events
    assert "reuseRunBtn" in events
    assert "產生中" in js
    assert "僅 DOCX" not in js
    assert "own_pleading_word" in reuse_html
    assert "document-reuse-select" in events
    assert "loadDocumentReuse" in events
    assert "documentReuse" in state
    assert "selectedDocument" in state
    assert "selectedReuseDocument" not in state
