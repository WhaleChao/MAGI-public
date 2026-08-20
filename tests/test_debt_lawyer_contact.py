from __future__ import annotations

import re
import zipfile
from pathlib import Path
from types import SimpleNamespace

import pytest
from docx import Document
from flask import Flask, render_template

from api import debt_document_generator as debt
from api.law_firm_contact import resolve_lawyer_contact
from src.supplement_core.docx_builder import build_supplement_docx


ROOT = Path(__file__).resolve().parents[1]

CONTACT_ENV = (
    "MAGI_DEFAULT_DEBT_LAWYER",
    "MAGI_CONSUMER_DEBT_LAWYER",
    "MAGI_PUBLIC_LAWYER_NAME",
    "MAGI_LAWYER_NAME",
    "MAGI_DEFAULT_LAWYER",
    "MAGI_LAF_DEFAULT_LAWYER_NAME",
    "MAGI_PUBLIC_LAWYER_ADDRESS",
    "MAGI_LAW_FIRM_ADDRESS",
    "MAGI_PUBLIC_LAWYER_PHONE",
    "MAGI_LAW_FIRM_PHONE",
    "MAGI_PUBLIC_LAWYER_MOBILE",
    "MAGI_LAW_FIRM_MOBILE",
)


@pytest.fixture(autouse=True)
def _isolate_lawyer_contact_environment(monkeypatch) -> None:
    for name in CONTACT_ENV:
        monkeypatch.delenv(name, raising=False)
    monkeypatch.setenv("MAGI_LAF_BRANCH_PROFILE_DB", "0")


def _all_text(document: Document) -> str:
    values = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            values.extend(cell.text for cell in row.cells)
    return "\n".join(values)


def test_contact_resolution_prefers_payload_then_public_environment(
    monkeypatch,
) -> None:
    for name in CONTACT_ENV:
        monkeypatch.delenv(name, raising=False)
    monkeypatch.setenv("MAGI_LAW_FIRM_ADDRESS", "環境合成地址")
    monkeypatch.setenv("MAGI_LAW_FIRM_PHONE", "02-0000-0000")
    monkeypatch.setenv("MAGI_LAW_FIRM_MOBILE", "0900-000-000")

    values = resolve_lawyer_contact({"lawyer_name": "合成律師"})

    assert values == {
        "LAWYER_NAME": "合成律師",
        "LAWYER_ADDRESS": "環境合成地址",
        "LAWYER_PHONE": "02-0000-0000",
        "LAWYER_MOBILE": "0900-000-000",
    }


def test_contact_resolution_uses_shared_debt_setting_precedence(
    monkeypatch,
) -> None:
    from api import laf_branch_profiles

    profile_calls = 0

    def profile():
        nonlocal profile_calls
        profile_calls += 1
        return SimpleNamespace(lawyer_name="不應採用")

    monkeypatch.setenv("MAGI_DEFAULT_DEBT_LAWYER", "環境消債律師")
    monkeypatch.setattr(
        laf_branch_profiles,
        "get_law_firm_profile",
        profile,
    )
    settings = {
        "default_debt_lawyer": "設定消債律師",
        "consumer_debt_lawyer": "次順位消債律師",
    }

    from_setting = resolve_lawyer_contact(
        {},
        requested_fields=("LAWYER_NAME",),
        settings_getter=lambda key, default="": settings.get(key, default),
    )
    from_payload = resolve_lawyer_contact(
        {"lawyer_name": "人工覆寫律師"},
        requested_fields=("LAWYER_NAME",),
        settings_getter=lambda key, default="": settings.get(key, default),
    )

    assert from_setting == {"LAWYER_NAME": "設定消債律師"}
    assert from_payload == {"LAWYER_NAME": "人工覆寫律師"}
    assert profile_calls == 0


def test_requested_name_only_does_not_load_unused_profile_fields(monkeypatch) -> None:
    from api import laf_branch_profiles

    profile_calls = 0

    def profile():
        nonlocal profile_calls
        profile_calls += 1
        return SimpleNamespace(lawyer_name="不應採用")

    monkeypatch.setattr(
        laf_branch_profiles,
        "get_law_firm_profile",
        profile,
    )

    values = resolve_lawyer_contact(
        {"lawyer_name": "人工覆寫律師"},
        requested_fields=("LAWYER_NAME",),
    )

    assert values == {"LAWYER_NAME": "人工覆寫律師"}
    assert profile_calls == 0


def test_demo_lawyer_and_bare_client_mobile_are_not_contact_overrides(
    monkeypatch,
) -> None:
    from api import laf_branch_profiles

    monkeypatch.setattr(
        laf_branch_profiles,
        "get_law_firm_profile",
        lambda: SimpleNamespace(
            lawyer_name="正式消債律師",
            address_line="合成市事務所路一號",
            phone="02-0000-0000",
            mobile="0911-000-000",
        ),
    )

    values = resolve_lawyer_contact(
        {"lawyer_name": "測試律師", "mobile": "0999-CLIENT-ONLY"}
    )

    assert values["LAWYER_NAME"] == "正式消債律師"
    assert values["LAWYER_MOBILE"] == "0911-000-000"
    assert "0999-CLIENT-ONLY" not in values.values()


def test_public_seed_profile_becomes_explicit_human_fill_markers(
    monkeypatch,
) -> None:
    from api import laf_branch_profiles

    for name in CONTACT_ENV:
        monkeypatch.delenv(name, raising=False)
    monkeypatch.setattr(
        laf_branch_profiles,
        "get_law_firm_profile",
        lambda: SimpleNamespace(
            lawyer_name="受任律師",
            address_line="範例事務所地址",
            phone="事務所電話",
            mobile="",
        ),
    )

    assert resolve_lawyer_contact({}) == {
        "LAWYER_NAME": "請填律師姓名",
        "LAWYER_ADDRESS": "請填律師地址",
        "LAWYER_PHONE": "請填律師電話",
        "LAWYER_MOBILE": "請填律師手機",
    }


def test_report_template_replaces_all_contact_placeholders(
    tmp_path: Path,
    monkeypatch,
) -> None:
    template_dir = tmp_path / "templates"
    template_dir.mkdir()
    template = Document()
    table = template.add_table(rows=4, cols=1)
    for cell, placeholder in zip(
        (row.cells[0] for row in table.rows),
        ("LAWYER_NAME", "LAWYER_ADDRESS", "LAWYER_PHONE", "LAWYER_MOBILE"),
    ):
        cell.text = placeholder
    template.save(template_dir / "D.docx")
    monkeypatch.setattr(debt, "_TEMPLATE_DIR", str(template_dir))
    contact = {
        "LAWYER_NAME": "合成律師",
        "LAWYER_ADDRESS": "合成市測試路一號",
        "LAWYER_PHONE": "02-0000-0000",
        "LAWYER_MOBILE": "0900-000-000",
    }

    document = debt.generate_report(contact)
    text = _all_text(document)

    assert not any(placeholder in text for placeholder in contact)
    assert all(value in text for value in contact.values())


def test_application_replaces_both_lawyer_cells_from_same_contact_contract(
    tmp_path: Path,
    monkeypatch,
) -> None:
    template_dir = tmp_path / "application-template"
    template_dir.mkdir()
    template = Document()
    dimensions = ((2, 3), (2, 3), (1, 1), (1, 2), (1, 1), (2, 1), (2, 3))
    for rows, columns in dimensions:
        template.add_table(rows=rows, cols=columns)
    template.tables[0].cell(1, 1).text = "LAWYER_NAME"
    template.tables[6].cell(1, 2).text = "LAWYER_NAME"
    template.save(template_dir / "A.docx")
    monkeypatch.setattr(debt, "_TEMPLATE_DIR", str(template_dir))

    document = debt.generate_application(
        {
            "name": "合成聲請人",
            "address": "合成市聲請路一號",
            "lawyer_name": "合成律師",
        }
    )

    assert document.tables[0].cell(1, 1).text == "合成律師"
    assert document.tables[6].cell(1, 2).text == "合成律師"
    assert "LAWYER_NAME" not in _all_text(document)


def test_supplement_template_and_api_output_use_same_contact_contract(
    tmp_path: Path,
) -> None:
    contact = {
        "lawyer_name": "合成律師",
        "lawyer_address": "合成市補件路二號",
        "lawyer_phone": "03-000-0000",
        "lawyer_mobile": "0911-000-000",
    }
    template = tmp_path / "D_supplement.docx"
    document = Document()
    document.add_paragraph("謹就消費者債務清理事件陳報事：")
    for placeholder in (
        "LAWYER_NAME",
        "LAWYER_ADDRESS",
        "LAWYER_PHONE",
        "LAWYER_MOBILE",
    ):
        document.add_paragraph(placeholder)
    document.save(template)
    output = tmp_path / "output/supplement.docx"

    result = build_supplement_docx(
        {**contact, "parties": ["合成聲請人"], "court": "合成地方法院"},
        {"case_meta": {}, "items": []},
        [],
        procedure="更生",
        brief_seq=1,
        output_path=str(output),
        template_path=str(template),
    )
    rendered = _all_text(Document(result["output_path"]))
    expected = {
        "合成律師",
        "合成市補件路二號",
        "03-000-0000",
        "0911-000-000",
    }
    assert not set(result["filled_fields"]) & {
        "LAWYER_NAME",
        "LAWYER_ADDRESS",
        "LAWYER_PHONE",
        "LAWYER_MOBILE",
    }
    assert not expected & set(result["filled_fields"].values())
    assert all(value in rendered for value in expected)
    assert "LAWYER_" not in rendered

    api_document = debt.generate_supplement(contact)
    api_text = _all_text(api_document)
    assert all(value in api_text for value in expected)


def test_report_and_supplement_forms_expose_optional_contact_overrides() -> None:
    expected = {
        "lawyer_name",
        "lawyer_address",
        "lawyer_phone",
        "lawyer_mobile",
    }
    for form_type in ("report", "supplement"):
        schema = debt.get_form_schema(form_type)
        assert expected <= {field["key"] for field in schema["fields"]}
    application = debt.get_form_schema("application")
    assert "lawyer_name" in {field["key"] for field in application["fields"]}


def test_rendered_osc_debt_page_submits_all_manual_contact_overrides() -> None:
    app = Flask(
        __name__,
        static_folder=str(ROOT / "static"),
        template_folder=str(ROOT / "templates"),
    )
    with app.test_request_context("/osc/debt"):
        html = render_template("osc_debt.html", user=None)

    expected_ids = {
        "app-lawyer-name",
        "rpt-lawyer-name",
        "rpt-lawyer-address",
        "rpt-lawyer-phone",
        "rpt-lawyer-mobile",
        "sup-lawyer-name",
        "sup-lawyer-address",
        "sup-lawyer-phone",
        "sup-lawyer-mobile",
        "batch-lawyer-name",
        "batch-lawyer-address",
        "batch-lawyer-phone",
        "batch-lawyer-mobile",
    }
    for element_id in expected_ids:
        assert f'id="{element_id}"' in html

    payload_bindings = {
        "lawyer_name": (
            "app-lawyer-name",
            "rpt-lawyer-name",
            "sup-lawyer-name",
            "batch-lawyer-name",
        ),
        "lawyer_address": (
            "rpt-lawyer-address",
            "sup-lawyer-address",
            "batch-lawyer-address",
        ),
        "lawyer_phone": (
            "rpt-lawyer-phone",
            "sup-lawyer-phone",
            "batch-lawyer-phone",
        ),
        "lawyer_mobile": (
            "rpt-lawyer-mobile",
            "sup-lawyer-mobile",
            "batch-lawyer-mobile",
        ),
    }
    for payload_key, element_ids in payload_bindings.items():
        for element_id in element_ids:
            assert re.search(
                rf"{payload_key}:\s*document\.getElementById\('{element_id}'\)\.value",
                html,
            )


def test_osc_debt_prepares_shared_setting_and_preserves_manual_overrides(
    monkeypatch,
) -> None:
    from api.blueprints import osc_debt

    monkeypatch.setattr(
        osc_debt,
        "_debt_setting_value",
        lambda key, default="": (
            "設定消債律師" if key == "default_debt_lawyer" else default
        ),
    )
    contact = {
        "lawyer_address": "合成市事務所路二號",
        "lawyer_phone": "02-0000-0000",
        "lawyer_mobile": "0922-000-000",
    }

    from_setting = osc_debt._prepare_debt_document_data("report", contact)
    from_manual = osc_debt._prepare_debt_document_data(
        "report",
        {**contact, "lawyer_name": "人工覆寫律師"},
    )

    assert from_setting["lawyer_name"] == "設定消債律師"
    assert from_manual["lawyer_name"] == "人工覆寫律師"
    assert from_setting["lawyer_mobile"] == "0922-000-000"


def test_shipped_debt_templates_are_deidentified_placeholder_assets() -> None:
    application = ROOT / "integrations/debt_robot/document/A.docx"
    reports = (
        ROOT / "integrations/debt_robot/document/D.docx",
        ROOT / "data/templates/D_supplement.docx",
    )
    all_templates = (
        application,
        ROOT / "integrations/debt_robot/document/B.docx",
        ROOT / "integrations/debt_robot/document/C.docx",
        *reports,
    )

    assert _all_text(Document(application)).count("LAWYER_NAME") >= 2
    for path in reports:
        text = _all_text(Document(path))
        assert all(
            placeholder in text
            for placeholder in (
                "LAWYER_NAME",
                "LAWYER_ADDRESS",
                "LAWYER_PHONE",
                "LAWYER_MOBILE",
            )
        )

    mobile = re.compile(r"(?<!\d)09\d{2}[- ]?\d{3}[- ]?\d{3}(?!\d)")
    for path in all_templates:
        document = Document(path)
        assert not mobile.search(_all_text(document))
        assert not document.core_properties.author
        assert not document.core_properties.last_modified_by
        with zipfile.ZipFile(path) as archive:
            assert "docProps/custom.xml" not in archive.namelist()
            document_xml = archive.read("word/document.xml")
            assert b"w:rsid" not in document_xml


def test_shipped_application_and_report_replace_every_contact_placeholder() -> None:
    contact = {
        "lawyer_name": "合成律師",
        "lawyer_address": "合成市完整路三號",
        "lawyer_phone": "04-0000-0000",
        "lawyer_mobile": "0922-000-000",
    }
    application = debt.generate_application(
        {
            **contact,
            "name": "合成聲請人",
            "address": "合成市聲請路四號",
        }
    )
    application_text = _all_text(application)
    assert application_text.count("合成律師") >= 2
    assert "LAWYER_" not in application_text

    report = debt.generate_report(
        {
            **contact,
            "A1": "一",
            "A2": "合成年度測字第一號",
            "A3": "測",
            "A4": "合成聲請人",
            "E1": "合成地方法院",
        }
    )
    report_text = _all_text(report)
    assert all(value in report_text for value in contact.values())
    assert "LAWYER_" not in report_text
