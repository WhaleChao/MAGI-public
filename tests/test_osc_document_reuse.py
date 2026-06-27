from pathlib import Path

import pytest
from docx import Document

from api.osc.document_reuse import build_pleading_index, reuse_docx_document


def _add_split_runs(paragraph, parts):
    for part in parts:
        paragraph.add_run(part)


def _make_source_docx(path: Path) -> None:
    doc = Document()

    header = doc.sections[0].header.paragraphs[0]
    _add_split_runs(header, ["臺灣", "臺北", "地方法院"])

    paragraph = doc.add_paragraph()
    _add_split_runs(paragraph, ["案號：113年度", "訴字", "第100號", "　股別：義股"])

    doc.add_paragraph("內部案號：OSC-113-001")
    doc.add_paragraph("案由：損害賠償")

    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "原告"
    table.cell(0, 1).text = "王小明"
    table.cell(1, 0).text = "被告"
    table.cell(1, 1).text = "李大華"

    footer = doc.sections[0].footer.paragraphs[0]
    footer.text = "此致 臺灣臺北地方法院"

    doc.save(path)


def _all_doc_text(path: Path) -> str:
    doc = Document(path)
    parts = []

    def collect(container):
        parts.extend(p.text for p in container.paragraphs)
        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    collect(cell)

    collect(doc)
    for section in doc.sections:
        collect(section.header)
        collect(section.footer)
    return "\n".join(parts)


def test_reuse_docx_replaces_split_runs_tables_headers_and_footers(tmp_path):
    source = tmp_path / "20260101 民事準備書狀.docx"
    output_dir = tmp_path / "out"
    _make_source_docx(source)

    source_case = {
        "court_case_no": "113年度訴字第100號",
        "case_number": "OSC-113-001",
        "court_division": "義股",
        "plaintiff_name": "王小明",
        "defendant_name": "李大華",
        "court_name": "臺灣臺北地方法院",
        "case_reason": "損害賠償",
    }
    target_case = {
        "court_case_no": "115年度重訴字第888號",
        "case_number": "OSC-115-009",
        "court_division": "忠股",
        "plaintiff_name": "陳新明",
        "defendant_name": "林新華",
        "court_name": "臺灣新北地方法院",
        "case_reason": "返還借款",
    }

    result = reuse_docx_document(
        source,
        source_case,
        target_case,
        suggested_filename="改作書狀.docx",
        output_dir=output_dir,
    )

    output = Path(result["output_path"])
    assert output.exists()
    assert result["file_name"] == "改作書狀.docx"
    assert result["replacement_count"] >= 7
    assert {item["field"] for item in result["replacements"] if item["count"]} >= {
        "court_case_no",
        "internal_case_no",
        "division",
        "plaintiff",
        "defendant",
        "institution",
        "case_reason",
    }

    text = _all_doc_text(output)
    for expected in (
        "115年度重訴字第888號",
        "OSC-115-009",
        "忠股",
        "陳新明",
        "林新華",
        "臺灣新北地方法院",
        "返還借款",
    ):
        assert expected in text
    for stale in (
        "113年度訴字第100號",
        "OSC-113-001",
        "義股",
        "王小明",
        "李大華",
        "臺灣臺北地方法院",
        "損害賠償",
    ):
        assert stale not in text

    source_text = _all_doc_text(source)
    assert "113年度訴字第100號" in source_text
    assert "臺灣臺北地方法院" in source_text
    assert "115年度重訴字第888號" not in source_text


def test_reuse_docx_never_overwrites_existing_output(tmp_path):
    source = tmp_path / "民事準備書狀.docx"
    output_dir = tmp_path / "out"
    output_dir.mkdir()
    _make_source_docx(source)

    existing = output_dir / "重用.docx"
    _make_source_docx(existing)
    before = existing.read_bytes()

    result = reuse_docx_document(
        source,
        {"court_case_no": "113年度訴字第100號"},
        {"court_case_no": "115年度重訴字第888號"},
        suggested_filename="重用.docx",
        output_dir=output_dir,
    )

    output = Path(result["output_path"])
    assert output.name == "重用 (1).docx"
    assert output.exists()
    assert existing.read_bytes() == before


def test_reuse_docx_rejects_non_docx_source(tmp_path):
    source = tmp_path / "書狀.pdf"
    source.write_bytes(b"%PDF-1.4\n")

    with pytest.raises(ValueError, match="Only .docx"):
        reuse_docx_document(source, {}, {})


def test_build_pleading_index_lists_docx_pleadings_newest_first(tmp_path):
    old_doc = tmp_path / "20250101 民事準備書狀.docx"
    new_doc = tmp_path / "20260101 民事聲請狀.docx"
    ignored = tmp_path / "筆錄.docx"
    _make_source_docx(old_doc)
    _make_source_docx(new_doc)
    _make_source_docx(ignored)
    old_doc.touch()
    new_doc.touch()

    index = build_pleading_index(tmp_path)

    assert [item["file_name"] for item in index] == [
        "20260101 民事聲請狀.docx",
        "20250101 民事準備書狀.docx",
    ]

