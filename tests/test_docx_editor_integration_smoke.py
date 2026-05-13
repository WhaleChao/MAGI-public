"""
tests/test_docx_editor_integration_smoke.py

Phase 10: end-to-end smoke test — generate_docx → cmd_extract → cmd_find → apply_tracked_edits round-trip
"""

import io
import os
import sys
import json
import tempfile
import pytest

_SKILL_DIR = os.path.join(os.path.dirname(__file__), "..", "skills", "docx-editor")
sys.path.insert(0, _SKILL_DIR)
sys.path.insert(0, os.path.join(_SKILL_DIR, "lib"))

from lib.generator import generate_docx, GenerateDocxRequest, SectionSpec, TableSpec
from lib.tracked_edits import apply_tracked_edits, EditInput

import importlib.util as _ilu
_spec = _ilu.spec_from_file_location("docx_action", os.path.join(_SKILL_DIR, "action.py"))
_action_mod = _ilu.module_from_spec(_spec)
_spec.loader.exec_module(_action_mod)
cmd_extract = _action_mod.cmd_extract
cmd_find = _action_mod.cmd_find
cmd_apply = _action_mod.cmd_apply


# ── Full round-trip ─────────────────────────────────────────────────────────────

def test_full_roundtrip():
    """generate_docx → write to temp → cmd_extract → cmd_find → apply_tracked_edits."""

    # Step 1: generate_docx
    req = GenerateDocxRequest(
        title="測試文件",
        sections=[
            SectionSpec(heading="事實陳述", level=1, content="甲方同意於期限內付款。乙方確認收訖。"),
            SectionSpec(heading="法律分析", level=1, content="依民法第184條，當事人應負責。"),
        ],
    )
    docx_bytes = generate_docx(req)
    assert len(docx_bytes) > 500

    # Step 2: write to temp file
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf:
        tf.write(docx_bytes)
        tmp_path = tf.name

    try:
        # Step 3: cmd_extract
        extract_result = cmd_extract(tmp_path)
        assert extract_result["paragraph_count"] > 0
        full_text = extract_result["text"]
        assert "甲方同意" in full_text
        assert "民法" in full_text

        # Step 4: cmd_find
        find_result = cmd_find(tmp_path, "甲方")
        assert find_result["total"] >= 1
        assert find_result["matches"][0]["match"] == "甲方"

        # Step 5: apply_tracked_edits
        edit = EditInput(
            find="甲方同意",
            replace="原告同意",
            context_before="",
            context_after="於期限內",
            reason="依律師指令改稱謂",
        )
        result = apply_tracked_edits(docx_bytes, [edit], author="MAGI")
        assert len(result.errors) == 0
        assert len(result.changes) == 1
        # deleted_text may be the full find or a fragment (depending on run splitting)
        assert "甲方" in result.changes[0].deleted_text or "同意" in result.changes[0].deleted_text
        assert "原告" in result.changes[0].inserted_text or "同意" in result.changes[0].inserted_text

        # Step 6: verify output docx is readable
        with io.BytesIO(result.bytes) as buf:
            from docx import Document
            doc = Document(buf)
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "測試文件" in all_text

    finally:
        os.unlink(tmp_path)


# ── generate + landscape round-trip ──────────────────────────────────────────

def test_generate_landscape_roundtrip():
    """generate_docx landscape=True → readable → orientation swapped."""
    from docx.enum.section import WD_ORIENT

    req = GenerateDocxRequest(
        title="橫式報告",
        sections=[SectionSpec(heading="費用表", level=1, table=TableSpec(
            headers=["項目", "金額"],
            rows=[["律師費", "50000"]],
        ))],
        landscape=True,
    )
    docx_bytes = generate_docx(req)
    doc = _round_trip(docx_bytes)
    section = doc.sections[0]
    assert section.page_width > section.page_height


# ── citation parse integration ─────────────────────────────────────────────────

def test_citation_format_roundtrip():
    """parse_citations + render_citations_for_telegram integration."""
    from skills.bridge.citation_format import parse_citations, render_citations_for_telegram

    answer = """依據[1]，被告應負責。

<CITATIONS>
[{"ref": 1, "doc_id": "合約書", "page": "3", "quote": "被告應負損害賠償責任"}]
</CITATIONS>"""

    parsed = parse_citations(answer)
    assert len(parsed.citations) == 1
    assert parsed.citations[0].doc_id == "合約書"

    rendered = render_citations_for_telegram(parsed)
    assert "📄 引用：" in rendered
    assert "合約書" in rendered
    assert "<CITATIONS>" not in rendered


def _round_trip(docx_bytes: bytes):
    from docx import Document
    import io
    return Document(io.BytesIO(docx_bytes))
