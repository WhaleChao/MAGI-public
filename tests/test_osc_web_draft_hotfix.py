from __future__ import annotations

import os
from pathlib import Path
from urllib.parse import quote

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from flask import Flask
from flask_login import LoginManager, UserMixin, login_user


class _Operator(UserMixin):
    id = "osc-draft-hotfix-operator"
    role = "operator"


def test_runtime_exports_are_authorized_without_widening_runtime_root(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    from api.osc import utils

    runtime_root = tmp_path / "runtime"
    exports = runtime_root / "shared" / "exports"
    exports.mkdir(parents=True)
    generated = exports / "generated.docx"
    generated.write_bytes(b"generated-by-osc")
    outside = runtime_root / "private-state.json"
    outside.write_text("private", encoding="utf-8")

    monkeypatch.delenv("MAGI_V3_OFFLINE_CERTIFICATION", raising=False)
    monkeypatch.delenv("PAPERCLIP_FILEMANAGER_TEST_BASE", raising=False)
    monkeypatch.setenv("MAGI_EXPORTS_DIR", str(exports))
    monkeypatch.setattr(utils, "default_synology_share_roots", lambda **_kwargs: [])

    roots = utils._osc_allowed_local_roots()
    assert os.path.realpath(exports) in roots
    assert utils._osc_is_safe_local_path(str(generated)) is True
    assert utils._osc_is_safe_local_path(str(outside)) is False

    escaping_link = exports / "escape.json"
    escaping_link.symlink_to(outside)
    assert utils._osc_is_safe_local_path(str(escaping_link)) is False


def test_authenticated_web_download_reads_its_own_runtime_export(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    exports = tmp_path / "runtime" / "shared" / "exports"
    exports.mkdir(parents=True)
    generated = exports / "民事準備書狀_合成.docx"
    payload = b"synthetic-osc-docx"
    generated.write_bytes(payload)

    monkeypatch.delenv("MAGI_V3_OFFLINE_CERTIFICATION", raising=False)
    monkeypatch.delenv("PAPERCLIP_FILEMANAGER_TEST_BASE", raising=False)
    monkeypatch.setenv("MAGI_EXPORTS_DIR", str(exports))

    from api.blueprints import osc_cases

    monkeypatch.setattr(osc_cases, "_osc_audit_file_event", lambda *_a, **_k: None)

    app = Flask(__name__)
    app.config.update(TESTING=True, SECRET_KEY="osc-draft-hotfix")
    login = LoginManager(app)

    @login.user_loader
    def _load_user(user_id: str):
        return _Operator() if user_id == _Operator.id else None

    @app.post("/__osc_hotfix_login")
    def _login():
        login_user(_Operator())
        return {"ok": True}

    app.register_blueprint(osc_cases.osc_bp)
    client = app.test_client()
    assert client.post("/__osc_hotfix_login").status_code == 200

    response = client.get(
        f"/api/osc/files/content?path={quote(str(generated), safe='')}"
    )
    assert response.status_code == 200
    assert response.data == payload
    assert response.headers["Content-Disposition"].startswith("attachment;")


def test_web_pleading_export_keeps_osc_generation_and_office_format_contract(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    from api import startup
    from api.blueprints import osc_cases

    # The web workflow retains the desktop OSC.py document choices and prompt
    # fields, while adding citation and quality gates before export.
    assert osc_cases._OSC_DRAFT_DOC_TYPES == [
        "民事起訴狀",
        "民事答辯狀",
        "民事準備書狀",
        "民事上訴狀",
        "民事聲請狀",
        "刑事告訴狀",
        "刑事答辯狀",
        "刑事上訴狀",
        "刑事聲請狀",
        "刑事陳報狀",
        "行政起訴狀",
        "行政答辯狀",
        "抗告狀",
        "聲明異議狀",
        "強制執行聲請狀",
        "假扣押聲請狀",
        "假處分聲請狀",
        "支付命令聲請狀",
        "本票裁定聲請狀",
    ]
    for field in (
        "{doc_type}",
        "{case_number}",
        "{division}",
        "{court_name}",
        "{reason}",
        "{plaintiff}",
        "{defendant}",
        "{case_facts}",
        "{legal_insights}",
        "{reference_style}",
    ):
        assert field in osc_cases._OSC_DRAFT_PROMPT_TEMPLATE

    exports = tmp_path / "exports"
    monkeypatch.setattr(startup, "EXPORTS_DIR", str(exports))
    monkeypatch.setattr(startup, "_find_nas_pleading_style_template", lambda _title: "")
    result = startup._export_form_docx(
        """民事準備書狀
案號：合成年度訴字第一號
股別：合成股
原告：測試原告
被告：測試被告

為損害賠償事件，提出準備書狀事：
一、此為不含真實個資的格式驗證內容。

此致
臺灣合成地方法院 公鑒
具狀人：測試原告
中華民國一一五年八月二十五日""",
        "osc-web-format-contract",
        title="民事準備書狀",
    )

    assert result["success"] is True
    document = Document(result["path"])
    section = document.sections[0]
    assert section.page_width.cm == pytest.approx(21.0, abs=0.02)
    assert section.page_height.cm == pytest.approx(29.7, abs=0.02)
    assert section.top_margin.cm == pytest.approx(2.54, abs=0.02)
    assert section.bottom_margin.cm == pytest.approx(2.54, abs=0.02)
    assert section.left_margin.cm == pytest.approx(3.17, abs=0.02)
    assert section.right_margin.cm == pytest.approx(3.17, abs=0.02)

    title = document.paragraphs[0]
    assert title.text == "民事準備書狀"
    assert title.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert title.runs[0].bold is True
    assert title.runs[0].font.size.pt == pytest.approx(26.0)

    meta_text = [
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    ]
    assert {"案號", "股別", "原告", "被告"} <= set(meta_text)

    body = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith("為損害賠償事件")
    )
    assert body.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
    assert body.paragraph_format.first_line_indent.pt == pytest.approx(32.0)
    signature = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith("具狀人")
    )
    assert signature.alignment == WD_ALIGN_PARAGRAPH.RIGHT
