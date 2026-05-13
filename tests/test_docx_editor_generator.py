"""
tests/test_docx_editor_generator.py

Phase 4 tests: generate_docx() + GenerateDocxRequest / SectionSpec / TableSpec
"""

import io
import os
import sys
import warnings
import pytest

# Ensure skill lib is importable
_SKILL_DIR = os.path.join(os.path.dirname(__file__), "..", "skills", "docx-editor")
sys.path.insert(0, _SKILL_DIR)

from lib.generator import (
    GenerateDocxRequest,
    SectionSpec,
    TableSpec,
    generate_docx,
)


def _round_trip(docx_bytes: bytes):
    """Helper: round-trip bytes → python-docx Document."""
    from docx import Document
    buf = io.BytesIO(docx_bytes)
    return Document(buf)


# ── Test 1: simple heading + content ──────────────────────────────────────────

def test_simple_sections_readable():
    """generate_docx with title + heading + content produces valid .docx."""
    req = GenerateDocxRequest(
        title="測試文件",
        sections=[
            SectionSpec(heading="第一章", level=1, content="第一段內容。\n\n第二段繼續。"),
            SectionSpec(heading="第一節", level=2, content="細節說明。"),
        ],
    )
    docx_bytes = generate_docx(req)
    assert len(docx_bytes) > 500

    doc = _round_trip(docx_bytes)
    # Extract all text
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "測試文件" in all_text
    assert "第一章" in all_text
    assert "第一節" in all_text
    assert "第一段內容" in all_text
    assert "第二段繼續" in all_text
    assert "細節說明" in all_text


# ── Test 2: heading level jump raises ValueError ───────────────────────────────

def test_heading_level_jump_raises():
    """Level jump (1 → 3) must raise ValueError."""
    req = GenerateDocxRequest(
        title="跳級文件",
        sections=[
            SectionSpec(heading="章", level=1),
            SectionSpec(heading="小小節", level=3),  # jump: 1→3 invalid
        ],
    )
    with pytest.raises(ValueError, match="Heading level jump"):
        generate_docx(req)


# ── Test 3: landscape mode ─────────────────────────────────────────────────────

def test_landscape_orientation():
    """landscape=True should swap page width/height."""
    from docx.enum.section import WD_ORIENT

    req = GenerateDocxRequest(
        title="橫式文件",
        sections=[SectionSpec(heading="Landscape Test", level=1)],
        landscape=True,
    )
    docx_bytes = generate_docx(req)
    doc = _round_trip(docx_bytes)
    section = doc.sections[0]
    # In landscape, width > height
    assert section.page_width > section.page_height


# ── Test 4: table contents ─────────────────────────────────────────────────────

def test_table_contents():
    """Table headers and rows appear in the document."""
    req = GenerateDocxRequest(
        title="表格文件",
        sections=[
            SectionSpec(
                heading="費用明細",
                level=1,
                table=TableSpec(
                    headers=["項目", "金額", "說明"],
                    rows=[
                        ["服務費", "5000", "法律諮詢"],
                        ["文件費", "1000", "書狀製作"],
                    ],
                ),
            )
        ],
    )
    docx_bytes = generate_docx(req)
    doc = _round_trip(docx_bytes)

    # Find tables
    assert len(doc.tables) == 1
    tbl = doc.tables[0]
    # Header row
    header_cells = [cell.text for cell in tbl.rows[0].cells]
    assert "項目" in header_cells
    assert "金額" in header_cells
    assert "說明" in header_cells
    # Data rows
    all_cell_text = " ".join(
        cell.text for row in tbl.rows[1:] for cell in row.cells
    )
    assert "服務費" in all_cell_text
    assert "5000" in all_cell_text
    assert "文件費" in all_cell_text


# ── Test 5: table column mismatch raises ValueError ────────────────────────────

def test_table_col_mismatch_raises():
    """Table row with wrong column count raises ValueError."""
    req = GenerateDocxRequest(
        title="錯誤表格",
        sections=[
            SectionSpec(
                table=TableSpec(
                    headers=["A", "B", "C"],
                    rows=[["x", "y"]],  # missing column
                )
            )
        ],
    )
    with pytest.raises(ValueError, match="Table row 0 has 2 columns, expected 3"):
        generate_docx(req)


# ── Test 6: page_break section ────────────────────────────────────────────────

def test_page_break_section():
    """page_break=True in a section produces a page break element in docx."""
    req = GenerateDocxRequest(
        title="分頁文件",
        sections=[
            SectionSpec(heading="第一章", level=1, content="第一章內容。"),
            SectionSpec(heading="第二章", level=1, content="第二章內容。", page_break=True),
        ],
    )
    docx_bytes = generate_docx(req)
    # Just verify the docx is readable and contains both headings
    doc = _round_trip(docx_bytes)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "第一章" in all_text
    assert "第二章" in all_text
    # Also verify the raw XML contains a page break
    import zipfile
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
        xml = zf.read("word/document.xml").decode("utf-8")
    assert "w:lastRenderedPageBreak" in xml or "w:pageBreak" in xml or "w:br" in xml


# ── Test 7: numbered prefix heading warning ───────────────────────────────────

def test_numbered_prefix_heading_warns():
    """Heading with '1. ' prefix should warn but not raise."""
    req = GenerateDocxRequest(
        title="編號標題警告",
        sections=[SectionSpec(heading="1. 引言", level=1)],
    )
    with warnings.catch_warnings(record=True) as w:
        warnings.simplefilter("always")
        docx_bytes = generate_docx(req)
        assert any("numbering prefix" in str(warning.message) for warning in w)
    doc = _round_trip(docx_bytes)
    assert "1. 引言" in "\n".join(p.text for p in doc.paragraphs)


# ── Test 8: empty sections list ───────────────────────────────────────────────

def test_empty_sections():
    """Empty sections list should still produce a valid docx with just the title."""
    req = GenerateDocxRequest(title="空文件", sections=[])
    docx_bytes = generate_docx(req)
    assert len(docx_bytes) > 100
    doc = _round_trip(docx_bytes)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "空文件" in all_text


# ── Test 9: multi-level heading sequence OK ───────────────────────────────────

def test_valid_heading_level_sequence():
    """Valid 1→2→3→2→3 sequence should not raise."""
    req = GenerateDocxRequest(
        title="多層標題",
        sections=[
            SectionSpec(heading="章", level=1),
            SectionSpec(heading="節", level=2),
            SectionSpec(heading="款", level=3),
            SectionSpec(heading="節2", level=2),
            SectionSpec(heading="款2", level=3),
        ],
    )
    docx_bytes = generate_docx(req)
    doc = _round_trip(docx_bytes)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "節2" in all_text
