"""
生成 docx-editor 測試 fixture 檔案。

執行方式：
    python tests/fixtures/docx_editor/generate_fixtures.py
"""

import io
import os
import zipfile

import docx
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))


def save_docx(doc: Document, filename: str):
    path = os.path.join(OUTPUT_DIR, filename)
    doc.save(path)
    print(f"Created: {path}")
    return path


def make_simple():
    """單段純文字 (happy path)"""
    doc = Document()
    doc.add_paragraph("Hello World")
    return save_docx(doc, "simple.docx")


def make_multi_run():
    """一段內多 runs（測 run splitting）"""
    doc = Document()
    p = doc.add_paragraph()
    # Add multiple runs with different formatting
    run1 = p.add_run("The quick ")
    run2 = p.add_run("brown fox")
    run2.bold = True
    run3 = p.add_run(" jumps over")
    run4 = p.add_run(" the lazy dog")
    run4.italic = True
    return save_docx(doc, "multi_run.docx")


def make_multi_paragraph():
    """多段（測 cross-paragraph anchor）"""
    doc = Document()
    doc.add_paragraph("First paragraph with introduction text.")
    doc.add_paragraph("Second paragraph contains the important clause: the defendant shall pay.")
    doc.add_paragraph("Third paragraph with conclusion.")
    return save_docx(doc, "multi_paragraph.docx")


def make_with_table():
    """含表格（確認不破壞表格）"""
    doc = Document()
    doc.add_paragraph("Document with a table below:")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Item A"
    table.cell(1, 1).text = "100"
    doc.add_paragraph("Text after the table: please review above.")
    return save_docx(doc, "with_table.docx")


def make_with_existing_changes():
    """已含 tracked changes（測 ID 衝突）"""
    doc = Document()
    p_el = doc.add_paragraph("Original text here.").runs[0]._element.getparent()

    # Manually inject a w:ins tracked change with w:id="5"
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    # Build: <w:ins w:id="5" w:author="Previous" w:date="2026-01-01T00:00:00Z">
    #          <w:r><w:t> inserted</w:t></w:r>
    #        </w:ins>
    ins_el = OxmlElement("w:ins")
    ins_el.set(qn("w:id"), "5")
    ins_el.set(qn("w:author"), "Previous Author")
    ins_el.set(qn("w:date"), "2026-01-01T00:00:00Z")
    r_el = OxmlElement("w:r")
    t_el = OxmlElement("w:t")
    t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t_el.text = " (previously inserted)"
    r_el.append(t_el)
    ins_el.append(r_el)
    p_el.append(ins_el)

    # Build: <w:del w:id="3" w:author="Previous" w:date="2026-01-01T00:00:00Z">
    #          <w:r><w:delText>deleted word</w:delText></w:r>
    #        </w:del>
    del_el = OxmlElement("w:del")
    del_el.set(qn("w:id"), "3")
    del_el.set(qn("w:author"), "Previous Author")
    del_el.set(qn("w:date"), "2026-01-01T00:00:00Z")
    r2_el = OxmlElement("w:r")
    dt_el = OxmlElement("w:delText")
    dt_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    dt_el.text = "deleted"
    r2_el.append(dt_el)
    del_el.append(r2_el)
    # Add a second paragraph for deletions
    p2 = doc.add_paragraph()
    p2._element.append(del_el)
    # Add a run after
    r3_el = OxmlElement("w:r")
    t3_el = OxmlElement("w:t")
    t3_el.text = " text after deletion."
    r3_el.append(t3_el)
    p2._element.append(r3_el)

    return save_docx(doc, "with_existing_changes.docx")


if __name__ == "__main__":
    make_simple()
    make_multi_run()
    make_multi_paragraph()
    make_with_table()
    make_with_existing_changes()
    print("\nAll fixtures generated successfully.")
