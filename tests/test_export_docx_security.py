from __future__ import annotations

import importlib
import json
import shutil
import subprocess
from pathlib import Path

import pytest
from docx import Document


def _reload_export_docx(monkeypatch, tmp_path):
    monkeypatch.setenv("MAGI_EXPORTS_DIR", str(tmp_path / "exports"))
    for name in (
        "MAGI_V3_RELEASE_ID",
        "MAGI_V3_DEPLOYMENT_MODE",
        "MAGI_V3_RELEASE_MANIFEST",
    ):
        monkeypatch.delenv(name, raising=False)
    import skills.ops.export_docx as export_docx

    return importlib.reload(export_docx)


def test_export_bilingual_docx_rejects_path_traversal_filename(monkeypatch, tmp_path):
    export_docx = _reload_export_docx(monkeypatch, tmp_path)

    result = export_docx.export_bilingual_docx(
        [{"page": 1, "source": "A", "target": "B"}],
        filename="../escape.docx",
    )

    assert result["success"] is False
    assert "basename" in result["error"]
    assert not (tmp_path / "escape.docx").exists()


def test_export_summary_docx_rejects_absolute_filename(monkeypatch, tmp_path):
    export_docx = _reload_export_docx(monkeypatch, tmp_path)

    result = export_docx.export_summary_docx(
        [{"heading": "H", "summary": "S", "excerpt": "E"}],
        filename=str(tmp_path / "escape.docx"),
    )

    assert result["success"] is False
    assert "basename" in result["error"]


def test_export_transcript_docx_rejects_non_docx_filename(monkeypatch, tmp_path):
    export_docx = _reload_export_docx(monkeypatch, tmp_path)

    result = export_docx.export_transcript_docx(
        [{"speaker": "A", "time": "00:00", "content": "hello"}],
        filename="transcript.txt",
    )

    assert result["success"] is False
    assert "must end with .docx" in result["error"]


@pytest.mark.parametrize("mode", ["bilingual", "transcript", "summary"])
def test_node_module_failure_uses_valid_python_docx_fallback(
    monkeypatch,
    tmp_path,
    mode,
):
    export_docx = _reload_export_docx(monkeypatch, tmp_path)
    monkeypatch.setattr(
        export_docx.subprocess,
        "run",
        lambda *args, **kwargs: subprocess.CompletedProcess(
            args=args[0],
            returncode=1,
            stdout="",
            stderr="Error: Cannot find module 'docx'",
        ),
    )

    if mode == "bilingual":
        result = export_docx.export_bilingual_docx(
            [{"page": 1, "source": "Source\x00 text", "target": "中文譯文"}],
            title="雙語對照",
            subtitle="fallback",
            filename="fallback-bilingual.docx",
        )
        expected = {"雙語對照", "Source text", "中文譯文", "原文", "翻譯"}
    elif mode == "transcript":
        result = export_docx.export_transcript_docx(
            [{"speaker": "法官", "time": "10:30", "content": "請陳述意見。"}],
            title="庭審逐字稿",
            case_info="合成年度訴字第一號",
            filename="fallback-transcript.docx",
        )
        expected = {"庭審逐字稿", "法官", "10:30", "請陳述意見。", "發言人"}
    else:
        result = export_docx.export_summary_docx(
            [{"heading": "第一部分", "summary": "合成摘要", "excerpt": "合成原文節錄"}],
            title="文件摘要",
            filename="fallback-summary.docx",
        )
        expected = {"文件摘要", "第一部分", "合成摘要", "合成原文節錄", "原文節錄"}

    path = Path(result["path"])
    assert result["success"] is True
    assert result["generator"] == "python-docx-fallback"
    assert result["fallback_code"] == "node_backend_unavailable"
    assert "fallback_reason" not in result
    assert path.is_file()
    assert path.parent == (tmp_path / "exports").resolve()
    assert result["validation"]["table_rows"] == 2
    document = Document(path)
    text = "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    )
    assert expected <= set(text.splitlines())


def test_node_success_remains_preferred_over_python_fallback(monkeypatch, tmp_path):
    export_docx = _reload_export_docx(monkeypatch, tmp_path)

    def successful_node(args, **kwargs):
        payload = json.loads(Path(args[-1]).read_text(encoding="utf-8"))
        document = export_docx._build_python_docx(payload)
        document.save(payload["out_path"])
        return subprocess.CompletedProcess(args=args, returncode=0, stdout="", stderr="")

    monkeypatch.setattr(export_docx.subprocess, "run", successful_node)
    monkeypatch.setattr(
        export_docx,
        "_run_python_docx_fallback",
        lambda *args, **kwargs: pytest.fail("fallback must not run after Node success"),
    )

    result = export_docx.export_summary_docx(
        [{"heading": "H", "summary": "S", "excerpt": "E"}],
        filename="node-success.docx",
    )

    assert result["success"] is True
    assert result["generator"] == "node-docx"
    assert result["validation"]["mode"] == "summary"
    assert result["validation"]["table_rows"] == 2
    assert len(Document(result["path"]).tables) == 1


def test_available_docx_generator_satisfies_structural_contract(monkeypatch, tmp_path):
    export_docx = _reload_export_docx(monkeypatch, tmp_path)
    node = export_docx._find_node()
    node_available = bool(export_docx._find_node_path()) and bool(
        Path(node).is_file() if Path(node).is_absolute() else shutil.which(node)
    )

    result = export_docx.export_bilingual_docx(
        [{"page": 1, "source": "real Node source", "target": "真實 Node 譯文"}],
        title="真實 Node 契約驗證",
        filename="real-node.docx",
    )

    assert result["success"] is True
    assert result["generator"] == (
        "node-docx" if node_available else "python-docx-fallback"
    )
    assert result["validation"]["mode"] == "bilingual"
    assert result["validation"]["table_rows"] == 2
    assert result["validation"]["table_columns"] == 3


def test_paragraph_only_node_candidate_is_rejected_and_falls_back(monkeypatch, tmp_path):
    export_docx = _reload_export_docx(monkeypatch, tmp_path)

    def paragraph_only_node(args, **kwargs):
        payload = json.loads(Path(args[-1]).read_text(encoding="utf-8"))
        document = Document()
        document.add_paragraph("not the requested summary table")
        document.save(payload["out_path"])
        return subprocess.CompletedProcess(args=args, returncode=0, stdout="", stderr="")

    monkeypatch.setattr(export_docx.subprocess, "run", paragraph_only_node)
    result = export_docx.export_summary_docx(
        [{"heading": "H", "summary": "S", "excerpt": "E"}],
        filename="node-wrong-shape.docx",
    )

    assert result["success"] is True
    assert result["generator"] == "python-docx-fallback"
    assert result["validation"]["mode"] == "summary"
    assert result["validation"]["table_rows"] == 2


def test_existing_symlink_cannot_redirect_fallback_outside_exports(monkeypatch, tmp_path):
    export_docx = _reload_export_docx(monkeypatch, tmp_path)
    exports = tmp_path / "exports"
    exports.mkdir()
    outside = tmp_path / "outside.docx"
    (exports / "linked.docx").symlink_to(outside)

    result = export_docx.export_bilingual_docx(
        [{"page": 1, "source": "A", "target": "B"}],
        filename="linked.docx",
    )

    assert result["success"] is False
    assert "escapes exports directory" in result["error"]
    assert not outside.exists()


def test_sealed_release_uses_fallback_without_probing_global_node(monkeypatch, tmp_path):
    export_docx = _reload_export_docx(monkeypatch, tmp_path)
    monkeypatch.setenv("MAGI_V3_RELEASE_ID", "sealed-test-release")
    monkeypatch.setattr(
        export_docx,
        "_find_node",
        lambda: pytest.fail("sealed release must not probe a host Node binary"),
    )
    monkeypatch.setattr(
        export_docx,
        "_find_node_path",
        lambda: pytest.fail("sealed release must not probe host/global node_modules"),
    )
    monkeypatch.setattr(
        export_docx.subprocess,
        "run",
        lambda *args, **kwargs: pytest.fail("sealed release must not invoke Node"),
    )

    result = export_docx.export_summary_docx(
        [{"heading": "第一部分", "summary": "摘要", "excerpt": "節錄"}],
        filename="sealed-fallback.docx",
    )

    assert result["success"] is True
    assert result["generator"] == "python-docx-fallback"
    assert result["fallback_code"] == "node_backend_unavailable"
    assert Path(result["path"]).is_file()


def test_sealed_fallback_failure_preserves_existing_target(monkeypatch, tmp_path):
    export_docx = _reload_export_docx(monkeypatch, tmp_path)
    monkeypatch.setenv("MAGI_V3_RELEASE_ID", "sealed-test-release")
    target = tmp_path / "exports" / "existing.docx"
    target.parent.mkdir(parents=True)
    old_document = Document()
    old_document.add_paragraph("preserve this existing document")
    old_document.save(target)
    original = target.read_bytes()
    monkeypatch.setattr(
        export_docx,
        "_build_python_docx",
        lambda _data: (_ for _ in ()).throw(ImportError("synthetic python-docx failure")),
    )

    result = export_docx.export_bilingual_docx(
        [{"page": 1, "source": "A", "target": "B"}],
        filename="existing.docx",
    )

    assert result["success"] is False
    assert result["error_code"] == "docx_export_failed"
    assert result["error"] == "DOCX export failed"
    assert "synthetic" not in str(result)
    assert target.read_bytes() == original


def test_sealed_directory_swap_fails_without_writing_outside(monkeypatch, tmp_path):
    export_docx = _reload_export_docx(monkeypatch, tmp_path)
    monkeypatch.setenv("MAGI_V3_RELEASE_ID", "sealed-test-release")
    exports = tmp_path / "exports"
    parked = tmp_path / "exports-parked"
    outside = tmp_path / "outside"
    outside.mkdir()
    original_builder = export_docx._build_python_docx

    def swap_parent_after_validation(data):
        document = original_builder(data)
        exports.rename(parked)
        exports.symlink_to(outside, target_is_directory=True)
        return document

    monkeypatch.setattr(export_docx, "_build_python_docx", swap_parent_after_validation)
    result = export_docx.export_summary_docx(
        [{"heading": "H", "summary": "S", "excerpt": "E"}],
        filename="redirected.docx",
    )

    assert result["success"] is False
    assert not (outside / "redirected.docx").exists()
    assert not list(outside.glob("*.docx"))
    assert not (parked / "redirected.docx").exists()


def test_node_and_fallback_failure_preserve_existing_target(monkeypatch, tmp_path):
    export_docx = _reload_export_docx(monkeypatch, tmp_path)
    target = tmp_path / "exports" / "existing-node.docx"
    target.parent.mkdir(parents=True)
    old_document = Document()
    old_document.add_paragraph("preserve this existing Node-era document")
    old_document.save(target)
    original = target.read_bytes()
    monkeypatch.setattr(
        export_docx.subprocess,
        "run",
        lambda *args, **kwargs: subprocess.CompletedProcess(
            args=args[0], returncode=1, stdout="", stderr="synthetic Node failure"
        ),
    )
    monkeypatch.setattr(
        export_docx,
        "_build_python_docx",
        lambda _data: (_ for _ in ()).throw(ImportError("synthetic fallback failure")),
    )

    result = export_docx.export_summary_docx(
        [{"heading": "H", "summary": "S", "excerpt": "E"}],
        filename="existing-node.docx",
    )

    assert result["success"] is False
    assert target.read_bytes() == original


def test_node_directory_swap_fails_without_outside_write_or_old_file_loss(
    monkeypatch,
    tmp_path,
):
    export_docx = _reload_export_docx(monkeypatch, tmp_path)
    exports = tmp_path / "exports"
    exports.mkdir()
    parked = tmp_path / "exports-parked"
    outside = tmp_path / "outside"
    outside.mkdir()
    target = exports / "existing-node-swap.docx"
    old_document = Document()
    old_document.add_paragraph("preserve across Node directory swap")
    old_document.save(target)
    original = target.read_bytes()

    def node_then_swap(args, **kwargs):
        payload = json.loads(Path(args[-1]).read_text(encoding="utf-8"))
        document = export_docx._build_python_docx(payload)
        document.save(payload["out_path"])
        exports.rename(parked)
        exports.symlink_to(outside, target_is_directory=True)
        return subprocess.CompletedProcess(args=args, returncode=0, stdout="", stderr="")

    monkeypatch.setattr(export_docx.subprocess, "run", node_then_swap)
    result = export_docx.export_summary_docx(
        [{"heading": "H", "summary": "S", "excerpt": "E"}],
        filename="existing-node-swap.docx",
    )

    assert result["success"] is False
    assert result["error_code"] == "docx_export_directory_changed"
    assert not list(outside.glob("*.docx"))
    assert (parked / "existing-node-swap.docx").read_bytes() == original


def test_node_same_path_directory_replacement_cannot_reopen_fallback(
    monkeypatch,
    tmp_path,
):
    export_docx = _reload_export_docx(monkeypatch, tmp_path)
    exports = tmp_path / "exports"
    exports.mkdir()
    parked = tmp_path / "exports-parked"
    target = exports / "existing-node-replaced.docx"
    old_document = Document()
    old_document.add_paragraph("preserve across same-path directory replacement")
    old_document.save(target)
    original = target.read_bytes()

    def node_then_replace_directory(args, **kwargs):
        payload = json.loads(Path(args[-1]).read_text(encoding="utf-8"))
        document = export_docx._build_python_docx(payload)
        document.save(payload["out_path"])
        exports.rename(parked)
        exports.mkdir()
        return subprocess.CompletedProcess(
            args=args,
            returncode=1,
            stdout="",
            stderr="synthetic Node failure after same-path replacement",
        )

    monkeypatch.setattr(export_docx.subprocess, "run", node_then_replace_directory)
    result = export_docx.export_summary_docx(
        [{"heading": "H", "summary": "S", "excerpt": "E"}],
        filename="existing-node-replaced.docx",
    )

    assert result["success"] is False
    assert result["error_code"] == "docx_export_directory_changed"
    assert not list(exports.glob("*.docx"))
    assert (parked / "existing-node-replaced.docx").read_bytes() == original
